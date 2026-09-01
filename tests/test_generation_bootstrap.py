from pathlib import Path
import threading
import time

import pytest
from fastapi.testclient import TestClient

from conftest import events


def _checkpoint_config(engine):
    conf = {"engine": {
        "kind": "s2", "s2_key": "test-key", "generation_mode": "fast",
        "s2_base_url": engine.S2_DEFAULT_BASE, "s2_model": engine.S2_DEFAULT_MODEL,
        "s2_verify_ssl": True, "s2_wire": "auto",
    }}
    conf["setup"] = {
        "model_ids": [engine.S2_DEFAULT_MODEL, engine.S2_QUALITY_MODEL],
        "text_verified_model_ids": [engine.S2_DEFAULT_MODEL, engine.S2_QUALITY_MODEL],
        "tested_connection_fingerprint": engine._connection_fingerprint(conf),
    }
    return conf


def test_default_generation_preflight_checks_opencode_without_sowork(engine, job, monkeypatch):
    checked = []

    def resolve(name, _eng=None):
        checked.append(name)
        return None

    monkeypatch.setattr(engine, "resolve_cli", resolve)
    conf = {"engine": {"kind": "s2", "s2_key": "test-key"}}

    result = engine.generation_preflight(str(job), conf)

    assert checked == ["opencode"]
    assert result["ok"] is False
    assert result["repair"] == "opencode"
    assert result["checks"][-1]["id"] == "runtime"
    assert result["checks"][-1]["status"] == "repairing"


def test_bootstrap_checks_report_generation_ready_with_estimates_and_actions(engine, tmp_path, monkeypatch):
    skill = tmp_path / "skill"
    skill.mkdir()
    (skill / "SKILL.md").write_text("# rules\n", encoding="utf-8")
    conf = _checkpoint_config(engine)
    conf["engine"]["skill_dir"] = str(skill)
    engine.write_json(engine.conf_path(), conf)
    monkeypatch.setattr(engine, "resolve_cli", lambda name, _eng=None: "/managed/opencode" if name == "opencode" else None)

    with TestClient(engine.app) as client:
        response = client.get("/v1/bootstrap/checks")

    assert response.status_code == 200
    body = response.json()
    assert body["status"] == "ready"
    assert body["ready_for_generation"] is True
    assert [item["id"] for item in body["checks"]] == [
        "storage", "parsers", "skill", "runtime", "connection"
    ]
    assert all(item["status"] == "pass" for item in body["checks"])
    assert all(item["estimate_seconds"][0] <= item["estimate_seconds"][1] for item in body["checks"])
    assert all(set(item["repair"]) == {"kind", "action", "label"} for item in body["checks"])


def test_bootstrap_checks_make_missing_opencode_repairable_without_hiding_other_results(engine, tmp_path, monkeypatch):
    skill = tmp_path / "skill"
    skill.mkdir()
    (skill / "SKILL.md").write_text("# rules\n", encoding="utf-8")
    conf = _checkpoint_config(engine)
    conf["engine"]["skill_dir"] = str(skill)
    engine.write_json(engine.conf_path(), conf)
    monkeypatch.setattr(engine, "resolve_cli", lambda *_args, **_kwargs: None)

    with TestClient(engine.app) as client:
        body = client.get("/v1/bootstrap/checks").json()

    by_id = {item["id"]: item for item in body["checks"]}
    assert body["status"] == "repairable"
    assert body["ready_for_generation"] is False
    assert by_id["runtime"]["status"] == "repairable"
    assert by_id["runtime"]["repair"] == {
        "kind": "automatic", "action": "repair_opencode", "label": "一键修复生成组件"
    }
    assert by_id["storage"]["status"] == "pass"
    assert by_id["connection"]["status"] == "pass"


def test_missing_default_shell_repairs_inside_job_and_rechecks_path(engine, job, monkeypatch):
    state = {"ready": False, "provisioned": []}

    def resolve(name, _eng=None):
        assert name == "opencode"
        return "/tmp/opencode-cli" if state["ready"] else None

    def provision(which):
        state["provisioned"].append(which)
        state["ready"] = True
        engine.PROV.update({"state": "done", "pct": 100, "note": "已安装并校验", "error": ""})

    monkeypatch.setattr(engine, "resolve_cli", resolve)
    monkeypatch.setattr(engine, "_provision_codex", provision)
    engine.PROV.update({"state": "idle", "pct": 0, "note": "", "error": ""})

    ok, note = engine.ensure_default_shell(str(job), {"kind": "s2"})

    assert ok is True
    assert state["provisioned"] == ["opencode"]
    assert "已就绪" in note
    worklog = [line for event in events(job) if event.get("type") == "worklog" for line in event["lines"]]
    assert any("修复生成组件" in line for line in worklog)
    assert any("生成组件已就绪" in line for line in worklog)


def test_pipeline_attempt_receives_only_declared_inputs(engine, job, tmp_path):
    engine.generation_pipeline.initialize(
        job, run_id="copy-inputs", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
    )
    (job / "招标文件_解析版.md").write_text("公开解析文本", encoding="utf-8")
    (job / "你的要求.md").write_text("公开补充要求", encoding="utf-8")
    (job / "不应暴露.txt").write_text("本机秘密", encoding="utf-8")
    outside = tmp_path / "outside-secret.txt"
    outside.write_text("目录外秘密", encoding="utf-8")
    (job / "素材").mkdir()
    (job / "素材" / "允许.txt").write_text("允许素材", encoding="utf-8")
    (job / "素材" / "越界.txt").symlink_to(outside)
    target = tmp_path / "attempt"
    target.mkdir()
    node = {"id": "chapter_write:01"}

    engine._pipeline_copy_inputs(str(job), node, str(target))

    assert (target / "招标文件_解析版.md").is_file()
    assert (target / "你的要求.md").is_file()
    assert (target / "素材" / "允许.txt").is_file()
    assert not (target / "不应暴露.txt").exists()
    assert not (target / "素材" / "越界.txt").exists()


def test_dependent_chapter_copies_and_hashes_its_dependency_outputs(engine, job, tmp_path):
    engine.generation_pipeline.initialize(
        job, run_id="dependency-readset", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "old", "title": "旧节点", "output": "旧节点.md"}],
    )
    for name in ("投标文件组成.md", "评分点响应矩阵.md", "废标风险清单.md"):
        (job / name).write_text("# 规划\n\n" + "逐项响应。" * 40, encoding="utf-8")
    engine.write_json(str(job / "response_plan.json"), {"chapters": [
        {"id": "a", "title": "A", "output": "A.md", "basis": ["条款A"],
         "scoring_points": [], "material_slots": [], "dependencies": []},
        {"id": "b", "title": "B", "output": "B.md", "basis": ["条款B"],
         "scoring_points": [], "material_slots": [], "dependencies": ["a"]},
    ]})
    engine.generation_pipeline.start_node(job, "response_plan", input_digest="plan")
    engine.generation_pipeline.complete_node(job, "response_plan", input_digest="plan")
    state = engine.generation_pipeline.apply_response_plan(job)
    upstream = job / "A.md"
    upstream.write_text("上游版本一" * 80, encoding="utf-8")
    downstream = next(node for node in state["nodes"] if node["id"] == "chapter_write:b")
    target = tmp_path / "dependent-attempt"
    target.mkdir()

    engine._pipeline_copy_inputs(str(job), downstream, str(target))
    names = engine._pipeline_declared_input_names(str(job), downstream, state)
    before = engine.generation_pipeline.file_digest([job / name for name in names])
    upstream.write_text("上游版本二" * 80, encoding="utf-8")
    after = engine.generation_pipeline.file_digest([job / name for name in names])

    assert (target / "A.md").is_file()
    assert before != after


def test_launch_reports_environment_preparation_before_business_step_one(engine, tmp_path, monkeypatch):
    job = Path(engine.jpath("preflight-job"))
    job.mkdir(parents=True)
    (job / "采购文件.md").write_text("# 采购文件\n", encoding="utf-8")
    engine.write_json(str(job / "任务.json"), {"name": "预检任务", "tender": "采购文件.md", "staged": True})
    engine.write_json(engine.conf_path(), {"engine": {"kind": "s2", "s2_key": ""}})
    started = []
    monkeypatch.setattr(engine, "_start_reserved_worker", lambda *args: started.append(args))

    result = engine._launch_job("preflight-job", str(job), "auto")

    assert result["mode"] == "mock"
    progress = [event for event in events(job) if event.get("type") == "progress"]
    assert progress[-1]["step"] == 0
    assert progress[-1]["stage"] == "正在检查生成环境"
    assert started


def test_diagnostics_exposes_preflight_shell_source_and_repair_state(engine, job, monkeypatch):
    shell = "/managed/data/bin/opencode-cli"
    monkeypatch.setattr(engine, "resolve_cli", lambda name, _eng=None: shell if name == "opencode" else None)
    engine.write_json(engine.conf_path(), {"engine": {"kind": "s2", "s2_key": "test-key"}})
    engine.write_json(str(job / "preflight.json"), {
        "phase": "environment", "repair": "opencode", "checked_at": "2026-08-20T12:00:00",
        "checks": [{"id": "runtime", "status": "repairing", "message": "正在修复"}],
    })
    engine.PROV.update({"state": "running", "which": "opencode", "pct": 42,
                        "note": "正在下载", "path": "", "error": ""})

    result = engine._diagnostic_snapshot("job-1")

    assert result["runtime_source"] == shell
    assert result["provision"] == {"state": "running", "which": "opencode", "pct": 42,
                                    "note": "正在下载", "error": ""}
    assert result["job"]["preflight"]["repair"] == "opencode"
    assert result["job"]["last_activity"]


def test_job_flow_expands_environment_preflight_before_business_steps(engine, job):
    preflight = {
        "phase": "environment",
        "checked_at": "2026-08-21T09:00:00",
        "checks": [
            {"id": "storage", "label": "任务目录", "status": "pass", "message": "任务文件可以保存"},
            {"id": "skill", "label": "写作规则", "status": "pass", "message": "写作规则已就绪"},
            {"id": "connection", "label": "模型连接", "status": "pending", "message": "正在建立连接"},
            {"id": "runtime", "label": "生成组件", "status": "repairing", "message": "正在自动修复"},
        ],
    }
    engine.write_json(str(job / "preflight.json"), preflight)
    progress = {"type": "progress", "step": 0, "total": 12, "pct": 1,
                "stage": "正在修复生成组件", "ts": "2026-08-21T09:00:01"}

    flow = engine.job_flow(str(job), "running", {}, progress, {})

    assert flow["version"] == 2
    assert flow["current_phase"] == "environment"
    assert flow["checkpoint"] == {"step": 0, "label": "任务文件已保存"}
    assert flow["recoverable"] is True
    assert [phase["id"] for phase in flow["phases"]] == [
        "environment", "parse", "plan", "write", "assemble", "deliver"
    ]
    environment = flow["phases"][0]
    assert environment["state"] == "active"
    assert environment["evidence"] == "preflight.json"
    assert [check["state"] for check in environment["checks"]] == [
        "done", "done", "active", "active"
    ]


def test_job_flow_uses_verified_business_checkpoint_and_keeps_interrupted_job_recoverable(engine, job):
    progress = engine.read_json(str(job / "progress.json"), {})
    outcome = {"state": "stopped", "reason": "上游连接中断", "ts": "2026-08-21T09:10:00"}
    engine.write_json(str(job / "outcome.json"), outcome)

    flow = engine.job_flow(str(job), "stopped", {"oc_session": "session-1"}, progress, outcome)

    assert flow["current_phase"] == "parse"
    assert flow["checkpoint"]["step"] == 1
    assert flow["checkpoint"]["label"] == "体检素材"
    assert flow["recoverable"] is True
    assert flow["phases"][0]["state"] == "done"
    assert flow["phases"][1]["state"] == "attention"
    assert flow["phases"][1]["detail"] == "上游连接中断"


def test_job_listing_and_diagnostics_share_the_same_persisted_flow(engine, job, monkeypatch):
    monkeypatch.setattr(engine, "resolve_cli", lambda *_args, **_kwargs: "/managed/opencode-cli")
    listed = engine.list_jobs()
    diagnostic = engine._diagnostic_snapshot("job-1")

    assert listed[0]["flow"]["checkpoint"]["step"] == 1
    listed_flow = dict(listed[0]["flow"])
    diagnostic_flow = dict(diagnostic["job"]["flow"])
    listed_silence = listed_flow.pop("silence_seconds")
    diagnostic_silence = diagnostic_flow.pop("silence_seconds")
    assert diagnostic_flow == listed_flow
    assert 0 <= diagnostic_silence - listed_silence <= 1


def test_s2_launch_uses_checkpoint_pipeline_instead_of_legacy_long_session(engine, tmp_path, monkeypatch):
    job = Path(engine.jpath("pipeline-launch"))
    job.mkdir(parents=True)
    (job / "采购文件.md").write_text("# 采购需求\n" + "软件平台建设与运维服务。" * 20, encoding="utf-8")
    engine.write_json(str(job / "任务.json"), {
        "name": "流水线启动", "tender": "采购文件.md", "staged": True,
        "template_snapshot": {"package": {"outline": [
            {"title": "项目理解", "purpose": "说明建设目标"},
            {"title": "实施方案", "purpose": "说明实施计划"},
        ]}},
    })
    engine.write_json(engine.conf_path(), _checkpoint_config(engine))
    started = []
    monkeypatch.setattr(engine, "generation_preflight", lambda *_args: {"ok": True, "repair": False})
    monkeypatch.setattr(engine, "config_agent_cmd", lambda: ["opencode", "run", engine.AGENT_PROMPT])
    monkeypatch.setattr(engine, "_start_reserved_worker", lambda *args: started.append(args))

    result = engine._launch_job("pipeline-launch", str(job), "auto")

    assert result["mode"] == "agent"
    assert started and started[0][2] is engine.generation_pipeline_worker
    state = engine.generation_pipeline.load(job)
    assert state["mode"] == "fast"
    manifest = engine.read_json(str(job / "任务.json"), {}).get("skill_manifest") or {}
    assert manifest.get("run_id") == engine.read_json(str(job / "任务.json"), {}).get("run_id")
    assert manifest.get("manifest_sha256")
    assert manifest.get("execution_path") == "opencode"
    assert [node["title"] for node in state["nodes"] if node["id"].startswith("chapter_write:")] == [
        "项目理解", "实施方案", "技术应答偏离表", "商务偏离表",
    ]


def test_model_node_uses_non_stream_completion_and_injects_skill_contract(
    engine, tmp_path, monkeypatch
):
    job = Path(engine.jpath("direct-model-node"))
    job.mkdir(parents=True)
    skill_dir = tmp_path / "skill"
    skill_dir.mkdir()
    (skill_dir / "SKILL.md").write_text(
        "# 投标写作规则\n\n不得编造资质；未知事实标记为需补充。\n", encoding="utf-8"
    )
    conf = _checkpoint_config(engine)
    conf["engine"]["skill_dir"] = str(skill_dir)
    engine.write_json(engine.conf_path(), conf)
    engine.write_json(str(job / "任务.json"), {
        "name": "直连节点", "run_id": "run-direct", "tender": "采购文件.md",
    })
    engine._set_skill_manifest(
        str(job), str(skill_dir), "direct-model", True, "direct_model_completion"
    )
    (job / "招标文件_解析版.md").write_text(
        "# 采购要求\n\n必须提供实施、验收与运维方案。", encoding="utf-8"
    )
    captured = {}

    def fake_openai(base, key, path, payload=None, **kwargs):
        captured.update({"base": base, "key": key, "path": path,
                         "payload": payload, "kwargs": kwargs})
        return {"choices": [{"finish_reason": "stop", "message": {"content": (
            '# 实施方案\n\n' + '逐项响应招标文件并提供实施、验收和运维安排。' * 20
        )}}]}

    def fake_stream(base, key, payload, **kwargs):
        # v0.20.4 起节点走流式优先(空闲超时替代总超时,防网关掐长连接);
        # 请求体与非流式同构,这里复用同一份捕获与回包。
        return fake_openai(base, key, "/chat/completions", payload, **kwargs)

    monkeypatch.setattr(engine, "_openai_req", fake_openai)
    monkeypatch.setattr(engine, "_openai_stream_req", fake_stream)
    monkeypatch.setattr(
        engine, "_pipeline_declared_input_names",
        lambda *_args, **_kwargs: ["招标文件_解析版.md"],
    )
    monkeypatch.setattr(
        engine, "oc_serve", lambda *_args, **_kwargs: (_ for _ in ()).throw(
            AssertionError("direct model nodes must not start OpenCode")
        )
    )
    node = {
        "id": "chapter_write:01", "title": "实施方案", "outputs": ["章节_01.md"],
        "min_chars": 120, "attempt": 1, "attempt_serial": 1,
    }

    engine._pipeline_model_runner(
        str(job), node, "必须使用 write/edit 工具写文件。", engine.S2_DEFAULT_MODEL
    )

    assert captured["path"] == "/chat/completions"
    assert captured["payload"]["model"] == engine.S2_DEFAULT_MODEL
    assert captured["payload"]["max_tokens"] == 4800
    assert "不得编造资质" in captured["payload"]["messages"][0]["content"]
    assert "至少 2500 个中文字符" in captured["payload"]["messages"][1]["content"]
    assert "禁止反复使用“信息项｜内容”" in captured["payload"]["messages"][1]["content"]
    assert "write/edit" not in captured["payload"]["messages"][1]["content"]
    assert (job / "章节_01.md").is_file()
    assert engine.skill_evidence(str(job))["state"] == "verified"


def test_truncated_chapter_continues_once_without_discarding_first_response(
    engine, tmp_path, monkeypatch
):
    job = Path(engine.jpath("continued-model-node"))
    job.mkdir(parents=True)
    skill_dir = tmp_path / "continued-skill"
    skill_dir.mkdir()
    (skill_dir / "SKILL.md").write_text(
        "# 投标写作规则\n\n逐项响应，未知事实标记为需补充。\n", encoding="utf-8"
    )
    conf = _checkpoint_config(engine)
    conf["engine"]["skill_dir"] = str(skill_dir)
    engine.write_json(engine.conf_path(), conf)
    engine.write_json(str(job / "任务.json"), {
        "name": "断点续写", "run_id": "run-continued", "tender": "采购文件.md",
    })
    engine._set_skill_manifest(
        str(job), str(skill_dir), "continued-model", True, "direct_model_completion"
    )
    (job / "招标文件_解析版.md").write_text(
        "# 采购要求\n\n必须提供完整实施、验收与运维方案。", encoding="utf-8"
    )
    calls = []
    # 同上:变化的是句尾序号,断言里的「第一部分逐项响应」前缀原样保留
    first = "# 实施方案\n\n" + "".join(
        "第一部分逐项响应招标要求，并给出验收依据（第%d项）。" % j for j in range(1, 141))
    second = "## 运维与验收\n\n第二部分补齐运维安排和量化验收标准。" * 12

    def fake_openai(_base, _key, _path, payload=None, **_kwargs):
        calls.append(payload)
        if len(calls) == 1:
            return {"choices": [{"finish_reason": "length", "message": {"content": first}}]}
        return {"choices": [{"finish_reason": "stop", "message": {"content": second}}]}

    monkeypatch.setattr(engine, "_openai_req", fake_openai)
    monkeypatch.setattr(engine, "_openai_stream_req",
                        lambda base, key, payload, **kw: fake_openai(base, key, "/chat/completions", payload))
    monkeypatch.setattr(
        engine, "_pipeline_declared_input_names",
        lambda *_args, **_kwargs: ["招标文件_解析版.md"],
    )
    node = {
        "id": "chapter_write:01", "title": "实施方案", "outputs": ["章节_01.md"],
        "min_chars": 120, "attempt": 1, "attempt_serial": 1,
    }

    engine._pipeline_model_runner(
        str(job), node, "撰写完整实施方案。", engine.S2_DEFAULT_MODEL
    )

    body = (job / "章节_01.md").read_text(encoding="utf-8")
    assert len(calls) == 2
    assert calls[0]["max_tokens"] == 4800
    assert calls[1]["max_tokens"] == 2400
    assert "第一部分逐项响应" in body
    assert "第二部分补齐" in body
    assert "不要重复已经完成的内容" in calls[1]["messages"][-1]["content"]


def test_response_plan_is_built_locally_without_calling_the_model(
    engine, tmp_path, monkeypatch
):
    job = Path(engine.jpath("local-response-plan"))
    job.mkdir(parents=True)
    skill_dir = tmp_path / "local-plan-skill"
    skill_dir.mkdir()
    (skill_dir / "SKILL.md").write_text(
        "# 投标写作规则\n\n评分点和废标风险必须逐项核对，不得编造资质。\n", encoding="utf-8"
    )
    conf = _checkpoint_config(engine)
    conf["engine"]["skill_dir"] = str(skill_dir)
    engine.write_json(engine.conf_path(), conf)
    engine.write_json(str(job / "任务.json"), {
        "name": "本地响应规划", "run_id": "run-local-plan", "tender": "采购文件.md",
    })
    engine.generation_pipeline.initialize(
        job, run_id="run-local-plan", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[
            {"id": "01", "title": "项目理解", "output": "章节_01_项目理解.md"},
            {"id": "02", "title": "实施验收", "output": "章节_02_实施验收.md"},
        ],
    )
    (job / "招标文件_解析版.md").write_text(
        "# 采购要求\n\n必须提供实施和验收方案。\n\n如弄虚作假，将取消投标资格。\n",
        encoding="utf-8",
    )
    engine._set_skill_manifest(
        str(job), str(skill_dir), "local-plan", True, "direct_model_completion"
    )
    node = next(
        item for item in engine.generation_pipeline.load(job)["nodes"]
        if item["id"] == "response_plan"
    )
    node.update({"attempt": 1, "attempt_serial": 1})
    monkeypatch.setattr(
        engine, "_openai_req",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(
            AssertionError("response planning must be deterministic and local")
        ),
    )
    monkeypatch.setattr(
        engine, "_openai_stream_req",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(
            AssertionError("response planning must be deterministic and local")
        ),
    )

    engine._pipeline_model_runner(
        str(job), node, "提取响应规划", engine.S2_DEFAULT_MODEL
    )

    plan = engine.read_json(str(job / "response_plan.json"), {})
    assert [item["title"] for item in plan["chapters"]] == ["项目理解", "实施验收"]
    assert "取消投标资格" in (job / "废标风险清单.md").read_text(encoding="utf-8")
    assert (job / "评分点响应矩阵.md").is_file()
    assert engine.skill_evidence(str(job))["state"] == "verified"


def test_pipeline_retry_message_reports_output_contract_failure_instead_of_connection(
    engine, monkeypatch
):
    emitted = []
    monkeypatch.setattr(engine, "emit", lambda _job, event: emitted.append(event))

    engine._pipeline_event("/tmp/job", "retry", {
        "id": "chapter_write:01", "title": "项目理解", "attempt": 1,
        "max_attempts": 3, "error_code": "model_output_truncated",
    })

    assert "输出过长" in emitted[0]["text"]
    assert "连接中断" not in emitted[0]["text"]


def test_truncated_markdown_keeps_complete_body_and_drops_incomplete_tail(engine):
    body = (
        "# 实施方案\n\n"
        + "本节逐项响应招标要求，明确实施动作与验收依据。" * 50
        + "\n\n## 未完成的小节\n\n这是被长度上限截断的半句"
    )

    accepted = engine._pipeline_complete_truncated_markdown(body, min_chars=120)

    assert accepted.endswith("。")
    assert "这是被长度上限截断的半句" not in accepted
    assert engine._pipeline_complete_truncated_markdown("内容过短", min_chars=120) == ""


def test_unverified_bidder_fact_claims_are_marked_without_breaking_tables(engine):
    body = (
        "我方具备大型系统集成经验，近三年完成三项类似工程。\n\n"
        "| 序号 | 响应 | 证据 |\n|---|---|---|\n"
        "| 1 | 我方拥有广泛的金融客户群 | 案例证明 |\n"
        "| 2 | 我方在过去三年完成三项类似项目 | 合同与验收材料〔需补充〕 |\n"
        "我方注册资本为〔需补充〕万元。\n"
    )

    safe = engine._pipeline_mark_unverified_claims(body)

    assert safe.count("〔需补充：请核实本项投标人事实并提供证据〕") == 4
    assert "| 1 | 〔需补充：请核实本项投标人事实并提供证据〕我方拥有" in safe
    assert "| 2 | 〔需补充：请核实本项投标人事实并提供证据〕我方在过去三年" in safe
    assert "〔需补充：请核实本项投标人事实并提供证据〕我方注册资本为〔需补充〕万元" in safe


def test_later_unverified_claim_is_not_protected_by_earlier_marker(engine):
    body = "客户资料〔需补充〕；我方具有一级资质，公司拥有全国服务网络，我方持有安全认证。"

    safe = engine._pipeline_mark_unverified_claims(body)

    assert "〔需补充：请核实本项投标人事实并提供证据〕我方具有一级资质" in safe
    assert "〔需补充：请核实本项投标人事实并提供证据〕公司拥有全国服务网络" in safe
    assert "〔需补充：请核实本项投标人事实并提供证据〕我方持有安全认证" in safe


def test_standard_pipeline_preserves_model_review_as_separate_artifact(engine, job):
    state = engine.generation_pipeline.initialize(
        job, run_id="standard-review", mode="standard",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
    )

    review = next(node for node in state["nodes"] if node["id"] == "quality_review")
    assert review["outputs"] == ["模型复核报告.md"]


def test_resumed_pipeline_rejects_changed_gateway_but_accepts_new_key(engine, job):
    state = engine.generation_pipeline.initialize(
        job, run_id="frozen-transport", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
        credential_fingerprint=engine.generation_pipeline.credential_fingerprint("old-key"),
        base_url="https://gateway.example/v1",
    )
    assert state["model_routes"]["base_url"] == "https://gateway.example/v1"

    # 网关变化仍然硬拦截:换网关等于换执行环境,检查点必须作废
    with pytest.raises(engine.generation_pipeline.NodeExecutionError) as exc:
        engine._pipeline_validate_execution_contract(job, {
            "base_url": "https://other-gateway.example/v1", "api_key": "old-key"
        })
    assert exc.value.code == "execution_contract_changed"

    # 同网关换 Key 是用户失败后最常见的自救,不再作废检查点;指纹换到新 Key 上继续
    engine._pipeline_validate_execution_contract(job, {
        "base_url": "https://gateway.example/v1", "api_key": "new-key"
    })
    updated = engine.generation_pipeline.load(job)
    assert updated["model_routes"]["credential_fingerprint"] == \
        engine.generation_pipeline.credential_fingerprint("new-key")


def test_repeated_generic_table_headers_are_specialized_by_section(engine):
    body = """# 第一章 资格响应

## 案例一
| 信息项 | 内容 |
|---|---|
| 客户 | 〔需补充〕 |

## 案例二
| 信息项 | 内容 |
|---|---|
| 客户 | 〔需补充〕 |

## 公司基本情况
| 信息项 | 内容 |
|---|---|
| 名称 | 〔需补充〕 |
"""

    safe = engine._pipeline_specialize_repeated_table_headers(body)

    assert "| 信息项 | 内容 |" not in safe
    assert "| 案例一信息项 | 待补充/响应内容 |" in safe
    assert "| 案例二信息项 | 待补充/响应内容 |" in safe
    assert "| 公司基本情况信息项 | 待补充/响应内容 |" in safe


def test_compact_response_plan_is_expanded_to_four_valid_checkpoint_files(engine, tmp_path):
    job = tmp_path / "compact-plan"
    job.mkdir()
    engine.generation_pipeline.initialize(
        job, run_id="run-plan", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[
            {"id": "01", "title": "项目理解与总体方案", "output": "章节_01_总体方案.md"},
            {"id": "02", "title": "实施与验收", "output": "章节_02_实施验收.md"},
        ],
    )
    node = next(item for item in engine.generation_pipeline.load(job)["nodes"]
                if item["id"] == "response_plan")
    attempt = engine.generation_pipeline.attempt_directory(job, {**node, "attempt": 1})
    compact = {
        "composition": ["投标函与资格证明", "技术响应文件", "商务响应与报价文件"],
        "scoring_points": [{
            "requirement": "总体方案完整性", "score": "未知", "location": "项目理解与总体方案",
            "evidence": "方案正文", "gap": "〔需补充〕最终评分分值",
        }],
        "risks": [{
            "category": "资格", "requirement": "提供有效营业执照", "risk": "缺失可能废标",
            "action": "提交前人工核验并盖章",
        }],
        "chapter_guidance": [{
            "title": "项目理解与总体方案", "basis": ["采购需求"],
            "scoring_points": ["总体方案完整性"], "material_slots": ["项目案例"],
        }],
    }

    engine._pipeline_write_response_plan(str(job), str(attempt), compact)

    engine.generation_pipeline.validate_outputs(attempt, node)
    plan = engine.read_json(str(attempt / "response_plan.json"), {})
    assert [item["id"] for item in plan["chapters"]] == ["01", "02"]
    assert "总体方案完整性" in (attempt / "评分点响应矩阵.md").read_text(encoding="utf-8")
    assert "缺失可能废标" in (attempt / "废标风险清单.md").read_text(encoding="utf-8")


def test_model_json_parser_skips_reasoning_braces_and_finds_fenced_contract(engine):
    raw = (
        '分析过程中的示例 {not-json} 不属于结果。\n'
        '```json\n{"composition":["技术响应"],"scoring_points":[],"risks":[]}\n```\n'
        '以上为结果。'
    )

    parsed = engine._pipeline_json_object(raw)

    assert parsed["composition"] == ["技术响应"]


def test_skill_contract_keeps_fact_boundaries_but_drops_tool_scripts(engine):
    raw = (
        '# 投标 Skill\n必须遵守事实边界，不得编造资质。\n'
        '```python\nprint("tool script must not reach model")\n```\n'
        '## 成本控制\n评分点和废标风险必须逐项核对。\n'
        '## 边界\n未知客户事实标记为需补充。\n'
    )

    compiled = engine._pipeline_skill_contract(raw)

    assert "不得编造资质" in compiled
    assert "废标风险必须逐项核对" in compiled
    assert "tool script must not reach model" not in compiled


def test_pipeline_worker_runs_local_parse_short_nodes_and_word_checkpoint(engine, tmp_path, monkeypatch):
    job = Path(engine.jpath("pipeline-worker"))
    job.mkdir(parents=True)
    (job / "采购文件.md").write_text("# 采购需求\n" + "软件平台建设、部署、培训、验收和运维服务。" * 30, encoding="utf-8")
    meta = {
        "name": "流水线跑通", "tender": "采购文件.md", "staged": False, "run_id": "run-1",
        "template_snapshot": {"package": {"outline": [
            {"title": "项目理解"}, {"title": "实施方案"},
        ]}},
    }
    engine.write_json(str(job / "任务.json"), meta)
    engine.write_json(engine.conf_path(), _checkpoint_config(engine))
    engine._initialize_generation_pipeline(str(job), meta, engine.read_json(engine.conf_path(), {}))
    calls = []

    def model_runner(path, node, _prompt, model):
        calls.append((node["id"], model))
        for output in node["outputs"]:
            if output == "response_plan.json":
                engine.write_json(str(Path(path) / output), {"chapters": [
                    {"id": "plan_a", "title": "响应总述", "output": "规划_A_响应总述.md",
                     "basis": ["采购需求"], "scoring_points": ["总体方案"],
                     "material_slots": [], "dependencies": []},
                    {"id": "plan_b", "title": "实施交付", "output": "规划_B_实施交付.md",
                     "basis": ["交付要求"], "scoring_points": ["实施计划"],
                     "material_slots": ["实施案例"],
                     "dependencies": ["plan_a"]},
                ]})
            else:
                # 假模型也得写「像样的正文」:同一句连着重复 120 遍正是复读退化的形状,
                # 章节门禁会(正确地)把它判成废稿,stub 也就不再代表正常产出了。
                text = "# %s\n\n%s" % (node["title"], "".join(
                    "逐项依据招标要求响应第%d条。" % j for j in range(1, 121)))
                (Path(path) / output).write_text(text, encoding="utf-8")

    def export_word(path, known, force=False):
        from docx import Document
        document = Document()
        document.add_heading("投标文件", level=1)
        document.add_paragraph("逐项响应招标要求并提供可核验依据。" * 30)
        document.save(Path(path) / "投标文件_整册.docx")
        known.add("投标文件_整册.docx")
        return ["投标文件_整册.docx"]

    monkeypatch.setattr(engine, "_pipeline_model_runner", model_runner)
    monkeypatch.setattr(engine, "ensure_docx", export_word)
    def format_ok(path, _word=''):
        (Path(path) / "Word格式自检报告.md").write_text(
            "# 投标文件格式自检报告\n\n- 结论：✅ 全部通过（1 项）\n", encoding="utf-8")
        return {"status": "pass"}
    monkeypatch.setattr(engine, "word_format_audit", format_ok)
    monkeypatch.setattr(engine, "settle", lambda path, commit=True, **_kwargs: {
        "state": "done" if commit else "ready",
    })
    monkeypatch.setattr(engine, "delivery_summary", lambda path: {"ready": True})

    engine.generation_pipeline_worker(str(job))

    state = engine.generation_pipeline.load(job)
    assert state["state"] == "done"
    assert all(node["state"] == "done" for node in state["nodes"])
    assert calls[0][0] == "response_plan"
    assert {node for node, _model in calls[1:]} == {
        "chapter_write:plan_a", "chapter_write:plan_b",
        "chapter_write:technical_deviation", "chapter_write:business_deviation",
    }
    assert [node for node, _model in calls].index("chapter_write:plan_a") < [
        node for node, _model in calls].index("chapter_write:plan_b")
    assert {model for _node, model in calls} == {engine.S2_DEFAULT_MODEL}
    assert (job / "招标文件_解析版.md").is_file()
    assert (job / "投标文件_整册.md").is_file()
    assert (job / "投标文件_整册.docx").stat().st_size > 1024
    assert (job / "Word格式自检报告.md").is_file()


def test_pipeline_offline_e2e_reaches_real_word_delivery_gate(engine, monkeypatch):
    job = Path(engine.jpath("pipeline-real-word"))
    job.mkdir(parents=True)
    (job / "采购文件.md").write_text(
        "# 采购需求\n" + "统一平台建设、实施、培训、验收和运维要求。" * 80,
        encoding="utf-8",
    )
    meta = {
        "name": "真实Word门禁", "tender": "采购文件.md", "staged": False, "run_id": "real-word",
        "template_snapshot": {"package": {"outline": [
            {"title": "项目理解"}, {"title": "实施方案"},
        ]}},
    }
    engine.write_json(str(job / "任务.json"), meta)
    conf = _checkpoint_config(engine)
    engine.write_json(engine.conf_path(), conf)
    engine._initialize_generation_pipeline(str(job), meta, conf)

    def model_runner(path, node, _prompt, _model):
        title = str(node.get("title") or "响应")
        for output in node["outputs"]:
            target = Path(path) / output
            if output == "response_plan.json":
                engine.write_json(str(target), {"chapters": [
                    {"id": "01", "title": "项目理解", "output": "章节_01_项目理解.md",
                     "basis": ["采购需求"], "scoring_points": ["方案完整性"],
                     "material_slots": [], "dependencies": []}
                ]})
            elif "偏离表" in output:
                target.write_text(
                    f"# {title}\n\n| 序号 | 招标要求 | 投标响应 | 偏离情况 | 依据/证据 | 备注 |\n"
                    "|---|---|---|---|---|---|\n"
                    "| 1 | 按期交付 | 完全响应 | 无偏离 | 招标解析版 | 已核对 |\n"
                    + (title + "核对说明。") * 40,
                    encoding="utf-8",
                )
            else:
                sections = [
                    f"## {title}要点{i}\n\n{title}第{i}项给出独立的实施动作、交付标准、核对依据和风险控制。"
                    + (f"{title}细化措施{i}。" * 80)
                    for i in range(1, 9)
                ]
                target.write_text("# " + title + "\n\n" + "\n\n".join(sections), encoding="utf-8")

    monkeypatch.setattr(engine, "_pipeline_model_runner", model_runner)
    engine.generation_pipeline_worker(str(job))

    state = engine.generation_pipeline.load(job)
    word = job / "投标文件_整册.docx"
    assert state["state"] == "done"
    assert engine.read_json(str(job / "outcome.json"), {})["state"] == "done"
    assert engine.delivery_summary(str(job))["ready"] is True
    assert engine.delivery_summary(str(job))["format"]["status"] == "pass"
    assert engine._valid_docx(str(word)) is True


def test_resume_pipeline_does_not_require_old_opencode_session(engine, job, monkeypatch):
    meta = engine.read_json(str(job / "任务.json"), {})
    meta.update({"run_id": "resume-run", "staged": False})
    engine.write_json(str(job / "任务.json"), meta)
    engine.generation_pipeline.initialize(
        job, run_id="resume-run", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01_项目理解.md"}],
    )
    started = []
    monkeypatch.setattr(engine, "_start_reserved_worker", lambda *args: started.append(args))

    with TestClient(engine.app) as client:
        response = client.post(f"/v1/jobs/{job.name}/resume")

    assert response.status_code == 200
    assert response.json()["pipeline"] is True
    assert started and started[0][2] is engine.generation_pipeline_worker


def test_resume_finalizes_done_pipeline_after_crash_before_outcome_commit(engine, job, monkeypatch):
    meta = engine.read_json(str(job / "任务.json"), {})
    meta.update({"run_id": "commit-pending", "staged": False})
    engine.write_json(str(job / "任务.json"), meta)
    state = engine.generation_pipeline.initialize(
        job, run_id="commit-pending", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
    )
    for node in state["nodes"]:
        node["state"] = "done"
        node["finished_at"] = engine.now()
    state["state"] = "done"
    state["recoverable"] = False
    state["current_nodes"] = []
    engine.write_json(str(job / "pipeline.json"), state)
    monkeypatch.setattr(engine, "settle", lambda path, **_kwargs: {"state": "done"})
    started = []
    monkeypatch.setattr(engine, "_start_reserved_worker", lambda *args: started.append(args))

    with TestClient(engine.app) as client:
        response = client.post(f"/v1/jobs/{job.name}/resume")

    assert response.status_code == 200
    assert response.json()["finalized"] is True
    assert started == []
    assert engine._is_running(job.name) is False


def test_commit_pending_pipeline_exposes_resume_action(engine, job):
    state = engine.generation_pipeline.initialize(
        job, run_id="commit-pending-ui", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
    )
    for node in state["nodes"]:
        node["state"] = "done"
    state.update({"state": "done", "recoverable": False, "current_nodes": []})
    engine.write_json(str(job / "pipeline.json"), state)

    assert "resume" in engine.job_can(str(job), "stopped", {})


def test_job_flow_uses_pipeline_node_truth_when_legacy_progress_is_stale(engine, job):
    meta = engine.read_json(str(job / "任务.json"), {})
    engine.generation_pipeline.initialize(
        job, run_id="flow-run", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01_项目理解.md"}],
    )
    for name in ("投标文件组成.md", "评分点响应矩阵.md", "废标风险清单.md"):
        (job / name).write_text("# 规划\n" + "逐项响应。" * 40, encoding="utf-8")
    engine.write_json(str(job / "response_plan.json"), {"chapters": [
        {"id": "01", "title": "项目理解", "output": "章节_01_项目理解.md",
         "basis": ["采购需求"], "scoring_points": [],
         "material_slots": [], "dependencies": []}
    ]})
    engine.generation_pipeline.start_node(job, "response_plan", input_digest="plan")
    engine.generation_pipeline.complete_node(job, "response_plan", input_digest="plan")
    engine.generation_pipeline.start_node(job, "chapter_write:01", input_digest="chapter")
    stale = {"type": "progress", "stage": "正在读招标文件", "step": 1, "pct": 2, "total": 12}

    flow = engine.job_flow(str(job), "running", meta, stale, {})

    assert flow["current_phase"] == "write"
    assert "项目理解" in flow["current_action"]
    assert flow["pipeline"]["current_nodes"] == ["chapter_write:01"]
    write_phase = next(phase for phase in flow["phases"] if phase["id"] == "write")
    parse_phase = next(phase for phase in flow["phases"] if phase["id"] == "parse")
    assert any(check["id"] == "chapter_write:01" and check["state"] == "active"
               for check in write_phase["checks"])
    assert parse_phase["expected_seconds"] == 90
    assert write_phase["expected_seconds"] == 600
    assert parse_phase["estimate_source"] == "pipeline_reference"


def test_settle_never_commits_done_before_delivery_is_ready(engine, job, monkeypatch):
    committed = []
    halted = []
    monkeypatch.setattr(engine, "harvest", lambda *_args: None)
    monkeypatch.setattr(engine, "list_deliverables", lambda *_args: ["投标文件_整册.docx"])
    monkeypatch.setattr(engine, "_body_mds", lambda *_args: [])
    monkeypatch.setattr(engine, "_body_docxs", lambda *_args: ["投标文件_整册.docx"])
    monkeypatch.setattr(engine, "quality_audit", lambda *_args: {
        "status": "pass", "level": "green", "summary": "质检通过",
    })
    monkeypatch.setattr(engine, "delivery_summary", lambda *_args, **_kwargs: {
        "ready": False, "word": {"present": True}, "toc": {"status": "fail"},
    })
    monkeypatch.setattr(engine, "_commit_done", lambda *_args: committed.append(True) or True)
    monkeypatch.setattr(engine, "halt", lambda _job, why: halted.append(why))

    result = engine.settle(str(job), known={"投标文件_整册.docx"})

    assert result["state"] == "stopped"
    assert committed == []
    assert halted == ["已停止（出件检查未通过）"]


def test_pipeline_outcome_done_is_not_user_visible_until_delivery_node_is_done(engine, job, monkeypatch):
    engine.generation_pipeline.initialize(
        job, run_id="split-terminal", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
    )
    engine.write_json(str(job / "outcome.json"), {
        "state": "done", "word": "投标文件_整册.docx", "ts": engine.now(),
    })
    monkeypatch.setattr(engine, "_body_docxs", lambda *_args: ["投标文件_整册.docx"])

    assert engine.job_state(str(job)) == "stopped"


def test_interrupting_pipeline_sessions_attempts_every_session(engine, job, monkeypatch):
    engine.PIPELINE_SESSIONS[job.name] = {"session-a", "session-b"}
    calls = []

    def interrupt(sid):
        calls.append(sid)
        return sid == "session-b"

    monkeypatch.setattr(engine, "_interrupt_or_finished", interrupt)

    assert engine._pipeline_interrupt_sessions(str(job)) is False
    assert set(calls) == {"session-a", "session-b"}


def test_stop_sets_cancel_tombstone_before_interrupting_sessions(engine, job, monkeypatch):
    owner = "owner-1"
    engine.RUNNING[job.name] = owner
    observed = []
    monkeypatch.setattr(engine, "_pipeline_interrupt_sessions", lambda _job: observed.append(
        engine.CANCEL.get(job.name)) or True)
    monkeypatch.setattr(engine, "_take_proc", lambda *_args: None)

    ok, requested = engine._stop_running_owner(str(job), job.name, owner)

    assert ok is True
    assert requested is True
    assert observed == [owner]


def test_node_retry_requires_and_deduplicates_idempotency_key(engine, job, monkeypatch):
    engine.generation_pipeline.initialize(
        job, run_id="retry-api", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01_项目理解.md"}],
    )
    engine.generation_pipeline.fail_node(
        job, "chapter_write:01", "invalid_configuration", retryable=False)
    started = []
    monkeypatch.setattr(engine, "_start_reserved_worker", lambda *args: started.append(args))

    with TestClient(engine.app) as client:
        missing = client.post(f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry")
        first = client.post(
            f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry",
            headers={"Idempotency-Key": "retry-request-0001"},
        )
        duplicate = client.post(
            f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry",
            headers={"Idempotency-Key": "retry-request-0001"},
        )

    assert missing.status_code == 400
    assert first.status_code == 200
    assert first.json()["deduplicated"] is False
    assert duplicate.status_code == 200
    assert duplicate.json()["deduplicated"] is True
    assert len(started) == 1


def test_concurrent_same_node_retry_key_is_deduplicated_before_side_effect(engine, job, monkeypatch):
    engine.generation_pipeline.initialize(
        job, run_id="retry-race", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
    )
    engine.generation_pipeline.fail_node(job, "chapter_write:01", "blocked", retryable=False)
    entered = threading.Event()
    release = threading.Event()
    original_retry = engine.generation_pipeline.retry_node
    side_effects = []

    def delayed_retry(*args):
        side_effects.append(1)
        entered.set()
        release.wait(2)
        return original_retry(*args)

    monkeypatch.setattr(engine.generation_pipeline, "retry_node", delayed_retry)
    monkeypatch.setattr(engine, "_start_reserved_worker", lambda *_args: None)
    responses = []

    def request_retry():
        with TestClient(engine.app) as client:
            responses.append(client.post(
                f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry",
                headers={"Idempotency-Key": "same-concurrent-retry"},
            ))

    first = threading.Thread(target=request_retry)
    first.start()
    assert entered.wait(1)
    second = threading.Thread(target=request_retry)
    second.start(); second.join(1)
    release.set(); first.join(1)

    assert sorted(response.status_code for response in responses) == [200, 202]
    assert sum(bool(response.json().get("deduplicated")) for response in responses) == 1
    assert next(response for response in responses if response.status_code == 202).json()["pending"] is True
    assert side_effects == [1]


def test_active_dispatch_lease_prevents_ultrafast_worker_double_dispatch(engine, job, monkeypatch):
    engine.generation_pipeline.initialize(
        job, run_id="retry-fast-worker", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
    )
    engine.generation_pipeline.fail_node(job, "chapter_write:01", "blocked", retryable=False)
    key = "fast-worker-retry-key"
    engine.write_json(str(job / ".node_retry_requests.json"), {
        "chapter_write:01|" + key: {
            "node_id": "chapter_write:01", "accepted_at": engine.now(),
            "status": "dispatching", "lease_until": time.time() + 30,
            "dispatch_token": "finished-owner",
        }
    })
    started = []
    monkeypatch.setattr(engine, "_start_reserved_worker", lambda *args: started.append(args))

    with TestClient(engine.app) as client:
        response = client.post(
            f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry",
            headers={"Idempotency-Key": key})

    assert response.status_code == 202
    assert response.json()["pending"] is True
    assert response.json()["deduplicated"] is True
    assert started == []


def test_failed_retry_dispatch_is_not_reported_as_accepted_on_duplicate(engine, job, monkeypatch):
    engine.generation_pipeline.initialize(
        job, run_id="retry-dispatch-failure", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
    )
    engine.generation_pipeline.fail_node(job, "chapter_write:01", "blocked", retryable=False)
    monkeypatch.setattr(engine, "_start_reserved_worker",
                        lambda *_args: (_ for _ in ()).throw(RuntimeError("dispatch boom")))
    headers = {"Idempotency-Key": "retry-dispatch-failure-key"}

    with TestClient(engine.app, raise_server_exceptions=False) as client:
        first = client.post(f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry", headers=headers)
        duplicate = client.post(f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry", headers=headers)

    assert first.status_code == 500
    assert first.json()["code"] == "retry_dispatch_failed"
    assert duplicate.status_code == 500
    assert duplicate.json()["ok"] is False
    assert duplicate.json()["code"] == "retry_dispatch_failed"
    assert duplicate.json()["deduplicated"] is True


def test_new_retry_key_recovers_after_previous_dispatch_failure(engine, job, monkeypatch):
    engine.generation_pipeline.initialize(
        job, run_id="retry-new-key", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
    )
    engine.generation_pipeline.fail_node(job, "chapter_write:01", "blocked", retryable=False)
    calls = []

    def dispatch(*args):
        calls.append(args)
        if len(calls) == 1:
            raise RuntimeError("dispatch boom")

    monkeypatch.setattr(engine, "_start_reserved_worker", dispatch)
    with TestClient(engine.app, raise_server_exceptions=False) as client:
        first = client.post(
            f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry",
            headers={"Idempotency-Key": "retry-dispatch-old-key"})
        second = client.post(
            f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry",
            headers={"Idempotency-Key": "retry-dispatch-new-key"})

    assert first.status_code == 500
    assert second.status_code == 200
    assert len(calls) == 2


def test_started_worker_is_not_redispatched_when_acceptance_receipt_write_fails(engine, job, monkeypatch):
    engine.generation_pipeline.initialize(
        job, run_id="retry-receipt-failure", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
    )
    engine.generation_pipeline.fail_node(job, "chapter_write:01", "blocked", retryable=False)
    starts = []

    def ultrafast_dispatch(_base, owner, *_args):
        starts.append(owner)
        engine._release_running(job.name, owner)

    original_write = engine.write_json
    ledger_path = str(job / ".node_retry_requests.json")
    ledger_writes = 0

    def fail_acceptance_receipt(path, data):
        nonlocal ledger_writes
        if str(path) == ledger_path:
            ledger_writes += 1
            if ledger_writes == 2:
                raise OSError("disk busy after worker started")
        return original_write(path, data)

    monkeypatch.setattr(engine, "_start_reserved_worker", ultrafast_dispatch)
    monkeypatch.setattr(engine, "write_json", fail_acceptance_receipt)
    with TestClient(engine.app) as client:
        first = client.post(
            f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry",
            headers={"Idempotency-Key": "retry-receipt-old-key"})
        second = client.post(
            f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry",
            headers={"Idempotency-Key": "retry-receipt-new-key"})

    assert first.status_code == 202
    assert first.json()["code"] == "retry_dispatch_confirmation_pending"
    assert second.status_code == 202
    assert len(starts) == 1


def test_stale_dispatching_retry_intent_is_reconciled_instead_of_fake_success(engine, job, monkeypatch):
    engine.generation_pipeline.initialize(
        job, run_id="retry-stale-intent", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "01", "title": "项目理解", "output": "章节_01.md"}],
    )
    engine.generation_pipeline.fail_node(job, "chapter_write:01", "blocked", retryable=False)
    key = "stale-dispatching-retry-key"
    engine.write_json(str(job / ".node_retry_requests.json"), {
        "chapter_write:01|" + key: {
            "node_id": "chapter_write:01", "accepted_at": engine.now(), "status": "dispatching",
        }
    })
    started = []
    monkeypatch.setattr(engine, "_start_reserved_worker", lambda *args: started.append(args))

    with TestClient(engine.app) as client:
        response = client.post(
            f"/v1/jobs/{job.name}/nodes/chapter_write:01/retry",
            headers={"Idempotency-Key": key})

    assert response.status_code == 200
    assert response.json()["deduplicated"] is False
    assert len(started) == 1


def test_deterministic_source_parse_failure_is_blocked_not_auto_retried(engine, tmp_path):
    job = Path(engine.jpath("bad-source"))
    job.mkdir(parents=True)
    (job / "空白.txt").write_text(" \n", encoding="utf-8")
    meta = {"name": "空白采购文件", "tender": "空白.txt", "run_id": "bad-source"}
    engine.write_json(str(job / "任务.json"), meta)
    conf = _checkpoint_config(engine)
    engine.write_json(engine.conf_path(), conf)
    engine._initialize_generation_pipeline(str(job), meta, conf)

    engine.generation_pipeline_worker(str(job))

    state = engine.generation_pipeline.load(job)
    source = next(node for node in state["nodes"] if node["id"] == "source_parse")
    assert source["state"] == "blocked"
    assert state["state"] == "blocked"
    assert engine.generation_pipeline.summary(job)["retrying"] == 0
