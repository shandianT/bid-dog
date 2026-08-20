import io
import zipfile

import pytest
from docx import Document

from template_engine import (builtin_templates, compare_instruction_coverage,
                             compile_template_instructions, derive_template,
                             extract_document_structure, recommend_template,
                             validate_package)


@pytest.mark.parametrize(
    ("text", "expected"),
    [
        ("软件平台建设，包含系统集成、数据安全、信创适配和三年运维", "government-it"),
        ("设备采购清单，要求品牌型号、技术参数、供货安装和质保", "goods"),
        ("施工总承包，评审施工组织设计、工期、质量安全和文明施工", "construction"),
        ("驻场运维外包服务，按SLA、响应时间和人员配置考核", "service"),
        ("规划咨询研究课题，要求调研方法、成果报告和专家团队", "consulting"),
    ],
)
def test_scene_recommendation_uses_procurement_content(text, expected):
    result = recommend_template("招标文件.md", text, builtin_templates())

    assert result["template_id"] == expected
    assert result["confidence"] >= 0.4
    assert result["reasons"]


def test_scene_recommendation_without_scene_evidence_uses_stable_general_fallback():
    result = recommend_template("附件.md", "请按采购文件要求响应。", builtin_templates())

    assert result["template_id"] == "government"
    assert result["confidence"] == 0.25
    assert "未识别到强场景特征" in result["reasons"][0]


def test_docx_table_content_participates_in_scene_recommendation():
    document = Document()
    document.add_paragraph("采购需求详见下表")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "序号"
    table.rows[0].cells[1].text = "采购内容"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "软件平台建设、系统集成、数据安全、信创适配和运维"
    buffer = io.BytesIO()
    document.save(buffer)

    draft = derive_template("", "优秀标书.docx", buffer.getvalue(), builtin_templates())

    assert draft["recommendation"]["template_id"] == "government-it"


def test_empty_uploaded_document_is_not_reported_as_a_successful_derivation():
    with pytest.raises(ValueError, match="未识别到可复用目录或表格结构"):
        derive_template("", "空白.txt", b"", builtin_templates())


def test_thin_uploaded_structure_is_returned_as_not_ready():
    draft = derive_template("", "片段.md", "# 技术方案\n只有一个章节。".encode(), builtin_templates())

    assert draft["validation"]["ready"] is False
    assert draft["validation"]["checks"]["source_structure"] is False
    assert draft["extraction"]["heading_count"] == 1


def test_uploaded_headings_are_generalized_before_the_template_is_returned():
    document = Document()
    for heading in (
        "深圳市某局软件平台项目技术方案",
        "张三项目经理团队配置",
        "ABC银行内部预算报价说明",
        "某公司数据安全实施方案",
    ):
        document.add_heading(heading, level=1)
    buffer = io.BytesIO()
    document.save(buffer)

    draft = derive_template("", "优秀标书.docx", buffer.getvalue(), builtin_templates())
    serialized = str(draft)

    assert draft["validation"]["ready"] is True
    assert "深圳市某局" not in serialized
    assert "张三" not in serialized
    assert "ABC银行" not in serialized
    assert "某公司" not in serialized


def test_docx_archive_with_extreme_expansion_is_rejected_before_parsing():
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b"A" * (2 * 1024 * 1024))

    with pytest.raises(ValueError, match="压缩比例异常"):
        extract_document_structure("异常.docx", buffer.getvalue())


def test_compiled_scene_package_has_materially_better_instruction_coverage_than_legacy_prompt():
    template = next(item for item in builtin_templates() if item["id"] == "construction")
    legacy = compare_instruction_coverage(template["prompt"], template)
    compiled_text = compile_template_instructions(template)
    compiled = compare_instruction_coverage(compiled_text, template)

    assert compiled["coverage"] >= 0.95
    assert compiled["coverage"] - legacy["coverage"] >= 0.70
    assert "招标文件原文高于模板" in compiled_text
    assert "不得猜测或编造" in compiled_text


def test_scene_instructions_cover_an_external_construction_rubric_better_than_the_legacy_prompt():
    template = next(item for item in builtin_templates() if item["id"] == "construction")
    rubric = [
        "项目理解", "施工组织设计", "进度计划", "资源配置", "质量管理", "安全文明", "应急预案",
        "项目管理人员表", "工程技术偏离表", "证明材料", "〔需补充〕", "招标文件原文高于模板",
    ]
    legacy_text = "围绕施工组织设计、工期计划、资源配置、质量安全和应急预案编制投标文件，逐项对应工程量、技术标准和评审要点。"
    compiled_text = compile_template_instructions(template)
    coverage = lambda text: sum(term in text for term in rubric) / len(rubric)

    assert coverage(compiled_text) >= 0.9
    assert coverage(compiled_text) - coverage(legacy_text) >= 0.4


def test_compiler_keeps_quality_rules_when_a_custom_package_is_oversized():
    template = next(item for item in builtin_templates() if item["id"] == "construction")
    template["package"]["scoring_focus"] = [
        {"name": "超长评分项%d" % index, "checks": [("检查内容" * 80) + str(index)]}
        for index in range(60)
    ]
    template["package"]["quality_rules"].append("必须保留的关键质检规则")

    compiled = compile_template_instructions(template)

    assert len(compiled) <= 20000
    assert "必须保留的关键质检规则" in compiled


def test_compiler_turns_word_format_settings_into_explicit_instructions():
    template = next(item for item in builtin_templates() if item["id"] == "construction")

    compiled = compile_template_instructions(template)

    assert "生成封面" in compiled
    assert "生成可更新目录" in compiled
    assert "跨页表格重复表头" in compiled
    assert "宽表必要时使用横向页面" in compiled
    assert "招标文件指定目录时保持原顺序" in compiled


def test_every_builtin_package_passes_the_same_quality_contract():
    for template in builtin_templates():
        validation = validate_package(template["package"])
        assert validation["ready"] is True, (template["id"], validation)
        assert validation["score"] == 100, (template["id"], validation)


def test_missing_required_package_dimension_never_counts_as_ready_even_with_high_score():
    template = next(item for item in builtin_templates() if item["id"] == "construction")
    template["package"]["sources"] = []

    validation = validate_package(template["package"])

    assert validation["score"] >= 70
    assert validation["ready"] is False
    assert "sources" in validation["warnings"]
