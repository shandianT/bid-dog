"""P2:素材按章检索喂料、《评标索引》/《投标人补料清单》、模型复核闭环。"""
import json
from pathlib import Path

import pytest
from conftest import events

import prompts


def _checkpoint_config(engine):
    conf = {"engine": {
        "kind": "s2", "s2_key": "test-key", "generation_mode": "fast",
        "s2_base_url": engine.S2_DEFAULT_BASE, "s2_model": engine.S2_DEFAULT_MODEL,
        "s2_verify_ssl": True, "s2_wire": "auto",
    }}
    conf["setup"] = {
        "model_ids": [engine.S2_DEFAULT_MODEL, engine.S2_QUALITY_MODEL],
        "text_verified_model_ids": [engine.S2_DEFAULT_MODEL, engine.S2_QUALITY_MODEL],
        "tested_connection_fingerprint": engine._connection_fingerprint(conf),
    }
    return conf


TITLES = ("技术方案", "售后服务方案", "项目实施与验收")


def _pipeline_job(engine, name, titles=TITLES):
    job = Path(engine.jpath(name))
    job.mkdir(parents=True)
    engine.generation_pipeline.initialize(
        job, run_id="run-" + name, mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[{"id": "%02d" % i, "title": t, "output": "章节_%02d_%s.md" % (i, t)}
                  for i, t in enumerate(titles, 1)])
    return job


def _materials(job):
    mat = job / "素材"
    (mat / "章节模板").mkdir(parents=True)
    (mat / "产品能力表.md").write_text("# 产品能力表\n\n| 能力 | 参数 |\n|---|---|\n| 并发 | 1 万 |\n", encoding="utf-8")
    (mat / "公司介绍.md").write_text("# 公司介绍\n\n成立于 2010 年。\n", encoding="utf-8")
    (mat / "logo.png").write_bytes(b"\x89PNG\r\n\x1a\n" + b"0" * 64)
    (mat / "章节模板" / "售后服务模板.md").write_text("# 售后服务\n\n7×24 小时响应,2 小时到场。\n", encoding="utf-8")
    (mat / "章节模板" / "施工组织模板.md").write_text("# 施工组织\n\n土建、机电安装的组织。\n", encoding="utf-8")
    return mat


# ------------------------------------------------------------------ 素材按章检索
def test_material_terms_come_from_title_slots_and_scoring_points_with_bigrams():
    import engine_v1 as engine
    terms = engine._material_terms({"title": "售后服务方案", "material_slots": ["服务承诺"],
                                    "scoring_points": ["响应时间 4 小时"], "basis": ["招标文件_解析版.md"]})
    assert {"售后服务方案", "售后", "服务", "方案", "服务承诺", "承诺", "响应时间", "小时"} <= terms
    assert "4" not in terms and "" not in terms


def test_material_selection_keeps_root_files_and_ranks_templates_by_chapter_terms(engine):
    job = _pipeline_job(engine, "mat-select")
    mat = _materials(job)
    node = {"id": "chapter_write:02", "title": "售后服务方案", "material_slots": ["服务承诺"],
            "scoring_points": ["响应时间"]}
    chosen, skipped = engine._pipeline_select_materials(str(job), node, engine._pipeline_material_paths(str(job)))
    names = [Path(p).relative_to(mat).as_posix() for p in chosen]
    assert names == ["产品能力表.md", "公司介绍.md", "章节模板/售后服务模板.md"]
    assert skipped == 2          # logo.png(二进制)+ 施工组织模板(与本章无关)


def test_material_selection_stops_at_budget_and_keeps_the_better_ranked_template(engine, monkeypatch):
    job = _pipeline_job(engine, "mat-budget")
    mat = job / "素材" / "章节模板"
    mat.mkdir(parents=True)
    (mat / "售后服务方案模板.md").write_text("零" * 2000, encoding="utf-8")          # 文件名命中,6000 字节
    (mat / "附件B.md").write_text("服务" + "零" * 2000, encoding="utf-8")            # 正文命中一次,6006 字节
    monkeypatch.setenv("BIDDOG_CHAPTER_MATERIAL_CHARS", "10000")
    node = {"id": "chapter_write:02", "title": "售后服务方案"}
    chosen, skipped = engine._pipeline_select_materials(str(job), node, engine._pipeline_material_paths(str(job)))
    assert [Path(p).name for p in chosen] == ["售后服务方案模板.md"] and skipped == 1
    monkeypatch.setenv("BIDDOG_CHAPTER_MATERIAL_CHARS", "5")                        # 预算不许小于 1 万
    chosen, _skipped = engine._pipeline_select_materials(str(job), node, engine._pipeline_material_paths(str(job)))
    assert [Path(p).name for p in chosen] == ["售后服务方案模板.md"]


def test_copy_inputs_orders_declared_inputs_before_selected_materials(engine, tmp_path):
    job = _pipeline_job(engine, "copy-order")
    _materials(job)
    declared = ["招标文件_解析版.md", "你的要求.md", "投标文件组成.md", "评分点响应矩阵.md", "废标风险清单.md"]
    for name in declared:
        (job / name).write_text("# %s\n\n内容。\n" % name, encoding="utf-8")
    engine.write_json(str(job / "response_plan.json"), {"chapters": [], "source": "local"})
    engine.write_json(engine.conf_path(), _checkpoint_config(engine))
    state = engine.generation_pipeline.load(job)
    node = {**next(n for n in state["nodes"] if n["id"] == "chapter_write:02"), "material_slots": ["服务承诺"]}
    target = tmp_path / "attempt"
    target.mkdir()

    ordered = engine._pipeline_copy_inputs(str(job), node, str(target))

    assert ordered[:6] == declared + ["response_plan.json"]
    assert ordered[6:] == ["素材/产品能力表.md", "素材/公司介绍.md", "素材/章节模板/售后服务模板.md"]
    assert (target / "SKILL.md").is_file() and (target / "素材" / "章节模板" / "售后服务模板.md").is_file()
    assert not (target / "素材" / "logo.png").exists()
    assert not (target / "素材" / "章节模板" / "施工组织模板.md").exists()
    diagnostics = [json.loads(line) for line in (job / "diagnostics.jsonl").read_text(encoding="utf-8").splitlines()]
    hit = next(d for d in diagnostics if d["code"] == "material_selected")
    assert hit["level"] == "info" and hit["context"]["node_id"] == "chapter_write:02" and "跳过 2 份" in hit["detail"]


# ------------------------------------------------------------------ 评标索引
MATRIX = ("# 评分点响应矩阵\n\n> 规划来源:模型逐项核对(m)\n\n"
          "| 序号 | 评分项 | 分值 | 响应位置 | 评估标准/证据 | 缺口 |\n|---|---|---:|---|---|---|\n"
          "| 1 | 售后服务 | 10 分 | 第2章 | 承诺 2 小时到场 | 无 |\n"
          "| 2 | 实施方案 | 5 分 | 项目实施与验收 | 里程碑计划 | 无 |\n"
          "| 3 | 类似业绩 | 3 分 | 〔需补充〕 | 三年内 2 个合同 | 缺案例 |\n")


def test_scoring_index_is_built_only_from_a_model_checked_plan(engine):
    job = _pipeline_job(engine, "index")
    (job / "评分点响应矩阵.md").write_text(MATRIX, encoding="utf-8")
    engine.write_json(str(job / "response_plan.json"), {"chapters": [], "source": "local"})
    assert engine._build_scoring_index(str(job)) == ""        # 本地索引印出来是一张全〔需补充〕的表,不出

    engine.write_json(str(job / "response_plan.json"), {"chapters": [], "source": "model", "model": "m"})
    index = engine._build_scoring_index(str(job))
    assert index.startswith("# 评标索引")
    rows = [line for line in index.splitlines() if line.startswith("| ") and not line.startswith("| 序号")]
    assert rows == ["| 1 | 售后服务 | 10 分 | 承诺 2 小时到场 | 售后服务方案 |",
                    "| 2 | 实施方案 | 5 分 | 里程碑计划 | 项目实施与验收 |",
                    "| 3 | 类似业绩 | 3 分 | 三年内 2 个合同 | 〔需补充〕 |"]
    assert "评分点合计 18 分" in index


# ------------------------------------------------------------------ 补料清单
def test_supplement_list_collects_placeholders_per_chapter_and_risk_actions(engine):
    job = _pipeline_job(engine, "supplement")
    (job / "章节_01_技术方案.md").write_text(
        "# 技术方案\n\n正文〔需补充:公司案例〕再来一次〔需补充:公司案例〕以及〔参数待核实〕。\n", encoding="utf-8")
    (job / "章节_02_售后服务方案.md").write_text(
        "# 售后服务方案\n\n〔此处粘贴:营业执照〕{{图:IMG-01}}〔配图建议:服务网点分布〕\n", encoding="utf-8")
    (job / "废标风险清单.md").write_text(
        "# 废标风险清单\n\n| 序号 | 类别 | 招标要求 | 依据 | 提交前动作 |\n|---|---|---|---|---|\n"
        "| 1 | 资格 | 提供有效营业执照 | 招标文件 3.1 | 复印件加盖公章 |\n", encoding="utf-8")

    text = engine._build_supplement_list(str(job), ["章节_01_技术方案.md", "章节_02_售后服务方案.md", "缺失.md"])

    assert text.startswith("# 投标人补料清单")
    assert [line for line in text.splitlines() if line.endswith("| 待处理 |")] == [
        "| 1 | 技术方案 | 需补充 | 公司案例 | 待处理 |",
        "| 2 | 技术方案 | 参数待核实 | — | 待处理 |",
        "| 3 | 售后服务方案 | 此处粘贴 | 营业执照 | 待处理 |",
        "| 4 | 售后服务方案 | 配图建议 | 服务网点分布 | 待处理 |",
    ]
    assert "| 1 | 资格 | 提供有效营业执照 | 复印件加盖公章 | 待核对 |" in text
    assert "人工确认节点" in text

    empty = engine._build_supplement_list(str(job), [])
    assert "正文未留占位" in empty and "复印件加盖公章" in empty
    (job / "废标风险清单.md").unlink()
    assert "未生成废标风险清单" in engine._build_supplement_list(str(job), [])


def test_index_and_supplement_files_are_never_mistaken_for_the_body(engine):
    job = _pipeline_job(engine, "not-body")
    names = ["评标索引.md", "投标人补料清单.md", "投标文件_整册.md"]
    assert engine._body_mds(str(job), names) == ["投标文件_整册.md"]


def test_assemble_puts_the_scoring_index_first_and_the_supplement_list_last(engine, monkeypatch):
    job = Path(engine.jpath("assemble-p2"))
    job.mkdir(parents=True)
    (job / "采购文件.md").write_text("# 采购需求\n" + "统一平台建设、实施、培训、验收和运维要求。" * 80, encoding="utf-8")
    titles = ["总体方案", "售后服务方案", "培训方案"]
    meta = {"name": "整册", "tender": "采购文件.md", "staged": False, "run_id": "assemble",
            "template_snapshot": {"package": {"outline": [{"title": t} for t in titles]}}}
    engine.write_json(str(job / "任务.json"), meta)
    conf = _checkpoint_config(engine)
    engine.write_json(engine.conf_path(), conf)
    engine._initialize_generation_pipeline(str(job), meta, conf)

    def model_runner(path, node, _prompt, _model):
        title = str(node.get("title") or "响应")
        if node["id"] == "response_plan":
            for output in node["outputs"]:
                if output == "response_plan.json":
                    engine.write_json(str(Path(path) / output), {"source": "model", "model": "m", "chapters": [
                        {"id": "%02d" % i, "title": t, "output": "章节_%02d_%s.md" % (i, t),
                         "basis": ["需求"], "scoring_points": [], "material_slots": [], "dependencies": []}
                        for i, t in enumerate(titles, 1)]})
                elif output == "评分点响应矩阵.md":
                    (Path(path) / output).write_text(MATRIX, encoding="utf-8")
                elif output == "废标风险清单.md":
                    (Path(path) / output).write_text(
                        "# 废标风险清单\n\n| 序号 | 类别 | 招标要求 | 依据 | 提交前动作 |\n|---|---|---|---|---|\n"
                        "| 1 | 资格 | 提供有效营业执照 | 3.1 | 复印件加盖公章 |\n", encoding="utf-8")
                else:
                    (Path(path) / output).write_text("# %s\n\n" % output + "规划行。" * 60, encoding="utf-8")
            return
        for output in node["outputs"]:
            body = ("| 序号 | 招标要求 | 投标响应 | 偏离情况 | 依据/证据 | 备注 |\n|---|---|---|---|---|---|\n| 1 | 按期 | 完全响应 | 无偏离 | 解析版 | 已核对 |\n"
                    if "偏离表" in output else "")
            hole = "〔需补充:讲师简历〕" if title == "培训方案" else ""
            (Path(path) / output).write_text("# %s\n\n%s%s%s" % (title, body, hole, "".join(
                "逐项依据招标要求响应第%d条。" % j for j in range(1, 121))), encoding="utf-8")

    def export_word(path, known, force=False):
        from docx import Document
        document = Document()
        document.add_heading("投标文件", level=1)
        document.add_paragraph("逐项响应招标要求并提供可核验依据。" * 30)
        document.save(Path(path) / "投标文件_整册.docx")
        known.add("投标文件_整册.docx")
        return ["投标文件_整册.docx"]

    def format_ok(path, _word=""):
        (Path(path) / "Word格式自检报告.md").write_text(
            "# 投标文件格式自检报告\n\n- 结论：✅ 全部通过（1 项）\n", encoding="utf-8")
        return {"status": "pass"}

    monkeypatch.setattr(engine, "_pipeline_model_runner", model_runner)
    monkeypatch.setattr(engine, "ensure_docx", export_word)
    monkeypatch.setattr(engine, "word_format_audit", format_ok)
    monkeypatch.setattr(engine, "settle", lambda path, commit=True, **_k: {"state": "done" if commit else "ready"})
    monkeypatch.setattr(engine, "delivery_summary", lambda path: {"ready": True})

    engine.generation_pipeline_worker(str(job))

    assert engine.generation_pipeline.load(job)["state"] == "done"
    book = (job / "投标文件_整册.md").read_text(encoding="utf-8")
    assert book.startswith("# 评标索引")
    assert "| 1 | 售后服务 | 10 分 | 承诺 2 小时到场 | 售后服务方案 |" in book
    assert book.rstrip().endswith("方可作为正式投标文件提交。")
    assert book.index("# 总体方案") < book.index("# 培训方案") < book.index("# 投标人补料清单")
    supplement = (job / "投标人补料清单.md").read_text(encoding="utf-8")
    assert "| 1 | 培训方案 | 需补充 | 讲师简历 | 待处理 |" in supplement
    assert "| 1 | 资格 | 提供有效营业执照 | 复印件加盖公章 | 待核对 |" in supplement
    assert (job / "评标索引.md").read_text(encoding="utf-8").startswith("# 评标索引")
    artifacts = [e["name"] for e in events(job) if e.get("type") == "artifact"]
    assert "评标索引.md" in artifacts and "投标人补料清单.md" in artifacts


# ------------------------------------------------------------------ 模型复核闭环
REVIEW = ("总体结论:补料后可提交。\n\n"
          "| 章节 | 问题 | 级别 | 修订建议 |\n|---|---|---|---|\n"
          "| 售后服务方案 | 漏答「响应时间」评分点 | 必办 | 补一段 2 小时到场承诺并引用评分办法 3.2 |\n"
          "| 第1章 | 表格用了泛化表头 | 建议 | 按用途命名列 |\n"
          "| 无 | 无 | — | — |\n")


def test_review_task_pins_a_table_the_engine_can_parse(engine):
    task = prompts.review_task()
    assert prompts.REVIEW_TABLE_HEADER in task and "必办" in task and "建议" in task
    assert engine._pipeline_direct_task("job", {"id": "quality_review"}) == task


def test_review_report_rows_become_actionable_gaps(engine):
    job = _pipeline_job(engine, "review")
    assert engine._review_report_gaps(str(job)) == []
    (job / "模型复核报告.md").write_text(REVIEW, encoding="utf-8")

    gaps = engine._review_report_gaps(str(job))

    assert [g["level"] for g in gaps] == ["red", "yellow"]
    assert gaps[0]["title"] == "复核:漏答「响应时间」评分点"
    redo = gaps[0]["actions"][0]
    assert redo["act"] == "redo" and redo["label"] == "重做这一章"
    assert redo["param"].startswith("重写章节「售后服务方案」:漏答「响应时间」评分点;补一段 2 小时到场承诺")
    assert gaps[0]["actions"][-1] == {"act": "open_artifact", "label": "打开复核报告", "file": "模型复核报告.md"}
    assert gaps[1]["detail"] == "第1章 · 按用途命名列"
    assert gaps[1]["actions"][0]["act"] == "redo" and "「技术方案」" in gaps[1]["actions"][0]["param"]


class _GreenGate:
    def audit(self, _path, **_kw):
        return {"level": "green", "plan": [], "items": [], "images": [], "chapters": ["a"], "total_chars": 5000}

    def apply_fix(self, *_args, **_kwargs):
        return 0

    def write_report(self, path, *_args, **_kwargs):
        Path(path).write_text("# 成品质检报告\n", encoding="utf-8")


def _delivered_job(engine, name):
    from docx import Document
    job = _pipeline_job(engine, name)
    engine.write_json(str(job / "任务.json"), {"name": name, "tender": "采购文件.md", "created_at": engine.now()})
    (job / "投标文件_整册.md").write_text("# 投标文件\n\n" + "正文。" * 800, encoding="utf-8")
    document = Document()
    document.add_heading("投标文件", level=1)
    document.add_paragraph("逐项响应招标要求并提供可核验依据。" * 30)
    document.save(job / "投标文件_整册.docx")
    return job


@pytest.mark.parametrize("level, health, status, counts", [
    ("必办", "red", "fail", (1, 0)),
    ("建议", "yellow", "warning", (0, 1)),
])
def test_quality_audit_folds_review_findings_into_the_pre_delivery_check(engine, monkeypatch, level, health, status, counts):
    job = _delivered_job(engine, "audit-" + health)
    (job / "模型复核报告.md").write_text(
        "| 章节 | 问题 | 级别 | 修订建议 |\n|---|---|---|---|\n| 售后服务方案 | 漏答响应时间 | %s | 补一段承诺 |\n" % level,
        encoding="utf-8")
    monkeypatch.setattr(engine, "_skill_module", lambda name: _GreenGate() if name == "quality_gate" else None)

    result = engine.quality_audit(str(job), set(engine.list_deliverables(str(job))))

    assert result["status"] == status and result["level"] == health
    healths = [e for e in events(job) if e.get("type") == "health"]
    assert healths and healths[-1]["level"] == health
    gap = next(g for g in healths[-1]["gaps"] if g["title"].startswith("复核:"))
    assert gap["level"] == health and gap["actions"][0]["act"] == "redo"
    assert "「售后服务方案」" in gap["actions"][0]["param"]
    message = [e for e in events(job) if e.get("type") == "message"][-1]
    assert "模型复核:必办 %d 项、建议 %d 项" % counts in message["text"]


def test_quality_audit_stays_green_without_a_review_report(engine, monkeypatch):
    job = _delivered_job(engine, "audit-clean")
    monkeypatch.setattr(engine, "_skill_module", lambda name: _GreenGate() if name == "quality_gate" else None)
    result = engine.quality_audit(str(job), set(engine.list_deliverables(str(job))))
    assert result["status"] == "pass" and result["level"] == "green"
    assert not [e for e in events(job) if e.get("type") == "health"]
