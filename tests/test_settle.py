from pathlib import Path
import hashlib
import os

import pytest
from docx import Document
from fastapi.testclient import TestClient

from conftest import events


def _texts(job):
    return [str(e.get("text") or "") for e in events(job) if e.get("type") in ("message", "error")]


def _latest(job, kind):
    return [e for e in events(job) if e.get("type") == kind][-1]


def _write_body_docx(path, paragraphs=80):
    document = Document()
    document.add_heading("投标文件", level=1)
    for index in range(paragraphs):
        document.add_paragraph("第%d项完整响应，满足招标文件要求并提供实施说明。" % (index + 1))
    document.save(path)


def _sha256(path):
    return hashlib.sha256(Path(path).read_bytes()).hexdigest()


def test_body_word_is_the_only_direct_success(engine, job, monkeypatch):
    _write_body_docx(job / "投标文件_技术标.docx")
    monkeypatch.setattr(engine, "quality_audit", lambda *_: None)
    monkeypatch.setattr(engine, "delivery_summary", lambda *_args, **_kwargs: {"ready": True})

    result = engine.settle(str(job))

    assert result["state"] == "done"
    assert engine.job_state(str(job)) == "done"
    assert _latest(job, "progress")["step"] == 12
    assert any("任务完成，已整理好" in text for text in _texts(job))


@pytest.mark.parametrize("payload", [b"x", b"PK" + b"not-a-real-word-file" * 300])
def test_corrupt_or_fake_docx_never_opens_the_delivery_gate(engine, job, payload):
    (job / "投标文件_技术标.docx").write_bytes(payload)

    result = engine.settle(str(job))

    assert result["state"] == "stopped"
    assert engine.job_state(str(job)) == "stopped"


def test_word_format_audit_flags_style_misses_as_warn_not_delivery_failure(engine, job):
    # 0.20.5 起:Word 能打开、检查能跑完,样式项不符 = warn(提交前人工确认),
    # 不再把整单判失败——「生成了 Word 还报错」是真机反馈的直接来源。
    _write_body_docx(job / "投标文件_技术标.docx", paragraphs=2)

    result = engine.word_format_audit(str(job), "投标文件_技术标.docx")

    assert result["status"] == "warn"
    assert result["failed"] > 0
    assert "不符" in result["summary"]
    report = (job / "Word格式自检报告.md").read_text(encoding="utf-8")
    assert "❌" in report

    # 检查根本跑不成(文件损坏)才是 fail:这才是真正的交付失败
    (job / "投标文件_技术标.docx").write_bytes(b"not a real docx")
    broken = engine.word_format_audit(str(job), "投标文件_技术标.docx")
    assert broken["status"] == "fail"


def test_word_format_report_is_bound_to_the_exact_docx_bytes(engine, job):
    word = job / "投标文件_技术标.docx"
    _write_body_docx(word, paragraphs=10)
    digest = engine._file_digest(str(word))
    (job / "Word格式自检报告.md").write_text(
        "# 格式报告\n\n- SHA-256：`%s`\n- 结论：✅ 全部通过（1 项）\n" % digest,
        encoding="utf-8",
    )
    assert engine._word_format_status(str(job), word.name)["status"] == "pass"

    document = Document(word)
    document.add_paragraph("质检后改写 Word")
    document.save(word)

    assert engine._word_format_status(str(job), word.name)["status"] == "stale"


def test_delivery_cache_cannot_hide_same_size_same_mtime_word_replacement(engine, job):
    word = job / "投标文件_技术标.docx"
    _write_body_docx(word, paragraphs=10)
    digest = engine._file_digest(str(word))
    report = job / "Word格式自检报告.md"
    report.write_text(
        "# 格式报告\n\n- SHA-256：`%s`\n- 结论：✅ 全部通过（1 项）\n" % digest,
        encoding="utf-8",
    )
    engine.generation_pipeline.initialize(
        job, run_id="delivery-cache", mode="fast",
        model_routes={"fast": "fast", "quality": "quality"}, chapters=[],
    )
    first = engine.delivery_summary(str(job), quality={
        "status": "pass", "level": "green", "summary": "关键检查已通过"})
    assert first["format"]["status"] == "pass"
    cached_signature = engine._delivery_signature(str(job))
    original = word.read_bytes()
    original_stat = word.stat()
    changed = bytearray(original)
    changed[-1] ^= 1
    word.write_bytes(changed)
    os.utime(word, ns=(original_stat.st_atime_ns, original_stat.st_mtime_ns))
    assert engine._delivery_signature(str(job)) != cached_signature

    second = engine.delivery_summary(str(job))

    assert second["ready"] is False
    # 替换后的 Word 不能藏在旧的绿色结论后面:stale 触发现场重检,
    # 结论必须绑定当前字节——绝不允许直接沿用 pass。
    assert second["format"]["status"] != "pass"
    assert second["format"]["word_sha256_bound"] is True


def test_failed_redo_cannot_reuse_the_previous_word_as_new_success(engine, job):
    word = job / "投标文件_技术标.docx"
    _write_body_docx(word)
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["redo_baseline"] = {"docx": {word.name: _sha256(word)}, "md": {}}
    engine.write_json(str(job / "任务.json"), meta)

    result = engine.settle(str(job), stop_reason="已停止（定向重做失败）")

    assert result["state"] == "stopped"
    assert engine.job_state(str(job)) == "stopped"


@pytest.mark.parametrize("instruction", ["只改项目进度安排", "重写" * 1000])
def test_mock_redo_captures_baseline_before_dispatch(engine, job, monkeypatch, instruction):
    word = job / "投标文件_技术标.docx"
    _write_body_docx(word)
    engine.write_json(
        str(job / "outcome.json"),
        {"state": "done", "word": word.name, "ts": engine.now()},
    )
    monkeypatch.setattr(engine, "config_agent_cmd", lambda: None)
    monkeypatch.setattr(engine.time, "sleep", lambda _seconds: None)

    class ImmediateThread:
        def __init__(self, target, args=(), **_kwargs):
            self.target = target
            self.args = args

        def start(self):
            self.target(*self.args)

    class ImmediateThreading:
        Thread = ImmediateThread

    monkeypatch.setattr(engine, "threading", ImmediateThreading)
    with TestClient(engine.app) as client:
        response = client.post("/v1/jobs/job-1/redo", json={"instruction": instruction})

    assert response.status_code == 200
    assert _sha256(word) == engine.read_json(str(job / "任务.json"), {})["redo_baseline"]["docx"][word.name]
    assert engine.job_state(str(job)) == "stopped"
    assert engine.read_json(str(job / "outcome.json"), {})["state"] == "stopped"
    assert not any(e.get("pct") == 100 and e.get("step") == 12 for e in events(job))


def test_body_markdown_must_export_word_before_success(engine, job, monkeypatch):
    (job / "投标文件_技术标.md").write_text("# 正文\n" + "完整响应。" * 800, encoding="utf-8")

    def make_word(job_path, known):
        name = "投标文件_技术标.docx"
        _write_body_docx(Path(job_path, name))
        known.add(name)
        return [name]

    monkeypatch.setattr(engine, "ensure_docx", make_word)
    monkeypatch.setattr(engine, "quality_audit", lambda *_: None)
    monkeypatch.setattr(engine, "delivery_summary", lambda *_args, **_kwargs: {"ready": True})

    result = engine.settle(str(job))

    assert result["state"] == "done"
    assert engine.job_state(str(job)) == "done"
    assert any("任务完成，已整理好" in text for text in _texts(job))


def test_analysis_only_is_stopped_with_real_progress_and_log_tail(engine, job):
    (job / "招标文件_解析版.md").write_text("解析依据" * 300, encoding="utf-8")
    (job / "run.log").write_text("\n".join("TAIL-%02d" % i for i in range(1, 13)), encoding="utf-8")

    result = engine.settle(str(job))

    assert result["state"] == "stopped"
    assert engine.job_state(str(job)) == "stopped"
    progress = _latest(job, "progress")
    assert progress["step"] == 1
    assert progress["pct"] == 2
    assert "生成中断" in progress["stage"]
    errors = [e for e in events(job) if e.get("type") == "error"]
    assert "TAIL-05" in errors[-1]["text"]
    assert "TAIL-12" in errors[-1]["text"]
    assert {a["act"] for a in errors[-1]["actions"]} >= {"open_log", "rerun"}
    assert not any("任务完成，已整理好" in text for text in _texts(job))
    health = _latest(job, "health")
    assert health["level"] == "red"
    assert "没有可交付的 Word" in health["summary"]
    assert "修复" not in str(health)


def test_empty_output_is_stopped_without_fake_completion(engine, job):
    result = engine.settle(str(job))

    assert result["state"] == "stopped"
    assert engine.job_state(str(job)) == "stopped"
    assert "没有产出" in _latest(job, "progress")["stage"]
    assert not any("任务完成，已整理好" in text for text in _texts(job))


def test_markdown_export_failure_never_becomes_done(engine, job, monkeypatch):
    (job / "投标文件_技术标.md").write_text("# 正文\n" + "完整响应。" * 800, encoding="utf-8")
    monkeypatch.setattr(engine, "ensure_docx", lambda *_: [])

    result = engine.settle(str(job))

    assert result["state"] == "stopped"
    assert engine.job_state(str(job)) == "stopped"
    health = _latest(job, "health")
    assert "没有可交付的 Word" in health["summary"]
    assert any(a["act"] == "export_docx" for g in health["gaps"] for a in g.get("actions", []))
    assert not any("任务完成，已整理好" in text for text in _texts(job))


def test_partial_chapters_are_not_misclassified_as_complete_body(engine, job, monkeypatch):
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["oc_session"] = "session-can-resume"
    engine.write_json(str(job / "任务.json"), meta)
    for number, title in ((1, "项目理解"), (4, "实施方案"), (5, "资格证明文件")):
        (job / ("第%d章_%s.md" % (number, title))).write_text(
            "# %s\n" % title + "章节正文。" * 400, encoding="utf-8"
        )
    (job / "run.log").write_text(
        "\x1b[91mError:\x1b[0m Error reading stream: stream idle timeout: "
        "no data received within configured window\n",
        encoding="utf-8",
    )

    def unexpected_export(*_args, **_kwargs):
        raise AssertionError("分章过程稿不得触发完整正文 Word 导出")

    monkeypatch.setattr(engine, "ensure_docx", unexpected_export)

    result = engine.settle(str(job), stop_reason="已停止（连接中断，内容已保留）")

    assert result["state"] == "stopped"
    assert engine.job_state(str(job)) == "stopped"
    error = _latest(job, "error")
    assert "已生成 3 个章节" in error["text"]
    assert "还没有汇总成完整正文和最终 Word" in error["text"]
    assert "正文稿已经生成" not in error["text"]
    assert "Word 导出失败" not in error["text"]
    assert "\x1b" not in error["text"]
    assert {action["act"] for action in error["actions"]} >= {"resume", "open_log"}


def test_legacy_word_export_error_is_corrected_when_only_chapters_exist(engine, job):
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["oc_session"] = "legacy-session"
    engine.write_json(str(job / "任务.json"), meta)
    (job / "第5章_资格证明文件.md").write_text("章节正文。" * 400, encoding="utf-8")

    safe = engine.sanitize_event(job, {
        "type": "error",
        "text": "正文稿已经生成，但最终 Word 导出失败，因此这单仍未完成。",
        "actions": [{"act": "export_docx", "label": "重试导出 Word"}],
    })

    assert "已生成 1 个章节" in safe["text"]
    assert "Word 导出失败" not in safe["text"]
    assert {action["act"] for action in safe["actions"]} >= {"resume", "open_log"}


@pytest.mark.parametrize("name", ["招标文件_解析版.md", "成品质检报告.md", "废标风险清单.md"])
def test_analysis_files_are_never_offered_for_repair(engine, job, name):
    (job / name).write_text("分析依据" * 500, encoding="utf-8")

    engine.quality_audit(str(job), set(engine.list_deliverables(str(job))))

    emitted = events(job)
    assert not any(a.get("act") == "repair" for e in emitted for a in e.get("actions", []))
    assert not any(
        a.get("act") == "repair"
        for e in emitted
        for gap in e.get("gaps", [])
        for a in gap.get("actions", [])
    )
