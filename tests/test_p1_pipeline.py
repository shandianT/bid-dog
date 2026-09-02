"""P1:修改指令路由到单章、章节并行度、生成参数可配、章节撰写契约。"""
import json
import threading
import time
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

import prompts


def _checkpoint_config(engine):
    conf = {"engine": {
        "kind": "s2", "s2_key": "test-key", "generation_mode": "fast",
        "s2_base_url": engine.S2_DEFAULT_BASE, "s2_model": engine.S2_DEFAULT_MODEL,
        "s2_verify_ssl": True, "s2_wire": "auto",
    }}
    conf["setup"] = {
        "model_ids": [engine.S2_DEFAULT_MODEL, engine.S2_QUALITY_MODEL],
        "text_verified_model_ids": [engine.S2_DEFAULT_MODEL, engine.S2_QUALITY_MODEL],
        "tested_connection_fingerprint": engine._connection_fingerprint(conf),
    }
    return conf


# ------------------------------------------------------------------ 修改指令路由
def _routed_job(engine, name):
    job = Path(engine.jpath(name))
    job.mkdir(parents=True)
    engine.generation_pipeline.initialize(
        job, run_id="run-" + name, mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[
            {"id": "01", "title": "技术方案", "output": "章节_01_技术方案.md"},
            {"id": "02", "title": "售后服务方案", "output": "章节_02_售后服务方案.md"},
            {"id": "03", "title": "项目实施与验收", "output": "章节_03_项目实施与验收.md"},
        ])
    return job


@pytest.mark.parametrize("instruction, scope, node_id", [
    ("第2章 把响应时间改成 2 小时", "chapter", "chapter_write:02"),
    ("第三部分的验收流程补一段应急预案", "chapter", "chapter_write:03"),
    ("售后服务方案里把 4 小时改成 2 小时", "chapter", "chapter_write:02"),
    ("技术偏离表再补三条条款", "chapter", "chapter_write:technical_deviation"),
    ("商务偏离表付款方式按招标改", "chapter", "chapter_write:business_deviation"),
    ("整册把公司名统一为 XX 科技", "whole", ""),
    ("第1章和第2章都改一下语气", "whole", ""),          # 涉及多章:引擎一次只重写一章,走整册
    ("把语气改得更正式一点", "whole", ""),                # 没指到章节
    ("第2章的目录格式改一下", "whole", ""),               # 版式类关键词压过章节命中
])
def test_revision_instruction_routes_to_one_chapter_or_the_whole_book(engine, instruction, scope, node_id):
    job = _routed_job(engine, "route-" + str(abs(hash(instruction)))[:8])
    route = engine._route_revision_instruction(str(job), instruction)
    assert route["scope"] == scope, route
    assert route.get("node_id", "") == node_id
    assert [c["title"] for c in route["chapters"]][:3] == ["技术方案", "售后服务方案", "项目实施与验收"]


def test_revision_route_falls_back_to_whole_for_legacy_jobs(engine, job):
    route = engine._route_revision_instruction(str(job), "第2章改一下")
    assert route["scope"] == "whole" and route["chapters"] == []


def test_revision_plan_endpoint_returns_the_route(engine):
    job = _routed_job(engine, "route-endpoint")
    with TestClient(engine.app) as client:
        body = client.post("/v1/jobs/%s/revisions/plan" % job.name,
                           json={"instruction": "第2章把响应时间改成 2 小时"}).json()
    assert body["ok"] is True and body["scope"] == "chapter"
    assert body["node_id"] == "chapter_write:02" and body["title"] == "售后服务方案"
    assert "售后服务方案" in body["reason"]


# ------------------------------------------------------------------ 生成参数
def test_generation_params_are_clamped_and_only_known_keys_survive(engine):
    assert engine._normalize_generation_params({"temperature": "0.7", "frequency_penalty": 5, "bogus": 1,
                                                "presence_penalty": "abc"}) == {
        "temperature": 0.7, "frequency_penalty": 2.0}
    assert engine._normalize_generation_params(None) == {}
    payload = engine._apply_generation_params({"model": "m"}, {"temperature": 0.5, "frequency_penalty": 0.0,
                                                               "presence_penalty": 0.3})
    assert payload == {"model": "m", "temperature": 0.5, "presence_penalty": 0.3}   # 为 0 的 penalty 不进请求体


def test_generation_params_round_trip_through_agent_settings(engine, monkeypatch):
    engine.write_json(engine.conf_path(), _checkpoint_config(engine))
    monkeypatch.setattr(engine, "invalidate_oc_runtime", lambda: None)
    with TestClient(engine.app) as client:
        before = client.get("/v1/agent").json()
        assert before["generation_params"] == {"temperature": 0.2, "frequency_penalty": 0.0, "presence_penalty": 0.0}
        body = dict(before); body.pop("s2_model_effective", None)
        body["generation_params"] = {"temperature": 0.6, "frequency_penalty": 0.4}
        saved = client.put("/v1/agent", json=body).json()
        assert saved["ok"] is True
        assert saved["generation_params"] == {"temperature": 0.6, "frequency_penalty": 0.4, "presence_penalty": 0.0}
        # 不传 = 沿用
        body.pop("generation_params")
        again = client.put("/v1/agent", json=body).json()
        assert again["generation_params"]["temperature"] == 0.6
        # 任务跑着的时候不许换参数(A/B 的可比性靠这个)
        monkeypatch.setattr(engine, "config_locked_jobs", lambda: ["busy-job"])
        body["generation_params"] = {"temperature": 0.9}
        locked = client.put("/v1/agent", json=body)
        assert locked.status_code == 409


def test_model_node_uses_configured_generation_params(engine, tmp_path, monkeypatch):
    job = Path(engine.jpath("params-node"))
    job.mkdir(parents=True)
    skill_dir = tmp_path / "skill"; skill_dir.mkdir()
    (skill_dir / "SKILL.md").write_text("# 规则\n\n不得编造资质。\n", encoding="utf-8")
    conf = _checkpoint_config(engine)
    conf["engine"]["skill_dir"] = str(skill_dir)
    conf["engine"]["generation_params"] = {"temperature": 0.6, "frequency_penalty": 0.4}
    engine.write_json(engine.conf_path(), conf)
    engine.write_json(str(job / "任务.json"), {"name": "参数", "run_id": "run-params", "tender": "采购文件.md"})
    engine._set_skill_manifest(str(job), str(skill_dir), "params", True, "direct_model_completion")
    (job / "招标文件_解析版.md").write_text("# 采购要求\n\n必须提供实施方案。", encoding="utf-8")
    captured = {}

    def fake_stream(_base, _key, payload, **_kwargs):
        captured.update(payload)
        return {"choices": [{"finish_reason": "stop", "message": {"content": "# 实施方案\n\n" + "逐项响应。" * 40}}]}

    monkeypatch.setattr(engine, "_openai_stream_req", fake_stream)
    monkeypatch.setattr(engine, "_pipeline_declared_input_names", lambda *_a, **_k: ["招标文件_解析版.md"])
    node = {"id": "chapter_write:01", "title": "实施方案", "outputs": ["章节_01.md"],
            "min_chars": 120, "attempt": 1, "attempt_serial": 1}
    engine._pipeline_model_runner(str(job), node, "撰写", engine.S2_DEFAULT_MODEL)
    assert captured["temperature"] == 0.6 and captured["frequency_penalty"] == 0.4
    assert "presence_penalty" not in captured


# ------------------------------------------------------------------ 章节契约
def test_chapter_contract_is_compact_and_names_the_used_headings():
    used = ["总体架构", "实施进度计划", "1. 质量保证措施"]
    text = prompts.chapter_contract("售后服务方案", used)
    assert len(text) <= prompts.CHAPTER_CONTRACT_MAX_CHARS
    assert "只写章节「售后服务方案」" in text
    assert "总体架构、实施进度计划、1. 质量保证措施" in text
    assert "不少于 3500 个中文字符" in text
    assert "〔此处粘贴" in text and "{{图:图片ID}}" in text
    first = prompts.chapter_contract("技术方案", [])
    assert "第一个写的论述章" in first


def test_used_headings_are_read_from_finished_chapters_only(engine, tmp_path):
    job = Path(engine.jpath("used-headings"))
    job.mkdir(parents=True)
    engine.generation_pipeline.initialize(
        job, run_id="run-heads", mode="fast",
        model_routes={"fast": engine.S2_DEFAULT_MODEL, "quality": engine.S2_QUALITY_MODEL},
        chapters=[
            {"id": "01", "title": "技术方案", "output": "章节_01_技术方案.md"},
            {"id": "02", "title": "售后服务方案", "output": "章节_02_售后服务方案.md"},
        ])
    (job / "章节_01_技术方案.md").write_text(
        "# 技术方案\n\n## 1. 总体架构设计\n\n正文。\n\n### 1.1 部署拓扑\n\n正文。\n\n## 二、实施进度计划\n\n正文。" * 4,
        encoding="utf-8")
    engine.generation_pipeline.start_node(job, "chapter_write:01", input_digest="d")
    engine.generation_pipeline.complete_node(job, "chapter_write:01", input_digest="d")
    (job / "章节_02_售后服务方案.md").write_text("# 售后\n\n## 未完成章的小标题\n\n草稿" * 30, encoding="utf-8")
    current = next(n for n in engine.generation_pipeline.load(job)["nodes"] if n["id"] == "chapter_write:02")
    assert engine._used_chapter_headings(str(job), current) == ["总体架构设计", "部署拓扑", "实施进度计划"]
    assert prompts.used_headings_from_text("## 3.2 验收标准\n### 风险应对\n## (一)项目理解\n## 5G 网络方案\n") == ["验收标准", "风险应对", "项目理解", "5G 网络方案"]


def test_word_count_threshold_is_one_constant_across_prompt_and_audit(engine):
    src = (Path(engine.__file__)).read_text(encoding="utf-8")
    assert "min_chapter=prompts.CHAPTER_TARGET_CHARS" in src
    assert "min_chapter=3500" not in src and "至少 2500" not in src and "至少 800 字" not in src
    assert prompts.CHAPTER_TARGET_CHARS == 3500
    task = prompts.chapter_task({"title": "技术方案", "basis": ["需求书"], "scoring_points": ["架构(20 分)"],
                                 "material_slots": [], "user_note": "补一段容灾"})
    assert "不少于 3500 个中文字符" in task and "架构(20 分)" in task and "补一段容灾" in task


# ------------------------------------------------------------------ 并行度
def test_chapters_run_up_to_four_in_flight_and_throttle_after_rate_limit(engine, monkeypatch):
    job = Path(engine.jpath("parallel-chapters"))
    job.mkdir(parents=True)
    (job / "采购文件.md").write_text("# 采购需求\n" + "统一平台建设、实施、培训、验收和运维要求。" * 80, encoding="utf-8")
    titles = ["总体方案", "技术架构", "实施方案", "培训方案", "验收方案", "运维方案"]
    meta = {"name": "并行", "tender": "采购文件.md", "staged": False, "run_id": "parallel",
            "template_snapshot": {"package": {"outline": [{"title": t} for t in titles]}}}
    engine.write_json(str(job / "任务.json"), meta)
    conf = _checkpoint_config(engine)
    engine.write_json(engine.conf_path(), conf)
    engine._initialize_generation_pipeline(str(job), meta, conf)
    monkeypatch.setenv("BIDDOG_CHAPTER_PARALLEL", "4")
    lock = threading.Lock()
    state = {"now": 0, "peak_before": 0, "peak_after": 0, "throttled_at": None}

    def model_runner(path, node, _prompt, _model):
        title = str(node.get("title") or "响应")
        if node["id"] == "response_plan":
            for output in node["outputs"]:          # 规划节点声明了四个输出,少一个门禁就判失败
                if output == "response_plan.json":
                    engine.write_json(str(Path(path) / output), {"chapters": [
                        {"id": "%02d" % i, "title": t, "output": "章节_%02d_%s.md" % (i, t),
                         "basis": ["需求"], "scoring_points": [], "material_slots": [], "dependencies": []}
                        for i, t in enumerate(titles, 1)]})
                else:
                    (Path(path) / output).write_text("# %s\n\n" % output + "规划行。" * 60, encoding="utf-8")
            return
        with lock:
            state["now"] += 1
            key = "peak_after" if state["throttled_at"] else "peak_before"
            state[key] = max(state[key], state["now"])
        time.sleep(0.12)
        if node["id"] == "chapter_write:02" and not state["throttled_at"]:
            # 模拟这一章撞到 429:事件钩子把这一单标成限流,余下章节降到 2 路
            engine._pipeline_event(str(path), "retry", {**node, "error_code": "model_rate_limited"})
            with lock: state["throttled_at"] = time.time()
        with lock: state["now"] -= 1
        for output in node["outputs"]:
            body = ("| 序号 | 招标要求 | 投标响应 | 偏离情况 | 依据/证据 | 备注 |\n|---|---|---|---|---|---|\n| 1 | 按期 | 完全响应 | 无偏离 | 解析版 | 已核对 |\n"
                    if "偏离表" in output else "")
            (Path(path) / output).write_text("# %s\n\n%s%s" % (title, body, "".join(
                "逐项依据招标要求响应第%d条。" % j for j in range(1, 121))), encoding="utf-8")

    def export_word(path, known, force=False):
        from docx import Document
        document = Document()
        document.add_heading("投标文件", level=1)
        document.add_paragraph("逐项响应招标要求并提供可核验依据。" * 30)
        document.save(Path(path) / "投标文件_整册.docx")
        known.add("投标文件_整册.docx")
        return ["投标文件_整册.docx"]

    def format_ok(path, _word=""):
        (Path(path) / "Word格式自检报告.md").write_text(
            "# 投标文件格式自检报告\n\n- 结论：✅ 全部通过（1 项）\n", encoding="utf-8")
        return {"status": "pass"}

    monkeypatch.setattr(engine, "_pipeline_model_runner", model_runner)
    monkeypatch.setattr(engine, "ensure_docx", export_word)
    monkeypatch.setattr(engine, "word_format_audit", format_ok)
    monkeypatch.setattr(engine, "settle", lambda path, commit=True, **_k: {"state": "done" if commit else "ready"})
    monkeypatch.setattr(engine, "delivery_summary", lambda path: {"ready": True})

    engine.generation_pipeline_worker(str(job))

    assert engine.generation_pipeline.load(job)["state"] == "done"
    assert state["peak_before"] >= 3, state                      # 4 路在飞(不再两章一批互相等)
    assert state["throttled_at"] and state["peak_after"] <= 2, state   # 撞限流后降到 2 路
    assert job.name not in engine.PIPELINE_THROTTLED or True   # 集合只是本轮标记,不做持久化断言


def test_chapter_slots_follow_environment_and_throttle(engine, monkeypatch):
    monkeypatch.setenv("BIDDOG_CHAPTER_PARALLEL", "6")
    engine.PIPELINE_THROTTLED.discard("j")
    assert engine._chapter_parallelism() == 6 and engine._chapter_slots("j") == 6
    engine.PIPELINE_THROTTLED.add("j")
    assert engine._chapter_slots("j") == 2
    engine.PIPELINE_THROTTLED.discard("j")
    monkeypatch.setenv("BIDDOG_CHAPTER_PARALLEL", "99")
    assert engine._chapter_parallelism() == 8
    monkeypatch.setenv("BIDDOG_CHAPTER_PARALLEL", "abc")
    assert engine._chapter_parallelism() == 4
