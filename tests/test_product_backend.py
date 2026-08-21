import io
import json
import threading
import zipfile
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from conftest import events


def _job_by_id(client, job_id="job-1"):
    return next(item for item in client.get("/v1/jobs").json() if item["job_id"] == job_id)


def _write_word(path, with_toc=False):
    document = Document()
    document.add_heading("投标文件", level=1)
    if with_toc:
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        field = OxmlElement("w:instrText")
        field.set(qn("xml:space"), "preserve")
        field.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(field)
    for index in range(80):
        document.add_paragraph("第%d项完整响应，满足采购要求并提供实施说明。" % (index + 1))
    document.save(path)


def test_job_list_exposes_five_state_presentation_and_operational_summary(engine, job, monkeypatch):
    monkeypatch.setattr(engine, "now", lambda: "2026-08-08 12:00:00")
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["created_at"] = "2026-08-08 11:55:00"
    meta["staged"] = True
    engine.write_json(str(job / "任务.json"), meta)

    with TestClient(engine.app) as client:
        item = _job_by_id(client)

    assert item["state"] == "staged"  # legacy API remains intact
    assert item["presentation"] == {"code": "preparing", "label": "准备中"}
    assert item["status"] == "准备中"
    assert item["current_action"]
    assert item["last_activity_at"]
    assert item["eta"] >= 0
    assert item["elapsed"] >= 0
    assert item["usage"]["calls"] == 0


def test_legacy_unknown_is_presented_as_failed_not_unknown(engine, job):
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["staged"] = False
    engine.write_json(str(job / "任务.json"), meta)
    engine.write_json(
        str(job / "progress.json"),
        {"type": "progress", "stage": "生成中", "pct": 8, "step": 1, "total": 12, "ts": engine.now()},
    )

    with TestClient(engine.app) as client:
        item = _job_by_id(client)

    assert item["state"] == "unknown"
    assert item["presentation"]["code"] == "incomplete"
    assert item["presentation"]["label"] == "未完成"
    assert "状态不明" not in item["current_action"]


def test_cli_compatibility_runtime_disables_pause_with_a_friendly_reason(engine, job):
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["staged"] = False
    engine.write_json(str(job / "任务.json"), meta)
    engine.write_json(
        str(job / "runtime.json"),
        {
            "execution_path": "cli_compat",
            "can_pause": False,
            "pause_disabled_reason": "稳定兼容模式运行时暂不支持暂停；如需中止，请使用停止。",
        },
    )
    owner = engine._reserve_running(job.name)
    try:
        with TestClient(engine.app) as client:
            item = _job_by_id(client)
    finally:
        engine._release_running(job.name, owner)

    assert "pause" not in item["can"]
    assert item["runtime"]["mode"] == "compatibility"
    assert item["runtime"]["capabilities"]["pause"]["enabled"] is False
    assert "暂不支持暂停" in item["runtime"]["capabilities"]["pause"]["reason"]


def test_probe_fallback_uses_only_friendly_copy_and_keeps_redacted_diagnostics(engine, job, monkeypatch):
    monkeypatch.setattr(engine, "oc_serve", lambda: "http://127.0.0.1:1")
    monkeypatch.setattr(engine, "oc_session", lambda *_args: "session-1")
    monkeypatch.setattr(engine, "oc_probe_once", lambda: (False, "执行外壳探活 90 秒没有完整回复"))

    result = engine.oc_run(str(job), "work")

    assert result == engine.OC_RUN_FALLBACK
    user_texts = [e.get("text") for e in events(job) if e.get("type") == "message"]
    assert user_texts[-1] == (
        "主连接响应较慢，已切换稳定通道继续；仍使用同一模型和同一套要求，不会降低内容标准。"
    )
    assert "探活" not in user_texts[-1]
    diagnostics = [json.loads(line) for line in (job / "diagnostics.jsonl").read_text(encoding="utf-8").splitlines()]
    assert diagnostics[-1]["code"] == "opencode_probe_failed"
    assert "90 秒" in diagnostics[-1]["detail"]


def test_opencode_server_run_drains_agent_events_while_running(engine, job, monkeypatch):
    (job / engine.AGENT_EVENTS_FILE).write_text(
        json.dumps({"type": "progress", "stage": "体检素材", "pct": 8, "step": 1, "total": 12}) + "\n",
        encoding="utf-8",
    )
    monkeypatch.setattr(engine, "oc_serve", lambda: "http://127.0.0.1:1")
    monkeypatch.setattr(engine, "oc_session", lambda *_args: "session-1")
    monkeypatch.setattr(engine, "oc_probe_once", lambda: (True, ""))
    monkeypatch.setattr(engine, "oc_send", lambda *_args, **_kwargs: (True, {}))
    monkeypatch.setattr(engine, "oc_watch", lambda *_args, **_kwargs: None)
    monkeypatch.setattr(engine, "oc_turn", lambda _sid: (True, ""))
    monkeypatch.setattr(engine, "OC_QUIET", 0)
    monkeypatch.setattr(engine.time, "sleep", lambda _seconds: None)

    assert engine.oc_run(str(job), "work") == engine.OC_RUN_COMPLETED
    assert engine.read_json(str(job / "progress.json"), {})["stage"] == "体检素材"


def test_delivery_summary_reports_word_toc_deviations_and_key_checks(engine, job):
    word = job / "投标文件_技术标.docx"
    _write_word(word, with_toc=True)
    (job / "技术应答偏离表.md").write_text(
        "|序号|要求|响应|\n|---|---|---|\n|1|技术要求|无偏离|\n", encoding="utf-8"
    )
    (job / "商务偏离表.md").write_text(
        "|序号|要求|响应|\n|---|---|---|\n|1|商务要求|无偏离|\n", encoding="utf-8"
    )
    (job / "成品质检报告.md").write_text("# 成品质检报告\n\n✅ 全部关键检查通过\n", encoding="utf-8")
    engine.write_json(str(job / "outcome.json"), {"state": "done", "word": word.name, "ts": engine.now()})

    with TestClient(engine.app) as client:
        delivery = _job_by_id(client)["delivery"]

    assert delivery["word"]["present"] is True
    assert delivery["word"]["name"] == word.name
    assert delivery["toc"]["status"] == "pass"
    assert delivery["deviations"]["status"] == "pass"
    assert delivery["deviations"]["technical"]["present"] is True
    assert delivery["deviations"]["business"]["present"] is True
    assert delivery["checks"]["status"] == "pass"
    assert delivery["ready"] is True


def test_delivery_summary_never_guesses_missing_toc_or_deviation_complete(engine, job):
    word = job / "投标文件_技术标.docx"
    _write_word(word, with_toc=False)
    (job / "技术应答偏离表.md").write_text("|序号|响应|\n|---|---|\n", encoding="utf-8")
    engine.write_json(str(job / "outcome.json"), {"state": "done", "word": word.name, "ts": engine.now()})

    with TestClient(engine.app) as client:
        item = _job_by_id(client)

    assert item["delivery"]["toc"]["status"] == "fail"
    assert item["delivery"]["deviations"]["status"] == "fail"
    assert item["delivery"]["checks"]["status"] in {"fail", "unknown"}
    assert item["delivery"]["ready"] is False
    assert item["presentation"]["code"] != "completed"


def test_delivery_summary_cache_reuses_unchanged_docx_and_invalidates_on_change(engine, job, monkeypatch):
    word = job / "投标文件_完整.docx"
    _write_word(word, with_toc=True)
    (job / "技术应答偏离表.md").write_text("|项|响应|\n|---|---|\n|1|满足|\n", encoding="utf-8")
    (job / "商务偏离表.md").write_text("|项|响应|\n|---|---|\n|1|满足|\n", encoding="utf-8")
    (job / "成品质检报告.md").write_text("✅ 通过", encoding="utf-8")
    original = engine._docx_has_toc
    calls = []
    monkeypatch.setattr(engine, "_docx_has_toc", lambda path: (calls.append(path) or original(path)))

    first = engine.delivery_summary(str(job))
    second = engine.delivery_summary(str(job))
    assert first == second
    assert len(calls) == 1

    _write_word(word, with_toc=False)
    changed = engine.delivery_summary(str(job))
    assert len(calls) == 2
    assert changed["toc"]["status"] == "fail"


def test_delivery_change_cannot_reuse_a_stale_green_quality_verdict(engine, job):
    word = job / "投标文件_完整.docx"
    _write_word(word, with_toc=True)
    (job / "技术应答偏离表.md").write_text("|项|响应|\n|---|---|\n|1|满足|\n", encoding="utf-8")
    (job / "商务偏离表.md").write_text("|项|响应|\n|---|---|\n|1|满足|\n", encoding="utf-8")
    report = job / "成品质检报告.md"
    report.write_text("✅ 通过", encoding="utf-8")

    first = engine.delivery_summary(
        str(job), {"status": "pass", "level": "green", "summary": "关键检查已通过"}
    )
    assert first["ready"] is True

    report.write_text("🔴 未通过：目录与正文不一致", encoding="utf-8")
    changed = engine.delivery_summary(str(job))
    assert changed["checks"]["status"] == "fail"
    assert changed["checks"]["level"] == "red"
    assert changed["ready"] is False


def test_quality_gate_runs_before_done_and_a_red_result_cannot_commit_success(engine, job, monkeypatch):
    word = job / "投标文件_技术标.docx"
    _write_word(word, with_toc=True)
    calls = []

    def rejected(*_args):
        calls.append("quality")
        return {"status": "fail", "level": "red", "summary": "关键检查未通过"}

    monkeypatch.setattr(engine, "quality_audit", rejected)
    result = engine.settle(str(job))

    assert calls == ["quality"]
    assert result["state"] == "stopped"
    assert engine.read_json(str(job / "outcome.json"), {})["state"] == "stopped"
    assert not any(e.get("type") == "progress" and e.get("pct") == 100 for e in events(job))


def _prepare_running_oc(engine, monkeypatch):
    monkeypatch.setattr(engine, "oc_serve", lambda: "http://127.0.0.1:1")
    monkeypatch.setattr(engine, "oc_session", lambda *_args: "session-1")
    monkeypatch.setattr(engine, "oc_probe_once", lambda: (True, ""))
    monkeypatch.setattr(engine, "oc_send", lambda *_args, **_kwargs: (True, {}))
    monkeypatch.setattr(engine, "oc_watch", lambda *_args, **_kwargs: None)
    monkeypatch.setattr(engine.time, "sleep", lambda _seconds: None)


def test_post_dispatch_error_is_calm_and_offers_real_resume_not_rerun(engine, job, monkeypatch):
    _prepare_running_oc(engine, monkeypatch)
    monkeypatch.setattr(engine, "oc_turn", lambda _sid: (True, "synthetic low-level stream error"))

    assert engine.oc_run(str(job), "work", allow_cli_fallback=False) == engine.OC_RUN_INTERRUPTED

    errors = [event for event in events(job) if event.get("type") == "error"]
    assert errors
    assert "执行外壳" not in errors[-1]["text"]
    action = next(action for action in errors[-1]["actions"] if action["label"] == "从已保存内容继续")
    assert action["act"] == "resume"
    diagnostic = json.loads((job / "diagnostics.jsonl").read_text(encoding="utf-8").splitlines()[-1])
    assert diagnostic["code"] == "opencode_turn_interrupted"
    assert "synthetic low-level" in diagnostic["detail"]


def test_outer_post_dispatch_exception_keeps_technical_detail_only_in_diagnostics(engine, job, monkeypatch):
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["oc_session"] = "session-1"
    engine.write_json(str(job / "任务.json"), meta)
    monkeypatch.setattr(engine, "oc_run", lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError("low-level opencode crash")))
    monkeypatch.setattr(engine, "ensure_default_shell", lambda *_args, **_kwargs: (True, "ready"))
    monkeypatch.setattr(engine, "settle", lambda *_args, **_kwargs: {"state": "stopped"})

    engine.agent_via_server_or_cli(str(job), "prompt", ["unused"])

    error = [item for item in events(job) if item.get("type") == "error"][-1]
    assert error["text"] == "连接意外中断，任务已安全停下；已生成的内容都已保留。"
    assert {action["act"] for action in error["actions"]} >= {"resume", "open_log"}
    assert not any(word in error["text"] for word in ("执行外壳", "OpenCode", "CLI", "探活", "兼容"))
    diagnostic = json.loads((job / "diagnostics.jsonl").read_text(encoding="utf-8").splitlines()[-1])
    assert diagnostic["code"] == "generation_worker_exception"
    assert "low-level opencode crash" in diagnostic["detail"]


def test_stall_is_calm_and_technical_timeout_is_only_in_diagnostics(engine, job, monkeypatch):
    _prepare_running_oc(engine, monkeypatch)
    monkeypatch.setattr(engine, "OC_STALL", 0)
    monkeypatch.setattr(engine, "oc_turn", lambda _sid: (False, ""))

    assert engine.oc_run(str(job), "work", allow_cli_fallback=False) == engine.OC_RUN_INTERRUPTED

    texts = [str(event.get("text") or "") for event in events(job) if event.get("type") == "message"]
    assert texts[-1] == "连接暂时没有响应，任务已安全停下；已生成的内容都已保留。"
    assert "执行外壳" not in texts[-1]
    diagnostic = json.loads((job / "diagnostics.jsonl").read_text(encoding="utf-8").splitlines()[-1])
    assert diagnostic["code"] == "opencode_stalled"


def test_server_stall_before_body_automatically_falls_back_to_cli_once(engine, job, monkeypatch):
    """Catch the customer failure where parsing exists but the managed model session goes silent."""
    _prepare_running_oc(engine, monkeypatch)
    (job / "招标文件_解析版.md").write_text("# 解析结果\n", encoding="utf-8")
    monkeypatch.setattr(engine, "OC_STALL", 0)
    monkeypatch.setattr(engine, "OC_SLOW", 999)
    monkeypatch.setattr(engine, "oc_turn", lambda _sid: (False, ""))

    assert engine.oc_run(str(job), "work") == engine.OC_RUN_FALLBACK

    runtime = engine.read_json(str(job / "runtime.json"), {})
    assert runtime["execution_path"] == "cli_compat"
    assert runtime["fallback_count"] == 1
    texts = [str(event.get("text") or "") for event in events(job)]
    assert sum("已切换稳定通道继续" in text for text in texts) == 1


def test_server_stall_after_body_never_replays_the_job(engine, job, monkeypatch):
    _prepare_running_oc(engine, monkeypatch)
    (job / "投标文件_技术标.md").write_text("# 已写正文\n\n不得从头覆盖。\n", encoding="utf-8")
    monkeypatch.setattr(engine, "OC_STALL", 0)
    monkeypatch.setattr(engine, "OC_SLOW", 999)
    monkeypatch.setattr(engine, "oc_turn", lambda _sid: (False, ""))

    assert engine.oc_run(str(job), "work") == engine.OC_RUN_INTERRUPTED
    runtime = engine.read_json(str(job / "runtime.json"), {})
    assert runtime.get("execution_path") != "cli_compat"
    assert "不得从头覆盖" in (job / "投标文件_技术标.md").read_text(encoding="utf-8")


def test_preflight_warns_when_macos_runs_from_translocated_dmg(engine, job, monkeypatch):
    monkeypatch.setattr(engine.sys, "platform", "darwin")
    monkeypatch.setattr(engine.sys, "executable", "/private/var/folders/x/AppTranslocation/ABC/d/中标狗.app/Contents/MacOS/bid-engine")
    monkeypatch.setattr(engine, "resolve_cli", lambda *_a, **_k: "/bundle/opencode-cli")
    engine.write_json(engine.conf_path(), {
        "engine": {"kind": "s2", "s2_key": "runtime-test-key", "s2_model": "senseaudio-s2"}
    })

    result = engine.generation_preflight(str(job))
    location = next(item for item in result["checks"] if item["id"] == "app_location")
    assert location["status"] == "warning"
    assert "应用程序" in location["message"]


def test_agent_self_test_uses_the_same_managed_session_path_as_real_jobs(engine, monkeypatch):
    """A passing short CLI probe must not hide a broken production server-session path."""
    engine.write_json(engine.conf_path(), {
        "engine": {
            "kind": "s2", "s2_base_url": "https://gateway.invalid/v1",
            "s2_key": "runtime-test-key", "s2_model": "senseaudio-s2",
            "s2_verify_ssl": True,
        }
    })
    monkeypatch.setattr(engine, "_openai_req", lambda *_a, **_k: {"data": [{"id": "senseaudio-s2"}]})
    monkeypatch.setattr(engine, "resolve_cli", lambda *_a, **_k: "/synthetic/opencode")
    monkeypatch.setattr(engine, "oc_serve", lambda *_a, **_k: "http://127.0.0.1:12345")
    monkeypatch.setattr(engine, "oc_probe", lambda: (False, "managed session did not finish"))

    class PassingCliProbe:
        returncode = 0
        stdout = "中标狗连接成功"
        stderr = ""

    monkeypatch.setattr(engine, "_tracked_detached_run", lambda *_a, **_k: PassingCliProbe())
    with TestClient(engine.app) as client:
        response = client.post("/v1/agent/test")

    assert response.status_code == 200
    result = response.json()
    assert result["ok"] is False
    assert result["execution_path"] == "opencode_server"
    assert "正式生成链路" in result["error"]


def test_flow_exposes_actual_expected_and_remaining_time_for_each_phase(engine, job, monkeypatch):
    monkeypatch.setattr(engine, "now", lambda: "2026-08-21 10:05:00")
    engine.write_json(str(job / "任务.json"), {
        "name": "耗时测试", "tender": "招标文件.docx", "created_at": "2026-08-21 10:00:00",
    })
    (job / "events.jsonl").write_text(
        "\n".join([
            json.dumps({"ts": "2026-08-21 10:00:00", "type": "progress", "stage": "环境检查", "step": 0, "pct": 0}),
            json.dumps({"ts": "2026-08-21 10:01:00", "type": "progress", "stage": "读取招标文件", "step": 1, "pct": 2}),
        ]) + "\n", encoding="utf-8",
    )
    engine.write_json(str(job / "progress.json"), {
        "type": "progress", "stage": "读取招标文件", "step": 1, "pct": 2,
        "total": 12, "ts": "2026-08-21 10:01:00",
    })
    owner = engine._reserve_running(job.name)
    try:
        with TestClient(engine.app) as client:
            item = _job_by_id(client)
    finally:
        engine._release_running(job.name, owner)

    phases = {phase["id"]: phase for phase in item["flow"]["phases"]}
    assert phases["environment"]["elapsed_seconds"] == 60
    assert phases["environment"]["expected_seconds"] == 60
    assert phases["parse"]["elapsed_seconds"] == 240
    assert phases["parse"]["expected_seconds"] == 390
    assert phases["parse"]["remaining_seconds"] == 150
    assert phases["parse"]["estimate_source"] == "reference"
    assert phases["write"]["expected_seconds"] == 1200
    assert item["flow"]["remaining_seconds"] == 2160
    assert item["eta_seconds"] == 2160
    assert item["flow"]["stalled"] is True
    assert item["flow"]["current_action"] == "模型响应偏慢，正在持续检查连接"


def test_flow_silence_warning_is_not_hidden_by_a_long_historical_phase(engine, job, monkeypatch):
    monkeypatch.setattr(engine, "now", lambda: "2026-08-21 10:05:00")
    monkeypatch.setattr(engine, "stage_stats", lambda: {
        "stages": [{"step": 1, "avg_s": 1200, "from_history": True}],
    })
    engine.write_json(str(job / "任务.json"), {
        "name": "慢阶段心跳测试", "tender": "招标文件.docx",
        "created_at": "2026-08-21 10:00:00",
    })
    (job / "招标文件_解析版.md").write_text("# 解析结果\n", encoding="utf-8")
    (job / "events.jsonl").write_text(
        json.dumps({"ts": "2026-08-21 10:01:00", "type": "progress",
                    "stage": "读取招标文件", "step": 1, "pct": 2}) + "\n",
        encoding="utf-8",
    )
    progress = {"type": "progress", "stage": "读取招标文件", "step": 1,
                "pct": 2, "total": 12, "ts": "2026-08-21 10:01:00"}

    flow = engine.job_flow(str(job), state="running", prog=progress)

    assert flow["expected_seconds"] == 1200
    assert flow["silence_seconds"] == 240
    assert flow["stalled"] is True
    assert flow["current_action"] == "模型响应偏慢，正在持续检查连接"


def test_flow_explains_analysis_after_the_parsed_source_is_already_on_disk(engine, job, monkeypatch):
    monkeypatch.setattr(engine, "now", lambda: "2026-08-21 10:02:00")
    engine.write_json(str(job / "任务.json"), {
        "name": "解析进度测试", "tender": "招标文件.docx",
        "created_at": "2026-08-21 10:00:00",
    })
    (job / "招标文件_解析版.md").write_text("# 已提取的招标正文\n", encoding="utf-8")
    progress = {"type": "progress", "stage": "已就绪，正在读招标文件", "step": 1,
                "pct": 2, "total": 12, "ts": "2026-08-21 10:01:00"}

    flow = engine.job_flow(str(job), state="running", prog=progress)

    assert flow["current_phase"] == "parse"
    assert flow["current_action"] == "招标正文已提取，正在识别目录、条款、评分项和废标条件"
    parse = next(phase for phase in flow["phases"] if phase["id"] == "parse")
    assert parse["detail"] == flow["current_action"]


def test_flow_advances_from_verified_analysis_files_even_without_agent_progress_events(engine, job):
    (job / "投标文件组成.md").write_text("组成分析\n" * 80, encoding="utf-8")
    engine.write_json(str(job / "word_format_spec.json"), {})
    (job / "格式要求摘要.md").write_text("格式要求\n" * 8, encoding="utf-8")
    (job / "评分点响应矩阵.md").write_text("评分项\n" * 8, encoding="utf-8")
    (job / "废标风险清单.md").write_text("废标风险\n" * 8, encoding="utf-8")
    progress = {"type": "progress", "stage": "已就绪，正在读招标文件", "step": 1,
                "pct": 2, "total": 12, "ts": engine.now()}

    flow = engine.job_flow(str(job), state="running", prog=progress)

    assert flow["checkpoint"] == {"step": 5, "label": "评分废标"}
    assert flow["current_phase"] == "plan"
    assert flow["phases"][1]["state"] == "done"


def test_runtime_pause_reason_uses_no_implementation_terms(engine, job):
    engine.write_json(
        str(job / "runtime.json"),
        {"execution_path": "cli_compat", "can_pause": False},
    )
    reason = engine.job_runtime(str(job))["capabilities"]["pause"]["reason"]
    assert "暂不支持暂停" in reason
    assert not any(word in reason for word in ("执行外壳", "OpenCode", "CLI", "探活", "兼容"))


def test_current_action_normalizes_technical_stage_text(engine, job):
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["staged"] = False
    engine.write_json(str(job / "任务.json"), meta)
    engine.write_json(
        str(job / "progress.json"),
        {"type": "progress", "stage": "执行外壳探活中", "pct": 2, "step": 1, "total": 12, "ts": engine.now()},
    )
    owner = engine._reserve_running(job.name)
    try:
        with TestClient(engine.app) as client:
            item = _job_by_id(client)
    finally:
        engine._release_running(job.name, owner)

    assert item["current_action"] == "正在建立稳定连接"
    assert not any(word in item["current_action"] for word in ("执行外壳", "OpenCode", "CLI", "探活", "兼容"))


def test_failed_probe_is_short_cached_per_configuration_fingerprint(engine, monkeypatch):
    calls = []
    config = {
        "engine": {"kind": "s2", "s2_key": "secret-a", "s2_model": "senseaudio-s2"}
    }
    engine.write_json(engine.conf_path(), config)
    monkeypatch.setattr(engine, "oc_probe", lambda: (calls.append(1) or False, "slow"))

    assert engine.oc_probe_once() == (False, "slow")
    assert engine.oc_probe_once() == (False, "slow")
    assert len(calls) == 1

    config["engine"]["s2_model"] = "deepseek-v4-flash"
    engine.write_json(engine.conf_path(), config)
    assert engine.oc_probe_once() == (False, "slow")
    assert len(calls) == 2


def test_archive_and_project_metadata_survive_legacy_task_json_writes(engine, job):
    stale = engine.read_json(str(job / "任务.json"), {})
    with TestClient(engine.app) as client:
        response = client.patch(
            "/v1/jobs/%s" % job.name,
            json={"archived": True, "project_id": "project-a", "name": "项目 A 技术标"},
        )
        assert response.status_code == 200

    # A running worker may still hold an old 任务.json snapshot. Its write must not
    # erase user organization metadata.
    stale["paused"] = True
    engine.write_json(str(job / "任务.json"), stale)

    with TestClient(engine.app) as client:
        item = _job_by_id(client)
        assert item["archived_at"]
        assert item["project_id"] == "project-a"
        assert item["name"] == "项目 A 技术标"
        restored = client.patch("/v1/jobs/%s" % job.name, json={"archived": False})
        assert restored.status_code == 200
        assert restored.json()["archived_at"] == ""


def test_concurrent_product_metadata_patches_keep_disjoint_fields(engine, job):
    path = str(job / "product.json")
    barrier = threading.Barrier(3)

    def update(values):
        barrier.wait()
        engine.patch_json(path, values)

    first = threading.Thread(target=update, args=({"archived_at": "2026-08-08 10:00:00"},))
    second = threading.Thread(target=update, args=({"project_id": "project-b"},))
    first.start()
    second.start()
    barrier.wait()
    first.join()
    second.join()

    assert engine.read_json(path, {}) == {
        "archived_at": "2026-08-08 10:00:00",
        "project_id": "project-b",
    }


def test_rerun_copies_all_per_task_materials(engine, job, monkeypatch):
    (job / "素材" / "nested").mkdir(parents=True)
    (job / "素材" / "nested" / "资质.txt").write_text("qualification", encoding="utf-8")
    (job / "参考资料").mkdir()
    (job / "参考资料" / "旧标书.docx").write_bytes(b"reference")
    (job / "你的要求.md").write_text("保持品牌口径", encoding="utf-8")
    source_meta = engine.read_json(str(job / "任务.json"), {})
    source_meta["template_id"] = "government"
    source_meta["template_snapshot"] = {
        "id": "government",
        "name": "政府采购",
        "prompt": "逐条响应政府采购要求",
        "settings": {"quality_gate": True},
    }
    engine.write_json(str(job / "任务.json"), source_meta)
    engine.write_json(str(job / "product.json"), {"project_id": "project-a"})
    monkeypatch.setattr(
        engine,
        "_launch_job",
        lambda jid, _path, _mock="auto": {"job_id": jid, "mode": "staged"},
    )

    with TestClient(engine.app) as client:
        response = client.post("/v1/jobs/%s/rerun" % job.name)

    assert response.status_code == 200
    new_job = Path(engine.jpath(response.json()["job_id"]))
    assert (new_job / "素材" / "nested" / "资质.txt").read_text(encoding="utf-8") == "qualification"
    assert (new_job / "参考资料" / "旧标书.docx").read_bytes() == b"reference"
    assert (new_job / "你的要求.md").read_text(encoding="utf-8") == "保持品牌口径"
    assert engine.read_json(str(new_job / "product.json"), {})["project_id"] == "project-a"
    rerun_meta = engine.read_json(str(new_job / "任务.json"), {})
    assert rerun_meta["template_snapshot"] == source_meta["template_snapshot"]


def test_bulk_archive_restore_and_deliverables_only_zip(engine, job):
    (job / "投标文件_完整.docx").write_bytes(b"word-result")
    (job / "响应矩阵.md").write_text("matrix", encoding="utf-8")
    (job / "run.log").write_text("private runtime", encoding="utf-8")
    (job / "招标文件_解析版.md").write_text("internal parsed source", encoding="utf-8")

    with TestClient(engine.app) as client:
        archived = client.post(
            "/v1/jobs/bulk", json={"action": "archive", "job_ids": [job.name]}
        )
        assert archived.status_code == 200
        assert archived.json()["succeeded"] == [job.name]
        assert _job_by_id(client)["archived_at"]

        exported = client.post("/v1/jobs/export", json={"job_ids": [job.name]})
        assert exported.status_code == 200
        restored = client.post(
            "/v1/jobs/bulk", json={"action": "restore", "job_ids": [job.name]}
        )
        assert restored.status_code == 200
        assert _job_by_id(client)["archived_at"] == ""

    archive = zipfile.ZipFile(io.BytesIO(exported.content))
    names = set(archive.namelist())
    assert any(name.endswith("投标文件_完整.docx") for name in names)
    assert any(name.endswith("响应矩阵.md") for name in names)
    assert not any(name.endswith("招标文件.docx") for name in names)
    assert not any(name.endswith("任务.json") or name.endswith("run.log") for name in names)
    assert not any("解析版" in name for name in names)


def test_job_list_scope_filters_archived_and_keeps_default_all(engine, job):
    other = Path(engine.jpath("job-2")); other.mkdir(parents=True)
    (other / "招标文件.docx").write_bytes(b"tender")
    engine.write_json(str(other / "任务.json"), {"name": "任务二", "tender": "招标文件.docx", "staged": True})
    engine.write_json(str(job / "product.json"), {"archived_at": engine.now(), "project_id": "p1"})
    engine.write_json(str(other / "product.json"), {"project_id": "p2"})

    with TestClient(engine.app) as client:
        assert {item["job_id"] for item in client.get("/v1/jobs").json()} == {"job-1", "job-2"}
        assert [item["job_id"] for item in client.get("/v1/jobs", params={"scope": "archived"}).json()] == ["job-1"]
        assert [item["job_id"] for item in client.get("/v1/jobs", params={"scope": "active"}).json()] == ["job-2"]
        assert [item["job_id"] for item in client.get("/v1/jobs", params={"project_id": "p2"}).json()] == ["job-2"]


def test_default_templates_and_crud(engine):
    with TestClient(engine.app) as client:
        defaults = client.get("/v1/templates")
        assert defaults.status_code == 200
        assert [(item["id"], item["name"]) for item in defaults.json()[:3]] == [
            ("government", "政府采购"),
            ("construction", "工程施工"),
            ("service", "服务类投标"),
        ]
        assert all(item["description"] and item["prompt"] and item["settings"] for item in defaults.json()[:3])

        created = client.post(
            "/v1/templates",
            json={"name": "医疗器械标", "description": "适合设备采购", "prompt": "突出设备参数",
                  "settings": {"quality_gate": True, "include_deviation_tables": True}},
        )
        assert created.status_code == 200
        template_id = created.json()["id"]
        changed = client.put(
            "/v1/templates/%s" % template_id,
            json={"name": "医疗设备标", "prompt": "逐项核对设备参数"},
        )
        assert changed.status_code == 200
        assert changed.json()["name"] == "医疗设备标"
        assert client.delete("/v1/templates/%s" % template_id).status_code == 200
        assert all(item["id"] != template_id for item in client.get("/v1/templates").json())


def test_builtin_templates_are_complete_scene_packages(engine):
    with TestClient(engine.app) as client:
        templates = {item["id"]: item for item in client.get("/v1/templates").json()}

    assert {"government", "government-it", "goods", "construction", "service", "consulting"} <= set(templates)
    for template_id in ("government", "government-it", "goods", "construction", "service", "consulting"):
        package = templates[template_id]["package"]
        assert package["schema_version"] == 1
        assert package["priority_rule"] == "招标文件原文高于模板；冲突时以招标文件为准"
        assert len(package["outline"]) >= 5
        assert package["scoring_focus"]
        assert package["tables"]
        assert package["material_slots"]
        assert package["quality_rules"]
        assert package["sources"]
        assert all(source["title"] and source["issuer"] and source["url"] for source in package["sources"])


def test_auto_template_recommendation_compiles_scene_rules_into_job_prompt(engine):
    tender = (
        "# 工程施工总承包招标文件\n"
        "评分内容包括施工组织设计、施工进度计划、质量保证、安全文明施工、项目管理机构和应急预案。\n"
    ).encode("utf-8")

    with TestClient(engine.app) as client:
        response = client.post(
            "/v1/jobs",
            files={"tender": ("工程施工采购.md", tender, "text/markdown")},
            data={"template_id": "auto", "start": "0", "name": "工程施工测试", "prompt": "重点逐项对应评分分值。"},
        )

    assert response.status_code == 200
    job_id = response.json()["job_id"]
    stored = engine.read_json(str(Path(engine.jpath(job_id)) / "任务.json"), {})
    assert stored["template_id"] == "construction"
    assert stored["template_snapshot"]["package"]["scene"]["category"] == "engineering_construction"
    assert stored["template_recommendation"]["confidence"] > 0
    prompt = stored["prompt"]
    assert "招标文件原文高于模板" in prompt
    assert "施工组织设计" in prompt
    assert "进度计划" in prompt
    assert "质量安全" in prompt
    assert "重点逐项对应评分分值" in prompt


@pytest.mark.parametrize("requested", ["auto", "default", ""])
def test_default_template_aliases_all_freeze_a_real_scene_snapshot(engine, requested):
    """Regression: legacy/default clients must not create a template-less task or 400."""
    tender = (
        "# 工程施工总承包招标文件\n"
        "施工组织设计、施工进度计划、质量保证、安全文明施工和项目管理机构。\n"
    ).encode("utf-8")

    with TestClient(engine.app) as client:
        response = client.post(
            "/v1/jobs",
            files={"tender": ("工程施工采购.md", tender, "text/markdown")},
            data={"template_id": requested, "start": "0", "name": "模板兼容测试"},
        )

    assert response.status_code == 200
    stored = engine.read_json(str(Path(engine.jpath(response.json()["job_id"])) / "任务.json"), {})
    assert stored["template_id"] == "construction"
    assert stored["template_recommendation"]["template_id"] == "construction"
    assert stored["template_snapshot"]["id"] == "construction"
    assert stored["template_snapshot"]["package"]["outline"]


def test_auto_template_recommendation_uses_main_document_selected_from_files(engine):
    tender = (
        "# 工程施工总承包招标文件\n"
        "施工组织设计、施工进度计划、质量安全、文明施工和工程量清单。\n"
    ).encode("utf-8")

    with TestClient(engine.app) as client:
        response = client.post(
            "/v1/jobs",
            files=[("files", ("工程施工采购.md", tender, "text/markdown"))],
            data={"template_id": "auto", "start": "0", "name": "仅文件列表上传", "relpaths": "[\"工程施工采购.md\"]"},
        )

    assert response.status_code == 200
    stored = engine.read_json(str(Path(engine.jpath(response.json()["job_id"])) / "任务.json"), {})
    assert stored["template_id"] == "construction"


def test_template_preview_and_job_use_the_same_scene_hint(engine):
    tender = "# 采购文件\n按采购要求响应。".encode("utf-8")
    scene_hint = "本项目是软件信息化平台建设，包含系统集成、数据安全、信创和运维。"

    with TestClient(engine.app) as client:
        preview = client.post(
            "/v1/templates/recommend",
            files={"file": ("采购文件.md", tender, "text/markdown")},
            data={"scene_hint": scene_hint},
        )
        job = client.post(
            "/v1/jobs",
            files={"tender": ("采购文件.md", tender, "text/markdown")},
            data={"template_id": "auto", "start": "0", "name": "一致性测试", "prompt": scene_hint},
        )

    assert preview.status_code == 200
    assert job.status_code == 200
    stored = engine.read_json(str(Path(engine.jpath(job.json()["job_id"])) / "任务.json"), {})
    assert preview.json()["template_id"] == "government-it"
    assert stored["template_id"] == preview.json()["template_id"]


def test_uploaded_good_bid_derives_reviewable_template_without_copying_body_facts(engine):
    document = Document()
    document.add_heading("投标文件", level=1)
    document.add_heading("资格与符合性响应", level=1)
    document.add_paragraph("客户名称：深圳市某局，投标报价人民币123456元。")
    document.add_heading("技术方案", level=1)
    document.add_heading("实施进度计划", level=2)
    document.add_paragraph("本项目专属实施内容，不应复制进入模板。")
    document.add_heading("质量保证与售后服务", level=1)
    table = document.add_table(rows=2, cols=3)
    for index, value in enumerate(("序号", "招标要求", "投标响应")):
        table.rows[0].cells[index].text = value
    project_data = document.add_table(rows=1, cols=3)
    for index, value in enumerate(("某单位总部", "平台一期项目", "13800138000")):
        project_data.rows[0].cells[index].text = value
    buffer = io.BytesIO()
    document.save(buffer)

    with TestClient(engine.app) as client:
        response = client.post(
            "/v1/templates/derive",
            files={"file": ("优秀历史标书.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"name": "信息化项目历史模板"},
        )

    assert response.status_code == 200
    draft = response.json()
    assert draft["requires_review"] is True
    assert draft["validation"]["score"] >= 70
    assert draft["package"]["generated_from"]["kind"] == "uploaded_bid"
    outline_titles = [item["title"] for item in draft["package"]["outline"]]
    assert "资格与符合性响应" in outline_titles
    assert "功能与技术响应" in outline_titles
    assert "实施进度与工期保障" in outline_titles
    payload = json.dumps(draft, ensure_ascii=False)
    assert "深圳市某局" not in payload
    assert "123456" not in payload
    assert "某单位总部" not in payload
    assert "13800138000" not in payload
    assert "本项目专属实施内容" not in payload


def test_thin_uploaded_template_cannot_bypass_review_by_calling_save_api(engine):
    with TestClient(engine.app) as client:
        derived = client.post(
            "/v1/templates/derive",
            files={"file": ("片段.md", "# 技术方案\n只有一个章节。".encode(), "text/markdown")},
        )
        draft = derived.json()
        saved = client.post(
            "/v1/templates",
            json={key: draft[key] for key in ("name", "description", "prompt", "settings", "package")},
        )

    assert derived.status_code == 200
    assert draft["validation"]["ready"] is False
    assert saved.status_code == 400
    assert "结构不足" in saved.json()["error"]


def test_job_creation_saves_an_immutable_template_snapshot(engine):
    with TestClient(engine.app) as client:
        response = client.post(
            "/v1/jobs",
            files={"tender": ("采购文件.docx", b"tender", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"template_id": "government", "project_id": "project-government",
                  "start": "0", "name": "采购项目"},
        )
        assert response.status_code == 200
        job_id = response.json()["job_id"]
        snapshot_before = engine.read_json(str(Path(engine.jpath(job_id)) / "任务.json"), {})["template_snapshot"]
        client.put("/v1/templates/government", json={"prompt": "未来的新模板内容"})

    stored = engine.read_json(str(Path(engine.jpath(job_id)) / "任务.json"), {})
    assert snapshot_before["id"] == "government"
    assert snapshot_before["name"] == "政府采购"
    assert snapshot_before["settings"]["quality_gate"] is True
    assert stored["template_snapshot"] == snapshot_before
    with TestClient(engine.app) as client:
        listed = _job_by_id(client, job_id)
    assert listed["project_id"] == "project-government"
    assert listed["template_id"] == "government"


def test_one_click_diagnostics_is_safe_and_includes_job_bundle_link(engine, job):
    key = "s" + "k-" + "test-secret-value-123456789"
    engine.write_json(
        engine.conf_path(),
        {"engine": {"kind": "s2", "s2_key": key, "s2_model": "model-a"}},
    )
    (job / "diagnostics.jsonl").write_text(
        json.dumps({"code": "network", "detail": "Authorization: Bearer " + key}) + "\n",
        encoding="utf-8",
    )

    with TestClient(engine.app) as client:
        response = client.get("/v1/diagnostics", params={"jid": job.name})

    assert response.status_code == 200
    text = json.dumps(response.json(), ensure_ascii=False)
    assert key not in text
    assert response.json()["job"]["bundle_url"].endswith("/%s/bundle" % job.name)
    assert response.json()["checks"]

    with TestClient(engine.app) as client:
        posted = client.post("/v1/diagnostics", json={"job_id": job.name})
    assert posted.status_code == 200
    assert key not in posted.text
    assert posted.json()["job"]["job_id"] == job.name


def test_setup_connect_invalid_key_is_not_saved_or_echoed(engine, monkeypatch):
    key = "s" + "k-" + "invalid-secret-value-123456789"
    monkeypatch.setattr(engine, "setup_connection_probe", lambda _conf: (False, "401 bad key " + key))

    with TestClient(engine.app) as client:
        initial = client.get("/v1/setup").json()
        response = client.post("/v1/setup/connect", json={"key": key})

    assert initial["needed"] is True
    assert response.status_code == 400
    assert key not in response.text
    assert engine.s2_conf(engine.read_json(engine.conf_path(), {}))["api_key"] == ""


def test_first_run_recommends_fast_mode_and_flash_model(engine):
    with TestClient(engine.app) as client:
        status = client.get("/v1/setup").json()

    assert status["recommended"] == {
        "engine": "opencode",
        "mode": "fast",
        "model": "deepseek-v4-flash",
    }


def test_setup_connect_complete_and_key_change_requires_retest(engine, monkeypatch):
    key_a = "s" + "k-" + "valid-secret-value-aaaaaaaa"
    key_b = "s" + "k-" + "valid-secret-value-bbbbbbbb"
    monkeypatch.setattr(
        engine,
        "setup_connection_probe",
        lambda _conf: (True, "", [engine.S2_DEFAULT_MODEL, "qwen-vl-max"]),
    )

    with TestClient(engine.app) as client:
        connected = client.post("/v1/setup/connect", json={"key": key_a})
        assert connected.status_code == 200
        assert key_a not in connected.text
        assert connected.json()["mode"] == "fast"
        assert connected.json()["model"] == "deepseek-v4-flash"
        completed = client.post("/v1/setup/complete")
        assert completed.status_code == 200
        assert client.get("/v1/setup").json()["needed"] is False

        # Simulate a later settings change outside the setup wizard.
        conf = engine.read_json(engine.conf_path(), {})
        conf["engine"]["s2_key"] = key_b
        engine.write_json(engine.conf_path(), conf)
        changed = client.get("/v1/setup").json()
        assert changed["needed"] is True
        assert changed["connected"] is False
        assert client.post("/v1/setup/complete").status_code == 409

        reconnected = client.post("/v1/setup/connect", json={"key": key_b})
        assert reconnected.status_code == 200
        assert client.post("/v1/setup/complete").status_code == 200

    stored = engine.read_json(engine.conf_path(), {})
    assert stored["engine"]["kind"] == "s2"
    assert stored["engine"]["mode"] == "agents"
    assert stored["engine"]["s2_model"] == engine.S2_DEFAULT_MODEL
    provider = next(item for item in stored["providers"] if item["id"] == "setup-s2")
    assert provider["model"] == engine.S2_DEFAULT_MODEL
    assert provider["vision_model"] == "qwen-vl-max"
    assert stored["routing"]["default"] == "setup-s2"


def test_setup_failure_leaves_entire_previous_configuration_unchanged(engine, monkeypatch):
    old_key = "s" + "k-" + "old-value-aaaaaaaaaaaaaaaa"
    new_key = "s" + "k-" + "new-value-bbbbbbbbbbbbbbbb"
    original = {"engine": {"kind": "codex"}, "providers": [{"id": "old", "api_key": old_key}],
                "routing": {"default": "old"}}
    engine.write_json(engine.conf_path(), original)
    monkeypatch.setattr(engine, "setup_connection_probe", lambda _conf: (False, "invalid", []))
    with TestClient(engine.app) as client:
        response = client.post("/v1/setup/connect", json={"key": new_key})
    assert response.status_code == 400
    assert engine.read_json(engine.conf_path(), {}) == original


def test_setup_skips_legacy_user_with_jobs(engine, job):
    with TestClient(engine.app) as client:
        status = client.get("/v1/setup").json()

    assert status["needed"] is False
    assert status["legacy_skipped"] is True


def test_revision_creates_versioned_child_without_mutating_parent(engine, job, monkeypatch):
    (job / "素材").mkdir()
    (job / "素材" / "公司介绍.md").write_text("facts", encoding="utf-8")
    (job / "投标文件_完整.docx").write_bytes(b"version-one")
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["template_id"] = "service"
    meta["template_snapshot"] = {
        "id": "service", "name": "服务类投标", "prompt": "强调服务交付",
        "settings": {"quality_gate": True},
    }
    engine.write_json(str(job / "任务.json"), meta)
    parent_before = {p.name: p.read_bytes() for p in job.iterdir() if p.is_file()}
    monkeypatch.setattr(
        engine,
        "_launch_job",
        lambda jid, _path, _mock="auto": {"job_id": jid, "mode": "staged"},
    )

    with TestClient(engine.app) as client:
        response = client.post(
            "/v1/jobs/%s/revisions" % job.name,
            json={"instruction": "加强实施方案", "start": False},
        )
        revisions = client.get("/v1/jobs/%s/revisions" % job.name)

    assert response.status_code == 200
    child_id = response.json()["job_id"]
    child = Path(engine.jpath(child_id))
    product = engine.read_json(str(child / "product.json"), {})
    assert product["parent_job_id"] == job.name
    assert product["root_job_id"] == job.name
    assert product["version"] == 2
    assert engine.read_json(str(child / "任务.json"), {})["template_snapshot"] == meta["template_snapshot"]
    assert (child / "投标文件_完整.docx").read_bytes() == b"version-one"
    assert "加强实施方案" in (child / "你的要求.md").read_text(encoding="utf-8")
    assert revisions.status_code == 200
    assert [item["job_id"] for item in revisions.json()] == [child_id]
    assert parent_before == {p.name: p.read_bytes() for p in job.iterdir() if p.is_file()}


def test_concurrent_revisions_get_unique_contiguous_versions(engine, job, monkeypatch):
    original_family = engine._revision_family

    def slow_family(root_id):
        rows = original_family(root_id)
        engine.time.sleep(0.04)
        return rows

    monkeypatch.setattr(engine, "_revision_family", slow_family)
    barrier = threading.Barrier(3)
    responses = []

    def create(instruction):
        barrier.wait()
        with TestClient(engine.app) as client:
            responses.append(client.post(
                "/v1/jobs/%s/revisions" % job.name,
                json={"instruction": instruction, "start": False},
            ))

    threads = [threading.Thread(target=create, args=("修改 %d" % index,)) for index in (1, 2)]
    for thread in threads: thread.start()
    barrier.wait()
    for thread in threads: thread.join()
    assert all(response.status_code == 200 for response in responses)
    assert sorted(response.json()["version"] for response in responses) == [2, 3]


def test_opencode_usage_is_collected_from_session_messages(engine, job, monkeypatch):
    def api(path, *_args, **_kwargs):
        if path.endswith("/message"):
            return 200, [
                {"type": "assistant", "tokens": {"input": 120, "output": 30}, "cost": 0.12},
                {"type": "assistant", "tokens": {"input": 80, "output": 20}, "cost": 0.08},
            ]
        return 200, {}

    monkeypatch.setattr(engine, "oc_api", api)
    engine.collect_oc_usage(str(job), "session-usage")

    usage = engine.read_json(str(job / "usage.json"), {})
    assert usage["calls"] == 2
    assert usage["input_tokens"] == 200
    assert usage["output_tokens"] == 50
    assert usage["total_tokens"] == 250
    assert usage["estimated_cost"] == 0.2
    assert usage["currency"] == "USD"


def test_usage_without_cost_does_not_invent_price_or_currency(engine, job, monkeypatch):
    monkeypatch.setattr(
        engine,
        "oc_api",
        lambda *_args, **_kwargs: (200, [{"type": "assistant", "tokens": {"input": 10, "output": 4}}]),
    )
    engine.collect_oc_usage(str(job), "session-no-cost")
    usage = engine._job_usage(str(job))
    assert usage["estimated_cost"] is None
    assert usage["currency"] is None
