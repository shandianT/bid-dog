import importlib.util
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SKILL_ARCHIVE = ROOT / "bidmultiagenttao_v5.3.zip"


def _load_builder(tmp_path):
    with zipfile.ZipFile(SKILL_ARCHIVE) as archive:
        archive.extractall(tmp_path)
    refs = tmp_path / "bid-multiagent-tao" / "references"
    sys.path.insert(0, str(refs))
    try:
        spec = importlib.util.spec_from_file_location(
            "bundled_build_tender_docx", refs / "build_tender_docx.py"
        )
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        return module
    finally:
        sys.path.remove(str(refs))


def _extract_refs(tmp_path):
    with zipfile.ZipFile(SKILL_ARCHIVE) as archive:
        archive.extractall(tmp_path)
    return tmp_path / "bid-multiagent-tao" / "references"


def _twips_to_cm(value):
    return int(value) / 567.0


def test_narrative_table_uses_readable_landscape_layout(tmp_path):
    """Long-response tables must not be squeezed into one-character-wide portrait columns."""
    builder = _load_builder(tmp_path)
    doc = Document()
    spec = builder.load_spec(None)
    builder.style_doc(doc, spec)
    rows = [
        ["序号", "招标文件依据", "响应说明", "证明材料", "缺口说明"],
        [
            "1",
            "招标文件要求中标单位的人员必须到岗到位，不得转让或分包。",
            "我方承诺中标后项目人员全部到岗到位，不转让、不分包。",
            "投标承诺书、拟投入本项目人员一览表。",
            "无缺口，已完全响应。",
        ],
    ]

    table = builder.add_table(doc, rows, spec)

    assert len(doc.sections) == 3, "长文本多列表应自动切换为横版，并在表后恢复竖版"
    assert doc.sections[1].orientation == WD_ORIENT.LANDSCAPE
    grid = table._tbl.find(qn("w:tblGrid"))
    widths = [_twips_to_cm(item.get(qn("w:w"))) for item in grid.findall(qn("w:gridCol"))]
    assert widths[0] >= 1.25, "“序号”不应被压成一字一行"
    assert min(widths[1:]) >= 3.0, "叙述列需有可读的最小宽度"


def test_table_paragraphs_override_body_indentation_and_spacing(tmp_path):
    builder = _load_builder(tmp_path)
    doc = Document()
    spec = builder.load_spec(None)
    builder.style_doc(doc, spec)
    table = builder.add_table(
        doc,
        [["序号", "要求内容"], ["1", "本单元格文字不应继承正文首行缩进和一点五倍行距。"]],
        spec,
    )

    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            ppr = paragraph._p.get_or_add_pPr()
            indent = ppr.find(qn("w:ind"))
            assert indent is not None
            assert indent.get(qn("w:firstLine")) == "0"
            # Normal 正文用 firstLineChars=200（两字符缩进）；仅写 firstLine=0
            # 在部分 Word/WPS/LibreOffice 兼容路径中仍会继承字符缩进。
            assert indent.get(qn("w:firstLineChars")) == "0"
            assert indent.get(qn("w:left")) == "0"
            assert indent.get(qn("w:right")) == "0"
            spacing = ppr.find(qn("w:spacing"))
            assert spacing is not None
            assert spacing.get(qn("w:before")) == "0"
            assert spacing.get(qn("w:after")) == "0"
            assert spacing.get(qn("w:line")) == "276"
            assert spacing.get(qn("w:lineRule")) == "auto"


def test_all_table_cells_are_left_aligned(tmp_path):
    builder = _load_builder(tmp_path)
    doc = Document()
    spec = builder.load_spec(None)
    builder.style_doc(doc, spec)
    table = builder.add_table(
        doc,
        [["序号", "偏离状态", "缺口说明"], ["1", "无偏离", "完全响应"]],
        spec,
    )

    for row in table.rows:
        for cell in row.cells:
            assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_format_gate_rejects_table_paragraph_that_reinherits_body_indent(tmp_path):
    refs = _extract_refs(tmp_path)
    output = tmp_path / "table.docx"
    built = subprocess.run(
        [
            sys.executable,
            str(refs / "build_tender_docx.py"),
            str(ROOT / "tests" / "fixtures" / "table_layout_validation.md"),
            str(output),
            "--title",
            "表格排版优化验证",
        ],
        capture_output=True,
        text=True,
    )
    assert built.returncode == 0, built.stderr
    document = Document(output)
    document.tables[0].cell(1, 1).paragraphs[0].paragraph_format.first_line_indent = Pt(24)
    document.save(output)

    checked = subprocess.run(
        [
            sys.executable,
            str(refs / "check_docx_format.py"),
            str(output),
            "--title",
            "表格排版优化验证",
        ],
        capture_output=True,
        text=True,
    )
    assert checked.returncode == 1
    assert "表格·段落" in checked.stdout


def test_format_gate_rejects_center_aligned_table_cell(tmp_path):
    refs = _extract_refs(tmp_path)
    output = tmp_path / "table-center.docx"
    built = subprocess.run(
        [
            sys.executable,
            str(refs / "build_tender_docx.py"),
            str(ROOT / "tests" / "fixtures" / "table_layout_validation.md"),
            str(output),
            "--title",
            "表格排版优化验证",
        ],
        capture_output=True,
        text=True,
    )
    assert built.returncode == 0, built.stderr
    document = Document(output)
    document.tables[0].cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(output)

    checked = subprocess.run(
        [
            sys.executable,
            str(refs / "check_docx_format.py"),
            str(output),
            "--title",
            "表格排版优化验证",
        ],
        capture_output=True,
        text=True,
    )
    assert checked.returncode == 1
    assert "表格·对齐" in checked.stdout
