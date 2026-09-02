"""P3:Word 真预览——docx → HTML 只读渲染,以及它的端点。"""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from fastapi.testclient import TestClient

PNG_1x1 = bytes.fromhex(
    '89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489'
    '0000000d4944415478da63f8cfc0f01f0005000101f0d4d4b70000000049454e44ae426082')


def _build(path, pages=1):
    document = Document()
    document.add_heading('投标文件', level=1)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run('某某项目技术标').bold = True
    body = document.add_paragraph('逐项响应招标要求并提供可核验依据,')
    body.add_run('引用条款 3.2').italic = True
    table = document.add_table(rows=3, cols=3)
    for col, text in enumerate(('序号', '招标要求', '响应')): table.cell(0, col).text = text
    table.cell(1, 0).text = '1'
    table.cell(1, 1).merge(table.cell(1, 2)).text = '按期交付'          # 横向合并 → colspan
    table.cell(2, 0).text = '2'
    table.cell(2, 1).text = '质保 3 年'
    table.cell(2, 2).text = '完全响应'
    document.add_picture(io.BytesIO(PNG_1x1), width=Inches(0.5))
    for _ in range(pages):
        document.add_page_break()
        document.add_heading('售后服务方案', level=2)
        document.add_paragraph('7×24 小时响应 <脚本> & 转义。')
    document.save(path)


def test_docx_preview_keeps_structure_and_escapes_text(engine, tmp_path):
    path = tmp_path / '投标文件_整册.docx'
    _build(str(path))
    html, stats = engine._docx_preview_html(str(path))
    assert '<h1>投标文件</h1>' in html
    assert '<p class="a-center"><b>某某项目技术标</b></p>' in html
    assert '逐项响应招标要求并提供可核验依据,<i>引用条款 3.2</i>' in html
    assert '<td colspan="2">' in html and '按期交付' in html and html.count('<tr>') == 3
    assert '<img src="data:image/png;base64,' in html
    assert '<hr class="pb">' in html and '<h2>售后服务方案</h2>' in html
    assert '&lt;脚本&gt; &amp; 转义' in html and '<脚本>' not in html
    assert stats['tables'] == 1 and stats['images'] == 1 and stats['paragraphs'] >= 6
    assert stats['truncated'] is False and stats['images_skipped'] == 0


def test_docx_preview_stops_at_the_block_budget(engine, tmp_path, monkeypatch):
    path = tmp_path / '长文.docx'
    _build(str(path), pages=6)
    monkeypatch.setattr(engine, 'DOCX_PREVIEW_MAX_BLOCKS', 8)
    html, stats = engine._docx_preview_html(str(path))
    assert stats['truncated'] is True and '预览到此为止' in html


def test_artifact_html_endpoint_only_serves_registered_word_files(engine, job):
    _build(str(job / '投标文件_整册.docx'))
    (job / '说明.md').write_text('# 说明\n', encoding='utf-8')
    with TestClient(engine.app) as client:
        good = client.get('/v1/jobs/%s/artifacts/投标文件_整册.docx/html' % job.name)
        assert good.status_code == 200
        body = good.json()
        assert body['ok'] is True and body['name'] == '投标文件_整册.docx'
        assert '<h1>投标文件</h1>' in body['html'] and body['stats']['tables'] == 1
        assert client.get('/v1/jobs/%s/artifacts/说明.md/html' % job.name).status_code == 400
        assert client.get('/v1/jobs/%s/artifacts/不存在.docx/html' % job.name).status_code == 404
        assert client.get('/v1/jobs/%s/artifacts/..%%2F任务.json/html' % job.name).status_code in (404, 422)
        # 招标原件是个假 docx(fixture 写的是纯文本):解析失败要报 422,不能 500
        assert client.get('/v1/jobs/%s/artifacts/招标文件.docx/html' % job.name).status_code in (404, 422)
