import hashlib
import json
import os
import threading
import time

import pytest
from docx import Document
from fastapi.testclient import TestClient

from conftest import events


SKILL_EVENT_FILE = "skill_events.jsonl"


def _sha256_bytes(value):
    return hashlib.sha256(value).hexdigest()


def _skill_manifest(tmp_path, run_id, *, execution_path="opencode", present=True,
                    injected=True, accepted=True):
    """Build a synthetic manifest without persisting the raw fixture path."""
    skill_dir = tmp_path / ("synthetic-skill-" + run_id)
    skill_path = skill_dir / "SKILL.md"
    if present:
        skill_dir.mkdir(parents=True, exist_ok=True)
        skill_path.write_text("# Synthetic skill\n\nFixture-only instructions.\n", encoding="utf-8")
        manifest_sha256 = _sha256_bytes(skill_path.read_bytes())
    else:
        manifest_sha256 = ""
    path_sha256 = _sha256_bytes(os.path.realpath(skill_path).encode("utf-8"))
    manifest = {
        "run_id": run_id,
        "version": "synthetic-1",
        "manifest_sha256": manifest_sha256,
        "path_sha256": path_sha256,
        "injected": bool(injected),
        "accepted": bool(accepted),
        "execution_path": execution_path,
        "file_present": bool(present),
    }
    return manifest, skill_path


def _set_current_run(engine, job, run_id, manifest=None):
    meta_path = job / "任务.json"
    meta = engine.read_json(str(meta_path), {})
    meta["run_id"] = run_id
    if manifest is None:
        meta.pop("skill_manifest", None)
    else:
        meta["skill_manifest"] = manifest
    engine.write_json(str(meta_path), meta)


def _write_skill_events(engine, job, *records):
    for record in records:
        engine._append_skill_event(str(job), record)


def _read_event(manifest, *, run_id=None, status="completed", path_sha256=None,
                manifest_sha256=None, event_type="read_manifest"):
    return {
        "run_id": run_id or manifest["run_id"],
        "type": event_type,
        "status": status,
        "path_sha256": path_sha256 or manifest["path_sha256"],
        "manifest_sha256": manifest_sha256 or manifest["manifest_sha256"],
        "ts": "2026-08-07T00:00:00Z",
    }


def _oc_read_event(skill_path, session_id, *, event_session=None, offset=None,
                   limit=None, truncated=False):
    lines = skill_path.read_text(encoding="utf-8").splitlines()
    tool_input = {"filePath": str(skill_path)}
    if offset is not None:
        tool_input["offset"] = offset
    if limit is not None:
        tool_input["limit"] = limit
    output = "<content>\n" + "\n".join(
        "%d: %s" % (index + 1, line) for index, line in enumerate(lines)
    ) + "\n\n(End of file - total %d lines)\n</content>" % len(lines)
    return {
        "type": "message.part.updated",
        "properties": {
            "sessionID": event_session or session_id,
            "part": {
                "type": "tool",
                "tool": "read",
                "sessionID": event_session or session_id,
                "state": {
                    "status": "completed",
                    "input": tool_input,
                    "output": output,
                    "metadata": {"truncated": bool(truncated)},
                },
            }
        },
    }


def _oc_next_called(skill_path, session_id, call_id="call-read", *, offset=1, limit=2000):
    return {
        "type": "session.next.tool.called",
        "data": {
            "sessionID": session_id,
            "assistantMessageID": "message-current",
            "callID": call_id,
            "tool": "read",
            "input": {"path": str(skill_path), "offset": offset, "limit": limit},
            "provider": {"executed": False},
        },
    }


def _oc_next_success(skill_path, session_id, call_id="call-read", *, truncated=False,
                     offset=1, next_offset=None):
    structured = {
        "type": "text-page",
        "content": skill_path.read_text(encoding="utf-8"),
        "mime": "text/markdown",
        "offset": offset,
        "truncated": bool(truncated),
    }
    if next_offset is not None:
        structured["next"] = next_offset
    return {
        "type": "session.next.tool.success",
        "data": {
            "sessionID": session_id,
            "assistantMessageID": "message-current",
            "callID": call_id,
            "structured": structured,
            "content": [],
            "outputPaths": [],
            "provider": {"executed": False},
        },
    }


def _write_body_docx(path):
    document = Document()
    document.add_heading("投标文件", level=1)
    for index in range(80):
        document.add_paragraph("第%d项完整响应，满足招标文件要求并提供实施说明。" % (index + 1))
    document.save(path)


@pytest.mark.parametrize("execution_path", ["builtin", "opencode"])
def test_accepted_skill_without_runtime_read_is_unverifiable_not_missing(
    engine, job, tmp_path, execution_path
):
    run_id = "run-current"
    manifest, raw_path = _skill_manifest(tmp_path, run_id, execution_path=execution_path)
    _set_current_run(engine, job, run_id, manifest)

    result = engine.skill_evidence(str(job))

    assert result["state"] == "unverifiable"
    assert result["ok"] is False
    assert result["why"]
    assert not any(word in result["why"] for word in ("没使用", "未使用", "没有检测到技能包被使用"))
    assert str(raw_path) not in json.dumps(result, ensure_ascii=False)


def test_unverifiable_skill_completion_message_never_claims_the_skill_was_not_used(
    engine, job, tmp_path, monkeypatch
):
    run_id = "run-current"
    manifest, _ = _skill_manifest(tmp_path, run_id, execution_path="opencode")
    _set_current_run(engine, job, run_id, manifest)
    _write_body_docx(job / "投标文件_技术标.docx")
    monkeypatch.setattr(engine, "quality_audit", lambda *_args, **_kwargs: None)

    result = engine.settle(str(job))

    assert result["state"] == "done"
    text = "\n".join(
        str(event.get("text") or "")
        for event in events(job)
        if event.get("type") in ("message", "error")
    )
    assert "无法确认" in text or "无法核验" in text
    assert "没有检测到技能包被使用" not in text
    assert "没使用" not in text
    assert "未使用" not in text


def test_completed_current_run_read_with_exact_manifest_hashes_is_verified(engine, job, tmp_path):
    run_id = "run-current"
    manifest, raw_path = _skill_manifest(tmp_path, run_id)
    _set_current_run(engine, job, run_id, manifest)
    _write_skill_events(engine, job, _read_event(manifest))

    result = engine.skill_evidence(str(job))

    assert result["state"] == "verified"
    assert result["ok"] is True
    assert result["hits"]
    assert str(raw_path) not in json.dumps(result, ensure_ascii=False)


@pytest.mark.parametrize(
    "event_override",
    [
        {"path_sha256": _sha256_bytes(b"another-synthetic-path")},
        {"manifest_sha256": _sha256_bytes(b"another-synthetic-manifest")},
        {"status": "failed"},
        {"status": "running"},
        {"run_id": "run-old"},
        {"event_type": "read"},
    ],
    ids=[
        "path-mismatch",
        "manifest-mismatch",
        "failed-read",
        "running-read",
        "old-run",
        "wrong-event-type",
    ],
)
def test_non_exact_or_non_completed_read_never_verifies_current_run(
    engine, job, tmp_path, event_override
):
    run_id = "run-current"
    manifest, _ = _skill_manifest(tmp_path, run_id)
    _set_current_run(engine, job, run_id, manifest)
    _write_skill_events(engine, job, _read_event(manifest, **event_override))

    result = engine.skill_evidence(str(job))

    assert result["state"] == "unverifiable"
    assert result["ok"] is False


def test_old_exact_evidence_cannot_cover_a_new_run(engine, job, tmp_path):
    old_manifest, _ = _skill_manifest(tmp_path, "run-old")
    current_manifest, _ = _skill_manifest(tmp_path, "run-current")
    _set_current_run(engine, job, "run-current", current_manifest)
    _write_skill_events(engine, job, _read_event(old_manifest))

    result = engine.skill_evidence(str(job))

    assert result["state"] == "unverifiable"
    assert result["ok"] is False


def test_log_echo_of_skill_filename_is_not_verified_evidence(engine, job, tmp_path):
    run_id = "run-current"
    manifest, _ = _skill_manifest(tmp_path, run_id)
    _set_current_run(engine, job, run_id, manifest)
    (job / "run.log").write_text(
        "Prompt says to read SKILL.md and run build_tender_docx later.\n",
        encoding="utf-8",
    )

    result = engine.skill_evidence(str(job))

    assert result["state"] == "unverifiable"
    assert result["ok"] is False


@pytest.mark.parametrize(
    "manifest_mode",
    ["absent", "file-missing", "not-injected", "not-accepted"],
)
def test_missing_or_undelivered_skill_is_missing(engine, job, tmp_path, manifest_mode):
    run_id = "run-current"
    if manifest_mode == "absent":
        manifest = None
    else:
        manifest, _ = _skill_manifest(
            tmp_path,
            run_id,
            present=manifest_mode != "file-missing",
            injected=manifest_mode != "not-injected",
            accepted=manifest_mode != "not-accepted",
        )
    _set_current_run(engine, job, run_id, manifest)

    result = engine.skill_evidence(str(job))

    assert result["state"] == "missing"
    assert result["ok"] is False
    assert result["why"]


@pytest.mark.parametrize("expected_state", ["verified", "unverifiable", "missing"])
def test_job_log_exposes_the_skill_state(engine, job, tmp_path, expected_state):
    run_id = "run-current"
    if expected_state == "missing":
        manifest = None
    else:
        manifest, _ = _skill_manifest(tmp_path, run_id)
    _set_current_run(engine, job, run_id, manifest)
    if expected_state == "verified":
        _write_skill_events(engine, job, _read_event(manifest))

    result = engine.job_log(job.name)

    assert result["skill_state"] == expected_state
    assert result["skill_used"] is (expected_state == "verified")


def test_new_launch_rotates_run_id_before_work_starts(engine, job, monkeypatch):
    _set_current_run(engine, job, "run-old", None)
    monkeypatch.setattr(engine, "mock_agent", lambda *_args, **_kwargs: None)

    engine._launch_job(job.name, str(job), mock="1")

    current = engine.read_json(str(job / "任务.json"), {}).get("run_id")
    assert current
    assert current != "run-old"


def test_redo_rotates_run_id_before_work_starts(engine, job, monkeypatch):
    _set_current_run(engine, job, "run-old", None)
    monkeypatch.setattr(engine, "config_agent_cmd", lambda: "synthetic-agent")
    monkeypatch.setattr(engine, "real_agent", lambda *_args, **_kwargs: None)

    with TestClient(engine.app) as client:
        response = client.post(
            "/v1/jobs/%s/redo" % job.name,
            json={"instruction": "只重做第三章"},
        )

    assert response.status_code == 200
    current = engine.read_json(str(job / "任务.json"), {}).get("run_id")
    assert current
    assert current != "run-old"


def test_resume_keeps_the_same_run_id(engine, job, monkeypatch):
    run_id = "run-current"
    _set_current_run(engine, job, run_id, None)
    meta_path = job / "任务.json"
    meta = engine.read_json(str(meta_path), {})
    meta["oc_session"] = "session-current"
    engine.write_json(str(meta_path), meta)
    monkeypatch.setattr(engine, "oc_serve", lambda: "http://127.0.0.1:9")
    monkeypatch.setattr(engine, "oc_api", lambda *_args, **_kwargs: (200, {"id": "session-current"}))
    monkeypatch.setattr(engine, "resume_worker", lambda *_args, **_kwargs: None)

    with TestClient(engine.app) as client:
        response = client.post("/v1/jobs/%s/resume" % job.name)

    assert response.status_code == 200
    assert engine.read_json(str(meta_path), {}).get("run_id") == run_id


def test_unsigned_job_directory_event_cannot_forge_verified_state(engine, job, tmp_path):
    """The agent can write its cwd, so an unsigned JSON line must never be trusted."""
    run_id = "run-current"
    manifest, _ = _skill_manifest(tmp_path, run_id)
    _set_current_run(engine, job, run_id, manifest)
    forged = _read_event(manifest)
    (job / SKILL_EVENT_FILE).write_text(
        json.dumps(forged, ensure_ascii=False) + "\n", encoding="utf-8"
    )

    result = engine.skill_evidence(str(job))

    assert result["state"] == "unverifiable"
    assert result["ok"] is False


def test_legacy_job_without_structured_receipt_is_unverifiable_not_missing(engine, job):
    """Pre-0.18.2 jobs never had run_id/manifest, so absence proves nothing."""
    meta_path = job / "任务.json"
    meta = engine.read_json(str(meta_path), {})
    meta.pop("run_id", None)
    meta.pop("skill_manifest", None)
    engine.write_json(str(meta_path), meta)

    result = engine.skill_evidence(str(job))

    assert result["state"] == "unverifiable"
    assert "旧" in result["why"] or "历史" in result["why"]


def test_rerun_uses_launch_path_and_creates_current_skill_receipt(engine, job, monkeypatch):
    monkeypatch.setattr(engine, "config_agent_cmd", lambda: "synthetic-agent {skill}")
    monkeypatch.setattr(engine, "real_agent", lambda *_args, **_kwargs: None)

    with TestClient(engine.app) as client:
        response = client.post("/v1/jobs/%s/rerun" % job.name)

    assert response.status_code == 200
    body = response.json()
    meta = engine.read_json(str(engine.jpath(body["job_id"])) + "/任务.json", {})
    assert meta.get("run_id")
    assert isinstance(meta.get("skill_manifest"), dict)
    assert meta["skill_manifest"].get("run_id") == meta["run_id"]


def test_observed_full_read_requires_the_current_opencode_session(engine, job, tmp_path):
    run_id = "run-current"
    session_id = "session-current"
    manifest, skill_path = _skill_manifest(tmp_path, run_id)
    _set_current_run(engine, job, run_id, manifest)
    meta_path = job / "任务.json"
    meta = engine.read_json(str(meta_path), {})
    meta["oc_session"] = session_id
    engine.write_json(str(meta_path), meta)

    observed = engine._observe_oc_skill_event(
        str(job), _oc_read_event(skill_path, session_id), session_id
    )

    assert observed is True
    assert engine.skill_evidence(str(job))["state"] == "verified"


@pytest.mark.parametrize(
    "event_kwargs",
    [
        {"event_session": "session-other"},
        {"offset": 2},
        {"limit": 1},
        {"truncated": True},
    ],
    ids=["wrong-session", "partial-offset", "partial-limit", "truncated-output"],
)
def test_wrong_session_or_partial_read_is_not_observed(
    engine, job, tmp_path, event_kwargs
):
    run_id = "run-current"
    session_id = "session-current"
    manifest, skill_path = _skill_manifest(tmp_path, run_id)
    _set_current_run(engine, job, run_id, manifest)
    meta_path = job / "任务.json"
    meta = engine.read_json(str(meta_path), {})
    meta["oc_session"] = session_id
    engine.write_json(str(meta_path), meta)

    observed = engine._observe_oc_skill_event(
        str(job), _oc_read_event(skill_path, session_id, **event_kwargs), session_id
    )

    assert observed is False
    assert engine.skill_evidence(str(job))["state"] == "unverifiable"


def test_durable_opencode_read_call_and_success_verify_the_current_run(
    engine, job, tmp_path
):
    run_id = "run-current"
    session_id = "session-current"
    manifest, skill_path = _skill_manifest(tmp_path, run_id)
    _set_current_run(engine, job, run_id, manifest)
    meta_path = job / "任务.json"
    meta = engine.read_json(str(meta_path), {})
    meta["oc_session"] = session_id
    engine.write_json(str(meta_path), meta)

    assert engine._observe_oc_skill_event(
        str(job), _oc_next_called(skill_path, session_id), session_id
    ) is False
    assert engine._observe_oc_skill_event(
        str(job), _oc_next_success(skill_path, session_id), session_id
    ) is True
    assert engine.skill_evidence(str(job))["state"] == "verified"


def test_durable_partial_read_success_never_verifies(engine, job, tmp_path):
    run_id = "run-current"
    session_id = "session-current"
    manifest, skill_path = _skill_manifest(tmp_path, run_id)
    _set_current_run(engine, job, run_id, manifest)
    meta_path = job / "任务.json"
    meta = engine.read_json(str(meta_path), {})
    meta["oc_session"] = session_id
    engine.write_json(str(meta_path), meta)

    assert engine._observe_oc_skill_event(
        str(job), _oc_next_called(skill_path, session_id, limit=1), session_id
    ) is False
    assert engine._observe_oc_skill_event(
        str(job), _oc_next_success(
            skill_path, session_id, truncated=True, next_offset=2
        ), session_id
    ) is False
    assert engine.skill_evidence(str(job))["state"] == "unverifiable"


def test_job_log_replays_opencode_evidence_after_engine_restart(
    engine, job, tmp_path, monkeypatch
):
    run_id = "run-current"
    session_id = "session-current"
    manifest, skill_path = _skill_manifest(tmp_path, run_id)
    _set_current_run(engine, job, run_id, manifest)
    meta_path = job / "任务.json"
    meta = engine.read_json(str(meta_path), {})
    meta["oc_session"] = session_id
    engine.write_json(str(meta_path), meta)
    engine._append_skill_event(str(job), _read_event(manifest))
    assert engine.skill_evidence(str(job))["state"] == "verified"

    # A real App restart rotates the in-memory signing key. The durable OpenCode
    # session is the trusted source used to re-issue a current-process receipt.
    monkeypatch.setattr(engine, "_SKILL_EVENT_KEY", b"new-engine-process-key" * 2)
    assert engine.skill_evidence(str(job))["state"] == "unverifiable"
    payloads = [
        _oc_next_called(skill_path, session_id),
        _oc_next_success(skill_path, session_id),
    ]
    # Durable evidence belongs to the historical run. An App/skill upgrade, or
    # switching today's engine to Codex, must not invalidate that old receipt.
    skill_path.write_text("# Upgraded skill\n\nDifferent instructions.\n", encoding="utf-8")
    engine.write_json(
        engine.conf_path(),
        {"engine": {"kind": "codex", "cli_path": "/synthetic/codex"}},
    )
    replay_engines = []

    class FakeEventStream:
        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        def __iter__(self):
            return iter(
                [("data: " + json.dumps(item, ensure_ascii=False) + "\n").encode("utf-8")
                 for item in payloads]
            )

    def replay_server(eng=None):
        replay_engines.append(dict(eng or {}))
        return "http://127.0.0.1:9"

    monkeypatch.setattr(engine, "oc_serve", replay_server)
    monkeypatch.setattr(engine.urllib.request, "urlopen", lambda *_a, **_k: FakeEventStream())

    result = engine.job_log(job.name)

    assert result["skill_state"] == "verified"
    assert result["skill_used"] is True
    assert replay_engines and replay_engines[-1].get("kind") == "s2"


def test_launch_reserves_the_job_before_worker_thread_runs(engine, job, monkeypatch):
    entered = threading.Event()
    release = threading.Event()
    calls = []

    def blocked_mock(job_path):
        calls.append(job_path)
        entered.set()
        release.wait(2)

    monkeypatch.setattr(engine, "mock_agent", blocked_mock)
    first = engine._launch_job(job.name, str(job), mock="1")
    assert entered.wait(1)
    first_run = engine.read_json(str(job / "任务.json"), {}).get("run_id")

    second = engine._launch_job(job.name, str(job), mock="1")
    second_run = engine.read_json(str(job / "任务.json"), {}).get("run_id")
    release.set()
    time.sleep(0.05)

    assert first["mode"] == "mock"
    assert second["mode"] == "running"
    assert second.get("already_running") is True
    assert second_run == first_run
    assert len(calls) == 1


def test_redo_reserves_the_job_before_worker_thread_runs(engine, job, monkeypatch):
    entered = threading.Event()
    release = threading.Event()
    calls = []

    def blocked_redo(job_path, instruction):
        calls.append((job_path, instruction))
        entered.set()
        release.wait(2)

    monkeypatch.setattr(engine, "config_agent_cmd", lambda: None)
    monkeypatch.setattr(engine, "mock_redo", blocked_redo)

    with TestClient(engine.app) as client:
        first = client.post(
            "/v1/jobs/%s/redo" % job.name, json={"instruction": "只重做第三章"}
        )
        assert first.status_code == 200
        assert entered.wait(1)
        first_run = engine.read_json(str(job / "任务.json"), {}).get("run_id")
        second = client.post(
            "/v1/jobs/%s/redo" % job.name, json={"instruction": "只重做第四章"}
        )
        second_run = engine.read_json(str(job / "任务.json"), {}).get("run_id")

    release.set()
    time.sleep(0.05)
    assert second.status_code == 409
    assert second_run == first_run
    assert len(calls) == 1


def test_skill_version_metadata_never_persists_arbitrary_text(engine, tmp_path):
    skill_dir = tmp_path / "custom-skill"
    skill_dir.mkdir()
    (skill_dir / "SKILL.md").write_text("# Safe instructions\n", encoding="utf-8")
    raw_version = "customer private note with spaces"
    (skill_dir / ".skill_version").write_text(raw_version, encoding="utf-8")

    manifest = engine._skill_manifest(
        "run-current", str(skill_dir), "dispatch", True, "opencode"
    )

    encoded = json.dumps(manifest, ensure_ascii=False)
    assert raw_version not in encoded
    assert manifest["version"].startswith("custom-")


@pytest.mark.parametrize(
    "raw_version",
    [
        "ghp_" + "A" * 36,
        "sk-" + "B" * 37,
        "private-token-without-spaces",
    ],
)
def test_skill_version_only_persists_numeric_version_labels(engine, tmp_path, raw_version):
    suffix = hashlib.sha256(raw_version.encode()).hexdigest()[:8]
    skill_dir = tmp_path / ("custom-skill-" + suffix)
    skill_dir.mkdir()
    (skill_dir / "SKILL.md").write_text("# Safe instructions\n", encoding="utf-8")
    (skill_dir / ".skill_version").write_text(raw_version, encoding="utf-8")

    manifest = engine._skill_manifest(
        "run-current", str(skill_dir), "dispatch", True, "opencode"
    )

    assert raw_version not in json.dumps(manifest, ensure_ascii=False)
    assert manifest["version"].startswith("custom-")


def test_skill_version_invalid_utf8_is_hashed_without_aborting_launch(engine, tmp_path):
    skill_dir = tmp_path / "binary-version-skill"
    skill_dir.mkdir()
    (skill_dir / "SKILL.md").write_text("# Safe instructions\n", encoding="utf-8")
    (skill_dir / ".skill_version").write_bytes(b"\xff\xfesecret")

    manifest = engine._skill_manifest(
        "run-current", str(skill_dir), "dispatch", True, "opencode"
    )

    assert manifest["version"].startswith("custom-")


def test_read_tail_passes_a_strict_read_limit(engine, monkeypatch):
    class GrowingFile:
        def __init__(self):
            self.read_sizes = []

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        def seek(self, *_args):
            return None

        def tell(self):
            return 100

        def read(self, size=None):
            self.read_sizes.append(size)
            return b"x" * (size or 1000)

    growing = GrowingFile()
    monkeypatch.setattr("builtins.open", lambda *_a, **_k: growing)

    assert engine.read_tail("synthetic-growing.log", 17) == "x" * 17
    assert growing.read_sizes == [17]


def test_reserved_worker_holds_slot_through_server_fallback_cleanup(engine, job, monkeypatch):
    entered_fallback = threading.Event()
    release_fallback = threading.Event()

    monkeypatch.setattr(engine, "oc_run", lambda *_a, **_k: engine.OC_RUN_FALLBACK)
    monkeypatch.setattr(engine, "ensure_default_shell", lambda *_args, **_kwargs: (True, "ready"))

    def blocked_real_agent(*_args):
        entered_fallback.set()
        release_fallback.wait(2)

    monkeypatch.setattr(engine, "real_agent", blocked_real_agent)
    owner = engine._reserve_running(job.name)
    assert owner
    worker = engine._start_reserved_worker(
        job.name, owner, engine.agent_via_server_or_cli,
        str(job), "synthetic prompt", ["synthetic-cli"],
    )

    assert entered_fallback.wait(1)
    assert engine._reserve_running(job.name) is None
    release_fallback.set()
    worker.join(1)
    assert job.name not in engine.RUNNING


def test_stop_keeps_reservation_until_worker_exits(engine, job, monkeypatch):
    entered = threading.Event()
    release = threading.Event()

    def blocked_worker():
        entered.set()
        release.wait(2)

    owner = engine._reserve_running(job.name)
    worker = engine._start_reserved_worker(job.name, owner, blocked_worker)
    assert entered.wait(1)
    monkeypatch.setattr(engine, "kill_tree", lambda *_args: None)

    result = engine.stop_job(job.name)

    assert result["ok"] is True
    assert engine._reserve_running(job.name) is None
    release.set()
    worker.join(1)
    assert job.name not in engine.RUNNING
    next_owner = engine._reserve_running(job.name)
    assert next_owner and not engine._cancel_requested(job.name, next_owner)
    engine._release_running(job.name, next_owner)


def test_late_owner_cannot_release_a_new_generation(engine, job):
    first = engine._reserve_running(job.name)
    assert engine._release_running(job.name, first) is True
    second = engine._reserve_running(job.name)
    assert second and second != first

    assert engine._release_running(job.name, first) is False
    assert job.name in engine.RUNNING
    assert engine._release_running(job.name, second) is True


def test_start_rejection_does_not_clobber_staged_metadata(engine, job):
    meta_path = job / "任务.json"
    meta = engine.read_json(str(meta_path), {})
    meta.update({"staged": True, "run_id": "first-owner-run"})
    engine.write_json(str(meta_path), meta)
    owner = engine._reserve_running(job.name)
    assert owner

    with TestClient(engine.app) as client:
        response = client.post("/v1/jobs/%s/start?mock=1" % job.name)

    after = engine.read_json(str(meta_path), {})
    engine._release_running(job.name, owner)
    assert response.status_code == 200
    assert response.json().get("already_running") is True
    assert after.get("staged") is True
    assert after.get("run_id") == "first-owner-run"


def test_delete_reports_failure_when_job_directory_remains(engine, job, monkeypatch):
    monkeypatch.setattr(engine.shutil, "rmtree", lambda *_a, **_k: None)
    monkeypatch.setattr(engine.time, "sleep", lambda *_a: None)

    with TestClient(engine.app) as client:
        response = client.delete("/v1/jobs/%s" % job.name)

    assert response.status_code == 500
    assert response.json().get("ok") is False
    assert job.is_dir()


def test_delete_running_opencode_job_interrupts_session_before_removal(
    engine, job, monkeypatch
):
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["oc_session"] = "session-delete"
    engine.write_json(str(job / "任务.json"), meta)
    engine.OC["base"] = "http://127.0.0.1:9"
    interrupted = []
    monkeypatch.setattr(engine, "oc_interrupt", lambda sid: interrupted.append(sid) or True)

    def cancellable_worker():
        while not engine._cancel_requested(job.name):
            time.sleep(0.01)

    owner = engine._reserve_running(job.name)
    worker = engine._start_reserved_worker(job.name, owner, cancellable_worker)
    with TestClient(engine.app) as client:
        response = client.delete("/v1/jobs/%s" % job.name)
    worker.join(1)

    assert response.status_code == 200
    assert interrupted == ["session-delete"]
    assert not job.exists()
    assert job.name not in engine.RUNNING


def test_delete_keeps_job_when_opencode_interrupt_is_unconfirmed(
    engine, job, monkeypatch
):
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["oc_session"] = "session-busy"
    engine.write_json(str(job / "任务.json"), meta)
    engine.OC["base"] = "http://127.0.0.1:9"
    release = threading.Event()
    monkeypatch.setattr(engine, "oc_interrupt", lambda _sid: False)
    monkeypatch.setattr(engine, "oc_turn", lambda _sid: (False, ""))

    owner = engine._reserve_running(job.name)
    worker = engine._start_reserved_worker(job.name, owner, lambda: release.wait(2))
    with TestClient(engine.app) as client:
        response = client.delete("/v1/jobs/%s" % job.name)

    assert response.status_code == 502
    assert job.exists()
    assert not engine._cancel_requested(job.name, owner)
    release.set()
    worker.join(1)


def test_stop_does_not_claim_success_when_opencode_interrupt_fails(
    engine, job, monkeypatch
):
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["oc_session"] = "session-busy"
    engine.write_json(str(job / "任务.json"), meta)
    engine.OC["base"] = "http://127.0.0.1:9"
    release = threading.Event()
    monkeypatch.setattr(engine, "oc_interrupt", lambda _sid: False)
    monkeypatch.setattr(engine, "oc_turn", lambda _sid: (False, ""))

    owner = engine._reserve_running(job.name)
    worker = engine._start_reserved_worker(job.name, owner, lambda: release.wait(2))
    result = engine.stop_job(job.name)

    assert result.status_code == 502
    assert not engine._cancel_requested(job.name, owner)
    assert engine.read_json(str(job / "outcome.json"), {}) == {}
    release.set()
    worker.join(1)


def test_cli_wait_kills_process_when_current_owner_is_cancelled(engine, job, monkeypatch):
    class FakeProcess:
        @staticmethod
        def poll():
            return None

    proc = FakeProcess()
    killed = []
    owner = engine._reserve_running(job.name)
    assert engine._request_cancel(job.name, owner)
    monkeypatch.setattr(engine, "kill_tree", lambda target: killed.append(target))

    result = engine.wait_cli_process(proc, str(job), poll_seconds=0.01)

    assert result["status"] == "cancelled"
    assert killed == [proc]
    engine._release_running(job.name, owner)


def test_opencode_evidence_replay_blocks_new_launch_admission(
    engine, job, tmp_path, monkeypatch
):
    run_id = "run-replay-lock"
    session_id = "session-replay-lock"
    manifest, _skill_path = _skill_manifest(tmp_path, run_id)
    _set_current_run(engine, job, run_id, manifest)
    meta = engine.read_json(str(job / "任务.json"), {})
    meta["oc_session"] = session_id
    engine.write_json(str(job / "任务.json"), meta)
    entered = threading.Event()
    release = threading.Event()

    def blocked_server(_eng=None):
        entered.set()
        release.wait(2)
        return ""

    monkeypatch.setattr(engine, "oc_serve", blocked_server)
    replay = threading.Thread(
        target=engine._replay_oc_skill_evidence, args=(str(job),), daemon=True
    )
    replay.start()
    assert entered.wait(1)
    assert engine._reserve_running("new-job") is None
    release.set()
    replay.join(1)
    owner = engine._reserve_running("new-job")
    assert owner
    engine._release_running("new-job", owner)


def test_create_during_evidence_replay_is_staged_and_returns_retryable_error(
    engine, tmp_path
):
    engine.OC_REPLAYING = True
    jobs_dir = tmp_path / "jobs"

    with TestClient(engine.app) as client:
        response = client.post(
            "/v1/jobs",
            files={"tender": ("招标文件.docx", b"synthetic tender")},
            data={"name": "replay admission", "start": "1", "mock": "1"},
        )

    created = [path for path in jobs_dir.iterdir() if path.is_dir()]
    assert response.status_code == 503
    assert response.json().get("retryable") is True
    assert len(created) == 1
    meta = engine.read_json(str(created[0] / "任务.json"), {})
    progress = engine.read_json(str(created[0] / "progress.json"), {})
    assert meta.get("staged") is True
    assert progress.get("step") == 0
    assert created[0].name not in engine.RUNNING


def test_rerun_during_evidence_replay_is_staged_instead_of_fake_running(
    engine, job
):
    engine.OC_REPLAYING = True

    with TestClient(engine.app) as client:
        response = client.post("/v1/jobs/%s/rerun" % job.name)

    reruns = [path for path in job.parent.iterdir() if path.is_dir() and path != job]
    assert response.status_code == 503
    assert len(reruns) == 1
    assert engine.read_json(str(reruns[0] / "任务.json"), {}).get("staged") is True
    assert reruns[0].name not in engine.RUNNING


def test_stop_and_done_commit_have_one_atomic_terminal_winner(
    engine, job, monkeypatch
):
    word = job / "投标文件_技术标.docx"
    document = Document()
    document.add_heading("投标文件", level=1)
    for index in range(80):
        document.add_paragraph("第%d项完整响应，满足招标文件要求。" % index)
    document.save(word)
    entered = threading.Event()
    release = threading.Event()
    calls = 0
    original_body_docxs = engine._body_docxs

    def block_before_terminal_commit(*args, **kwargs):
        nonlocal calls
        calls += 1
        result = original_body_docxs(*args, **kwargs)
        if calls == 2:
            entered.set()
            release.wait(2)
        return result

    monkeypatch.setattr(engine, "_body_docxs", block_before_terminal_commit)
    monkeypatch.setattr(engine, "quality_audit", lambda *_args: None)
    owner = engine._reserve_running(job.name)
    worker = threading.Thread(target=engine.settle, args=(str(job),), daemon=True)
    worker.start()
    assert entered.wait(1)

    stopped = engine.stop_job(job.name)
    assert stopped["ok"] is True
    release.set()
    worker.join(2)

    assert engine.read_json(str(job / "outcome.json"), {}).get("state") == "stopped"
    assert not any(
        event.get("type") == "progress" and event.get("step") == 12
        for event in events(job)
    )
    engine._release_running(job.name, owner)


def test_chat_pause_never_claims_success_without_real_control(engine, job, monkeypatch):
    monkeypatch.setattr(engine, "config_agent_cmd", lambda: ["synthetic-agent"])
    owner = engine._reserve_running(job.name)
    before = engine.read_json(str(job / "任务.json"), {})

    handled = engine.route_command(str(job), job.name, "暂停一下", True)

    after = engine.read_json(str(job / "任务.json"), {})
    messages = [
        event.get("text", "") for event in events(job)
        if event.get("type") == "message" and event.get("role") == "agent"
    ]
    assert handled is True
    assert after.get("paused") == before.get("paused")
    assert not engine._cancel_requested(job.name, owner)
    assert messages and "暂停按钮" in messages[-1]
    assert "已暂停" not in messages[-1]
    engine._release_running(job.name, owner)


def test_shutdown_admission_does_not_restage_an_already_running_job(engine, job):
    meta_path = job / "任务.json"
    before_meta = engine.read_json(str(meta_path), {})
    before_progress = engine.read_json(str(job / "progress.json"), {})
    owner = engine._reserve_running(job.name)
    engine.SHUTTING_DOWN = True

    with TestClient(engine.app) as client:
        response = client.post("/v1/jobs/%s/start?mock=1" % job.name)

    after_meta = engine.read_json(str(meta_path), {})
    after_progress = engine.read_json(str(job / "progress.json"), {})
    assert response.status_code == 200
    assert response.json().get("already_running") is True
    assert after_meta == before_meta
    assert after_progress == before_progress
    engine.SHUTTING_DOWN = False
    engine._release_running(job.name, owner)


def test_mock_pause_is_rejected_without_stranding_the_running_owner(
    engine, job, monkeypatch
):
    monkeypatch.setattr(engine, "config_agent_cmd", lambda: None)
    owner = engine._reserve_running(job.name)
    before = engine.read_json(str(job / "任务.json"), {})

    with TestClient(engine.app) as client:
        response = client.post(
            "/v1/jobs/%s/control" % job.name, json={"action": "pause"}
        )

    after = engine.read_json(str(job / "任务.json"), {})
    assert response.status_code == 400
    assert response.json().get("ok") is False
    assert after.get("paused") == before.get("paused")
    assert engine._owner_running(job.name, owner)
    assert not engine._cancel_requested(job.name, owner)
    assert engine.read_json(str(job / "outcome.json"), {}) == {}
    engine._release_running(job.name, owner)
