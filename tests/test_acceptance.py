import json

import pytest
from docx import Document

import acceptance


def _write_sample(path, text):
    document = Document()
    document.add_heading(text, level=1)
    for index in range(20):
        document.add_paragraph("%s-%d-采购要求和响应条款" % (text, index))
    document.save(path)


def test_acceptance_manifest_rejects_duplicate_documents(tmp_path):
    first = tmp_path / "采购文件一.docx"
    second = tmp_path / "采购文件二.docx"
    third = tmp_path / "采购文件三.docx"
    _write_sample(first, "甲")
    second.write_bytes(first.read_bytes())
    _write_sample(third, "丙")
    manifest = tmp_path / "samples.json"
    manifest.write_text(
        json.dumps(
            {"samples": [
                {"name": "一", "path": str(first)},
                {"name": "二", "path": str(second)},
                {"name": "三", "path": str(third)},
            ]},
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="duplicates"):
        acceptance.load_samples(manifest)


def test_acceptance_word_gate_rejects_analysis_docx():
    assert acceptance.is_body_word("投标文件_技术标.docx")
    assert not acceptance.is_body_word("招标文件_解析版.docx")
    assert not acceptance.is_body_word("成品质检报告.docx")
