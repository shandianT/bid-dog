"""场景模板引擎。

外部调用只需要四个接口：列出内置模板、推荐模板、从优秀标书提炼模板草稿、
把模板编译成任务指令。复杂的场景规则、脱敏和质量评分都收在本模块内部。
"""

from __future__ import annotations

import copy
import hashlib
import io
import json
import re
import zipfile
from typing import Any, Dict, Iterable, List, Sequence, Tuple


PRIORITY_RULE = "招标文件原文高于模板；冲突时以招标文件为准"
MAX_DOCX_ENTRIES = 5000
MAX_DOCX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 200

SOURCE_GOVERNMENT_87 = {
    "title": "政府采购货物和服务招标投标管理办法（财政部令第87号）",
    "issuer": "中华人民共和国财政部",
    "version": "2017",
    "url": "https://www.mof.gov.cn/gp/xxgkml/tfs/201707/t20170718_2652766.htm",
}
SOURCE_GOVERNMENT_DEMAND = {
    "title": "政府采购需求管理办法（财库〔2021〕22号）",
    "issuer": "中华人民共和国财政部",
    "version": "2021",
    "url": "https://www.mof.gov.cn/gkml/caizhengwengao/wg2021/wg202005/202109/t20210917_3753625.htm",
}
SOURCE_STANDARD_PURCHASE = {
    "title": "标准设备、材料、勘察、设计和监理招标文件（发改法规〔2017〕1606号）",
    "issuer": "国家发展和改革委员会等九部门",
    "version": "2017",
    "url": "https://zfxxgk.ndrc.gov.cn/web/iteminfo.jsp?id=2872",
}
SOURCE_STANDARD_CONSTRUCTION = {
    "title": "标准施工招标资格预审文件和标准施工招标文件",
    "issuer": "国家发展和改革委员会等九部门",
    "version": "2007（2013修订）",
    "url": "https://www.ndrc.gov.cn/xxgk/zcfb/tz/200712/t20071221_965545.html",
}


def _outline(title: str, purpose: str, evidence: Sequence[str] = (), required: bool = True) -> Dict[str, Any]:
    return {
        "title": title,
        "purpose": purpose,
        "required": required,
        "evidence": list(evidence),
    }


def _table(name: str, columns: Sequence[str], required: bool = True) -> Dict[str, Any]:
    return {"name": name, "columns": list(columns), "required": required}


def _slot(name: str, evidence: Sequence[str], required: bool = False) -> Dict[str, Any]:
    return {
        "name": name,
        "evidence": list(evidence),
        "required": required,
        "missing_policy": "缺失时标记〔需补充〕，不得猜测或编造",
    }


def _package(
    category: str,
    subtype: str,
    labels: Sequence[str],
    keywords: Sequence[str],
    outline: Sequence[Dict[str, Any]],
    scoring_focus: Sequence[Dict[str, Any]],
    tables: Sequence[Dict[str, Any]],
    material_slots: Sequence[Dict[str, Any]],
    quality_rules: Sequence[str],
    sources: Sequence[Dict[str, str]],
) -> Dict[str, Any]:
    return {
        "schema_version": 1,
        "priority_rule": PRIORITY_RULE,
        "scene": {
            "category": category,
            "subtype": subtype,
            "labels": list(labels),
            "match_keywords": list(keywords),
        },
        "outline": list(outline),
        "scoring_focus": list(scoring_focus),
        "tables": list(tables),
        "material_slots": list(material_slots),
        "formatting": {
            "include_cover": True,
            "include_toc": True,
            "repeat_table_headers": True,
            "wide_tables_landscape": True,
            "keep_source_heading_order_when_required": True,
        },
        "quality_rules": list(quality_rules),
        "sources": list(sources),
        "generated_from": {"kind": "builtin"},
    }


COMMON_QUALITY_RULES = [
    "每个资格、符合性和评分要求必须映射到明确章节或表格行",
    "每项结论必须能追溯到招标原文或用户素材；无出处不得输出为事实",
    "报价、投标人名称、资质有效期、承诺和签章必须由人确认",
    "缺失材料统一标记〔需补充〕，不得编造资质、案例、人员或参数",
    "禁止在不同章节重复套用同一组小标题或大段模板话术",
    "最终 Word 必须通过目录、页码、字体、表格、图片落位和应答覆盖检查",
]


BUILTIN_TEMPLATES: List[Dict[str, Any]] = [
    {
        "id": "government",
        "name": "政府采购",
        "description": "适合未明确细分类型的政府采购货物与服务项目",
        "prompt": "逐条覆盖资格、符合性、实质性要求和评分办法，所有证明材料明确出处。",
        "settings": {"mode": "standard", "quality_gate": True, "include_toc": True, "include_deviation_tables": True},
        "package": _package(
            "government_procurement", "general", ["政府采购", "综合"],
            ["政府采购", "公开招标", "竞争性磋商", "资格审查", "符合性审查", "综合评分"],
            [
                _outline("投标函与授权文件", "完成法定签署、授权和投标承诺", ["法定代表人身份证明", "授权委托书"]),
                _outline("资格证明文件", "逐项对应资格条件并给出证明材料索引", ["营业执照", "信用记录", "政策声明"]),
                _outline("符合性与实质性响应", "逐条响应无效投标和不可偏离条款", ["招标文件符合性条款"]),
                _outline("报价与商务响应", "按报价口径、付款、交付和合同条件响应", ["报价表", "商务偏离表"]),
                _outline("技术与服务方案", "围绕采购目标、技术要求和服务标准组织方案", ["技术参数", "服务需求"]),
                _outline("实施、验收与售后", "说明实施计划、验收方法、培训和售后保障", ["项目计划", "验收标准"]),
                _outline("评分点响应索引", "把每个评分因素映射到页码、材料和证据", ["评分办法"]),
            ],
            [
                {"name": "资格与符合性", "checks": ["资格条件", "无效投标条款", "政策证明"]},
                {"name": "评分办法", "checks": ["评分因素", "分值", "客观证据", "响应位置"]},
                {"name": "实质性要求", "checks": ["星号条款", "不可偏离项", "交付验收"]},
            ],
            [_table("资格审查响应表", ["序号", "资格要求", "响应说明", "证明材料", "页码"]),
             _table("技术偏离表", ["序号", "招标要求", "投标响应", "偏离情况", "证明材料"]),
             _table("商务偏离表", ["序号", "商务条款", "投标响应", "偏离情况", "说明"])],
            [_slot("主体与授权材料", ["营业执照", "法定代表人身份证明", "授权委托书"], True),
             _slot("资格与政策证明", ["信用记录", "中小企业声明", "税收社保证明"]),
             _slot("业绩与案例", ["合同关键页", "验收或中标证明"]),
             _slot("报价与承诺", ["报价表", "服务承诺", "投标有效期"], True)],
            COMMON_QUALITY_RULES, [SOURCE_GOVERNMENT_87, SOURCE_GOVERNMENT_DEMAND],
        ),
        "builtin": True,
    },
    {
        "id": "government-it",
        "name": "政府采购·软件与信息化",
        "description": "适合软件平台、系统集成、数字化建设和信息安全项目",
        "prompt": "以评分点为主线展开总体架构、功能参数、实施迁移、安全合规和运维服务。",
        "settings": {"mode": "standard", "quality_gate": True, "include_toc": True, "include_deviation_tables": True},
        "package": _package(
            "government_procurement", "software_it", ["政府采购", "软件", "信息化"],
            ["软件", "信息化", "平台建设", "系统集成", "网络安全", "数据安全", "信创", "运维"],
            [
                _outline("资格与符合性响应", "完成资格、政策和实质性条款逐项响应"),
                _outline("项目理解与总体方案", "说明建设目标、范围、用户和总体路线", ["采购需求", "现状说明"]),
                _outline("总体架构与技术路线", "给出业务、应用、数据、技术和部署架构", ["架构要求", "技术标准"]),
                _outline("功能与技术参数响应", "逐项对应功能、性能、兼容和接口要求", ["技术参数表"]),
                _outline("实施、迁移与集成", "说明里程碑、数据迁移、接口联调和上线切换", ["工期要求", "系统清单"]),
                _outline("安全与合规", "覆盖数据、网络、权限、日志、备份和等保要求", ["安全条款"]),
                _outline("测试、验收与培训", "把测试用例、验收指标和培训交付物对应到要求", ["验收标准"]),
                _outline("运维与售后服务", "说明 SLA、响应时限、巡检、应急和持续服务", ["服务要求"]),
                _outline("团队、业绩与评分证据", "只引用素材库中可核验的人员、证书和案例", ["评分办法"]),
            ],
            [{"name": "功能参数", "checks": ["逐项参数", "偏离情况", "证明材料"]},
             {"name": "实施可行性", "checks": ["里程碑", "迁移", "联调", "风险"]},
             {"name": "安全与服务", "checks": ["安全控制", "SLA", "应急", "验收"]}],
            [_table("功能技术响应表", ["序号", "招标要求", "投标响应", "实现方式", "证明材料", "偏离情况"]),
             _table("实施里程碑计划", ["阶段", "工作内容", "交付物", "责任角色", "完成时间"]),
             _table("SLA响应表", ["服务事项", "服务等级", "响应时间", "恢复时间", "考核方式"])],
            [_slot("软件产品与技术材料", ["产品白皮书", "功能截图", "检测或兼容证明"]),
             _slot("信息化项目案例", ["合同关键页", "验收证明"]),
             _slot("项目团队", ["人员简历", "证书", "社保证明"]),
             _slot("安全与合规材料", ["等保", "安全认证", "数据处理说明"])],
            COMMON_QUALITY_RULES + ["功能参数不得只写“满足”，必须给出实现方式或证据位置"],
            [SOURCE_GOVERNMENT_87, SOURCE_GOVERNMENT_DEMAND],
        ),
        "builtin": True,
    },
    {
        "id": "goods",
        "name": "设备与货物采购",
        "description": "适合设备、材料、成品货物及配套安装采购",
        "prompt": "逐项核对核心产品、技术规格、供货安装、验收、质保和备品备件要求。",
        "settings": {"mode": "standard", "quality_gate": True, "include_toc": True, "include_deviation_tables": True},
        "package": _package(
            "goods_procurement", "equipment_material", ["设备", "货物", "材料"],
            ["设备采购", "货物采购", "材料采购", "供货", "安装调试", "技术参数", "质保", "备品备件"],
            [
                _outline("资格与商务文件", "完成资格、授权、报价与商务条款响应"),
                _outline("货物清单与分项报价", "保持品目、数量、单位和报价口径一致", ["采购清单"]),
                _outline("技术规格逐项响应", "逐参数给出品牌型号、数值和证明材料", ["技术规格"]),
                _outline("供货与进度计划", "说明生产备货、运输、到货和风险保障", ["交货期"]),
                _outline("安装调试与验收", "说明安装条件、调试步骤、测试和验收依据", ["验收标准"]),
                _outline("质量保证与售后", "说明质保期、维修响应、备件和培训", ["售后要求"]),
                _outline("制造商与产品证明", "整理授权、检测、认证和产品资料", ["证明材料要求"]),
            ],
            [{"name": "核心参数", "checks": ["品牌型号", "参数值", "证明页", "偏离"]},
             {"name": "履约能力", "checks": ["供货期", "安装", "验收", "质保"]}],
            [_table("货物分项报价表", ["序号", "货物名称", "品牌型号", "数量", "单价", "合价"]),
             _table("技术规格偏离表", ["序号", "招标参数", "投标参数", "偏离情况", "证明材料页码"]),
             _table("供货安装计划", ["阶段", "工作内容", "时间", "责任人", "交付物"])],
            [_slot("产品技术资料", ["彩页", "检测报告", "技术白皮书"]),
             _slot("制造商材料", ["制造商授权", "售后承诺"]),
             _slot("供货业绩", ["同类合同", "验收证明"]),
             _slot("报价信息", ["分项报价", "税率", "运保安装费用"], True)],
            COMMON_QUALITY_RULES + ["品牌、型号和技术参数必须来自用户确认或产品证明，严禁猜测"],
            [SOURCE_STANDARD_PURCHASE, SOURCE_GOVERNMENT_87],
        ),
        "builtin": True,
    },
    {
        "id": "construction",
        "name": "工程施工",
        "description": "适合施工总承包、专业工程、装修和改造项目",
        "prompt": "围绕施工组织设计、进度资源、质量安全、项目管理和工程量要求编制。",
        "settings": {"mode": "standard", "quality_gate": True, "include_toc": True, "include_deviation_tables": True},
        "package": _package(
            "engineering_construction", "construction_renovation", ["工程", "施工", "改造"],
            ["施工总承包", "工程施工", "改造工程", "施工组织设计", "工程量清单", "工期", "质量安全", "文明施工"],
            [
                _outline("资格、商务与投标函", "完成企业资质、人员资格、投标函和商务响应"),
                _outline("项目理解与施工部署", "说明工程特点、重难点、总体部署和施工分区", ["图纸", "工程量清单"]),
                _outline("施工组织设计与专项方案", "按专业和关键工序给出可执行施工方法", ["技术标准", "施工要求"]),
                _outline("进度计划与工期保障", "提供里程碑、横道或网络计划及纠偏措施", ["计划工期"]),
                _outline("资源配置与项目管理机构", "对应机械、材料、劳动力和管理人员配置", ["人员要求"]),
                _outline("质量管理与检验试验", "说明质量目标、控制点、检验批和成品保护", ["质量标准"]),
                _outline("安全文明与环境保护", "覆盖危大工程、消防、临电、环保和文明施工", ["安全要求"]),
                _outline("风险与应急预案", "识别工期、质量、安全、现场和供应风险并给出处置流程"),
                _outline("工程量、技术与商务响应", "逐项核对清单、技术标准、合同和偏离情况"),
            ],
            [{"name": "施工组织", "checks": ["重难点", "施工方法", "工序衔接", "现场布置"]},
             {"name": "进度资源", "checks": ["工期", "里程碑", "机械", "劳动力"]},
             {"name": "质量安全", "checks": ["质量控制", "安全措施", "文明施工", "应急"]}],
            [_table("主要施工方法响应表", ["分部分项", "招标要求", "施工方法", "质量控制点", "验收标准"]),
             _table("施工进度计划", ["阶段", "工作内容", "开始时间", "完成时间", "资源配置"]),
             _table("项目管理人员表", ["岗位", "姓名", "资格证书", "类似经验", "职责"]),
             _table("工程技术偏离表", ["序号", "标准或清单要求", "投标响应", "偏离情况", "说明"])],
            [_slot("企业施工资质", ["资质证书", "安全生产许可证"], True),
             _slot("项目管理人员", ["建造师证", "职称证", "安全证", "社保"]),
             _slot("类似工程业绩", ["合同", "中标通知", "验收证明"]),
             _slot("施工资源", ["机械设备", "劳动力计划", "材料供应计划"])],
            COMMON_QUALITY_RULES + ["工程类章节数量随招标目录和评分办法展开，不得套用IT服务类的少章节限制"],
            [SOURCE_STANDARD_CONSTRUCTION, SOURCE_STANDARD_PURCHASE],
        ),
        "builtin": True,
    },
    {
        "id": "service",
        "name": "服务类投标",
        "description": "适合运维、外包、运营、物业及持续性服务项目",
        "prompt": "围绕服务流程、组织团队、SLA、质量考核、应急响应和持续改进编制。",
        "settings": {"mode": "standard", "quality_gate": True, "include_toc": True, "include_deviation_tables": True},
        "package": _package(
            "service_delivery", "operations_outsourcing", ["服务", "运维", "运营", "外包"],
            ["运维服务", "运营服务", "外包服务", "物业服务", "SLA", "服务流程", "响应时间", "考核指标"],
            [
                _outline("资格与商务响应", "完成资格、报价、合同和服务承诺响应"),
                _outline("需求理解与服务目标", "把服务范围、对象、边界和目标说清楚", ["服务需求"]),
                _outline("总体服务方案与流程", "按受理、执行、反馈、闭环设计服务流程"),
                _outline("组织架构与人员配置", "说明岗位、班次、能力、替补和培训机制", ["人员要求"]),
                _outline("SLA与考核指标", "逐项响应服务等级、时限、质量和考核扣罚", ["考核办法"]),
                _outline("质量控制与持续改进", "说明检查、复盘、满意度和改进闭环"),
                _outline("风险、应急与连续性", "覆盖突发事件、人员缺口、系统故障和业务连续性"),
                _outline("案例、团队与评分证据", "按评分项引用可核验业绩和人员材料", ["评分办法"]),
            ],
            [{"name": "服务方案", "checks": ["流程", "覆盖范围", "交付物", "可执行性"]},
             {"name": "团队能力", "checks": ["岗位", "人数", "证书", "经验", "替补"]},
             {"name": "服务保障", "checks": ["SLA", "质量", "应急", "持续改进"]}],
            [_table("SLA响应表", ["服务事项", "指标", "目标值", "响应时限", "考核方式"]),
             _table("人员配置表", ["岗位", "人数", "资格要求", "职责", "到岗安排"]),
             _table("服务风险与应急表", ["风险事件", "触发条件", "响应动作", "责任角色", "恢复目标"])],
            [_slot("服务团队", ["人员简历", "资格证书", "社保或劳动关系"]),
             _slot("同类服务业绩", ["服务合同", "验收或评价证明"]),
             _slot("服务工具与制度", ["服务流程", "工单工具", "质量制度"]),
             _slot("报价与服务承诺", ["费用明细", "SLA承诺", "人员投入承诺"], True)],
            COMMON_QUALITY_RULES + ["服务方案必须给出角色、动作、输入、输出和时限，不能只写原则"],
            [SOURCE_GOVERNMENT_87, SOURCE_GOVERNMENT_DEMAND],
        ),
        "builtin": True,
    },
    {
        "id": "consulting",
        "name": "咨询、规划与研究服务",
        "description": "适合咨询、规划、评估、设计和研究类项目",
        "prompt": "围绕问题理解、方法论、调研分析、里程碑、成果体系和专家团队编制。",
        "settings": {"mode": "standard", "quality_gate": True, "include_toc": True, "include_deviation_tables": True},
        "package": _package(
            "professional_service", "consulting_research", ["咨询", "规划", "研究", "设计"],
            ["咨询服务", "规划编制", "研究课题", "评估服务", "设计服务", "调研", "方法论", "成果报告"],
            [
                _outline("资格与商务响应", "完成资格、报价、合同和承诺响应"),
                _outline("项目背景与问题理解", "说明业务背景、核心问题、利益相关方和目标", ["采购需求"]),
                _outline("工作方法与技术路线", "把方法、步骤、数据、工具和判断标准串联起来"),
                _outline("调研、分析与论证计划", "说明样本、访谈、数据来源、分析和验证方法"),
                _outline("进度计划与里程碑", "把阶段任务、交付物、评审和修改闭环对应到工期"),
                _outline("项目团队与专家机制", "说明角色分工、投入、专家审查和替补机制"),
                _outline("质量保证与成果验收", "说明内部评审、版本控制、成果标准和验收方法"),
                _outline("保密、数据与风险管理", "覆盖资料保密、数据使用、进度和成果风险"),
                _outline("类似业绩与评分证据", "按评分项引用合同、成果和评价证明", ["评分办法"]),
            ],
            [{"name": "方法论", "checks": ["问题", "方法", "数据", "验证", "成果"]},
             {"name": "项目组织", "checks": ["团队", "分工", "投入", "专家机制"]},
             {"name": "交付质量", "checks": ["里程碑", "成果清单", "评审", "验收"]}],
            [_table("工作任务与成果表", ["阶段", "工作任务", "方法", "交付成果", "验收标准"]),
             _table("进度里程碑表", ["里程碑", "时间", "输入", "输出", "评审人"]),
             _table("项目团队表", ["角色", "人员", "专业能力", "投入时间", "职责"])],
            [_slot("咨询团队", ["人员简历", "职称或资格", "项目经验"]),
             _slot("类似项目业绩", ["合同", "成果摘要", "客户评价或验收"]),
             _slot("方法与工具", ["方法论说明", "调研工具", "分析框架"]),
             _slot("成果与报价", ["成果清单", "费用构成", "知识产权和保密承诺"], True)],
            COMMON_QUALITY_RULES + ["方法论必须落到本项目的输入、动作、输出和验收，禁止只列通用框架名"],
            [SOURCE_GOVERNMENT_DEMAND, SOURCE_STANDARD_PURCHASE],
        ),
        "builtin": True,
    },
]


def builtin_templates() -> List[Dict[str, Any]]:
    # 前三项顺序是已发布界面和旧客户端的兼容合同；新增细分场景排在其后。
    order = {name: index for index, name in enumerate(
             ("government", "construction", "service", "government-it", "goods", "consulting"))}
    return copy.deepcopy(sorted(BUILTIN_TEMPLATES, key=lambda item: order.get(item["id"], 99)))


def _clean_text(value: Any, limit: int = 300) -> str:
    value = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", str(value or ""))
    return re.sub(r"\s+", " ", value).strip()[:limit]


def _clean_title(value: Any) -> str:
    title = _clean_text(value, 100)
    title = re.sub(r"^\s*(?:第[一二三四五六七八九十百]+[章节篇部分]|[0-9一二三四五六七八九十]+[.、．)])\s*", "", title)
    if re.search(r"(?:报价|金额|客户名称|项目名称)[:：]", title) or re.search(r"\d{6,}", title):
        return ""
    return title


def _bounded_int(value: Any, upper: int) -> int:
    try:
        return max(0, min(int(value or 0), upper))
    except (TypeError, ValueError):
        return 0


def _unique(items: Iterable[str], limit: int = 30) -> List[str]:
    out: List[str] = []
    for item in items:
        cleaned = _clean_text(item, 100)
        if cleaned and cleaned not in out:
            out.append(cleaned)
        if len(out) >= limit:
            break
    return out


def _safe_table_headers(items: Iterable[str]) -> List[str]:
    """只保留看起来像字段名的首行，避免把历史项目数据误当成模板表头。"""
    raw_items = [_clean_text(item, 100) for item in items]
    cleaned_items = [_clean_title(item) for item in raw_items]
    if any(raw and not cleaned for raw, cleaned in zip(raw_items, cleaned_items)):
        return []
    if len(cleaned_items) < 2:
        return []
    header_words = (
        "序号", "名称", "内容", "要求", "响应", "参数", "指标", "单位", "数量", "单价", "总价", "备注",
        "证明", "人员", "岗位", "职责", "偏离", "分值", "评分", "材料", "证书", "时间", "阶段", "交付",
        "风险", "措施", "条款", "页码", "来源", "成果", "服务", "工作", "说明",
    )
    header_like = sum(any(word in cell for word in header_words) for cell in cleaned_items)
    if header_like < max(1, (len(cleaned_items) + 1) // 2):
        return []
    mapped: List[str] = []
    column_rules = [
        (("序号", "编号"), "序号"), (("名称", "对象"), "名称"), (("要求", "标准", "条款"), "要求或标准"),
        (("响应", "说明"), "响应说明"), (("参数", "指标"), "参数或指标"), (("数量",), "数量"),
        (("单位",), "单位"), (("单价",), "单价"), (("总价", "合价", "金额"), "总价"),
        (("人员", "岗位", "角色"), "人员或岗位"), (("职责",), "职责"), (("偏离",), "偏离情况"),
        (("分值", "评分"), "评分信息"), (("材料", "证明", "证书"), "证明材料"),
        (("时间", "日期", "阶段"), "时间或阶段"), (("交付", "成果"), "交付成果"),
        (("风险",), "风险"), (("措施", "动作"), "措施"), (("备注",), "备注"),
        (("服务", "工作", "内容"), "工作内容"), (("页码",), "页码"), (("来源",), "来源"),
    ]
    for index, cell in enumerate(cleaned_items, 1):
        generic = next((label for words, label in column_rules if any(word in cell for word in words)), "列%d" % index)
        if generic in mapped:
            generic = "%s%d" % (generic, index)
        mapped.append(generic)
    return mapped


def _guard_docx_archive(blob: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(blob)) as archive:
            entries = archive.infolist()
            total_size = sum(max(0, item.file_size) for item in entries)
            total_compressed = sum(max(0, item.compress_size) for item in entries)
    except (zipfile.BadZipFile, OSError) as exc:
        raise ValueError("Word文件无法读取，请确认文件未损坏") from exc
    if len(entries) > MAX_DOCX_ENTRIES or total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
        raise ValueError("Word文件展开后过大，请精简图片或拆分后重试")
    if total_compressed and total_size / total_compressed > MAX_DOCX_COMPRESSION_RATIO:
        raise ValueError("Word文件压缩比例异常，请另存为新文件后重试")


def _generic_uploaded_heading(title: str, index: int) -> str:
    """历史标书标题只保留章节语义，不把单位、项目或人员名称带入模板。"""
    rules = [
        (("投标函", "授权"), "投标函与授权文件"),
        (("资格", "符合性"), "资格与符合性响应"),
        (("偏离",), "偏离响应"),
        (("报价", "商务", "合同"), "报价与商务响应"),
        (("理解", "需求"), "项目理解与需求分析"),
        (("架构", "技术路线", "总体方案"), "总体方案与技术路线"),
        (("功能", "参数", "技术响应", "技术方案"), "功能与技术响应"),
        (("施工",), "施工组织与实施"),
        (("进度", "工期", "里程碑"), "实施进度与工期保障"),
        (("质量", "验收", "测试"), "质量保证与验收"),
        (("安全", "保密", "数据"), "安全与保密"),
        (("服务", "售后", "运维", "SLA"), "服务与售后保障"),
        (("团队", "人员", "组织", "项目经理"), "团队与组织保障"),
        (("风险", "应急"), "风险与应急"),
        (("案例", "业绩"), "案例与业绩证明"),
    ]
    return next((label for words, label in rules if any(word in title for word in words)), "自定义章节%d" % index)


def normalize_package(package: Any) -> Dict[str, Any]:
    raw = package if isinstance(package, dict) else {}
    scene = raw.get("scene") if isinstance(raw.get("scene"), dict) else {}
    formatting = raw.get("formatting") if isinstance(raw.get("formatting"), dict) else {}
    normalized = {
        "schema_version": 1,
        "priority_rule": PRIORITY_RULE,
        "scene": {
            "category": _clean_text(scene.get("category"), 80) or "custom",
            "subtype": _clean_text(scene.get("subtype"), 80) or "custom",
            "labels": _unique(scene.get("labels") or [], 12),
            "match_keywords": _unique(scene.get("match_keywords") or [], 30),
        },
        "outline": [],
        "scoring_focus": [],
        "tables": [],
        "material_slots": [],
        "formatting": {
            key: bool(formatting.get(key)) for key in (
                "include_cover", "include_toc", "repeat_table_headers", "wide_tables_landscape",
                "keep_source_heading_order_when_required",
            )
        },
        "quality_rules": _unique(raw.get("quality_rules") or [], 30),
        "sources": [],
        "generated_from": dict(raw.get("generated_from")) if isinstance(raw.get("generated_from"), dict) else {"kind": "custom"},
    }
    for item in raw.get("outline") or []:
        if not isinstance(item, dict):
            continue
        title = _clean_title(item.get("title"))
        if title:
            normalized["outline"].append({
                "title": title,
                "purpose": _clean_text(item.get("purpose"), 300),
                "required": bool(item.get("required", True)),
                "evidence": _unique(item.get("evidence") or [], 12),
            })
        if len(normalized["outline"]) >= 40:
            break
    for item in raw.get("scoring_focus") or []:
        if isinstance(item, dict) and _clean_title(item.get("name")):
            normalized["scoring_focus"].append({
                "name": _clean_title(item.get("name")),
                "checks": _unique(item.get("checks") or [], 16),
            })
        if len(normalized["scoring_focus"]) >= 20:
            break
    for item in raw.get("tables") or []:
        if isinstance(item, dict) and _clean_title(item.get("name")):
            normalized["tables"].append({
                "name": _clean_title(item.get("name")),
                "columns": _unique(item.get("columns") or [], 16),
                "required": bool(item.get("required", True)),
            })
        if len(normalized["tables"]) >= 20:
            break
    for item in raw.get("material_slots") or []:
        if isinstance(item, dict) and _clean_title(item.get("name")):
            normalized["material_slots"].append({
                "name": _clean_title(item.get("name")),
                "evidence": _unique(item.get("evidence") or [], 16),
                "required": bool(item.get("required", False)),
                "missing_policy": "缺失时标记〔需补充〕，不得猜测或编造",
            })
        if len(normalized["material_slots"]) >= 20:
            break
    for item in raw.get("sources") or []:
        if not isinstance(item, dict):
            continue
        title, issuer, url = _clean_text(item.get("title"), 200), _clean_text(item.get("issuer"), 120), _clean_text(item.get("url"), 500)
        if title and issuer and (url.startswith("https://") or url.startswith("user-upload://")):
            normalized["sources"].append({"title": title, "issuer": issuer, "version": _clean_text(item.get("version"), 80), "url": url})
        if len(normalized["sources"]) >= 20:
            break
    generated_raw = normalized["generated_from"]
    normalized["generated_from"] = {
        "kind": _clean_text(generated_raw.get("kind"), 40) or "custom",
        **({"source_name": _clean_text(generated_raw.get("source_name"), 120)} if generated_raw.get("source_name") else {}),
        **({"source_sha256": _clean_text(generated_raw.get("source_sha256"), 64)} if generated_raw.get("source_sha256") else {}),
    }
    if normalized["generated_from"]["kind"] == "uploaded_bid":
        normalized["generated_from"].update({
            "source_structure_ready": bool(generated_raw.get("source_structure_ready")),
            "heading_count": _bounded_int(generated_raw.get("heading_count"), 40),
            "table_count": _bounded_int(generated_raw.get("table_count"), 12),
        })
    return normalized


def validate_package(package: Any) -> Dict[str, Any]:
    value = normalize_package(package)
    checks = {
        "scene": bool(value["scene"]["category"] and value["scene"]["match_keywords"]),
        "outline": len(value["outline"]) >= 5,
        "scoring_focus": bool(value["scoring_focus"]),
        "tables": bool(value["tables"]),
        "material_slots": bool(value["material_slots"]),
        "quality_rules": len(value["quality_rules"]) >= 4,
        "sources": bool(value["sources"]),
    }
    score = round(sum(1 for ok in checks.values() if ok) / len(checks) * 100)
    return {"score": score, "checks": checks, "ready": all(checks.values()), "warnings": [name for name, ok in checks.items() if not ok]}


def extract_document_structure(filename: str, blob: bytes) -> Tuple[List[str], List[List[str]], str]:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    headings: List[str] = []
    tables: List[List[str]] = []
    body_parts: List[str] = []
    if ext == "docx":
        from docx import Document

        _guard_docx_archive(blob)
        document = Document(io.BytesIO(blob))
        for paragraph in document.paragraphs:
            text = _clean_text(paragraph.text, 500)
            if not text:
                continue
            body_parts.append(text)
            style = _clean_text(getattr(getattr(paragraph, "style", None), "name", ""), 80).lower()
            numbered = bool(re.match(r"^(?:第[一二三四五六七八九十百]+[章节篇]|\d+(?:\.\d+){0,3}[、.．\s])", text))
            if "heading" in style or "标题" in style or numbered:
                title = _clean_title(text)
                if title:
                    headings.append(title)
        table_text_chars = 0
        for table in document.tables[:12]:
            if not table.rows:
                continue
            headers = _safe_table_headers(cell.text for cell in table.rows[0].cells)
            if len(headers) >= 2:
                tables.append(headers)
            for row in table.rows[:20]:
                for cell in row.cells[:20]:
                    text = _clean_text(cell.text, 200)
                    if not text:
                        continue
                    remaining = 30000 - table_text_chars
                    if remaining <= 0:
                        break
                    body_parts.append(text[:remaining])
                    table_text_chars += min(len(text), remaining)
                if table_text_chars >= 30000:
                    break
            if table_text_chars >= 30000:
                break
    elif ext in ("md", "markdown", "txt"):
        text = blob.decode("utf-8", errors="ignore")
        body_parts.append(text[:100000])
        for line in text.splitlines():
            stripped = line.strip()
            match = re.match(r"^#{1,6}\s+(.+)$", stripped)
            if match:
                title = _clean_title(match.group(1))
            elif re.match(r"^(?:第[一二三四五六七八九十百]+[章节篇]|\d+(?:\.\d+){0,3}[、.．\s])", stripped):
                title = _clean_title(stripped)
            else:
                title = ""
            if title:
                headings.append(title)
        for line in text.splitlines():
            if line.count("|") >= 3 and "---" not in line:
                cells = _safe_table_headers(cell.strip() for cell in line.strip("|").split("|"))
                if len(cells) >= 2:
                    tables.append(cells)
                    break
    elif ext == "pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(blob))
            text = "\n".join((page.extract_text() or "") for page in reader.pages[:200])
            body_parts.append(text[:100000])
            for line in text.splitlines():
                if re.match(r"^(?:第[一二三四五六七八九十百]+[章节篇]|\d+(?:\.\d+){0,3}[、.．\s])", line.strip()):
                    title = _clean_title(line)
                    if title:
                        headings.append(title)
        except Exception as exc:
            raise ValueError("PDF未能提取文字；请先做OCR或另存为Word后再生成模板") from exc
    else:
        raise ValueError("暂支持 .docx、.pdf、.md 和 .txt")
    return _unique(headings, 40), tables[:12], "\n".join(body_parts)[:120000]


def recommend_template(filename: str, text: str, templates: Sequence[Dict[str, Any]]) -> Dict[str, Any]:
    haystack = (filename + "\n" + text).lower()
    ranked: List[Tuple[float, Dict[str, Any], List[str]]] = []
    for template in templates:
        package = normalize_package(template.get("package"))
        keywords = package["scene"].get("match_keywords") or []
        hits = [word for word in keywords if word.lower() in haystack]
        score = sum(max(1.0, min(len(word), 8) / 2.0) for word in hits)
        if template.get("id") == "government" and "政府采购" in haystack:
            score += 1.0
        ranked.append((score, template, hits))
    ranked.sort(key=lambda row: row[0], reverse=True)
    score, selected, hits = ranked[0] if ranked else (0.0, {"id": "government", "name": "政府采购·综合"}, [])
    if score == 0:
        selected = next((item for item in templates if item.get("id") == "government"), selected)
        hits = []
    confidence = round(min(0.98, 0.25 + score / (score + 12.0) * 0.73), 2) if score else 0.25
    return {
        "template_id": selected.get("id", "government"),
        "template_name": selected.get("name", "政府采购·综合"),
        "confidence": confidence,
        "reasons": hits[:6] or ["未识别到强场景特征，使用综合模板并以招标原文为准"],
    }


def _purpose_for_title(title: str) -> str:
    rules = [
        (("资格", "符合性"), "逐项对应资格、符合性和实质性要求，并建立证明材料索引"),
        (("技术", "方案", "实施"), "围绕本章对应的采购要求和评分点给出可执行方案与证据"),
        (("进度", "工期", "里程碑"), "明确阶段、时间、责任角色、交付物和纠偏措施"),
        (("质量", "售后", "服务"), "明确质量目标、服务动作、时限、检查和持续改进"),
        (("安全", "应急", "风险"), "识别风险、触发条件、责任角色、处置动作和恢复目标"),
        (("团队", "人员", "机构"), "按岗位说明人员能力、投入、职责和可核验证明"),
        (("报价", "商务", "合同"), "逐条响应报价口径、商务条件和合同条款，保留人工确认"),
    ]
    for keys, purpose in rules:
        if any(key in title for key in keys):
            return purpose
    return "按招标文件对应要求组织本章，明确响应、证据、缺口和人工确认项"


def derive_template(name: str, filename: str, blob: bytes, templates: Sequence[Dict[str, Any]], scene_hint: str = "") -> Dict[str, Any]:
    headings, table_headers, body = extract_document_structure(filename, blob)
    if not headings and not table_headers:
        raise ValueError("未识别到可复用目录或表格结构；请上传带标题样式或规范表头的标书")
    recommendation = recommend_template(filename, scene_hint + "\n" + body, templates)
    base = next((item for item in templates if item.get("id") == recommendation["template_id"]), templates[0])
    package = normalize_package(copy.deepcopy(base.get("package")))
    source_headings = [title for title in headings if title and not re.search(r"(?:目录|投标文件$|封面)", title)]
    safe_headings: List[str] = []
    for index, title in enumerate(source_headings, 1):
        generic = _generic_uploaded_heading(title, index)
        if generic not in safe_headings:
            safe_headings.append(generic)
    enough_structure = len(safe_headings) >= 4 or (len(safe_headings) >= 2 and bool(table_headers))
    if enough_structure:
        derived_outline = [
            {"title": title, "purpose": _purpose_for_title(title), "required": True, "evidence": []}
            for title in safe_headings[:30]
        ]
        for item in package["outline"]:
            if len(derived_outline) >= 5:
                break
            if item.get("title") not in {entry["title"] for entry in derived_outline}:
                derived_outline.append(item)
        package["outline"] = derived_outline
    if table_headers:
        derived_tables = []
        for index, headers in enumerate(table_headers[:8], 1):
            generic_name = "历史标书表格结构%d" % index
            if any("偏离" in cell for cell in headers):
                generic_name = "偏离响应表"
            elif any("报价" in cell or "单价" in cell for cell in headers):
                generic_name = "报价明细表"
            elif any("人员" in cell or "岗位" in cell for cell in headers):
                generic_name = "人员配置表"
            derived_tables.append({"name": generic_name, "columns": headers, "required": False})
        existing = {item["name"] for item in package["tables"]}
        package["tables"].extend(item for item in derived_tables if item["name"] not in existing)
    source_ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "document"
    package["generated_from"] = {
        "kind": "uploaded_bid",
        "source_name": "用户上传的优秀历史标书.%s" % source_ext,
        "source_sha256": hashlib.sha256(blob).hexdigest(),
        "source_structure_ready": enough_structure,
        "heading_count": len(safe_headings),
        "table_count": len(table_headers),
    }
    package["sources"].append({
        "title": "用户上传的优秀历史标书（仅提取结构，不复制正文事实）",
        "issuer": "用户上传",
        "version": "待用户确认",
        "url": "user-upload://" + hashlib.sha256(blob).hexdigest()[:16],
    })
    package = normalize_package(package)
    validation = validate_package(package)
    validation["checks"]["source_structure"] = enough_structure
    validation["score"] = round(sum(1 for ok in validation["checks"].values() if ok) / len(validation["checks"]) * 100)
    validation["ready"] = all(validation["checks"].values())
    validation["warnings"] = [key for key, ok in validation["checks"].items() if not ok]
    extraction = {
        "heading_count": len(safe_headings),
        "table_count": len(table_headers),
        "text_chars": len(body),
        "source_structure_ready": enough_structure,
    }
    template = {
        "name": _clean_text(name, 80) or (_clean_text(base.get("name"), 60) + "·历史结构模板"),
        "description": "由用户上传的优秀历史标书提炼；仅保留目录、表格结构和场景规则，保存前需人工确认",
        "prompt": "按已确认的场景模板包组织投标文件；历史标书只用于结构参考，不得复制项目专属事实。",
        "settings": dict(base.get("settings") or {}),
        "package": package,
        "requires_review": True,
        "validation": validation,
        "extraction": extraction,
        "recommendation": recommendation,
    }
    return template


def compile_template_instructions(template: Dict[str, Any]) -> str:
    package = normalize_package(template.get("package"))
    lines = [
        "# 场景模板包：%s" % _clean_text(template.get("name"), 100),
        "",
        "## 使用优先级",
        "- %s。" % PRIORITY_RULE,
        "- 模板只规定组织、响应和检查方法；不得覆盖招标文件，也不得补造用户未提供的事实。",
    ]
    if template.get("prompt"):
        lines += ["", "## 场景目标", "- " + _clean_text(template.get("prompt"), 1200)]
    lines += ["", "## 质检规则"]
    lines.extend("- " + rule for rule in package["quality_rules"])
    format_rules = [
        ("include_cover", "生成封面，项目名称、投标人和日期等未确认字段保留待确认"),
        ("include_toc", "生成可更新目录，并在定稿前刷新标题层级与页码"),
        ("repeat_table_headers", "跨页表格重复表头，避免续页字段含义丢失"),
        ("wide_tables_landscape", "宽表必要时使用横向页面，保证字段完整可读"),
        ("keep_source_heading_order_when_required", "招标文件指定目录时保持原顺序，模板目录只作补充"),
    ]
    enabled_format_rules = [text for key, text in format_rules if package["formatting"].get(key)]
    if enabled_format_rules:
        lines += ["", "## Word格式"]
        lines.extend("- " + rule for rule in enabled_format_rules)
    lines += ["", "## 建议目录与写作要求"]
    for index, item in enumerate(package["outline"], 1):
        evidence = "；证据：" + "、".join(item["evidence"]) if item.get("evidence") else ""
        lines.append("%d. %s：%s%s" % (index, item["title"], item.get("purpose") or "按招标原文响应", evidence))
    lines += ["", "## 评分响应"]
    for item in package["scoring_focus"]:
        lines.append("- %s：%s" % (item["name"], "、".join(item["checks"])))
    lines += ["", "## 必备表格"]
    for item in package["tables"]:
        lines.append("- %s：%s" % (item["name"], "｜".join(item["columns"])))
    lines += ["", "## 材料槽位"]
    for item in package["material_slots"]:
        lines.append("- %s：%s；%s" % (item["name"], "、".join(item["evidence"]), item["missing_policy"]))
    return "\n".join(lines)[:20000]


def compare_instruction_coverage(text: str, template: Dict[str, Any]) -> Dict[str, Any]:
    """离线 A/B 指标：模板关键概念在任务指令中的显式覆盖率。"""
    package = normalize_package(template.get("package"))
    expected = _unique(
        [item["title"] for item in package["outline"]]
        + [item["name"] for item in package["scoring_focus"]]
        + [item["name"] for item in package["tables"]]
        + [item["name"] for item in package["material_slots"]],
        80,
    )
    hit = [item for item in expected if item in text]
    return {"expected": len(expected), "hit": len(hit), "coverage": round(len(hit) / len(expected), 3) if expected else 0.0, "missing": [item for item in expected if item not in hit]}
