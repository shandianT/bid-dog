"""Durable, evidence-based generation pipeline state.

This module intentionally has no FastAPI or OpenCode dependency.  It owns the
small on-disk contract that lets the engine resume one bounded node without
replaying already completed tender work.
"""

from __future__ import annotations

import datetime
import hashlib
import json
import os
import re
import shutil
import threading
import time
import unicodedata
import urllib.parse
import zipfile
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional


PIPELINE_VERSION = 2
PIPELINE_FILE = "pipeline.json"
MAX_ATTEMPTS = 3
SOURCE_MIN_CHARS = 80
SOURCE_MAX_CHARS = 2_000_000
SOURCE_MAX_FILE_BYTES = 100 * 1024 * 1024
DOCX_MAX_ENTRIES = 5_000
DOCX_MAX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024
DOCX_MAX_COMPRESSION_RATIO = 200
PDF_MAX_PAGES = 500
PDF_MAX_PAGE_CHARS = 40_000
VALID_STATES = {
    "pending",
    "running",
    "done",
    "retry_wait",
    "failed",
    "blocked",
    "skipped",
}


class PipelineError(RuntimeError):
    pass


class OutputValidationError(PipelineError):
    pass


class SourceParseError(PipelineError):
    pass


class NodeExecutionError(PipelineError):
    def __init__(self, code: str, *, retryable: bool, detail: str = ""):
        super().__init__(detail or code)
        self.code = str(code or "node_failed")[:120]
        self.retryable = bool(retryable)
        self.detail = str(detail or code)


_LOCKS: Dict[str, threading.RLock] = {}
_LOCKS_GUARD = threading.Lock()


def _now() -> str:
    return datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _lock(path: Path) -> threading.RLock:
    key = str(path.resolve())
    with _LOCKS_GUARD:
        return _LOCKS.setdefault(key, threading.RLock())


def _pipeline_path(job: os.PathLike[str] | str) -> Path:
    return Path(job) / PIPELINE_FILE


def _read(path: Path) -> Dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except (OSError, ValueError):
        return {}


def _write(path: Path, state: Dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_name(path.name + ".tmp")
    tmp.write_text(json.dumps(state, ensure_ascii=False, indent=1), encoding="utf-8")
    os.replace(tmp, path)


def _write_text(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_name(path.name + ".tmp")
    tmp.write_text(value, encoding="utf-8")
    os.replace(tmp, path)


def load(job: os.PathLike[str] | str) -> Dict[str, Any]:
    return _read(_pipeline_path(job))


def credential_fingerprint(secret: str) -> str:
    return hashlib.sha256(str(secret or "").encode("utf-8")).hexdigest()[:16]


def safe_base_url(value: str) -> str:
    """Persist a routable URL without userinfo, query credentials, or fragments."""
    try:
        parts = urllib.parse.urlsplit(str(value or "").strip())
        if parts.scheme not in ("http", "https") or not parts.hostname:
            return ""
        host = parts.hostname
        if ":" in host and not host.startswith("["):
            host = "[%s]" % host
        netloc = host + ((":" + str(parts.port)) if parts.port else "")
        return urllib.parse.urlunsplit((parts.scheme, netloc, parts.path.rstrip("/"), "", ""))
    except (TypeError, ValueError):
        return ""


def file_digest(paths: Iterable[os.PathLike[str] | str]) -> str:
    digest = hashlib.sha256()
    for value in sorted(str(Path(item)) for item in paths):
        path = Path(value)
        digest.update(path.name.encode("utf-8"))
        if path.is_file():
            with path.open("rb") as source:
                for chunk in iter(lambda: source.read(1024 * 1024), b""):
                    digest.update(chunk)
        else:
            digest.update(b"<missing>")
    return digest.hexdigest()


def _guard_docx_archive(path: Path) -> None:
    try:
        with zipfile.ZipFile(path) as archive:
            entries = archive.infolist()
            if len(entries) > DOCX_MAX_ENTRIES:
                raise SourceParseError("DOCX 文件条目过多，已停止解析")
            total_size = sum(max(0, int(item.file_size)) for item in entries)
            total_compressed = sum(max(1, int(item.compress_size)) for item in entries)
            if total_size > DOCX_MAX_UNCOMPRESSED_BYTES:
                raise SourceParseError("DOCX 解压后体积过大，已停止解析")
            if total_size and total_size / total_compressed > DOCX_MAX_COMPRESSION_RATIO:
                raise SourceParseError("DOCX 压缩比异常，已停止解析")
    except SourceParseError:
        raise
    except (OSError, zipfile.BadZipFile) as exc:
        raise SourceParseError("DOCX 文件损坏或格式不正确") from exc


def _clean_source_text(value: str) -> str:
    value = str(value or "").replace("\x00", "")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{4,}", "\n\n\n", value)
    return value.strip()[:SOURCE_MAX_CHARS]


def _heading_count(value: str) -> int:
    pattern = re.compile(
        r"^(?:#{1,6}\s+|第[一二三四五六七八九十百零〇0-9]+[章节部分篇]\s*|[0-9]+(?:\.[0-9]+)*[、.．]\s*)",
        re.MULTILINE,
    )
    return len(pattern.findall(value or ""))


def _extract_docx(path: Path) -> Dict[str, Any]:
    _guard_docx_archive(path)
    try:
        from docx import Document

        document = Document(str(path))
    except SourceParseError:
        raise
    except Exception as exc:
        raise SourceParseError("DOCX 正文读取失败") from exc

    parts: List[str] = []
    paragraph_count = 0
    for paragraph in document.paragraphs:
        text = _clean_source_text(paragraph.text)
        if text:
            paragraph_count += 1
            parts.append(text)
    table_count = 0
    cell_count = 0
    for table in document.tables[:200]:
        table_count += 1
        rows: List[str] = []
        for row in table.rows[:500]:
            cells = [_clean_source_text(cell.text)[:1000] for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                cell_count += len(cells)
                rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))
    text = _clean_source_text("\n\n".join(parts))
    return {
        "text": text,
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "cell_count": cell_count,
        "page_count": 0,
    }


def _extract_pdf(path: Path) -> Dict[str, Any]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
    except Exception as exc:
        raise SourceParseError("PDF 文件损坏或缺少 pypdf 依赖") from exc
    if len(reader.pages) > PDF_MAX_PAGES:
        raise SourceParseError("PDF 页数超过本地解析上限")
    pages: List[str] = []
    remaining = SOURCE_MAX_CHARS
    for page in reader.pages:
        if remaining <= 0:
            break
        try:
            text = _clean_source_text(page.extract_text() or "")[:PDF_MAX_PAGE_CHARS]
            pages.append(text[:remaining])
            remaining -= len(pages[-1])
        except Exception:
            pages.append("")
    text = _clean_source_text("\n\n".join(page for page in pages if page))
    return {
        "text": text,
        "paragraph_count": len([line for line in text.splitlines() if line.strip()]),
        "table_count": 0,
        "cell_count": 0,
        "page_count": len(reader.pages),
    }


def _extract_text(path: Path) -> Dict[str, Any]:
    with path.open("rb") as source:
        data = source.read(SOURCE_MAX_FILE_BYTES + 1)
    if len(data) > SOURCE_MAX_FILE_BYTES:
        raise SourceParseError("文本文件超过本地解析上限")
    decoded = ""
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            decoded = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if not decoded:
        decoded = data.decode("utf-8", errors="ignore")
    text = _clean_source_text(decoded)
    return {
        "text": text,
        "paragraph_count": len([line for line in text.splitlines() if line.strip()]),
        "table_count": 0,
        "cell_count": 0,
        "page_count": 0,
    }


def parse_source(
    job: os.PathLike[str] | str,
    tender_name: str,
    *,
    cached_text: str = "",
    cached_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Extract a tender once and persist reusable Markdown plus a body-free manifest."""
    root = Path(job)
    safe_name = os.path.basename(str(tender_name or ""))
    source = root / safe_name
    if not safe_name or not source.is_file():
        raise SourceParseError("找不到招标文件")
    if source.stat().st_size > SOURCE_MAX_FILE_BYTES:
        raise SourceParseError("招标文件超过本地解析上限")
    extension = source.suffix.lower().lstrip(".")
    cached = _clean_source_text(cached_text)
    reused = bool((cached_metadata or {}).get("complete")) and len(cached) >= SOURCE_MIN_CHARS
    if reused:
        extracted: Dict[str, Any] = {
            "text": cached,
            "paragraph_count": int((cached_metadata or {}).get("paragraph_count") or 0),
            "table_count": int((cached_metadata or {}).get("table_count") or 0),
            "cell_count": int((cached_metadata or {}).get("cell_count") or 0),
            "page_count": int((cached_metadata or {}).get("page_count") or 0),
        }
    elif extension == "docx":
        extracted = _extract_docx(source)
    elif extension == "pdf":
        extracted = _extract_pdf(source)
    elif extension in {"txt", "md"}:
        extracted = _extract_text(source)
    else:
        raise SourceParseError("暂不支持本地解析该文件格式：.%s" % (extension or "未知"))

    body = _clean_source_text(extracted.get("text") or "")
    if len(body) < SOURCE_MIN_CHARS:
        suffix = "；扫描版 PDF 请先 OCR" if extension == "pdf" else ""
        raise SourceParseError("未提取到足够正文%s" % suffix)
    heading_count = int((cached_metadata or {}).get("heading_count") or _heading_count(body))
    manifest = {
        "version": 1,
        "source_name": safe_name,
        "source_type": extension,
        "sha256": file_digest([source]),
        "char_count": len(body),
        "heading_count": heading_count,
        "paragraph_count": int(extracted.get("paragraph_count") or 0),
        "table_count": int(extracted.get("table_count") or 0),
        "cell_count": int(extracted.get("cell_count") or 0),
        "page_count": int(extracted.get("page_count") or 0),
        "reused_extract": reused,
        "complete": True,
        "parsed_at": _now(),
    }
    markdown = "# 招标文件解析版\n\n> 来源：%s\n\n%s\n" % (safe_name, body)
    _write_text(root / "招标文件_解析版.md", markdown)
    _write(root / "source_manifest.json", manifest)
    return manifest


def _node(
    node_id: str,
    kind: str,
    outputs: Iterable[str],
    *,
    model_tier: str = "",
    title: str = "",
    min_chars: int = 0,
    dependencies: Optional[Iterable[str]] = None,
) -> Dict[str, Any]:
    return {
        "id": node_id,
        "kind": kind,
        "state": "pending",
        "model_tier": model_tier,
        "title": title,
        "outputs": list(outputs),
        "min_chars": int(min_chars or 0),
        "attempt": 0,
        "attempt_serial": 0,
        "max_attempts": MAX_ATTEMPTS,
        "dependencies": list(dependencies or []),
        "input_digest": "",
        "started_at": "",
        "finished_at": "",
        "last_activity_at": "",
        "error_code": "",
        "retry_after_seconds": 0,
    }


def initialize(
    job: os.PathLike[str] | str,
    *,
    run_id: str,
    mode: str,
    model_routes: Dict[str, str],
    chapters: Iterable[Dict[str, Any]],
    credential_fingerprint: str = "",
    base_url: str = "",
) -> Dict[str, Any]:
    """Create a v1 pipeline without persisting a credential or source body."""
    mode = "standard" if str(mode).lower() == "standard" else "fast"
    routes = {
        "fast": str(model_routes.get("fast") or ""),
        "quality": str(model_routes.get("quality") or ""),
    }
    if credential_fingerprint:
        routes["credential_fingerprint"] = str(credential_fingerprint)
    if base_url:
        routes["base_url"] = safe_base_url(base_url)

    nodes: List[Dict[str, Any]] = [
        _node(
            "source_parse",
            "local",
            ["招标文件_解析版.md", "source_manifest.json"],
            title="本地解析招标文件",
            min_chars=80,
        ),
        _node(
            "response_plan",
            "model",
            ["投标文件组成.md", "评分点响应矩阵.md", "废标风险清单.md", "response_plan.json"],
            model_tier="fast",
            title="提取响应规划",
            min_chars=80,
        ),
    ]
    for chapter in chapters:
        chapter_id = re.sub(r"[^A-Za-z0-9_-]", "", str(chapter.get("id") or ""))
        output = os.path.basename(str(chapter.get("output") or ""))
        if not chapter_id or not output or output in {
            item for existing in nodes for item in existing.get("outputs", [])
        }:
            raise PipelineError("章节节点缺少唯一 id 或输出文件")
        nodes.append(
            _node(
                "chapter_write:" + chapter_id,
                "model",
                [output],
                model_tier="fast",
                title=str(chapter.get("title") or chapter_id)[:120],
                min_chars=120,
            )
        )
    for deviation_id, title, output in (
        ("technical_deviation", "技术应答偏离表", "技术应答偏离表.md"),
        ("business_deviation", "商务偏离表", "商务偏离表.md"),
    ):
        if output not in {item for existing in nodes for item in existing.get("outputs", [])}:
            nodes.append(
                _node(
                    "chapter_write:" + deviation_id,
                    "model",
                    [output],
                    model_tier="fast",
                    title=title,
                    min_chars=120,
                )
            )
    nodes.extend(
        [
            _node(
                "assemble",
                "local",
                ["投标文件_整册.md"],
                title="汇总整册",
                min_chars=200,
            ),
            _node(
                "quality_review",
                "model" if mode == "standard" else "local",
                # 标准模式的模型复核是独立证据；后续确定性质检仍会生成
                # 《成品质检报告.md》，两份报告不能互相覆盖。
                ["模型复核报告.md" if mode == "standard" else "成品质检报告.md"],
                model_tier="quality" if mode == "standard" else "",
                title="交付质检",
                min_chars=40,
            ),
            _node(
                "word_export",
                "local",
                ["投标文件_整册.docx", "Word格式自检报告.md"],
                title="导出 Word",
            ),
            _node(
                "delivery_gate",
                "local",
                [],
                title="交付门禁",
            ),
        ]
    )
    state = {
        "version": PIPELINE_VERSION,
        "run_id": str(run_id),
        "mode": mode,
        "state": "pending",
        "recoverable": True,
        "current_nodes": [],
        "model_routes": routes,
        "nodes": nodes,
        "created_at": _now(),
        "updated_at": _now(),
    }
    path = _pipeline_path(job)
    with _lock(path):
        _write(path, state)
    return state


def _find_node(state: Dict[str, Any], node_id: str) -> Dict[str, Any]:
    for node in state.get("nodes") or []:
        if node.get("id") == node_id:
            return node
    raise PipelineError("流水线节点不存在：%s" % node_id)


def dependency_outputs(state: Dict[str, Any], node_id: str) -> List[str]:
    """Return the transitive declared outputs that one node is allowed to read."""
    by_id = {str(node.get("id") or ""): node for node in state.get("nodes") or []}
    result: List[str] = []
    seen = set()

    def visit(current_id: str) -> None:
        if current_id in seen:
            return
        seen.add(current_id)
        current = by_id.get(current_id)
        if not current:
            raise PipelineError("节点依赖不存在：%s" % current_id)
        for dependency in current.get("dependencies") or []:
            visit(str(dependency))
        for output in current.get("outputs") or []:
            name = os.path.basename(str(output))
            if name and name not in result:
                result.append(name)

    node = by_id.get(str(node_id))
    if not node:
        raise PipelineError("流水线节点不存在：%s" % node_id)
    for dependency in node.get("dependencies") or []:
        visit(str(dependency))
    return result


def _reject_unsafe_unicode(value: Any) -> None:
    if isinstance(value, str):
        if any(0xD800 <= ord(char) <= 0xDFFF for char in value):
            raise OutputValidationError("响应规划包含无效 Unicode 字符")
        try:
            value.encode("utf-8")
        except UnicodeError as exc:
            raise OutputValidationError("响应规划包含无效 Unicode 字符") from exc
    elif isinstance(value, dict):
        for key, item in value.items():
            _reject_unsafe_unicode(key)
            _reject_unsafe_unicode(item)
    elif isinstance(value, (list, tuple)):
        for item in value:
            _reject_unsafe_unicode(item)


def _validated_plan_chapters(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    raw = data.get("chapters") if isinstance(data, dict) else None
    if not isinstance(raw, list) or not raw or len(raw) > 12:
        raise OutputValidationError("响应规划章节数量必须为 1 到 12")
    _reject_unsafe_unicode(raw)
    chapters: List[Dict[str, Any]] = []
    ids = set()
    outputs = set()
    reserved_outputs = {value.casefold() for value in {
        "招标文件_解析版.md", "投标文件组成.md", "评分点响应矩阵.md", "废标风险清单.md",
        "技术应答偏离表.md", "商务偏离表.md", "投标文件_整册.md", "成品质检报告.md",
    }}
    reserved_ids = {"technical_deviation", "business_deviation"}
    windows_devices = {"con", "prn", "aux", "nul"} | {
        "%s%d" % (prefix, number) for prefix in ("com", "lpt") for number in range(1, 10)
    } | {prefix + suffix for prefix in ("com", "lpt") for suffix in ("¹", "²", "³")}
    for item in raw:
        if not isinstance(item, dict):
            raise OutputValidationError("响应规划章节必须是对象")
        chapter_id = str(item.get("id") or "").strip()
        title = unicodedata.normalize("NFC", str(item.get("title") or "").strip())[:120]
        output = unicodedata.normalize("NFC", str(item.get("output") or "").strip())
        basis = item.get("basis")
        basis_values = ([unicodedata.normalize("NFC", str(value).strip())
                         for value in basis if str(value).strip()]
                        if isinstance(basis, list) else
                        [unicodedata.normalize("NFC", str(basis or "").strip())])
        folded_id = chapter_id.casefold()
        if (not re.fullmatch(r"[A-Za-z0-9_-]{1,40}", chapter_id)
                or folded_id in ids or folded_id in reserved_ids):
            raise OutputValidationError("响应规划章节 id 不安全或重复")
        if not title or not basis_values[0]:
            raise OutputValidationError("响应规划章节缺少标题或招标依据")
        folded_output = output.casefold()
        # Windows strips trailing spaces/dots from a path component before resolving
        # device names, so ``AUX .md`` is just as unsafe as ``AUX.md``.  Reject the
        # complete Win32 forbidden character/control set even when running on macOS.
        stem = output.rsplit(".", 1)[0].rstrip(" .").casefold()
        device_base = stem.split(".", 1)[0].rstrip(" .")
        utf8_bytes = len(output.encode("utf-8"))
        utf16_units = len(output.encode("utf-16-le")) // 2
        if (len(output) > 120 or utf8_bytes > 240 or utf16_units > 120
                or output != os.path.basename(output) or "/" in output or "\\" in output
                or any(ord(char) < 32 or char in '<>:"/\\|?*' for char in output)
                or output.endswith((" ", ".")) or device_base in windows_devices
                or not folded_output.endswith(".md") or folded_output in outputs
                or folded_output in reserved_outputs):
            raise OutputValidationError("响应规划章节输出文件不安全或重复")
        if not isinstance(item.get("dependencies"), list):
            raise OutputValidationError("响应规划章节 dependencies 必须是数组")
        if not isinstance(item.get("material_slots"), list):
            raise OutputValidationError("响应规划章节 material_slots 必须是数组")
        if not isinstance(item.get("scoring_points"), list):
            raise OutputValidationError("响应规划章节 scoring_points 必须是数组")
        dependencies = [str(value).strip() for value in item["dependencies"]]
        if any(not value for value in dependencies) or len(dependencies) != len(set(dependencies)):
            raise OutputValidationError("响应规划章节依赖不安全或重复")
        material_slots = [unicodedata.normalize("NFC", str(value).strip())[:120]
                          for value in item["material_slots"]
                          if str(value).strip()]
        scoring_points = [unicodedata.normalize("NFC", str(value).strip())[:240]
                          for value in item["scoring_points"]
                          if str(value).strip()]
        chapters.append({
            "id": chapter_id,
            "title": title,
            "output": output,
            "basis": basis_values[:20],
            "dependencies": dependencies,
            "material_slots": material_slots[:20],
            "scoring_points": scoring_points[:40],
        })
        ids.add(folded_id); outputs.add(folded_output)
    for chapter in chapters:
        if any(dep.casefold() not in ids or dep.casefold() == chapter["id"].casefold()
               for dep in chapter["dependencies"]):
            raise OutputValidationError("响应规划章节依赖不存在或自循环")
    canonical_ids = {chapter["id"].casefold(): chapter["id"] for chapter in chapters}
    for chapter in chapters:
        chapter["dependencies"] = [canonical_ids[dep.casefold()] for dep in chapter["dependencies"]]
    graph = {chapter["id"]: set(chapter["dependencies"]) for chapter in chapters}
    resolved = set()
    while len(resolved) < len(graph):
        ready = {node_id for node_id, deps in graph.items()
                 if node_id not in resolved and deps <= resolved}
        if not ready:
            raise OutputValidationError("响应规划章节依赖存在环")
        resolved |= ready
    return chapters


def apply_response_plan(job: os.PathLike[str] | str) -> Dict[str, Any]:
    """Atomically replace provisional template chapters with the validated plan DAG."""
    root = Path(job)
    data = _read(root / "response_plan.json")
    chapters = _validated_plan_chapters(data)

    def apply(state):
        response = _find_node(state, "response_plan")
        if response.get("state") != "done":
            raise PipelineError("响应规划尚未完成")
        old_chapters = {
            str(node.get("id") or ""): node
            for node in state.get("nodes") or []
            if str(node.get("id") or "").startswith("chapter_write:")
        }
        existing = [node for node in state.get("nodes") or []
                    if not str(node.get("id") or "").startswith("chapter_write:")]
        insert_at = next((index for index, node in enumerate(existing)
                          if node.get("id") == "assemble"), len(existing))
        dynamic = []
        for chapter in chapters:
            node_id = "chapter_write:" + chapter["id"]
            dependencies = ["chapter_write:" + dep for dep in chapter["dependencies"]]
            previous = old_chapters.get(node_id)
            if (previous and previous.get("outputs") == [chapter["output"]]
                    and previous.get("title") == chapter["title"]
                    and previous.get("dependencies") == dependencies
                    and previous.get("basis") == chapter["basis"]
                    and previous.get("material_slots") == chapter["material_slots"]
                    and previous.get("scoring_points") == chapter["scoring_points"]):
                dynamic.append(previous)
            else:
                planned_node = _node(
                    node_id, "model", [chapter["output"]], model_tier="fast",
                    title=chapter["title"], min_chars=120, dependencies=dependencies,
                )
                planned_node["basis"] = chapter["basis"]
                planned_node["material_slots"] = chapter["material_slots"]
                planned_node["scoring_points"] = chapter["scoring_points"]
                dynamic.append(planned_node)
        for deviation_id, title, output in (
            ("technical_deviation", "技术应答偏离表", "技术应答偏离表.md"),
            ("business_deviation", "商务偏离表", "商务偏离表.md"),
        ):
            node_id = "chapter_write:" + deviation_id
            previous = old_chapters.get(node_id)
            if (previous and previous.get("outputs") == [output]
                    and previous.get("title") == title):
                dynamic.append(previous)
            else:
                dynamic.append(_node(
                    node_id, "model", [output], model_tier="fast",
                    title=title, min_chars=120,
                ))
        state["nodes"] = existing[:insert_at] + dynamic + existing[insert_at:]
        state["plan_applied_at"] = _now()
        return state

    _, state = _mutate(root, apply)
    return state


def _output_text(path: Path) -> str:
    if path.suffix.lower() in (".md", ".txt", ".json"):
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            return ""
    return ""


def validate_outputs(job: os.PathLike[str] | str, node: Dict[str, Any]) -> None:
    root = Path(job)
    outputs = node.get("outputs") or []
    if not outputs:
        return
    for output in outputs:
        path = root / os.path.basename(str(output))
        if not path.is_file():
            raise OutputValidationError("输出文件不存在：%s" % path.name)
        if path.suffix.lower() == ".docx":
            if path.stat().st_size < 1024:
                raise OutputValidationError("Word 输出无效：%s" % path.name)
            try:
                with zipfile.ZipFile(path) as archive:
                    names = set(archive.namelist())
                    if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                        raise OutputValidationError("Word 输出结构无效：%s" % path.name)
                    document_xml = archive.read("word/document.xml")
                    visible = re.sub(rb"<[^>]+>", b"", document_xml)
                    if len(visible.strip()) < 80:
                        raise OutputValidationError("Word 正文内容不足：%s" % path.name)
            except zipfile.BadZipFile as exc:
                raise OutputValidationError("Word 输出不是有效 DOCX：%s" % path.name) from exc
            continue
        minimum = int(node.get("min_chars") or 0) if path.suffix.lower() in (".md", ".txt") else 0
        if minimum and len(_output_text(path).strip()) < minimum:
            raise OutputValidationError("输出内容不足：%s" % path.name)
    if node.get("id") == "response_plan":
        data = _read(root / "response_plan.json")
        _validated_plan_chapters(data)


def attempt_directory(job: os.PathLike[str] | str, node: Dict[str, Any]) -> Path:
    safe_id = re.sub(r"[^A-Za-z0-9_.-]", "_", str(node.get("id") or "node"))
    attempt = max(1, int(node.get("attempt_serial") or node.get("attempt") or 1))
    target = Path(job) / ".pipeline_attempts" / safe_id / ("attempt-%d" % attempt) / "work"
    target.mkdir(parents=True, exist_ok=True)
    return target


def promote_outputs(
    job: os.PathLike[str] | str,
    attempt_root: os.PathLike[str] | str,
    node: Dict[str, Any],
) -> None:
    """Validate one isolated attempt, then atomically promote declared files."""
    root = Path(job)
    source_root = Path(attempt_root)
    validate_outputs(source_root, node)
    for output in node.get("outputs") or []:
        name = os.path.basename(str(output))
        source = source_root / name
        target = root / name
        tmp = target.with_name(target.name + ".promote.tmp")
        shutil.copy2(source, tmp)
        os.replace(tmp, target)


def _archive_outputs(root: Path, node: Dict[str, Any]) -> None:
    existing = [root / os.path.basename(str(name)) for name in node.get("outputs") or []]
    existing = [path for path in existing if path.is_file()]
    if not existing:
        return
    safe_id = re.sub(r"[^A-Za-z0-9_.-]", "_", str(node.get("id") or "node"))
    target = root / ".pipeline_attempts" / safe_id / ("attempt-%d" % max(1, int(node.get("attempt") or 1)))
    target.mkdir(parents=True, exist_ok=True)
    for path in existing:
        shutil.copy2(path, target / path.name)


def _refresh_state(state: Dict[str, Any]) -> None:
    nodes = state.get("nodes") or []
    running = [node["id"] for node in nodes if node.get("state") == "running"]
    state["current_nodes"] = running
    if nodes and all(node.get("state") in ("done", "skipped") for node in nodes):
        state["state"] = "done"
        state["recoverable"] = False
    elif any(node.get("state") == "blocked" for node in nodes):
        state["state"] = "blocked"
        state["recoverable"] = False
    elif any(node.get("state") == "failed" for node in nodes):
        state["state"] = "failed"
        state["recoverable"] = False
    elif any(node.get("state") == "retry_wait" for node in nodes):
        state["state"] = "pending"
        state["recoverable"] = any(
            node.get("state") == "retry_wait"
            and int(node.get("attempt") or 0) < int(node.get("max_attempts") or MAX_ATTEMPTS)
            for node in nodes
        )
    elif running:
        state["state"] = "running"
        state["recoverable"] = True
    else:
        state["state"] = "pending"
        state["recoverable"] = any(node.get("state") == "pending" for node in nodes)
    state["updated_at"] = _now()


def _mutate(job: os.PathLike[str] | str, callback):
    path = _pipeline_path(job)
    with _lock(path):
        state = _read(path)
        if state.get("version") != PIPELINE_VERSION:
            raise PipelineError("任务没有可用的流水线状态")
        result = callback(state)
        _refresh_state(state)
        _write(path, state)
        return result, state


def migrate(job: os.PathLike[str] | str) -> Dict[str, Any]:
    """Upgrade durable v1 state without replaying completed chapter nodes."""
    root = Path(job)
    path = _pipeline_path(root)
    with _lock(path):
        state = _read(path)
        version = int(state.get("version") or 0)
        if version == PIPELINE_VERSION:
            return state
        if version != 1:
            raise PipelineError("任务没有可迁移的流水线状态")
        if state.get("mode") == "standard":
            for node in state.get("nodes") or []:
                if node.get("id") != "quality_review":
                    continue
                if node.get("outputs") == ["成品质检报告.md"]:
                    node["outputs"] = ["模型复核报告.md"]
                    if node.get("state") == "done":
                        # 旧文件可能已被后续确定性质检覆盖，无法证明仍是 S2
                        # 原始复核。只重跑这一节点，绝不把脚本报告冒充模型证据。
                        node.update({
                            "state": "pending", "attempt": 0,
                            "input_digest": "", "started_at": "",
                            "finished_at": "", "last_activity_at": "",
                            "error_code": "", "retry_after_seconds": 0,
                        })
        state["version"] = PIPELINE_VERSION
        _refresh_state(state)
        _write(path, state)
        return state


def start_node(job: os.PathLike[str] | str, node_id: str, *, input_digest: str) -> Dict[str, Any]:
    root = Path(job)

    def apply(state):
        node = _find_node(state, node_id)
        digest = str(input_digest or "")
        if node.get("state") == "done" and node.get("input_digest") == digest:
            validate_outputs(root, node)
            return dict(node)
        if node.get("state") == "blocked":
            raise PipelineError("节点已阻断，需要先修复配置或人工重试：%s" % node_id)
        if node.get("state") == "done" and node.get("input_digest") != digest:
            _archive_outputs(root, node)
            node["attempt"] = 0
        if int(node.get("attempt") or 0) >= int(node.get("max_attempts") or MAX_ATTEMPTS):
            raise PipelineError("节点已达到最大尝试次数：%s" % node_id)
        node.update(
            {
                "state": "running",
                "attempt": int(node.get("attempt") or 0) + 1,
                "attempt_serial": int(node.get("attempt_serial") or 0) + 1,
                "input_digest": digest,
                "started_at": _now(),
                "finished_at": "",
                "last_activity_at": _now(),
                "error_code": "",
                "retry_after_seconds": 0,
            }
        )
        return dict(node)

    result, _ = _mutate(root, apply)
    return result


def touch_node(job: os.PathLike[str] | str, node_id: str) -> Dict[str, Any]:
    def apply(state):
        node = _find_node(state, node_id)
        if node.get("state") == "running":
            node["last_activity_at"] = _now()
        return dict(node)

    result, _ = _mutate(job, apply)
    return result


def complete_node(job: os.PathLike[str] | str, node_id: str, *, input_digest: str) -> Dict[str, Any]:
    root = Path(job)

    def apply(state):
        node = _find_node(state, node_id)
        if node.get("input_digest") != str(input_digest or ""):
            raise PipelineError("节点输入已变化，不能提交旧结果")
        validate_outputs(root, node)
        node.update(
            {
                "state": "done",
                "finished_at": _now(),
                "last_activity_at": _now(),
                "error_code": "",
                "retry_after_seconds": 0,
            }
        )
        return dict(node)

    result, _ = _mutate(root, apply)
    return result


def fail_node(
    job: os.PathLike[str] | str,
    node_id: str,
    error_code: str,
    *,
    retryable: bool,
    retry_after: int = 0,
) -> Dict[str, Any]:
    def apply(state):
        node = _find_node(state, node_id)
        attempts = int(node.get("attempt") or 0)
        can_retry = bool(retryable and attempts < int(node.get("max_attempts") or MAX_ATTEMPTS))
        node.update(
            {
                "state": "retry_wait" if can_retry else ("blocked" if not retryable else "failed"),
                "finished_at": _now(),
                "last_activity_at": _now(),
                "error_code": str(error_code or "node_failed")[:120],
                "retry_after_seconds": max(0, int(retry_after or 0)) if can_retry else 0,
            }
        )
        return dict(node)

    result, _ = _mutate(job, apply)
    return result


def run_model_node(
    job: os.PathLike[str] | str,
    node_id: str,
    runner,
    *,
    prompt: str,
    input_digest: str = "",
    sleep=time.sleep,
    on_event=None,
) -> Dict[str, Any]:
    """Run one bounded model node and retry only that node.

    ``runner`` receives a frozen node snapshot, prompt, and concrete model id.
    It must write the declared outputs into the job directory before returning.
    """
    root = Path(job)
    state = load(root)
    node = _find_node(state, node_id)
    if node.get("kind") != "model":
        raise PipelineError("节点不是模型节点：%s" % node_id)
    digest = str(input_digest or hashlib.sha256(prompt.encode("utf-8")).hexdigest())
    route = str(node.get("model_tier") or "fast")
    model = str((state.get("model_routes") or {}).get(route) or "")
    if not model:
        raise NodeExecutionError("model_route_missing", retryable=False, detail="模型路由未配置")

    while True:
        started = start_node(root, node_id, input_digest=digest)
        if started.get("state") == "done":
            return started
        if on_event:
            on_event("started", dict(started))
        try:
            runner(dict(started), prompt, model)
            completed = complete_node(root, node_id, input_digest=digest)
            if on_event:
                on_event("done", dict(completed))
            return completed
        except NodeExecutionError as exc:
            code, retryable = exc.code, exc.retryable
            caught: Exception = exc
        except OutputValidationError as exc:
            code, retryable, caught = "output_validation_failed", True, exc
        except Exception as exc:
            code, retryable, caught = "node_runner_exception", False, exc

        if code == "cancelled":
            cancel_node(root, node_id)
            if on_event:
                on_event("cancelled", dict(started))
            if isinstance(caught, NodeExecutionError):
                raise caught
            raise NodeExecutionError(code, retryable=True, detail=str(caught)) from caught
        attempt = int(started.get("attempt") or 1)
        retry_after = min(30, 2 ** attempt) if retryable else 0
        failed = fail_node(root, node_id, code, retryable=retryable, retry_after=retry_after)
        if on_event:
            on_event("retry" if failed.get("state") == "retry_wait" else "failed", dict(failed))
        if failed.get("state") != "retry_wait":
            if isinstance(caught, NodeExecutionError):
                raise caught
            raise NodeExecutionError(code, retryable=False, detail=str(caught)) from caught
        sleep(retry_after)


def cancel_node(job: os.PathLike[str] | str, node_id: str) -> Dict[str, Any]:
    """Return an interrupted node to a recoverable checkpoint without consuming an attempt."""
    def apply(state):
        node = _find_node(state, node_id)
        if node.get("state") in ("running", "retry_wait"):
            node.update({
                "state": "pending",
                "attempt": max(0, int(node.get("attempt") or 0) - 1),
                "started_at": "",
                "finished_at": "",
                "last_activity_at": "",
                "error_code": "",
                "retry_after_seconds": 0,
            })
        return dict(node)

    result, _ = _mutate(job, apply)
    return result


def recover(job: os.PathLike[str] | str) -> Dict[str, Any]:
    def apply(state):
        for node in state.get("nodes") or []:
            attempts = int(node.get("attempt") or 0)
            maximum = int(node.get("max_attempts") or MAX_ATTEMPTS)
            if node.get("state") in ("running", "retry_wait") and attempts < maximum:
                node.update(
                    {
                        "state": "pending",
                        "started_at": "",
                        "finished_at": "",
                        "retry_after_seconds": 0,
                    }
                )
        state["recoverable"] = any(
            node.get("state") == "pending"
            or (node.get("state") in ("running", "retry_wait")
                and int(node.get("attempt") or 0) < int(node.get("max_attempts") or MAX_ATTEMPTS))
            for node in state.get("nodes") or []
        )
        return state

    _, state = _mutate(job, apply)
    return state


def retry_node(job: os.PathLike[str] | str, node_id: str) -> Dict[str, Any]:
    """Explicit human retry after a permanent/max-attempt stop."""
    def apply(state):
        node = _find_node(state, node_id)
        if node.get("state") not in ("failed", "blocked"):
            raise PipelineError("当前节点不需要人工重试：%s" % node_id)
        node.update(
            {
                "state": "pending",
                "attempt": 0,
                "started_at": "",
                "finished_at": "",
                "last_activity_at": "",
                "error_code": "",
                "retry_after_seconds": 0,
            }
        )
        return dict(node)

    result, _ = _mutate(job, apply)
    return result


def summary(job: os.PathLike[str] | str) -> Dict[str, Any]:
    state = load(job)
    nodes = state.get("nodes") or []
    counts = {status: 0 for status in VALID_STATES}
    for node in nodes:
        status = node.get("state")
        if status in counts:
            counts[status] += 1
    retry_node: Optional[Dict[str, Any]] = next(
        (node for node in nodes if node.get("state") == "retry_wait"), None
    )
    retry = None
    if retry_node:
        retry = {
            "node_id": retry_node.get("id"),
            "attempt": int(retry_node.get("attempt") or 0),
            "max_attempts": int(retry_node.get("max_attempts") or MAX_ATTEMPTS),
            "retry_after_seconds": int(retry_node.get("retry_after_seconds") or 0),
            "error_code": retry_node.get("error_code") or "",
        }
    problem_node: Optional[Dict[str, Any]] = next(
        (node for node in nodes if node.get("state") in ("failed", "blocked")), None
    )
    problem = None
    if problem_node:
        problem = {
            "node_id": problem_node.get("id"),
            "title": problem_node.get("title") or problem_node.get("id"),
            "state": problem_node.get("state"),
            "attempt": int(problem_node.get("attempt") or 0),
            "max_attempts": int(problem_node.get("max_attempts") or MAX_ATTEMPTS),
            "error_code": problem_node.get("error_code") or "",
        }
    return {
        "version": state.get("version"),
        "state": state.get("state") or "",
        "mode": state.get("mode") or "",
        "total": len(nodes),
        "done": counts["done"] + counts["skipped"],
        "running": counts["running"],
        "waiting": counts["pending"],
        "retrying": counts["retry_wait"],
        "failed": counts["failed"] + counts["blocked"],
        "current_nodes": [node.get("id") for node in nodes if node.get("state") == "running"],
        "retry": retry,
        "problem": problem,
        "recoverable": bool(state.get("recoverable")),
        "updated_at": state.get("updated_at") or "",
    }
