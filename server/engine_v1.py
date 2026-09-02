#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
中标狗 · 本地引擎(v1 协议实现)
任务=目录:jobs/<id>/ 下的文件即全部状态(任务.json / progress.json / events.jsonl / chat.jsonl / 交付物)
运行:pip install fastapi uvicorn python-multipart python-docx pypdf certifi && python3 engine_v1.py   # 127.0.0.1:8080
真实 agent:环境变量 AGENT_CMD 命令模板(占位符 {tender}/{out}/{materials}),不配则跑内置 mock 流程。
打包:pyinstaller -F engine_v1.py → 作为 Tauri sidecar 随安装包分发(见 BUILD.md)。
"""
import os, re, sys, ssl, json, glob, time, signal, hashlib, hmac, secrets, base64, contextvars, uuid, shlex, shutil, zipfile, threading, subprocess, datetime, asyncio, socket, http.client, urllib.request, urllib.error, urllib.parse, io, platform
import xml.etree.ElementTree as ET
from collections import deque
from concurrent.futures import ThreadPoolExecutor, as_completed, wait, FIRST_COMPLETED
from threading import Lock as _ThreadLock, RLock as _ThreadRLock
from typing import List
from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse, HTMLResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from template_engine import (builtin_templates, compare_instruction_coverage,
                             compile_template_instructions, derive_template,
                             extract_document_structure, normalize_package,
                             recommend_template, validate_package)
import generation_pipeline
import prompts

def _configure_stdio_utf8():
    """Windows GUI/重定向日志常落到 cp1252；启动横幅含中文时不能因此让整个 sidecar 崩溃。"""
    for name in ('stdout', 'stderr'):
        stream = getattr(sys, name, None)
        reconfigure = getattr(stream, 'reconfigure', None)
        if not reconfigure: continue
        try: reconfigure(encoding='utf-8', errors='backslashreplace')
        except (AttributeError, OSError, ValueError): pass

if os.name == 'nt':
    _configure_stdio_utf8()

ENGINE_VERSION = '0.22.0'
MAX_TEMPLATE_UPLOAD_BYTES = 50 * 1024 * 1024
AUTHOR = 'FDE-家涛'
ENGINE_FEATURES = ['probe_models', 'chat_test', 'agent_binding', 'assets_ingest', 'attachments', 'rerun', 'job_cancel', 'assets_dir_config', 'cli_autofind', 'sowork_engine', 'agent_test',
                   'provider_delete', 'job_delete', 'vision_index', 'artifact_open', 'job_folder_open', 'chat_control', 'job_redo', 'job_stop', 'job_log', 'skill_evidence',
                   's2_engine', 'responses_relay', 'codex_bundled', 'agent_provision', 'preset_config',
                   'quality_gate', 'job_start', 'job_resume', 'multi_file_job', 'models_cache', 'async_sse', 's2_quick_setup',
                   'opencode_engine', 'relay_chat_passthrough', 'opencode_bundled', 'dual_shell_provision',
                   'worklog_stream', 'stage_eta', 'job_presentation', 'runtime_capabilities',
                   'delivery_summary', 'job_usage', 'job_archive', 'job_projects', 'job_bulk_actions',
                   'deliverables_zip', 'task_templates', 'one_click_diagnostics', 'job_revisions',
                   'setup_onboarding', 'scene_template_packages', 'template_derivation',
                   'incremental_model_stream', 'automatic_session_recovery',
                   'checkpoint_generation_pipeline', 'local_tender_parse',
                   'bootstrap_health']
HERE = os.path.dirname(os.path.abspath(__file__))
def _data_root():
    env = os.environ.get('BID_HOME')
    if env: return env
    base = os.path.join(os.path.expanduser('~'), 'Documents')
    new, old = os.path.join(base, '中标狗'), os.path.join(base, '标书助手')  # 旧品牌目录名,勿被全局改名脚本替换
    if not os.path.isdir(new) and os.path.isdir(old):
        try: os.rename(old, new)          # 产品改名:老用户数据目录整体自动迁移
        except Exception: return old      # 挪不动(权限/被占用)就继续用老目录,数据优先
    return new
DATA = _data_root()
MULTIUSER = os.environ.get('BID_MULTIUSER') == '1'   # 云端网页模式:每位访客独立工作区
PASSWORD = os.environ.get('BID_PASSWORD', '')        # 设置后需口令才能访问(公网必设)
ALLOW_AGENT_CONFIG = os.environ.get('BID_ALLOW_AGENT_CONFIG') == '1'  # 云端默认禁止网页改生成引擎命令
_ws = contextvars.ContextVar('ws', default='')

def ws_root():
    w = _ws.get()
    return os.path.join(DATA, 'workspaces', w) if (MULTIUSER and w) else DATA

def _mk(p): os.makedirs(p, exist_ok=True); return p
def jobs_dir():   return _mk(os.path.join(ws_root(), 'jobs'))
def conf_path():  return os.path.join(_mk(ws_root()), 'config.json')

def assets_default(): return os.path.join(ws_root(), '素材库')

def assets_dir():
    """素材库位置:默认在数据目录,可在「素材库」面板改到任意文件夹(如公司共享盘)"""
    conf = json_quiet(conf_path())
    d = (conf.get('assets_dir') or '').strip()
    if d:
        d = os.path.expanduser(d)
        try: return _mk(d)
        except Exception: pass          # 配置的路径建不出来(盘符不在了等)→ 回落默认,不让功能瘫掉
    return _mk(assets_default())

def json_quiet(p):
    try:
        with open(p, encoding='utf-8') as f: return json.load(f)
    except Exception: return {}

def ensure_preset():
    """预置配置:给某个客户定制发包时,把 preset_config.json 放在引擎旁边(或 BID_PRESET 指定),
    首次启动自动作为 config.json 种子——客户拿到手连 Key 都不用填。
    只在还没有配置时生效,绝不覆盖用户已有设置;公开发布包里不放这个文件。"""
    cp = os.path.join(_mk(DATA), 'config.json')
    if os.path.isfile(cp): return
    cands = [os.environ.get('BID_PRESET', '')]
    meipass = getattr(sys, '_MEIPASS', None)
    if meipass: cands.append(os.path.join(meipass, 'preset_config.json'))
    cands += [os.path.join(os.path.dirname(os.path.abspath(sys.executable)), 'preset_config.json'),
              os.path.join(HERE, 'preset_config.json')]
    for c in cands:
        if c and os.path.isfile(c):
            data = json_quiet(c)
            if data:
                write_json(cp, data)
                return

app = FastAPI(title='bid-dog-engine')
_DESKTOP_ORIGINS = {'tauri://localhost', 'http://tauri.localhost', 'https://tauri.localhost'}
_DESKTOP_ORIGINS.update(x.strip() for x in os.environ.get('BID_ALLOWED_ORIGINS', '').split(',') if x.strip())
# *.localhost 整族按 RFC 6761/浏览器实现恒解析回环,放行它不增加 DNS 重绑定面。
# 为什么需要:新前端(app-next)是 ES module,规范规定 module 脚本一律按 CORS 模式抓取,
# 同源也带 Origin 头——tauri.localhost:端口 下的静态资源请求因此会进这道闸。
_LOOPBACK_ORIGIN = re.compile(r'^https?://(?:127\.0\.0\.1|(?:[a-z0-9-]+\.)*localhost)(?::\d+)?$', re.I)

def origin_allowed(origin):
    """只让本机页面/Tauri 壳跨域访问桌面引擎，阻断恶意网页读取 Key 或发起本机命令。"""
    value = str(origin or '').strip()
    return not value or value in _DESKTOP_ORIGINS or bool(_LOOPBACK_ORIGIN.fullmatch(value))

app.add_middleware(
    CORSMiddleware,
    allow_origins=sorted(_DESKTOP_ORIGINS),
    allow_origin_regex=r'^https?://(?:127\.0\.0\.1|(?:[a-z0-9-]+\.)*localhost)(?::\d+)?$',
    allow_methods=['*'],
    allow_headers=['*'],
    expose_headers=['X-Request-ID'],
    allow_credentials=True,
)

# ---------- 云端网页模式:口令门 + 访客工作区隔离(桌面版不设 BID_PASSWORD 则完全不生效) ----------
def _tok(pw): return hashlib.sha256(('bid-assistant|' + pw).encode()).hexdigest()[:32]

LOGIN_HTML = """<!doctype html><meta charset=utf-8><title>中标狗</title><link rel="icon" type="image/svg+xml" href="data:image/svg+xml,%3Csvg xmlns=%22http://www.w3.org/2000/svg%22 width=%221024%22 height=%221024%22 viewBox=%220 0 1024 1024%22%3E%3Cdefs%3E%3ClinearGradient id=%22bg%22 x1=%220%22 y1=%220%22 x2=%220%22 y2=%221%22%3E%3Cstop offset=%220%22 stop-color=%22%231274e0%22/%3E%3Cstop offset=%221%22 stop-color=%22%230a55aa%22/%3E%3C/linearGradient%3E%3C/defs%3E%3Crect width=%221024%22 height=%221024%22 rx=%22232%22 fill=%22url(%23bg)%22/%3E%3Cg transform=%22translate(512 512) scale(1.07) translate(-512 -446)%22%3E%3Cg fill=%22%23ffffff%22%3E%3Cpath d=%22M 344 332 L 322 172 Q 318 140 350 152 L 472 218 Q 430 252 344 332 Z%22/%3E%3Cpath d=%22M 680 332 L 702 172 Q 706 140 674 152 L 552 218 Q 594 252 680 332 Z%22/%3E%3Cellipse cx=%22512%22 cy=%22490%22 rx=%22270%22 ry=%22250%22/%3E%3C/g%3E%3Cg fill=%22%230a55aa%22%3E%3Ccircle cx=%22418%22 cy=%22448%22 r=%2232%22/%3E%3Ccircle cx=%22606%22 cy=%22448%22 r=%2232%22/%3E%3Cellipse cx=%22512%22 cy=%22560%22 rx=%2250%22 ry=%2238%22/%3E%3C/g%3E%3Cpath d=%22M 512 596 L 512 640 Q 512 706 470 706 Q 442 706 442 676 Q 442 662 452 654%22 fill=%22none%22/%3E%3Cpath d=%22M 456 618 Q 456 700 512 700 Q 568 700 568 618 Q 540 636 512 636 Q 484 636 456 618 Z%22 fill=%22%23ff8da1%22/%3E%3C/g%3E%3C/svg%3E">
<style>body{font-family:-apple-system,"PingFang SC","Segoe UI",sans-serif;display:flex;align-items:center;
justify-content:center;height:100vh;margin:0;background:#fbfbfd;color:#1d1d1f}
form{background:#fff;padding:38px 40px;border-radius:16px;box-shadow:0 10px 40px rgba(0,0,0,.08);width:320px}
h1{font:600 20px/1.3 inherit;margin:0 0 8px}p{font:400 13px/1.6 inherit;color:#6e6e73;margin:0 0 20px}
input{width:100%;box-sizing:border-box;border:1px solid rgba(0,0,0,.12);border-radius:9px;padding:11px 12px;font:400 14px inherit;outline:none}
input:focus{border-color:#0a63c9}button{width:100%;margin-top:12px;border:0;border-radius:999px;background:#0a63c9;
color:#fff;font:500 14px inherit;padding:11px;cursor:pointer}button:hover{background:#0a55aa}
.e{color:#c0392b;font:400 12.5px inherit;margin-top:10px}</style>
<form method=post action=/login><h1>中标狗</h1><p>请输入访问口令</p>
<input name=password type=password placeholder="访问口令" autofocus><button>进入</button>__ERR__</form>"""

def login_page(err='', code=200):
    return HTMLResponse(LOGIN_HTML.replace('__ERR__', '<div class=e>口令不正确</div>' if err else ''), status_code=code)

@app.post('/login')
async def login(request: Request):
    form = await request.form()
    if PASSWORD and form.get('password') != PASSWORD:
        return login_page(err='1', code=401)
    r = RedirectResponse('/', status_code=303)
    r.set_cookie('bid_auth', _tok(PASSWORD), max_age=30 * 86400, httponly=True, samesite='lax')
    r.set_cookie('bid_uid', request.cookies.get('bid_uid') or secrets.token_hex(8), max_age=365 * 86400, httponly=True, samesite='lax')
    return r

@app.middleware('http')
async def gate(request: Request, call_next):
    supplied_request_id = str(request.headers.get('x-request-id') or '')[:80]
    request_id = (supplied_request_id if re.fullmatch(r'[A-Za-z0-9._:-]{8,80}', supplied_request_id)
                  else uuid.uuid4().hex)
    request.state.request_id = request_id
    def tracked(response):
        response.headers['X-Request-ID'] = request_id
        return response
    uid = request.cookies.get('bid_uid') or (secrets.token_hex(8) if MULTIUSER else '')
    _ws.set(uid)
    origin = request.headers.get('origin', '')
    if origin and not origin_allowed(origin):
        return tracked(JSONResponse({'error': 'forbidden origin', 'request_id': request_id}, status_code=403))
    # /v1/relay/* 是给本机 Codex 用的,它没有浏览器 Cookie,自己带 relay_token 鉴权,不走口令门
    if (PASSWORD and request.url.path not in ('/login', '/v1/health')
            and not request.url.path.startswith('/v1/relay/')
            and request.cookies.get('bid_auth') != _tok(PASSWORD)):
        if request.url.path.startswith('/v1/'):
            return tracked(JSONResponse({'error': 'unauthorized', 'login': '/',
                                         'request_id': request_id}, status_code=401))
        return login_page()
    resp = await call_next(request)
    if MULTIUSER and uid and not request.cookies.get('bid_uid'):
        resp.set_cookie('bid_uid', uid, max_age=365 * 86400, httponly=True, samesite='lax')
    return tracked(resp)

def now(): return datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
def jpath(jid): return os.path.join(jobs_dir(), os.path.basename(jid))

_TS_FORMAT = '%Y-%m-%d %H:%M:%S'

def _parse_ts(value):
    try: return datetime.datetime.strptime(str(value or ''), _TS_FORMAT)
    except (TypeError, ValueError): return None

def _format_ts(value):
    return value.strftime(_TS_FORMAT) if isinstance(value, datetime.datetime) else ''

def read_json(p, dft):
    try:
        with open(p, encoding='utf-8') as f: return json.load(f)
    except Exception: return dft
def write_json(p, obj):
    with _json_lock(p):
        tmp = p + '.tmp'
        with open(tmp, 'w', encoding='utf-8') as f: json.dump(obj, f, ensure_ascii=False, indent=1)
        os.replace(tmp, p)

_JSON_LOCKS = {}
_JSON_LOCKS_GUARD = _ThreadLock()

def _json_lock(path):
    key = os.path.realpath(path)
    with _JSON_LOCKS_GUARD:
        return _JSON_LOCKS.setdefault(key, _ThreadRLock())

def patch_json(path, changes=None, remove=()):
    """Atomically read/merge/write a JSON object so UI metadata cannot clobber runtime fields."""
    with _json_lock(path):
        current = read_json(path, {})
        if not isinstance(current, dict): current = {}
        current.update(changes or {})
        for key in remove: current.pop(key, None)
        write_json(path, current)
        return current

def update_runtime(job, **changes):
    changes['updated_at'] = now()
    return patch_json(os.path.join(job, 'runtime.json'), changes)

def append_diagnostic(job, code, detail, level='warning', **context):
    """Persist redacted technical evidence separately from calm user-facing events."""
    if not os.path.isdir(job): return
    record = {'ts': now(), 'code': str(code or 'unknown'), 'level': level,
              'detail': _redact_runtime(str(detail or ''))}
    if context: record['context'] = _redact_runtime(context)
    path = os.path.join(job, 'diagnostics.jsonl')
    with _json_lock(path):
        try:
            with open(path, 'a', encoding='utf-8') as target:
                target.write(json.dumps(record, ensure_ascii=False) + '\n')
        except FileNotFoundError:
            pass

def compatibility_fallback(job, code, detail):
    reason = '当前稳定运行方式暂不支持暂停；如需中止，请使用停止。'
    runtime = read_json(os.path.join(job, 'runtime.json'), {})
    fallback_count = max(0, int(runtime.get('fallback_count') or 0)) + 1
    update_runtime(job, execution_path='cli_compat', can_pause=False,
                   pause_disabled_reason=reason, fallback_reason_code=code,
                   fallback_count=fallback_count)
    append_diagnostic(job, code, detail, fallback='cli_compat')
    emit(job, {'type': 'message', 'role': 'agent',
               'text': '主连接响应较慢，已切换稳定通道继续；'
                       '仍使用同一模型和同一套要求，不会降低内容标准。'})

_SECRET_RE = re.compile(r'(?i)(?:' + 's' + 'k' + r')-[a-z0-9_-]{16,}')
_NAMED_SECRET_RE = re.compile(
    r'(?i)\b(api[_-]?key|access[_-]?token|refresh[_-]?token|password|secret)\b(\s*[:=]\s*)["\']?([^\s,"\']+)'
)
_BEARER_RE = re.compile(r'(?i)\b(Bearer)(\s+)[A-Za-z0-9._~+/=-]{12,}')
_URL_RE = re.compile(r'(?i)https?://[^\s<>"\']+')

def _sanitize_urls(value):
    """URL 只保留协议、主机、端口和路径，禁止诊断中泄露 userinfo/query token。"""
    def replace(match):
        raw = match.group(0)
        suffix = ''
        while raw and raw[-1] in '),.;]}':
            suffix = raw[-1] + suffix
            raw = raw[:-1]
        return generation_pipeline.safe_base_url(raw) + suffix
    return _URL_RE.sub(replace, str(value or ''))

def redact(value):
    """事件、诊断包和用户可见错误的统一脱敏入口。配置原件和 run.log 不在这里改写。"""
    if isinstance(value, dict): return {k: redact(v) for k, v in value.items()}
    if isinstance(value, list): return [redact(v) for v in value]
    if isinstance(value, tuple): return tuple(redact(v) for v in value)
    if isinstance(value, str):
        value = _sanitize_urls(value)
        value = _SECRET_RE.sub('[API Key 已隐藏]', value)
        value = _NAMED_SECRET_RE.sub(lambda m: m.group(1) + m.group(2) + '[凭据已隐藏]', value)
        return _BEARER_RE.sub(lambda m: m.group(1) + m.group(2) + '[凭据已隐藏]', value)
    return value

def _safe_secret_text(value, secrets_to_hide=()):
    """用户可见文本先按本次请求的真实凭据精确替换，再走通用 Key 模式脱敏。"""
    text = str(value or '')
    for secret in secrets_to_hide or ():
        secret = str(secret or '')
        if secret: text = text.replace(secret, '[API Key 已隐藏]')
    return redact(text)

def _http_error_detail(error, secrets_to_hide=(), limit=400):
    """上游错误正文不回显；网关可能把 Authorization 原样塞进 message。"""
    try: size = len(error.read())
    except Exception: size = 0
    return '上游错误正文已隐藏%s' % (('(%d bytes)' % size) if size else '')

def _sensitive_config_key(key):
    norm = re.sub(r'[^a-z0-9]', '', str(key or '').lower())
    return (norm in ('key', 'apikey', 'token', 'accesstoken', 'refreshtoken', 'password', 'secret',
                     'credential', 'credentials', 's2key', 'relaytoken')
            or norm.endswith(('apikey', 'accesstoken', 'refreshtoken', 'password', 'secret')))

def _configured_secrets(conf=None):
    """收集本地已知凭据，供日志/诊断/状态响应做精确值替换。"""
    conf = conf if conf is not None else read_json(conf_path(), {})
    found = set()
    def collect_scalars(value):
        if isinstance(value, dict):
            for child in value.values(): collect_scalars(child)
        elif isinstance(value, list):
            for child in value: collect_scalars(child)
        elif isinstance(value, (str, int, float)):
            text = str(value).strip()
            if text: found.add(text)
    def walk(value, key=''):
        if _sensitive_config_key(key):
            collect_scalars(value)
            return
        if isinstance(value, dict):
            for child_key, child in value.items(): walk(child, child_key)
        elif isinstance(value, list):
            for child in value: walk(child, key)
    walk(conf)
    eng = conf.get('engine') or {}
    for field in ('cmd', 'env'):
        text = str(eng.get(field) or '').strip()
        if text:
            found.add(text)
            for line in text.splitlines():
                if '=' in line:
                    name, value = line.split('=', 1)
                    if _sensitive_config_key(name): collect_scalars(value.strip().strip('"\''))
    return tuple(sorted(found, key=len, reverse=True))

def _redact_runtime(value, conf=None):
    secrets_to_hide = _configured_secrets(conf)
    def clean(item):
        if isinstance(item, dict): return {k: clean(v) for k, v in item.items()}
        if isinstance(item, list): return [clean(v) for v in item]
        if isinstance(item, tuple): return tuple(clean(v) for v in item)
        if isinstance(item, str): return _strip_terminal_controls(_safe_secret_text(item, secrets_to_hide))
        return item
    return clean(value)

_TERMINAL_ESCAPE_RE = re.compile(
    r'\x1b(?:\[[0-?]*[ -/]*[@-~]|\][^\x07]*(?:\x07|\x1b\\)?)'
)

def _strip_terminal_controls(value):
    """日志可以保留换行和制表符，但不能把 ANSI/OSC 控制码带进产品 UI。"""
    text = _TERMINAL_ESCAPE_RE.sub('', str(value or ''))
    return ''.join(ch for ch in text if ch in ('\n', '\r', '\t') or ord(ch) >= 32)

def emit(job, ev):
    """事件即真相:UI 只消费引擎验证后的 events.jsonl。"""
    if not os.path.isdir(job): return
    ev = _redact_runtime(dict(ev))
    if ev.get('type') == 'error' and not ev.get('actions'):
        ev['actions'] = [{'act': 'open_log', 'label': '查看运行日志'},
                         {'act': 'rerun', 'label': '重跑本任务'}]
    # 所有进度都重新按磁盘证据验算；字典里的 verified 只是数据，绝不是信任凭证。
    if ev.get('type') == 'progress':
        try: ev = sanitize_event(job, ev)
        except Exception: pass
    ev['ts'] = now()
    try:
        event_path = os.path.join(job, 'events.jsonl')
        # 会话 watcher、worker 和控制请求会并发写事件。单次 f.write 不是跨线程
        # 的协议边界；用同一文件锁保证每条 JSONL 不被另一条从中间插断。
        with _json_lock(event_path):
            with open(event_path, 'a', encoding='utf-8') as f:
                f.write(json.dumps(ev, ensure_ascii=False) + '\n')
        if ev['type'] == 'progress':
            write_json(os.path.join(job, 'progress.json'), ev)
    except FileNotFoundError:
        pass

STAGES = ['体检素材', '图片入库', '读懂组成', '提取格式', '评分废标', '拆解分工',
          '分章撰写', '逐条应答', '汇总成册', '配图复核', '自查体检', '出Word质检']

# agent 只允许写这个“声明流”；引擎核验后再转写 canonical events.jsonl。
AGENT_EVENTS_FILE = 'agent_events.jsonl'

# 这是产品阶段与最小磁盘证据的公开契约。实际判定在 evidence_for_step()，保持一一对应。
STAGE_EVIDENCE = [
    {'step': 1, 'stage': STAGES[0], 'evidence': '招标文件或解析收据'},
    {'step': 2, 'stage': STAGES[1], 'evidence': '图片索引或无图片收据'},
    {'step': 3, 'stage': STAGES[2], 'evidence': '投标文件组成分析'},
    {'step': 4, 'stage': STAGES[3], 'evidence': '可解析格式规范与摘要'},
    {'step': 5, 'stage': STAGES[4], 'evidence': '评分矩阵与废标风险清单'},
    {'step': 6, 'stage': STAGES[5], 'evidence': '响应矩阵'},
    {'step': 7, 'stage': STAGES[6], 'evidence': '大纲对应的章节稿'},
    {'step': 8, 'stage': STAGES[7], 'evidence': '商务/技术偏离表'},
    {'step': 9, 'stage': STAGES[8], 'evidence': '正文 Markdown'},
    {'step': 10, 'stage': STAGES[9], 'evidence': '配图复核或无图收据'},
    {'step': 11, 'stage': STAGES[10], 'evidence': '投标文件自检报告'},
    {'step': 12, 'stage': STAGES[11], 'evidence': '非空正文 Word'},
]

def mock_agent(job):
    """内置模拟 agent:走完 12 阶段,发问一次,产出样例交付物 —— 演示与前端联调用"""
    _mock_agent(job)

def _cancelled(job): return _cancel_requested(os.path.basename(job)) or not os.path.isdir(job)

def _mock_agent(job):
    emit(job, {'type': 'message', 'role': 'agent', 'text': '收到招标文件,开始读。读完我会把结构、评分表和要你确认的事列出来。'})
    for i, st in enumerate(STAGES):
        if _cancelled(job): return            # 任务被删:立即收工
        if read_json(os.path.join(job, '任务.json'), {}).get('paused'):
            while read_json(os.path.join(job, '任务.json'), {}).get('paused'):
                if _cancelled(job): return
                time.sleep(1)
        pct = int((i + 0.5) / len(STAGES) * 100)
        emit(job, {'type': 'progress', 'stage': st, 'pct': pct, 'step': i + 1, 'total': len(STAGES)})
        emit(job, {'type': 'worklog', 'lines': ['[%s] 开始:读取上一步产物,校验输入齐全' % st,
                                                '$ 正在执行 %s 相关脚本与分析…' % st]})
        if st == '评分废标':
            emit(job, {'type': 'question', 'id': 'q1', 'text': '评分表要求安全生产许可证在有效期内;素材库那份 2026-06 到期。等你上传新证,还是按「即将换证」写?',
                       'options': ['等我上传新证', '按即将换证写']})
        if st == '分章撰写':
            emit(job, {'type': 'message', 'role': 'agent', 'text': '六章已分给子 agent 并行写,写完逐章汇报。'})
            for ch in range(1, 7):
                time.sleep(1.2); emit(job, {'type': 'chapter', 'no': ch, 'state': 'done', 'pages': 6 + ch * 3})
        time.sleep(1.5)
    for fn, txt in [('投标文件_技术标.md', '# 投标文件(技术标)\n\n(mock 交付物,接真实 agent 后为完整标书)\n'),
                    ('投标文件自检报告.md', '# 自检报告\n\n结论:仅可作初稿\n\n## 需人工确认\n- 投标人名称\n- 报价\n- 安全生产许可证新件\n')]:
        open(os.path.join(job, fn), 'w', encoding='utf-8').write(txt)
        emit(job, {'type': 'artifact', 'name': fn})
    settle(job)

DELIVER_EXT = ('.md', '.docx', '.xlsx', '.pdf')

NOT_DELIVERABLE = ('你的要求.md',)   # 输入件,不是产出:算进交付物会让「跑完却没产出」的告警失灵

def list_deliverables(job):
    meta = read_json(os.path.join(job, '任务.json'), {})
    tender = meta.get('tender', '')
    out = []
    for fn in sorted(os.listdir(job)):
        if fn.endswith(DELIVER_EXT) and not fn.startswith(('_', '.')) \
                and fn != tender and fn not in NOT_DELIVERABLE:
            out.append(fn)
    return out

def artifact_info(fn):
    """给 UI 的稳定产物说明:不靠前端猜文件用途,同名规则在桌面/网页都一致。"""
    low, ext = fn.lower(), os.path.splitext(fn)[1].lower()
    if ext in ('.docx', '.xlsx', '.pdf'):
        group, kind, rank = 0, {'.docx': 'WORD', '.xlsx': 'EXCEL', '.pdf': 'PDF'}[ext], 10
        if ext == '.docx' and ('完整' in fn or '投标文件' in fn):
            purpose, rank = '主要交付文件：可编辑的完整投标文件，提交前请人工复核、签字和盖章。', 0
        elif ext == '.xlsx':
            purpose = '可编辑的表格附件，用于核对清单、报价或响应数据。'
        elif ext == '.pdf':
            purpose = '便于预览或提交的 PDF 文件，请核对页面与盖章要求。'
        else:
            purpose = 'Word 交付文件，可直接打开检查内容和排版。'
    elif any(k in fn for k in ('自检', '门禁', '检查报告', '废标风险', '补料清单')):
        group, kind, rank = 1, ('报告' if '报告' in fn or '自检' in fn else '清单'), 10
        if '格式' in fn or '门禁' in fn:
            purpose, rank = '格式检查结果：核对字体、页边距、目录、页码与表格版式。', 0
        elif '补料' in fn:
            purpose, rank = '待补材料清单：列出提交前仍需补齐或确认的事实材料。', 1
        elif '废标' in fn:
            purpose, rank = '高风险事项清单：提交前逐条排除可能导致废标的问题。', 2
        else:
            purpose, rank = '内容自检报告：查看缺失项、风险项和需要人工确认的结论。', 0
    elif any(k in fn for k in ('响应矩阵', '响应对照', '偏离表', '评分', '格式要求', '解析版', '配图清单', '组成')):
        group, kind, rank = 2, '分析', 10
        if '响应矩阵' in fn or '响应对照' in fn:
            purpose, rank = '响应依据：对照招标要求、评分点与标书章节，检查是否逐条覆盖。', 0
        elif '偏离表' in fn:
            purpose, rank = '偏离核对：集中查看商务或技术响应与招标要求的差异。', 1
        elif '评分' in fn:
            purpose, rank = '评分分析：核对评分标准、得分证据和对应章节。', 2
        elif '解析版' in fn:
            purpose = '招标文件解析文本，供 Agent 分析和追溯原始要求。'
        elif '配图' in fn:
            purpose = '配图计划：说明图片素材及其建议插入位置。'
        else:
            purpose = '分析依据文件，用于复核要求、格式和响应关系。'
    else:
        group, kind, rank = 3, ('稿件' if ext == '.md' else ext.lstrip('.').upper()), 10
        purpose = '过程稿件：用于追溯和继续编辑，通常无需作为最终交付提交。'
    return {'group': group, 'kind': kind, 'purpose': purpose, 'rank': rank}

def artifact_summary(job, names=None):
    names = sorted(names or list_deliverables(job),
                   key=lambda n: (artifact_info(n)['group'], artifact_info(n)['rank'], n))
    body_words = _body_docxs(job, names) if '_body_docxs' in globals() else []
    primary = body_words[0] if body_words else None
    checks = [n for n in names if artifact_info(n)['group'] == 1]
    if not primary:
        # 客户报障:「跑完了,没有生成标书」。之前这里照样播报「任务完成,已整理好 N 个文件」——
        # 因为 md 稿件也算交付物,一堆 .md 就让 known 非空,完成态与告警全都失效。
        # 缺最终 Word 就是没交付,必须当面说,并给一键补出的按钮(零 token,不重跑)。
        body = _body_mds(job, names)
        lines = ['⚠ **跑完了,但没有生成最终 Word 标书。**', '']
        if body:
            lines += ['正文稿已经写出来了(`%s`),缺的只是最后一步「导出 Word」——'
                      '模型跳过了 build_tender_docx.py。点下面的按钮可以直接补出来,'
                      '**不消耗额度、不用重跑**。' % '`、`'.join(body[:3]), '']
        else:
            lines += ['任务目录里连正文稿(投标文件_*.md)都没有,说明生成环节就没写成。'
                      '请看运行日志确认是哪一步停的,然后重跑本任务。', '']
        lines.append('已产出的 %d 个文件仍在右侧“已产出”里,可以逐个查看。' % len(names))
        actions = ([{'act': 'export_docx', 'label': '立即补出 Word'}] if body else []) \
            + [{'act': 'open_log', 'label': '查看运行日志'},
               {'act': 'rerun', 'label': '重跑本任务'},
               {'act': 'open_job_folder', 'label': '打开任务文件夹'}]
        return '\n'.join(lines), actions
    lines = ['任务完成，已整理好 **%d 个文件**。' % len(names), '', '**建议按这个顺序查看：**']
    if primary: lines.append('1. `%s` — 最终 Word，先检查内容与排版。' % primary)
    if checks: lines.append('%d. `%s` — 查看风险、缺项和人工确认事项。' % (2 if primary else 1, checks[0]))
    lines += ['', '其他分析依据和过程稿件已收纳在右侧“已产出”，点击文件旁的“打开”即可直接查看。']
    actions = []
    if primary: actions.append({'act': 'open_artifact', 'label': '打开最终 Word', 'file': primary})
    actions.append({'act': 'open_job_folder', 'label': '打开任务文件夹'})
    return '\n'.join(lines), actions

# 12 步默认预估秒数(没有本机历史时的参考值;有历史后用滚动平均逐步替代)
STAGE_DEFAULT_S = [60, 120, 150, 60, 150, 150, 900, 300, 180, 60, 120, 150]

def stage_stats_path(): return os.path.join(DATA, 'stage_stats.json')

def record_stage(step, dur):
    """某一步真实耗时入库:EWMA 滚动平均,预计时间越用越准"""
    if not (1 <= step <= 12) or dur <= 0 or dur > 3 * 3600: return
    st = read_json(stage_stats_path(), {})
    k = str(step)
    cur = st.get(k) or {}
    avg = cur.get('avg')
    st[k] = {'avg': round(dur if avg is None else 0.6 * avg + 0.4 * dur, 1), 'n': int(cur.get('n', 0)) + 1}
    try: write_json(stage_stats_path(), st)
    except Exception: pass

@app.get('/v1/stats/stages')
def stage_stats():
    """前端算「预计还剩」用:每步平均耗时(本机历史优先,缺的用默认参考值)"""
    st = read_json(stage_stats_path(), {})
    avgs = []
    for i in range(1, 13):
        h = st.get(str(i))
        avgs.append({'step': i, 'avg_s': (h or {}).get('avg') or STAGE_DEFAULT_S[i - 1],
                     'from_history': bool(h)})
    return {'stages': avgs}

def _last_json_event(job):
    """Return the last complete canonical event without loading an unbounded log."""
    path = os.path.join(job, 'events.jsonl')
    try:
        with open(path, 'rb') as source:
            source.seek(0, os.SEEK_END)
            size = source.tell()
            source.seek(max(0, size - 65536))
            lines = source.read().decode('utf-8', 'ignore').splitlines()
    except OSError:
        return {}
    for line in reversed(lines):
        try: return json.loads(line)
        except (TypeError, ValueError): continue
    return {}

def _job_last_activity(job, meta=None, prog=None, outcome=None):
    meta = meta if isinstance(meta, dict) else read_json(os.path.join(job, '任务.json'), {})
    prog = prog if isinstance(prog, dict) else read_json(os.path.join(job, 'progress.json'), {})
    outcome = outcome if isinstance(outcome, dict) else read_json(os.path.join(job, 'outcome.json'), {})
    runtime = read_json(os.path.join(job, 'runtime.json'), {})
    candidates = [meta.get('created_at'), prog.get('ts'), outcome.get('ts'), runtime.get('updated_at'),
                  _last_json_event(job).get('ts')]
    parsed = [dt for dt in (_parse_ts(value) for value in candidates) if dt]
    if parsed: return _format_ts(max(parsed))
    try: return _format_ts(datetime.datetime.fromtimestamp(os.path.getmtime(job)))
    except OSError: return now()

def _job_elapsed(job, state, meta=None, outcome=None, last_activity=''):
    meta = meta if isinstance(meta, dict) else read_json(os.path.join(job, '任务.json'), {})
    outcome = outcome if isinstance(outcome, dict) else read_json(os.path.join(job, 'outcome.json'), {})
    snapshot = meta.get('engine_snapshot') if isinstance(meta.get('engine_snapshot'), dict) else {}
    start = _parse_ts(snapshot.get('started_at') or meta.get('created_at'))
    if not start: return 0
    terminal = state in ('done', 'stopped', 'unknown')
    end = _parse_ts(outcome.get('ts')) if terminal else _parse_ts(now())
    end = end or _parse_ts(last_activity) or start
    return max(0, int((end - start).total_seconds()))

def _job_eta(state, prog):
    if state in ('done', 'stopped', 'unknown'): return 0
    try: step = max(0, min(12, int((prog or {}).get('step') or 0)))
    except (TypeError, ValueError): step = 0
    rows = (stage_stats() or {}).get('stages') or []
    return max(0, int(sum(float(row.get('avg_s') or 0) for row in rows if int(row.get('step') or 0) > step)))

def _job_usage(job, meta=None):
    meta = meta if isinstance(meta, dict) else read_json(os.path.join(job, '任务.json'), {})
    snapshot = meta.get('engine_snapshot') if isinstance(meta.get('engine_snapshot'), dict) else {}
    raw = read_json(os.path.join(job, 'usage.json'), {})
    return {'model': raw.get('model') or snapshot.get('model') or '',
            'calls': max(0, int(raw.get('calls') or 0)),
            'input_tokens': max(0, int(raw.get('input_tokens') or 0)),
            'output_tokens': max(0, int(raw.get('output_tokens') or 0)),
            'total_tokens': max(0, int(raw.get('total_tokens') or 0)),
            'estimated_cost': raw.get('estimated_cost'),
            'currency': raw.get('currency') or None}

# 工作台词过滤:agent 的原始输出噪声很大,只留人能看懂的动作行
_ANSI_RE = re.compile(r'\x1b\[[0-9;]*m')
_WORKLOG_SKIP = re.compile(r'^(session id|reasoning|workdir|model:|provider|approval|sandbox|tokens used|'
                           r'-{5,}|> build|user$|codex$|exec$|thinking$|mcp|WARNING|warning:|To view this session)')

def worklog_clean(chunk):
    out = []
    for ln in chunk.splitlines():
        ln = _ANSI_RE.sub('', ln).strip()
        if not ln or _WORKLOG_SKIP.match(ln): continue
        if len(ln) > 160: ln = ln[:157] + '…'
        out.append(ln)
    return out

RUNNING = {}      # jid → 本轮 owner token；完整收尾前一直占位，旧 worker 不能释放新一轮
RUNNING_LOCK = threading.Lock()
SHUTTING_DOWN = False
EXITING = False
SHUTDOWN_GENERATION = 0
OC_REPLAYING = False
JOB_CONTROL = {} # jid → stop/delete/pause 控制令牌；控制动作期间禁止下一轮抢跑
CANCEL = {}      # jid → owner token；取消只作用于指定运行代际
TERMINAL_OWNERS = {} # jid → 已原子提交完成态的 owner；停止不能再把同一代伪装成已停止
PROCS = {}        # jid → 真实 agent 的 Popen 句柄(删除任务时用来杀进程)
PROC_OWNERS = {}  # jid → 与 PROCS 同一代的 owner token
RUNNING_THREADS = {}  # jid → 本轮 worker 线程句柄;用于识别「线程已死但 owner 未释放」的僵尸占位

def _reserve_running_reason(base):
    """返回 (owner, reason)。调用方据此区分同单重复点击与全局暂时禁入。"""
    with RUNNING_LOCK:
        # 同一任务已经占位时必须先返回 running；shutdown/replay 不能借一次重复点击
        # 把真实活跃任务改写成 staged。
        if base in RUNNING:
            # 僵尸占位自愈:worker 线程已死(极端崩溃没走到 finally 释放)时,
            # RUNNING 里的 owner 会把这单永远焊死在「正在跑」——重试/重跑全被拒,
            # 用户只能重启应用。线程对象在且已死,当场回收,让这次预约继续。
            t = RUNNING_THREADS.get(base)
            if t is not None and not t.is_alive():
                RUNNING.pop(base, None)
                CANCEL.pop(base, None)
                TERMINAL_OWNERS.pop(base, None)
                RUNNING_THREADS.pop(base, None)
            else:
                return None, 'running'
        if EXITING: return None, 'exiting'
        if SHUTTING_DOWN: return None, 'shutdown'
        if OC_REPLAYING: return None, 'replay'
        if base in JOB_CONTROL: return None, 'control'
        owner = secrets.token_hex(16)
        CANCEL.pop(base, None)  # 历史异常退出留下的脏标记不能污染新一轮
        TERMINAL_OWNERS.pop(base, None)
        RUNNING[base] = owner
        return owner, ''

def _reserve_running(base):
    """在启动线程前原子占位；快速双击/并发请求只能有一个派发者成功。"""
    return _reserve_running_reason(base)[0]

def _release_running(base, owner):
    """只允许预约本轮的 owner 释放；迟到的旧 worker 永远碰不到新一轮。"""
    with RUNNING_LOCK:
        if not owner or RUNNING.get(base) != owner: return False
        if CANCEL.get(base) == owner: CANCEL.pop(base, None)
        if TERMINAL_OWNERS.get(base) == owner: TERMINAL_OWNERS.pop(base, None)
        RUNNING_THREADS.pop(base, None)
        del RUNNING[base]
        return True

def _is_running(base):
    with RUNNING_LOCK: return base in RUNNING

def _has_running():
    with RUNNING_LOCK: return bool(RUNNING)

def _running_snapshot():
    with RUNNING_LOCK: return tuple(RUNNING)

def _running_owner(base):
    with RUNNING_LOCK: return RUNNING.get(base)

def _begin_job_control(base):
    """给 stop/delete/pause 加 tombstone，避免旧轮结束后新轮在控制动作中间抢跑。"""
    with RUNNING_LOCK:
        # replay 会读持久会话与任务证据；退出提交后也不允许再开始文件破坏性操作。
        if SHUTTING_DOWN or EXITING or OC_REPLAYING or base in JOB_CONTROL:
            return None, None
        token = secrets.token_hex(16)
        JOB_CONTROL[base] = token
        return token, RUNNING.get(base)

def _end_job_control(base, token):
    with RUNNING_LOCK:
        if token and JOB_CONTROL.get(base) == token:
            JOB_CONTROL.pop(base, None)

def _request_cancel(base, owner):
    with RUNNING_LOCK:
        if not owner or RUNNING.get(base) != owner: return False
        # 完成态提交与取消请求共用这把锁：谁先拿到锁谁成为这一代唯一终态。
        if TERMINAL_OWNERS.get(base) == owner: return False
        CANCEL[base] = owner
        return True

def _cancel_requested(base, owner=None):
    with RUNNING_LOCK:
        current = RUNNING.get(base)
        expected = owner or current
        return bool(current and expected == current and CANCEL.get(base) == expected)

def _owner_running(base, owner):
    with RUNNING_LOCK: return bool(owner and RUNNING.get(base) == owner)

def _commit_done(job, word):
    """原子提交 done；若 stop/pause 已拿到本代取消权，绝不再覆盖为完成。"""
    base = os.path.basename(job)
    with RUNNING_LOCK:
        owner = RUNNING.get(base)
        if owner and CANCEL.get(base) == owner:
            return False
        if owner:
            TERMINAL_OWNERS[base] = owner
        try:
            write_json(os.path.join(job, 'outcome.json'),
                       {'state': 'done', 'word': word, 'ts': now()})
            emit(job, {'type': 'progress', 'stage': '完成', 'pct': 100, 'step': 12,
                       'total': 12, 'terminal': True, 'verified': True})
        except Exception:
            if owner and TERMINAL_OWNERS.get(base) == owner:
                TERMINAL_OWNERS.pop(base, None)
            raise
        return True

def _register_proc(base, proc):
    with RUNNING_LOCK:
        PROCS[base] = proc
        PROC_OWNERS[base] = RUNNING.get(base)

def _take_proc(base, owner):
    """stop/delete 只拿走同一 owner 的 CLI，绝不误杀后来一轮。"""
    with RUNNING_LOCK:
        if PROC_OWNERS.get(base) != owner: return None
        PROC_OWNERS.pop(base, None)
        return PROCS.pop(base, None)

def _clear_proc(base, proc):
    with RUNNING_LOCK:
        if PROCS.get(base) is proc:
            PROCS.pop(base, None)
            PROC_OWNERS.pop(base, None)

def _begin_oc_replay():
    global OC_REPLAYING
    with RUNNING_LOCK:
        if SHUTTING_DOWN or EXITING or OC_REPLAYING or RUNNING or JOB_CONTROL: return False
        OC_REPLAYING = True
        return True

def _end_oc_replay():
    global OC_REPLAYING
    with RUNNING_LOCK: OC_REPLAYING = False

def _run_reserved_worker(base, owner, target, args):
    try:
        target(*args)
    finally:
        # owner 的生命周期覆盖 fallback、收拢产物、补 Word、settle 与质检；只在全部结束后释放。
        _release_running(base, owner)

def _start_reserved_worker(base, owner, target, *args):
    thread = threading.Thread(target=_run_reserved_worker,
                              args=(base, owner, target, args), daemon=True)
    with RUNNING_LOCK:
        if RUNNING.get(base) == owner:
            RUNNING_THREADS[base] = thread   # 登记本代线程,供僵尸占位自愈判断
    thread.start()
    return thread

# GUI 方式启动的 App(Finder/Dock 双击)继承不到 shell 的 PATH,Homebrew/npm 装的 CLI 会"明明装了却找不到"。
# 解决:按增补 PATH 查找 + 扫常见安装位置,拿绝对路径执行。
EXTRA_BIN = [os.path.expanduser(p) for p in (
    '/opt/homebrew/bin', '/usr/local/bin', '~/.local/bin', '~/bin', '~/.npm-global/bin',
    '~/.claude/local', '~/.codex/bin', '~/.nvm/current/bin', '/opt/node22/bin')]
if os.name == 'nt':
    EXTRA_BIN += [os.path.join(os.environ.get('APPDATA', ''), 'npm'),
                  os.path.join(os.environ.get('LOCALAPPDATA', ''), 'Programs', 'claude')]

def aug_path_env():
    env = {**os.environ, **shell_env_snapshot()}
    env['PATH'] = os.pathsep.join([p for p in EXTRA_BIN if os.path.isdir(p)]
                                  + [env.get('PATH', ''), os.environ.get('PATH', '')])
    return env

# SoWork(商汤)桌面端自带的 openclaw CLI:装在 App 包内,不在 PATH 上,得按包路径找
# 子进程与本会话解绑:POSIX 用新会话;Windows 不开控制台窗口
DETACH = {'start_new_session': True} if os.name != 'nt' else {'creationflags': 0x08000000}
DETACHED_CHILDREN = {}
DETACHED_CHILDREN_LOCK = threading.Lock()

SOWORK_GLOBS = [
    # 应用名可能是「商汤输入法SoWork.app」「SoWork.app」等,CLI 也可能不在 Resources/cli 下,
    # 所以既按常见路径直取,也在包内递归找一层
    '/Applications/*SoWork*.app/Contents/Resources/cli/openclaw',
    '/Applications/*商汤*.app/Contents/Resources/cli/openclaw',
    os.path.expanduser('~/Applications/*SoWork*.app/Contents/Resources/cli/openclaw'),
    os.path.expanduser('~/Applications/*商汤*.app/Contents/Resources/cli/openclaw'),
    '/Applications/*SoWork*.app/Contents/**/openclaw',
    '/Applications/*商汤*.app/Contents/**/openclaw',
    os.path.expanduser('~/Applications/*SoWork*.app/Contents/**/openclaw'),
    os.path.expanduser('~/Applications/*商汤*.app/Contents/**/openclaw'),
    os.path.join(os.environ.get('LOCALAPPDATA', '') or '_', 'Programs', '**', 'openclaw.exe'),
    os.path.join(os.environ.get('APPDATA', '') or '_', '**', 'openclaw.exe'),
]

def bundled_cli(base):
    """安装包内置的执行外壳(Tauri externalBin,打包后去掉平台三元组后缀,与引擎同目录):
    优先级最高的"客户免装任何东西"来源;其次是「一键安装」下载到数据目录的那份(版本钉过、测过兼容)"""
    exe = base + ('.exe' if os.name == 'nt' else '')
    dirs = [os.path.dirname(os.path.abspath(sys.executable)),   # PyInstaller sidecar:与主程序同目录
            os.path.dirname(os.path.abspath(sys.argv[0] or '.')), HERE]
    for d in dirs:
        c = os.path.join(d, exe)
        if os.path.isfile(c): return c
    c = os.path.join(DATA, 'bin', exe)                          # 「一键安装」落点
    return c if os.path.isfile(c) else None

def bundled_codex():
    return bundled_cli('codex-cli')

CLI_FAMILY = {'codex': ('codex',), 'claude': ('claude',), 'opencode': ('opencode',),
              'sowork': ('openclaw', 'sowork'), 'openclaw': ('openclaw', 'sowork')}

def resolve_cli(name, eng=None):
    """找 CLI 的绝对路径:显式配置(仅当文件名与工具同族)> 内置/一键安装 > 增补 PATH > 常见安装位置。
    显式路径带族校验的原因:「CLI 路径」是共享字段,切换引擎后旧路径会指向另一个工具——
    曾把 codex 二进制当 opencode 跑出 unexpected argument '--auto'(真机事故)。名字对不上就当没填。"""
    if eng and (eng.get('cli_path') or '').strip():
        p = os.path.expanduser(eng['cli_path'].strip())
        fam = CLI_FAMILY.get(name)
        base = os.path.basename(p).lower()
        if os.path.isfile(p) and (not fam or any(k in base for k in fam)):
            return p
    if name == 'codex':
        b = bundled_cli('codex-cli')
        if b: return b
    if name == 'opencode':
        b = bundled_cli('opencode-cli')
        if b: return b
    if name in ('sowork', 'openclaw'):
        for pat in SOWORK_GLOBS:
            if not pat: continue
            hits = sorted(glob.glob(pat, recursive=True))
            if hits: return hits[0]
        name = 'openclaw'
    p = shutil.which(name, path=aug_path_env()['PATH'])
    if p: return p
    exe = name + ('.exe' if os.name == 'nt' else '')
    for d in EXTRA_BIN:
        c = os.path.join(d, exe)
        if os.path.isfile(c): return c
    return None

_SHELL_ENV = {}
_SHELL_ENV_DONE = [False]
SKIP_ENV_KEYS = {'_', 'PWD', 'OLDPWD', 'SHLVL', 'PS1', 'TERM', 'COLUMNS', 'LINES'}

def shell_env_snapshot():
    """GUI 双击启动的 App 不继承终端环境(PATH、CLI 网关配置等全丢)。
    跑一次「登录+交互」shell 抓取用户真实环境变量再合并进子进程——
    比把命令塞进 shell 里执行更稳:能读到 .zshrc,又不会把 rc 的输出混进 agent 结果。"""
    if _SHELL_ENV_DONE[0] or os.name == 'nt': return _SHELL_ENV
    _SHELL_ENV_DONE[0] = True
    sh = os.environ.get('SHELL') or '/bin/zsh'
    if not os.path.isfile(sh): sh = '/bin/bash'
    for args in ([sh, '-lic', 'env'], [sh, '-lc', 'env']):     # 交互式读不到就退回登录式
        try:
            r = subprocess.run(args, capture_output=True, text=True, timeout=15, stdin=subprocess.DEVNULL)
            for ln in (r.stdout or '').splitlines():
                if '=' not in ln: continue
                k, v = ln.split('=', 1)
                if k.isidentifier() and k not in SKIP_ENV_KEYS: _SHELL_ENV[k] = v
            if _SHELL_ENV: break
        except Exception: pass
    return _SHELL_ENV

def login_shell_wrap(cmd, eng):
    """环境已由 shell_env_snapshot 继承,命令直接执行即可(保留此函数供调用点统一)"""
    return cmd

def agent_env(eng=None):
    """子进程环境 = 应用环境 < 登录 shell 快照 < 增补 PATH < 用户填的附加环境变量(优先级从低到高)"""
    eng = eng or {}
    env = {**os.environ}
    if eng.get('login_shell') is not False:
        env.update(shell_env_snapshot())          # 终端里有、App 里没有的,以终端为准
    paths = [p for p in EXTRA_BIN if os.path.isdir(p)]
    env['PATH'] = os.pathsep.join(paths + [env.get('PATH', ''), os.environ.get('PATH', '')])
    # 「自动(默认)」= s2 = OpenCode,和显式选 opencode 走同一套环境:隔离 XDG + 生成 provider 配置
    # + 直通端点口令。**这里漏一个 kind,外壳就拿不到我们的 provider**,表现是它报
    # 「UnknownError / Unexpected server error」——完全看不出是配置没送到(真机踩过)。
    if (eng.get('kind') or '') in ('s2', 'opencode'):
        # 出错不再静默吞掉:配置没送到时 opencode 只会甩一句
        # 「UnknownError / Unexpected server error」,完全看不出是我们这边没生成配置。
        # 把真实原因写进环境变量,报错翻译层据此说人话(见 real_agent / agent_test)。
        try: env.update(opencode_env(eng))
        except Exception as e:
            env['BIDDOG_SHELL_CONF_ERR'] = '内置生成组件配置失败:%s' % e
    # 显式选择 codex = 使用用户自己的 Codex CLI 订阅与登录态。这里不能注入中标狗生成的
    # CODEX_HOME / S2 Key，否则「切回 Codex」表面上切了引擎，实际仍在花发放方的额度，
    # 还会覆盖用户原有的 ~/.codex 配置。
    for ln in (eng.get('env') or '').splitlines():
        ln = ln.strip()
        if not ln or ln.startswith('#') or '=' not in ln: continue
        k, v = ln.split('=', 1)
        env[k.strip()] = os.path.expanduser(v.strip())         # 用户显式指定的最高优先
    return env

def ensure_line_ts(job):
    """给 agent 自写的、没有 ts 的事件行补时间戳并持久化。
       由 watcher 每几秒调用一次 → 即使没人开着页面,阶段耗时也是真实的。"""
    path = os.path.join(job, 'events.jsonl'); tsf = os.path.join(job, '.line_ts.json')
    if not os.path.isfile(path): return {}
    index = read_json(tsf, {}); dirty = False
    try: lines = open(path, encoding='utf-8').read().splitlines()
    except Exception: return index
    for i, ln in enumerate(lines):
        if str(i) in index: continue
        try: e = json.loads(ln)
        except Exception: continue
        if e.get('ts'): continue
        index[str(i)] = now(); dirty = True
    if dirty:
        try: write_json(tsf, index)
        except Exception: pass
    return index

def merged_materials(job):
    """本次任务真正该用的素材目录。

    以前是二选一:`mdir if os.path.isdir(mdir) else assets_dir()` ——
    只要用户在新建任务向导里拖进**任何一个**公司文件,就建出了 job/素材/,
    于是**整个全局素材库(公司介绍、产品能力表、资质与案例、应答要点、图片索引)
    在这一单里全部失联**,agent 满篇写〔需补充〕。
    用户越认真传材料,产出越差 —— 这条极其反直觉,而且藏在一行代码里。

    现在改成合并:把全局库里任务级没有的文件补进 job/素材/,同名一律以任务级为准
    (这一单专门传的东西优先)。合并只补不覆盖,可重复调用。"""
    mdir = os.path.join(job, '素材')
    glob_dir = assets_dir()
    if not os.path.isdir(mdir):
        return glob_dir if os.path.isdir(glob_dir) else mdir
    if os.path.isdir(glob_dir) and os.path.realpath(glob_dir) != os.path.realpath(mdir):
        for root, dirs, files in os.walk(glob_dir):
            # 原始标书/ 是入库存档的大二进制原件,章节内容早已拆进 章节模板/;
            # 整目录拷进每个任务只会把素材推向节点上限,生成用不到它
            dirs[:] = [d for d in dirs if not d.startswith(('.', '_')) and d not in ('原件', '原始标书')]
            rel = os.path.relpath(root, glob_dir)
            dst_root = mdir if rel == '.' else os.path.join(mdir, rel)
            for fn in files:
                if fn.startswith(('.', '_')) or fn == '入库流水.jsonl': continue
                dst = os.path.join(dst_root, fn)
                if os.path.exists(dst): continue          # 任务级同名文件优先,不覆盖
                try:
                    os.makedirs(dst_root, exist_ok=True)
                    shutil.copy2(os.path.join(root, fn), dst)
                except Exception: pass
    return mdir

def _backflow_job_materials(job, rels, tag):
    """把本单导入的素材回流到全局素材库 任务导入/<任务名>/(反馈:任务材料应自动落库)。

    md5 对照 入库流水.jsonl 去重——同一份文件多单复用不会堆副本;图片顺带进 图片/
    并登记占位索引,与「AI 识图打标」共用一条链路。回流失败只记流水,不影响任务。"""
    lib = assets_dir()
    log_path = os.path.join(lib, '入库流水.jsonl')
    seen = set()
    try:
        for ln in open(log_path, encoding='utf-8', errors='ignore').read().splitlines():
            try: seen.add(json.loads(ln).get('md5'))
            except Exception: continue
    except OSError:
        pass
    safe_tag = (''.join(c for c in str(tag) if c not in '\\/:*?"<>|').strip() or '未命名任务')[:20]
    copied, imgs = 0, []
    for rel in rels[:200]:
        src = os.path.join(job, '素材', rel)
        if not os.path.isfile(src): continue
        try:
            with open(src, 'rb') as fh: blob = fh.read()
        except OSError: continue
        digest = hashlib.md5(blob).hexdigest()
        if digest in seen: continue
        seen.add(digest)
        ext = os.path.splitext(rel)[1].lower()
        try:
            if ext in ('.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp'):
                d = os.path.join(lib, '图片'); os.makedirs(d, exist_ok=True)
                fn2 = os.path.basename(rel)
                dst = os.path.join(d, fn2)
                if not os.path.exists(dst):
                    shutil.copy2(src, dst); imgs.append(fn2)
            else:
                d = os.path.join(lib, '任务导入', safe_tag)
                os.makedirs(d, exist_ok=True)
                dst = os.path.join(d, os.path.basename(rel))
                if not os.path.exists(dst): shutil.copy2(src, dst)
            copied += 1
            with open(log_path, 'a', encoding='utf-8') as lf:
                lf.write(json.dumps({'ts': now(), 'source': '任务导入:' + os.path.basename(rel),
                                     'md5': digest, 'from_job': os.path.basename(job),
                                     'sections': 0, 'images': 1 if ext in ('.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp') else 0,
                                     'categories': {}}, ensure_ascii=False) + '\n')
        except Exception:
            continue
    if imgs:
        try:
            entries = read_json(os.path.join(lib, '图片索引.json'), [])
            have = {e.get('file') for e in entries}
            for f in imgs:
                if f not in have:
                    entries.append({'id': os.path.splitext(f)[0], 'file': f, 'caption': '〔待AI识图〕',
                                    'category': '未识别', 'ocr': '', 'keywords': '来自任务导入', 'anchor': ''})
            write_image_index(entries)
        except Exception:
            pass

def _skill_module(name):
    """从技能包 references/ 里加载确定性脚本(quality_gate / build_tender_docx 等)。
    单一事实来源:引擎不复制这些逻辑,升级技能包 = 引擎门禁同步升级。"""
    sd = skill_dir_conf()
    if not os.path.isfile(os.path.join(sd, 'SKILL.md')): sd = ensure_skill()
    path = os.path.join(sd, 'references', name + '.py')
    if not os.path.isfile(path): return None
    try:
        import importlib.util
        refs = os.path.dirname(path)
        if refs not in sys.path: sys.path.insert(0, refs)   # 脚本之间互相 import(tender_images 等)
        spec = importlib.util.spec_from_file_location('skillref_' + name, path)
        mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)
        return mod
    except Exception:
        return None

def _job_find(job, *pats):
    """在任务目录(含一层子目录)找第一个命中文件"""
    for base in [job] + [os.path.join(job, d) for d in sorted(os.listdir(job))
                         if os.path.isdir(os.path.join(job, d)) and not d.startswith(('.', '_'))]:
        for pat in pats:
            hits = sorted(glob.glob(os.path.join(base, pat)))
            if hits: return hits[0]
    return None

# 收拢产物时跳过的目录。「参考资料」是客户自己上传的输入件,收上来会被当成我们的产出,
# 既污染交付清单,又让「跑完却没产出」的告警失灵——和 NOT_DELIVERABLE 是同一个坑。
HARVEST_SKIP = {'素材', '章节', '参考资料', 'materials', 'node_modules', '__pycache__',
                'venv', '.venv', 'bid-multiagent-tao', 'references', 'images', '图片', 'assets'}

def harvest(job, depth=3):
    """把写进子目录的交付物收回任务根目录。

    AGENT_PROMPT 第 1 条写着「所有产物直接写入 outDir,不要另建下层输出目录」,
    但模型照建不误——`outDir/投标文件/技术标.docx` 这种。以前只扫一层、
    而且文件名必须以「投标」开头,漏掉的就等于没产出:任务跑完客户一份标书都看不到
    (完成播报里列的「② 产物写到别处了」说的就是这个,只是当时没真的去捞)。
    现在按目录深度递归,凡是交付扩展名的文件都往上收。"""
    moved = []
    def walk(d, lv):
        if lv > depth: return
        try: entries = sorted(os.listdir(d))
        except Exception: return
        for name in entries:
            p = os.path.join(d, name)
            if os.path.isdir(p):
                if name in HARVEST_SKIP or name.startswith(('.', '_')): continue
                walk(p, lv + 1)
            elif name.endswith(DELIVER_EXT) and not name.startswith(('_', '.')) \
                    and name not in NOT_DELIVERABLE:
                dst = os.path.join(job, name)
                if os.path.exists(dst): continue     # 根目录已有同名:根目录那份才是终稿,不覆盖
                try: shutil.move(p, dst); moved.append(name)
                except Exception: pass
    for name in sorted(os.listdir(job)):
        p = os.path.join(job, name)
        if os.path.isdir(p) and name not in HARVEST_SKIP and not name.startswith(('.', '_')):
            walk(p, 1)
    return moved

# 正文稿识别:这些关键词一出现就不是正文,是报告/分析/中间件
_NOT_BODY = ('自检', '清洗', '报告', '质检', '门禁', '矩阵', '偏离表', '解析版',
             '配图清单', '补料', '废标', '组成', '索引', '格式要求', '大纲', '定向重做说明', '.bak')

_CHAPTER_FILE_RE = re.compile(
    r'^(?:第[0-9一二三四五六七八九十百]+(?:[-—至到][0-9一二三四五六七八九十百]+)?章|章节[_\-\s]*[0-9一二三四五六七八九十百]+)'
)

def _is_chapter_fragment(name):
    return bool(_CHAPTER_FILE_RE.match(os.path.basename(str(name or ''))))

def _chapter_mds(job, known=None):
    names = list(known) if known is not None else list_deliverables(job)
    return [fn for fn in sorted(names)
            if fn.lower().endswith('.md') and _is_chapter_fragment(fn)
            and os.path.isfile(os.path.join(job, fn))]

def _body_mds(job, known=None):
    """挑出「可以拿去导出 Word 的正文稿」。

    原来只认 fn.startswith('投标'),模型把正文命名成《技术标.md》《XX项目标书.md》时
    全链路都当它不存在:质检退回旧门禁、修复稿不重建 docx、缺 Word 也不告警。
    现在放宽到常见命名,再兜底"最大的那份 md",宁可多认一份也不要漏掉正文。"""
    names = list(known) if known is not None else list_deliverables(job)
    mds = [fn for fn in sorted(names)
           if fn.endswith('.md') and not _is_chapter_fragment(fn)
           and not any(k in fn for k in _NOT_BODY)]
    hit = [fn for fn in mds if fn.startswith('投标')]
    if not hit:
        hit = [fn for fn in mds if any(k in fn for k in ('投标文件', '技术标', '商务标', '标书', '方案'))]
    if not hit and mds:
        # 还是认不出来:取最大的一份。门槛 1200 字节 ≈ 400 汉字——再小的多半是清单/说明而不是正文,
        # 硬导出只会给客户一份空壳 Word;报告清单类文件名前面已经排除过了,这里的误判风险很低。
        big = max(mds, key=lambda f: os.path.getsize(os.path.join(job, f)))
        if os.path.getsize(os.path.join(job, big)) > 1200: hit = [big]
    return hit

def _body_docxs(job, known=None):
    """只认正文 Word；招标原件和分析/报告类 docx 都不能把任务推成完成。"""
    names = list(known) if known is not None else list_deliverables(job)
    out = []
    for fn in sorted(names):
        if not fn.lower().endswith('.docx') or _is_chapter_fragment(fn) or any(k in fn for k in _NOT_BODY): continue
        if not any(k in fn for k in ('投标', '技术标', '商务标', '标书', '方案', '响应文件')): continue
        if _valid_docx(os.path.join(job, fn)): out.append(fn)
    return out

def _valid_docx(path, min_text=100):
    """校验正文 Word 的 ZIP 结构和可见文字，拒绝空壳、截断包和仅改扩展名的伪文件。"""
    try:
        if not os.path.isfile(path) or os.path.getsize(path) < 1024 or not zipfile.is_zipfile(path):
            return False
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            if '[Content_Types].xml' not in names or 'word/document.xml' not in names:
                return False
            info = archive.getinfo('word/document.xml')
            if info.file_size <= 0 or info.file_size > 20 * 1024 * 1024:
                return False
            root = ET.fromstring(archive.read(info))
        tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
        visible = ''.join(node.text or '' for node in root.iter(tag))
        return len(re.sub(r'\s+', '', visible)) >= int(min_text)
    except (OSError, KeyError, zipfile.BadZipFile, ET.ParseError, RuntimeError, ValueError):
        return False

def _docx_has_toc(path):
    """A visible heading called 目录 is not enough; require a real Word TOC field."""
    try:
        with zipfile.ZipFile(path) as archive:
            root = ET.fromstring(archive.read('word/document.xml'))
        instructions = ''.join(node.text or '' for node in root.iter()
                               if str(node.tag).endswith('}instrText'))
        return bool(re.search(r'\bTOC\b', instructions, re.I))
    except (OSError, KeyError, zipfile.BadZipFile, ET.ParseError):
        return False

def _tabular_rows(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == '.xlsx':
            with zipfile.ZipFile(path) as archive:
                rows = 0
                for name in archive.namelist():
                    if name.startswith('xl/worksheets/sheet') and name.endswith('.xml'):
                        root = ET.fromstring(archive.read(name))
                        rows = max(rows, sum(1 for node in root.iter() if str(node.tag).endswith('}row')) - 1)
                return max(0, rows)
        if ext == '.docx':
            with zipfile.ZipFile(path) as archive:
                root = ET.fromstring(archive.read('word/document.xml'))
            return max(0, sum(1 for node in root.iter() if str(node.tag).endswith('}tr')) - 1)
        text = open(path, encoding='utf-8', errors='ignore').read()
        table = [line for line in text.splitlines() if line.strip().startswith('|')
                 and not re.fullmatch(r'[|:\-\s]+', line.strip())]
        return max(0, len(table) - 1)
    except (OSError, KeyError, zipfile.BadZipFile, ET.ParseError):
        return 0

def _deviation_item(job, needles):
    names = [name for name in list_deliverables(job) if any(needle in name for needle in needles)]
    rows = sum(_tabular_rows(os.path.join(job, name)) for name in names)
    return {'present': bool(names), 'files': names, 'rows': rows,
            'status': 'pass' if names and rows > 0 else 'fail'}

def _quality_result_from_disk(job, signature=''):
    saved = read_json(os.path.join(job, 'delivery.json'), {})
    quality = saved.get('quality') if isinstance(saved, dict) else None
    # A quality verdict only applies to the exact set of deliverables it checked.
    # If Word/report/table bytes changed afterwards, never carry an old green result
    # into the new delivery summary; parse the current report or fall back to unknown.
    if (isinstance(quality, dict) and quality.get('status') and signature
            and saved.get('quality_signature') == signature):
        return quality
    report = next((name for name in list_deliverables(job) if '成品质检报告' in name), '')
    if not report:
        return {'status': 'unknown', 'level': 'unknown', 'summary': '关键检查结果尚未生成'}
    try: text = open(os.path.join(job, report), encoding='utf-8', errors='ignore').read()
    except OSError: text = ''
    if '🔴' in text or '未通过' in text:
        return {'status': 'fail', 'level': 'red', 'summary': '关键检查有必须处理项', 'report': report}
    if '🟡' in text:
        return {'status': 'warning', 'level': 'yellow', 'summary': '关键检查有建议确认项', 'report': report}
    return {'status': 'pass', 'level': 'green', 'summary': '关键检查已通过', 'report': report}

_DELIVERY_DIGEST_CACHE = {}
_DELIVERY_DIGEST_LOCK = threading.Lock()
_DELIVERY_DIGEST_LAST_PRUNE = 0.0
_DELIVERY_DIGEST_PRUNE_QUEUE = deque()
_DELIVERY_DIGEST_GENERATIONS = {}
_DELIVERY_DIGEST_GENERATION = 0

def _remember_delivery_digest(path, identity, digest):
    global _DELIVERY_DIGEST_LAST_PRUNE, _DELIVERY_DIGEST_GENERATION
    prune_candidates = []
    with _DELIVERY_DIGEST_LOCK:
        # Every changed digest receives a new generation. Older queued tokens then
        # become harmless tombstones and cannot delete a recreated path.
        _DELIVERY_DIGEST_GENERATION += 1
        generation = _DELIVERY_DIGEST_GENERATION
        _DELIVERY_DIGEST_GENERATIONS[path] = generation
        _DELIVERY_DIGEST_PRUNE_QUEUE.append((path, generation))
        _DELIVERY_DIGEST_CACHE[path] = (identity, digest)
        current = time.monotonic()
        if (len(_DELIVERY_DIGEST_CACHE) > 4096
                and current - _DELIVERY_DIGEST_LAST_PRUNE >= 60):
            _DELIVERY_DIGEST_LAST_PRUNE = current
            scanned = 0
            # Hard-bound total queue work, not only valid candidates: a backlog of
            # obsolete generation tokens must not monopolize the global lock.
            while _DELIVERY_DIGEST_PRUNE_QUEUE and scanned < 256:
                cached_path, generation = _DELIVERY_DIGEST_PRUNE_QUEUE.popleft()
                scanned += 1
                if (_DELIVERY_DIGEST_GENERATIONS.get(cached_path) != generation
                        or cached_path not in _DELIVERY_DIGEST_CACHE):
                    continue
                prune_candidates.append((cached_path, generation))
                _DELIVERY_DIGEST_PRUNE_QUEUE.append((cached_path, generation))
    # Filesystem calls are deliberately outside the global cache lock.
    stale = [(cached_path, generation) for cached_path, generation in prune_candidates
             if not os.path.exists(cached_path)]
    if stale:
        with _DELIVERY_DIGEST_LOCK:
            for cached_path, generation in stale:
                if _DELIVERY_DIGEST_GENERATIONS.get(cached_path) != generation:
                    continue
                _DELIVERY_DIGEST_CACHE.pop(cached_path, None)
                _DELIVERY_DIGEST_GENERATIONS.pop(cached_path, None)

def _windows_change_time(path):
    """Read NTFS ChangeTime; Python st_ctime is still creation time on Windows."""
    import ctypes
    from ctypes import wintypes

    class FILE_BASIC_INFO(ctypes.Structure):
        _fields_ = [
            ('CreationTime', ctypes.c_longlong), ('LastAccessTime', ctypes.c_longlong),
            ('LastWriteTime', ctypes.c_longlong), ('ChangeTime', ctypes.c_longlong),
            ('FileAttributes', wintypes.DWORD),
        ]

    kernel32 = ctypes.WinDLL('kernel32', use_last_error=True)
    create_file = kernel32.CreateFileW
    create_file.argtypes = [wintypes.LPCWSTR, wintypes.DWORD, wintypes.DWORD,
                            wintypes.LPVOID, wintypes.DWORD, wintypes.DWORD,
                            wintypes.HANDLE]
    create_file.restype = wintypes.HANDLE
    handle = create_file(path, 0, 0x00000007, None, 3, 0, None)
    invalid = wintypes.HANDLE(-1).value
    if handle == invalid:
        raise ctypes.WinError(ctypes.get_last_error())
    try:
        info = FILE_BASIC_INFO()
        get_info = kernel32.GetFileInformationByHandleEx
        get_info.argtypes = [wintypes.HANDLE, ctypes.c_int, wintypes.LPVOID, wintypes.DWORD]
        get_info.restype = wintypes.BOOL
        if not get_info(handle, 0, ctypes.byref(info), ctypes.sizeof(info)):
            raise ctypes.WinError(ctypes.get_last_error())
        return int(info.ChangeTime)
    finally:
        kernel32.CloseHandle(handle)


def _delivery_change_marker(path, stat):
    if os.name != 'nt':
        return int(getattr(stat, 'st_ctime_ns', stat.st_ctime * 1e9)), True
    try:
        return _windows_change_time(path), True
    except (OSError, AttributeError, ImportError):
        # If native metadata is unavailable, correctness wins over caching.
        return 0, False


def _delivery_file_fingerprint(path):
    """Return a content identity without rereading unchanged large artifacts every poll."""
    absolute = os.path.abspath(path)
    for _ in range(2):
        before = os.stat(absolute)
        before_change, cacheable = _delivery_change_marker(absolute, before)
        identity = (
            int(getattr(before, 'st_dev', 0)), int(getattr(before, 'st_ino', 0)),
            int(before.st_size),
            int(getattr(before, 'st_mtime_ns', before.st_mtime * 1e9)),
            before_change,
        )
        if cacheable:
            with _DELIVERY_DIGEST_LOCK:
                cached = _DELIVERY_DIGEST_CACHE.get(absolute)
                if cached and cached[0] == identity:
                    return identity, cached[1]
        digest = _file_digest(absolute)
        after = os.stat(absolute)
        after_change, after_cacheable = _delivery_change_marker(absolute, after)
        after_identity = (
            int(getattr(after, 'st_dev', 0)), int(getattr(after, 'st_ino', 0)),
            int(after.st_size),
            int(getattr(after, 'st_mtime_ns', after.st_mtime * 1e9)),
            after_change,
        )
        if identity != after_identity:
            continue
        if cacheable and after_cacheable:
            # Keep live entries so fixed-order task scans keep hitting the cache. Missing
            # paths are reclaimed through a bounded rotating queue without starvation.
            _remember_delivery_digest(absolute, identity, digest)
        return identity, digest
    raise OSError('交付文件在校验期间持续变化')


def _delivery_signature(job):
    rows = []
    for name in list_deliverables(job):
        path = os.path.join(job, name)
        try:
            identity, content_digest = _delivery_file_fingerprint(path)
            # ready 同时依赖 Word、格式报告、偏离表与品质报告；所有交付物都绑定
            # 内容身份，避免同大小/同 mtime 替换绕过门禁。
            rows.append((name,) + identity + (content_digest,))
        except OSError:
            rows.append((name, -1, -1, -1, -1, -1, ''))
    return hashlib.sha256(json.dumps(rows, ensure_ascii=False, separators=(',', ':')).encode()).hexdigest()

def delivery_summary(job, quality=None):
    signature = _delivery_signature(job)
    cache_path = os.path.join(job, 'delivery.json')
    cached = read_json(cache_path, {})
    pipeline_state = generation_pipeline.load(job)
    format_required = pipeline_state.get('version') == generation_pipeline.PIPELINE_VERSION
    cached_summary = cached.get('summary') if isinstance(cached, dict) else None
    if (quality is None and isinstance(cached, dict) and cached.get('signature') == signature
            and isinstance(cached_summary, dict)
            and (not format_required or isinstance(cached_summary.get('format'), dict))):
        return cached_summary
    words = _body_docxs(job)
    primary = words[0] if words else ''
    jid = urllib.parse.quote(os.path.basename(job), safe='')
    word = {'present': bool(primary), 'name': primary,
            'url': ('/v1/jobs/%s/artifacts/%s' % (jid, urllib.parse.quote(primary, safe='')) if primary else '')}
    has_toc = bool(primary and _docx_has_toc(os.path.join(job, primary)))
    toc = {'status': ('pass' if has_toc else 'fail'), 'present': has_toc}
    technical = _deviation_item(job, ('技术应答偏离表', '技术偏离表'))
    business = _deviation_item(job, ('商务偏离表',))
    deviations = {'status': 'pass' if technical['status'] == business['status'] == 'pass' else 'fail',
                  'technical': technical, 'business': business,
                  'total_rows': technical['rows'] + business['rows']}
    quality = quality if isinstance(quality, dict) else _quality_result_from_disk(job, signature)
    word_format = (_word_format_status(job) if format_required else
                   {'status': 'not_required', 'present': False, 'report': ''})
    if format_required and word['present'] and word_format.get('status') in ('stale', 'missing'):
        # 质检/手动重试重建过 Word 后旧报告过期(或报告丢失)——Word 就在眼前,
        # 当场重跑一次格式检查即可,不能因为报告旧就把整单判红(真机「生成了 Word 还报错」路径之一)。
        word_format_audit(job, primary)
        word_format = _word_format_status(job, primary)
    format_ok = (not format_required) or word_format['status'] in ('pass', 'warn')
    components = [word['present'], toc['status'] == 'pass', deviations['status'] == 'pass',
                  quality.get('status') in ('pass', 'warning'), format_ok]
    ready = all(components)
    if (not word['present'] or toc['status'] == 'fail' or deviations['status'] == 'fail'
            or quality.get('status') == 'fail' or not format_ok):
        check_status = 'fail'
    elif quality.get('status') in ('unknown', 'warning'):
        check_status = quality.get('status')
    elif format_required and word_format['status'] == 'warn':
        check_status = 'warning'
    else:
        check_status = 'pass'
    format_note = ('；Word 格式自检有不符项，提交前请按报告核对'
                   if format_required and word_format['status'] == 'warn' else '')
    checks = {'status': check_status, 'level': quality.get('level') or ('green' if ready else 'red'),
              'summary': (quality.get('summary') or ('关键检查已通过' if ready else '关键检查仍有待确认项')) + format_note}
    summary = {'word': word, 'toc': toc, 'deviations': deviations, 'format': word_format,
               'checks': checks, 'ready': ready}
    try:
        patch_json(cache_path, {'signature': signature, 'summary': summary,
                                'quality': quality, 'quality_signature': signature,
                                'updated_at': now()})
    except OSError:
        pass
    return summary

def _file_digest(path):
    try:
        h = hashlib.sha256()
        with open(path, 'rb') as source:
            for chunk in iter(lambda: source.read(1024 * 1024), b''): h.update(chunk)
        return h.hexdigest()
    except OSError:
        return ''

def _named_files(job, needles, suffix=None, min_size=1):
    try: names = os.listdir(job)
    except OSError: return []
    hits = []
    for fn in names:
        if suffix and not fn.lower().endswith(suffix): continue
        if not any(n in fn for n in needles): continue
        try:
            if os.path.isfile(os.path.join(job, fn)) and os.path.getsize(os.path.join(job, fn)) >= min_size:
                hits.append(fn)
        except OSError:
            pass
    return hits

def evidence_for_step(job, step):
    """返回某一步的磁盘证据是否成立；不采信模型文本或百分比。"""
    meta = read_json(os.path.join(job, '任务.json'), {})
    names = list_deliverables(job)
    if step == 1:
        tender = meta.get('tender')
        return bool((tender and os.path.isfile(os.path.join(job, tender))) or
                    _named_files(job, ('解析版',), '.md', 20))
    if step == 2:
        # 旧任务可能没有单独写出“无图片”证据，却已产出下一步的组成分析。
        # 后续结构证据成立时，图片盘点不应永久卡住整条进度链。
        return bool(_named_files(job, ('图片索引', '无图片', '图片检查'), min_size=10) or
                    _named_files(job, ('投标文件组成', '组成分析'), '.md', 500))
    if step == 3:
        return bool(_named_files(job, ('投标文件组成', '组成分析'), '.md', 500))
    if step == 4:
        spec = _job_find(job, 'word_format_spec.json')
        valid = bool(spec and isinstance(read_json(spec, None), dict))
        return valid and bool(_named_files(job, ('格式', '版式'), '.md', 20))
    if step == 5:
        return bool(_named_files(job, ('评分点响应矩阵', '评分矩阵'), min_size=20) and
                    _named_files(job, ('废标风险',), min_size=20))
    if step == 6:
        # “评分点响应矩阵”是第 5 步证据，不能再被模糊匹配成第 6 步。
        response_files = _named_files(job, ('响应矩阵', '响应对照'), '.md', 20)
        return any('评分' not in name for name in response_files)
    if step == 7:
        outline = _named_files(job, ('大纲', '目录'), '.md', 20)
        chapters = _named_files(job, ('第', '章节'), '.md', 100)
        return bool(outline and chapters)
    if step == 8:
        return bool(_named_files(job, ('技术应答偏离表', '技术偏离表'), min_size=20) and
                    _named_files(job, ('商务偏离表',), min_size=20))
    if step == 9: return bool(_body_mds(job, names))
    if step == 10:
        return bool(_named_files(job, ('成品质检报告', '配图复核', '无图'), min_size=20))
    if step == 11: return bool(_named_files(job, ('投标文件自检报告', '自检报告'), '.md', 20))
    if step == 12: return bool(_body_docxs(job, names))
    return False

PIPELINE_STEP_BY_NODE = {'source_parse': 4, 'response_plan': 6, 'assemble': 9,
                         'quality_review': 11, 'word_export': 12, 'delivery_gate': 12}

def _pipeline_checkpoint_step(job):
    """流水线节点状态换算成 12 步进度——这是引擎亲手写盘的权威证据。

    文件启发式(evidence_for_step)找的是智能体模式的产物(图片索引、格式规范等),
    分段流水线不产这些文件,证据链会在缺口处截断,把真实进度钳回「检查资料 1/12」
    (真机反馈:进度显示与实际不符的主要来源)。节点状态本身就是磁盘证据,直接换算。"""
    try:
        state = generation_pipeline.load(job)
    except Exception:
        return 0
    if state.get('version') != generation_pipeline.PIPELINE_VERSION: return 0
    nodes = state.get('nodes') or []
    best = 0
    for node in nodes:
        if node.get('state') != 'done': continue
        nid = str(node.get('id') or '')
        best = max(best, 8 if nid.startswith('chapter_write:') else PIPELINE_STEP_BY_NODE.get(nid, 0))
    running = next((n for n in nodes if n.get('state') in ('running', 'retry_wait')), None)
    if running:
        nid = str(running.get('id') or '')
        best = max(best, 7 if nid.startswith('chapter_write:')
                   else max(0, PIPELINE_STEP_BY_NODE.get(nid, 1) - 1))
    return best

def verified_step(job, claimed=12):
    """阶段必须连续成立；缺第 N 步时，后面的孤立文件不能把进度跳过去。
    流水线模式下节点检查点是更强的证据,取两者较大值。"""
    try: claimed = max(0, min(12, int(claimed or 0)))
    except Exception: claimed = 0
    done = 0
    for step in range(1, claimed + 1):
        if not evidence_for_step(job, step): break
        done = step
    return min(claimed, max(done, _pipeline_checkpoint_step(job)))

def observed_activity_step(job):
    """识别正在发生的最远活动，不把它冒充连续验收通过的 checkpoint。"""
    names = list_deliverables(job)
    if _body_docxs(job, names): return 12
    if _named_files(job, ('投标文件自检报告', '自检报告'), '.md', 20): return 11
    if _named_files(job, ('成品质检报告', '配图复核', '无图'), min_size=20): return 10
    if _body_mds(job, names): return 9
    if (_named_files(job, ('技术应答偏离表', '技术偏离表'), min_size=20) or
            _named_files(job, ('商务偏离表',), min_size=20)): return 8
    if _chapter_mds(job, names): return 7
    response_files = _named_files(job, ('响应矩阵', '响应对照'), '.md', 20)
    if any('评分' not in name for name in response_files): return 6
    if (_named_files(job, ('评分点响应矩阵', '评分矩阵'), min_size=20) or
            _named_files(job, ('废标风险',), min_size=20)): return 5
    if (_job_find(job, 'word_format_spec.json') or
            _named_files(job, ('格式', '版式'), '.md', 20)): return 4
    if _named_files(job, ('投标文件组成', '组成分析'), '.md', 20): return 3
    if _named_files(job, ('图片索引', '无图片', '图片检查'), min_size=10): return 2
    if evidence_for_step(job, 1): return 1
    return 0

def sanitize_event(job, ev):
    """净化历史/第三方写入的进度；verified 字段本身永远不构成信任。"""
    safe = redact(dict(ev or {}))
    if safe.get('type') != 'progress':
        # v0.19.6 及更早版本会把“若干分章稿”误认成“完整正文”，继而把连接中断
        # 写成 Word 导出失败。SSE 回放时纠正旧结论，升级后无需用户重跑才能看懂现状。
        if safe.get('type') == 'error' and _chapter_mds(job) and not _body_mds(job):
            text = str(safe.get('text') or '')
            if '正文稿已经生成' in text or 'Word 导出失败' in text:
                suffix = ('\n\n运行日志' + text.split('\n\n运行日志', 1)[1]
                          if '\n\n运行日志' in text else '')
                count = len(_chapter_mds(job))
                safe['text'] = ('分章撰写中断：已生成 %d 个章节，但还没有汇总成完整正文和最终 Word。'
                                '已经写出的内容均已保留。' % count) + suffix
                meta = read_json(os.path.join(job, '任务.json'), {})
                safe['actions'] = ([{'act': 'resume', 'label': '从已保存内容继续'}]
                                   if meta.get('oc_session') else [{'act': 'rerun', 'label': '重新生成'}])
                safe['actions'].append({'act': 'open_log', 'label': '查看运行日志'})
        return safe
    safe.pop('verified', None)
    try: claimed = max(0, min(12, int(safe.get('step') or 0)))
    except Exception: claimed = 0
    try: pct = int(safe.get('pct') or 0)
    except Exception: pct = 0
    # settle() 会先持久化 done outcome，再发最终事件。回放时也只凭 outcome + 有效正文 Word
    # 恢复 12/12；任何 agent 手写的 verified:true 都会走下面的证据钳制。
    outcome = read_json(os.path.join(job, 'outcome.json'), {})
    delivery_words = _body_docxs(job)
    meta = read_json(os.path.join(job, '任务.json'), {})
    redo = meta.get('redo_baseline') if isinstance(meta.get('redo_baseline'), dict) else None
    if redo:
        old = (redo.get('docx') or {})
        delivery_words = [fn for fn in delivery_words
                          if old.get(fn) != _file_digest(os.path.join(job, fn))]
    if claimed >= 12 and pct >= 100 and outcome.get('state') == 'done' and delivery_words:
        safe.update({'step': 12, 'total': 12, 'pct': 100, 'verified': True})
        return safe
    step = verified_step(job, claimed)
    max_pct = min(99, int(step * 100 / 12))
    safe.update({'step': step, 'total': 12, 'pct': max(0, min(pct, max_pct)), 'verified': True})
    if claimed > step:
        safe['stage'] = STAGES[step - 1] if step else '准备中'
    return safe

def ingest_agent_event(job, ev):
    """把 agent_events.jsonl 的声明验证后写入 canonical 事件流。"""
    raw = redact(dict(ev or {}))
    # verified 是引擎内部结论，不允许声明流自行携带；否则一句 JSON 就能绕过全部磁盘证据。
    raw.pop('verified', None)
    safe = sanitize_event(job, raw)
    if safe.get('type') == 'progress': safe['verified'] = True
    emit(job, safe)
    try: claimed, accepted = int(raw.get('step') or 0), int(safe.get('step') or 0)
    except Exception: claimed, accepted = 0, 0
    if raw.get('type') == 'progress' and claimed > accepted:
        emit(job, {'type': 'worklog', 'lines': [
            '⚠ 它声称完成第 %d 步，但对应产物还没出现；进度已按磁盘证据停在第 %d 步。' %
            (claimed, accepted)
        ]})
    return safe

def drain_agent_events(job, offset=0):
    """从声明流按行续读；尾部半行解析失败时不前移，下一轮写完整后再处理。"""
    path = os.path.join(job, AGENT_EVENTS_FILE)
    try: lines = open(path, encoding='utf-8').read().splitlines()
    except Exception: return offset, []
    accepted = []
    while offset < len(lines):
        try: raw = json.loads(lines[offset])
        except Exception: break
        offset += 1
        accepted.append(ingest_agent_event(job, raw))
    return offset, accepted

def ensure_docx(job, known, force=False):
    """兜底导出 Word:有正文 md、却一个 docx 都没有时,引擎自己把 Word 出出来。

    客户报障的真实场景:任务跑到 100%「完成」,右侧一堆 .md,就是没有标书。
    根因是 agent 写完正文就收工,没跑 build_tender_docx.py;而引擎这边
    只有「质检修复动作真实发生」时才会重建 docx,green 直接放过去了。
    导出是确定性脚本、零 token,没有任何理由不兜。返回新出的文件名列表。"""
    if not force and _body_docxs(job, known): return []
    body = _body_mds(job, known)
    if not body: return []
    bt = _skill_module('build_tender_docx')
    if not bt:
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '⚠ 正文已写好但没出 Word,而技能包里的导出脚本 build_tender_docx.py 也没找到。'
                           '请到「设置 · 生成引擎」确认技能包路径,或重跑本任务。'})
        return []
    mat = merged_materials(job)
    spec_f = _job_find(job, 'word_format_spec.json')
    meta = read_json(os.path.join(job, '任务.json'), {})
    made = []
    for fn in body:
        stem = os.path.splitext(fn)[0]
        target = stem + '.docx'
        argv = [os.path.join(job, fn), os.path.join(job, target),
                '--title', str(meta.get('name') or stem), '--images-dir', mat]
        if spec_f: argv += ['--format-spec', spec_f]
        old_argv = sys.argv
        try:
            sys.argv = ['build_tender_docx.py'] + argv
            try: bt.main()
            except SystemExit as e:
                if e.code: raise RuntimeError('导出脚本退出码 %s' % e.code)
        except Exception as e:
            append_diagnostic(job, 'word_export_failed', str(e), level='error', file=fn)
            friendly = ('导出组件缺少 Word 模板(多因系统清理了临时文件),重启应用后点「重试导出 Word」即可恢复'
                        if ('Package not found' in str(e) or '_MEI' in str(e)) else str(e))
            emit(job, {'type': 'message', 'role': 'agent',
                       'text': '⚠ 自动补出 Word 失败(%s):`%s` 的正文是好的,'
                               '可以先打开 md 查看。' % (friendly, fn)})
            continue
        finally:
            sys.argv = old_argv
        if os.path.isfile(os.path.join(job, target)):
            made.append(target); known.add(target)
            emit(job, {'type': 'artifact', 'name': target})
    if made:
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '📄 模型写完正文就收工了,没执行最后一步导出。已由引擎**自动补出 Word**:`%s`'
                           '(确定性脚本,零额度消耗;封面/目录/页眉页脚/图片落位都按格式规范生成)。'
                           % '`、`'.join(made),
                   'actions': [{'act': 'open_artifact', 'label': '打开 Word', 'file': made[0]}]})
    return made


def word_format_audit(job, word_name=''):
    """执行技能包中的确定性 Word 格式门禁。

    不调用子进程：PyInstaller 环境下 ``sys.executable script.py`` 会重启引擎。
    这里直接调用同一脚本的检查函数，失败项与脚本退出码 1 同源。
    """
    words = _body_docxs(job)
    word_name = word_name or (words[0] if words else '')
    report_name = 'Word格式自检报告.md'
    report_path = os.path.join(job, report_name)
    checker = _skill_module('check_docx_format')
    if not word_name or not checker:
        detail = '未找到待检 Word' if not word_name else '技能包缺少 check_docx_format.py'
        open(report_path, 'w', encoding='utf-8').write(
            '# 投标文件格式自检报告\n\n- 结论：❌ 未通过\n- 原因：%s\n' % detail)
        return {'status': 'fail', 'report': report_name, 'failed': 1, 'summary': detail}
    word_path = os.path.join(job, word_name)
    spec_path = _job_find(job, 'word_format_spec.json')
    meta = read_json(os.path.join(job, '任务.json'), {})
    try:
        spec = checker.load_spec(spec_path)
        document = checker.Document(word_path)
        checks = checker.Checker()
        front = spec['cover'].get('enabled', True) or spec['toc'].get('enabled', True)
        restart = bool(spec['page_numbering'].get('body_restart', True)) and front
        body_index = checker.check_page_numbering(document, spec, restart, checks)
        checker.check_page(document, spec, body_index, checks)
        checker.check_text_style(
            document.styles['Normal'], spec['body'], '正文', checks,
            indent_expected=spec['body'].get('first_line_indent_chars', 2))
        checker.check_headings(document, spec, checks)
        checker.check_header_footer(
            document, spec, body_index, restart, str(meta.get('name') or ''), checks)
        checker.check_toc(document, spec, checks)
        checker.check_cover(document, spec, str(meta.get('name') or ''), checks)
        checker.check_tables(document, spec, checks)
        checker.check_images(document, checks)
        checker.check_body_direct_overrides(document, spec, checks)
        checker.check_conventions(spec, checks, bool(spec_path))
        checker.write_report(report_path, word_path, spec_path, checks)
        # 格式结论必须绑定受检 Word 的确切字节；后续质检若重建
        # Word，旧绿色报告会立即变为 stale，不得开启交付。
        with open(report_path, 'a', encoding='utf-8') as report:
            report.write('\n- SHA-256：`%s`\n' % _file_digest(word_path))
        failed = len(checks.failed)
        # Word 能打开、检查能跑完,只是样式项不符 → warn,不是 fail:
        # 文件明明已导出且可用,把整单打成"未完成"是真机反馈里典型的
        # 「生成了 Word 还报错」——样式差异属于提交前人工确认项,不是交付失败。
        return {'status': 'pass' if failed == 0 else 'warn', 'report': report_name,
                'failed': failed,
                'summary': ('Word 格式门禁已通过' if failed == 0
                            else 'Word 已导出;格式自检有 %d 项不符(不影响打开使用),详见报告' % failed)}
    except Exception as exc:
        safe = redact(str(exc))
        open(report_path, 'w', encoding='utf-8').write(
            '# 投标文件格式自检报告\n\n- 结论：❌ 未通过\n- 原因：格式检查执行失败：%s\n' % safe)
        return {'status': 'fail', 'report': report_name, 'failed': 1,
                'summary': '格式检查执行失败'}


def _word_format_status(job, word_name=''):
    report_name = 'Word格式自检报告.md'
    try:
        text = open(os.path.join(job, report_name), encoding='utf-8', errors='ignore').read()
    except OSError:
        return {'status': 'missing', 'present': False, 'report': report_name}
    passed = '结论：✅ 全部通过' in text
    digest_match = re.search(r'- SHA-256：`([0-9a-f]{64})`', text, re.I)
    words = _body_docxs(job)
    word_name = word_name or (words[0] if words else '')
    current_digest = _file_digest(os.path.join(job, word_name)) if word_name else ''
    bound = bool(digest_match and current_digest
                 and hmac.compare_digest(digest_match.group(1).lower(), current_digest.lower()))
    # 报告有 SHA 绑定说明检查真的跑完了:未通过=样式项不符(warn);
    # 没有 SHA 的失败报告是「检查没跑成」(fail)。passed 但绑定失效 = 报告过期(stale)。
    if passed:
        status = 'pass' if bound else 'stale'
    else:
        status = 'warn' if bound else 'fail'
    return {'status': status, 'present': True, 'report': report_name,
            'word_sha256_bound': bound}

@app.post('/v1/jobs/{jid}/export_docx')
def export_docx(jid: str):
    """手动补出 Word:任务跑完没有 docx 时,前端「立即补出 Word」按钮调它。"""
    job = jpath(jid)
    if not os.path.isdir(job):      # 任务被删了/换了机器:要返 404,别让 os.listdir 抛成 500
        return JSONResponse({'ok': False, 'error': '任务不存在(可能已被删除)'}, 404)
    known = set(list_deliverables(job))
    made = ensure_docx(job, known)
    if made: return {'ok': True, 'made': made}
    if any(fn.lower().endswith('.docx') for fn in known):
        return {'ok': True, 'made': [], 'error': '任务里已经有 Word 了,不用再补'}
    # 有正文却没出成 Word = 导出组件出错,不是「没有正文」。旧文案把用户指向
    # 「重跑本任务」——既烧额度又治不了模板缺失,必须把真实原因和轻量动作给出来。
    errs = _last_export_errors(job)
    if _body_mds(job, known):
        hint = ('Word 导出组件执行失败:%s。' % errs[0] if errs else 'Word 导出组件执行失败。')
        if errs and ('Package not found' in errs[0] or '_MEI' in errs[0]):
            hint += ' 多因系统清理了临时文件——重启应用后点「重试导出 Word」即可,无需重跑任务。'
        return JSONResponse({'ok': False, 'code': 'word_export_failed', 'error': hint}, 500)
    return JSONResponse({'ok': False, 'error': '没有找到可导出的正文稿(投标文件_*.md),请重跑本任务'}, 400)

def _last_export_errors(job, limit=3):
    """diagnostics.jsonl 里最近的 word_export_failed 明细(新→旧),给导出失败响应用。"""
    out = []
    try:
        for ln in reversed(open(os.path.join(job, 'diagnostics.jsonl'),
                                encoding='utf-8', errors='ignore').read().splitlines()[-100:]):
            try: rec = json.loads(ln)
            except Exception: continue
            if rec.get('code') == 'word_export_failed':
                out.append(str(rec.get('detail') or '')[:200])
                if len(out) >= limit: break
    except OSError:
        pass
    return out

def quality_audit(job, known):
    """完成后的确定性质检 + 自动修复 + 重出 Word(引擎兜底,不管模型有没有自觉跑过):
    图片按索引锚点搬正/补插/剔除、重复段折叠、按章字数、应答覆盖率 → 《成品质检报告》。
    修复动作真实发生时,用修复稿重建 docx(build_tender_docx,零 token)——给客户看的必须是修好的 Word。"""
    # 最高优先级红灯：没有正文 Word 就不是“待优化”，而是没有交付。分析件绝不进入内容修复。
    if not _body_docxs(job, known):
        body = _body_mds(job, known)
        if body:
            detail = '正文 Markdown 已存在，但 Word 导出没有成功。'
            actions = [{'act': 'export_docx', 'label': '立即补出 Word'},
                       {'act': 'open_log', 'label': '查看运行日志'},
                       {'act': 'rerun', 'label': '重跑本任务'}]
        else:
            detail = '当前只有解析/分析文件，没有可提交的正文 Word。'
            actions = [{'act': 'open_log', 'label': '查看运行日志'},
                       {'act': 'rerun', 'label': '重跑本任务'}]
        emit(job, {'type': 'health', 'level': 'red', 'summary': '没有可交付的 Word，任务未完成',
                   'gaps': [{'level': 'red', 'title': '没出 Word，未完成', 'detail': detail,
                             'actions': actions}]})
        return {'status': 'fail', 'level': 'red', 'summary': '没有可交付的 Word'}
    qg = _skill_module('quality_gate')
    if not qg: return content_gate(job, known)          # 老技能目录没有质检脚本:退回旧内容门禁
    mat = merged_materials(job)
    score = _job_find(job, '评分点响应矩阵.md')
    devs = [p for p in [_job_find(job, '技术应答偏离表.md'), _job_find(job, '商务偏离表.md')] if p]
    mds = _body_mds(job, known)     # 与兜底导出共用同一套正文识别,避免命名一变全链路失灵
    if not mds: return content_gate(job, known)
    worst, fixed_total, lines, audit_errors = 'green', 0, [], []
    red_items = []          # 具体红项:出件前检查面板要给每条挂上「重做这一章」,不能只丢一句"见报告"
    order = {'green': 0, 'yellow': 1, 'red': 2}
    for fn in mds:
        mp = os.path.join(job, fn)
        try:
            res = qg.audit(mp, materials=mat, min_chapter=prompts.CHAPTER_TARGET_CHARS, score=score, deviations=devs)
            fixed, pre_images = None, None
            if res['plan'] or any(l == 'red' and ('打散' in t or '重复' in t) for l, t in res['items']):
                pre_images = res['images']
                fixed = qg.apply_fix(mp, res)
                fixed_total += fixed or 0
                res = qg.audit(mp, materials=mat, min_chapter=prompts.CHAPTER_TARGET_CHARS, score=score, deviations=devs)
            try: qg.write_report(os.path.join(job, '成品质检报告.md'), fn, res, fixed, images_pre=pre_images)
            except TypeError: qg.write_report(os.path.join(job, '成品质检报告.md'), fn, res, fixed)
            if order[res['level']] > order[worst]: worst = res['level']
            icon = {'green': '✅', 'yellow': '🟡', 'red': '🔴'}[res['level']]
            lines.append('%s %s:%d 章共 %d 字' % (icon, fn, len(res['chapters']), res['total_chars']))
            for l, t in res['items'][:4]:
                lines.append(('  🔴 ' if l == 'red' else '  🟡 ') + t)
            red_items += [t for l, t in res['items'] if l == 'red']
            if fixed: lines.append('  🔧 已自动修复 %d 处(图片落位/重复段),原稿备份 *.bak.md' % fixed)
        except Exception as e:
            lines.append('⚠ %s 质检异常:%s' % (fn, e)); audit_errors.append(str(e)); continue
    emit(job, {'type': 'artifact', 'name': '成品质检报告.md'})
    # 修复真的动了内容 → 用修复稿重建 Word(客户拿到手的必须是修好的那份)
    if fixed_total:
        bt = _skill_module('build_tender_docx')
        spec_f = _job_find(job, 'word_format_spec.json')
        meta = read_json(os.path.join(job, '任务.json'), {})
        for fn in mds:
            stem = os.path.splitext(fn)[0]
            # 目标 docx:同名优先;唯一一个 docx 交付物也认(模型常把 投标文件_技术标.md 出成 技术标.docx)
            docxs = [x for x in known if x.endswith('.docx')]
            target = (stem + '.docx') if (stem + '.docx') in known else (docxs[0] if len(docxs) == 1 else stem + '.docx')
            if not bt: break
            try:
                argv = [os.path.join(job, fn), os.path.join(job, target),
                        '--title', str(meta.get('name') or stem), '--images-dir', mat]
                if spec_f: argv += ['--format-spec', spec_f]
                sub = (meta.get('name') or '')
                old_argv = sys.argv
                sys.argv = ['build_tender_docx.py'] + argv
                try: bt.main()
                finally: sys.argv = old_argv
                emit(job, {'type': 'artifact', 'name': target})
                lines.append('🔁 已用修复稿重新导出 %s(图片已按锚点落位)' % target)
            except SystemExit:
                pass
            except Exception as e:
                lines.append('⚠ 重建 %s 失败:%s(md 修复稿仍有效)' % (target, e))
    # 标准模式的模型复核报告:必办 → 红,建议 → 黄,和确定性质检的结论合在同一张出件前检查里
    review_gaps = _review_report_gaps(job)
    review_red = [g for g in review_gaps if g['level'] == 'red']
    if review_gaps:
        lines.append('🧐 模型复核:必办 %d 项、建议 %d 项(见《模型复核报告.md》)'
                     % (len(review_red), len(review_gaps) - len(review_red)))
        if review_red: worst = 'red'
        elif worst == 'green': worst = 'yellow'
    verdict = {'green': '✅ 成品质检通过', 'yellow': '🟡 成品质检:有建议项(见《成品质检报告.md》)',
               'red': '🔴 成品质检:有必须处理项——按报告处理或对相应章节「定向重做」后再交付'}[worst]
    emit(job, {'type': 'message', 'role': 'agent', 'text': verdict + '\n\n' + '\n'.join(lines[:14]),
               'actions': [{'act': 'open_artifact', 'label': '打开质检报告', 'file': '成品质检报告.md'}]
                          + ([{'act': 'open_redo', 'label': '定向重做不达标章节'}] if worst == 'red' else [])})
    if worst == 'red':
        # 每条 gap 自带可执行动作。以前只给一条「见《成品质检报告.md》」,前端画一颗「处理」按钮
        # 却没绑任何事件——客户点了没反应,卡在那儿不知道该干嘛(真实报障)。
        # 现在:动作由后端指定(它才知道该怎么修),前端只负责执行;没有动作的 gap 不画按钮。
        gaps = [{'level': 'red', 'title': '逐项详情见《成品质检报告.md》',
                 'detail': '字数/图片落位/重复段/应答覆盖率都在报告里,逐条对着改',
                 'actions': [{'act': 'open_artifact', 'label': '打开报告', 'file': '成品质检报告.md'}]}]
        # 把质检抓到的具体红项也放进来,每条直接挂「定向重做这一章」——用户不用自己抄章节名
        for t in red_items[:6]:
            g = {'level': 'red', 'title': t, 'detail': '按报告要求补足后再交付'}
            ch = re.search(r'章节「([^」]+)」', t)
            if ch:
                g['actions'] = [{'act': 'redo', 'label': '重做这一章',
                                 'param': '重写章节「%s」:%s' % (ch.group(1), t)}]
            gaps.append(g)
        gaps.extend(review_gaps)
        gaps.append({'level': 'yellow', 'title': '也可以对整册不达标章节一起重做',
                     'detail': '只重做这些章节,其余产物保留,完成后自动重新汇总并更新自检',
                     'actions': [{'act': 'open_redo', 'label': '定向重做'}]})
        emit(job, {'type': 'health', 'level': 'red', 'summary': '成品质检有必须处理项', 'gaps': gaps})
    elif review_gaps:
        emit(job, {'type': 'health', 'level': 'yellow', 'summary': '模型复核有建议项',
                   'gaps': review_gaps + [{'level': 'yellow', 'title': '逐项详情见《模型复核报告.md》',
                                           'detail': '建议项改了更好,不改也能提交',
                                           'actions': [{'act': 'open_artifact', 'label': '打开复核报告',
                                                        'file': '模型复核报告.md'}]}]})
    if audit_errors:
        append_diagnostic(job, 'quality_audit_partial_error', '; '.join(audit_errors), level='error')
        return {'status': 'unknown', 'level': 'unknown', 'summary': '部分关键检查没有完成'}
    return {'status': ('fail' if worst == 'red' else ('warning' if worst == 'yellow' else 'pass')),
            'level': worst,
            'summary': {'red': '关键检查有必须处理项', 'yellow': '关键检查有建议确认项',
                        'green': '关键检查已通过'}[worst],
            'report': '成品质检报告.md'}

def content_gate(job, names):
    """内容门禁:格式自检只管版式,这里管正文是否被逐字打散/重复灌注,坏了不静默交付"""
    try:
        import doc_quality as dq
    except Exception:
        return {'status': 'unknown', 'level': 'unknown', 'summary': '内容检查组件不可用'}
    gaps, bad = [], []
    # 章节稿也要扫:复读循环是**某一章**写坏了,只扫合并后的整册虽然也能看见,
    # 却说不出是哪一章——而这类问题的处理动作恰恰是「重做那一章」,得有章名才点得动。
    bodies = set(_body_mds(job, names)) | set(_body_docxs(job, names)) | set(_chapter_mds(job, names))
    for fn in sorted(bodies):
        try:
            r = dq.detect(dq.read_any(os.path.join(job, fn)))
        except Exception:
            continue
        for i in r.get('issues', []):
            bad.append(fn)
            # 每条都挂上真能点的「一键修复」:这类问题的处理动作就是清洗,不用让用户自己去找按钮
            acts = [{'act': 'repair', 'label': '一键修复'}]
            # 复读循环是例外:清洗只能把几百遍折成一遍,折完这一章仍旧没有内容。
            # 首选动作必须是「重做这一章」,否则修完看着干净、交出去还是废稿。
            if i.get('code') == 'loop':
                acts = [{'act': 'redo', 'label': '重做这一章',
                         'param': '重写章节「%s」' % os.path.splitext(fn)[0]},
                        {'act': 'repair', 'label': '仅折叠复读'}]
            gaps.append({'level': 'red', 'title': '%s:%s' % (fn, i['title']), 'detail': i['detail'],
                         'actions': acts})
    if gaps:
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '⚠ 内容门禁未通过:%s 存在正文被逐字打散或整段重复灌注的问题,直接交付会是废稿。'
                           '可在「出件前检查」点「一键修复」自动清洗后重新出 Word。' % '、'.join(sorted(set(bad))),
                   'actions': [{'act': 'repair', 'label': '一键修复内容异常'}]})
        emit(job, {'type': 'health', 'level': 'red', 'summary': '内容异常,需修复后再出件', 'gaps': gaps[:8]})
        return {'status': 'fail', 'level': 'red', 'summary': '内容异常，需要修复后再出件'}
    return {'status': 'pass', 'level': 'green', 'summary': '关键检查已通过'}

@app.post('/v1/jobs/{jid}/repair')
def repair_job(jid: str):
    """一键修复:清洗 md 交付物(逐字碎片合并、重复样板折叠),生成 *_已清洗.md"""
    import doc_quality as dq
    job = jpath(jid); fixed = []
    names = list_deliverables(job)
    bodies = set(_body_mds(job, names)) | set(_body_docxs(job, names))
    for fn in sorted(bodies):
        if not fn.endswith(('.md', '.docx')) or '已清洗' in fn: continue
        p = os.path.join(job, fn)
        try:
            src = dq.read_any(p)
            if dq.detect(src).get('ok'): continue
            out = os.path.join(job, os.path.splitext(fn)[0] + '_已清洗.md')
            open(out, 'w', encoding='utf-8').write(dq.repair(src))
            fixed.append(os.path.basename(out))
            emit(job, {'type': 'artifact', 'name': os.path.basename(out)})
        except Exception as e:
            emit(job, {'type': 'error', 'text': '修复 %s 失败:%s' % (fn, e)})
    emit(job, {'type': 'message', 'role': 'agent',
               'text': ('已清洗:%s。请核对内容后用清洗版重新出 Word。' % '、'.join(fixed)) if fixed else '未发现需要清洗的内容异常。'})
    return {'ok': True, 'fixed': fixed}

def kill_tree(p):
    """子进程跑在独立会话里:必须杀整个进程组,而且父进程先死不代表子孙也死了——
    所以先记住进程组号,SIGTERM 之后无条件再补一次 SIGKILL(组已消失时忽略即可)。"""
    if not p: return
    try: pgid = os.getpgid(p.pid)
    except Exception: pgid = None
    if os.name == 'nt':
        # Popen.kill 只杀直接子进程；CLI/执行外壳还会拉工具子进程，必须按精确 PID 杀整棵树。
        try:
            subprocess.run(['taskkill', '/PID', str(int(p.pid)), '/T', '/F'],
                           stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                           timeout=10, check=False, creationflags=0x08000000)
        except Exception: pass
        try: p.kill()
        except Exception: pass
        return
    for sig, wait_s in ((signal.SIGTERM, 4), (signal.SIGKILL, 2)):
        if pgid:
            try: os.killpg(pgid, sig)
            except ProcessLookupError: pass
            except Exception: pass
        try: p.wait(wait_s)
        except Exception: pass
        if sig is signal.SIGTERM and pgid:
            try:            # 组里还有活口(父死子活)就继续走 SIGKILL
                os.killpg(pgid, 0)
            except Exception:
                return

def _tracked_detached_run(cmd, timeout, **kwargs):
    """subprocess.run 的可清理版本：关 App 时能找回并终止测试/校验子进程。"""
    capture = bool(kwargs.pop('capture_output', False))
    if capture:
        kwargs.setdefault('stdout', subprocess.PIPE)
        kwargs.setdefault('stderr', subprocess.PIPE)
    with RUNNING_LOCK:
        if SHUTTING_DOWN or EXITING:
            raise RuntimeError('应用正在退出，未启动新的外部命令')
        proc = subprocess.Popen(cmd, **kwargs, **DETACH)
        with DETACHED_CHILDREN_LOCK:
            DETACHED_CHILDREN[id(proc)] = proc
    try:
        try:
            stdout, stderr = proc.communicate(timeout=timeout)
        except subprocess.TimeoutExpired as exc:
            kill_tree(proc)
            try: stdout, stderr = proc.communicate(timeout=1)
            except Exception: stdout, stderr = exc.output, exc.stderr
            raise subprocess.TimeoutExpired(cmd, timeout, output=stdout, stderr=stderr)
        return subprocess.CompletedProcess(cmd, proc.returncode, stdout, stderr)
    finally:
        with DETACHED_CHILDREN_LOCK:
            DETACHED_CHILDREN.pop(id(proc), None)

def _kill_tracked_detached_children():
    with DETACHED_CHILDREN_LOCK:
        children = list(DETACHED_CHILDREN.values())
        DETACHED_CHILDREN.clear()
    for proc in children:
        try: kill_tree(proc)
        except Exception: pass

CLI_STALL_SECONDS = float(os.environ.get('BIDDOG_CLI_STALL_SECONDS', os.environ.get('BID_CLI_STALL_SECONDS', 15 * 60)))
CLI_STALL_POLL = float(os.environ.get('BIDDOG_CLI_STALL_POLL', os.environ.get('BID_CLI_STALL_POLL', 2)))
CLI_TOTAL_SECONDS = float(os.environ.get('BIDDOG_CLI_TOTAL_SECONDS', os.environ.get('BID_CLI_TOTAL_SECONDS', 3 * 3600)))
_ACTIVITY_CONTROL = {'events.jsonl', AGENT_EVENTS_FILE, 'progress.json', '任务.json',
                     'outcome.json', '.line_ts.json', 'answers.jsonl', 'chat.jsonl'}

def cli_activity_signature(job):
    """CLI 活跃指纹：日志 + 产物文件本身。修改已有正文也算，轮询控制文件不算。"""
    meta = read_json(os.path.join(job, '任务.json'), {})
    tender = os.path.basename(meta.get('tender') or '')
    items = []
    log_path = os.path.join(job, 'run.log')
    try:
        st = os.stat(log_path); log_sig = (st.st_size, st.st_mtime_ns)
    except OSError: log_sig = (0, 0)
    for root, dirs, files in os.walk(job):
        dirs[:] = [d for d in dirs if d not in ('node_modules', '__pycache__', '.git')]
        for fn in files:
            if fn == tender or fn == 'run.log' or fn in _ACTIVITY_CONTROL or fn.endswith('.tmp'): continue
            p = os.path.join(root, fn)
            try:
                st = os.stat(p)
                items.append((os.path.relpath(p, job), st.st_size, st.st_mtime_ns))
            except OSError:
                pass
    return (log_sig, tuple(sorted(items)))

# 测试与诊断工具使用的稳定名字（此前派工单草案采用了下划线前缀）。
_job_activity_signature = cli_activity_signature

def wait_cli_process(proc, job, stall_seconds=None, total_seconds=None, poll_seconds=None):
    """等待 CLI，同时执行 15 分钟无真实文件活动与 3 小时总闸。"""
    stall_seconds = CLI_STALL_SECONDS if stall_seconds is None else float(stall_seconds)
    total_seconds = CLI_TOTAL_SECONDS if total_seconds is None else float(total_seconds)
    poll_seconds = CLI_STALL_POLL if poll_seconds is None else float(poll_seconds)
    started = changed = time.monotonic()
    sig = cli_activity_signature(job)
    base = os.path.basename(job)
    while True:
        rc = proc.poll()
        if rc is not None: return {'status': 'exited', 'rc': rc}
        if _cancel_requested(base):
            kill_tree(proc)
            return {'status': 'cancelled', 'rc': proc.poll()}
        cur = cli_activity_signature(job)
        now_m = time.monotonic()
        if cur != sig:
            sig, changed = cur, now_m
        if now_m - changed >= stall_seconds:
            kill_tree(proc)
            return {'status': 'stalled', 'rc': proc.poll()}
        if now_m - started >= total_seconds:
            kill_tree(proc)
            return {'status': 'timeout', 'rc': proc.poll()}
        time.sleep(max(0.01, poll_seconds))

def read_tail(path, n=800):
    try:
        with open(path, 'rb') as f:
            f.seek(0, os.SEEK_END)
            size = f.tell()
            limit = min(size, max(0, int(n)))
            f.seek(max(0, size - limit))
            # 文件可能在 tell/seek 后继续增长；read(limit) 仍保证内存上界。
            return f.read(limit).decode('utf-8', 'ignore')
    except Exception:
        return ''

def _tail_lines(path, count=8):
    try:
        lines = open(path, encoding='utf-8', errors='ignore').read().splitlines()
    except Exception:
        return []
    return [_redact_runtime(x) for x in lines[-count:]]

def settle(job, known=None, stop_reason=None, commit=True):
    """统一出件闸门。CLI、OpenCode server、mock 都只能从这里进入终态。"""
    if not os.path.isdir(job): return {'state': 'stopped', 'reason': '任务目录不存在'}
    harvest(job)
    known = set(list_deliverables(job)) if known is None else set(known) | set(list_deliverables(job))
    for fn in sorted(known): emit(job, {'type': 'artifact', 'name': fn})

    meta = read_json(os.path.join(job, '任务.json'), {})
    redo = meta.get('redo_baseline') if isinstance(meta.get('redo_baseline'), dict) else None
    redo_docs = (redo or {}).get('docx') or {}
    redo_mds = (redo or {}).get('md') or {}
    body_now = _body_mds(job, known)
    words_now = _body_docxs(job, known)
    changed_words = [fn for fn in words_now
                     if redo_docs.get(fn) != _file_digest(os.path.join(job, fn))] if redo else words_now
    changed_mds = [fn for fn in body_now
                   if redo_mds.get(fn) != _file_digest(os.path.join(job, fn))] if redo else body_now

    if body_now and (not words_now or (redo and not changed_words and changed_mds)):
        try: ensure_docx(job, known)
        except Exception as e:
            emit(job, {'type': 'message', 'role': 'agent',
                       'text': '⚠ 自动导出 Word 没有成功：%s。正文稿已保留，可直接重试导出。' % redact(str(e)),
                       'actions': [{'act': 'export_docx', 'label': '重试导出 Word'},
                                   {'act': 'open_log', 'label': '查看运行日志'}]})
        # 定向重做只有正文真的变化时才允许覆盖旧 Word；否则旧稿不能替本次失败兜底。
        if redo and not changed_words and changed_mds:
            try: ensure_docx(job, known, force=True)
            except Exception: pass
        known |= set(list_deliverables(job))

    words = _body_docxs(job, known)
    if redo:
        words = [fn for fn in words if redo_docs.get(fn) != _file_digest(os.path.join(job, fn))]
    if words:
        outcome = read_json(os.path.join(job, 'outcome.json'), {})
        if outcome.get('state') == 'done' and commit:
            return {'state': 'done', 'word': words[0], 'artifacts': sorted(known)}
        try:
            quality = quality_audit(job, known)
            # Third-party/legacy hooks returned None; keep them compatible while the built-in gate is structured.
            if not isinstance(quality, dict):
                quality = {'status': 'pass', 'level': 'green', 'summary': '关键检查已完成'}
        except Exception as e:
            quality = {'status': 'unknown', 'level': 'unknown', 'summary': '关键检查没有完成'}
            append_diagnostic(job, 'quality_gate_error', str(e), level='error')
            emit(job, {'type': 'message', 'role': 'agent',
                       'text': 'Word 已生成，但关键检查没有完成；请先查看检查结果再交付。',
                       'actions': [{'act': 'open_log', 'label': '查看诊断详情'}]})
        known |= set(list_deliverables(job))
        patch_json(os.path.join(job, 'delivery.json'), {'updated_at': now(), 'quality': quality},
                   remove=('signature', 'summary'))
        if quality.get('status') in ('fail', 'unknown'):
            why = '已停止（关键检查未通过）' if quality.get('status') == 'fail' else '已停止（关键检查未完成）'
            halt(job, why)
            return {'state': 'stopped', 'reason': why, 'word': words[0],
                    'artifacts': sorted(known), 'delivery': delivery_summary(job, quality)}
        pipeline_state = generation_pipeline.load(job)
        if pipeline_state.get('version') == generation_pipeline.PIPELINE_VERSION:
            # quality_audit 可能修改 Markdown 并重建 Word；因此格式门禁必须
            # 在所有内容修复之后重跑，并用 SHA-256 绑定最终 Word。
            format_result = word_format_audit(job, words[0])
            known |= set(list_deliverables(job))
            if format_result.get('status') == 'fail':
                # 只有「检查没跑成/Word 不可用」才拦交付;样式项不符(warn)不把整单打停——
                # Word 明明已生成可用,判停就是真机反馈的「生成了 Word 还报错」。
                why = '已停止（Word 格式检查无法完成）'
                halt(job, why)
                return {'state': 'stopped', 'reason': why, 'word': words[0],
                        'artifacts': sorted(known), 'delivery': delivery_summary(job, quality)}
            if format_result.get('status') == 'warn':
                emit(job, {'type': 'message', 'role': 'agent',
                           'text': '%s。这些属于提交前人工确认项,在「出件前检查」里逐条看即可。' %
                                   (format_result.get('summary') or 'Word 格式自检有不符项')})
        delivery = delivery_summary(job, quality)
        if not delivery.get('ready'):
            why = '已停止（出件检查未通过）'
            halt(job, why)
            return {'state': 'stopped', 'reason': why, 'word': words[0],
                    'artifacts': sorted(known), 'delivery': delivery}
        if not commit:
            return {'state': 'ready', 'word': words[0], 'artifacts': sorted(known),
                    'delivery': delivery}
        if redo:
            meta.pop('redo_baseline', None)
            write_json(os.path.join(job, '任务.json'), meta)
        if not _commit_done(job, words[0]):
            # stop/pause 已经原子取得本轮终态；由控制请求写 stopped/paused，旧 worker 不得翻盘。
            return {'state': 'stopped', 'reason': '本轮已收到停止或暂停请求',
                    'artifacts': sorted(known)}
        summary, actions = artifact_summary(job, known)
        used = skill_evidence(job)
        if used['state'] == 'unverifiable':
            summary += ('\n\n⚠ **技能包运行证据暂时无法核验**（%s）请人工复核内容、'
                        '响应矩阵和格式门禁，再决定是否交付。' % used['why'])
            actions = (actions or []) + [{'act': 'open_log', 'label': '查看运行日志'}]
        elif used['state'] == 'missing':
            summary += ('\n\n⚠ **写作规则没有完整载入**（%s）请在设置中重新测试连接，并人工复核内容、'
                        '响应矩阵和格式检查，再决定是否交付。' % used['why'])
            actions = (actions or []) + [{'act': 'open_log', 'label': '查看运行日志'}]
        emit(job, {'type': 'message', 'role': 'agent', 'text': summary, 'actions': actions})
        emit(job, {'type': 'skill_used', 'state': used['state'], 'ok': used['ok'],
                   'hits': used['hits'], 'why': used['why']})
        return {'state': 'done', 'word': words[0], 'artifacts': sorted(known),
                'delivery': delivery_summary(job, quality)}

    # 没有 Word：最高优先级健康红灯，绝不把解析件或分章过程稿包装成完成。
    body = _body_mds(job, known)
    chapters = _chapter_mds(job, known)
    if redo:
        emit(job, {'type': 'health', 'level': 'red',
                   'summary': '定向重做没有更新可交付的 Word，任务未完成',
                   'gaps': [{'level': 'red', 'title': '旧 Word 不能替本次重做算成功',
                             'detail': '旧稿仍完整保留；请查看日志后继续或重新定向重做。',
                             'actions': [{'act': 'open_log', 'label': '查看运行日志'},
                                         {'act': 'open_redo', 'label': '重新定向重做'}]}]})
    elif chapters and not body:
        resume = ([{'act': 'resume', 'label': '从已保存内容继续'}]
                  if meta.get('oc_session') else [{'act': 'rerun', 'label': '重新生成'}])
        emit(job, {'type': 'health', 'level': 'red',
                   'summary': '分章撰写中断，已保留 %d 个章节' % len(chapters),
                   'gaps': [{'level': 'red', 'title': '正文尚未汇总，最终 Word 未生成',
                             'detail': '可从已保存章节继续，不需要从招标解析重新开始。',
                             'actions': resume + [{'act': 'open_log', 'label': '查看运行日志'}]}]})
    else:
        try: quality_audit(job, known)
        except Exception: pass
    analysis = [fn for fn in known if fn.endswith(DELIVER_EXT)]
    lines = _tail_lines(os.path.join(job, 'run.log'), 8)
    if redo:
        why = stop_reason or '已停止（定向重做未更新正文 Word）'
        lead = '本次定向重做没有生成或更新正文 Word；旧 Word 已保留，但不能算作本次成功。'
        actions = [{'act': 'open_log', 'label': '查看运行日志'},
                   {'act': 'open_redo', 'label': '重新定向重做'}]
    elif body:
        why = '已停止（生成中断：Word 导出失败）'
        lead = '正文稿已经生成，但最终 Word 导出失败，因此这单仍未完成。'
        actions = [{'act': 'export_docx', 'label': '重试导出 Word'},
                   {'act': 'open_log', 'label': '查看运行日志'},
                   {'act': 'rerun', 'label': '重跑本任务'}]
    elif chapters:
        why = stop_reason or '已停止（撰写中断，内容已保留）'
        lead = ('分章撰写中断：已生成 %d 个章节，但还没有汇总成完整正文和最终 Word。'
                '已经写出的内容均已保留。' % len(chapters))
        actions = ([{'act': 'resume', 'label': '从已保存内容继续'}]
                   if meta.get('oc_session') else [{'act': 'rerun', 'label': '重新生成'}])
        actions.append({'act': 'open_log', 'label': '查看运行日志'})
    elif analysis:
        why = '已停止（生成中断：只有分析文件）'
        lead = '生成中断：目前只有解析/分析文件，没有正文和最终 Word。'
        actions = [{'act': 'open_log', 'label': '查看运行日志'},
                   {'act': 'rerun', 'label': '重跑本任务'}]
    else:
        why = '已停止（没有产出）'
        lead = '执行结束但没有产出任何交付文件。'
        actions = [{'act': 'open_log', 'label': '查看运行日志'},
                   {'act': 'rerun', 'label': '重跑本任务'}]
    if stop_reason: why = stop_reason
    detail = lead + (('\n\n运行日志最后 %d 行：\n%s' % (len(lines), '\n'.join(lines))) if lines else '')
    emit(job, {'type': 'error', 'text': detail, 'actions': actions})
    halt(job, why)
    return {'state': 'stopped', 'reason': why, 'artifacts': sorted(known)}

SKILL_EVENT_FILE = 'skill_events.jsonl'
_SKILL_EVENT_KEY = secrets.token_bytes(32)  # 仅存于本次引擎进程内；任务 cwd 中的 agent 拿不到
_SKILL_EVENT_FIELDS = ('run_id', 'type', 'status', 'path_sha256', 'manifest_sha256', 'ts')
_OC_SKILL_READS = {}  # durable event 的 tool.called → tool.success，仅在内存关联原始路径/内容

def _sha256_text(value):
    return hashlib.sha256(str(value or '').encode('utf-8')).hexdigest()

def _new_run_id():
    """一次 launch/redo 对应一个证据域；resume 则继续沿用原证据域。"""
    return uuid.uuid4().hex

def _dispatch_digest(dispatch):
    """只冻结派发内容的指纹，绝不把原始 prompt/cmd/path 写入任务元数据。"""
    if isinstance(dispatch, (list, tuple)):
        value = json.dumps([str(x) for x in dispatch], ensure_ascii=False, separators=(',', ':'))
    else:
        value = str(dispatch or '')
    return _sha256_text(value) if value else ''

def _dispatch_contains_skill(dispatch, skill_dir):
    values = dispatch if isinstance(dispatch, (list, tuple)) else [dispatch]
    text = '\0'.join(str(x or '') for x in values)
    skill_dir = os.path.realpath(skill_dir or '')
    skill_path = os.path.join(skill_dir, 'SKILL.md') if skill_dir else ''
    return bool(skill_dir and (skill_dir in text or skill_path in text))

def _skill_manifest(run_id, skill_dir, dispatch, injected, execution_path):
    """构造可落盘的非敏感技能派发收据：只有版本、布尔值与不可逆哈希。"""
    skill_dir = os.path.realpath(skill_dir or '')
    skill_path = os.path.realpath(os.path.join(skill_dir, 'SKILL.md')) if skill_dir else ''
    present = bool(skill_path and os.path.isfile(skill_path))
    version = ''
    if skill_dir:
        try:
            raw_bytes = open(os.path.join(skill_dir, '.skill_version'), 'rb').read(80)
            try: raw_version = raw_bytes.decode('ascii').strip()
            except UnicodeDecodeError: raw_version = ''
            # 只允许明显的数字版本标签原样落盘（兼容内置技能版与标准产品版）；
            # PAT/Key/备注等任意 token 一律只留不可逆摘要。
            safe_version = re.fullmatch(
                r'v?(?:0|[1-9]\d*)\.(?:0|[1-9]\d*)(?:\.(?:0|[1-9]\d*))?'
                r'(?:-[0-9A-Za-z.-]+)?(?:\+[0-9A-Za-z.-]+)?', raw_version
            )
            version = (raw_version if safe_version and len(raw_version) <= 40
                       else ('custom-' + hashlib.sha256(raw_bytes).hexdigest()[:12]
                             if raw_bytes else ''))
        except OSError:
            version = ''
    return {
        'run_id': run_id,
        'version': version or 'unversioned',
        'manifest_sha256': _file_digest(skill_path) if present else '',
        'path_sha256': _sha256_text(skill_path) if skill_path else '',
        'dispatch_sha256': _dispatch_digest(dispatch),
        'injected': bool(injected and present),
        'accepted': False,
        'execution_path': str(execution_path or ''),
        'file_present': present,
    }

def _set_skill_manifest(job, skill_dir, dispatch, injected, execution_path):
    meta_path = os.path.join(job, '任务.json')
    meta = read_json(meta_path, {})
    run_id = str(meta.get('run_id') or '')
    if not run_id: return None
    manifest = _skill_manifest(run_id, skill_dir, dispatch, injected, execution_path)
    meta['skill_manifest'] = manifest
    write_json(meta_path, meta)
    return manifest

def _mark_skill_accepted(job, execution_path=None):
    """只有外壳真实接收派发（HTTP 2xx/Popen 成功）后才把 accepted 置真。"""
    meta_path = os.path.join(job, '任务.json')
    meta = read_json(meta_path, {})
    manifest = meta.get('skill_manifest')
    if not isinstance(manifest, dict) or manifest.get('run_id') != meta.get('run_id'): return False
    manifest = dict(manifest)
    manifest['accepted'] = True
    if execution_path: manifest['execution_path'] = str(execution_path)
    meta['skill_manifest'] = manifest
    write_json(meta_path, meta)
    return True

def _skill_event_auth(record):
    payload = {k: record.get(k) for k in _SKILL_EVENT_FIELDS if record.get(k) not in (None, '')}
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(',', ':')).encode('utf-8')
    return hmac.new(_SKILL_EVENT_KEY, raw, hashlib.sha256).hexdigest()

def _skill_events(job):
    """只返回本次引擎进程签发的事件；agent 可写任务目录，但不能伪造进程内 HMAC。"""
    try:
        # 证据文件正常只有几行。被任务进程灌入大量垃圾时只看末尾，宁可降级为 unverifiable，
        # 也不能让日志接口无界读内存。
        lines = read_tail(os.path.join(job, SKILL_EVENT_FILE), 1024 * 1024).splitlines()
    except OSError:
        return []
    records = []
    for line in lines:
        try:
            record = json.loads(line)
            signature = str(record.get('auth') or '') if isinstance(record, dict) else ''
            if (isinstance(record, dict) and signature
                    and hmac.compare_digest(signature, _skill_event_auth(record))):
                records.append(record)
        except (TypeError, ValueError):
            pass
    return records

def _append_skill_event(job, record):
    """结构化证据只允许安全字段；调用方不得传原始工具参数。"""
    safe = {k: record.get(k) for k in _SKILL_EVENT_FIELDS if record.get(k) not in (None, '')}
    if not safe.get('ts'): safe['ts'] = now()
    # SSE 断线重连会重放历史；同一轮完全相同的读取证据只落一次。
    identity = tuple(safe.get(k) for k in (
        'run_id', 'type', 'status', 'path_sha256', 'manifest_sha256'
    ))
    if any(tuple(old.get(k) for k in (
            'run_id', 'type', 'status', 'path_sha256', 'manifest_sha256'
        )) == identity for old in _skill_events(job)):
        return
    safe['auth'] = _skill_event_auth(safe)
    try:
        with open(os.path.join(job, SKILL_EVENT_FILE), 'a', encoding='utf-8') as f:
            f.write(json.dumps(safe, ensure_ascii=False) + '\n')
    except OSError:
        pass

def _read_output_covers_file(state, path):
    """OpenCode 1.18.13–1.18.18 read 是按行返回；只有从首行连续覆盖到末行才算完整读取。"""
    tool_input = state.get('input') or {}
    try:
        offset = int(tool_input.get('offset', 1))
        limit = int(tool_input.get('limit', 2000))
    except (TypeError, ValueError):
        return False
    if offset != 1: return False
    try:
        with open(path, encoding='utf-8', errors='replace') as f:
            line_count = sum(1 for _ in f)
    except OSError:
        return False
    if limit < line_count: return False
    metadata = state.get('metadata') or {}
    if metadata.get('truncated') is True: return False
    numbered = []
    for line in str(state.get('output') or '').splitlines():
        match = re.match(r'^\s*(\d+):(?: |$)', line)
        if match: numbered.append(int(match.group(1)))
    if line_count == 0: return bool(re.search(r'End of file|0\s+lines', str(state.get('output') or ''), re.I))
    expected = list(range(1, line_count + 1))
    return any(numbered[i:i + line_count] == expected
               for i in range(max(0, len(numbered) - line_count + 1)))

def _observe_oc_skill_event(job, event, sid):
    """从 OpenCode SSE 观测精确的完整 read；原始路径与内容只在进程内短暂关联。"""
    if not isinstance(event, dict): return False
    event_type = str(event.get('type') or '')
    # 内置 OpenCode 1.18.13–1.18.18 的单会话 event 端点回放 durable `session.next.*`，
    # called 持有 path，success 才持有完整 content，必须用 sessionID + callID 关联。
    if event_type in ('session.next.tool.called', 'session.next.tool.success',
                      'session.next.tool.failed'):
        data = event.get('data') if isinstance(event.get('data'), dict) else event.get('properties')
        data = data if isinstance(data, dict) else {}
        event_sid = str(data.get('sessionID') or '')
        call_id = str(data.get('callID') or '')
        meta = read_json(os.path.join(job, '任务.json'), {})
        manifest = meta.get('skill_manifest') or {}
        run_id = str(meta.get('run_id') or '')
        if (not sid or event_sid != sid or not call_id or not run_id
                or str(meta.get('oc_session') or '') != sid
                or manifest.get('run_id') != run_id):
            return False
        pending_key = (os.path.realpath(job), run_id, sid, call_id)
        if event_type == 'session.next.tool.failed':
            _OC_SKILL_READS.pop(pending_key, None)
            return False
        if event_type == 'session.next.tool.called':
            if str(data.get('tool') or '').lower() != 'read': return False
            tool_input = data.get('input') or {}
            raw_path = tool_input.get('path') or tool_input.get('filePath')
            if not isinstance(raw_path, str) or not raw_path.strip(): return False
            try: offset = int(tool_input.get('offset', 1))
            except (TypeError, ValueError): return False
            if offset != 1: return False
            candidate = os.path.expanduser(raw_path.strip())
            if not os.path.isabs(candidate): candidate = os.path.join(job, candidate)
            candidate = os.path.realpath(candidate)
            path_sha256 = _sha256_text(candidate)
            manifest_sha256 = str(manifest.get('manifest_sha256') or '')
            if path_sha256 != manifest.get('path_sha256') or not manifest_sha256:
                return False
            if tool_input.get('limit') is not None:
                try:
                    if int(tool_input.get('limit')) <= 0: return False
                except (TypeError, ValueError):
                    return False
            # 有界缓存，异常事件洪泛时丢最旧项；丢证据只会降级 amber，不会误判 verified。
            while len(_OC_SKILL_READS) >= 256:
                _OC_SKILL_READS.pop(next(iter(_OC_SKILL_READS)))
            _OC_SKILL_READS[pending_key] = {
                'path_sha256': path_sha256, 'manifest_sha256': manifest_sha256,
            }
            return False
        pending = _OC_SKILL_READS.pop(pending_key, None)
        structured = data.get('structured') or {}
        content = structured.get('content') if isinstance(structured, dict) else None
        if (not pending or not isinstance(content, str)
                or structured.get('truncated') is True or structured.get('next') is not None
                or _sha256_text(content) != pending.get('manifest_sha256')
                or pending.get('path_sha256') != manifest.get('path_sha256')
                or pending.get('manifest_sha256') != manifest.get('manifest_sha256')):
            return False
        _append_skill_event(job, {
            'run_id': run_id, 'type': 'read_manifest', 'status': 'completed',
            'path_sha256': pending['path_sha256'],
            'manifest_sha256': pending['manifest_sha256'], 'ts': now(),
        })
        return True

    # 兼容标准 v2 `message.part.updated` 事件；当前 1.18.13–1.18.18 的 durable 端点不走此形状。
    if not event_type.endswith('message.part.updated'): return False
    properties = event.get('properties') or {}
    part = properties.get('part') or {}
    event_sid = str(properties.get('sessionID') or '')
    part_sid = str(part.get('sessionID') or '')
    if not sid or event_sid != sid or part_sid != sid: return False
    state = part.get('state') or {}
    if part.get('type') != 'tool' or str(part.get('tool') or '').lower() != 'read': return False
    if str(state.get('status') or '').lower() != 'completed': return False
    raw_path = (state.get('input') or {}).get('filePath')
    if not isinstance(raw_path, str) or not raw_path.strip(): return False

    meta = read_json(os.path.join(job, '任务.json'), {})
    manifest = meta.get('skill_manifest') or {}
    run_id = str(meta.get('run_id') or '')
    if (not run_id or manifest.get('run_id') != run_id
            or str(meta.get('oc_session') or '') != sid): return False
    candidate = os.path.expanduser(raw_path.strip())
    if not os.path.isabs(candidate): candidate = os.path.join(job, candidate)
    candidate = os.path.realpath(candidate)
    path_sha256 = _sha256_text(candidate)
    manifest_sha256 = _file_digest(candidate)
    if (path_sha256 != manifest.get('path_sha256')
            or not manifest_sha256
            or manifest_sha256 != manifest.get('manifest_sha256')):
        return False
    if not _read_output_covers_file(state, candidate): return False
    _append_skill_event(job, {
        'run_id': run_id, 'type': 'read_manifest', 'status': 'completed',
        'path_sha256': path_sha256, 'manifest_sha256': manifest_sha256, 'ts': now(),
    })
    return True

def _replay_oc_skill_evidence(job, timeout=4):
    """引擎重启后从 OpenCode 的持久会话重放证据，再用本进程临时密钥重新签发。"""
    if not _begin_oc_replay(): return False
    try:
        meta = read_json(os.path.join(job, '任务.json'), {})
        if not isinstance(meta, dict): return False
        manifest = meta.get('skill_manifest')
        sid = str(meta.get('oc_session') or '')
        if (not isinstance(manifest, dict) or manifest.get('execution_path') != 'opencode'
                or manifest.get('run_id') != meta.get('run_id') or not sid):
            return False
        try:
            # 历史会话固定保存在中标狗隔离的 OpenCode XDG 目录；用户后来切到
            # Codex/Claude 也必须从同一目录重放，不能跟着当前 kind 改数据库。
            replay_eng = dict((read_json(conf_path(), {}) or {}).get('engine') or {})
            replay_eng['kind'] = 's2'
            replay_eng.pop('cli_path', None)
            base_url = oc_serve(replay_eng)
            if not base_url: return False
            url = base_url + '/api/session/%s/event' % urllib.parse.quote(sid, safe='')
            req = urllib.request.Request(url, headers=oc_auth())
            deadline = time.monotonic() + max(1, min(int(timeout or 4), 10))
            count = 0
            with urllib.request.urlopen(req, timeout=max(1, min(int(timeout or 4), 10))) as response:
                for raw in response:
                    if time.monotonic() >= deadline or count >= 20000: break
                    line = raw.decode('utf-8', 'ignore').strip()
                    if not line.startswith('data:'): continue
                    count += 1
                    try: event = json.loads(line[5:].strip())
                    except (TypeError, ValueError): continue
                    if _observe_oc_skill_event(job, event, sid): return True
        except Exception:
            return False
        return False
    finally:
        _end_oc_replay()

def skill_evidence(job):
    """返回 verified / unverifiable / missing，避免把“看不见”误说成“没使用”。"""
    meta = read_json(os.path.join(job, '任务.json'), {})
    run_id = str(meta.get('run_id') or '')
    manifest = meta.get('skill_manifest')
    if not run_id and not isinstance(manifest, dict):
        why = ('任务尚未启动，当前没有运行证据。' if meta.get('staged') else
               '这是旧版本创建的历史任务，当时没有结构化技能证据，无法倒推是否读取。')
        return {'state': 'unverifiable', 'ok': False, 'hits': [], 'why': why}
    missing = []
    if not run_id: missing.append('本轮运行标识缺失')
    if not isinstance(manifest, dict):
        missing.append('没有技能派发收据')
        manifest = {}
    if manifest.get('run_id') != run_id: missing.append('技能派发收据不属于本轮')
    if not manifest.get('file_present') or not manifest.get('manifest_sha256') or not manifest.get('path_sha256'):
        missing.append('技能清单文件缺失或无法计算指纹')
    if not manifest.get('injected'): missing.append('技能指令没有注入本轮派发')
    if not manifest.get('accepted'): missing.append('生成服务没有确认收到本轮写作规则')
    if not manifest.get('execution_path'): missing.append('执行路径没有冻结')
    if missing:
        return {'state': 'missing', 'ok': False, 'hits': [],
                'why': '写作规则没有完整载入：%s。' % '；'.join(dict.fromkeys(missing))}

    for record in _skill_events(job):
        if (record.get('run_id') == run_id
                and record.get('type') == 'read_manifest'
                and record.get('status') == 'completed'
                and record.get('path_sha256') == manifest.get('path_sha256')
                and record.get('manifest_sha256') == manifest.get('manifest_sha256')):
            return {'state': 'verified', 'ok': True,
                    'hits': ['本轮已核验 SKILL.md 精确读取事件'], 'why': ''}
    return {'state': 'unverifiable', 'ok': False, 'hits': [],
            'why': '写作规则已送达，但本轮没有留下可核验的完整读取记录，无法确认实际读取情况。'}

# ==================== OpenCode server 模式 ====================
# 以前是 `opencode run --auto`:喂一条 prompt、等进程死,中间什么都看不见、停不掉、崩了从头来。
# server 模式把这三件事都解开了。以下端点名与请求体全部从固定版本的 /doc
# 实拉并跑通。1.18.18 的正式接口（/session）与兼容接口（/api/session）
# 不能混用：兼容创建接口会静默忽略 directory，导致节点在数据根目录里执行。
OC = {'proc': None, 'port': 0, 'base': '', 'pw': '', 'fingerprint': ''}     # 常驻 server 句柄
OC_LOCK = threading.Lock()
PIPELINE_SESSIONS = {}  # job id -> {OpenCode short-session ids}; stop/delete must interrupt all
PIPELINE_SESSIONS_LOCK = threading.Lock()

def _free_port():
    s = socket.socket(); s.bind(('127.0.0.1', 0)); p = s.getsockname()[1]; s.close(); return p

def oc_auth():
    """server 的访问凭证。不设密码时 opencode 自己会警告 `server is unsecured` ——
    那不是空话:这个 server 的 bash 与 edit 是全放行的,本机任何一个进程只要
    POST 一下就能以当前用户身份执行任意命令。所以每次起 server 现生成一把随机口令。
    实测认证方式是 HTTP Basic,**用户名必须正好是 opencode**(试过 x/admin/空,全 401),
    SSE 事件流同样要带。"""
    if not OC['pw']: return {}
    tok = base64.b64encode(('opencode:' + OC['pw']).encode()).decode()
    return {'Authorization': 'Basic ' + tok}

def oc_api(path, data=None, timeout=60, method=None):
    """打 opencode server。返回 (status, 解包后的 body)。它的响应统一裹在 data 里。"""
    if not OC['base']: return 0, None
    hdr = {'Content-Type': 'application/json'}; hdr.update(oc_auth())
    req = urllib.request.Request(
        OC['base'] + path, data=json.dumps(data).encode() if data is not None else None,
        headers=hdr, method=method or ('POST' if data is not None else 'GET'))
    try:
        with urllib.request.urlopen(req, timeout=timeout) as r:
            body = r.read().decode('utf-8', 'ignore')
            out = json.loads(body) if body.strip() else None
            if isinstance(out, dict) and 'data' in out and len(out) <= 2: out = out['data']
            return r.status, out
    except urllib.error.HTTPError as e:
        return e.code, _http_error_detail(e)
    except Exception as e:
        return -1, str(e)[:200]

def oc_model(conf=None):
    up = s2_conf(conf or read_json(conf_path(), {}))
    return {'providerID': 'biddog-s2', 'modelID': up['model']}


def oc_create_session(directory, title):
    """Create a 1.18.18 session whose working directory is provably isolated."""
    expected = os.path.abspath(str(directory or DATA))
    path = '/session?directory=%s' % urllib.parse.quote(expected, safe='')
    status, body = oc_api(path, {'title': str(title or '')[:180]}, timeout=30)
    if status not in (200, 201) or not isinstance(body, dict) or not body.get('id'):
        return ''
    actual = str(body.get('directory') or '')
    # OpenCode canonicalizes platform aliases (macOS: /tmp -> /private/tmp;
    # Windows can also normalize drive-letter case). Compare filesystem identity,
    # while still sending the caller's absolute directory to the formal API.
    expected_identity = os.path.normcase(os.path.realpath(expected))
    actual_identity = os.path.normcase(os.path.realpath(os.path.abspath(actual))) if actual else ''
    if actual_identity and actual_identity != expected_identity:
        return ''
    return str(body['id'])

def oc_config_fingerprint(conf=None):
    """常驻外壳/探活缓存的配置身份；只存 Key 摘要，不存秘密。"""
    conf = conf or read_json(conf_path(), {})
    eng = conf.get('engine') or {}
    up = s2_conf(conf)
    raw = {'kind': eng.get('kind', 's2'), 'model': up['model'], 'base': up['base_url'],
           'verify': bool(up['verify_ssl']), 'wire': up['wire'],
           'cli': eng.get('cli_path') or '',
           'credential': hashlib.sha256((up.get('api_key') or '').encode()).hexdigest()[:16]}
    return hashlib.sha256(json.dumps(raw, sort_keys=True).encode()).hexdigest()

def session_runtime_compatible(snapshot, conf=None):
    """续做只能回到创建该会话的同一运行身份。

    新任务保存整体指纹（包含凭据的不可逆摘要，不存 Key）；旧任务没有
    该字段时，至少比对当时已落盘的模型、网关、协议和 TLS 策略。
    """
    snapshot = snapshot if isinstance(snapshot, dict) else {}
    conf = conf or read_json(conf_path(), {})
    frozen = str(snapshot.get('runtime_fingerprint') or '')
    if frozen:
        return hmac.compare_digest(frozen, oc_config_fingerprint(conf))
    eng = conf.get('engine') or {}
    up = s2_conf(conf)
    current = {'kind': eng.get('kind', 's2'), 'model': up['model'],
               'base_url': up['base_url'], 'wire': up['wire'],
               'verify_ssl': bool(up['verify_ssl'])}
    for key, value in current.items():
        if key in snapshot and snapshot.get(key) not in (None, '') and snapshot.get(key) != value:
            return False
    return True

def invalidate_oc_runtime():
    """设置变化后关掉旧 server 并清探活；下一单会按新配置干净重启。"""
    with OC_LOCK:
        p = OC.get('proc')
        if p:
            try:
                if p.poll() is None: kill_tree(p)
            except Exception: pass
        OC.update({'proc': None, 'port': 0, 'base': '', 'pw': '', 'fingerprint': ''})
        try: _OC_PROBED.update({'ok': False, 'why': '', 'ts': 0.0, 'fingerprint': ''})
        except NameError: pass

def oc_serve(eng=None):
    """确保常驻 opencode server 活着,返回 base url;起不来返回 ''(上层回落 CLI 模式)。

    **healthy 不等于能用**:第一次接这个模式时踩过 —— 环境变量没传进去,
    /global/health 照样返回 healthy,真正调模型才 401,而 opencode 把它包成一句
    `UnknownError / Unexpected server error`,完全看不出是 Key 没送到。
    所以起完必须真调一次模型探活。"""
    with OC_LOCK:
        conf = read_json(conf_path(), {})
        if eng is not None:
            conf = dict(conf); conf['engine'] = eng
        fingerprint = oc_config_fingerprint(conf)
        if (OC['proc'] and OC['proc'].poll() is None and OC['base']
                and OC.get('fingerprint') == fingerprint):
            return OC['base']
        if OC.get('proc'):
            try:
                if OC['proc'].poll() is None: kill_tree(OC['proc'])
            except Exception: pass
            OC.update({'proc': None, 'port': 0, 'base': '', 'pw': '', 'fingerprint': ''})
        eng = eng if eng is not None else (conf.get('engine') or {})
        cli = resolve_cli('opencode', eng)
        if not cli: return ''
        port = _free_port()
        env = agent_env(eng)
        if env.get('BIDDOG_SHELL_CONF_ERR'): return ''
        pw = secrets.token_urlsafe(24)
        env['OPENCODE_SERVER_PASSWORD'] = pw
        logf = open(os.path.join(DATA, 'opencode-server.log'), 'a', encoding='utf-8')
        try:
            p = subprocess.Popen([cli, 'serve', '--port', str(port), '--hostname', '127.0.0.1'],
                                 stdout=logf, stderr=logf, stdin=subprocess.DEVNULL,
                                 env=env, cwd=DATA, **DETACH)
        except Exception:
            logf.close()
            return ''
        logf.close()  # 子进程已复制句柄；父进程不应为常驻 server 永久泄漏一个文件描述符
        OC.update({'proc': p, 'port': port, 'base': 'http://127.0.0.1:%d' % port,
                   'pw': pw, 'fingerprint': fingerprint})
        for _ in range(60):
            if p.poll() is not None: OC.update({'proc': None, 'base': '', 'pw': '', 'fingerprint': ''}); return ''
            st, _b = oc_api('/global/health', timeout=3)
            if st == 200: break
            time.sleep(0.5)
        else:
            kill_tree(p); OC.update({'proc': None, 'base': '', 'pw': '', 'fingerprint': ''}); return ''
        return OC['base']

def oc_probe():
    """真调一次模型确认整条链路通(healthy 骗不了这一步)。返回 (ok, 说明)。"""
    sid = oc_create_session(DATA, '连通性探活')
    if not sid: return False, '建不出目录隔离会话'
    pst, _reply = oc_api('/session/%s/prompt_async' % sid,
                         {'model': oc_model(),
                          'parts': [{'type': 'text', 'text': '只回两个字:可用'}]}, timeout=90)
    if pst not in (200, 201, 202, 204):
        return False, '执行外壳探活请求失败(HTTP %s)' % pst
    # 必须等它真的答完再看有没有报错。以前这里用 oc_busy 判,而 oc_busy 恒为 False,
    # 于是一秒就往下走、在回复还没到的时候读消息、没看见 error 就宣布「通了」——
    # Key 是错的也照样通过,这个探活等于白做。
    done = False
    for _ in range(90):
        time.sleep(1)
        done, turn_error = oc_turn(sid)
        if turn_error: return False, turn_error
        if done: break
    if not done: return False, '执行外壳探活 90 秒没有完整回复'
    ms, order = oc_messages(sid)
    if not ms:
        return False, '执行外壳探活没有拿到有效回复'
    for m in ms:
        info = oc_message_info(m)
        err = oc_message_error(info)
        if err:
            if '401' in err or 'incorrect API key' in err:
                return False, 'Key 没送到执行外壳或 Key 无效(原文:%s)' % err[:120]
            return False, err[:160]
    top = oc_latest_assistant(ms, order)
    if not top or top.get('finish') not in OC_CLEAN_FINISH:
        return False, '执行外壳探活没有干净收尾'
    return True, ''

# 一轮只有 stop 算干净收尾；长度截断、内容拦截、错误和取消都必须显式失败。
# 'tool-calls' 不在里面 —— 那是「这段话说完了,接着还要调工具」。
OC_CLEAN_FINISH = {'stop'}
OC_FAILED_FINISH = {'length', 'content-filter', 'content_filter', 'error',
                    'aborted', 'abort', 'canceled', 'cancelled'}
OC_FINISH = OC_CLEAN_FINISH | OC_FAILED_FINISH
OC_QUIET = 25      # 已 finish 且事件流再静这么多秒,才算整单收工
OC_SLOW = float(os.environ.get('BIDDOG_OC_SLOW_SECONDS', 180))
OC_STALL = float(os.environ.get('BIDDOG_OC_STALL_SECONDS', 8 * 60))
OC_AUTO_RECOVER = max(0, min(6, int(os.environ.get('BIDDOG_OC_AUTO_RECOVER', 3))))
# 整单自动恢复的硬上限:OC_AUTO_RECOVER 是「连败」预算(任务有真实进展就清零重来),
# TOTAL 是整单封顶。旧版整单只给 2 次、几小时的任务断两次流就永久停——
# 「6 次生成只成 1 次」的主因之一;新策略偶发断流可多次自愈,持续故障也不会热循环烧 Key。
OC_AUTO_RECOVER_TOTAL = max(OC_AUTO_RECOVER, min(20, int(os.environ.get('BIDDOG_OC_AUTO_RECOVER_TOTAL', 12))))

AUTO_RECOVERY_PROMPT = (
    '上一轮模型连接中断了，现在从已保存内容继续。先检查任务目录，已经写好的不要重写、'
    '不要推翻；只补没完成的章节和表格，然后继续汇总、质检并导出最终 Word。')


def _recoverable_turn_error(error):
    """网络/网关的瞬时故障都算可恢复——这类错误重试大概率能过,直接停等于把整单交还给用户。
    Key 无效、权限、配额彻底用尽这类确定性失败不在其列(重试也没用,该报就报)。"""
    text = str(error or '').lower()
    return any(marker in text for marker in (
        'stream idle timeout', 'no data received within configured window',
        'connection reset', 'unexpected_eof', 'unexpected eof',
        'stream disconnected before completion', 'eof occurred in violation',
        'incompleteread', 'timed out', 'timeout',
        # 以下为真机/网关常见瞬时故障口径(429 限流、5xx 网关抖动、连接层断开)
        'econnreset', 'econnrefused', 'socket hang up', 'fetch failed',
        'network error', 'connection error', 'connection aborted',
        'remote end closed', 'server disconnected', 'temporarily unavailable',
        'too many requests', 'rate limit', 'overloaded', 'over capacity',
        'bad gateway', 'service unavailable', 'gateway timeout',
        ' 429', ' 502', ' 503', ' 504', '(429', '(502', '(503', '(504',
        'http 429', 'http 502', 'http 503', 'http 504',
        'status 429', 'status 502', 'status 503', 'status 504'))

def _server_fallback_safe(job):
    """Only replay through the stable CLI path before any bid body exists, once per job."""
    runtime = read_json(os.path.join(job, 'runtime.json'), {})
    if int(runtime.get('fallback_count') or 0) >= 1: return False
    return not _body_mds(job) and not _body_docxs(job)


def oc_message_info(message):
    if isinstance(message, dict) and isinstance(message.get('info'), dict):
        return message['info']
    return message if isinstance(message, dict) else {}


def oc_message_error(info):
    error = info.get('error') if isinstance(info, dict) else None
    if not isinstance(error, dict): return str(error or '')
    if error.get('message'): return str(error['message'])
    data = error.get('data') if isinstance(error.get('data'), dict) else {}
    return str(data.get('message') or '')


def oc_messages(sid):
    """Return message rows plus ordering for the pinned OpenCode contract.

    1.18.18's standard endpoint returns chronological ``{info, parts}`` rows.
    The legacy compatibility endpoint may still be present and returns newest
    first.  Empty standard results are not proof of an empty session, so fall
    back before declaring the turn unfinished.
    """
    st, rows = oc_api('/session/%s/message' % sid, timeout=30)
    if st == 200 and isinstance(rows, list) and rows:
        order = 'chronological' if any(isinstance(row, dict) and isinstance(row.get('info'), dict)
                                       for row in rows) else 'legacy'
        return rows, order
    st, rows = oc_api('/api/session/%s/message' % sid, timeout=30)
    return (rows, 'legacy') if st == 200 and isinstance(rows, list) else ([], 'legacy')


def oc_latest_assistant(rows, order):
    infos = [oc_message_info(row) for row in rows]
    assistants = [info for info in infos
                  if (info.get('role') or info.get('type')) == 'assistant']
    if not assistants: return {}
    return assistants[-1] if order == 'chronological' else assistants[0]

OC_TURN_SIG = {}   # sid → 最近一次 oc_turn 观察到的消息活动签名(第三路心跳的旁路通道)

def oc_turn(sid):
    """这一轮跑完了没有,返回 (done, 出错说明)。

    **不要用 GET /session/status 判忙** —— 1.18.13 与 1.18.18 上它都恒返回 `{}`。这不是猜:
    真机连续 150 秒每 10 秒取一次,15 次全是空的。第一版 oc_run 就栽在这儿:
    busy 恒为 False,8 秒空转直接宣布「跑完了、任务目录里没有生成任何交付物」,
    而那会儿 agent 正在读素材、转招标文件,后面还有十来步要走。
    schema 里那个 session.idle 事件也指望不上,这版**一次都不发**(实测整轮 0 次)。

    真正能判的是消息列表；适配层会兼容新版正序与旧版倒序返回 ——
        finish=None          还在生成
        finish='tool-calls'  这段说完了,后面还要调工具,没完
        finish='stop'        这一轮真收工了

    顺带把「消息条数 + 最新 assistant 消息摘要」写进 OC_TURN_SIG[sid]:SSE 半开、
    长章节又暂时不写文件时,事件流和文件签名双双静默——但消息体在长(流式 token
    在涨),签名一直在变。oc_run 以它作第三路心跳,不再把正常长跑误判成卡死、
    亲手打断好端端的一轮(「任务中断误报」反馈的主根因)。"""
    ms, order = oc_messages(sid)
    if not ms:
        return False, ''          # 取不到就当没跑完:宁可多等,也不能误判成「没产出」
    top = oc_latest_assistant(ms, order)
    try:
        OC_TURN_SIG[sid] = '%d:%s' % (len(ms), hashlib.md5(
            json.dumps(top, sort_keys=True, ensure_ascii=False, default=str).encode('utf-8', 'ignore')).hexdigest())
    except Exception:
        OC_TURN_SIG[sid] = str(len(ms))
    if not top: return False, ''
    err = oc_message_error(top)
    if err: return True, str(err)[:300]
    finish = top.get('finish')
    if finish in OC_CLEAN_FINISH: return True, ''
    if finish == 'length': return True, '执行外壳输出达到长度上限，本轮没有完整收尾'
    if finish in ('content-filter', 'content_filter'): return True, '执行外壳回复被内容策略拦截，本轮未完成'
    if finish in ('aborted', 'abort', 'canceled', 'cancelled'): return True, '执行外壳本轮已取消，未算完成'
    if finish == 'error': return True, '执行外壳本轮报错，未算完成'
    return False, ''

def collect_oc_usage(job, sid):
    """Persist session-level model usage for the task card and diagnostics.

    OpenCode exposes usage on assistant messages. Re-reading the complete session
    makes this idempotent across resume/reconnect instead of double-counting calls.
    """
    if not sid or not os.path.isdir(job): return _job_usage(job)
    messages, _order = oc_messages(sid)
    if not messages: return _job_usage(job)
    calls = input_tokens = output_tokens = 0
    estimated_cost = 0.0
    has_cost = False
    for message in messages:
        if not isinstance(message, dict): continue
        info = message.get('info') if isinstance(message.get('info'), dict) else message
        if info.get('type') != 'assistant': continue
        tokens = info.get('tokens') if isinstance(info.get('tokens'), dict) else {}
        calls += 1
        input_tokens += int(tokens.get('input') or tokens.get('prompt') or tokens.get('input_tokens') or 0)
        output_tokens += int(tokens.get('output') or tokens.get('completion') or tokens.get('output_tokens') or 0)
        try:
            if info.get('cost') is not None:
                estimated_cost += float(info.get('cost') or 0)
                has_cost = True
        except (TypeError, ValueError):
            pass
    snapshot = (read_json(os.path.join(job, '任务.json'), {}) or {}).get('engine_snapshot') or {}
    usage = {'model': snapshot.get('model') or (oc_model() or {}).get('modelID') or '',
             'calls': calls, 'input_tokens': input_tokens, 'output_tokens': output_tokens,
             'total_tokens': input_tokens + output_tokens,
             'estimated_cost': round(estimated_cost, 6) if has_cost else None,
             'currency': 'USD' if has_cost else None, 'updated_at': now()}
    write_json(os.path.join(job, 'usage.json'), usage)
    return usage

def oc_session(job, directory):
    """给任务建/取会话。会话 id 落进 任务.json —— 崩了、关窗了都能凭它续跑。"""
    meta = read_json(os.path.join(job, '任务.json'), {})
    sid = meta.get('oc_session')
    if sid:
        st, s = oc_api('/api/session/%s' % sid, timeout=20)
        if st == 200 and isinstance(s, dict) and s.get('id'): return sid
    sid = oc_create_session(directory, meta.get('name') or os.path.basename(job))
    if sid:
        meta['oc_session'] = sid
        write_json(os.path.join(job, '任务.json'), meta)
    return sid

def oc_send(sid, text, delivery='queue', model=None, job=None):
    """给会话投一条消息。**这个接口是异步的:0 秒返回 200,不等跑完。**

    无显式 model 时使用兼容接口，delivery 控制运行中补充要求:
      queue = 排队,等当前这轮跑完再处理(实测第一条完整跑完、第二条随后被回应,两条都不丢)
      steer = 立刻引导当前这轮
    产品负责人定的「想到什么先记下来、下个环节带上」就是 queue。"""
    if model:
        # 1.18.18 的正式异步端点才接受 model + parts。旧
        # /api/.../prompt 的 schema 只收 prompt/delivery，把 model 静默丢掉后
        # 会回落到 OpenCode 默认模型，最终误报 Missing API key。
        path = '/session/%s/prompt_async' % sid
        payload = {'model': model, 'parts': [{'type': 'text', 'text': text}]}
    else:
        # 运行中补充要求继承该会话已固定的模型，并保留
        # 兼容端点的 queue/steer 交付语义。
        path = '/api/session/%s/prompt' % sid
        payload = {'delivery': delivery, 'prompt': {'text': text}}
    st, b = oc_api(path, payload, timeout=60)
    ok = st in (200, 201, 202, 204)
    if ok and job: _mark_skill_accepted(job, 'opencode')
    return ok, b

def oc_interrupt(sid):
    st, _ = oc_api('/api/session/%s/interrupt' % sid, {}, timeout=20)
    return st in (200, 201, 202, 204)

# 事件 → 人话。opencode 的事件类型很细,但用户要看的只有「它此刻在动哪个文件、跑什么命令」。
_OC_TOOL_CN = {'read': '读', 'write': '写', 'edit': '改', 'bash': '执行', 'glob': '找文件',
               'grep': '搜内容', 'task': '派子 agent', 'skill': '用技能', 'webfetch': '抓网页',
               'todowrite': '记待办', 'question': '问你一个问题'}

def _oc_line(ev):
    """把一条 opencode 事件翻成一行台词;不值得显示的返回 None。"""
    t = ev.get('type') or ''
    p = ev.get('properties') or {}
    if t.endswith('message.part.updated') or t == 'message.part.updated':
        part = p.get('part') or {}
        if part.get('type') != 'tool': return None
        stt = (part.get('state') or {})
        if stt.get('status') != 'running': return None
        tool = part.get('tool') or ''
        inp = stt.get('input') or {}
        arg = (inp.get('filePath') or inp.get('command') or inp.get('pattern')
               or inp.get('description') or inp.get('prompt') or '')
        arg = str(arg).replace('\n', ' ')[:90]
        return '%s %s' % (_OC_TOOL_CN.get(tool, tool), arg) if arg else _OC_TOOL_CN.get(tool, tool)
    if 'reasoning.started' in t: return '思考中…'
    if 'agent.switched' in t: return '切换角色:%s' % (p.get('agent') or '')
    if 'session.error' in t or t == 'session.error':
        return '⚠ %s' % str((p.get('error') or {}).get('message') or p)[:140]
    return None

def oc_watch(job, sid, stop, beat=None):
    """订阅单会话的 SSE,把事件流转成我们的台词 / 提问 / 报错,顺带当心跳。

    这是「它正在做什么」第一次对 opencode 真的有内容 —— 之前那块靠增量读 stdout,
    而 opencode 的子 agent 输出根本不走 stdout(实测 run.log 停在 8KB 不动)。

    beat 是给 oc_run 判收工用的心跳:每收到一条事件就盖个时间戳。光看消息 finish
    会在排队的两轮之间踩空,得配合「事件流也静下来了」一起判。"""
    buf, asked, last, flushed, admitted = [], set(), '', time.time(), 0
    fail_streak, last_fail_note = 0, 0.0
    # 外壳版本行为不一致：有的重连会从第 1 条完整重放，有的只发新事件。
    # 用事件内容指纹去重，不再用“本连接序号 <= 上次总数”判断，否则
    # new-only 连接的所有新事件都会被当作历史跳过。
    seen_keys, seen_order = set(), []
    def seen_before(ev):
        raw = json.dumps(ev, ensure_ascii=False, sort_keys=True, separators=(',', ':'))
        key = hashlib.sha256(raw.encode('utf-8')).hexdigest()
        if key in seen_keys: return True
        seen_keys.add(key); seen_order.append(key)
        if len(seen_order) > 5000:
            seen_keys.discard(seen_order.pop(0))
        return False
    def flush():
        nonlocal buf, flushed
        if buf: emit(job, {'type': 'worklog', 'lines': buf[-8:]}); buf = []
        flushed = time.time()
    while not stop.is_set():
        try:
            # 每次重连现取 base 与口令:server 重启换端口/口令后,旧 url 会 401 死循环、
            # 台词永久静音(裸 except 又吞掉一切,用户只看到「像死机了」)。
            # timeout 是 socket 级的:半开连接(对端悄悄断了)最多静默这么久就会抛错重连。
            # 旧值 600 秒比 OC_STALL(480 秒)还长,SSE 假死会先触发误报中断;180 秒能在
            # 判卡死之前把事件流救回来。正常长思考期间超时重连无害——事件按内容指纹去重。
            url = OC['base'] + '/api/session/%s/event' % sid
            req = urllib.request.Request(url, headers=oc_auth())
            with urllib.request.urlopen(req, timeout=180) as r:
                fail_streak = 0
                for raw in r:
                    if stop.is_set(): break
                    line = raw.decode('utf-8', 'ignore').strip()
                    if not line.startswith('data:'): continue
                    try: ev = json.loads(line[5:].strip())
                    except Exception: continue
                    if seen_before(ev): continue
                    if beat is not None: beat['ts'] = time.time()
                    # 证据判断与工作台词分离：只把 completed + read + 精确 SKILL.md 指纹
                    # 写成结构化收据，原始 filePath/command/prompt 均不落盘。
                    _observe_oc_skill_event(job, ev, sid)
                    t = ev.get('type') or ''
                    if 'question.asked' in t:
                        q = (ev.get('properties') or {})
                        rid = q.get('id') or q.get('requestID')
                        if rid and rid not in asked:
                            asked.add(rid); oc_emit_question(job, sid, rid, q)
                        continue
                    if 'prompt.admitted' in t:
                        # 第一条是我们自己派下去的活,别对着用户说「收到你补充的要求」——
                        # 他什么都还没补充呢。从第二条起才是真的插话。
                        admitted += 1
                        if admitted > 1: buf.append('收到你补充的要求,已排进队列')
                        continue
                    ln = _oc_line(ev)
                    if ln and ln != last: buf.append(ln); last = ln
                    # 攒够三条、或者隔了三秒就发一次:只按条数发的话,
                    # 它慢下来时台词会长时间一片空白,看着像死机了
                    if buf and (len(buf) >= 3 or time.time() - flushed >= 3): flush()
        except Exception as e:
            fail_streak += 1
            if time.time() - last_fail_note > 60:      # 同类错误 60 秒只记一次,不刷诊断流水
                last_fail_note = time.time()
                append_diagnostic(job, 'oc_watch_reconnect', str(e)[:200], attempt=fail_streak)
            if fail_streak == 5:
                emit(job, {'type': 'worklog',
                           'lines': ['实时台词通道暂时不可用,正在重连;生成进度不受影响']})
        flush()
        if not stop.is_set(): time.sleep(2)   # 断了就重连,别让台词永久静音

def oc_emit_question(job, sid, rid, q):
    """agent 主动提问 → 推给前端。

    这是七条「界面承诺、链路没接」里唯一能靠 opencode 直接治好的一条:
    以前 agent 问了、用户答了,答案写进 answers.jsonl 就没人读了。
    现在答案会经 /question/{id}/reply 真的回到它的上下文里。"""
    qs = q.get('questions') or []
    first = qs[0] if qs else {}
    text = first.get('question') or q.get('question') or '需要你确认一件事'
    opts = [o.get('label') if isinstance(o, dict) else str(o)
            for o in (first.get('options') or [])]
    meta = read_json(os.path.join(job, '任务.json'), {})
    pend = meta.setdefault('oc_questions', {})
    pend[rid] = {'session': sid, 'text': text, 'ts': now()}
    write_json(os.path.join(job, '任务.json'), meta)
    # 只报 agent 真正给的选项。「我来输入」是界面自己的自救入口,由前端统一补 ——
    # 放在这里等于让每一个发问题事件的地方都记得拼一次,漏一处用户就没法回答了。
    emit(job, {'type': 'question', 'id': rid, 'text': text, 'options': opts})

def oc_answer(job, rid, text):
    """把用户的回答送回 agent。返回 (ok, 说明)。"""
    meta = read_json(os.path.join(job, '任务.json'), {})
    rec = (meta.get('oc_questions') or {}).get(rid)
    if not rec: return False, '这个问题不在待回答清单里(可能已经答过或已过期)'
    sid = rec.get('session')
    st, b = oc_api('/api/session/%s/question/%s/reply' % (sid, rid),
                   {'answers': [[text]]}, timeout=30)
    if st not in (200, 201, 202, 204):
        st, b = oc_api('/question/%s/reply' % rid, {'answers': [[text]]}, timeout=30)
    ok = st in (200, 201, 202, 204)
    if ok:
        meta.setdefault('oc_questions', {}).pop(rid, None)
        write_json(os.path.join(job, '任务.json'), meta)
    return ok, ('' if ok else '回传失败(HTTP %s):%s' % (st, str(b)[:120]))

def halt(job, why):
    """任务终止:保留它跑到哪一步,不要把进度清零。

    四种终态(手动停止 / 未能启动 / 没有产出 / agent 异常退出)以前一律 pct=0 step=0,
    于是前端看到「pct 0 且没到 100」就当成运行中,一直转圈、一直挂在「进行中」——
    跑挂了跟正在跑长得一模一样,用户会一直等下去,等到投标截止。
    停在第几步、跑了多久,是用户决定「重跑还是接着改」的依据,必须留着。"""
    prev = read_json(os.path.join(job, 'progress.json'), {})
    write_json(os.path.join(job, 'outcome.json'), {'state': 'stopped', 'reason': why, 'ts': now()})
    safe = sanitize_event(job, {'type': 'progress', 'stage': why,
                                'pct': min(int(prev.get('pct') or 0), 99),
                                'step': int(prev.get('step') or 0), 'total': 12})
    safe.update({'stage': why, 'terminal': True, 'verified': True})
    emit(job, safe)

OC_RUN_FALLBACK = 'fallback'
OC_RUN_COMPLETED = 'completed'
OC_RUN_INTERRUPTED = 'interrupted'
OC_RUN_CANCELLED = 'cancelled'

def oc_run(job, prompt, allow_cli_fallback=True):
    """用 OpenCode server 模式跑一单，并区分“未派发”与“派发后中断”。

    跟 CLI 模式(`run --auto`)的区别:
      · 看得见 —— SSE 事件流转成真实台词(CLI 模式下 opencode 的子 agent 根本不写 stdout)
      · 停得掉 —— interrupt 优雅中断,半截产物完整落盘,不是杀进程组
      · 问得着 —— agent 的提问推给用户,答案经 /question/reply 真的回到它上下文
      · 崩了能续 —— 会话 id 落在 任务.json,重开应用凭它接着做
    """
    base = os.path.basename(job)
    update_runtime(job, execution_path='opencode_starting', can_pause=False,
                   pause_disabled_reason='正在建立稳定连接，暂时不能暂停。')
    def fallback(code, detail):
        if allow_cli_fallback:
            compatibility_fallback(job, code, detail)
        else:
            append_diagnostic(job, code, detail, fallback='disabled_for_resume')
        return OC_RUN_FALLBACK
    if _cancel_requested(base): return OC_RUN_CANCELLED
    if not oc_serve(): return fallback('opencode_server_unavailable', 'OpenCode server did not become healthy')
    if _cancel_requested(base): return OC_RUN_CANCELLED
    sid = oc_session(job, job)
    if not sid: return fallback('opencode_session_unavailable', 'OpenCode session could not be created')
    if _cancel_requested(base):
        oc_interrupt(sid)
        return OC_RUN_CANCELLED
    ok, why = oc_probe_once()
    if _cancel_requested(base):
        oc_interrupt(sid)
        return OC_RUN_CANCELLED
    if not ok:
        mark_connection_check(job, False, '连接探活失败:%s' % (why or '未知原因'))
        return fallback('opencode_probe_failed', why)
    mark_connection_check(job, True, '连接已验证,生成进行中')
    update_runtime(job, execution_path='opencode_server', can_pause=True,
                   pause_disabled_reason='', session_id=sid, auto_recovery_count=0)
    stop = threading.Event()
    beat = {'ts': time.time()}          # oc_watch 每收到一条事件就刷新
    threading.Thread(target=oc_watch, args=(job, sid, stop, beat), daemon=True).start()
    snap = read_json(os.path.join(job, '任务.json'), {}).get('engine_snapshot') or {}
    pinned_model = {'providerID': 'biddog-s2', 'modelID': snap.get('model')} if snap.get('model') else oc_model()
    if _cancel_requested(base):
        stop.set(); oc_interrupt(sid); return OC_RUN_CANCELLED
    sent, send_detail = oc_send(sid, prompt, delivery='queue', model=pinned_model, job=job)
    if not sent:
        # 请求已经尝试派发，无法证明服务端完全没收到；禁止从头自动重放。
        append_diagnostic(job, 'opencode_dispatch_unconfirmed', send_detail,
                          level='error', session_id=sid)
        emit(job, {'type': 'error',
                   'text': '连接意外中断，任务已安全停下；已生成的内容都已保留。',
                   'actions': [{'act': 'resume', 'label': '从已保存内容继续'},
                               {'act': 'open_log', 'label': '查看诊断详情'}]})
        stop.set(); return OC_RUN_INTERRUPTED
    emit(job, {'type': 'message', 'role': 'agent',
               'text': '已经开始生成。过程中你随时可以在这里补要求——'
                       '**会排到当前这一步之后生效**,不会打断它。'})
    result = OC_RUN_INTERRUPTED
    agent_line = 0
    current_step = 0
    step_started = time.time()
    job_sig = cli_activity_signature(job)
    last_turn_sig = ''
    slow_notified = False

    auto_recoveries = 0       # 连败计数:任务出现真实进展(文件落盘)就清零
    total_recoveries = 0      # 整单累计,封顶 OC_AUTO_RECOVER_TOTAL
    recovery_wait_cycles = 0

    def auto_recover(reason, interrupt=False):
        """Resume the same durable session after a transient slow-stream failure.

        Replaying the original prompt or switching to CLI would redo expensive analysis and may
        overwrite chapters. A bounded continuation prompt instead inspects durable artifacts and
        resumes only the missing work.
        """
        nonlocal auto_recoveries, total_recoveries, recovery_wait_cycles, slow_notified, job_sig
        if auto_recoveries >= OC_AUTO_RECOVER: return False
        if total_recoveries >= OC_AUTO_RECOVER_TOTAL: return False
        if interrupt: oc_interrupt(sid)
        sent_again, detail = oc_send(sid, AUTO_RECOVERY_PROMPT, delivery='queue',
                                     model=pinned_model, job=job)
        if not sent_again:
            append_diagnostic(job, 'opencode_auto_recovery_dispatch_failed', detail,
                              level='warning', session_id=sid, attempt=auto_recoveries + 1)
            return False
        auto_recoveries += 1
        total_recoveries += 1
        # OpenCode may briefly keep the failed assistant message at the top while admitting the
        # continuation prompt. Give that durable queue five polling cycles before judging it again,
        # otherwise one stale error can burn both retries in two seconds.
        recovery_wait_cycles = 5
        slow_notified = False
        beat['ts'] = time.time()
        job_sig = cli_activity_signature(job)
        update_runtime(job, execution_path='opencode_server', can_pause=True,
                       pause_disabled_reason='', session_id=sid,
                       auto_recovery_count=auto_recoveries,
                       last_auto_recovery_at=now())
        append_diagnostic(job, 'opencode_auto_recovery', reason,
                          session_id=sid, attempt=auto_recoveries)
        mark_connection_check(job, True, '连接已自动恢复(本单第 %d 次续跑)' % total_recoveries)
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '模型连接刚才短暂中断，已自动从保存位置继续（本单第 %d 次自动续跑）。'
                           '已落盘内容不会重写。' % total_recoveries})
        return True

    def drain_progress():
        nonlocal agent_line, current_step, step_started
        agent_line, accepted = drain_agent_events(job, agent_line)
        for safe in accepted:
            if safe.get('type') != 'progress': continue
            try: step = int(safe.get('step') or 0)
            except (TypeError, ValueError): step = 0
            if step and step != current_step:
                if current_step: record_stage(current_step, min(time.time() - step_started, 3 * 3600))
                current_step, step_started = step, time.time()
        return accepted
    try:
        for _ in range(3 * 3600):
            if _cancel_requested(base):
                oc_interrupt(sid)  # 停本地监控之前先停 server 会话，不能留它后台继续耗 Key/写文件
                result = OC_RUN_CANCELLED
                break
            time.sleep(1)
            drain_progress()
            current_sig = cli_activity_signature(job)
            if current_sig != job_sig:
                job_sig = current_sig
                beat['ts'] = time.time()
                # 文件真的在落盘 = 上一段连败已经翻篇,连败预算清零;
                # 长任务里偶发的多次断流因此都能自愈,不会两次就把整单焊死。
                auto_recoveries = 0
            quiet = time.time() - beat['ts']
            if slow_notified and quiet < OC_SLOW: slow_notified = False   # 恢复活动后,下次变慢还要再提醒
            if recovery_wait_cycles:
                recovery_wait_cycles -= 1
                done, err = False, ''
            else:
                done, err = oc_turn(sid)
                # 第三路心跳:SSE 半开 + 长章节暂不写文件时,只有消息体在长。
                # 它变了就说明模型还在正常吐字,绝不能按「卡死」打断。
                turn_sig = OC_TURN_SIG.get(sid, '')
                if turn_sig and turn_sig != last_turn_sig:
                    last_turn_sig = turn_sig
                    beat['ts'] = time.time()
            # 用户可能在本轮请求等待期间点了停止。停止意图必须高于迟到的连接错误，
            # 否则会误报“已切换稳定通道继续”，甚至启动第二条执行路径。
            if _cancel_requested(base):
                oc_interrupt(sid)
                result = OC_RUN_CANCELLED
                break
            if err:
                if _recoverable_turn_error(err) and auto_recover(err):
                    continue
                mark_connection_check(job, False, '连接中断:%s' % str(err)[:120])
                oc_interrupt(sid)
                if allow_cli_fallback and _server_fallback_safe(job):
                    fallback('opencode_turn_interrupted', err)
                    result = OC_RUN_FALLBACK
                else:
                    append_diagnostic(job, 'opencode_turn_interrupted', err,
                                      level='error', session_id=sid)
                    emit(job, {'type': 'error',
                               'text': '连接中断，任务已安全停下；已生成的内容都已保留。',
                               'actions': [{'act': 'resume', 'label': '从已保存内容继续'},
                                           {'act': 'open_log', 'label': '查看诊断详情'}]})
                break
            # 收工要两条同时成立:最新一条 assistant 已经 finish,**并且**事件流也静了。
            # 只看 finish 会在排队的两轮之间踩空 —— 第一轮 stop 的那一瞬,
            # queue 里的第二条还没起来,会被当成整单跑完。
            if done and quiet >= OC_QUIET:
                result = OC_RUN_COMPLETED
                break
            if quiet >= OC_SLOW and not slow_notified:
                slow_notified = True
                append_diagnostic(job, 'opencode_slow',
                                  'No meaningful session or file activity for %.1f seconds' % quiet,
                                  session_id=sid)
                emit(job, {'type': 'message', 'role': 'agent',
                           'text': '模型响应比平时慢，正在持续检查连接；已有内容会自动保留。'})
            # 事件流彻底没动静又没 finish:多半卡住了,当面说,别干等三小时
            if quiet >= OC_STALL:
                detail = 'No meaningful session or file activity for %.1f seconds' % quiet
                if auto_recover(detail, interrupt=True):
                    continue
                mark_connection_check(job, False, '连接长时间无响应,任务已安全停下')
                oc_interrupt(sid)
                if allow_cli_fallback and _server_fallback_safe(job):
                    fallback('opencode_stalled', detail)
                    result = OC_RUN_FALLBACK
                else:
                    append_diagnostic(job, 'opencode_stalled', detail,
                                      level='error', session_id=sid)
                    emit(job, {'type': 'message', 'role': 'agent',
                               'text': '连接暂时没有响应，任务已安全停下；已生成的内容都已保留。',
                               'actions': [{'act': 'resume', 'label': '从已保存内容继续'},
                                           {'act': 'open_log', 'label': '查看诊断详情'}]})
                break
    finally:
        drain_progress()
        if current_step: record_stage(current_step, min(time.time() - step_started, 3 * 3600))
        collect_oc_usage(job, sid)
        OC_TURN_SIG.pop(sid, None)
        stop.set()
    return result


# ---------- 生成参数(可配,给 A/B 用) ----------
# 复读循环最爱的组合是「低温 + 无 penalty + 长输出」。默认值不动(没有证据不改默认),
# 但必须能配:tools/model_ab.py 跑同一份招标文件对比 temperature / frequency_penalty,
# 换参数才有数字可依。penalty 为 0 时不放进请求体,免得少数网关拒收未知字段。
GEN_PARAM_BOUNDS = {'temperature': (0.0, 2.0, 0.2),
                    'frequency_penalty': (-2.0, 2.0, 0.0),
                    'presence_penalty': (-2.0, 2.0, 0.0)}


def _normalize_generation_params(raw):
    out = {}
    if not isinstance(raw, dict): return out
    for key, (low, high, _default) in GEN_PARAM_BOUNDS.items():
        if raw.get(key) in (None, ''): continue
        try: value = float(raw.get(key))
        except (TypeError, ValueError): continue
        out[key] = round(max(low, min(high, value)), 3)
    return out


def _generation_params(conf=None):
    conf = conf if isinstance(conf, dict) else read_json(conf_path(), {})
    eng = conf.get('engine') if isinstance(conf.get('engine'), dict) else {}
    chosen = _normalize_generation_params(eng.get('generation_params'))
    return {key: chosen.get(key, default) for key, (_low, _high, default) in GEN_PARAM_BOUNDS.items()}


def _apply_generation_params(payload, params):
    payload['temperature'] = float(params.get('temperature', 0.2))
    for key in ('frequency_penalty', 'presence_penalty'):
        if params.get(key): payload[key] = float(params[key])
    return payload


# ---------- 章节并行度 ----------
# 以前 batch = ready[:2]:两章一批、等两章都写完才发下一批。10 章 × 2–4 分钟就是 20 分钟串行。
# 现在是持续 N 路在飞:一章写完立刻补位;某一章撞到 429 限流,这一单余下的章节自动降到 2 路。
PIPELINE_THROTTLED = set()          # 本轮撞过限流的任务 id


def _chapter_parallelism():
    try: value = int(os.environ.get('BIDDOG_CHAPTER_PARALLEL', 4))
    except (TypeError, ValueError): value = 4
    return max(1, min(8, value))


def _chapter_slots(base):
    limit = _chapter_parallelism()
    return min(limit, 2) if base in PIPELINE_THROTTLED else limit


def _used_chapter_headings(job, node, state=None):
    """其他已写完章节的小标题清单:传给撰写员,本章不得重复(SKILL 第 4 步的纪律,流水线以前没做)。"""
    state = state or generation_pipeline.load(job)
    out = []
    for item in state.get('nodes') or []:
        if item.get('id') == node.get('id') or item.get('state') != 'done': continue
        if not str(item.get('id') or '').startswith('chapter_write:'): continue
        for name in item.get('outputs') or []:
            path = os.path.join(job, os.path.basename(str(name)))
            try: text = open(path, encoding='utf-8', errors='ignore').read(200000)
            except OSError: continue
            for heading in prompts.used_headings_from_text(text):
                if heading not in out: out.append(heading)
    return out[:prompts.CHAPTER_MAX_USED_HEADINGS]


def _pipeline_mode(conf=None):
    conf = conf if isinstance(conf, dict) else read_json(conf_path(), {})
    eng = conf.get('engine') or {}
    explicit = str(eng.get('generation_mode') or '').lower()
    if explicit in ('standard', 'quality'): return 'standard'
    if explicit == 'fast': return 'fast'
    return 'standard' if str(s2_conf(conf).get('model') or '') == S2_QUALITY_MODEL else 'fast'


def _safe_chapter_name(index, title):
    clean = re.sub(r'[\\/:*?"<>|\x00-\x1f]+', '-', str(title or '')).strip(' .-_')
    clean = re.sub(r'\s+', '', clean)[:40] or ('章节%d' % index)
    return '章节_%02d_%s.md' % (index, clean)


def _pipeline_chapters(meta):
    snapshot = meta.get('template_snapshot') if isinstance(meta.get('template_snapshot'), dict) else {}
    package = snapshot.get('package') if isinstance(snapshot.get('package'), dict) else {}
    outline = package.get('outline') if isinstance(package.get('outline'), list) else []
    if not outline:
        outline = [{'title': title} for title in (
            '资格与符合性响应', '项目理解与总体方案', '技术与服务响应',
            '实施与验收方案', '商务响应与承诺')]
    chapters = []
    for index, item in enumerate(outline[:12], 1):
        item = item if isinstance(item, dict) else {'title': str(item)}
        chapters.append({'id': '%02d' % index,
                         'title': str(item.get('title') or ('章节%d' % index))[:120],
                         'purpose': str(item.get('purpose') or '')[:500],
                         'evidence': [str(value)[:120] for value in (item.get('evidence') or [])[:12]],
                         'output': _safe_chapter_name(index, item.get('title'))})
    return chapters


def _initialize_generation_pipeline(job, meta, conf):
    up = s2_conf(conf)
    mode = _pipeline_mode(conf)
    verified = _verified_model_ids(conf)
    selected = str(up.get('model') or S2_DEFAULT_MODEL)
    if not verified:
        raise generation_pipeline.PipelineError('没有已验证的生成模型，请先在“模型接入”重新测试连接')
    selected_verified = selected if selected in verified else ''
    fast_model = (S2_DEFAULT_MODEL if S2_DEFAULT_MODEL in verified else
                  (selected_verified or verified[0]))
    quality_model = (S2_QUALITY_MODEL if S2_QUALITY_MODEL in verified else
                     (selected_verified or fast_model))
    state = generation_pipeline.initialize(
        job,
        run_id=str(meta.get('run_id') or _new_run_id()),
        mode=mode,
        model_routes={'fast': fast_model, 'quality': quality_model},
        chapters=_pipeline_chapters(meta),
        credential_fingerprint=generation_pipeline.credential_fingerprint(up.get('api_key') or ''),
        base_url=up.get('base_url') or '',
    )
    meta['pipeline_version'] = generation_pipeline.PIPELINE_VERSION
    meta['generation_mode'] = mode
    write_json(os.path.join(job, '任务.json'), meta)
    return state


def _pipeline_register_session(job, sid):
    with PIPELINE_SESSIONS_LOCK:
        PIPELINE_SESSIONS.setdefault(os.path.basename(job), set()).add(sid)


def _pipeline_unregister_session(job, sid):
    with PIPELINE_SESSIONS_LOCK:
        sessions = PIPELINE_SESSIONS.get(os.path.basename(job))
        if not sessions: return
        sessions.discard(sid)
        if not sessions: PIPELINE_SESSIONS.pop(os.path.basename(job), None)


def _pipeline_interrupt_sessions(job):
    with PIPELINE_SESSIONS_LOCK:
        sessions = list(PIPELINE_SESSIONS.get(os.path.basename(job), set()))
    results = [_interrupt_or_finished(sid) for sid in sessions]
    return all(results) if results else True


def _pipeline_declared_input_names(job, node, state=None):
    """One source of truth for model-node file inputs, including DAG dependencies."""
    node_id = str(node.get('id') or '')
    names = ['招标文件_解析版.md', '你的要求.md']
    if node_id != 'response_plan':
        names.extend(['投标文件组成.md', '评分点响应矩阵.md', '废标风险清单.md',
                      'response_plan.json'])
    state = state or generation_pipeline.load(job)
    names.extend(generation_pipeline.dependency_outputs(state, node_id))
    if node_id == 'quality_review':
        names.extend([os.path.basename(path) for path in glob.glob(os.path.join(job, '*.md'))])
    result = []
    for name in names:
        safe = os.path.basename(str(name or ''))
        if safe and safe not in result:
            result.append(safe)
    return result


def _pipeline_material_paths(job):
    """列出本任务可供节点使用的素材文件,超限**截断**而不是抛异常。

    旧版超过 1000 个文件/100MB 直接 raise,而 merged_materials 会把整个全局素材库
    补拷进 job/素材——素材库养得越大,任务越必然「分段生成异常」,且每次续跑都
    死在同一处(真机「传得越勤越失败」的隐蔽根源)。截断优先保留文本类规范素材,
    并记诊断供排查。"""
    materials = os.path.join(job, '素材')
    result = []
    if not os.path.isdir(materials):
        return result
    count = 0; total = 0; truncated = 0
    for parent, dirs, files in os.walk(materials, followlinks=False):
        dirs[:] = [name for name in dirs if not os.path.islink(os.path.join(parent, name))]
        for name in sorted(files):
            source = os.path.join(parent, name)
            if os.path.islink(source): continue
            try: size = os.path.getsize(source)
            except OSError: continue
            if count >= 1000 or total + size > 100 * 1024 * 1024:
                truncated += 1
                continue
            count += 1; total += size
            result.append(source)
    if truncated:
        append_diagnostic(job, 'material_limit_truncated',
                          '素材超过节点安全上限,已按序保留 %d 个、跳过 %d 个;'
                          '可在素材库精简后重跑' % (count, truncated))
    return result


# ---------- 素材按章检索 ----------
# 以前每个章节节点都把整个 素材/ 拷进尝试目录,再按文件名顺序读到 140k 预算——素材库越大,
# 排在后面的文件越被截光,而且每一章都在重复读同一堆无关的章节模板(token 成本大头)。
# 现在:根目录的规范素材(能力表/应答要点/公司介绍/资质案例/图片索引)永远带上;
# 子目录(章节模板/、任务导入/)按本章标题、素材槽位、评分点的关键词命中排序,装到预算为止;
# 图片等二进制不进节点目录——模型读不到,导出 Word 时从 job/素材 取。
MATERIAL_CANONICAL = ('产品能力表.md', '应答要点.md', '公司介绍.md', '产品资料.md', '资质与案例.md',
                      '图片索引.md', '图片索引.json')
MATERIAL_TEXT_EXTS = ('.md', '.txt', '.json', '.csv', '.docx', '.pdf', '.html', '.htm', '.doc')


def _material_terms(node):
    raw = [node.get('title')] + list(node.get('material_slots') or []) + \
          list(node.get('scoring_points') or []) + list(node.get('basis') or [])
    terms = set()
    for value in raw:
        for piece in re.split(r'[\s、,，;；:：()（）/／|｜\-—_.。\d〔〕「」]+', str(value or '')):
            piece = piece.strip()
            if len(piece) >= 2: terms.add(piece)
            # 中文长词整词常撞不上,再切 2-gram
            for i in range(len(piece) - 1):
                gram = piece[i:i + 2]
                if re.fullmatch(r'[\u4e00-\u9fff]{2}', gram): terms.add(gram)
    return terms


def _pipeline_select_materials(job, node, paths):
    """返回 (按相关性排好序的素材路径, 跳过数)。"""
    materials = os.path.join(job, '素材')
    budget = max(10000, int(os.environ.get('BIDDOG_CHAPTER_MATERIAL_CHARS', 60000)))
    canonical, ranked, skipped = [], [], 0
    terms = _material_terms(node)
    for source in paths:
        rel = os.path.relpath(source, materials).replace(os.sep, '/')
        name = os.path.basename(source)
        if os.path.splitext(name)[1].lower() not in MATERIAL_TEXT_EXTS:
            skipped += 1; continue
        if '/' not in rel or name in MATERIAL_CANONICAL:
            canonical.append(source); continue
        haystack = name + '\n' + _pipeline_model_file_text(source, 2000)
        score = sum((3 if term in name else 1) for term in terms if term in haystack)
        ranked.append((score, rel, source))
    ranked.sort(key=lambda item: (-item[0], item[1]))
    chosen, used = list(canonical), 0
    for score, _rel, source in ranked:
        if score <= 0:
            skipped += 1; continue
        try: size = os.path.getsize(source)
        except OSError: continue
        if used + size > budget:
            skipped += 1; continue
        chosen.append(source); used += size
    return chosen, skipped


def _pipeline_copy_inputs(job, node, target):
    """Materialize the exact read set for one model attempt; never use symlinks.

    返回按「该先给模型看什么」排好的相对路径:声明输入(解析版/要求/规划/前置章节)在前,
    规范素材其次,按相关性挑出来的章节模板最后——预算不够时被截掉的是最不相干的。"""
    node_id = str(node.get('id') or '')
    copied, ordered = set(), []
    for name in _pipeline_declared_input_names(job, node):
        safe = os.path.basename(str(name or ''))
        if not safe or safe in copied: continue
        source = os.path.join(job, safe)
        if os.path.isfile(source) and not os.path.islink(source):
            shutil.copy2(source, os.path.join(target, safe)); copied.add(safe); ordered.append(safe)
    skill_file = os.path.join(skill_dir_conf(read_json(conf_path(), {})), 'SKILL.md')
    if os.path.isfile(skill_file) and not os.path.islink(skill_file):
        shutil.copy2(skill_file, os.path.join(target, 'SKILL.md'))
    materials = os.path.join(job, '素材')
    if node_id.startswith('chapter_write:'):
        chosen, skipped = _pipeline_select_materials(job, node, _pipeline_material_paths(job))
        for source in chosen:
            rel = os.path.relpath(source, materials)
            destination = os.path.join(target, '素材', rel)
            os.makedirs(os.path.dirname(destination), exist_ok=True)
            shutil.copy2(source, destination)
            ordered.append(os.path.join('素材', rel).replace(os.sep, '/'))
        if skipped:
            append_diagnostic(job, 'material_selected',
                              '本章按相关性带入 %d 份素材,跳过 %d 份(二进制或与本章无关)' % (len(chosen), skipped),
                              level='info', node_id=node_id)
    return ordered


def _pipeline_model_file_text(path, budget):
    """Return bounded text for a model node; binary office files are parsed locally."""
    try:
        suffix = os.path.splitext(path)[1].lower()
        if suffix in ('.md', '.txt', '.json', '.csv'):
            return open(path, encoding='utf-8', errors='ignore').read(max(0, budget))
        if suffix in ('.docx', '.pdf'):
            blob = open(path, 'rb').read(MAX_TEMPLATE_UPLOAD_BYTES + 1)
            if len(blob) > MAX_TEMPLATE_UPLOAD_BYTES: return ''
            _headings, _tables, text = extract_document_structure(os.path.basename(path), blob)
            return str(text or '')[:max(0, budget)]
        if suffix in ('.html', '.htm'):
            blob = open(path, 'rb').read(MAX_TEMPLATE_UPLOAD_BYTES + 1)
            return html_to_text(blob)[:max(0, budget)]
        if suffix == '.doc':
            # 旧格式素材以前被静默丢弃——用户看着「上传成功」,内容却一个字进不了模型。
            # 伪 doc(zip)与 RTF 都能本地解;真 OLE 返回空,由上传侧提示转 docx。
            blob = open(path, 'rb').read(MAX_TEMPLATE_UPLOAD_BYTES + 1)
            if len(blob) > MAX_TEMPLATE_UPLOAD_BYTES: return ''
            if blob[:4] == b'PK\x03\x04':
                _headings, _tables, text = extract_document_structure(
                    os.path.basename(path) + 'x', blob)      # 加 x 让结构解析按 docx 分支走
                return str(text or '')[:max(0, budget)]
            if blob[:5] == b'{\\rtf':
                return rtf_to_text(blob)[:max(0, budget)]
            return ''
    except Exception:
        return ''
    return ''


def _pipeline_json_object(text):
    raw = str(text or '').strip()
    candidates = []
    try:
        value = json.loads(raw)
        if isinstance(value, dict): candidates.append(value)
    except (TypeError, ValueError):
        pass
    for match in re.finditer(r'```(?:json)?\s*(\{.*?\})\s*```', raw, re.I | re.S):
        try:
            value = json.loads(match.group(1))
            if isinstance(value, dict): candidates.append(value)
        except (TypeError, ValueError):
            pass
    decoder = json.JSONDecoder()
    for match in re.finditer(r'\{', raw):
        try:
            value, _end = decoder.raw_decode(raw[match.start():])
            if isinstance(value, dict): candidates.append(value)
        except (TypeError, ValueError):
            continue
    if not candidates: return None
    expected = {'files', 'composition', 'scoring_points', 'risks', 'chapter_guidance'}
    return max(candidates, key=lambda value: (len(expected & set(value)), len(value)))


def _pipeline_skill_contract(text, limit=6000):
    """Compile prose rules from the full Skill without sending tool scripts."""
    kept, fenced, priority, before_sections = [], False, False, True
    keywords = re.compile(r'必须|不得|严禁|事实|证据|评分|废标|偏离|质量|格式|字数|需补充|人工确认')
    for raw in str(text or '').splitlines():
        if raw.strip().startswith('```'):
            fenced = not fenced
            continue
        if fenced: continue
        line = raw.rstrip()
        if line.startswith('## '):
            before_sections = False
            priority = bool(re.search(r'成本控制|人工确认|边界', line))
        keep = (before_sections or priority or line.startswith('#') or keywords.search(line))
        if keep:
            kept.append(line[:800] if len(line) > 800 else line)
    value = '\n'.join(kept).strip()
    if len(value) <= limit: return value
    head = max(1, int(limit * 0.65)); tail = max(1, limit - head)
    return value[:head] + '\n\n[中间工具脚本已由引擎省略]\n\n' + value[-tail:]


def _pipeline_direct_task(job, node):
    node_id = str(node.get('id') or '')
    if node_id == 'response_plan':
        return ('分析招标要求，提取文件组成、评分点、废标风险和各模板章节的写作依据。'
                '只做响应规划，不写章节正文。')
    if node_id == 'chapter_write:technical_deviation':
        return ('生成技术应答偏离表，使用 Markdown 表格，列为“序号｜招标条款/技术要求｜'
                '投标响应｜偏离情况｜依据/证据｜备注”；逐项覆盖，未知事实标记〔需补充〕。')
    if node_id == 'chapter_write:business_deviation':
        return ('生成商务偏离表，使用 Markdown 表格，覆盖报价、付款、工期、质保、交付、签章等；'
                '列为“序号｜招标条款/商务要求｜投标响应｜偏离情况｜依据/证据｜备注”。')
    if node_id.startswith('chapter_write:'):
        return prompts.chapter_task(node)
    if node_id == 'quality_review':
        return prompts.review_task()
    return '完成节点“%s”，只输出本节点指定文件。' % (node.get('title') or node_id)


def _record_direct_skill_use(job, skill_path):
    meta = read_json(os.path.join(job, '任务.json'), {})
    manifest = meta.get('skill_manifest') if isinstance(meta.get('skill_manifest'), dict) else {}
    if (manifest.get('run_id') != meta.get('run_id')
            or manifest.get('path_sha256') != _sha256_text(os.path.realpath(skill_path))
            or manifest.get('manifest_sha256') != _file_digest(skill_path)):
        return False
    _mark_skill_accepted(job, 'direct_model_completion')
    _append_skill_event(job, {
        'run_id': meta.get('run_id'), 'type': 'read_manifest', 'status': 'completed',
        'path_sha256': manifest.get('path_sha256'),
        'manifest_sha256': manifest.get('manifest_sha256'), 'ts': now(),
    })
    return True


def _pipeline_cell(value):
    return str(value if value not in (None, '') else '〔需补充〕').replace('|', '｜').replace('\n', ' ')


def _pipeline_write_response_plan(job, attempt_root, compact, source='local', model='', notes=()):
    """Expand a small model analysis into the four deterministic planning artifacts.

    ``source`` 记录这份规划是谁给的:'local' = 本地关键词索引(候选,永远先落盘兜底),
    'model' = 模型逐项核对过的结构化结果。覆盖仪表和确认卡据此决定显示「真实覆盖率」
    还是「候选,待核对」——本地索引算出来的 0/N 不是事实,只会把人吓一跳。"""
    if not isinstance(compact, dict):
        raise generation_pipeline.NodeExecutionError('model_plan_invalid', retryable=True)
    state = generation_pipeline.load(job)
    provisional = [item for item in (state.get('nodes') or [])
                   if str(item.get('id') or '').startswith('chapter_write:')
                   and not str(item.get('id') or '').endswith(
                       (':technical_deviation', ':business_deviation'))]
    guidance = {}
    for item in compact.get('chapter_guidance') or []:
        if isinstance(item, dict) and item.get('title'):
            guidance[str(item['title']).strip()] = item
    chapters = []
    for item in provisional:
        title = str(item.get('title') or '')
        guide = guidance.get(title, {})
        chapters.append({
            'id': str(item.get('id') or '').split(':', 1)[-1],
            'title': title,
            'output': os.path.basename(str((item.get('outputs') or ['章节.md'])[0])),
            'basis': [str(value) for value in (guide.get('basis') or ['招标文件_解析版.md'])][:20],
            'scoring_points': [str(value) for value in (guide.get('scoring_points') or [])][:20],
            'material_slots': [str(value) for value in (guide.get('material_slots') or [])][:20],
            'dependencies': [str(value).split(':', 1)[-1]
                             for value in (item.get('dependencies') or [])][:20],
        })
    if not chapters:
        raise generation_pipeline.NodeExecutionError('model_plan_has_no_chapters', retryable=False)

    composition = [str(value) for value in (compact.get('composition') or []) if str(value).strip()][:40]
    source_line = (('> 规划来源:模型逐项核对(%s)' % model) if source == 'model'
                   else '> 规划来源:本地关键词索引(候选项,尚未经模型核对)')
    composition_lines = [
        '# 投标文件组成', '', source_line, '',
        '> 本目录根据场景模板与招标文件解析结果生成；客户专属事实和最终装订顺序须在提交前人工核对。', '',
        '## 文件组成建议', '',
    ]
    composition_lines.extend('- %s' % _pipeline_cell(value) for value in composition)
    composition_lines.extend(['', '## 正文章节与目的', '', '| 序号 | 章节 | 目的与依据 |', '|---|---|---|'])
    for index, chapter in enumerate(chapters, 1):
        basis = '；'.join(chapter['basis']) or '依据招标文件逐项响应'
        composition_lines.append('| %d | %s | %s |' % (index, _pipeline_cell(chapter['title']), _pipeline_cell(basis)))
    composition_lines.extend(['', '## 提交前人工确认', '',
                              '- 核对投标人名称、报价、资质、案例、人员、签章和日期。',
                              '- 所有〔需补充〕项完成后，方可作为正式投标文件提交。'])

    score_lines = ['# 评分点响应矩阵', '', source_line, '',
                   '> 分值或证据未在原文中明确时保留〔需补充〕，不得推测。', '']
    note_lines = ['> ' + str(item).strip() for item in (notes or []) if str(item).strip()]
    if note_lines:
        score_lines.extend(note_lines); score_lines.append('')
    score_lines.extend(['| 序号 | 原要求/评分点 | 分值 | 响应位置 | 证据 | 缺口 |',
                        '|---|---|---:|---|---|---|'])
    scores = [item for item in (compact.get('scoring_points') or []) if isinstance(item, dict)][:120]
    if not scores:
        scores = [{'requirement': '未识别到独立评分点', 'score': '未知',
                   'location': '全册复核', 'evidence': '招标文件_解析版.md',
                   'gap': '〔需补充〕人工核对评分办法'}]
    for index, item in enumerate(scores, 1):
        score_lines.append('| %d | %s | %s | %s | %s | %s |' % (
            index, _pipeline_cell(item.get('requirement')), _pipeline_cell(item.get('score') or '未知'),
            _pipeline_cell(item.get('location')), _pipeline_cell(item.get('evidence')),
            _pipeline_cell(item.get('gap'))))
    score_lines.extend(['', '## 使用说明', '',
                        '- 正文撰写按本矩阵逐项落位，质检阶段再次核对响应位置与证据。',
                        '- 评分分值、证明材料或客户事实缺失时必须保留〔需补充〕。'])

    risk_lines = ['# 废标风险清单', '', source_line, '',
                  '> 本清单用于提交前逐项人工核验，不替代招标文件原文和法律审查。', '',
                  '| 序号 | 类别 | 招标要求 | 风险 | 提交前动作 |', '|---|---|---|---|---|']
    risks = [item for item in (compact.get('risks') or []) if isinstance(item, dict)][:160]
    if not risks:
        risks = [{'category': '完整性', 'requirement': '逐项核对招标文件',
                  'risk': '遗漏实质性要求可能导致无效投标', 'action': '〔需补充〕人工复核'}]
    for index, item in enumerate(risks, 1):
        risk_lines.append('| %d | %s | %s | %s | %s |' % (
            index, _pipeline_cell(item.get('category')), _pipeline_cell(item.get('requirement')),
            _pipeline_cell(item.get('risk')), _pipeline_cell(item.get('action'))))
    risk_lines.extend(['', '## 门禁原则', '',
                       '- 资格、签章、报价、时限和实质性条款必须由投标负责人最终确认。',
                       '- 未取得证据的资质、案例、人员和承诺不得由模型补写。'])

    artifacts = {
        '投标文件组成.md': '\n'.join(composition_lines) + '\n',
        '评分点响应矩阵.md': '\n'.join(score_lines) + '\n',
        '废标风险清单.md': '\n'.join(risk_lines) + '\n',
        'response_plan.json': json.dumps({
            'chapters': chapters, 'source': source, 'model': str(model or ''),
            'notes': [str(item) for item in (notes or []) if str(item).strip()],
        }, ensure_ascii=False, indent=2) + '\n',
    }
    for name, text in artifacts.items():
        target = os.path.join(attempt_root, name)
        tmp = target + '.tmp'
        with open(tmp, 'w', encoding='utf-8') as handle: handle.write(text)
        os.replace(tmp, target)


def _pipeline_local_response_plan(job, attempt_root):
    """Build the non-creative response index locally and deterministically.

    The response plan is an index over the tender and the selected template.  It
    must never be a long-form generation dependency: a truncated model response
    used to strand the whole run before chapter writing had even begun.
    """
    parsed_path = os.path.join(attempt_root, '招标文件_解析版.md')
    parse_limit = max(300000, min(3000000, int(
        os.environ.get('BIDDOG_LOCAL_PLAN_MAX_CHARS', 1500000))))
    parsed = _pipeline_model_file_text(parsed_path, parse_limit)
    try:
        if os.path.getsize(parsed_path) > len(parsed.encode('utf-8')):
            append_diagnostic(
                job, 'local_plan_source_capped',
                '解析文本超过本地索引上限，已索引前 %d 字；交付前需人工核对尾部附件。' % parse_limit,
                level='warning')
    except OSError:
        pass
    lines = []
    for raw in parsed.splitlines():
        value = re.sub(r'\s+', ' ', str(raw or '')).strip(' #\t')
        if 8 <= len(value) <= 260 and value not in lines:
            lines.append(value)

    state = generation_pipeline.load(job)
    chapter_nodes = [node for node in (state.get('nodes') or [])
                     if str(node.get('id') or '').startswith('chapter_write:')]
    composition = ['投标函、资格证明及签章文件']
    composition.extend(str(node.get('title') or '') for node in chapter_nodes
                       if str(node.get('title') or '').strip())

    score_pattern = re.compile(r'(评分|评审|分值|得分|满分|\d+(?:\.\d+)?\s*分(?:\D|$))')
    risk_pattern = re.compile(r'(废标|无效投标|否决|取消.{0,12}资格|不得|必须|须提供|应提供)')
    all_score_lines = [line for line in lines if score_pattern.search(line)]
    all_risk_lines = [line for line in lines if risk_pattern.search(line)]
    score_lines = all_score_lines[:120]
    risk_lines = all_risk_lines[:160]
    if len(all_score_lines) > len(score_lines) or len(all_risk_lines) > len(risk_lines):
        append_diagnostic(
            job, 'local_plan_items_capped',
            '评分点或风险条目超过索引上限（评分 %d、风险 %d），已保留前 120/160 项。' %
            (len(all_score_lines), len(all_risk_lines)), level='warning')

    scores = [{
        'requirement': line,
        'score': (re.search(r'\d+(?:\.\d+)?\s*分', line).group(0)
                  if re.search(r'\d+(?:\.\d+)?\s*分', line) else '未知'),
        'location': '评分证据与对应章节',
        'evidence': '招标文件_解析版.md',
        'gap': '〔需补充〕提交前核对分值、响应位置和证据',
    } for line in score_lines]

    def risk_category(line):
        if re.search(r'(资格|资质|证明|案例)', line): return '资格与证据'
        if re.search(r'(签章|密封|报价|投标函)', line): return '形式与商务'
        if re.search(r'(工期|时限|日期|验收|质保)', line): return '交付与时限'
        return '实质性要求'

    risks = [{
        'category': risk_category(line),
        'requirement': line,
        'risk': '未逐项响应或缺少证据可能导致扣分、无效投标或取消资格',
        'action': '提交前由投标负责人对照原文核对；客户事实或证据缺失时保留〔需补充〕',
    } for line in risk_lines]

    guidance = []
    for node in chapter_nodes:
        title = str(node.get('title') or '')
        keywords = [word for word in re.split(r'[与、及/\s]+', title) if len(word) >= 2]
        basis = [line for line in lines if any(word in line for word in keywords)][:8]
        guidance.append({
            'title': title,
            'basis': basis or ['招标文件_解析版.md'],
            'scoring_points': [line for line in score_lines
                               if any(word in line for word in keywords)][:8],
            'material_slots': ['与本章要求对应的资质、案例、人员或承诺证据〔需补充〕'],
        })
    compact = {'composition': composition, 'scoring_points': scores,
               'risks': risks, 'chapter_guidance': guidance}
    _pipeline_write_response_plan(job, attempt_root, compact, source='local')
    return compact


# ---------- 响应规划:模型逐项核对(本地索引永远兜底) ----------
# 0.20.4 把规划改成纯本地,是为了「规划一步不能拖垮整单」(模型返回被截断,整单卡死在
# 章节撰写之前)。但纯本地的代价一直没人算:评分点矩阵只是含「分」字的行、「响应位置」
# 是一句常量、「缺口」全是〔需补充〕——于是覆盖仪表永远 0/N,按章补写一条都派不出去,
# 章节提示词里的「评分点」是含'分'字的句子而不是本章要拿的分。
# 现在分两步:本地索引先落盘(保底,零 token、秒级);模型只做「核对」——逐项提取评分点、
# 把每一项落到真实章节标题上、给出分值与缺口。核对失败/超时/格式不对,一律沿用本地索引
# 并记诊断,节点照样完成;界面按「候选,待核对」显示,而不是显示一个假的 0/N。
PLAN_OUTPUT_CONTRACT = (
    '只返回严格 JSON：{"composition":["文件组成"],'
    '"scoring_points":[{"requirement":"原要求","score":"分值或未知",'
    '"location":"响应章节标题","evidence":"所需证据","gap":"缺口或无"}],'
    '"risks":[{"category":"类别","requirement":"要求","risk":"风险",'
    '"action":"提交前动作"}],"chapter_guidance":[{"title":"章节标题原文",'
    '"basis":["招标依据"],"scoring_points":["评分点"],'
    '"material_slots":["所需素材"]}]}。评分点最多 120 项、风险最多 160 项、其余最多 40 项，'
    '不要返回章节正文。')
PLAN_MODEL_TRIES = 2
PLAN_SECTION_RX = re.compile(r'(评分|评标|评审|打分|资格|否决|废标|无效|实质性|投标文件的组成|投标文件组成|格式|装订|密封)')
_CN_DIGITS = {'零': 0, '〇': 0, '一': 1, '二': 2, '两': 2, '三': 3, '四': 4, '五': 5,
              '六': 6, '七': 7, '八': 8, '九': 9}


def _plan_model_enabled():
    return str(os.environ.get('BIDDOG_PLAN_MODEL', '1')).strip().lower() not in ('0', 'false', 'no', 'off')


def _cn_int(text):
    """'3' / '３' / '三' / '十二' → int;认不出返回 0。"""
    s = str(text or '').strip().translate(str.maketrans('０１２３４５６７８９', '0123456789'))
    if s.isdigit(): return int(s)
    if not s or any(ch not in _CN_DIGITS and ch != '十' for ch in s): return 0
    if '十' not in s: return _CN_DIGITS.get(s, 0) if len(s) == 1 else 0
    head, _, tail = s.partition('十')
    tens = _CN_DIGITS.get(head, 1) if head else 1
    ones = _CN_DIGITS.get(tail, 0) if tail else 0
    return tens * 10 + ones


def _plan_norm(text):
    return re.sub(r'[\s\u3000·•・,，。;；:：、()（）\[\]【】《》「」\-—_/\\|｜]+', '', str(text or '')).casefold()


def _plan_match_title(value, titles):
    """把「响应位置」落到真实章节标题上;落不到返回 ''。

    先精确,再「第 N 章」序号,最后才是包含关系——覆盖仪表、章节提示词都靠这一步
    才知道某个评分点归哪一章,匹配不上就只能算「还没落到章节」。"""
    raw = str(value or '').strip()
    if not raw or '需补充' in raw: return ''
    titles = [str(t) for t in (titles or []) if str(t).strip()]
    key = _plan_norm(raw)
    if not key or not titles: return ''
    normed = [(_plan_norm(t), t) for t in titles]
    for nk, title in normed:
        if nk and nk == key: return title
    m = (re.search(r'第\s*([0-9０-９零〇一二两三四五六七八九十]+)\s*(?:章|部分|节|篇)', raw)
         or re.fullmatch(r'\s*([0-9０-９]{1,2})\s*[.、．]?\s*', raw))
    if m:
        idx = _cn_int(m.group(1))
        if 1 <= idx <= len(titles): return titles[idx - 1]
    for nk, title in normed:
        if nk and len(nk) >= 2 and len(key) >= 2 and (nk in key or key in nk): return title
    return ''


def _plan_score(value):
    s = str(value if value is not None else '').strip()
    m = re.search(r'(\d+(?:\.\d+)?)', s.translate(str.maketrans('０１２３４５６７８９', '0123456789')))
    if not m: return '未知'
    n = float(m.group(1))
    if n <= 0 or n > 1000: return '未知'
    return ('%d 分' % n) if n.is_integer() else ('%s 分' % ('%.1f' % n).rstrip('0').rstrip('.'))


def _plan_score_value(text):
    m = re.search(r'\d+(?:\.\d+)?', str(text or ''))
    return float(m.group(0)) if m else 0.0


def _plan_score_sum(values):
    return sum(_plan_score_value(v) for v in values)


def _fmt_num(n):
    n = float(n or 0)
    return ('%d' % n) if n.is_integer() else ('%.1f' % n)


def _plan_expected_total(parsed):
    """从解析版里找评分办法的总分(「总分 100 分」「满分为100分」),找不到返回 None。"""
    for m in re.finditer(r'(?:总分|满分|合计|共计|总计)[^\d\n]{0,12}(\d{2,4})\s*分', str(parsed or '')):
        value = int(m.group(1))
        if 10 <= value <= 1000: return value
    return None


def _pipeline_plan_context(parsed, budget):
    """解析版超预算时,优先把评分/评标/资格/否决这些决定规划的章节整段喂给模型,再用文档开头补齐。"""
    text = str(parsed or '')
    if len(text) <= budget: return text
    blocks, current = [], []
    for line in text.splitlines():
        if line.startswith('#') and current:
            blocks.append('\n'.join(current)); current = []
        current.append(line)
    if current: blocks.append('\n'.join(current))
    picked, used, rest = [], 0, []
    for block in blocks:
        head = block.splitlines()[0] if block else ''
        if PLAN_SECTION_RX.search(head) and used + len(block) <= int(budget * 0.7):
            picked.append(block); used += len(block)
        else:
            rest.append(block)
    for block in rest:
        if used >= budget: break
        take = block[:budget - used]
        picked.append(take); used += len(take)
    return '\n\n'.join(picked)


def _plan_cell(value, limit):
    text = re.sub(r'\s+', ' ', str(value if value is not None else '')).strip()
    text = ''.join(ch for ch in text if not 0xD800 <= ord(ch) <= 0xDFFF)
    return text[:limit]


def _pipeline_normalize_plan(envelope, titles, candidate=None):
    """把模型回的规划 JSON 收成可信的结构:落位只认真实章节,分值统一成「N 分」,
    「无缺口」统一成「无」(空串会被写成〔需补充〕,那一项就永远算不上覆盖)。
    什么都没提取到返回 None,由调用方沿用本地索引。"""
    if not isinstance(envelope, dict): return None
    titles = [str(t) for t in (titles or []) if str(t).strip()]
    scores = []
    for item in (envelope.get('scoring_points') or [])[:200]:
        if not isinstance(item, dict): continue
        requirement = _plan_cell(item.get('requirement') or item.get('item') or item.get('name'), 240)
        if not requirement: continue
        gap = _plan_cell(item.get('gap'), 140)
        scores.append({'requirement': requirement, 'score': _plan_score(item.get('score')),
                       'location': _plan_match_title(item.get('location') or item.get('chapter'), titles),
                       'evidence': _plan_cell(item.get('evidence'), 140) or '招标文件_解析版.md',
                       'gap': '无' if (not gap or gap in COVERAGE_CLEAR_GAPS) else gap})
    risks = []
    for item in (envelope.get('risks') or [])[:200]:
        if not isinstance(item, dict): continue
        requirement = _plan_cell(item.get('requirement') or item.get('clause'), 240)
        if not requirement: continue
        risks.append({'category': _plan_cell(item.get('category'), 20) or '实质性要求',
                      'requirement': requirement,
                      'risk': _plan_cell(item.get('risk'), 140) or '未逐项响应可能导致无效投标',
                      'action': _plan_cell(item.get('action'), 140) or '提交前由投标负责人对照原文核对'})
    if not scores and not risks: return None
    composition = [_plan_cell(v, 120) for v in (envelope.get('composition') or []) if _plan_cell(v, 120)][:40]
    candidate = candidate if isinstance(candidate, dict) else {}
    cand_guidance = {str(g.get('title') or ''): g for g in (candidate.get('chapter_guidance') or [])
                     if isinstance(g, dict)}
    guidance, seen = {}, []
    for item in (envelope.get('chapter_guidance') or []):
        if not isinstance(item, dict): continue
        title = _plan_match_title(item.get('title'), titles)
        if not title or title in guidance: continue
        def _list(key, limit):
            raw = item.get(key)
            return ([_plan_cell(v, 240) for v in raw if _plan_cell(v, 240)][:limit]
                    if isinstance(raw, list) else [])
        guidance[title] = {'title': title, 'basis': _list('basis', 20),
                           'scoring_points': _list('scoring_points', 40),
                           'material_slots': _list('material_slots', 20)}
        seen.append(title)
    for title in titles:
        entry = guidance.get(title)
        if not entry:
            fallback = cand_guidance.get(title) or {}
            entry = guidance[title] = {'title': title,
                                       'basis': list(fallback.get('basis') or []),
                                       'scoring_points': list(fallback.get('scoring_points') or []),
                                       'material_slots': list(fallback.get('material_slots') or [])}
            seen.append(title)
        # 落到本章的评分点直接进本章提示词:这才是「本章要拿的分」
        for score in scores:
            if score['location'] != title: continue
            line = score['requirement'] + (('(%s)' % score['score']) if score['score'] != '未知' else '')
            if line not in entry['scoring_points']: entry['scoring_points'].append(line)
        entry['scoring_points'] = entry['scoring_points'][:40]
        if not entry['basis']: entry['basis'] = ['招标文件_解析版.md']
    if not composition: composition = list(candidate.get('composition') or [])
    return {'composition': composition, 'scoring_points': scores[:120], 'risks': risks[:160],
            'chapter_guidance': [guidance[t] for t in seen]}


def _start_node_heartbeat(job, node, label='模型生成中'):
    """生成期间的心跳:一次模型调用可长达数分钟且零事件,以前「静默 180 秒」就被
    误判成停滞(「模型响应偏慢」),用户读作中断。每 30 秒摸一次节点活动时间
    并发一行台词,界面上能看到它真的在干活。"""
    hb_stop = threading.Event()
    def _heartbeat():
        waited = 0
        while not hb_stop.wait(30):
            waited += 30
            try: generation_pipeline.touch_node(job, node.get('id'))
            except Exception: pass
            emit(job, {'type': 'worklog',
                       'lines': ['「%s」%s,已进行 %d 秒(连接正常)' % (node.get('title') or node.get('id'), label, waited)]})
    threading.Thread(target=_heartbeat, daemon=True).start()
    return hb_stop


def _pipeline_model_plan(job, node, model, up, attempt_root, candidate):
    """请模型逐项核对响应规划。成功返回 {'compact','notes'};任何失败返回 None(沿用本地索引)。

    这一步永远不抛(取消除外):它的存在前提就是「不能再拖垮整单」。"""
    base = os.path.basename(job)
    state = generation_pipeline.load(job)
    titles = [str(item.get('title') or '') for item in (state.get('nodes') or [])
              if str(item.get('id') or '').startswith('chapter_write:')
              and not str(item.get('id') or '').endswith((':technical_deviation', ':business_deviation'))
              and str(item.get('title') or '').strip()]
    parsed = _pipeline_model_file_text(os.path.join(attempt_root, '招标文件_解析版.md'), 3_000_000)
    budget = max(20000, int(os.environ.get('BIDDOG_PLAN_CONTEXT_CHARS', 120000)))
    context = _pipeline_plan_context(parsed, budget)
    request = _pipeline_model_file_text(os.path.join(attempt_root, '你的要求.md'), 6000)
    candidate = candidate if isinstance(candidate, dict) else {}
    cand_scores = [str(s.get('requirement') or '') for s in (candidate.get('scoring_points') or [])
                   if isinstance(s, dict) and s.get('requirement')][:120]
    cand_risks = [str(r.get('requirement') or '') for r in (candidate.get('risks') or [])
                  if isinstance(r, dict) and r.get('requirement')][:160]
    expected_total = _plan_expected_total(parsed)
    system = (
        '你是中标狗的响应规划核对员。任务:从招标文件里逐项提取评分点、废标/否决风险、投标文件组成,'
        '并把每个评分点落到给定的章节标题上。' + PLAN_OUTPUT_CONTRACT +
        '规则:1) location 只能从「章节目录」里原样照抄一个标题,确实落不到写空字符串;'
        '2) score 写「N 分」,原文没有分值写「未知」;各项分值合计应等于评分办法总分%s;'
        '3) gap 写「无」表示招标要求能被章节正文直接响应,有缺口时写清缺什么证据或素材;'
        '4) 评分点逐项来自评分办法/评标办法/详细评审表,一项不漏,也不编造;'
        '5) risks 按初步评审/资格审查/形式评审/响应性评审/否决条款逐行提取,每行一条;'
        '6) 本地索引候选只供参考,可增删改。不要输出章节正文,不要解释。'
        % ((',本文件识别到总分 %d 分' % expected_total) if expected_total else ''))
    user = ('# 章节目录(location 只能用这些标题)\n'
            + '\n'.join('%d. %s' % (index, title) for index, title in enumerate(titles, 1))
            + '\n\n# 本地索引候选(仅供核对,可增删改)\n## 评分相关条款\n'
            + '\n'.join('- ' + line for line in cand_scores)
            + '\n## 风险相关条款\n' + '\n'.join('- ' + line for line in cand_risks)
            + (('\n\n# 用户补充要求\n' + request) if request.strip() else '')
            + '\n\n# 招标文件解析版(节选)\n' + context)
    payload = {'model': model, 'stream': False,
               'messages': [{'role': 'system', 'content': system}, {'role': 'user', 'content': user}],
               'max_tokens': max(1200, min(8000, int(os.environ.get('BIDDOG_PLAN_MAX_TOKENS', 3200)))),
               'temperature': 0}
    emit(job, {'type': 'worklog',
               'lines': ['响应规划:本地索引 %d 条评分候选、%d 条风险候选,正在请模型逐项核对与落位'
                         % (len(cand_scores), len(cand_risks))]})
    reason = ''
    hb_stop = _start_node_heartbeat(job, node, '模型核对评分点与落位中')
    try:
        for attempt in range(1, PLAN_MODEL_TRIES + 1):
            if _cancel_requested(base):
                raise generation_pipeline.NodeExecutionError('cancelled', retryable=False)
            compact = None
            try:
                try:
                    response = _openai_stream_req(
                        up['base_url'], up['api_key'], payload,
                        idle_timeout=float(os.environ.get('BIDDOG_STREAM_IDLE_SECONDS', 90)),
                        total_cap=float(os.environ.get('BIDDOG_MODEL_NODE_TOTAL_SECONDS', 1800)),
                        verify=up.get('verify_ssl', True),
                        cancel_check=lambda: _cancel_requested(base))
                except _StreamCancelled:
                    raise generation_pipeline.NodeExecutionError('cancelled', retryable=False)
                except urllib.error.HTTPError as exc:
                    if exc.code not in (400, 404, 405, 422): raise
                    response = _openai_req(up['base_url'], up['api_key'], '/chat/completions', payload,
                                           timeout=float(os.environ.get('BIDDOG_MODEL_NODE_TIMEOUT_SECONDS', 240)),
                                           verify=up.get('verify_ssl', True))
                choice = ((response.get('choices') or [{}])[0] if isinstance(response, dict) else {}) or {}
                content = ((choice.get('message') or {}).get('content') if isinstance(choice, dict) else '') or ''
                compact = _pipeline_normalize_plan(_pipeline_json_object(content), titles, candidate)
                if not compact:
                    reason = '模型返回 %d 字,未形成可用的规划 JSON' % len(str(content))
            except generation_pipeline.NodeExecutionError:
                raise
            except urllib.error.HTTPError as exc:
                reason = 'HTTP %s %s' % (exc.code, _http_error_detail(exc, (up.get('api_key'),), 160))
            except Exception as exc:
                reason = net_hint(exc, (up.get('api_key'),))[:160]
            if compact:
                notes = []
                total = _plan_score_sum(item['score'] for item in compact['scoring_points'])
                if expected_total and total and abs(total - expected_total) > max(5, expected_total * 0.1):
                    notes.append('分值合计 %s 分,与评分办法总分 %d 分不一致,提交前请人工核对评分表'
                                 % (_fmt_num(total), expected_total))
                    append_diagnostic(job, 'plan_score_total_mismatch', notes[-1], level='warning',
                                      node_id=node.get('id'), model=model)
                located = sum(1 for item in compact['scoring_points'] if item['location'])
                summary = '评分点 %d 项(已落位 %d)、废标风险 %d 条' % (
                    len(compact['scoring_points']), located, len(compact['risks']))
                append_diagnostic(job, 'plan_model_checked', summary, level='info',
                                  node_id=node.get('id'), model=model)
                emit(job, {'type': 'worklog', 'lines': ['响应规划:模型核对完成,' + summary
                                                          + (';' + notes[0] if notes else '')]})
                return {'compact': compact, 'notes': notes}
            if attempt < PLAN_MODEL_TRIES:
                for _ in range(5):
                    if _cancel_requested(base): break
                    time.sleep(1)
    finally:
        hb_stop.set()
    append_diagnostic(job, 'plan_model_fallback',
                      '模型核对未成功(%s),沿用本地索引;评分点覆盖按候选显示' % (reason or '未知原因'),
                      level='warning', node_id=node.get('id'), model=model)
    emit(job, {'type': 'worklog',
               'lines': ['响应规划:模型核对未成功(%s),沿用本地索引,评分点覆盖将按候选显示'
                         % (reason or '未知原因')[:80]]})
    return None


def _pipeline_complete_truncated_markdown(content, min_chars=0):
    """Keep a substantial Markdown response at its last complete boundary."""
    text = str(content or '').strip()
    required = max(600, int(min_chars or 0) * 2)
    if len(text) < required:
        return ''
    sentence_end = max(text.rfind(mark) for mark in ('。', '！', '？', '；'))
    table_end = -1
    offset = 0
    for line in text.splitlines(True):
        offset += len(line)
        if line.strip().endswith('|') and line.strip().count('|') >= 2:
            table_end = offset - (0 if line.endswith('\n') else len(line) - len(line.rstrip()))
    boundary = max(sentence_end + 1 if sentence_end >= 0 else -1, table_end)
    if boundary < required:
        return ''
    return text[:boundary].rstrip()


def _pipeline_merge_markdown_continuation(body, continuation):
    """Append a continuation while removing a small exact overlap at the join."""
    head, tail = str(body or '').rstrip(), str(continuation or '').strip()
    if tail.startswith('```'):
        tail = re.sub(r'^```(?:markdown|md)?\s*', '', tail, flags=re.I)
        tail = re.sub(r'\s*```$', '', tail).strip()
    if not head or not tail:
        return head or tail
    for size in range(min(len(head), len(tail), 800), 31, -1):
        if head[-size:] == tail[:size]:
            tail = tail[size:].lstrip()
            break
    return head + ('\n\n' + tail if tail else '')


_UNVERIFIED_BIDDER_FACT = re.compile(
    r'(我方(?:具备|具有|拥有|持有|已(?:建立|取得|完成|通过|获得|成功)|'
    r'在.{0,30}(?:拥有|具备|有成功|完成|实施)|近三年.{0,30}(?:完成|实施)|'
    r'为依法|注册资本(?:为|不低于)|产品在.{0,30}(?:案例|应用))|'
    r'(?:上述|以上).{0,30}均为我方直接实施|'
    r'本公司(?:具备|拥有|通过|获得)|公司(?:具备|拥有|通过|获得)|上述人员全部为)'
)


def _pipeline_mark_unverified_claims(content):
    """Make unsupported bidder facts explicit without changing Markdown shape."""
    marker = '〔需补充：请核实本项投标人事实并提供证据〕'
    safe_lines = []
    for line in str(content or '').splitlines():
        matches = list(_UNVERIFIED_BIDDER_FACT.finditer(line))
        # 从右向左处理，避免插入标记后破坏后续匹配位置。只检查当前事实所在
        # 子句；前一事实已有标记，不能替后一事实背书。
        for match in reversed(matches):
            clause_start = max(
                line.rfind(mark, 0, match.start()) for mark in ('。', '；', ';', '，', ',', '|', '\n')
            ) + 1
            if '〔需补充' not in line[clause_start:match.start()]:
                line = line[:match.start()] + marker + line[match.start():]
        safe_lines.append(line)
    return '\n'.join(safe_lines)


def _pipeline_specialize_repeated_table_headers(content):
    """Replace repeated generic two-column headers with section-specific labels."""
    lines = str(content or '').splitlines()
    generic = re.compile(r'^\|\s*信息项\s*\|\s*内容\s*\|\s*$')
    if sum(1 for line in lines if generic.match(line.strip())) <= 2:
        return '\n'.join(lines)
    heading = '本节'
    used = {}
    for index, line in enumerate(lines):
        match = re.match(r'^#{2,6}\s+(.+?)\s*$', line.strip())
        if match:
            heading = re.sub(r'[`*_〔〕]', '', match.group(1)).strip()
            heading = re.sub(r'^\d+(?:\.\d+)*\s*', '', heading)[:18] or '本节'
        if not generic.match(line.strip()):
            continue
        used[heading] = used.get(heading, 0) + 1
        label = heading if used[heading] == 1 else '%s%d' % (heading, used[heading])
        lines[index] = '| %s信息项 | 待补充/响应内容 |' % label
    return '\n'.join(lines)


def _pipeline_validate_execution_contract(job, upstream):
    """Prevent a resumed job from silently switching gateways mid-run.

    只有**网关(base_url)**变化才作废检查点。Key 换新不再拦:任务失败后用户最常见的
    自救就是去「模型接入」重测/换 Key,旧版把这也判成契约变化,检查点直接作废、
    已写好的章节全部重打——「无法从断点续跑,必须重新开始」的隐蔽根源之一。
    同一网关下换 Key 是用户对自己账号的处置,接受并把指纹换到新 Key 上继续。"""
    state = generation_pipeline.load(job)
    routes = state.get('model_routes') if isinstance(state, dict) else {}
    if not isinstance(routes, dict) or not routes:
        return
    frozen_key = str(routes.get('credential_fingerprint') or '')
    current_key = generation_pipeline.credential_fingerprint(
        upstream.get('api_key') or '')
    frozen_base = str(routes.get('base_url') or '')
    current_base = generation_pipeline.safe_base_url(upstream.get('base_url') or '')
    if frozen_base and frozen_base != current_base:
        raise generation_pipeline.NodeExecutionError(
            'execution_contract_changed', retryable=False,
            detail='任务运行期间模型网关已变化；请切回原网关,或重跑本任务')
    if frozen_key and not hmac.compare_digest(frozen_key, current_key):
        try:
            def _accept(state2):
                r2 = state2.get('model_routes')
                if isinstance(r2, dict): r2['credential_fingerprint'] = current_key
                return state2
            generation_pipeline._mutate(job, _accept)
            append_diagnostic(job, 'execution_credential_updated',
                              '检测到同网关下更换了 Key,检查点继续有效', level='info')
            emit(job, {'type': 'message', 'role': 'agent',
                       'text': '检测到你更换了模型 Key(网关未变),已用新 Key 从检查点继续;已完成内容不受影响。'})
        except Exception:
            pass


class _StreamCancelled(Exception):
    """用户在流式读取期间请求了停止。"""

def _openai_stream_req(base, key, payload, idle_timeout=90, total_cap=1800,
                       verify=True, cancel_check=None):
    """流式 /chat/completions:逐 chunk 收内容,以「空闲超时」替代总超时。

    非流式长请求要等整段生成完才有首字节;国内中转网关/CDN/公司代理普遍在
    60-100 秒空闲后掐连接,章节生成动辄 2-4 分钟,必然被掐——这是「生成频繁中断」
    的最大元凶。流式下模型只要在吐字连接就有流量;连续 idle_timeout 秒没有新
    chunk 才判连接死亡。已收到的半截内容按 finish='length' 返回,复用现成的
    截断补全 + 续写逻辑,断流不再等于从头再来。返回与非流式同构的 dict。"""
    p2 = dict(payload); p2['stream'] = True
    hdr = {'Authorization': 'Bearer ' + (key or ''), 'User-Agent': 'bid-assistant/0.9',
           'Content-Type': 'application/json', 'Accept': 'text/event-stream'}
    req = urllib.request.Request((base or '').rstrip('/') + '/chat/completions',
                                 data=json.dumps(p2, ensure_ascii=False).encode('utf-8'), headers=hdr)
    kinds = ['default', 'certifi'] if verify else ['none']
    last = None
    for kind in kinds:
        try:
            ctx = _ssl_ctx(kind)
        except Exception:
            continue
        parts, finish, started = [], '', time.time()
        try:
            with _retry(lambda: urllib.request.urlopen(req, timeout=idle_timeout, context=ctx)) as r:
                for raw in r:
                    if cancel_check and cancel_check(): raise _StreamCancelled()
                    if time.time() - started > total_cap:
                        finish = finish or 'length'; break
                    line = raw.decode('utf-8', 'ignore').strip()
                    if not line.startswith('data:'): continue
                    data = line[5:].strip()
                    if data == '[DONE]': break
                    try: ev = json.loads(data)
                    except Exception: continue
                    ch = (ev.get('choices') or [{}])[0] if isinstance(ev, dict) else {}
                    delta = ch.get('delta') if isinstance(ch, dict) else {}
                    if isinstance(delta, dict) and delta.get('content'):
                        parts.append(str(delta['content']))
                    if isinstance(ch, dict) and ch.get('finish_reason'):
                        finish = str(ch['finish_reason'])
            return {'choices': [{'message': {'content': ''.join(parts)},
                                 'finish_reason': finish or ('stop' if parts else '')}],
                    'model': p2.get('model')}
        except _StreamCancelled:
            raise
        except urllib.error.HTTPError:
            raise                       # 4xx/5xx 交给上层归因(429 限流等)
        except urllib.error.URLError as e:
            if isinstance(getattr(e, 'reason', None), ssl.SSLCertVerificationError):
                last = e; continue      # 换 CA 再试一轮
            if len(''.join(parts)) >= 200:
                # 断流但已有实质内容:按截断处理,让续写逻辑接着补,别浪费已生成的
                return {'choices': [{'message': {'content': ''.join(parts)},
                                     'finish_reason': 'length'}], 'model': p2.get('model')}
            raise
        except Exception:
            if len(''.join(parts)) >= 200:
                return {'choices': [{'message': {'content': ''.join(parts)},
                                     'finish_reason': 'length'}], 'model': p2.get('model')}
            raise
    raise last

def _pipeline_model_runner(job, node, prompt, model):
    """Run a bounded node as one non-stream completion, then atomically promote files.

    The previous OpenCode tool loop repeatedly stalled after reading inputs.  A
    node already has a deterministic read/write contract, so giving the model
    filesystem tools only adds long-stream failure points.  The engine now reads
    the exact inputs, injects the Skill contract, and owns all writes itself.
    """
    base = os.path.basename(job)
    if _cancel_requested(base):
        raise generation_pipeline.NodeExecutionError('cancelled', retryable=False)
    conf = read_json(conf_path(), {})
    up = s2_conf(conf)
    if not up.get('api_key'):
        raise generation_pipeline.NodeExecutionError('model_auth_missing', retryable=False)
    _pipeline_validate_execution_contract(job, up)
    attempt_root = generation_pipeline.attempt_directory(job, node)
    ordered_inputs = _pipeline_copy_inputs(job, node, str(attempt_root))
    skill_path = os.path.join(str(attempt_root), 'SKILL.md')
    if not os.path.isfile(skill_path):
        raise generation_pipeline.NodeExecutionError('skill_contract_missing', retryable=False)
    skill_text = _pipeline_model_file_text(skill_path, 50000)
    if not skill_text.strip():
        raise generation_pipeline.NodeExecutionError('skill_contract_empty', retryable=False)
    skill_contract = _pipeline_skill_contract(skill_text)

    budget = max(30000, int(os.environ.get('BIDDOG_PIPELINE_CONTEXT_CHARS', 140000)))
    # 喂给模型的顺序 = _pipeline_copy_inputs 排好的顺序(声明输入 → 规范素材 → 相关章节模板),
    # 不再按文件名字母序:以前素材一多,招标解析版都可能被排在后面截掉。
    candidates = list(ordered_inputs or [])
    for parent, dirs, files in os.walk(str(attempt_root), followlinks=False):
        dirs[:] = [name for name in dirs if not os.path.islink(os.path.join(parent, name))]
        for name in sorted(files):
            rel = os.path.relpath(os.path.join(parent, name), str(attempt_root)).replace(os.sep, '/')
            if rel not in candidates: candidates.append(rel)
    sections = []
    for rel in candidates:
        path = os.path.join(str(attempt_root), rel)
        if not os.path.isfile(path) or os.path.islink(path): continue
        if os.path.realpath(path) == os.path.realpath(skill_path): continue
        if budget <= 0: break
        text = _pipeline_model_file_text(path, min(budget, 60000))
        if not text.strip(): continue
        sections.append('## 输入文件：%s\n%s' % (rel, text))
        budget -= len(text)

    outputs = [os.path.basename(str(name)) for name in (node.get('outputs') or [])]
    is_plan = node.get('id') == 'response_plan'
    if is_plan:
        # 先写本地索引(保底,永远先落盘);模型只做核对,核对不成也不拖垮整单。
        candidate = _pipeline_local_response_plan(job, str(attempt_root))
        if _plan_model_enabled():
            checked = _pipeline_model_plan(job, node, model, up, str(attempt_root), candidate)
            if checked:
                _pipeline_write_response_plan(job, str(attempt_root), checked['compact'],
                                              source='model', model=model, notes=checked['notes'])
        try:
            with RUNNING_LOCK:
                owner = RUNNING.get(base)
                if owner and CANCEL.get(base) == owner:
                    raise generation_pipeline.NodeExecutionError('cancelled', retryable=False)
                generation_pipeline.promote_outputs(job, attempt_root, node)
        except generation_pipeline.OutputValidationError as exc:
            raise generation_pipeline.NodeExecutionError(
                'node_output_invalid', retryable=True, detail=str(exc)) from exc
        _record_direct_skill_use(job, os.path.join(skill_dir_conf(conf), 'SKILL.md'))
        return
    single_markdown = bool(not is_plan and len(outputs) == 1
                           and os.path.splitext(outputs[0])[1].lower() in ('.md', '.txt'))
    output_contract = (
        PLAN_OUTPUT_CONTRACT
        if is_plan else (
        '只返回文件“%s”的完整 Markdown 正文，不要 JSON，不要代码围栏，不要解释。' % outputs[0]
        if single_markdown else
        '只返回一个严格 JSON 对象，格式为 {"files":{"文件名":"完整文件正文"}}。'
        'files 必须且只能包含：%s。Markdown 文件值必须是完整 Markdown 字符串；'
        'JSON 文件值可以是对象。' % '、'.join(outputs)))
    node_id_text = str(node.get('id') or '')
    is_prose_chapter = (node_id_text.startswith('chapter_write:')
                        and not node_id_text.endswith((':technical_deviation', ':business_deviation')))
    if is_prose_chapter:
        # 论述章:用结构化契约(prompts.py)替代被关键词正则打碎的 SKILL 摘要,
        # 并把其他章已用过的小标题传进去;SKILL 只留事实边界那一小段。
        system = (prompts.chapter_contract(node.get('title'), _used_chapter_headings(job, node))
                  + '\n\n<SKILL 事实边界摘要>\n%s\n</SKILL 事实边界摘要>\n%s' % (skill_contract[:2500], output_contract))
    else:
        system = (
            '你是中标狗的有界投标文件节点。以下 SKILL 是本轮必须遵守的写作契约。\n'
            '<SKILL>\n%s\n</SKILL>\n不要调用工具，不要解释过程。%s'
            '未知客户事实必须标记〔需补充〕，不得编造。' % (skill_contract, output_contract))
    user = _pipeline_direct_task(job, node) + '\n\n# 已由引擎确定性读取的输入\n' + '\n\n'.join(sections)
    chapter_tokens = max(1200, min(8000, int(
        os.environ.get('BIDDOG_CHAPTER_MAX_TOKENS', 4800))))
    payload = {
        'model': model, 'stream': False,
        'messages': [{'role': 'system', 'content': system}, {'role': 'user', 'content': user}],
        'max_tokens': 3200 if is_plan else (
            2500 if node.get('id') == 'quality_review' else chapter_tokens),
    }
    _apply_generation_params(payload, _generation_params(conf))
    hb_stop = _start_node_heartbeat(job, node)
    try:
        try:
            # 流式优先:非流式要等整段生成完才有首字节,60-100 秒空闲就被网关掐线。
            response = _openai_stream_req(up['base_url'], up['api_key'], payload,
                                          idle_timeout=float(os.environ.get('BIDDOG_STREAM_IDLE_SECONDS', 90)),
                                          total_cap=float(os.environ.get('BIDDOG_MODEL_NODE_TOTAL_SECONDS', 1800)),
                                          verify=up.get('verify_ssl', True),
                                          cancel_check=lambda: _cancel_requested(base))
        except _StreamCancelled:
            raise generation_pipeline.NodeExecutionError('cancelled', retryable=False)
        except urllib.error.HTTPError as exc:
            # 400/404/422 可能是网关不支持 stream:降级回非流式再试一次
            if exc.code in (400, 404, 405, 422):
                response = _openai_req(up['base_url'], up['api_key'], '/chat/completions', payload,
                                       timeout=float(os.environ.get('BIDDOG_MODEL_NODE_TIMEOUT_SECONDS', 240)),
                                       verify=up.get('verify_ssl', True))
            else:
                raise
    except generation_pipeline.NodeExecutionError:
        hb_stop.set()
        raise
    except urllib.error.HTTPError as exc:
        hb_stop.set()
        detail = _http_error_detail(exc, (up.get('api_key'),), 240)
        retryable = exc.code in (408, 409, 425, 429) or 500 <= exc.code < 600
        mark_connection_check(job, False, '模型请求失败(HTTP %s)' % exc.code)
        raise generation_pipeline.NodeExecutionError(
            'model_rate_limited' if exc.code == 429 else 'model_request_failed',
            retryable=retryable, detail='HTTP %s %s' % (exc.code, detail)) from exc
    except Exception as exc:
        hb_stop.set()
        mark_connection_check(job, False, '模型连接中断:%s' % net_hint(exc, (up.get('api_key'),))[:100])
        raise generation_pipeline.NodeExecutionError(
            'model_request_interrupted', retryable=True,
            detail=net_hint(exc, (up.get('api_key'),))) from exc
    hb_stop.set()
    if _cancel_requested(base):
        raise generation_pipeline.NodeExecutionError('cancelled', retryable=False)
    mark_connection_check(job, True, '连接正常,模型按节点生成中')
    choices = response.get('choices') if isinstance(response, dict) else None
    if not choices:
        detail = ((response.get('error') or {}).get('message')
                  if isinstance(response, dict) and isinstance(response.get('error'), dict) else '')
        append_diagnostic(job, 'model_response_without_choices', str(detail or 'empty choices')[:240],
                          level='warning', node_id=node.get('id'), model=model)
        raise generation_pipeline.NodeExecutionError('model_response_empty', retryable=True,
                                                       detail=str(detail or '')[:240])
    choice = (choices or [{}])[0]
    content = ((choice.get('message') or {}).get('content') if isinstance(choice, dict) else '')
    finish = str(choice.get('finish_reason') or '') if isinstance(choice, dict) else ''
    if finish and finish != 'stop':
        truncation_minimum = node.get('min_chars') or 0
        if (str(node.get('id') or '').startswith('chapter_write:')
                and node.get('id') not in (
                    'chapter_write:technical_deviation',
                    'chapter_write:business_deviation')):
            # 触顶后的首段必须足够完整，续写失败时也不能把几百字短稿晋升为
            # 已完成节点。正常 stop 响应仍由后续确定性质检按章节检查。
            truncation_minimum = max(1250, int(truncation_minimum or 0))
        accepted = (_pipeline_complete_truncated_markdown(content, truncation_minimum)
                    if single_markdown and finish == 'length' else '')
        if accepted:
            content = accepted
            continued = False
            if str(node.get('id') or '').startswith('chapter_write:'):
                continuation_tokens = max(800, min(4000, int(
                    os.environ.get('BIDDOG_CHAPTER_CONTINUATION_TOKENS', 2400))))
                continuation_payload = dict(payload)
                continuation_payload['max_tokens'] = continuation_tokens
                continuation_payload['messages'] = list(payload['messages']) + [
                    {'role': 'assistant', 'content': accepted},
                    {'role': 'user', 'content': (
                        '上文因单次输出容量结束。请从最后一个完整句之后继续，'
                        '不要重复已经完成的内容、标题或表格。补齐尚未覆盖的小节和验收内容；'
                        '只输出续写的 Markdown 正文。')},
                ]
                try:
                    _pipeline_validate_execution_contract(job, s2_conf(read_json(conf_path(), {})))
                    try:
                        followup = _openai_stream_req(
                            up['base_url'], up['api_key'], continuation_payload,
                            idle_timeout=float(os.environ.get('BIDDOG_STREAM_IDLE_SECONDS', 90)),
                            total_cap=float(os.environ.get('BIDDOG_MODEL_NODE_TOTAL_SECONDS', 1800)),
                            verify=up.get('verify_ssl', True),
                            cancel_check=lambda: _cancel_requested(base))
                    except _StreamCancelled:
                        raise generation_pipeline.NodeExecutionError('cancelled', retryable=False)
                    except urllib.error.HTTPError as _sexc:
                        if _sexc.code not in (400, 404, 405, 422): raise
                        followup = _openai_req(
                            up['base_url'], up['api_key'], '/chat/completions',
                            continuation_payload,
                            timeout=float(os.environ.get(
                                'BIDDOG_MODEL_NODE_TIMEOUT_SECONDS', 240)),
                            verify=up.get('verify_ssl', True))
                    if _cancel_requested(base):
                        raise generation_pipeline.NodeExecutionError('cancelled', retryable=False)
                    follow_choice = ((followup.get('choices') or [{}])[0]
                                     if isinstance(followup, dict) else {})
                    follow_content = ((follow_choice.get('message') or {}).get('content')
                                      if isinstance(follow_choice, dict) else '')
                    follow_finish = str(follow_choice.get('finish_reason') or '')
                    if follow_finish == 'length':
                        follow_content = _pipeline_complete_truncated_markdown(
                            follow_content, 0)
                    if str(follow_content or '').strip():
                        content = _pipeline_merge_markdown_continuation(
                            accepted, follow_content)
                        continued = True
                        append_diagnostic(
                            job, 'model_output_continued',
                            '首段触及容量后已从断点续写，合计保留 %d 字。' % len(content),
                            level='info', node_id=node.get('id'), model=model)
                except Exception as exc:
                    append_diagnostic(
                        job, 'model_continuation_failed_but_preserved',
                        '续写未返回，但首段 %d 字已保留：%s' % (
                            len(accepted), net_hint(exc, (up.get('api_key'),))[:160]),
                        level='warning', node_id=node.get('id'), model=model)
            if not continued:
                append_diagnostic(job, 'model_output_truncated_but_accepted',
                                  '已保留 %d 字完整正文，并移除末尾不完整句。' % len(content),
                                  level='warning', node_id=node.get('id'), model=model)
        else:
            raise generation_pipeline.NodeExecutionError('model_output_truncated', retryable=True,
                                                           detail='finish_reason=%s' % finish)
    if single_markdown:
        text = str(content or '').strip()
        if text.startswith('```'):
            text = re.sub(r'^```(?:markdown|md)?\s*', '', text, flags=re.I)
            text = re.sub(r'\s*```$', '', text).strip()
        if str(node.get('id') or '').startswith('chapter_write:'):
            text = _pipeline_mark_unverified_claims(text)
            text = _pipeline_specialize_repeated_table_headers(text)
        target = os.path.join(str(attempt_root), outputs[0]); tmp = target + '.tmp'
        with open(tmp, 'w', encoding='utf-8') as handle: handle.write(text.rstrip() + '\n')
        os.replace(tmp, target)
    else:
        envelope = _pipeline_json_object(content)
        if not isinstance(envelope, dict):
            raise generation_pipeline.NodeExecutionError('model_output_invalid_json', retryable=True)
        files = envelope.get('files')
        if not isinstance(files, dict):
            raise generation_pipeline.NodeExecutionError('model_output_invalid_json', retryable=True)
        if set(files) != set(outputs):
            raise generation_pipeline.NodeExecutionError(
                'model_output_contract_mismatch', retryable=True,
                detail='期望 %s，实际 %s' % (outputs, sorted(str(key) for key in files)))
        for name in outputs:
            value = files.get(name)
            text = (json.dumps(value, ensure_ascii=False, indent=2) if isinstance(value, (dict, list))
                    else str(value or ''))
            target = os.path.join(str(attempt_root), name)
            tmp = target + '.tmp'
            with open(tmp, 'w', encoding='utf-8') as handle: handle.write(text.rstrip() + '\n')
            os.replace(tmp, target)
    try:
        # 与 _request_cancel 共用 RUNNING_LOCK，使“最后一次取消检查 + 文件晋升”
        # 成为一个原子区间；停止请求先拿到锁时，迟到结果绝不会晋升。
        with RUNNING_LOCK:
            owner = RUNNING.get(base)
            if owner and CANCEL.get(base) == owner:
                raise generation_pipeline.NodeExecutionError('cancelled', retryable=False)
            generation_pipeline.promote_outputs(job, attempt_root, node)
    except generation_pipeline.OutputValidationError as exc:
        raise generation_pipeline.NodeExecutionError(
            'node_output_invalid', retryable=True, detail=str(exc)) from exc
    _record_direct_skill_use(job, os.path.join(skill_dir_conf(conf), 'SKILL.md'))


def _pipeline_prompt(job, node):
    outputs = [os.path.basename(str(name)) for name in (node.get('outputs') or [])]
    common = (
        '你正在执行中标狗的一个有界节点，不是整单自由发挥。工作目录是隔离的本次尝试目录。\n'
        '先读取当前目录的 `SKILL.md`，并遵守其中事实边界。不得编造客户名称、报价、资质、案例或人员。\n'
        '招标正文只读 `招标文件_解析版.md`；补充要求读 `你的要求.md`（若存在）。\n'
        '只完成本节点，不要启动子 agent，不要导出 Word，不要修改 pipeline.json。\n'
        '必须使用 write/edit 工具把内容真实写入下列指定文件；禁止只在聊天回复里给正文。'
        '结束前用 read 工具逐个检查文件存在且正文完整，缺一项就继续写，全部落盘后才能结束。\n')
    node_id = str(node.get('id') or '')
    if node_id == 'response_plan':
        return common + (
            '任务：基于招标解析版提取响应规划。必须分别写出：\n'
            '- `投标文件组成.md`：完整目录与每章目的；\n'
            '- `评分点响应矩阵.md`：原要求、分值（未知则标未知）、响应位置、证据、缺口；\n'
            '- `废标风险清单.md`：资格、签章、报价、时限、实质性条款风险。\n'
            '- `response_plan.json`：UTF-8 JSON，顶层含 chapters 数组；每项必须有 id、title、output，'
            '还必须有 basis（招标依据数组）、scoring_points（评分点数组）、'
            'material_slots（素材槽位数组）和 dependencies（前置章节 id 数组）；'
            '章节 id、输出文件必须唯一，依赖不得成环。\n'
            '每个文件至少 200 字，所有不确定项标〔需补充〕。')
    if node_id == 'chapter_write:technical_deviation':
        return common + (
            '任务：生成 `技术应答偏离表.md`。必须使用 Markdown 表格，列为“序号｜招标条款/技术要求｜'
            '投标响应｜偏离情况｜依据/证据｜备注”。逐项从解析版提取；未发现明确偏离时也要写一行“无偏离”，'
            '并注明核对依据，不能只写说明段落。唯一输出：`技术应答偏离表.md`。')
    if node_id == 'chapter_write:business_deviation':
        return common + (
            '任务：生成 `商务偏离表.md`。必须使用 Markdown 表格，列为“序号｜招标条款/商务要求｜'
            '投标响应｜偏离情况｜依据/证据｜备注”。逐项覆盖报价、付款、工期、质保、交付、签章等要求；'
            '未发现明确偏离时也要写一行“无偏离”，并注明核对依据。唯一输出：`商务偏离表.md`。')
    if node_id.startswith('chapter_write:'):
        basis = '、'.join(str(value) for value in (node.get('basis') or [])) or '以解析版为准'
        slots = '、'.join(str(value) for value in (node.get('material_slots') or [])) or '无指定素材槽位'
        scores = '、'.join(str(value) for value in (node.get('scoring_points') or [])) or '未标注独立评分点'
        dependency_files = '、'.join('`%s`' % name for name in
                                       generation_pipeline.dependency_outputs(
                                           generation_pipeline.load(job), node_id)) or '无'
        return common + (
            '任务：撰写章节“%s”。同时读取 `投标文件组成.md`、`评分点响应矩阵.md`、'
            '`废标风险清单.md` 和素材目录。\n'
            '本章规划依据：%s。对应评分点：%s。素材槽位：%s。'
            '前置章节产物：%s，必须读取并保持一致。\n'
            '唯一输出：`%s`。正文应逐项响应并标注证据来源；素材缺失写〔需补充〕，论述章不少于 %d 个中文字符。' %
            (node.get('title') or '本章', basis, scores, slots, dependency_files,
             outputs[0] if outputs else '章节.md', prompts.CHAPTER_TARGET_CHARS))
    if node_id == 'quality_review':
        output = outputs[0] if outputs else '模型复核报告.md'
        return common + ('任务：%s\n唯一输出：`%s`。' % (prompts.review_task(), output))
    raise generation_pipeline.PipelineError('没有该节点的模型提示：%s' % node_id)


def _pipeline_event(job, event, node):
    title = node.get('title') or node.get('id')
    if event == 'started':
        emit(job, {'type': 'worklog', 'lines': ['开始：%s（第 %d/%d 次）' %
              (title, int(node.get('attempt') or 1), int(node.get('max_attempts') or 3))]})
    elif event == 'done':
        emit(job, {'type': 'worklog', 'lines': ['完成：%s' % title]})
        for name in node.get('outputs') or []:
            emit(job, {'type': 'artifact', 'name': name})
    elif event == 'retry':
        error_code = str(node.get('error_code') or '')
        if error_code == 'model_rate_limited':
            PIPELINE_THROTTLED.add(os.path.basename(job))
        if error_code == 'model_output_truncated':
            reason = '模型输出过长，已缩小本节范围并重试'
        elif error_code in ('model_output_invalid_json', 'model_output_contract_mismatch',
                            'node_output_invalid', 'output_validation_failed'):
            reason = '模型返回格式不符合本节交付要求，正在重试'
        elif error_code in ('model_rate_limited', 'model_request_failed', 'model_response_empty'):
            reason = '上游模型暂时未完整响应，正在重试'
        else:
            reason = '连接暂时中断，正在重试'
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '“%s”%s（已尝试 %d/%d）；前面成果已保留。' %
                           (title, reason, int(node.get('attempt') or 0),
                            int(node.get('max_attempts') or 3))})


def _write_text_atomic(path, text):
    tmp = path + '.tmp'
    with open(tmp, 'w', encoding='utf-8') as handle: handle.write(text)
    os.replace(tmp, path)


def _plan_chapter_titles(job):
    try: state = generation_pipeline.load(job)
    except Exception: return []
    return [str(n.get('title') or '') for n in (state.get('nodes') or [])
            if str(n.get('id') or '').startswith('chapter_write:') and str(n.get('title') or '').strip()]


def _build_scoring_index(job):
    """《评标索引》:评委按它逐项翻——每个评分点对应到具体章节。只有规划经模型核对过才有资格
    出这张表:本地索引的「落位」是一句常量,印出来就是一张全是〔需补充〕的表。"""
    plan = read_json(os.path.join(job, 'response_plan.json'), {})
    if not isinstance(plan, dict) or plan.get('source') != 'model': return ''
    rows = _md_table_rows(os.path.join(job, '评分点响应矩阵.md'), 6)
    if not rows: return ''
    titles = _plan_chapter_titles(job)
    lines = ['# 评标索引', '',
             '> 评委按本表逐项翻阅:每个评分点对应到具体章节;页码以最终 Word 目录为准。', '',
             '| 序号 | 评分项 | 分值 | 评估标准 / 证据 | 对应章节 |', '|---|---|---:|---|---|']
    for index, row in enumerate(rows, 1):
        requirement, score, location, evidence = row[1], row[2], row[3], row[4]
        matched = _plan_match_title(location, titles) or ('' if '需补充' in location else location.strip())
        lines.append('| %d | %s | %s | %s | %s |' % (
            index, _pipeline_cell(requirement), _pipeline_cell(score), _pipeline_cell(evidence),
            _pipeline_cell(matched)))
    total = _plan_score_sum(row[2] for row in rows)
    if total: lines.extend(['', '评分点合计 %s 分。' % _fmt_num(total)])
    return '\n'.join(lines) + '\n'


_SUPPLEMENT_RX = re.compile(r'〔(需补充|此处粘贴|参数待核实|配图建议|配图待人工)[:：]?([^〕]{0,120})〕')


def _build_supplement_list(job, chapter_outputs):
    """《投标人补料清单》:正文里的占位(〔需补充〕〔此处粘贴〕〔参数待核实〕〔配图建议〕)按章汇总,
    加上废标风险清单里的提交前动作,再加固定的人工确认节点。自动初稿 ≠ 直接可投,这张表就是差距清单。"""
    items = []
    for name in chapter_outputs:
        try: text = open(os.path.join(job, name), encoding='utf-8', errors='ignore').read()
        except OSError: continue
        head = re.search(r'^#\s+(.+?)\s*$', text, re.M)
        chapter = head.group(1).strip() if head else os.path.splitext(name)[0]
        seen = set()
        for kind, detail in _SUPPLEMENT_RX.findall(text):
            key = (kind, detail.strip())
            if key in seen: continue
            seen.add(key)
            items.append((chapter, kind, detail.strip() or '—'))
    risks = _md_table_rows(os.path.join(job, '废标风险清单.md'), 5)
    lines = ['# 投标人补料清单', '',
             '> 自动初稿 ≠ 直接可投。下表是正文里留下的占位与提交前必办事项,逐条处理后再提交。', '',
             '## 一、正文占位(按章节)', '',
             '| 序号 | 章节 | 类型 | 内容 | 状态 |', '|---|---|---|---|---|']
    if items:
        for index, (chapter, kind, detail) in enumerate(items[:200], 1):
            lines.append('| %d | %s | %s | %s | 待处理 |' % (index, _pipeline_cell(chapter), kind, _pipeline_cell(detail)))
    else:
        lines.append('| — | — | — | 正文未留占位 | — |')
    lines.extend(['', '## 二、废标风险对应动作', '',
                  '| 序号 | 类别 | 招标要求 | 提交前动作 | 状态 |', '|---|---|---|---|---|'])
    if risks:
        for index, row in enumerate(risks[:80], 1):
            lines.append('| %d | %s | %s | %s | 待核对 |' % (
                index, _pipeline_cell(row[1]), _pipeline_cell(row[2]), _pipeline_cell(row[4])))
    else:
        lines.append('| — | — | 未生成废标风险清单 | 人工对照招标原文核对 | — |')
    lines.extend(['', '## 三、人工确认节点(不得代填)', '',
                  '- 投标人名称、报价、资质证书有效期、项目案例、签章/装订/密封/份数要求。',
                  '- 所有〔需补充〕项完成后,方可作为正式投标文件提交。'])
    return '\n'.join(lines) + '\n'


def _review_report_gaps(job):
    """把标准模式的《模型复核报告.md》解析成出件前检查里能点的条目:必办 → 红(挂「重做这一章」),
    建议 → 黄。以前这份报告只是右栏里一个文件,标准模式多花的那份钱没有闭环。"""
    path = os.path.join(job, '模型复核报告.md')
    if not os.path.isfile(path): return []
    titles = _plan_chapter_titles(job)
    gaps = []
    for row in _md_table_rows(path, 4)[:40]:
        chapter, problem, level, advice = row[0].strip(), row[1].strip(), row[2].strip(), row[3].strip()
        if chapter == '章节' or problem in ('', '无', '—', '-', '无问题'): continue
        matched = _plan_match_title(chapter, titles)
        actions = []
        if matched:
            actions.append({'act': 'redo', 'label': '重做这一章',
                            'param': '重写章节「%s」:%s;%s' % (matched, problem[:120], advice[:200])})
        actions.append({'act': 'open_artifact', 'label': '打开复核报告', 'file': '模型复核报告.md'})
        gaps.append({'level': 'red' if '必办' in level else 'yellow',
                     'title': '复核:%s' % problem[:80],
                     'detail': (('%s · ' % chapter) if chapter else '') + (advice[:200] or '见复核报告'),
                     'actions': actions})
    return gaps


def _pipeline_complete_local(job, node_id, digest, action, *, retryable=True):
    node = generation_pipeline.start_node(job, node_id, input_digest=digest)
    if node.get('state') == 'done': return node
    try:
        action(node)
        return generation_pipeline.complete_node(job, node_id, input_digest=digest)
    except Exception as exc:
        generation_pipeline.fail_node(job, node_id, 'local_node_failed', retryable=retryable)
        raise


def _pipeline_retry_sleep(base):
    """节点重试退避期间的可中断等待:分钟级退避里用户点「停止」必须秒级生效,
    不能睡满全程再响应(cancel 会在下一次节点启动检查时抛出)。"""
    def _sleep(seconds):
        deadline = time.time() + max(0.0, float(seconds or 0))
        while time.time() < deadline:
            if _cancel_requested(base): return
            time.sleep(min(1.0, max(0.0, deadline - time.time())))
    return _sleep

def _md_table_rows(path, cells):
    """Read data rows from one deterministic markdown table (skips header/separator)."""
    rows = []
    try:
        lines = open(path, encoding='utf-8', errors='ignore').read().splitlines()
    except OSError:
        return rows
    for ln in lines:
        ln = ln.strip()
        if not ln.startswith('|'): continue
        parts = [c.strip() for c in ln.strip('|').split('|')]
        if len(parts) < cells: continue
        if parts[0] in ('序号',) or set(parts[0]) <= set('-—: '): continue
        rows.append(parts)
    return rows

def _parse_confirm_summary(job):
    """Deterministic key facts for the pre-write confirmation card.

    Everything here is extracted locally from already-produced artifacts; a value
    we cannot find is reported as such instead of being guessed."""
    meta = read_json(os.path.join(job, '任务.json'), {})
    unknown = '未识别，请人工核对招标文件'
    text = ''
    try:
        text = open(os.path.join(job, '招标文件_解析版.md'), encoding='utf-8', errors='ignore').read()[:200000]
    except OSError:
        pass
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    def find_line(pattern, limit=140):
        rx = re.compile(pattern)
        for ln in lines:
            if rx.search(ln):
                return re.sub(r'^[#>*\-\d\.、\s]+', '', ln)[:limit]
        return ''
    deadline = find_line(r'(递交|提交|投标).{0,8}(截止|时间)') or find_line(r'开标时间') or unknown
    qual = find_line(r'资质.{0,30}(等级|要求|承包|资格)') or find_line(r'(施工总承包|安全生产许可证)') or unknown
    scoring_rows = _md_table_rows(os.path.join(job, '评分点响应矩阵.md'), 6)
    method_line = find_line(r'(综合评估法|综合评分法|评分办法|经评审的最低投标价)')
    plan = read_json(os.path.join(job, 'response_plan.json'), {})
    if isinstance(plan, dict) and plan.get('source') == 'local':
        # 本地索引只是按关键词挑出来的候选行,数出来的不是「评分点数」,别把它说成评分点。
        # 没有 response_plan.json 的是旧版/智能体路径的任务,矩阵本来就是模型写的,照旧。
        scoring = '识别到 %d 处评分相关条款(候选,待模型核对)' % len(scoring_rows)
    else:
        total = _plan_score_sum(row[2] for row in scoring_rows)
        scoring = ('共 %d 个评分点' % len(scoring_rows)) + ((' · 合计 %s 分' % _fmt_num(total)) if total else '')
    scoring += (' · ' + method_line) if method_line else ''
    veto_rows = _md_table_rows(os.path.join(job, '废标风险清单.md'), 5)
    cats = []
    for row in veto_rows:
        cat = row[1][:12]
        if cat and cat not in cats: cats.append(cat)
    veto = ('共 %d 条废标/否决风险' % len(veto_rows)) + (('：' + '、'.join(cats[:4])) if cats else '')
    return {
        'project': str(meta.get('name') or meta.get('title')
                       or os.path.basename(str(meta.get('tender') or '')) or '未命名项目')[:80],
        'deadline': deadline, 'qualification': qual,
        'scoring': scoring if scoring_rows else unknown,
        'veto': veto if veto_rows else unknown,
    }

CONFIRM_PARSE_ACCEPT = '确认无误，开始撰写'

def _confirm_gate_blocks(job, meta):
    """True when the job asked for pre-write confirmation and it has not happened."""
    if not meta.get('confirm_parse'): return False
    return not bool((read_json(os.path.join(job, '解析确认.json'), {}) or {}).get('confirmed'))

def generation_pipeline_worker(job):
    """Run local parse -> short model nodes -> deterministic Word, resumably."""
    base = os.path.basename(job)
    meta = read_json(os.path.join(job, '任务.json'), {})
    tender = os.path.basename(meta.get('tender') or '')
    source = os.path.join(job, tender)
    try:
        state = generation_pipeline.load(job)
        if state.get('version') == 1:
            state = generation_pipeline.migrate(job)
        if state.get('version') != generation_pipeline.PIPELINE_VERSION:
            state = _initialize_generation_pipeline(job, meta, read_json(conf_path(), {}))
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '已启动可恢复分段生成：本地解析一次，模型按节点执行；断线只重试当前节点。'})
        update_runtime(job, execution_path='checkpoint_pipeline', can_pause=False,
                       pause_disabled_reason='分段节点会自动保存；需要中止请使用停止。')

        source_digest = generation_pipeline.file_digest([source])
        def parse_action(_node):
            manifest = read_json(os.path.join(job, 'source_manifest.json'), {})
            if (manifest.get('sha256') != source_digest
                    or not os.path.isfile(os.path.join(job, '招标文件_解析版.md'))):
                generation_pipeline.parse_source(job, tender)
        _pipeline_complete_local(job, 'source_parse', source_digest, parse_action, retryable=False)
        emit(job, {'type': 'progress', 'stage': '招标文件已在本地解析，正在提取响应规划',
                   'pct': 12, 'step': 2, 'total': 12})

        request_file = os.path.join(job, '你的要求.md')
        plan_digest = generation_pipeline.file_digest(
            [os.path.join(job, '招标文件_解析版.md'), request_file])
        generation_pipeline.run_model_node(
            job, 'response_plan', lambda node, prompt, model: _pipeline_model_runner(job, node, prompt, model),
            prompt=_pipeline_prompt(job, next(n for n in state['nodes'] if n['id'] == 'response_plan')),
            input_digest=plan_digest, sleep=_pipeline_retry_sleep(base),
            on_event=lambda event, node: _pipeline_event(job, event, node))
        state = generation_pipeline.apply_response_plan(job)
        if _confirm_gate_blocks(job, meta):
            qid = 'confirm_parse:' + str(meta.get('run_id') or os.path.basename(job))
            emit(job, {'type': 'progress', 'stage': '响应规划已完成，等待确认关键信息后开始分章撰写',
                       'pct': 45, 'step': 6, 'total': 12})
            emit(job, {'type': 'question', 'id': qid, 'kind': 'confirm_parse',
                       'payload': _parse_confirm_summary(job),
                       'text': '响应规划已完成。开始分章撰写前，请确认解析出的关键信息——'
                               '发现读错了现在改一个字，胜过跑完 30 分钟后返工。',
                       'options': [CONFIRM_PARSE_ACCEPT]})
            update_runtime(job, execution_path='checkpoint_pipeline', can_pause=False,
                           pause_disabled_reason='正在等待确认解析结果。')
            # 落一个明确的可续跑终态:没有它任务会呈现为 unknown,
            # 右栏会误挂「没出 Word,未完成」红牌——明明只是等确认。
            write_json(os.path.join(job, 'outcome.json'),
                       {'state': 'stopped', 'reason': '等待确认解析结果（确认后自动继续）', 'ts': now()})
            return
        emit(job, {'type': 'progress', 'stage': '响应规划已完成，正在分章撰写',
                   'pct': 45, 'step': 6, 'total': 12})

        chapter_nodes = [node for node in state.get('nodes') or []
                         if str(node.get('id') or '').startswith('chapter_write:')]
        def write_chapter(node):
            current_state = generation_pipeline.load(job)
            inputs = [os.path.join(job, name) for name in
                      _pipeline_declared_input_names(job, node, current_state)]
            inputs.extend(_pipeline_material_paths(job))
            contract_fields = {key: node.get(key) for key in
                               ('title', 'basis', 'scoring_points', 'material_slots', 'dependencies')}
            # 重写标记只在有值时进入契约:0.20.4 升级上来的旧任务摘要保持不变,
            # 已完成章节不会因为升级被重跑;单章重写后摘要必然变化,触发精确重做。
            contract_fields.update({key: node.get(key) for key in ('user_note', 'rewrite_serial')
                                    if node.get(key)})
            contract = json.dumps(contract_fields, ensure_ascii=False, sort_keys=True)
            digest = generation_pipeline.file_digest(inputs) + ':' + _sha256_text(contract)
            return generation_pipeline.run_model_node(
                job, node['id'], lambda current, prompt, model: _pipeline_model_runner(job, current, prompt, model),
                prompt=_pipeline_prompt(job, node), input_digest=digest,
                sleep=_pipeline_retry_sleep(base),
                on_event=lambda event, current: _pipeline_event(job, event, current))
        # 每次恢复都从 DAG 根节点按序校验摘要；摘要未变的 done 节点会立即复用，
        # 依赖产物变化时下游才会被精确重做。
        # 持续 N 路在飞(默认 4):一章写完立刻补位,而不是两章一批互相等;撞到限流降到 2 路。
        PIPELINE_THROTTLED.discard(base)
        completed_ids = set()
        remaining = {node['id']: node for node in chapter_nodes}
        running = {}
        with ThreadPoolExecutor(max_workers=_chapter_parallelism()) as pool:
            while remaining or running:
                in_flight = {item['id'] for item in running.values()}
                ready = [node for node in remaining.values()
                         if set(node.get('dependencies') or []) <= completed_ids
                         and node['id'] not in in_flight]
                if not ready and not running:
                    raise generation_pipeline.PipelineError('章节依赖无法继续执行')
                for node in ready:
                    if len(running) >= _chapter_slots(base): break
                    remaining.pop(node['id'], None)
                    running[pool.submit(write_chapter, node)] = node
                if not running:
                    continue
                done, _pending = wait(list(running), return_when=FIRST_COMPLETED)
                for future in done:
                    node = running.pop(future)
                    future.result()
                    completed_ids.add(node['id'])

        emit(job, {'type': 'progress', 'stage': '各章已完成，正在汇总整册',
                   'pct': 78, 'step': 9, 'total': 12})
        state = generation_pipeline.load(job)
        chapter_outputs = [node['outputs'][0] for node in state.get('nodes') or []
                           if str(node.get('id') or '').startswith('chapter_write:')]
        assemble_digest = generation_pipeline.file_digest(
            [os.path.join(job, name) for name in chapter_outputs]
            + [os.path.join(job, name) for name in ('评分点响应矩阵.md', '废标风险清单.md', 'response_plan.json')
               if os.path.isfile(os.path.join(job, name))])
        def assemble(_node):
            pieces = []
            for name in chapter_outputs:
                text = open(os.path.join(job, name), encoding='utf-8', errors='ignore').read().strip()
                if text: pieces.append(text)
            # 整册级产物(SKILL 第 4.5/5 步要求、流水线以前没有):评标索引放整册最前,
            # 补料清单放最后;都另存一份单独文件,右栏能直接点开。
            index_md = _build_scoring_index(job)
            if index_md:
                _write_text_atomic(os.path.join(job, '评标索引.md'), index_md)
                emit(job, {'type': 'artifact', 'name': '评标索引.md'})
                pieces.insert(0, index_md.strip())
            supplement_md = _build_supplement_list(job, chapter_outputs)
            if supplement_md:
                _write_text_atomic(os.path.join(job, '投标人补料清单.md'), supplement_md)
                emit(job, {'type': 'artifact', 'name': '投标人补料清单.md'})
                pieces.append(supplement_md.strip())
            target = os.path.join(job, '投标文件_整册.md')
            tmp = target + '.tmp'
            open(tmp, 'w', encoding='utf-8').write('\n\n---\n\n'.join(pieces) + '\n')
            os.replace(tmp, target)
        _pipeline_complete_local(job, 'assemble', assemble_digest, assemble)

        state = generation_pipeline.load(job)
        quality_node = next(node for node in state['nodes'] if node['id'] == 'quality_review')
        quality_digest = generation_pipeline.file_digest(
            [os.path.join(job, '投标文件_整册.md'), os.path.join(job, '评分点响应矩阵.md'),
             os.path.join(job, '废标风险清单.md')])
        if quality_node.get('kind') == 'model':
            generation_pipeline.run_model_node(
                job, 'quality_review', lambda node, prompt, model: _pipeline_model_runner(job, node, prompt, model),
                prompt=_pipeline_prompt(job, quality_node), input_digest=quality_digest,
                sleep=_pipeline_retry_sleep(base),
                on_event=lambda event, node: _pipeline_event(job, event, node))
        else:
            def local_quality(_node):
                body = open(os.path.join(job, '投标文件_整册.md'), encoding='utf-8', errors='ignore').read()
                report = ('# 成品质检报告\n\n- 分章数量：%d\n- 整册字符数：%d\n'
                          '- 待补充标记：%d\n- 说明：最终 Word 导出后仍需经过确定性交付质检。\n' %
                          (len(chapter_outputs), len(body), body.count('〔需补充〕')))
                open(os.path.join(job, '成品质检报告.md'), 'w', encoding='utf-8').write(report)
            _pipeline_complete_local(job, 'quality_review', quality_digest, local_quality)

        emit(job, {'type': 'progress', 'stage': '整册已汇总，正在导出并检查 Word',
                   'pct': 90, 'step': 11, 'total': 12})
        word_digest = generation_pipeline.file_digest([os.path.join(job, '投标文件_整册.md')])
        def export_word(_node):
            ensure_docx(job, set(list_deliverables(job)), force=True)
            audit = word_format_audit(job, '投标文件_整册.docx')
            # 只有「检查没跑成/Word 不可用」才算节点失败;样式项不符(warn)如实告知,
            # 进「出件前检查」清单——Word 明明已导出可用,不能把整单打成未完成。
            if audit.get('status') == 'fail':
                raise generation_pipeline.OutputValidationError(
                    audit.get('summary') or 'Word 格式门禁未通过')
            if audit.get('status') == 'warn':
                emit(job, {'type': 'message', 'role': 'agent',
                           'text': '%s。这些属于提交前人工确认项,在「出件前检查」里逐条看即可。' %
                                   (audit.get('summary') or 'Word 格式自检有不符项')})
        # word_export 是幂等确定性步骤;Windows 上用户开着 Word 锁住目标文件时
        # 会抛 PermissionError,必须允许重试(关掉文件点重试即可),不能一锤 blocked。
        _pipeline_complete_local(job, 'word_export', word_digest, export_word, retryable=True)
        delivery_digest = generation_pipeline.file_digest([
            os.path.join(job, '投标文件_整册.docx'), os.path.join(job, '成品质检报告.md'),
            os.path.join(job, '技术应答偏离表.md'), os.path.join(job, '商务偏离表.md'),
            os.path.join(job, 'Word格式自检报告.md')])
        def delivery_gate(_node):
            result = settle(job, commit=False)
            delivery = delivery_summary(job)
            if not isinstance(result, dict) or result.get('state') != 'ready' or not delivery.get('ready'):
                raise generation_pipeline.OutputValidationError('最终交付门禁未通过')
        _pipeline_complete_local(job, 'delivery_gate', delivery_digest, delivery_gate, retryable=False)
        # 先落盘 delivery_gate=done，再提交用户可见 outcome=done。若在两者之间崩溃，
        # 任务仍显示未完成；下次恢复会复用已完成门禁并补提交，绝不会假完成。
        final = settle(job)
        if not isinstance(final, dict) or final.get('state') != 'done':
            raise generation_pipeline.OutputValidationError('交付终态提交失败')
    except generation_pipeline.SourceParseError as exc:
        append_diagnostic(job, 'source_parse_failed', str(exc), level='error')
        emit(job, {'type': 'error', 'text': '招标文件本地解析失败：%s' % exc,
                   'actions': [{'act': 'diagnose', 'label': '一键诊断'}]})
        settle(job, stop_reason='已停止（招标文件解析失败）')
    except generation_pipeline.NodeExecutionError as exc:
        if exc.code == 'cancelled': return
        append_diagnostic(job, exc.code, exc.detail, level='error')
        problem = generation_pipeline.summary(job).get('problem') or {}
        node_id = str(problem.get('node_id') or '当前节点')
        title = str(problem.get('title') or node_id)
        if exc.code == 'model_rate_limited':
            text = ('“%s”被上游模型限流或额度拦截；这不是本地引擎卡死。任务已停在检查点，'
                    '请稍后重试当前节点，或在模型接入中切换可用模型。' % title)
        else:
            text = ('“%s”连续重试后仍未完成，任务已停在可恢复检查点；已完成节点不会重跑。' % title)
        emit(job, {'type': 'error',
                   'text': text,
                   'actions': [{'act': 'retry_node', 'label': '重试当前节点', 'param': node_id},
                               {'act': 'open_log', 'label': '查看诊断详情'}]})
        # 终态归因要诚实:以前所有节点失败一律写「连接中断」,把模型输出不合格、
        # 限流等都报成断网,用户排查方向全被带偏(反馈 6 的一部分)。
        reason_map = {
            'model_request_interrupted': '已停止（模型连接中断，可重试当前节点续跑）',
            'model_request_failed': '已停止（模型连接中断，可重试当前节点续跑）',
            'model_rate_limited': '已停止（上游限流，稍后重试当前节点即可）',
            'model_response_empty': '已停止（模型返回为空，可重试当前节点）',
        }
        settle(job, stop_reason=reason_map.get(exc.code, '已停止（节点“%s”执行未通过，可重试该节点）' % title))
    except Exception as exc:
        append_diagnostic(job, 'pipeline_worker_exception', str(exc), level='error')
        problem = generation_pipeline.summary(job).get('problem') or {}
        retry_action = ({'act': 'retry_node', 'label': '重试当前节点',
                         'param': str(problem.get('node_id') or '')}
                        if problem.get('node_id') else
                        {'act': 'resume', 'label': '从检查点继续'})
        emit(job, {'type': 'error', 'text': '分段生成在当前节点停止，已完成内容均已保留。',
                   'actions': [retry_action,
                               {'act': 'open_log', 'label': '查看诊断详情'}]})
        settle(job, stop_reason='已停止（分段生成异常）')

_OC_PROBED = {'ok': False, 'why': '', 'ts': 0.0, 'fingerprint': ''}

def oc_probe_once(ttl=600, failure_ttl=300):
    """探活带缓存:每单都真调一次模型太浪费,但也不能一次都不探 ——
    healthy 返回 True 而 Key 没送到时,opencode 只会甩一句 UnknownError(踩过)。"""
    fingerprint = oc_config_fingerprint()
    same_config = _OC_PROBED.get('fingerprint') == fingerprint
    age = time.time() - float(_OC_PROBED.get('ts') or 0)
    cache_ttl = ttl if _OC_PROBED.get('ok') else failure_ttl
    if same_config and _OC_PROBED.get('ts') and age < cache_ttl:
        return bool(_OC_PROBED.get('ok')), str(_OC_PROBED.get('why') or '')
    ok, why = oc_probe()
    _OC_PROBED.update({'ok': ok, 'why': why, 'ts': time.time(), 'fingerprint': fingerprint})
    return ok, why

def real_agent(job, cmd):
    eng = read_json(conf_path(), {}).get('engine') or {}
    env = agent_env(eng)
    conf_err = env.pop('BIDDOG_SHELL_CONF_ERR', None)
    if conf_err:
        emit(job, {'type': 'error',
                   'text': '%s。这是本机配置问题，不是网关或 Key 的问题；多为数据目录不可写。' % conf_err,
                   'actions': [{'act': 'open_engine', 'label': '检查生成引擎设置'}]})
        halt(job, '已停止（生成服务配置未完成）')
        return
    env['CLAUDE_CODE_PRINT_BG_WAIT_CEILING_MS'] = '0'  # 增补 PATH/附加环境变量 + -p 模式等后台任务跑完
    base = os.path.basename(job)
    if _cancel_requested(base): return
    log = open(os.path.join(job, 'run.log'), 'a', encoding='utf-8')
    emit(job, {'type': 'message', 'role': 'agent', 'text': '生成任务已启动，进度会持续更新。'})
    stop = threading.Event(); known = set(list_deliverables(job)); agent_line = [0]
    def watcher():
        """运行中桥接:新交付物→artifact 事件;progress→progress.json;
        run.log 增量→流式「工作台词」事件(用户最缺的就是'它此刻在干嘛'的文字反馈);
        步进切换→上一步真实耗时入库(预计等待时间的数据来源)。"""
        log_path = os.path.join(job, 'run.log')
        log_off = [0]; cur_step = [0]; step_t0 = [time.time()]
        while not stop.wait(4):
            try:
                # agent 只能写声明流；逐行 JSON 成功解析后才前移 offset，半行不会因轮询被丢掉。
                agent_line[0], accepted = drain_agent_events(job, agent_line[0])
                for safe in accepted:
                    if safe.get('type') == 'progress':
                        stp = int(safe.get('step') or 0)
                        if stp != cur_step[0]:
                            if cur_step[0]: record_stage(cur_step[0], min(time.time() - step_t0[0], 600))
                            cur_step[0] = stp; step_t0[0] = time.time()
                # 工作台词:run.log 新增内容清洗后流出(每轮最多 8 行,防刷屏;单块事件,前端聚合显示)
                try:
                    sz = os.path.getsize(log_path)
                    if sz > log_off[0]:
                        with open(log_path, 'rb') as lf:
                            lf.seek(log_off[0]); chunk = lf.read(min(sz - log_off[0], 65536))
                        log_off[0] += len(chunk)
                        lines = worklog_clean(chunk.decode('utf-8', 'ignore'))
                        if lines:
                            emit(job, {'type': 'worklog', 'lines': lines[-8:]})
                except Exception: pass
                for fn in list_deliverables(job):
                    if fn not in known:
                        known.add(fn); emit(job, {'type': 'artifact', 'name': fn})
            except Exception: pass
    watcher_thread = threading.Thread(target=watcher, daemon=True)
    watcher_thread.start()
    rc = -1; spawn_err = ''; spawn_actions = None; stop_reason = None
    proc = None
    try:
        if _cancel_requested(base):
            stop.set(); watcher_thread.join(timeout=5); log.close(); return
        # 后台跑:stdin 关掉(CLI 不会等交互/弹窗)、新会话(不占当前终端会话,少被 macOS 当前台程序)
        proc = subprocess.Popen(cmd, shell=isinstance(cmd, str), cwd=job, stdin=subprocess.DEVNULL,
                                stdout=log, stderr=log, env=env, **DETACH)
        _register_proc(base, proc)
        _mark_skill_accepted(job, 'cli')
        waited = wait_cli_process(proc, job)
        rc = waited.get('rc') if waited.get('rc') is not None else -1
        if waited['status'] == 'stalled':
            spawn_err = ('生成进程连续 15 分钟没有日志或产物变化，已安全停止。'
                         '已有文件都保留，可查看日志后继续或重跑。')
            spawn_actions = [{'act': 'open_log', 'label': '查看运行日志'},
                             {'act': 'rerun', 'label': '重跑本任务'}]
            stop_reason = '已停止（长时间无进展）'
        elif waited['status'] == 'timeout':
            spawn_err = '生成总时长超过 3 小时，已安全停止；已有交付物保留。'
            spawn_actions = [{'act': 'open_log', 'label': '查看运行日志'},
                             {'act': 'rerun', 'label': '重跑本任务'}]
            stop_reason = '已停止（超过 3 小时总时限）'
    except FileNotFoundError:
        miss = os.path.basename(cmd.split()[0] if isinstance(cmd, str) else cmd[0])
        append_diagnostic(job, 'generator_command_missing', 'missing command: %s' % miss, level='error')
        spawn_err = '当前选择的生成方式不可用。可以修复生成设置，或先用内置演示确认操作流程。'
        spawn_actions = [{'act': 'mock_rerun', 'label': '先用内置演示把流程跑通'},
                         {'act': 'open_engine', 'label': '检查生成设置'}]
    except Exception as e:
        append_diagnostic(job, 'generator_runtime_exception', str(e), level='error')
        spawn_err = '生成过程遇到异常，已有文件都已保留；请一键诊断后重新生成。'
        spawn_actions = [{'act': 'diagnose', 'label': '一键诊断'},
                         {'act': 'rerun', 'label': '重新生成'}]
    finally:
        _clear_proc(base, proc)
    if _cancel_requested(base):              # 用户停止/删除了任务:控制端已经给出明确回执
        stop.set(); watcher_thread.join(timeout=5); log.close(); return
    if spawn_err:
        ev = {'type': 'error', 'text': spawn_err}
        if spawn_actions: ev['actions'] = spawn_actions
        emit(job, ev)
    stop.set(); watcher_thread.join(timeout=5)
    # 进程最后几毫秒写下的声明不能因 watcher 的 4 秒轮询周期丢掉。
    agent_line[0], _accepted = drain_agent_events(job, agent_line[0])
    log.close()
    harvest(job)
    for fn in list_deliverables(job):
        if fn not in known: known.add(fn); emit(job, {'type': 'artifact', 'name': fn})
    # 正文写了却没出 Word → 引擎零 token 补出来。放在分支判断之前:成功、异常退出、3 小时超时
    # 三条路都可能停在"正文写完、最后一步没跑"这个位置,救的动作是同一个。
    try: ensure_docx(job, known)
    except Exception as e:
        append_diagnostic(job, 'word_export_failed', str(e), level='error')
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '⚠ Word 导出没有完成；正文稿已保留，可查看诊断后再次导出。'})
    if spawn_err:
        return settle(job, known, stop_reason=stop_reason)
    if rc == 0 and known:
        return settle(job, known)
        summary, actions = artifact_summary(job, known)
        used = skill_evidence(job)
        if used['state'] == 'unverifiable':
            summary += ('\n\n⚠ **技能包运行证据暂时无法核验**(%s)请人工复核响应矩阵与格式门禁。' % used['why'])
            actions = (actions or []) + [{'act': 'open_log', 'label': '查看运行日志'}]
        elif used['state'] == 'missing':
            summary += ('\n\n⚠ **写作规则没有完整载入**(%s)请在设置中重新测试连接后重跑。' % used['why'])
            actions = (actions or []) + [{'act': 'open_engine', 'label': '去修生成引擎设置'},
                                         {'act': 'open_log', 'label': '查看运行日志'}]
        emit(job, {'type': 'message', 'role': 'agent', 'text': summary, 'actions': actions})
        emit(job, {'type': 'skill_used', 'state': used['state'], 'ok': used['ok'],
                   'hits': used['hits'], 'why': used['why']})
        try: quality_audit(job, known)
        except Exception as e:
            emit(job, {'type': 'message', 'role': 'agent', 'text': '⚠ 成品质检未能执行:%s' % e})
    elif rc == 0 and not known:
        return settle(job, known)
        # 正常退出却一个交付物都没有:以前这里静默,任务永远卡在最后一步
        tail = read_tail(os.path.join(job, 'run.log'), 700)
        append_diagnostic(job, 'no_deliverables_after_run', tail, level='error')
        emit(job, {'type': 'error',
                   'text': '生成过程已经结束，但没有产生可交付文件。已有输入和过程内容仍在，请检查生成设置后重新生成。',
                   'actions': [{'act': 'open_log', 'label': '查看运行日志'},
                               {'act': 'open_engine', 'label': '检查生成设置'},
                               {'act': 'open_job_folder', 'label': '打开任务文件夹'}]})
        halt(job, '已停止(没有产出)')
    elif rc != 0 and not spawn_err:
        tail = ''
        try: tail = open(os.path.join(job, 'run.log'), encoding='utf-8').read()[-600:]
        except Exception: pass
        low = tail.lower()
        if any(k in low for k in ('usage limit', 'rate limit', 'quota', 'insufficient_quota', 'plan limit',
                                  'exceeded your current', 'too many requests', '429')) or '额度' in tail:
            # 订阅额度耗尽:和「模型接入」的 API Key 是两个钱包,要向用户讲清
            append_diagnostic(job, 'generator_quota_exhausted', tail, level='error')
            emit(job, {'type': 'error',
                       'text': '当前生成方式的可用额度已经用完。等额度恢复后可重新生成；急用时可切换生成方式。',
                       'actions': [{'act': 'open_engine', 'label': '切换生成方式'},
                                   {'act': 'mock_rerun', 'label': '先用内置演示跑通流程'}]})
        elif any(k in tail for k in ('连接被对端掐断',)) or \
                any(k in low for k in ('connection reset by peer', 'unexpected_eof',
                                       'stream disconnected before completion', 'eof occurred in violation',
                                       'stream idle timeout', 'no data received within configured window')):
            # 上游网关把连接掐了。中转层已经自动重试过 RETRY_WAITS 次仍不通,不是配置问题,别让客户去翻设置。
            append_diagnostic(job, 'model_connection_interrupted', tail, level='error')
            meta_now = read_json(os.path.join(job, '任务.json'), {})
            retry_action = ({'act': 'resume', 'label': '从已保存内容继续'}
                            if meta_now.get('oc_session') else {'act': 'rerun', 'label': '重新生成'})
            emit(job, {'type': 'error',
                       'text': '**模型服务连接多次中断**，应用已自动重试 %d 次，任务仍未能完整结束。\n\n'
                               '已经写出的章节都保留着，重新生成时会继续利用这些内容。\n\n'
                               '如果连着几次都这样:\n'
                               '① 换个网络试一次（手机热点最快）；\n'
                               '② 过几分钟再试；\n'
                               '③ 急着稳定出件时先切回**标准模式**；极速模式更容易受长连接抖动影响。' % len(RETRY_WAITS),
                       'actions': [retry_action,
                                   {'act': 'open_engine', 'label': '切换模式'},
                                   {'act': 'open_log', 'label': '查看运行日志'}]})
        elif any(k in low for k in ('no api key found for provider', 'auth-profiles', 'auth profile',
                                    'not logged in', 'unauthorized', 'please login', 'please log in',
                                    'agents add', 'no credentials')):
            # 外部 CLI 自己的登录态失效(最典型:SoWork 的 openclaw 找不到 auth-profiles)。
            # 原文是英文 + 本机路径,客户看不懂也不知道能做什么 → 讲清是哪个引擎、为什么、以及一键换成我们的 Key。
            eng_now = (read_json(conf_path(), {}).get('engine') or {}).get('kind', 's2')
            who = {'sowork': 'SoWork(商汤)', 'claude': 'Claude Code', 'codex': 'Codex CLI'}.get(eng_now, '当前生成引擎')
            append_diagnostic(job, 'generator_authorization_missing', tail, level='error')
            emit(job, {'type': 'error',
                       'text': '**%s 没有可用的登录或授权**，所以任务没能开始（不是标书内容的问题）。\n\n'
                               '两条路,任选一条:\n'
                               '① **改用我们发给你的 Key**(推荐,不用登录任何账号):到「设置 · 模型接入」'
                               '把 Key 粘进最上面的快速接入卡,点「一键接入并测试」,然后重新生成。\n'
                               '② 继续用 %s:先在它自己的客户端里重新登录,再回到设置中测试连接，然后重新生成。' % (who, who),
                       'actions': [{'act': 'open_engine', 'label': '去改用我们的 Key'},
                                   {'act': 'mock_rerun', 'label': '先用内置演示跑通流程'},
                                   {'act': 'open_log', 'label': '查看运行日志'}]})
        else:
            append_diagnostic(job, 'generator_exit_failure', 'exit=%s\n%s' % (rc, tail), level='error')
            emit(job, {'type': 'error',
                       'text': '生成过程意外结束，已有文件都已保留。请一键诊断后重新生成。',
                       'actions': [{'act': 'diagnose', 'label': '一键诊断'},
                                   {'act': 'rerun', 'label': '重新生成'}]})
        settle(job, known, stop_reason=('已停止（连接中断，内容已保留）'
                                        if any(k in low for k in ('connection reset by peer', 'unexpected_eof',
                                                                  'stream disconnected before completion',
                                                                  'eof occurred in violation', 'stream idle timeout',
                                                                  'no data received within configured window'))
                                        else None))

# ---------- 生成引擎绑定(claude / codex / 自定义;打包版通过应用内设置,无需环境变量) ----------
AGENT_PROMPT = ('你是标书生成 agent。先完整阅读 {skill}/SKILL.md,然后严格按其流程执行:'
    'tenderPath={tender} outDir={out} materialsDir={materials} skillDir={skill} 。注意:'
    '1) 所有产物直接写入 outDir,不要另建下层输出目录;'
    '2) 招标文件若是 PDF/docx,先转成 招标文件_解析版.md 再分析;'
    '3) 全程同步执行,交付物全部落盘后才允许结束,禁止定时唤醒或把收尾留到以后;{mode}'
    '4) 每完成一个阶段,向 {out}/' + AGENT_EVENTS_FILE + ' 追加一行紧凑 JSON(这是声明流,引擎会按磁盘产物复核):'
    '{"type":"progress","stage":"<阶段名>","pct":<0-100>,"step":<序号>,"total":12};'
    '需要用户决策时追加 {"type":"question","id":"q<n>","text":"<问题>","options":["<选项1>","<选项2>"]} 并继续可并行工作;'
    '5) 最终交付物(投标文件_*.md、*.docx、投标文件自检报告.md)落在 outDir 根目录,出 Word 后必跑 check_docx_format.py 格式门禁;'
    '**没有 .docx 就等于这一单没交付**——客户要的是能直接提交的 Word,不是 md。'
    '正文写完必须亲自跑 build_tender_docx.py 出 Word,跑完 ls 确认文件真的在 outDir 根目录、体积不为 0,'
    '确认到了才允许结束;哪怕时间紧、哪怕正文还想再润色,也要先把 Word 出出来;'
    '6) **素材库是唯一事实来源**:开工前先 ls materialsDir 并通读其中的 公司介绍.md / 产品资料.md / 产品能力表.md / '
    '资质与案例.md / 应答要点.md / 图片索引.md 与 章节模板/ 目录;我方身份、产品能力、资质案例一律取自这里,'
    '缺什么写〔需补充〕,严禁编造,也不要另建空素材目录;'
    '**materialsDir 下若有 参考资料/ 目录**(投标人上传的过往中标标书等),开工前通读它,'
    '**照着它的章节组织方式、行文口径、详略分配写**——那是这家公司真实中过标的写法;'
    '但事实数据(公司名、业绩、资质、人员、报价)一律以素材库其余文件为准,'
    '**绝不许从参考件里搬**,那是另一个项目的,搬过来就是编造;'
    '7) **必须配图**:若 materialsDir/图片索引.md 存在,撰写时在讲到对应能力/架构/资质处独立成行打标 {{图:图片ID}}'
    '(ID 只能用索引里登记过的,按索引的"落位锚点"插),出 Word 时给 build_tender_docx.py 传 --images-dir "<materialsDir>",'
    '让图片真正插进文档;索引不存在则在正文标注〔配图建议:说明〕;'
    '8) **正文必须是完整段落**:严禁把一句话拆成一行一个字、严禁同一段落反复灌注多次(会被内容门禁判定为废稿);'
    '**篇幅按招标文件的分量走,不设统一字数目标**——给定一个数字,弱模型就会靠复制模板去凑'
    '(真机事故:给了「8—9 万字符」,它给 18 个章节套同一组小标题灌到 12 万字)。'
    '论述章写完用 wc -m 看一眼,明显偏薄的就地扩写;**格式件、证明件、表单本来就短,不要往里灌内容**;'
    '9) **出 Word 之前必跑质检脚本**:python3 {skill}/references/quality_gate.py <交付稿.md> '
    '--materials {materials} --fix ——它会按图片索引的锚点自动校正图片位置、补插漏图、剔除不存在的图片ID、'
    '折叠重复段落,并产出《成品质检报告.md》;跑完再出 Word。图片ID 只准用图片索引里登记过的,严禁自造。'
    '10) **开工第一件事:读 outDir/你的要求.md**(存在的话)——那是本次客户自己写的要求,'
    '优先级仅次于招标文件的强制条款与格式门禁,高于你自己的写作习惯;'
    '里面的角色设定、章节侧重、语气风格、必须生成的承诺函等,逐条落实(客户自己写了字数要求才按它来),'
    '并在《投标文件自检报告》里逐条说明落实情况(未落实的必须写明原因)。')

# 生成方式:agents=主控编排(SKILL.md 路径B,稳;分析与撰写环节并行提速);workflow=一条流水线并行(快,但更易半路失败)
MODE_AGENTS = ('**必须走 SKILL.md 的【路径 B】:由你作为主控编排子 agent 并亲自验收**,'
    '禁止调用 Workflow 工具跑 multiagent_workflow.js(并行流水线中途失败会整条断掉,产出不稳定)。'
    '**提速纪律(不许串行磨洋工):三个前置分析员必须一次性并行召唤;分章撰写必须并行召唤'
    '(每章一个子 agent 同时派出,平台不支持并行才逐章)——逐章串行是首要的时间浪费**。'
    '每步做完先核对产物文件是否真的写入、用 wc -m 核字数,不达标只重做那一章;'
    '汇总必须用 {skill}/references/assemble_tender.py 脚本拼装(禁止让模型整文重写输出),'
    '配图复核不再单独召唤子 agent——用第 9 条的 quality_gate.py 脚本替代(秒级、零 token);'
    '汇总成册和出 Word 必须由你亲自触发,不要假设子 agent 会自动汇总。')
MODE_WORKFLOW = ('可用 Workflow 工具跑 {skill}/references/multiagent_workflow.js(并行更快,但中途失败易整条中断);'
    '跑完必须逐项核对交付物是否齐全,缺什么就自己补做。')

def skill_dir_conf(conf=None):
    eng = (conf or read_json(conf_path(), {})).get('engine') or {}
    custom = eng.get('skill_dir') or os.environ.get('SKILL_DIR')
    return custom or ensure_skill()   # 托管默认目录必须走 ensure_skill:带版本标记,升级后自动刷新存量目录

SKILL_VERSION = '5.11'  # 技能包内容版本:表格全左对齐并清除字符缩进/语义列宽;存量目录自动刷新

def ensure_skill():
    """内置技能包:首次运行自动解压到数据目录;版本升级后老目录自动刷新,修复能到存量用户手里"""
    dst = os.path.join(DATA, 'bid-multiagent-tao')
    ver = os.path.join(dst, '.skill_version')
    try: cur = open(ver, encoding='utf-8').read().strip()
    except Exception: cur = ''
    # 版本标记之外再验一个关键文件:曾出过"旧 zip 抢先解压却打上新版本标记"的事故,标记从此不可全信
    if (os.path.isfile(os.path.join(dst, 'SKILL.md')) and cur == SKILL_VERSION
            and os.path.isfile(os.path.join(dst, 'references', 'quality_gate.py'))): return dst
    cands = []
    meipass = getattr(sys, '_MEIPASS', None)           # PyInstaller 打包版:zip 嵌在二进制里
    if meipass: cands.append(os.path.join(meipass, 'bidmultiagenttao_v5.3.zip'))
    cands += [os.path.join(HERE, 'bidmultiagenttao_v5.3.zip'),
              os.path.join(HERE, '..', 'bidmultiagenttao_v5.3.zip')]
    # 有多份同名 zip 时按修改时间取最新的:开发机上 server/ 常留着 CI 那步 cp 过来的旧副本,
    # 按顺序取第一个会让它抢在根目录的新包前面解压——技能包改了却"没生效",排查起来极费时间(踩过两次)
    exist = [z for z in cands if os.path.isfile(z)]
    if len(exist) > 1 and not meipass:
        exist.sort(key=lambda z: os.path.getmtime(z), reverse=True)
        cands = exist
    for z in cands:
        if os.path.isfile(z):
            try:
                zipfile.ZipFile(z).extractall(DATA)
                if os.path.isfile(os.path.join(dst, 'SKILL.md')):
                    open(ver, 'w', encoding='utf-8').write(SKILL_VERSION)
                    return dst
            except Exception: pass
    return dst

def _toml_s(v): return json.dumps(str(v), ensure_ascii=False)   # TOML 基本字符串与 JSON 字符串同形,借用转义

def codex_home_s2(eng=None):
    """给 S2 引擎单独造一个 CODEX_HOME:客户机器上原有的 ~/.codex(登录态、订阅额度)一个字都不动。
    每次运行都重写,配置永远和界面上填的一致——省掉「改了设置没生效」这类看不见的坑。"""
    eng = eng or (read_json(conf_path(), {}).get('engine') or {})
    d = _mk(os.path.join(ws_root(), 'codex_s2'))
    up = s2_conf()
    direct = (eng.get('s2_wire') or 'auto') == 'responses' and (eng.get('s2_direct') is True)
    base = up['base_url'] if direct else relay_base()
    toml = '\n'.join([
        '# 中标狗自动生成,请勿手改(每次运行会被覆盖)。作者:' + AUTHOR,
        'model = ' + _toml_s(up['model']),
        'model_provider = "biddog_s2"',
        'approval_policy = "never"',
        'sandbox_mode = "danger-full-access"',
        'hide_agent_reasoning = true',
        '',
        '[model_providers.biddog_s2]',
        'name = "中标狗 · S2"',
        'base_url = ' + _toml_s(base),
        'env_key = "BIDDOG_S2_KEY"',
        # 实测 codex-cli 0.146.0:wire_api = "chat" 已被移除,自定义网关只能走 responses,
        # 所以「翻译成 chat/completions」这件事必须由我们这一层做。
        'wire_api = "responses"',
        'request_max_retries = 2',
        'stream_max_retries = 2',
        ''])
    open(os.path.join(d, 'config.toml'), 'w', encoding='utf-8').write(toml)
    return d, base, direct

def opencode_home_s2(eng=None):
    """opencode 的隔离配置:provider 指向本机直通端点,Key 用 {env:} 引用中转口令。
    与用户自己的 opencode(XDG 目录)完全隔离,每次运行重写保持与界面一致。"""
    eng = eng or (read_json(conf_path(), {}).get('engine') or {})
    d = _mk(os.path.join(ws_root(), 'opencode_s2'))
    up = s2_conf()
    direct = (eng.get('s2_wire') or 'auto') == 'responses' and (eng.get('s2_direct') is True)
    base = (up['base_url'] if direct else relay_base())
    routed_models = []
    for model_id in (up['model'], S2_DEFAULT_MODEL, S2_QUALITY_MODEL):
        if model_id and model_id not in routed_models: routed_models.append(model_id)
    conf = {
        '$schema': 'https://opencode.ai/config.json',
        'provider': {'biddog-s2': {
            'npm': '@ai-sdk/openai-compatible', 'name': '中标狗 · S2',
            'options': {'baseURL': base, 'apiKey': '{env:BIDDOG_S2_KEY}'},
            'models': {model_id: {'name': model_id} for model_id in routed_models}}},
        # 默认流水线把每个节点放进独立目录，并把声明输入复制进去；模型不需要访问
        # 客户电脑的其他路径，也不需要 shell 或子 agent。旧的宽权限整单模式仅可通过
        # 明确的本机环境开关启用，不能成为安装包默认值。
        'permission': {
            'edit': 'allow', 'read': 'allow', 'glob': 'allow', 'grep': 'allow',
            'bash': ('allow' if os.environ.get('BIDDOG_UNSAFE_LEGACY_OPENCODE') == '1' else 'deny'),
            'task': ('allow' if os.environ.get('BIDDOG_UNSAFE_LEGACY_OPENCODE') == '1' else 'deny'),
            'external_directory': ('allow' if os.environ.get('BIDDOG_UNSAFE_LEGACY_OPENCODE') == '1' else 'deny'),
            'webfetch': 'deny', 'skill': 'deny'},
    }
    write_json(os.path.join(d, 'opencode.json'), conf)
    return d, base, direct

def opencode_env(eng):
    d, _base, direct = opencode_home_s2(eng)
    up = s2_conf()
    home = _mk(os.path.join(d, 'home'))
    no = ','.join(x for x in ['127.0.0.1', 'localhost', os.environ.get('NO_PROXY', ''), os.environ.get('no_proxy', '')] if x)
    return {'OPENCODE_CONFIG': os.path.join(d, 'opencode.json'),
            'BIDDOG_S2_KEY': (up['api_key'] if direct else relay_token()),
            'XDG_DATA_HOME': os.path.join(home, 'data'), 'XDG_CONFIG_HOME': os.path.join(home, 'cfg'),
            'XDG_CACHE_HOME': os.path.join(home, 'cache'),
            'NO_PROXY': no, 'no_proxy': no}

def s2_env(eng):
    """S2 引擎专用环境:指向我们自己的 CODEX_HOME;直连时给真 Key,走中转时只给本机中转口令(真 Key 不出引擎进程)"""
    d, _base, direct = codex_home_s2(eng)
    up = s2_conf()
    no = ','.join(x for x in ['127.0.0.1', 'localhost', os.environ.get('NO_PROXY', ''), os.environ.get('no_proxy', '')] if x)
    return {'CODEX_HOME': d, 'BIDDOG_S2_KEY': (up['api_key'] if direct else relay_token()),
            'NO_PROXY': no, 'no_proxy': no}   # 公司网络的代理变量会把「连本机中转」也代理掉,必须排除回环

def config_agent_cmd():
    """真实 agent 命令:env AGENT_CMD 最优先(字符串,shell 执行);否则按应用内设置生成(列表,免引号跨平台)"""
    env = os.environ.get('AGENT_CMD', '')
    if env: return env
    conf = read_json(conf_path(), {})
    eng = conf.get('engine') or {}
    kind = eng.get('kind', 's2')      # 默认「自动」:没 Key 时下一行会回落演示
    if kind == 'custom': return eng.get('cmd', '')
    # 「自动」的语义:有 Key 才产真实标书,没 Key 就返回空 → 上层自动回落内置演示流程(不报错、不空跑)
    if kind in ('s2', 'opencode') and not s2_conf(conf)['api_key']: return ''
    if kind not in ('claude', 'codex', 'sowork', 's2', 'opencode'): return ''
    sd = skill_dir_conf(conf)
    if not os.path.isfile(os.path.join(sd, 'SKILL.md')): sd = ensure_skill()
    mode = MODE_WORKFLOW if eng.get('mode') == 'workflow' else MODE_AGENTS   # 默认稳健的多子 agent
    prompt = AGENT_PROMPT.replace('{mode}', mode).replace('{skill}', sd)
    if kind == 'claude':
        sf = os.path.join(DATA, 'agent_settings.json')
        if not os.path.isfile(sf):
            write_json(sf, {'permissions': {'allow': ['Bash(*)', 'Read(*)', 'Write(*)', 'Edit(*)', 'MultiEdit(*)',
                'Glob(*)', 'Grep(*)', 'Task(*)', 'TodoWrite', 'LS(*)', 'WebFetch(*)', 'Workflow(*)', 'ToolSearch(*)', 'Skill(*)', 'NotebookEdit(*)']}})
        return [resolve_cli('claude', eng) or 'claude', '-p', prompt, '--settings', sf]
    if kind == 'sowork':
        # SoWork(商汤)自带 openclaw CLI:用本机已登录账号与网关,不另配 Key。
        # --thinking off:SenseAudio-S2 目前只支持 off,填 medium/high 会直接报错退出。
        # --session-key 带任务号:多任务/多人并行时会话不串。
        return [resolve_cli('sowork', eng) or 'openclaw', 'agent',
                '--agent', (eng.get('sowork_agent') or 'main'),
                '--session-key', 'bid-dog-{jobid}',
                '--thinking', (eng.get('thinking') or 'off'),
                '--timeout', str(int(eng.get('timeout') or 1800)),
                '--message', prompt]
    if kind in ('s2', 'opencode'):
        # 「自动(默认)」= OpenCode。拿真实招标文件做过同模型、同素材、同提示词的对比,
        # 它在唯一严格可比的那组里全面胜出:偏离表 112 行 vs 36 行(应答密度 6.2x vs 2.0x)、
        # 出了 codex 漏掉的《评标索引》、废标风险清单 19.7KB vs 0.7KB、
        # 三份里唯一质检判「通过」的,而模板化命中同为 0(是写透了不是灌水)。
        # 结构性原因:它有原生 task 子 agent,每章一个子 agent 独立写、各自有完整上下文预算;
        # codex 在单一上下文里串行写完六章,越往后预算越紧(实测它第 1 章只有 1320 字)。
        #
        # opencode 原生 OpenAI 兼容:baseURL 指本机直通端点即可,零协议翻译。--auto=非交互放行。
        # --dir {out}:把工作目录钉死在任务目录(它默认会向上找"项目根",可能钉错层)。
        up = s2_conf(conf)
        return [resolve_cli('opencode', eng) or 'opencode', 'run', '--auto', '--dir', '{out}',
                '-m', 'biddog-s2/' + up['model'], prompt]
    # codex:保留为显式可选项。切换到 OpenCode 只有一单样本支撑,客户真撞上问题要能一键切回来。
    # --skip-git-repo-check:任务目录不是 git 仓库也能跑
    # --dangerously-bypass-approvals-and-sandbox:非交互执行,免"信任目录/审批"卡住(本机自有目录)
    return [resolve_cli('codex', eng) or 'codex', 'exec', '--skip-git-repo-check', '--dangerously-bypass-approvals-and-sandbox', prompt]

def packaged_location_check(executable=None):
    path = os.path.abspath(str(executable or sys.executable or ''))
    if sys.platform == 'darwin' and '/AppTranslocation/' in path:
        return {'id': 'app_location', 'label': '应用位置', 'status': 'warning',
                'message': '当前从临时隔离位置运行；请把中标狗拖入“应用程序”后重新打开，路径会更稳定'}
    return {'id': 'app_location', 'label': '应用位置', 'status': 'pass',
            'message': '应用安装位置正常'}

def generation_preflight(job, conf=None):
    """Fast local checks shared by task startup and one-click diagnostics."""
    conf = conf or read_json(conf_path(), {})
    eng = conf.get('engine') if isinstance(conf.get('engine'), dict) else {}
    kind = str(eng.get('kind') or 's2')
    checks = []
    storage_ok = os.path.isdir(job) and os.access(job, os.W_OK)
    checks.append({'id': 'storage', 'label': '任务目录', 'status': 'pass' if storage_ok else 'fail',
                   'message': '任务文件可以保存' if storage_ok else '任务目录不可写'})
    checks.append(packaged_location_check())
    skill = skill_dir_conf(conf)
    skill_ok = os.path.isfile(os.path.join(skill, 'SKILL.md'))
    checks.append({'id': 'skill', 'label': '写作规则', 'status': 'pass' if skill_ok else 'fail',
                   'message': '写作规则已就绪' if skill_ok else '写作规则包缺失'})
    repair = ''
    if kind in ('s2', 'opencode'):
        if not s2_conf(conf)['api_key']:
            checks.append({'id': 'connection', 'label': '模型连接', 'status': 'demo',
                           'message': '未填写 Key，本轮使用内置演示流程'})
            checks.append({'id': 'runtime', 'label': '生成组件', 'status': 'skipped',
                           'message': '演示流程无需外部生成组件'})
        else:
            verified_models = _verified_model_ids(conf)
            checks.append({'id': 'connection', 'label': '模型连接',
                           'status': 'pending' if verified_models else 'fail',
                           'message': ('将在派发前建立模型连接' if verified_models else
                                       '当前网关或模型未验证，请先在“模型接入”测试连接')})
            shell = resolve_cli('opencode', eng)
            if shell:
                checks.append({'id': 'runtime', 'label': '生成组件', 'status': 'pass',
                               'message': '内置 OpenCode 已就绪', 'source': shell})
            else:
                repair = 'opencode'
                checks.append({'id': 'runtime', 'label': '生成组件', 'status': 'repairing',
                               'message': '缺少内置生成组件，将自动修复'})
    else:
        cli_name = {'codex': 'codex', 'claude': 'claude', 'sowork': 'sowork'}.get(kind)
        shell = resolve_cli(cli_name, eng) if cli_name else bool(config_agent_cmd())
        checks.append({'id': 'runtime', 'label': '生成组件', 'status': 'pass' if shell else 'fail',
                       'message': '所选生成方式已就绪' if shell else '所选生成方式不可用'})
    ok = storage_ok and skill_ok and not repair and not any(item['status'] == 'fail' for item in checks)
    result = {'ok': ok, 'kind': kind, 'phase': 'environment', 'repair': repair,
              'checks': checks, 'checked_at': now()}
    try: write_json(os.path.join(job, 'preflight.json'), result)
    except Exception: pass
    return result


def mark_connection_check(job, ok, detail=''):
    """把 preflight.json 里的「模型连接」检查项改成本次运行的真实结果。

    真机反馈 5:该项写的是启动时的静态判断(「将在派发前建立模型连接」),之后
    连接成功/断开都不再更新,任务导航栏的「模型连接」与实际 100% 脱节。
    现在:首次真实调用成功置 pass;断连/失败置 fail 并带原因;自动恢复后再置回。"""
    try:
        path = os.path.join(job, 'preflight.json')
        pf = read_json(path, {})
        changed = False
        for item in pf.get('checks') or []:
            if item.get('id') == 'connection':
                status = 'pass' if ok else 'fail'
                message = (detail or ('连接已建立' if ok else '连接中断'))[:160]
                if item.get('status') != status or item.get('message') != message:
                    item['status'], item['message'] = status, message
                    changed = True
        if changed:
            pf['checked_at'] = now()
            write_json(path, pf)
    except Exception:
        pass


def _bootstrap_repair(kind='none', action='', label='无需处理'):
    return {'kind': kind, 'action': action, 'label': label}


def _bootstrap_item(item_id, label, status, message, estimate_seconds, repair=None):
    return {'id': item_id, 'label': label, 'status': status, 'message': message,
            'estimate_seconds': list(estimate_seconds),
            'repair': repair or _bootstrap_repair()}


def bootstrap_checks(conf=None):
    """Return a fast, non-billing startup health report.

    This deliberately checks the exact local capabilities needed before a task
    starts.  It does not call the paid model path; a prior end-to-end probe is
    accepted only while its connection fingerprint still matches the current
    gateway, key, transport and model.
    """
    conf = conf or read_json(conf_path(), {})
    eng = conf.get('engine') if isinstance(conf.get('engine'), dict) else {}
    checks = []

    storage_ok = False
    try:
        root = _mk(ws_root())
        probe = os.path.join(root, '.bootstrap-write-%s' % secrets.token_hex(4))
        with open(probe, 'w', encoding='utf-8') as f:
            f.write('ok')
        os.remove(probe)
        storage_ok = True
    except OSError:
        storage_ok = False
    checks.append(_bootstrap_item(
        'storage', '保存位置', 'pass' if storage_ok else 'manual',
        '任务和结果可以保存' if storage_ok else '数据目录不可写，请检查磁盘和文件夹权限',
        (0, 1), None if storage_ok else
        _bootstrap_repair('manual', 'open_storage_help', '查看权限处理办法')))

    try:
        import docx as _docx  # noqa: F401
        import pypdf as _pypdf  # noqa: F401
        parsers_ok = True
    except ImportError:
        parsers_ok = False
    checks.append(_bootstrap_item(
        'parsers', '文档解析', 'pass' if parsers_ok else 'manual',
        'DOCX 和 PDF 解析能力已就绪' if parsers_ok else '安装包缺少文档解析组件，需要重新安装完整应用',
        (0, 2), None if parsers_ok else
        _bootstrap_repair('reinstall', 'reinstall_app', '重新安装完整组件')))

    skill = skill_dir_conf(conf)
    skill_ok = os.path.isfile(os.path.join(skill, 'SKILL.md'))
    checks.append(_bootstrap_item(
        'skill', '标书写作规则', 'pass' if skill_ok else 'repairable',
        '标书技能包已内置' if skill_ok else '标书技能包缺失，可以从安装包恢复',
        (0, 1), None if skill_ok else
        _bootstrap_repair('automatic', 'repair_skill', '一键恢复写作规则')))

    kind = str(eng.get('kind') or 's2')
    if kind in ('s2', 'opencode'):
        runtime_ok = bool(resolve_cli('opencode', eng))
        checks.append(_bootstrap_item(
            'runtime', '本地生成组件', 'pass' if runtime_ok else 'repairable',
            '内置 OpenCode 已就绪' if runtime_ok else '生成组件缺失或不可运行，可以在线恢复',
            (1, 3), None if runtime_ok else
            _bootstrap_repair('automatic', 'repair_opencode', '一键修复生成组件')))
        up = s2_conf(conf)
        if not up['api_key']:
            connection_status = 'manual'
            connection_message = '尚未填写模型 Key'
            connection_repair = _bootstrap_repair('manual', 'open_model_settings', '去接入模型')
        elif not _verified_model_ids(conf):
            connection_status = 'manual'
            connection_message = '模型设置尚未通过完整生成链路测试'
            connection_repair = _bootstrap_repair('manual', 'test_model', '测试模型连接')
        else:
            connection_status = 'pass'
            connection_message = '模型与完整生成链路已验证'
            connection_repair = None
        checks.append(_bootstrap_item('connection', '模型连接', connection_status,
                                      connection_message, (5, 20), connection_repair))
    else:
        cli_name = {'codex': 'codex', 'claude': 'claude', 'sowork': 'sowork'}.get(kind)
        runtime_ok = bool(resolve_cli(cli_name, eng)) if cli_name else bool(config_agent_cmd())
        checks.append(_bootstrap_item(
            'runtime', '本地生成组件', 'pass' if runtime_ok else 'manual',
            '所选生成方式已就绪' if runtime_ok else '所选生成方式不可用', (1, 3),
            None if runtime_ok else _bootstrap_repair('manual', 'open_model_settings', '检查生成方式')))
        checks.append(_bootstrap_item('connection', '模型连接', 'pass' if runtime_ok else 'manual',
                                      '生成方式已验证' if runtime_ok else '请先验证所选生成方式', (5, 20),
                                      None if runtime_ok else
                                      _bootstrap_repair('manual', 'test_model', '测试模型连接')))

    statuses = {item['status'] for item in checks}
    if 'manual' in statuses:
        status = 'manual'
    elif 'repairable' in statuses:
        status = 'repairable'
    else:
        status = 'ready'
    return {'status': status, 'ready_for_generation': status == 'ready',
            'checks': checks, 'checked_at': now()}


@app.get('/v1/bootstrap/checks')
def get_bootstrap_checks():
    return bootstrap_checks()

def ensure_default_shell(job, eng):
    """Repair the managed OpenCode shell in-place, with visible task progress."""
    if str((eng or {}).get('kind') or 's2') not in ('s2', 'opencode'):
        return True, '当前使用显式生成方式'
    ready = resolve_cli('opencode', eng)
    if ready: return True, '生成组件已就绪'
    emit(job, {'type': 'worklog', 'lines': ['检查到生成组件缺失，正在自动修复生成组件（无需安装 Python、Node 或 SoWork）']})
    emit(job, {'type': 'progress', 'stage': '正在修复生成组件', 'pct': 1, 'step': 0, 'total': 12})
    if PROV.get('state') != 'running':
        PROV.update({'state': 'idle', 'which': 'opencode', 'pct': 0, 'note': '', 'path': '', 'error': ''})
        worker = threading.Thread(target=_provision_codex, args=('opencode',), daemon=True)
        worker.start()
    else:
        worker = None
    last = None
    deadline = time.time() + 10 * 60
    while time.time() < deadline:
        state = str(PROV.get('state') or 'idle')
        snapshot = (state, int(PROV.get('pct') or 0), str(PROV.get('note') or ''))
        if snapshot != last:
            last = snapshot
            note = snapshot[2] or ('正在下载生成组件 %d%%' % snapshot[1])
            emit(job, {'type': 'worklog', 'lines': [note]})
        if state in ('done', 'error'): break
        if worker is not None and not worker.is_alive() and state != 'running': break
        time.sleep(.25)
    ready = resolve_cli('opencode', eng)
    if ready:
        emit(job, {'type': 'worklog', 'lines': ['生成组件已就绪，继续建立模型连接']})
        return True, '生成组件已就绪'
    why = str(PROV.get('error') or '自动修复没有完成')
    append_diagnostic(job, 'opencode_auto_repair_failed', why, level='error')
    return False, why

@app.get('/v1/agent')
def agent_status():
    conf = read_json(conf_path(), {})
    eng = conf.get('engine') or {}
    sd = skill_dir_conf(conf)
    if not os.path.isfile(os.path.join(sd, 'SKILL.md')): sd = ensure_skill()
    cl, cx, sw, oc = resolve_cli('claude', eng), resolve_cli('codex', eng), resolve_cli('sowork', eng), resolve_cli('opencode', eng)
    return {'kind': 'env' if os.environ.get('AGENT_CMD') else eng.get('kind', 's2'), 'mode': eng.get('mode', 'agents'),
            'generation_mode': _pipeline_mode(conf),
            # 自定义命令和环境变量也可能内嵌 Key：与 s2_key 一样只写不读。
            'cmd': '', 'cmd_set': bool((eng.get('cmd') or '').strip()),
            'skill_dir': sd, 'skill_ok': os.path.isfile(os.path.join(sd, 'SKILL.md')),
            'available': {'claude': bool(cl), 'codex': bool(cx), 'sowork': bool(sw), 'opencode': bool(oc)},
            'paths': {'claude': cl or '', 'codex': cx or '', 'sowork': sw or '', 'opencode': oc or ''},
            'cli_path': eng.get('cli_path', ''), 'env': '', 'env_set': bool((eng.get('env') or '').strip()),
            'login_shell': eng.get('login_shell', True), 'sowork_agent': eng.get('sowork_agent', 'main'),
            'thinking': eng.get('thinking', 'off'), 'timeout': eng.get('timeout', 1800),
            's2_base_url': eng.get('s2_base_url', ''), 's2_model': eng.get('s2_model', ''),
            's2_key_set': bool((eng.get('s2_key') or '').strip()),   # 只回是否已填,不把 Key 回传给页面
            's2_wire': eng.get('s2_wire', 'auto'), 's2_verify_ssl': eng.get('s2_verify_ssl', True),
            's2_defaults': {'base_url': S2_DEFAULT_BASE, 'model': S2_DEFAULT_MODEL},
            's2_borrowed': (not (eng.get('s2_key') or '').strip()) and bool(s2_conf(conf)['api_key']),
            'codex_bundled': bool(bundled_codex()), 'opencode_bundled': bool(bundled_cli('opencode-cli')),
            's2_model_effective': s2_conf(conf)['model'],
            'generation_params': _generation_params(conf)}

def config_locked_jobs():
    """只让正在执行或暂停的会话锁住全局模型。

    已停止的历史任务不应阻塞新任务切换默认模式；若其旧会话与新模型不一致，
    resume 入口会明确要求重跑，不会跨模型偷偷续写。
    """
    locked = set(_running_snapshot())
    try: ids = os.listdir(jobs_dir())
    except Exception: ids = []
    for jid in ids:
        job = jpath(jid)
        if not os.path.isdir(job): continue
        meta = read_json(os.path.join(job, '任务.json'), {})
        if meta.get('oc_session') and meta.get('paused'):
            locked.add(os.path.basename(jid))
    return sorted(locked)

@app.put('/v1/agent')
async def set_agent(req: Request):
    # 云端网页模式默认禁止:网页端可改命令 = 给访客开服务器远程执行
    if MULTIUSER and not ALLOW_AGENT_CONFIG:
        return JSONResponse({'ok': False, 'error': '云端部署已锁定生成引擎设置(由部署方通过环境变量配置)'}, 403)
    body = await req.json()
    conf = read_json(conf_path(), {'providers': [], 'routing': {}})
    old = conf.get('engine') or {}
    old_fingerprint = oc_config_fingerprint(conf)
    requested_generation_mode = str(body.get('generation_mode') or '').lower()
    if requested_generation_mode not in ('fast', 'standard'):
        requested_generation_mode = ('standard' if str(body.get('s2_model') or '') == S2_QUALITY_MODEL
                                     else str(old.get('generation_mode') or 'fast'))
    new_engine = {'kind': body.get('kind', 'mock'),
                      # 留空表示沿用已保存值，避免 GET 不回显后模式切换把秘密配置误清空。
                      'cmd': (body.get('cmd') or '').strip() or old.get('cmd', ''),
                      'skill_dir': body.get('skill_dir', ''), 'mode': body.get('mode', 'agents'),
                      'cli_path': body.get('cli_path', ''),
                      'env': (body.get('env') or '').strip() or old.get('env', ''),
                      'login_shell': body.get('login_shell', True),
                      'sowork_agent': body.get('sowork_agent', 'main'),
                      'thinking': body.get('thinking', 'off'),
                      'timeout': body.get('timeout', 1800),
                      'generation_mode': requested_generation_mode,
                      's2_base_url': body.get('s2_base_url', ''), 's2_model': body.get('s2_model', ''),
                      # Key 留空 = 沿用已存的(页面不回显 Key,不能因为"没传"就把它清掉)
                      's2_key': (body.get('s2_key') or '').strip() or old.get('s2_key', ''),
                      's2_wire': body.get('s2_wire', 'auto'), 's2_verify_ssl': body.get('s2_verify_ssl', True),
                      # 生成参数:传了就按传的(越界钳到范围内),没传沿用已存的
                      'generation_params': (_normalize_generation_params(body.get('generation_params'))
                                            if body.get('generation_params') is not None
                                            else (old.get('generation_params') or {}))}
    if body.get('cmd_clear'): new_engine['cmd'] = ''
    if body.get('env_clear'): new_engine['env'] = ''
    if body.get('s2_key_clear'): new_engine['s2_key'] = ''
    candidate = dict(conf); candidate['engine'] = new_engine
    new_fingerprint = oc_config_fingerprint(candidate)
    params_changed = _generation_params(candidate) != _generation_params(conf)
    if config_locked_jobs() and (new_fingerprint != old_fingerprint or params_changed):
        return JSONResponse({'ok': False,
                             'error': '还有任务正在生成。为避免同一任务中途换模型或参数，请等任务结束或先停止任务再切换。'}, 409)
    conf['engine'] = new_engine
    write_json(conf_path(), conf)
    if new_fingerprint != old_fingerprint: invalidate_oc_runtime()
    return {'ok': True, 's2_model_effective': s2_conf(conf)['model'],
            'generation_params': _generation_params(conf)}

@app.post('/v1/agent/test')
def agent_test():
    """一键连通性测试:用当前绑定的引擎发一条最小指令,把结果翻译成人话。
    客户不用开终端——这是「终端能跑、App 里不行」类问题的自助排查入口。"""
    conf = read_json(conf_path(), {})
    eng = conf.get('engine') or {}
    kind = 'env' if os.environ.get('AGENT_CMD') else eng.get('kind', 's2')
    if kind == 'mock':
        return {'ok': True, 'note': '当前是内置演示流程,无需外部 CLI(要产真实标书请选一个生成引擎)'}
    probe = '只回复六个字:中标狗连接成功'
    if kind == 'sowork':
        cli = resolve_cli('sowork', eng)
        if not cli: return {'ok': False, 'error': '没找到 SoWork 的 openclaw 命令。请确认已安装并登录 SoWork;'
                                                 '若装在非默认位置,把完整路径填到下面的「CLI 路径」。'}
        cmd = [cli, 'agent', '--agent', eng.get('sowork_agent') or 'main',
               '--session-key', 'bid-dog-selftest', '--thinking', eng.get('thinking') or 'off',
               '--timeout', '60', '--message', probe]
    elif kind == 'claude':
        cli = resolve_cli('claude', eng)
        if not cli: return {'ok': False, 'error': '没找到 claude 命令,请先安装 Claude Code CLI 并登录'}
        cmd = [cli, '-p', probe]
    elif kind == 'codex':
        cli = resolve_cli('codex', eng)
        if not cli: return {'ok': False, 'error': '没找到 codex 命令,请先安装 Codex CLI 并登录'}
        cmd = [cli, 'exec', '--skip-git-repo-check', '--dangerously-bypass-approvals-and-sandbox', probe]
    elif kind in ('s2', 'opencode'):
        # 三层分开验:①Key/网关能不能通 ②执行外壳在不在 ③整条链路(外壳→中转→网关)能不能跑通。
        # 分层报错是为了让客户自己看得懂卡在哪一层,而不是只看到一句"异常退出"。
        # s2 与 opencode 现在是同一个引擎(「自动(默认)」= OpenCode),自检合并成一条,
        # 免得两处分头维护、迟早说出两套不一样的话。
        up = s2_conf(conf)
        if not up['api_key']:
            return {'ok': False, 'error': '还没填 API Key——现在新建任务会跑内置演示流程(样例稿)。'
                                          '把你收到的那串 Key(sk-…)填进来即产真实标书;'
                                          '或先在「模型接入」里加好接入点,这里会自动借用。'}
        try:
            data = _openai_req(up['base_url'], up['api_key'], '/models', timeout=20, verify=up['verify_ssl'])
            ids = [m.get('id') for m in (data.get('data') or []) if m.get('id')]
            if ids and up['model'] not in ids:
                return {'ok': False, 'error': '网关通了,但你的套餐里没有模型「%s」。可用的是:%s' % (up['model'], '、'.join(ids[:8]))}
        except urllib.error.HTTPError as e:
            code = e.code
            msg = {401: 'Key 不对或已停用', 403: 'Key 没有这个网关的权限', 429: 'Key 的额度/频率超限'}.get(code, 'HTTP %s' % code)
            return {'ok': False, 'error': '连不上 S2 网关:%s。地址:%s' % (msg, up['base_url'])}
        except Exception as e:
            return {'ok': False, 'error': '连不上 S2 网关:%s' % net_hint(e, (up['api_key'],))}
        cli = resolve_cli('opencode', eng)
        if not cli:
            return {'ok': False, 'need_provision': True,
                    'error': '模型服务本身可以连接，只需修复内置生成组件。点下面的「一键修复生成组件」即可'
                             '(压缩约 60MB、解压约 170MB,不需要登录、不消耗任何订阅额度);'
                             '或手动 npm i -g opencode-ai。'}
        # 必须验证正式任务实际使用的常驻会话链路。短命令 `opencode run` 即使成功，
        # 也证明不了 serve → session → prompt → 完整收尾这条生产路径不会中断。
        t0 = time.time()
        if not oc_serve(eng):
            return {'ok': False, 'execution_path': 'opencode_server',
                    'error': '正式生成链路未能启动。内置生成组件存在，但常驻会话服务没有就绪。'}
        ok, why = oc_probe()
        latency = int((time.time() - t0) * 1000)
        if not ok:
            return {'ok': False, 'execution_path': 'opencode_server', 'latency_ms': latency,
                    'model': up['model'],
                    'error': '正式生成链路测试失败：%s' % (_safe_secret_text(why, (up['api_key'],)) or '模型会话没有完整返回')}
        # 切换极速/标准时，不能因 `/models` 里有这个名字就冻结路由。
        # 在完整 OpenCode 会话通过后，再对该模式所需文本模型逐一直调。
        probe_result = setup_connection_probe(conf)
        if not probe_result[0]:
            return {'ok': False, 'execution_path': 'opencode_server', 'latency_ms': latency,
                    'model': up['model'],
                    'error': _safe_secret_text(probe_result[1], (up['api_key'],)) or '文本模型直调未通过'}
        verified_ids = (probe_result[3] if len(probe_result) > 3 else [up['model']])
        _record_text_model_verification(conf, probe_result[2], verified_ids)
        return {'ok': True, 'execution_path': 'opencode_server', 'latency_ms': latency,
                'model': up['model'], 'reply': '正式生成链路可用'}
    else:   # custom / env:把命令里的占位符换成临时目录,只验证能否跑起来
        raw = os.environ.get('AGENT_CMD') or eng.get('cmd', '')
        if not raw: return {'ok': False, 'error': '还没填自定义命令'}
        tmp = _mk(os.path.join(DATA, '_selftest'))
        rep = lambda s: (s.replace('{tender}', os.path.join(tmp, 'probe.md')).replace('{out}', tmp)
                          .replace('{materials}', assets_dir()).replace('{jobid}', 'selftest'))
        open(os.path.join(tmp, 'probe.md'), 'w', encoding='utf-8').write('# 连通性测试\n')
        cmd = rep(raw) if isinstance(raw, str) else [rep(a) for a in raw]
    cmd = login_shell_wrap(cmd, eng)
    run_env = agent_env(eng)
    conf_err = run_env.pop('BIDDOG_SHELL_CONF_ERR', None)
    if conf_err:
        return {'ok': False, 'error': '%s。这是本机配置问题,不是网关或 Key 的问题;'
                                      '多为数据目录不可写。' % conf_err}
    t0 = time.time()
    try:
        r = _tracked_detached_run(
            cmd, shell=isinstance(cmd, str), capture_output=True, text=True,
            stdin=subprocess.DEVNULL, timeout=int(eng.get('test_timeout') or 120),
            env=run_env, cwd=DATA)
    except subprocess.TimeoutExpired:
        return {'ok': False, 'error': '测试超时(超过 2 分钟没有返回)。多为网关不通或 CLI 在等待登录/授权。'}
    except FileNotFoundError as e:
        return {'ok': False, 'error': '命令启动失败:%s。检查 CLI 路径是否正确。' % e}
    except Exception as e:
        return {'ok': False, 'error': '测试失败:%s' % e}
    out = ((r.stdout or '') + '\n' + (r.stderr or '')).strip()
    ms = int((time.time() - t0) * 1000)
    if r.returncode == 0:
        return {'ok': True, 'latency_ms': ms, 'reply': out[-300:] or '(无输出)'}
    low = out.lower()
    hint = ''
    if 'unknownerror' in low.replace(' ', '') and 'relay' not in low:
        hint = ('（执行外壳报了一个它自己也说不清的错。最常见的原因是它连不上本机中转层——'
                '请确认应用没有被安全软件拦截本机回环连接，然后重试。）')
    if 'thinking' in low and ('not supported' in low or 'unsupported' in low):
        hint = '该模型不支持所选思考等级,请把「思考等级」改成 off。'
    elif any(k in low for k in ('connection refused', 'gateway', 'econnrefused', 'timed out', 'tls', 'certificate')):
        hint = ('连不上本机网关。App 是双击启动的,拿不到你终端里的环境变量——'
                '请保持「用登录 shell 启动」勾选;仍不行就在「附加环境变量」里指定网关地址/配置路径。')
    elif 'no api key found for provider' in low or 'auth-profiles' in low or 'agents add' in low:
        # SoWork 的 openclaw 找不到该 agent 的授权档:原文是英文+本机路径,这里翻成人话并给出唯一省事的出路
        hint = ('这个引擎在本机没有可用的登录/授权(它把授权存在自己的 auth-profiles 里,现在是空的)。'
                '最省事的做法:回到「模型接入」最上面的快速接入卡,粘我们发给你的 Key 点「一键接入并测试」,'
                '切到「自动」引擎——不需要登录任何账号。要继续用它,就先在它自己的客户端里重新登录。')
    elif any(k in low for k in ('unauthorized', 'not logged in', 'login', '401')):
        hint = '未登录或授权失效,请先在对应 App/CLI 里登录后重试。'
    elif any(k in low for k in ('usage limit', 'quota', 'rate limit', '429')):
        hint = '该 CLI 的订阅额度用完了(与「模型接入」的 API Key 是两个独立额度)。'
    elif 'openclaw doctor' in low or ('mismatch' in low and 'config uses' in low):
        hint = ('这是 SoWork 本机安装自身的配置问题(插件配置不一致),与中标狗无关:'
                '终端跑一次 openclaw doctor 按提示修复,或重启/重装 SoWork;'
                '急用的话先把生成引擎切成「S2 模型」(不依赖 SoWork)。')
    if kind == 's2' and RELAY_LAST.get('error'):
        hint = '中转层报告:%s' % RELAY_LAST['error']      # 真实原因往往在上游,别让客户只看到 codex 的退出码
    return {'ok': False, 'error': ('退出码 %s。%s' % (r.returncode, hint)).strip(), 'reply': out[-300:]}

@app.post('/v1/jobs')
async def create_job(request: Request, tender: UploadFile = File(None), materials: UploadFile = File(None),
                     files: List[UploadFile] = File(None), relpaths: str = Form(''),
                     prompt: str = Form(''), name: str = Form(''), mock: str = Form('auto'),
                     start: str = Form('1'), template_id: str = Form(''),
                     project_id: str = Form(''), save_to_assets: str = Form('1'),
                     confirm_parse: str = Form('0')):
    """建任务(向导版约定):
    - tender = 招标文件(主件,永远落任务根目录——绝不进 素材/,素材库污染是内容变薄的根源之一)
    - files + relpaths = 参考素材(多文件/整文件夹,保留目录结构,落 素材/;相对路径做穿越防护)
    - start='0' 只暂存(任务状态=待开始),等 /v1/jobs/{jid}/start 再跑
    兼容旧调用:只传 tender(+materials zip)行为不变。"""
    gate = _generation_gate_response()
    if gate is not None: return gate
    raw_key = str(request.headers.get('idempotency-key') or '')
    idem_key = raw_key[:128] if re.fullmatch(r'[A-Za-z0-9._:-]{8,128}', raw_key) else ''
    create_ledger_path = os.path.join(jobs_dir(), '.create_requests.json')
    if idem_key:
        with _json_lock(create_ledger_path):
            create_ledger = read_json(create_ledger_path, {})
            prior = create_ledger.get(idem_key) if isinstance(create_ledger, dict) else None
            if isinstance(prior, dict) and os.path.isdir(jpath(prior.get('job_id') or '')):
                return {'job_id': prior['job_id'], 'mode': prior.get('mode', 'staged'),
                        'deduplicated': True}
    fl = [f for f in (files or []) if f and f.filename]
    if not (tender and tender.filename) and not fl:
        return JSONResponse({'error': '至少要有一个文件(招标文件)'}, 400)
    try: rels = json.loads(relpaths or '[]')
    except Exception: rels = []
    rels = [str(r or '') for r in rels] if isinstance(rels, list) else []
    while len(rels) < len(fl): rels.append('')
    doc_like = lambda fn: not fn.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tif', '.zip'))
    if not (tender and tender.filename) and fl:
        # 自动推荐前先确定真正主件；否则仅 files 的兼容调用只能按任务名猜场景。
        pick = next((i for i, f in enumerate(fl)
                     if re.search(r'招标|采购|磋商|询价|tender|rfp', os.path.basename(f.filename), re.I)
                     and doc_like(f.filename)), None)
        if pick is None: pick = next((i for i, f in enumerate(fl) if doc_like(f.filename)), 0)
        tender = fl.pop(pick); rels.pop(pick)
    template_snapshot = {}
    template_recommendation = {}
    sample_text = ''
    sample_metadata = {}
    requested_template_id = normalize_template_request(template_id)
    if requested_template_id == 'auto':
        blob = b''
        if tender and tender.filename:
            try:
                blob = await tender.read(MAX_TEMPLATE_UPLOAD_BYTES + 1)
                await tender.seek(0)
                if len(blob) > MAX_TEMPLATE_UPLOAD_BYTES:
                    return JSONResponse({'error': '自动分析暂支持50MB以内的主件；请压缩文件或手动选择模板'}, 413)
                sample_headings, sample_tables, sample_text = extract_document_structure(tender.filename, blob)
                sample_metadata = {'heading_count': len(sample_headings or []),
                                   'table_count': len(sample_tables or [])}
            except Exception:
                try: sample_text = blob.decode('utf-8', errors='ignore')[:120000]
                except Exception: sample_text = ''
        template_recommendation = recommend_template(
            getattr(tender, 'filename', '') or name, sample_text + '\n' + prompt, task_templates())
        template_id = str(template_recommendation.get('template_id') or 'government')
    elif requested_template_id:
        template_id = requested_template_id
        template_recommendation = {'template_id': template_id, 'confidence': 1.0,
                                   'reasons': ['用户明确选择']}
    else:
        template_id = ''
    if template_id:
        selected_template = get_task_template(template_id)
        if not selected_template: return JSONResponse({'error': '所选任务模板不存在'}, 400)
        template_snapshot = task_template_snapshot(selected_template)
        compiled = compile_template_instructions(template_snapshot)
        user_prompt = prompt.strip()
        prompt = compiled + (('\n\n# 用户补充要求\n' + user_prompt) if user_prompt else '')
    jid = datetime.datetime.now().strftime('%m%d-%H%M%S-') + uuid.uuid4().hex[:4]
    job = jpath(jid); os.makedirs(job)
    tname = os.path.basename(tender.filename or '招标文件.pdf')
    open(os.path.join(job, tname), 'wb').write(await tender.read())
    # 模板推荐只使用有限预览，不能把被截断的预览冒充完整解析结果。
    # 流水线会在后台对已保存主件执行一次完整、有预算的本地解析并固化结果。
    mdir = os.path.join(job, '素材')          # 注意:没有素材就不建目录,空 素材/ 会顶掉全局素材库的回落
    uploaded_materials = []   # 本单用户亲手导入的清单;任务详情据此与全局库合并件区分开(反馈 7)
    for i, f in enumerate(fl):
        rel = (rels[i] or f.filename or 'file').replace('\\', '/')
        rel = os.path.normpath(rel).replace(os.sep, '/')
        if rel.startswith(('..', '/')) or os.path.isabs(rel): rel = os.path.basename(f.filename or '') or 'file'
        while rel.startswith('./'): rel = rel[2:]
        parts = [x for x in rel.split('/') if x not in ('', '.', '..')]
        # 只剥一层:显式的 素材/ 前缀剥掉后其余层级原样保留;否则视为"拖入的文件夹名"剥掉顶层。
        # (曾经两条规则叠加把 素材/图片/x.png 剥成 x.png,图片索引全部失配——图不见了就是这来的)
        if parts and parts[0] == '素材': parts = parts[1:]
        elif len(parts) > 1: parts = parts[1:]
        rel = '/'.join(parts)
        rel = rel or os.path.basename(f.filename or 'file')
        dest = os.path.join(mdir, rel)
        os.makedirs(os.path.dirname(dest), exist_ok=True)
        open(dest, 'wb').write(await f.read())
        uploaded_materials.append(rel)
    if materials and materials.filename:
        z = os.path.join(job, '_m.zip'); open(z, 'wb').write(await materials.read())
        try:
            zf = zipfile.ZipFile(z)
            zf.extractall(_mk(mdir))
            uploaded_materials += [n for n in zf.namelist()
                                   if not n.endswith('/') and not os.path.basename(n).startswith('.')][:200]
        except Exception: pass
    if uploaded_materials and (save_to_assets or '1') != '0':
        # 反馈 3 落地:本单导入的材料自动回流素材库(任务导入/<任务名>/),下一单直接复用。
        # md5 与库内既有文件去重;失败不阻塞建单。
        threading.Thread(target=_backflow_job_materials,
                         args=(job, list(uploaded_materials), (name or tname or jid)[:20]),
                         daemon=True).start()
    # 先落 staged，再取得 admission 后由 _launch_job_reserved 原子切为运行；
    # 即使 replay/shutdown 暂时拒绝，也不会留下“看似运行但没有 worker”的孤儿。
    write_json(os.path.join(job, '任务.json'), {'name': name or tname, 'created_at': now(),
               'paused': False, 'staged': True, 'tender': tname, 'prompt': prompt,
               'template_id': template_id, 'template_snapshot': template_snapshot,
               'template_recommendation': template_recommendation,
               'confirm_parse': (confirm_parse or '0') != '0',
               'uploaded_materials': uploaded_materials[:200]})
    write_json(os.path.join(job, 'product.json'),
               {'name': name or tname, 'project_id': str(project_id or '').strip()[:120],
                'version': 1, 'root_job_id': jid, 'parent_job_id': '', 'created_at': now()})
    if idem_key:
        with _json_lock(create_ledger_path):
            create_ledger = read_json(create_ledger_path, {})
            if not isinstance(create_ledger, dict): create_ledger = {}
            create_ledger[idem_key] = {'job_id': jid, 'mode': 'staged', 'created_at': now()}
            write_json(create_ledger_path, dict(list(create_ledger.items())[-64:]))
    if prompt:
        # 落盘成文件,agent 才真的看得到(以前只发了一条聊天消息,界面写着「会作为生成指令的一部分」其实没进指令)
        open(os.path.join(job, '你的要求.md'), 'w', encoding='utf-8').write(prompt)
        # 短要求原样回显;长的(默认预填就有二十来行 markdown)只回一句摘要——
        # 整段灌进对话框既刷屏、又因为用户气泡是纯文本渲染而露出一堆 ** 和 #
        if len(prompt) <= 120:
            emit(job, {'type': 'message', 'role': 'user', 'text': prompt})
        else:
            head = next((l.strip(' #*').strip() for l in prompt.splitlines() if l.strip(' #*').strip()), '')
            emit(job, {'type': 'message', 'role': 'user',
                       'text': '已提交你的要求(%d 字,开头:%s…),完整内容存在任务目录的《你的要求.md》,'
                               'agent 开工先读它。' % (len(prompt), head[:24])})
    if start == '0':
        emit(job, {'type': 'progress', 'stage': '待开始(素材已就位,点「开始生成」)', 'pct': 0, 'step': 0, 'total': 12})
        return {'job_id': jid, 'mode': 'staged'}
    result = _launch_http_result(_launch_job(jid, job, mock))
    if idem_key and isinstance(result, dict):
        with _json_lock(create_ledger_path):
            create_ledger = read_json(create_ledger_path, {})
            if isinstance(create_ledger.get(idem_key), dict):
                create_ledger[idem_key]['mode'] = str(result.get('mode') or 'staged')
                write_json(create_ledger_path, create_ledger)
    return result

def _launch_blocked(jid, job, reason):
    """全局/控制态禁入时保留一个可见的待开始任务，绝不假报 running。"""
    messages = {
        'replay': '正在恢复运行日志证据，请稍后在任务列表点“开始生成”重试。',
        'control': '任务正在停止或删除，请稍后重试。',
        'shutdown': '应用正在安全退出，请重新打开后在任务列表点“开始生成”。',
        'exiting': '引擎正在退出，请等待应用重新连接后再开始生成。',
    }
    meta_path = os.path.join(job, '任务.json')
    meta = read_json(meta_path, {})
    # 只有本来就是待开始的新建/重跑任务才写 staged；历史/活跃任务绝不能被拒绝路径改状态。
    if meta.get('staged'):
        emit(job, {'type': 'progress', 'stage': '待开始（%s）' % messages.get(reason, '当前暂不能启动'),
                   'pct': 0, 'step': 0, 'total': 12})
    return {'ok': False, 'job_id': jid, 'mode': 'staged', 'blocked': reason,
            'retryable': True, 'error': messages.get(reason, '当前暂不能启动，请稍后重试。'),
            '_status': 409 if reason == 'control' else 503}

def _launch_http_result(result, include_ok=False):
    """把内部 admission 结果转换成 HTTP；`_status` 永不泄漏给前端。"""
    if isinstance(result, dict) and result.get('_status'):
        body = dict(result); status = int(body.pop('_status'))
        return JSONResponse(body, status)
    if include_ok and isinstance(result, dict):
        return {'ok': result.get('mode') != 'error', **result}
    return result

def _admission_failure(reason, running_error):
    errors = {
        'running': running_error,
        'control': '任务正在停止或删除，请稍后重试。',
        'replay': '正在恢复运行日志证据，请稍后重试。',
        'shutdown': '应用正在安全退出，请重新打开后重试。',
        'exiting': '引擎正在退出，请等待应用重新连接后重试。',
    }
    return JSONResponse({'ok': False, 'retryable': True,
                         'error': errors.get(reason, '当前暂不能启动，请稍后重试。'),
                         'blocked': reason},
                        409 if reason in ('running', 'control') else 503)

def _cancelled_before_dispatch():
    """控制请求在预检/准备期间取得了本轮取消权：没有 worker 可再启动。"""
    return {'ok': False, 'mode': 'cancelled', 'stopped': True, 'retryable': False,
            'error': '启动准备期间已收到停止请求，本轮没有派发生成。', '_status': 409}

def _launch_job(jid, job, mock='auto'):
    base = os.path.basename(jid)
    owner, reason = _reserve_running_reason(base)
    if not owner:
        if reason == 'running':
            return {'job_id': jid, 'mode': 'running', 'already_running': True}
        return _launch_blocked(jid, job, reason)
    try:
        result = _launch_job_reserved(jid, job, mock, owner)
    except Exception:
        _release_running(base, owner)
        raise
    if result.get('mode') in ('error', 'cancelled'): _release_running(base, owner)
    return result

def _launch_job_reserved(jid, job, mock, owner):
    """按当前引擎配置启动生成。主件路径只信 任务.json 的 tender 字段(建任务时就定死,不做猜测)。"""
    meta = read_json(os.path.join(job, '任务.json'), {})
    if _cancel_requested(os.path.basename(jid), owner):
        return _cancelled_before_dispatch()
    tpath = os.path.join(job, os.path.basename(meta.get('tender') or ''))
    if not os.path.isfile(tpath):
        emit(job, {'type': 'error', 'text': '找不到招标文件「%s」,请删除本任务重新创建。' % meta.get('tender')})
        return {'job_id': jid, 'mode': 'error'}
    close_pending_questions(job, meta)
    # launch 是全新证据轮次：轮换 run_id，并丢弃上一轮会话/证据引用。
    # resume 不经过这里，因此会保留原 run_id 与会话。
    meta['run_id'] = _new_run_id()
    meta['staged'] = False
    meta.pop('skill_manifest', None)
    meta.pop('oc_session', None)
    meta.pop('oc_questions', None)
    # 新开/重跑会重新结算；不能沿用上一次 stopped/done 终态。
    try: os.unlink(os.path.join(job, 'outcome.json'))
    except FileNotFoundError: pass
    conf_now = read_json(conf_path(), {})
    eng_now = conf_now.get('engine') or {}
    up_now = s2_conf(conf_now)
    meta['engine_snapshot'] = {'kind': eng_now.get('kind', 's2'), 'model': up_now['model'],
                               'base_url': up_now['base_url'], 'wire': up_now['wire'],
                               'verify_ssl': bool(up_now['verify_ssl']),
                               'runtime_fingerprint': oc_config_fingerprint(conf_now),
                               'started_at': now()}
    write_json(os.path.join(job, '任务.json'), meta)
    # 环境准备不是业务第 1 步。先给可见的 step=0 预检，真正建立连接后再进入读取招标文件。
    emit(job, {'type': 'progress', 'stage': '正在检查生成环境', 'pct': 0, 'step': 0, 'total': 12})
    preflight = generation_preflight(job, conf_now)
    emit(job, {'type': 'worklog', 'lines': ['保存任务文件完成', '识别场景模板完成',
              '正在检查生成组件' if preflight.get('repair') else '生成环境检查完成']})
    connection_problem = next((item for item in preflight.get('checks') or []
                               if item.get('id') == 'connection' and item.get('status') == 'fail'), None)
    if connection_problem:
        emit(job, {'type': 'error', 'text': connection_problem.get('message'),
                   'actions': [{'act': 'open_engine', 'label': '重新测试模型连接'},
                               {'act': 'diagnose', 'label': '一键诊断'}]})
        halt(job, '已停止（模型连接未验证）')
        return {'job_id': jid, 'mode': 'error', 'error': connection_problem.get('message')}
    if _cancel_requested(os.path.basename(jid), owner):
        return _cancelled_before_dispatch()
    agent_cmd = config_agent_cmd()
    use_mock = (mock == '1') or (mock == 'auto' and not agent_cmd)
    if use_mock:
        if _cancel_requested(os.path.basename(jid), owner):
            return _cancelled_before_dispatch()
        update_runtime(job, execution_path='demo', can_pause=False,
                       pause_disabled_reason='内置演示流程不支持暂停；可以等待完成或使用停止。')
        conf0 = read_json(conf_path(), {})
        if (conf0.get('engine') or {}).get('kind', 's2') in ('s2', 'opencode') and not s2_conf(conf0)['api_key']:
            emit(job, {'type': 'message', 'role': 'agent',
                       'text': '当前还没填 API Key,先用**内置演示流程**把全流程跑给你看(产出为样例稿)。'
                               '到「设置 · 模型接入」把 Key 填上再重跑本任务,产出的就是真实标书。',
                       'actions': [{'act': 'open_engine', 'label': '去填 Key'}]})
        _start_reserved_worker(os.path.basename(jid), owner, mock_agent, job)
    else:
        mat = merged_materials(job)     # 这一处最要紧:它决定 agent 命令行里的 materialsDir
        sd_path = skill_dir_conf(read_json(conf_path(), {}))
        if not os.path.isfile(os.path.join(sd_path, 'SKILL.md')): sd_path = ensure_skill()
        sub = lambda x: (x.replace('{tender}', tpath).replace('{out}', job)
                          .replace('{materials}', mat).replace('{jobid}', jid).replace('{skill}', sd_path))
        cmd = [sub(a) for a in agent_cmd] if isinstance(agent_cmd, list) else sub(agent_cmd)
        cmd = login_shell_wrap(cmd, (read_json(conf_path(), {}).get('engine') or {}))
        # OpenCode 引擎优先走 server 模式(看得见/停得掉/问得着/崩了能续);
        # 起不来或链路没通就自动回落 CLI 模式,行为跟以前一模一样 —— 不能因为新路子没通就跑不了单。
        kind_now = ((read_json(conf_path(), {}).get('engine') or {}).get('kind') or 's2')
        oc_prompt = sub(next((a for a in (agent_cmd if isinstance(agent_cmd, list) else [])
                              if 'SKILL.md' in str(a) or len(str(a)) > 300), ''))
        use_oc_server = bool(kind_now in ('s2', 'opencode') and oc_prompt
                             and not os.environ.get('BID_NO_OC_SERVER'))
        use_pipeline = bool(use_oc_server and not os.environ.get('BID_NO_CHECKPOINT_PIPELINE'))
        dispatch = oc_prompt if use_oc_server else cmd
        _set_skill_manifest(job, sd_path, dispatch,
                            _dispatch_contains_skill(dispatch, sd_path),
                            'opencode' if use_oc_server else 'cli')
        if use_oc_server:
            if _cancel_requested(os.path.basename(jid), owner):
                return _cancelled_before_dispatch()
            update_runtime(job, execution_path='opencode_starting', can_pause=False,
                           pause_disabled_reason='正在建立稳定连接，暂时不能暂停。')
            if use_pipeline:
                # `_set_skill_manifest` has just persisted a run-bound dispatch receipt.
                # Reusing the earlier in-memory metadata here would overwrite that receipt
                # and make a genuinely used Skill appear missing in diagnostics.
                pipeline_meta = read_json(os.path.join(job, '任务.json'), {})
                _initialize_generation_pipeline(job, pipeline_meta, conf_now)
                _start_reserved_worker(os.path.basename(jid), owner, generation_pipeline_worker, job)
            else:
                _start_reserved_worker(os.path.basename(jid), owner, agent_via_server_or_cli,
                                       job, oc_prompt, cmd)
        else:
            if _cancel_requested(os.path.basename(jid), owner):
                return _cancelled_before_dispatch()
            emit(job, {'type': 'progress', 'stage': '已就绪，正在读招标文件', 'pct': 2, 'step': 1, 'total': 12})
            update_runtime(job, execution_path='cli_compat', can_pause=False,
                           pause_disabled_reason='当前稳定运行方式暂不支持暂停；如需中止，请使用停止。')
            _start_reserved_worker(os.path.basename(jid), owner, real_agent, job, cmd)
    return {'job_id': jid, 'mode': 'mock' if use_mock else 'agent'}

def finish_job(job):
    """server 模式跑完后的收尾。刻意跟 CLI 模式(real_agent 里那段)保持同一套动作:
    收拢子目录产物 → 缺 Word 就零 token 补出来 → 完成播报 → 技能包证据 → 成品质检。
    两条路的收尾必须一致,否则会出现「换了条路跑,完成播报就不一样」这种说不清的差异。"""
    base = os.path.basename(job)
    if _cancel_requested(base): return
    return settle(job)
    harvest(job)
    known = set(list_deliverables(job))
    for fn in known: emit(job, {'type': 'artifact', 'name': fn})
    try: ensure_docx(job, known)
    except Exception as e:
        emit(job, {'type': 'message', 'role': 'agent', 'text': '⚠ 补出 Word 未能执行:%s' % e})
    if not known:
        emit(job, {'type': 'error',
                   'text': '生成过程结束了,但**没有生成任何可交付文件**。'
                           '请看运行日志确认它停在哪一步。',
                   'actions': [{'act': 'open_log', 'label': '查看运行日志'},
                               {'act': 'rerun', 'label': '重跑本任务'}]})
        halt(job, '已停止(没有产出)')
        return
    # 终态只能由 settle() 写入；以下旧收尾保留为迁移注释路径且不可达。
    summary, actions = artifact_summary(job, known)
    used = skill_evidence(job)
    if used['state'] == 'unverifiable':
        summary += ('\n\n⚠ **技能包运行证据暂时无法核验**(%s)请人工复核响应矩阵与格式门禁。' % used['why'])
        actions = (actions or []) + [{'act': 'open_log', 'label': '查看运行日志'}]
    elif used['state'] == 'missing':
        summary += ('\n\n⚠ **写作规则没有完整载入**(%s)请在设置中重新测试连接后重跑。' % used['why'])
        actions = (actions or []) + [{'act': 'open_log', 'label': '查看运行日志'}]
    emit(job, {'type': 'message', 'role': 'agent', 'text': summary, 'actions': actions})
    emit(job, {'type': 'skill_used', 'state': used['state'], 'ok': used['ok'],
               'hits': used['hits'], 'why': used['why']})
    try: quality_audit(job, known)
    except Exception as e:
        emit(job, {'type': 'message', 'role': 'agent', 'text': '⚠ 成品质检未能执行:%s' % e})

def agent_via_server_or_cli(job, prompt, cmd):
    """先试 server 模式,不成回落 CLI。两条路的收尾(收拢产物/补出 Word/质检/播报)必须一致 ——
    否则会出现「换了条路跑,完成播报就不一样」这种说不清的差异。"""
    base = os.path.basename(job)
    result = OC_RUN_INTERRUPTED
    eng = read_json(conf_path(), {}).get('engine') or {}
    ok, why = ensure_default_shell(job, eng)
    if not ok:
        emit(job, {'type': 'error',
                   'text': '生成组件自动修复未完成，任务已安全停下。请检查网络后重试。',
                   'actions': [{'act': 'diagnose', 'label': '一键诊断'},
                               {'act': 'open_engine', 'label': '检查生成设置'}]})
        settle(job, stop_reason='已停止（生成组件修复失败）')
        return
    if isinstance(cmd, list) and cmd and os.path.basename(str(cmd[0])).lower() in ('opencode', 'opencode.exe'):
        cmd[0] = resolve_cli('opencode', eng) or cmd[0]
    emit(job, {'type': 'progress', 'stage': '已就绪，正在读招标文件', 'pct': 2, 'step': 1, 'total': 12})
    emit(job, {'type': 'worklog', 'lines': ['生成组件已就绪', '正在建立模型连接', '开始读取招标文件']})
    try:
        result = oc_run(job, prompt)
    except Exception as e:
        append_diagnostic(job, 'generation_worker_exception', str(e), level='error')
        emit(job, {'type': 'error',
                   'text': '连接意外中断，任务已安全停下；已生成的内容都已保留。',
                   'actions': [{'act': 'resume', 'label': '从已保存内容继续'},
                               {'act': 'open_log', 'label': '查看诊断详情'}]})
    if result == OC_RUN_FALLBACK:
        if _cancel_requested(base): return
        real_agent(job, cmd)      # 回落仍沿用外层同一个 reservation，直到全部收尾完成
        return
    if result == OC_RUN_CANCELLED:
        return
    if result == OC_RUN_COMPLETED:
        finish_job(job)
        return
    latest = _last_json_event(job)
    if '已安全停下' not in str(latest.get('text') or ''):
        emit(job, {'type': 'error',
                   'text': '连接意外中断，任务已安全停下；已生成的内容都已保留。',
                   'actions': [{'act': 'resume', 'label': '从已保存内容继续'},
                               {'act': 'open_log', 'label': '查看诊断详情'}]})
    settle(job, stop_reason='已停止（连接中断，内容已保留）')

@app.post('/v1/jobs/{jid}/start')
def start_job(jid: str, mock: str = 'auto'):
    """启动「待开始」的暂存任务(向导里点「稍后开始」建出来的)"""
    gate = _generation_gate_response()
    if gate is not None: return gate
    job = jpath(jid)
    if not os.path.isdir(job): return JSONResponse({'ok': False, 'error': '任务不存在'}, 404)
    r = _launch_job(jid, job, mock)
    return _launch_http_result(r, include_ok=True)

def _interrupt_or_finished(sid):
    """中断 OpenCode 会话；已进入终态也算安全停住，网络/服务失败则返回 False。"""
    if not sid or not OC.get('base'): return True
    if oc_interrupt(sid): return True
    try:
        done, _why = oc_turn(sid)
        return bool(done)
    except Exception:
        return False

def _stop_running_owner(job, base, owner):
    """停止指定 owner 的 server/CLI/mock；控制 tombstone 必须由调用方持有。"""
    if not owner: return True, False
    requested = _request_cancel(base, owner)
    sid = (read_json(os.path.join(job, '任务.json'), {}) or {}).get('oc_session')
    server_active = bool(sid and OC.get('base'))
    interrupted = _interrupt_or_finished(sid) if server_active else True
    pipeline_interrupted = _pipeline_interrupt_sessions(job)
    proc = _take_proc(base, owner)
    # 有匹配 CLI 说明当前已走 fallback；旧 sid 中断失败不影响杀掉真正执行体。
    if ((server_active and not interrupted) or not pipeline_interrupted) and not proc:
        # 取消信号一旦发布就必须单调保留：某个 worker 可能已经观察到它并停下，
        # 此时“回滚取消”只会制造一个无法继续的半停状态。
        return False, requested
    if proc: kill_tree(proc)
    return True, requested

def _wait_owner_exit(base, owner, timeout=8):
    deadline = time.monotonic() + max(0.1, float(timeout))
    while _owner_running(base, owner) and time.monotonic() < deadline:
        time.sleep(0.05)
    return not _owner_running(base, owner)

@app.delete('/v1/jobs/{jid}')
def del_job(jid: str):
    job = jpath(jid)                      # jpath 已做 basename,防目录穿越
    if not os.path.isdir(job): return {'ok': True}   # 已经没了=删除成功,幂等
    base = os.path.basename(jid)
    control, owner = _begin_job_control(base)
    if not control:
        return JSONResponse({'ok': False, 'error': '任务正在执行另一个停止或删除操作，请稍后重试'}, 409)
    try:
        ok, _requested = _stop_running_owner(job, base, owner)
        if not ok:
            return JSONResponse({'ok': False, 'error': '任务暂时没有完全停止，因此没有删除；请稍后重试'}, 502)
        if owner and not _wait_owner_exit(base, owner):
            return JSONResponse({'ok': False, 'error': '生成进程尚未完全退出，任务未删除；请稍后重试'}, 409)
        for _ in range(3):                # 文件句柄释放仍可能晚几十毫秒，小重试后必须验真
            shutil.rmtree(job, ignore_errors=True)
            if not os.path.isdir(job): break
            time.sleep(0.4)
        if os.path.isdir(job):
            return JSONResponse({'ok': False, 'error': '任务目录仍被占用，删除没有完成；请稍后重试'}, 500)
        return {'ok': True}
    finally:
        _end_job_control(base, control)

def job_state(job, meta=None, prog=None):
    """任务当前处于哪个状态 —— 由引擎给出结论,不让前端从进度百分比反推。

    以前前端只有 `done = pct >= 100` 一个判据,于是「已停止(手动)」「已停止(没有产出)」
    「已停止(agent 异常退出)」「已停止(未能启动生成)」这四种终态(它们的 pct 都是 0)
    在界面上全部等于「运行中 0%」:一直转圈、一直挂在「进行中」分组里。
    用户会一直等下去——等到投标截止。"""
    meta = read_json(os.path.join(job, '任务.json'), {}) if meta is None else meta
    prog = read_json(os.path.join(job, 'progress.json'), {}) if prog is None else prog
    if meta.get('staged'): return 'staged'                       # 暂存,还没开跑
    if _is_running(os.path.basename(job)): return 'running'
    outcome = read_json(os.path.join(job, 'outcome.json'), {})
    pipeline_state = generation_pipeline.load(job)
    has_pipeline = pipeline_state.get('version') == generation_pipeline.PIPELINE_VERSION
    delivery_done = (not has_pipeline or (pipeline_state.get('state') == 'done' and any(
        node.get('id') == 'delivery_gate' and node.get('state') == 'done'
        for node in pipeline_state.get('nodes') or [])))
    if outcome.get('state') == 'done':
        return 'done' if _body_docxs(job) and delivery_done else 'stopped'
    if outcome.get('state') == 'stopped': return 'stopped'
    stage = str(prog.get('stage') or '')
    # 老版本遗留任务的迁移兜底：100% 仍必须有正文 Word；没有就按中断展示。
    if int(prog.get('pct') or 0) >= 100:
        return 'done' if _body_docxs(job) and delivery_done else 'stopped'
    # 暂停要跟「停止」分开:停止是这一单结束了,暂停是它还在半路、等着被接着做
    if meta.get('paused') and meta.get('oc_session'): return 'paused'
    if stage.startswith('已停止'): return 'stopped'
    if not prog: return 'staged'                                 # 从没跑过
    return 'unknown'      # 进程没了、进度也没到头:多半是引擎被杀或断电,单独标出来别装作在跑

# 每种状态下前端允许做什么。放在引擎侧,免得前端各处自己拼条件、拼错了就是死按钮
STATE_CAN = {'staged': ['start', 'delete'], 'running': ['stop', 'ask', 'delete'],
             'done': ['redo', 'rerun', 'ask', 'export', 'delete'],
             'paused': ['resume', 'stop', 'rerun', 'redo', 'ask', 'export', 'delete'],
             # 停止过、但留下了会话的,也能接着做 —— 由 /v1/jobs 逐单判断(见 list_jobs)
             'stopped': ['rerun', 'redo', 'ask', 'export', 'delete'],
             'unknown': ['rerun', 'ask', 'export', 'delete']}

def job_can(job, state, meta):
    """这一单此刻能做什么。stopped 且留下了会话的,额外给「继续做」——
    没有会话就不给:那颗按钮点下去只会失败,不如不出现。"""
    can = list(STATE_CAN.get(state, []))
    runtime = read_json(os.path.join(job, 'runtime.json'), {})
    can_pause = runtime.get('can_pause')
    if can_pause is None: can_pause = bool(meta.get('oc_session'))
    if state == 'running' and can_pause: can.insert(0, 'pause')
    pipeline = generation_pipeline.summary(job)
    pipeline_commit_pending = (pipeline.get('state') == 'done'
                               and read_json(os.path.join(job, 'outcome.json'), {}).get('state') != 'done')
    if (state in ('stopped', 'unknown') and pipeline.get('version') == generation_pipeline.PIPELINE_VERSION
            and (pipeline.get('recoverable') or pipeline_commit_pending)):
        can.insert(0, 'resume')
    # 引擎/应用异常退出时不会来得及写 stopped outcome，但已落盘的
    # OpenCode 会话仍是可恢复检查点。unknown 不能因为前端少一个按钮被迫重跑。
    if state in ('stopped', 'unknown') and meta.get('oc_session'): can.insert(0, 'resume')
    return list(dict.fromkeys(can))

def job_runtime(job, meta=None):
    meta = meta if isinstance(meta, dict) else read_json(os.path.join(job, '任务.json'), {})
    runtime = read_json(os.path.join(job, 'runtime.json'), {})
    path = str(runtime.get('execution_path') or ('opencode_server' if meta.get('oc_session') else ''))
    can_pause = runtime.get('can_pause')
    if can_pause is None: can_pause = bool(meta.get('oc_session'))
    reason = str(runtime.get('pause_disabled_reason') or '')
    if not can_pause and not reason:
        reason = '当前运行方式暂不支持暂停；如需中止，请使用停止。'
    if re.search(r'执行外壳|OpenCode|\bCLI\b|探活|兼容', reason, re.I):
        reason = '当前稳定运行方式暂不支持暂停；如需中止，请使用停止。'
    return {'mode': 'compatibility' if path in ('cli', 'cli_compat') else 'managed',
            'execution_path': path,
            'capabilities': {'pause': {'enabled': bool(can_pause), 'reason': reason}}}

PRESENTATION_STATES = {
    'preparing': '准备中',
    'generating': '生成中',
    'needs_input': '需要你确认',
    'completed': '已完成',
    'incomplete': '未完成',
}

def product_meta(job, task_meta=None):
    """Organization/version metadata lives apart from runtime-owned 任务.json.

    Workers frequently rewrite 任务.json from an earlier snapshot. Keeping these
    user-managed fields in a separate atomically patched document prevents an
    archive/project change from disappearing while a task is running.
    """
    task_meta = task_meta if isinstance(task_meta, dict) else read_json(os.path.join(job, '任务.json'), {})
    raw = read_json(os.path.join(job, 'product.json'), {})
    if not isinstance(raw, dict): raw = {}
    return {'name': str(raw.get('name') or task_meta.get('name') or os.path.basename(job)),
            'project_id': str(raw.get('project_id') or ''),
            'archived_at': str(raw.get('archived_at') or ''),
            'version': max(1, int(raw.get('version') or 1)),
            'parent_job_id': str(raw.get('parent_job_id') or ''),
            'root_job_id': str(raw.get('root_job_id') or os.path.basename(job))}

def update_product_meta(job, **changes):
    changes['updated_at'] = now()
    return patch_json(os.path.join(job, 'product.json'), changes)

DEFAULT_TASK_TEMPLATES = builtin_templates()

def templates_path(): return os.path.join(_mk(ws_root()), 'task_templates.json')

def task_templates():
    saved = read_json(templates_path(), {})
    items = saved.get('items') if isinstance(saved, dict) else []
    items = items if isinstance(items, list) else []
    overrides = {str(item.get('id')): item for item in items if isinstance(item, dict) and item.get('id')}
    out = []
    for default in DEFAULT_TASK_TEMPLATES:
        merged = dict(default); merged.update(overrides.pop(default['id'], {})); merged['builtin'] = True
        out.append(merged)
    out.extend(sorted((dict(item, builtin=False) for item in overrides.values()),
                      key=lambda item: str(item.get('created_at') or item.get('name') or '')))
    return out

def get_task_template(template_id):
    wanted = str(template_id or '')
    return next((item for item in task_templates() if item.get('id') == wanted), None)

def normalize_template_request(value):
    """Normalize legacy/default clients onto the one automatic selection path."""
    requested = str(value or '').strip()
    return 'auto' if requested.lower() in ('', 'auto', 'default') else requested

def task_template_snapshot(item):
    if not isinstance(item, dict): return {}
    safe = {key: item.get(key) for key in ('id', 'name', 'description', 'prompt', 'settings', 'package')}
    safe['settings'] = safe.get('settings') if isinstance(safe.get('settings'), dict) else {}
    safe['package'] = normalize_package(safe.get('package'))
    # JSON round-trip gives each task an immutable deep copy of nested settings.
    return json.loads(json.dumps(safe, ensure_ascii=False))

def _save_task_template(item):
    path = templates_path()
    with _json_lock(path):
        current = read_json(path, {})
        items = current.get('items') if isinstance(current, dict) else []
        items = [row for row in (items or []) if isinstance(row, dict) and row.get('id') != item['id']]
        items.append(item)
        write_json(path, {'items': items, 'updated_at': now()})
    return item

@app.get('/v1/templates')
def list_task_templates(): return task_templates()

@app.post('/v1/templates')
async def create_task_template(req: Request):
    body = await req.json()
    name = str(body.get('name') or '').strip()
    prompt = str(body.get('prompt') or '').strip()
    if not name or not prompt:
        return JSONResponse({'ok': False, 'error': '模板名称和生成要求不能为空'}, 400)
    base = get_task_template(str(body.get('base_template_id') or ''))
    package = normalize_package(body.get('package') if isinstance(body.get('package'), dict)
                                else ((base or {}).get('package') or {}))
    generated_from = package.get('generated_from') or {}
    if generated_from.get('kind') == 'uploaded_bid' and not generated_from.get('source_structure_ready'):
        return JSONResponse({'ok': False, 'error': '上传标书提取到的结构不足，请补足目录后再保存'}, 400)
    validation = validate_package(package)
    item = {'id': 'tpl-' + uuid.uuid4().hex[:10], 'name': name[:80],
            'description': str(body.get('description') or '').strip()[:300],
            'prompt': prompt[:20000],
            'settings': body.get('settings') if isinstance(body.get('settings'), dict) else {},
            'package': package, 'validation': validation,
            'builtin': False, 'created_at': now(), 'updated_at': now()}
    return _save_task_template(item)

@app.post('/v1/templates/recommend')
async def recommend_task_template(file: UploadFile = File(...), scene_hint: str = Form('')):
    blob = await file.read(MAX_TEMPLATE_UPLOAD_BYTES + 1)
    if len(blob) > MAX_TEMPLATE_UPLOAD_BYTES:
        return JSONResponse({'ok': False, 'error': '文件超过50MB，请先压缩或拆分'}, 413)
    try: _, _, text = extract_document_structure(file.filename or '招标文件', blob)
    except Exception: text = blob.decode('utf-8', errors='ignore')[:120000]
    result = recommend_template(file.filename or '招标文件', text + '\n' + scene_hint[:20000], task_templates())
    selected = get_task_template(result.get('template_id')) or {}
    result['template'] = task_template_snapshot(selected)
    return result

@app.post('/v1/templates/derive')
async def derive_task_template(file: UploadFile = File(...), name: str = Form(''), scene_hint: str = Form('')):
    blob = await file.read(MAX_TEMPLATE_UPLOAD_BYTES + 1)
    if len(blob) > MAX_TEMPLATE_UPLOAD_BYTES:
        return JSONResponse({'ok': False, 'error': '文件超过50MB，请先压缩或拆分'}, 413)
    try:
        return derive_template(name, file.filename or '优秀历史标书.docx', blob, task_templates(), scene_hint)
    except ValueError as exc:
        return JSONResponse({'ok': False, 'error': str(exc)}, 400)
    except Exception as exc:
        return JSONResponse({'ok': False, 'error': '模板提炼失败：%s' % str(exc)}, 400)

@app.put('/v1/templates/{template_id}')
async def update_task_template(template_id: str, req: Request):
    current = get_task_template(template_id)
    if not current: return JSONResponse({'ok': False, 'error': '模板不存在'}, 404)
    body = await req.json(); item = dict(current)
    for key, limit in (('name', 80), ('description', 300), ('prompt', 20000)):
        if key in body: item[key] = str(body.get(key) or '').strip()[:limit]
    if 'settings' in body:
        if not isinstance(body.get('settings'), dict):
            return JSONResponse({'ok': False, 'error': '模板默认设置必须是对象'}, 400)
        item['settings'] = body['settings']
    if 'package' in body:
        if not isinstance(body.get('package'), dict):
            return JSONResponse({'ok': False, 'error': '模板包必须是对象'}, 400)
        item['package'] = normalize_package(body['package'])
        item['validation'] = validate_package(item['package'])
    if not item.get('name') or not item.get('prompt'):
        return JSONResponse({'ok': False, 'error': '模板名称和生成要求不能为空'}, 400)
    item['updated_at'] = now()
    return _save_task_template(item)

@app.delete('/v1/templates/{template_id}')
def delete_task_template(template_id: str):
    current = get_task_template(template_id)
    if not current: return {'ok': True}
    if current.get('builtin'):
        return JSONResponse({'ok': False, 'error': '内置模板不能删除，可以直接修改内容'}, 400)
    path = templates_path()
    with _json_lock(path):
        saved = read_json(path, {})
        items = [item for item in (saved.get('items') or [])
                 if isinstance(item, dict) and item.get('id') != template_id]
        write_json(path, {'items': items, 'updated_at': now()})
    return {'ok': True}

def _pending_question_ids(job, meta=None):
    meta = meta if isinstance(meta, dict) else read_json(os.path.join(job, '任务.json'), {})
    pending = set((meta.get('oc_questions') or {}).keys()) if isinstance(meta.get('oc_questions'), dict) else set()
    path = os.path.join(job, 'events.jsonl')
    try:
        lines = open(path, encoding='utf-8', errors='ignore').read().splitlines()[-2000:]
    except OSError:
        lines = []
    for line in lines:
        try: event = json.loads(line)
        except (TypeError, ValueError): continue
        rid = str(event.get('id') or '')
        if event.get('type') == 'question' and rid: pending.add(rid)
        elif event.get('type') == 'question_closed' and rid: pending.discard(rid)
    return pending

def _pending_question_count(job, meta=None):
    return len(_pending_question_ids(job, meta))

def close_pending_questions(job, meta=None):
    """新的 run_id 开始前关闭上一轮未回答问题，防止重放旧事件时卡回旧会话。"""
    pending = sorted(_pending_question_ids(job, meta))
    for rid in pending:
        emit(job, {'type': 'question_closed', 'id': rid, 'reason': 'new_generation'})
    if pending:
        patch_json(os.path.join(job, '任务.json'), remove=('oc_questions',))
    return pending

def _friendly_current_action(stage, fallback='正在生成投标文件'):
    """Turn internal execution stages into calm, user-facing progress copy."""
    text = str(stage or '').strip()
    if not text:
        return fallback
    if re.search(r'执行外壳|OpenCode|\bCLI\b|探活|兼容(?:模式)?', text, re.I):
        if re.search(r'连接|启动|服务|探活|等待|响应|OpenCode|执行外壳|\bCLI\b', text, re.I):
            return '正在建立稳定连接'
        return fallback
    if '已派发' in text and '招标文件' in text:
        return '正在读取招标文件'
    return text

def job_presentation(job, state, meta, prog, delivery=None):
    pending = _pending_question_count(job, meta)
    if meta.get('paused') or pending:
        code = 'needs_input'
        action = ('任务已暂停，确认后可以继续生成' if meta.get('paused')
                  else '有 %d 项内容等待你确认' % pending)
    elif state == 'staged':
        code = 'preparing'
        action = _friendly_current_action(prog.get('stage'), '材料已就位，等待开始生成')
    elif state == 'running':
        code = 'generating'
        action = _friendly_current_action(prog.get('stage'), '正在生成投标文件')
    elif state == 'done':
        delivery = delivery if isinstance(delivery, dict) else delivery_summary(job)
        if delivery.get('ready'):
            code = 'completed'
            action = '交付文件已经生成，可以打开检查和下载'
        else:
            code = 'needs_input'
            action = str((delivery.get('checks') or {}).get('summary') or '交付文件已生成，还有检查项需要你确认')
    else:
        code = 'incomplete'
        action = ('任务意外中断，已生成的内容仍然保留' if state == 'unknown'
                  else _friendly_current_action(
                      (read_json(os.path.join(job, 'outcome.json'), {}) or {}).get('reason')
                      or prog.get('stage'), '任务未完成，可以查看原因后继续处理'))
    return {'code': code, 'label': PRESENTATION_STATES[code]}, action

FLOW_PHASES = [
    {'id': 'environment', 'label': '环境准备', 'first': 0, 'last': 0,
     'evidence': 'preflight.json', 'pending': '等待检查本地生成环境'},
    {'id': 'parse', 'label': '招标解析', 'first': 1, 'last': 4,
     'evidence': '招标文件、图片索引、组成与格式规范', 'pending': '等待读取招标文件'},
    {'id': 'plan', 'label': '响应规划', 'first': 5, 'last': 6,
     'evidence': '评分废标索引与响应矩阵', 'pending': '等待规划响应结构'},
    {'id': 'write', 'label': '并行撰写', 'first': 7, 'last': 8,
     'evidence': '章节稿与逐条偏离表', 'pending': '等待分章撰写'},
    {'id': 'assemble', 'label': 'Word 装配', 'first': 9, 'last': 10,
     'evidence': '正文与配图复核记录', 'pending': '等待汇总和装配'},
    {'id': 'deliver', 'label': '交付质检', 'first': 11, 'last': 12,
     'evidence': '自检报告与最终 Word', 'pending': '等待交付检查'},
]

PHASE_ENV_EXPECTED_S = 60

def _progress_timeline(job):
    """Return the first persisted timestamp for each observed progress event."""
    rows = []
    try:
        source = open(os.path.join(job, 'events.jsonl'), encoding='utf-8', errors='ignore')
    except OSError:
        return rows
    with source:
        for line in source:
            try: event = json.loads(line)
            except (TypeError, ValueError): continue
            if event.get('type') != 'progress': continue
            try: step = max(0, min(12, int(event.get('step') or 0)))
            except (TypeError, ValueError): continue
            stamp = _parse_ts(event.get('ts'))
            if stamp: rows.append((step, stamp))
    return sorted(rows, key=lambda item: item[1])

def _phase_expected(definition, stage_rows):
    if definition['id'] == 'environment': return PHASE_ENV_EXPECTED_S, 'reference', 0
    selected = [row for row in stage_rows
                if definition['first'] <= int(row.get('step') or 0) <= definition['last']]
    expected = int(sum(float(row.get('avg_s') or 0) for row in selected))
    history_count = sum(1 for row in selected if row.get('from_history'))
    source = 'history' if selected and history_count == len(selected) else ('mixed' if history_count else 'reference')
    return expected, source, history_count

def _flow_check_state(status):
    status = str(status or '').lower()
    if status in ('pass', 'demo', 'skipped'): return 'done'
    if status in ('pending', 'repairing', 'running'): return 'active'
    if status in ('warning', 'attention'): return 'attention'
    if status in ('fail', 'error'): return 'failed'
    return 'pending'

def job_flow(job, state=None, meta=None, prog=None, outcome=None):
    """Build the compact six-phase console exclusively from persisted job facts."""
    meta = meta if isinstance(meta, dict) else read_json(os.path.join(job, '任务.json'), {})
    prog = prog if isinstance(prog, dict) else read_json(os.path.join(job, 'progress.json'), {})
    if prog.get('type') == 'progress': prog = sanitize_event(job, prog)
    outcome = outcome if isinstance(outcome, dict) else read_json(os.path.join(job, 'outcome.json'), {})
    state = state or job_state(job, meta, prog)
    try: step = max(0, min(12, int(prog.get('step') or 0)))
    except (TypeError, ValueError): step = 0
    preflight = read_json(os.path.join(job, 'preflight.json'), {})
    # 展示进度同时参考落盘产物：事件丢失/延迟时，不让界面永久卡在第 1 步。
    # 但显式的 step=0 环境预检期间不跨过该阶段，即使目录里还有旧轮产物。
    if not (step == 0 and preflight and state == 'running'):
        step = max(step, verified_step(job, 12))
    checkpoint_step = step
    activity_step = (0 if step == 0 and preflight and state == 'running'
                     else max(checkpoint_step, observed_activity_step(job)))
    pipeline_state = generation_pipeline.load(job)
    pipeline_nodes = (pipeline_state.get('nodes') or []) if (
        pipeline_state.get('version') == generation_pipeline.PIPELINE_VERSION) else []
    pipeline_phase = {
        'source_parse': ('parse', 4), 'response_plan': ('plan', 6),
        'assemble': ('assemble', 9), 'quality_review': ('deliver', 11),
        'word_export': ('deliver', 12), 'delivery_gate': ('deliver', 12),
    }
    def node_phase(node):
        node_id = str(node.get('id') or '')
        return ('write', 8) if node_id.startswith('chapter_write:') else pipeline_phase.get(node_id, ('parse', 1))
    if pipeline_nodes:
        for node in pipeline_nodes:
            if node.get('state') in ('done', 'skipped'):
                checkpoint_step = max(checkpoint_step, node_phase(node)[1])
        active_pipeline_node = next((node for node in pipeline_nodes
                                     if node.get('state') in ('running', 'retry_wait', 'failed', 'blocked')), None)
        if active_pipeline_node:
            phase_id, phase_step = node_phase(active_pipeline_node)
            activity_step = {'parse': 1, 'plan': 5, 'write': 7,
                             'assemble': 9, 'deliver': 11}.get(phase_id, phase_step)
        else:
            activity_step = max(activity_step, checkpoint_step)
    stage_rows = (stage_stats() or {}).get('stages') or []
    timeline = _progress_timeline(job)
    clock = _parse_ts(now()) or datetime.datetime.now()
    created = _parse_ts(meta.get('created_at')) or (timeline[0][1] if timeline else clock)
    terminal_at = _parse_ts(outcome.get('ts') or (prog.get('ts') if state in ('done', 'stopped', 'unknown') else ''))

    def first_step_time(predicate):
        return next((stamp for observed_step, stamp in timeline if predicate(observed_step)), None)
    checks = []
    for item in (preflight.get('checks') or []):
        if not isinstance(item, dict): continue
        checks.append({'id': str(item.get('id') or ''), 'label': str(item.get('label') or ''),
                       'state': _flow_check_state(item.get('status')),
                       'detail': str(item.get('message') or '')})

    current_index = 0
    if activity_step:
        current_index = next((idx for idx, phase in enumerate(FLOW_PHASES)
                              if phase['first'] <= activity_step <= phase['last']), len(FLOW_PHASES) - 1)
    terminal_problem = state in ('paused', 'stopped', 'unknown')
    problem_state = 'attention' if state in ('paused', 'stopped') else 'failed'
    raw_reason = str(outcome.get('reason') or prog.get('stage') or '').strip()
    reason = _friendly_current_action(raw_reason, '') if raw_reason else ''
    chapters_now = _chapter_mds(job)
    if terminal_problem and chapters_now and not _body_mds(job) and 'Word 导出失败' in raw_reason:
        reason = '撰写中断，已保留 %d 个章节；可从已保存内容继续' % len(chapters_now)
    phases = []
    for idx, definition in enumerate(FLOW_PHASES):
        if definition['id'] == 'environment':
            phase_state = 'done' if checkpoint_step >= 1 or state == 'done' else 'active'
        elif state == 'done' or checkpoint_step > definition['last']:
            phase_state = 'done'
        elif idx == current_index:
            phase_state = 'active'
        elif idx < current_index:
            # 后续活动已真实落盘，但本阶段连续证据仍有缺件；必须提示补齐，不能伪装已完成。
            phase_state = 'attention'
        else:
            phase_state = 'pending'
        if terminal_problem and idx == current_index:
            phase_state = problem_state
        detail = definition['pending']
        if phase_state == 'done': detail = '已验证完成'
        elif idx == current_index:
            detail = reason or _friendly_current_action(prog.get('stage'), definition['pending'])
        phase = {'id': definition['id'], 'label': definition['label'], 'state': phase_state,
                 'detail': detail, 'evidence': definition['evidence']}
        expected, estimate_source, history_count = _phase_expected(definition, stage_rows)
        if pipeline_nodes:
            chapter_count = sum(1 for node in pipeline_nodes
                                if str(node.get('id') or '').startswith('chapter_write:'))
            pipeline_expected = {
                'environment': 60,
                'parse': 90,
                'plan': 240,
                'write': max(300, ((chapter_count + 1) // 2) * 300),
                'assemble': 120,
                'deliver': 420 if pipeline_state.get('mode') == 'standard' else 120,
            }
            expected = pipeline_expected.get(definition['id'], expected)
            estimate_source, history_count = 'pipeline_reference', 0
        if definition['id'] == 'environment':
            phase_started = created
            phase_finished = first_step_time(lambda observed: observed >= 1)
        else:
            phase_started = first_step_time(lambda observed: observed >= definition['first'])
            phase_finished = first_step_time(lambda observed: observed > definition['last'])
        if not phase_finished and phase_started:
            if idx == current_index and state == 'running': phase_finished = clock
            elif idx == current_index and terminal_problem: phase_finished = terminal_at or clock
            # 已完成的阶段没留下结束证据时,只认真实落盘的时间戳。以前会拿「此刻」兜底,
            # 于是隔夜再打开旧任务,环境准备就报出「实际 8 小时 · 通常 1 分钟」这种荒唐数字。
            elif phase_state == 'done': phase_finished = terminal_at or _parse_ts(prog.get('ts'))
        elapsed_seconds = None
        if phase_started and phase_finished:
            elapsed_seconds = max(0, int((phase_finished - phase_started).total_seconds()))
        # 环境准备是从「建任务」量到「第一步落盘」,中间包含用户去吃饭、隔天才点开始的空档。
        # 那段等待不是机器耗时,报出来只会让人以为程序卡了 8 小时;超出常规太多就不报实际值。
        if (definition['id'] == 'environment' and elapsed_seconds is not None
                and elapsed_seconds > max(expected * 20, 1800)):
            elapsed_seconds = None
        if phase_state == 'done': remaining_seconds = 0
        elif phase_state in ('attention', 'failed') and terminal_problem: remaining_seconds = 0
        elif elapsed_seconds is None: remaining_seconds = expected
        else: remaining_seconds = max(0, expected - elapsed_seconds)
        phase.update({'elapsed_seconds': elapsed_seconds, 'expected_seconds': expected,
                      'remaining_seconds': remaining_seconds,
                      'overdue_seconds': max(0, (elapsed_seconds or 0) - expected),
                      'estimate_source': estimate_source, 'history_steps': history_count})
        if definition['id'] == 'environment': phase['checks'] = checks
        if pipeline_nodes:
            node_checks = []
            for node in pipeline_nodes:
                if node_phase(node)[0] != definition['id']: continue
                raw_node_state = str(node.get('state') or 'pending')
                view_state = {'done': 'done', 'skipped': 'done', 'running': 'active',
                              'retry_wait': 'attention', 'failed': 'failed',
                              'blocked': 'failed'}.get(raw_node_state, 'pending')
                started_at = _parse_ts(node.get('started_at'))
                finished_at = _parse_ts(node.get('finished_at'))
                if started_at:
                    node_elapsed = max(0, int(((finished_at or clock) - started_at).total_seconds()))
                else:
                    node_elapsed = None
                detail_parts = []
                if node.get('attempt'): detail_parts.append('第 %d/%d 次' % (
                    int(node.get('attempt') or 0), int(node.get('max_attempts') or 3)))
                if node_elapsed is not None: detail_parts.append('已用 %d 秒' % node_elapsed)
                if node.get('error_code'): detail_parts.append(str(node.get('error_code')))
                node_checks.append({'id': str(node.get('id') or ''),
                                    'label': str(node.get('title') or node.get('id') or ''),
                                    'state': view_state, 'detail': ' · '.join(detail_parts),
                                    'attempt': int(node.get('attempt') or 0),
                                    'max_attempts': int(node.get('max_attempts') or 3),
                                    'elapsed_seconds': node_elapsed})
            if node_checks: phase['checks'] = node_checks
        phases.append(phase)

    checkpoint = {'step': checkpoint_step,
                  'label': (STAGES[checkpoint_step - 1] if checkpoint_step else '任务文件已保存')}
    last_activity = _parse_ts(_job_last_activity(job, meta, prog, outcome)) or created
    silence_seconds = max(0, int((clock - last_activity).total_seconds()))
    # 历史阶段可能很长，但“完全没有任何活动”的心跳警告不能因此被延后。
    # 这里与 oc_run 的慢响应提示使用同一阈值，界面与执行层才不会给出矛盾状态。
    slow_after = int(OC_SLOW)
    stopped_for_stall = terminal_problem and bool(re.search(r'连接中断|长时间无进展|没有响应', reason))
    stalled = bool((state == 'running' and silence_seconds >= slow_after) or stopped_for_stall)
    parsed_source_ready = bool(_named_files(job, ('解析版',), '.md', 20))
    if state == 'running' and activity_step <= 1 and parsed_source_ready and not stalled:
        parsed_action = '招标正文已提取，正在识别目录、条款、评分项和废标条件'
        phases[current_index]['detail'] = parsed_action
        reason = parsed_action
    if state == 'running' and activity_step > checkpoint_step:
        if activity_step in (7, 8):
            count = len(_chapter_mds(job))
            reason = ('正在分章撰写，已落盘 %d 个章节' % count) if count else '正在撰写章节和逐条响应'
        elif activity_step in (9, 10): reason = '正文已汇总，正在装配和检查 Word'
        elif activity_step >= 11: reason = '正在执行出件前检查'
        phases[current_index]['detail'] = reason
    if pipeline_nodes and active_pipeline_node:
        title = str(active_pipeline_node.get('title') or active_pipeline_node.get('id') or '当前节点')
        if active_pipeline_node.get('state') == 'retry_wait':
            reason = '%s连接中断，正在等待重试（%d/%d）' % (
                title, int(active_pipeline_node.get('attempt') or 0),
                int(active_pipeline_node.get('max_attempts') or 3))
        elif active_pipeline_node.get('state') in ('failed', 'blocked'):
            reason = '%s未完成，已停在可恢复检查点' % title
        else:
            reason = '正在执行：%s' % title
        phases[current_index]['detail'] = reason
    if stalled and state == 'running':
        phases[current_index]['state'] = 'attention'
        phases[current_index]['detail'] = '模型响应偏慢，正在持续检查连接'
    remaining = 0 if state in ('done', 'stopped', 'unknown', 'paused') else sum(
        int(phase.get('remaining_seconds') or 0) for phase in phases)
    current_timing = phases[current_index]
    pipeline_summary = generation_pipeline.summary(job) if pipeline_nodes else None
    return {'version': 2, 'current_phase': FLOW_PHASES[current_index]['id'],
            'current_action': phases[current_index]['detail'] if stalled else (reason or phases[current_index]['detail']),
            'checkpoint': checkpoint,
            'recoverable': bool(os.path.isdir(job) and os.path.isfile(os.path.join(job, '任务.json'))),
            'stalled': stalled, 'silence_seconds': silence_seconds,
            'elapsed_seconds': current_timing.get('elapsed_seconds'),
            'expected_seconds': current_timing.get('expected_seconds'),
            'remaining_seconds': remaining, 'pipeline': pipeline_summary,
            'phases': phases}

@app.patch('/v1/jobs/{jid}')
async def update_job(jid: str, req: Request):
    job = jpath(jid)
    if not os.path.isdir(job): return JSONResponse({'ok': False, 'error': '任务不存在'}, 404)
    body = await req.json()
    changes = {}
    if 'name' in body:
        name = str(body.get('name') or '').strip()
        if not name: return JSONResponse({'ok': False, 'error': '任务名称不能为空'}, 400)
        changes['name'] = name[:120]
    if 'project_id' in body:
        changes['project_id'] = str(body.get('project_id') or '').strip()[:120]
    if 'archived' in body:
        changes['archived_at'] = now() if bool(body.get('archived')) else ''
    if not changes:
        return JSONResponse({'ok': False, 'error': '没有可更新的字段'}, 400)
    update_product_meta(job, **changes)
    return {'ok': True, 'job_id': os.path.basename(jid), **product_meta(job)}

@app.get('/v1/jobs')
def list_jobs(scope: str = 'all', project_id: str = ''):
    out = []
    scope = str(scope or 'all').lower()
    if scope not in ('all', 'active', 'archived'): scope = 'all'
    project_id = str(project_id or '')
    for jid in sorted(os.listdir(jobs_dir()), reverse=True):
        job = jpath(jid)
        if not os.path.isdir(job): continue
        meta = read_json(os.path.join(job, '任务.json'), {})
        prog = read_json(os.path.join(job, 'progress.json'), {})
        if isinstance(prog, dict) and prog.get('type') == 'progress':
            prog = sanitize_event(job, prog)
        st = job_state(job, meta, prog)
        outcome = read_json(os.path.join(job, 'outcome.json'), {})
        delivery = delivery_summary(job)
        product = product_meta(job, meta)
        archived = bool(product.get('archived_at'))
        if scope == 'active' and archived: continue
        if scope == 'archived' and not archived: continue
        if project_id and product.get('project_id') != project_id: continue
        template_snapshot = meta.get('template_snapshot') if isinstance(meta.get('template_snapshot'), dict) else {}
        presentation, current_action = job_presentation(job, st, meta, prog, delivery)
        flow = job_flow(job, st, meta, prog, outcome)
        if st == 'running' and flow.get('current_action'):
            current_action = flow['current_action']
        last_activity = _job_last_activity(job, meta, prog, outcome)
        elapsed = _job_elapsed(job, st, meta, outcome, last_activity)
        eta = int(flow.get('remaining_seconds') or 0)
        out.append({'job_id': jid, 'name': product['name'], 'created_at': meta.get('created_at', ''),
                    'stage': prog.get('stage', '启动中'), 'pct': prog.get('pct', 0),
                    'staged': bool(meta.get('staged')),
                    'has_word': bool(_body_docxs(job)),
                    'state': st, 'can': job_can(job, st, meta),
                    'presentation': presentation, 'presentation_state': presentation['code'],
                    'status': presentation['label'], 'current_action': current_action,
                    'flow': flow,
                    'last_activity_at': last_activity,
                    'eta': eta, 'eta_seconds': eta, 'elapsed': elapsed, 'elapsed_seconds': elapsed,
                    'usage': _job_usage(job, meta), 'runtime': job_runtime(job, meta),
                    'delivery': delivery, 'archived_at': product['archived_at'],
                    'project_id': product['project_id'], 'version': product['version'],
                    'parent_job_id': product['parent_job_id'], 'root_job_id': product['root_job_id'],
                    'template_id': str(meta.get('template_id') or template_snapshot.get('id') or ''),
                    'template_name': str(template_snapshot.get('name') or '')})
    return out

def _json_response_payload(response):
    if isinstance(response, JSONResponse):
        try: body = json.loads(response.body.decode('utf-8'))
        except Exception: body = {'error': '操作失败'}
        return int(response.status_code), body
    return 200, response if isinstance(response, dict) else {'ok': False, 'error': '操作失败'}

@app.post('/v1/jobs/bulk')
async def bulk_jobs(req: Request):
    body = await req.json()
    action = str(body.get('action') or '').strip().lower()
    raw_ids = body.get('job_ids') or []
    if action not in ('archive', 'restore', 'rerun', 'delete'):
        return JSONResponse({'ok': False, 'error': '不支持的批量操作'}, 400)
    if not isinstance(raw_ids, list) or not raw_ids:
        return JSONResponse({'ok': False, 'error': '请选择至少一个任务'}, 400)
    job_ids = []
    for value in raw_ids[:100]:
        jid = str(value or '')
        if jid and os.path.basename(jid) == jid and jid not in job_ids: job_ids.append(jid)
    succeeded, created, failed = [], [], []
    for jid in job_ids:
        job = jpath(jid)
        if not os.path.isdir(job):
            failed.append({'job_id': jid, 'error': '任务不存在'})
            continue
        if action in ('archive', 'restore'):
            update_product_meta(job, archived_at=(now() if action == 'archive' else ''))
            succeeded.append(jid)
            continue
        if action == 'delete':
            status, result = _json_response_payload(del_job(jid))
        else:
            status, result = _json_response_payload(await rerun_job(jid))
        if status < 400 and result.get('ok', True):
            succeeded.append(jid)
            if result.get('job_id'): created.append(result['job_id'])
        else:
            failed.append({'job_id': jid, 'error': str(result.get('error') or '操作失败')})
    return {'ok': not failed, 'action': action, 'succeeded': succeeded,
            'created_job_ids': created, 'failed': failed}

@app.post('/v1/jobs/export')
async def export_jobs(req: Request):
    """Download only user deliverables; never include source tenders, keys or logs."""
    body = await req.json()
    raw_ids = body.get('job_ids') or []
    if not isinstance(raw_ids, list) or not raw_ids:
        return JSONResponse({'ok': False, 'error': '请选择至少一个任务'}, 400)
    buf = io.BytesIO(); count = 0
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as archive:
        for raw in raw_ids[:100]:
            jid = str(raw or '')
            if not jid or os.path.basename(jid) != jid: continue
            job = jpath(jid)
            if not os.path.isdir(job): continue
            for name in list_deliverables(job):
                low = name.lower()
                if ('解析版' in name or '招标文件解析' in name or '原文提取' in name
                        or 'tender_parsed' in low or low.startswith('parsed_')):
                    continue
                path = os.path.join(job, name)
                if not os.path.isfile(path): continue
                archive.write(path, '%s/%s' % (jid, os.path.basename(name)))
                count += 1
    if not count:
        return JSONResponse({'ok': False, 'error': '所选任务还没有可交付文件'}, 400)
    data = buf.getvalue()
    return StreamingResponse(io.BytesIO(data), media_type='application/zip',
                             headers={'Content-Disposition': 'attachment; filename="bid-dog-deliverables.zip"',
                                      'Content-Length': str(len(data)),
                                      'X-Deliverable-Count': str(count)})

@app.get('/v1/jobs/{jid}/events')
def events(jid: str, offset: int = 0, follow: bool = True):
    """SSE:从 offset 行起回放并持续跟踪 events.jsonl;
       真实 agent 自己写的事件行没有 ts,这里首次读到时补一个并持久化,
       否则每次重连/切换任务都会被当成"刚刚发生",耗时与预估全部归零。"""
    job = jpath(jid)
    path = os.path.join(job, 'events.jsonl')
    tsf = os.path.join(job, '.line_ts.json')
    def stamp(idx, line, index):
        cursor = idx + 1
        try: e = json.loads(line)
        except Exception:
            # 已换行但损坏的历史事件不能把无效 JSON 直接喂给页面；
            # 仍返回 cursor，让重连可以稳定跨过这一行。
            e = {'type': 'worklog', 'lines': ['检测到一条损坏的历史进度记录，已自动跳过。']}
        e = sanitize_event(job, e)
        if not e.get('ts'):
            key = str(idx)
            if key not in index:
                index[key] = now()
                try: write_json(tsf, index)
                except Exception: pass
            e['ts'] = index[key]
        e['_cursor'] = cursor
        return json.dumps(e, ensure_ascii=False)
    async def gen():
        # async 生成器:网页端每个连接不再占死一条线程池线程(多标签页/多次重连曾把线程池耗干,
        # 表现为"页面转半天然后断开、agent 请求全部卡住")。空转时 10 秒一个心跳防反代掐线。
        # 事件文件按字节偏移**增量**读:旧版每秒全量 read()+splitlines,长任务(数 MB 事件)
        # 乘以标签页数,机械盘上足以把界面状态拖到落后于实际进度(「状态不准」观感的一部分)。
        sent, idle = max(0, int(offset or 0)), 0
        index = read_json(tsf, {})
        lines, consumed, pending = [], 0, b''
        for _ in range(3600 * 8):
            burst = False
            if os.path.isfile(path):
                try:
                    size = os.path.getsize(path)
                    if size < consumed:            # 文件被重建:从头再读
                        lines, consumed, pending = [], 0, b''
                    if size > consumed:
                        with open(path, 'rb') as f:
                            f.seek(consumed)
                            chunk = f.read()
                        consumed += len(chunk)
                        parts = (pending + chunk).split(b'\n')
                        pending = parts.pop()      # 半截 JSON 不构成事件,等换行落盘再发
                        for raw_ln in parts:
                            s = raw_ln.decode('utf-8', 'ignore').strip('\r')
                            if s: lines.append(s)
                except Exception:
                    pass
                if sent > len(lines):
                    reset = {'type': 'stream_reset', 'cursor': 0, '_cursor': 0,
                             'reason': 'event_log_recreated', 'ts': now()}
                    yield 'id: 0\ndata: %s\n\n' % json.dumps(reset, ensure_ascii=False)
                    sent = 0; burst = True
                while sent < len(lines):
                    cursor = sent + 1
                    yield 'id: %d\ndata: %s\n\n' % (cursor, stamp(sent, lines[sent], index))
                    sent += 1; burst = True
            if not follow:
                break
            idle = 0 if burst else idle + 1
            if idle and idle % 10 == 0: yield ': ping\n\n'
            await asyncio.sleep(1)
        if follow:
            # 到达连接寿命上限:显式告知前端立即重连,而不是无声断掉靠 onerror 兜底
            rotate = {'type': 'stream_reset', 'cursor': sent, '_cursor': sent,
                      'reason': 'rotate', 'ts': now()}
            yield 'id: %d\ndata: %s\n\n' % (sent, json.dumps(rotate, ensure_ascii=False))
    return StreamingResponse(gen(), media_type='text/event-stream',
                             headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})

def job_context(job, limit=9000):
    """把任务的关键产出摘要拼成上下文,供任务结束后的问答使用"""
    ctx, budget = [], limit
    meta = read_json(os.path.join(job, '任务.json'), {})
    ctx.append('任务:%s(招标文件:%s)' % (meta.get('name', ''), meta.get('tender', '')))
    prog = read_json(os.path.join(job, 'progress.json'), {})
    if prog: ctx.append('进度:%s %s%%' % (prog.get('stage', ''), prog.get('pct', 0)))
    ctx.append('已产出:%s' % '、'.join(list_deliverables(job)) or '(无)')
    # 用户上传的参考资料优先进上下文(他多半就是想问这个)
    rd = os.path.join(job, '参考资料')
    if os.path.isdir(rd):
        for fn in sorted(os.listdir(rd))[:3]:
            if budget <= 0: break
            fp = os.path.join(rd, fn)
            try:
                import doc_quality as dq
                txt = dq.read_any(fp) if fn.lower().endswith(('.docx', '.md', '.txt')) else ''
            except Exception:
                txt = ''
            if txt:
                take = txt[:min(3000, budget)]
                ctx.append('--- 参考资料:%s ---\n%s' % (fn, take)); budget -= len(take)
    # 优先喂自检报告与关键分析件,再补正文
    order = ['自检', '响应矩阵', '废标', '评分', '偏离', '组成', '格式要求', '投标文件_']
    files = sorted(list_deliverables(job), key=lambda f: next((i for i, k in enumerate(order) if k in f), 99))
    for fn in files:
        if not fn.endswith('.md') or budget <= 0: continue
        try: txt = open(os.path.join(job, fn), encoding='utf-8', errors='ignore').read()
        except Exception: continue
        take = txt[:min(2500, budget)]
        ctx.append('--- %s ---\n%s' % (fn, take)); budget -= len(take)
    return '\n'.join(ctx)

def chat_reply(job, jid, question):
    """任务结束后的追问:用已配模型 + 任务产出上下文回答。无论成败最后都发 status idle,前端"正在回复"不悬空"""
    try:
        conf = read_json(conf_path(), {'providers': [], 'routing': {}})
        ps = conf.get('providers') or []
        p = next((x for x in ps if x.get('id') == (conf.get('routing') or {}).get('default')), ps[-1] if ps else None)
        if not p or not p.get('model'):
            emit(job, {'type': 'message', 'role': 'agent',
                       'text': '这条已记下。当前任务的 agent 已结束;要我基于已生成的内容直接回答,请先在「设置 · 模型接入」里配好模型。'})
            return
        hist = []
        try:
            for ln in open(os.path.join(job, 'events.jsonl'), encoding='utf-8').read().splitlines()[-40:]:
                e = json.loads(ln)
                if e.get('type') == 'message' and e.get('text') and e.get('role') in ('user', 'agent'):
                    hist.append({'role': 'user' if e.get('role') == 'user' else 'assistant', 'text': e['text'][:600]})
        except Exception: pass
        msgs = [{'role': 'system', 'content': '你是中标狗。基于下面这份任务的实际产出回答用户提问,'
                 '只依据材料、不编造;涉及投标人名称/报价/资质等未填项要如实说明需人工补充。\n\n' + job_context(job)}]
        msgs += [{'role': m['role'], 'content': m['text']} for m in hist[-8:]]
        msgs.append({'role': 'user', 'content': question})
        try:
            resp = _openai_req(p.get('base_url'), p.get('api_key'), '/chat/completions',
                               {'model': p['model'], 'messages': msgs, 'max_tokens': 900},
                               timeout=60, verify=p.get('verify_ssl', True))
            txt = ((resp.get('choices') or [{}])[0].get('message') or {}).get('content', '') or '(模型无回复)'
        except Exception as e:
            txt = '回答失败:%s' % net_hint(e, (p.get('api_key'),))
        emit(job, {'type': 'message', 'role': 'agent', 'text': txt})
    finally:
        emit(job, {'type': 'status', 'state': 'idle'})

@app.post('/v1/jobs/{jid}/attachments')
async def add_attachment(jid: str, file: UploadFile = File(...)):
    """给当前任务加参考资料(如过往中标标书),让 agent 能照着它的写法与口径写。

    以前这里只把文件存进 job/参考资料/ 再往 inbox.jsonl 写一条 —— 而 inbox.jsonl 没有
    任何消费者,`参考资料` 三个字在 AGENT_PROMPT 与整个技能包里**一次都没出现过**。
    界面上三处却都写着「AI 撰写时会参考它的写法」。用户传一份去年的中标标书想让它照着写,
    agent 一个字都读不到 —— 这是「界面承诺、链路没接」的第五例。

    现在改成确定性投递:把文件复制进 **materialsDir 下的 参考资料/**(agent 被明令
    开工先 ls materialsDir),并维护一份《参考资料清单.md》说明每份是什么、该怎么用。
    这条通道是 A 级的:引擎自己写、agent 一定看得到目录,不依赖任何模型纪律。"""
    job = jpath(jid)
    if not os.path.isdir(job): return JSONResponse({'ok': False, 'error': 'not found'}, 404)
    d = os.path.join(job, '参考资料'); os.makedirs(d, exist_ok=True)
    fn = os.path.basename(file.filename or '参考资料')
    blob = await file.read()
    open(os.path.join(d, fn), 'wb').write(blob)
    # 关键一步:进本任务的 素材/参考资料,agent 才够得着。
    # 绝不能走 merged_materials():job 没有 素材/ 目录时它返回的是**全局素材库**,
    # 这一单的参考标书会被写进库里、再被合并进之后每一个新任务——跨项目串档(真机隐患)。
    delivered = False
    try:
        md = os.path.join(job, '素材', '参考资料')
        os.makedirs(md, exist_ok=True)
        open(os.path.join(md, fn), 'wb').write(blob)
        idx = os.path.join(os.path.dirname(md), '参考资料清单.md')
        if not os.path.isfile(idx):
            open(idx, 'w', encoding='utf-8').write(
                '# 参考资料清单\n\n'
                '> 投标人上传的参考件(多为过往中标标书)。**只作行文与口径的参照**:\n'
                '> 章节怎么组织、话怎么说、详略怎么分配可以学它。\n'
                '> **事实数据一律以本素材库其余文件为准**,不得从参考件里搬公司名、\n'
                '> 业绩、资质、人员、报价——那是另一个项目的,搬过来就是编造。\n\n'
                '| 文件 | 上传时间 |\n|---|---|\n')
        open(idx, 'a', encoding='utf-8').write('| 参考资料/%s | %s |\n' % (fn, now()))
        delivered = True
    except Exception:
        pass
    emit(job, {'type': 'message', 'role': 'user', 'text': '（已上传参考资料:%s）' % fn})
    running = job_state(job) == 'running'
    if not delivered:
        txt = ('已存到任务的「参考资料」目录,但**没能放进素材目录**,这一单的 agent 可能读不到它。'
               '可以打开任务文件夹手动把它放进素材库。')
    elif running:
        # 正在跑的这一轮,agent 早就 ls 过素材目录了 —— 别再承诺「后续章节会参考」
        txt = ('已收到「%s」并放进素材目录。**这一轮多半已经错过了**(agent 开工时就读完了素材目录),'
               '它会在下次「重跑」或「定向重做」时被读到。现在也可以直接问我关于它的问题。' % fn)
    else:
        txt = ('已收到「%s」并放进素材目录,开跑时 agent 会读到它,照着它的写法与口径写'
               '(事实数据仍以素材库其余文件为准)。现在也可以直接问我关于它的问题。' % fn)
    emit(job, {'type': 'message', 'role': 'agent', 'text': txt})
    return {'ok': True, 'name': fn, 'delivered': delivered}

@app.get('/v1/jobs/{jid}/attachments')
def list_attachments(jid: str):
    """本任务的输入件清单:创建时导入的素材(kind=material)+ 事后添加的参考资料(kind=reference)。

    真机反馈 7:创建任务导入了 7 个素材文件,任务详情什么都不显示——旧版只列 job/参考资料/。
    创建导入件在 job/素材/,启动后又会被全局库合并进同一目录,没法靠列目录区分;
    任务.json 的 uploaded_materials 才是「用户本单亲手导入」的可靠清单。"""
    job = jpath(jid)
    out = []
    meta = read_json(os.path.join(job, '任务.json'), {})
    for rel in (meta.get('uploaded_materials') or [])[:200]:
        p = os.path.join(job, '素材', rel)
        if os.path.isfile(p):
            out.append({'name': rel, 'kind': 'material',
                        'size_kb': round(os.path.getsize(p) / 1024, 1)})
    d = os.path.join(job, '参考资料')
    if os.path.isdir(d):
        out += [{'name': fn, 'kind': 'reference',
                 'size_kb': round(os.path.getsize(os.path.join(d, fn)) / 1024, 1)}
                for fn in sorted(os.listdir(d)) if not fn.startswith('.')]
    return out

@app.post('/v1/jobs/{jid}/rerun')
async def rerun_job(jid: str, req: Request = None):
    """用同一份招标文件重开一个任务(素材库补好后重跑,不用再找原文件)"""
    gate = _generation_gate_response()
    if gate is not None: return gate
    old = jpath(jid)
    meta = read_json(os.path.join(old, '任务.json'), {})
    tname = meta.get('tender', '')
    tpath = os.path.join(old, tname)
    if not (tname and os.path.isfile(tpath)):
        return JSONResponse({'ok': False, 'error': '原任务的招标文件不在了'}, 404)
    raw_key = str(req.headers.get('idempotency-key') or '') if req is not None else ''
    idem_key = raw_key[:128] if re.fullmatch(r'[A-Za-z0-9._:-]{8,128}', raw_key) else ''
    ledger_path = os.path.join(old, '.rerun_requests.json')
    # 同一次 UI 操作的重试/双击只能创建一个子任务。锁覆盖“查账本→
    # 建目录→写账本”，因此两个并发请求不会同时看到空值。
    with _json_lock(ledger_path):
        ledger = read_json(ledger_path, {}) if idem_key else {}
        prior = ledger.get(idem_key) if isinstance(ledger, dict) and idem_key else None
        if isinstance(prior, dict) and os.path.isdir(jpath(prior.get('job_id') or '')):
            return {'ok': True, 'job_id': prior['job_id'], 'mode': prior.get('mode', 'staged'),
                    'deduplicated': True}
        nid = datetime.datetime.now().strftime('%m%d-%H%M%S-') + uuid.uuid4().hex[:4]
        nj = jpath(nid); os.makedirs(nj)
        try:
            shutil.copy2(tpath, os.path.join(nj, tname))
            # A rerun must be reproducible from this task's own inputs. Copy both the
            # wizard material tree and later reference uploads, including nested files.
            for dirname in ('素材', '参考资料'):
                source = os.path.join(old, dirname)
                if os.path.isdir(source): shutil.copytree(source, os.path.join(nj, dirname))
            requirements = os.path.join(old, '你的要求.md')
            if os.path.isfile(requirements):
                shutil.copy2(requirements, os.path.join(nj, '你的要求.md'))
            old_product = product_meta(old, meta)
            write_json(os.path.join(nj, '任务.json'),
                       {'name': (old_product.get('name') or tname) + ' · 重跑',
                        'created_at': now(), 'paused': False, 'staged': True,
                        'tender': tname, 'prompt': meta.get('prompt', ''),
                        'template_id': meta.get('template_id', ''),
                        'template_snapshot': meta.get('template_snapshot') or
                                             task_template_snapshot(get_task_template(meta.get('template_id')))})
            write_json(os.path.join(nj, 'product.json'),
                       {'project_id': old_product.get('project_id') or '', 'version': 1,
                        'root_job_id': nid, 'rerun_of': os.path.basename(jid), 'created_at': now()})
            if idem_key:
                ledger[idem_key] = {'job_id': nid, 'mode': 'staged', 'created_at': now()}
                # 只保留最近的操作键，避免长期使用后账本无限增长。
                items = list(ledger.items())[-32:]
                write_json(ledger_path, dict(items))
        except Exception:
            shutil.rmtree(nj, ignore_errors=True)
            raise
    # 重跑与新建/暂存启动必须走同一条派发路径：轮换 run_id、冻结模型与技能收据，
    # 再决定 OpenCode server / CLI / mock，避免旁路漏掉出件证据。
    try:
        result = _launch_job(nid, nj, 'auto')
    except Exception:
        # 派发在 worker 接管前异常：撤回子目录和幂等占位。否则同一请求
        # 重试会“成功”返回一个从未启动的毒化会话。
        if idem_key:
            with _json_lock(ledger_path):
                ledger = read_json(ledger_path, {})
                if isinstance(ledger.get(idem_key), dict) and ledger[idem_key].get('job_id') == nid:
                    ledger.pop(idem_key, None)
                    write_json(ledger_path, ledger)
        shutil.rmtree(nj, ignore_errors=True)
        raise
    payload = _launch_http_result(result, include_ok=True)
    if idem_key and isinstance(payload, dict):
        with _json_lock(ledger_path):
            ledger = read_json(ledger_path, {})
            if isinstance(ledger.get(idem_key), dict):
                ledger[idem_key]['mode'] = str(payload.get('mode') or 'staged')
                write_json(ledger_path, ledger)
    return payload

def _revision_family(root_id):
    rows = []
    try: ids = os.listdir(jobs_dir())
    except OSError: ids = []
    for candidate in ids:
        path = jpath(candidate)
        if not os.path.isdir(path): continue
        meta = read_json(os.path.join(path, '任务.json'), {})
        product = product_meta(path, meta)
        if product.get('root_job_id') != root_id or candidate == root_id: continue
        rows.append({'job_id': candidate, 'name': product['name'], 'version': product['version'],
                     'parent_job_id': product['parent_job_id'], 'created_at': meta.get('created_at', ''),
                     'state': job_state(path, meta)})
    return sorted(rows, key=lambda row: (int(row.get('version') or 1), row['created_at'], row['job_id']))

@app.get('/v1/jobs/{jid}/revisions')
def list_job_revisions(jid: str):
    job = jpath(jid)
    if not os.path.isdir(job): return JSONResponse({'ok': False, 'error': '任务不存在'}, 404)
    root_id = product_meta(job).get('root_job_id') or os.path.basename(jid)
    return _revision_family(root_id)

# ---------- 修改指令路由:指到某一章就只重写那一章,别整单重跑 ----------
# 「继续修改」以前一律建子任务从头跑全部节点:用户写「第三章售后响应时间改成 2 小时」,
# 要等 5–12 章全部重写完。现在先判范围:指令指到具体章节(第 N 章 / 章节标题 / 偏离表)
# 且不涉及整册与版式 → 走单章重写(旧稿进历史版本,汇总与 Word 自动跟着更新);
# 否则才出整册新版本。判定是确定性的,并把结果先给用户看,由用户最终选。
REVISION_WHOLE_RX = re.compile(
    r'(整册|整本|整份|全书|全部章节|所有章节|每一章|各章|全文|从头|重新生成|封面|目录|页眉|页脚|页码'
    r'|排版|版式|格式|字体|页边距|统一(?:改|换|替换|成|为)|投标人名称|公司名|公司名称)')


def _route_revision_instruction(job, instruction):
    text = str(instruction or '').strip()
    try:
        state = generation_pipeline.load(job)
    except Exception:
        state = {}
    if state.get('version') != generation_pipeline.PIPELINE_VERSION:
        return {'scope': 'whole', 'reason': '这一单用的是旧版生成流程,不支持单章重写', 'chapters': []}
    nodes = [item for item in (state.get('nodes') or [])
             if str(item.get('id') or '').startswith('chapter_write:')]
    chapters = [{'node_id': str(item.get('id') or ''), 'title': str(item.get('title') or '')} for item in nodes]
    prose = [item for item in chapters if not item['node_id'].endswith((':technical_deviation', ':business_deviation'))]
    hits = []
    def add(item):
        if item and item['node_id'] not in [h['node_id'] for h in hits]: hits.append(item)
    if re.search(r'技术.{0,4}偏离|技术应答', text):
        add(next((c for c in chapters if c['node_id'].endswith(':technical_deviation')), None))
    if re.search(r'商务.{0,4}偏离', text):
        add(next((c for c in chapters if c['node_id'].endswith(':business_deviation')), None))
    for m in re.finditer(r'第\s*([0-9０-９零〇一二两三四五六七八九十]+)\s*(?:章|部分|节|篇)', text):
        idx = _cn_int(m.group(1))
        if 1 <= idx <= len(prose): add(prose[idx - 1])
    key = _plan_norm(text)
    for item in sorted(prose, key=lambda c: -len(c['title'])):
        nk = _plan_norm(item['title'])
        if len(nk) >= 2 and (item['title'] in text or nk in key): add(item)
    whole = bool(REVISION_WHOLE_RX.search(text))
    if len(hits) == 1 and not whole:
        return {'scope': 'chapter', 'node_id': hits[0]['node_id'], 'title': hits[0]['title'],
                'reason': '指令指到了「%s」这一章' % hits[0]['title'], 'chapters': chapters}
    if len(hits) > 1:
        return {'scope': 'whole', 'reason': '指令涉及 %d 章(%s),引擎一次只能重写一章;整册新版本会一起改' % (
            len(hits), '、'.join(h['title'] for h in hits[:4])), 'chapters': chapters, 'hits': hits}
    if whole:
        return {'scope': 'whole', 'reason': '指令涉及整册或版式', 'chapters': chapters}
    return {'scope': 'whole', 'reason': '指令没有指到具体章节;要只改一章,请在下面选章', 'chapters': chapters}


@app.post('/v1/jobs/{jid}/revisions/plan')
async def plan_job_revision(jid: str, req: Request):
    job = jpath(jid)
    if not os.path.isdir(job): return JSONResponse({'ok': False, 'error': '任务不存在'}, 404)
    try: body = await req.json()
    except Exception: body = {}
    return {'ok': True, **_route_revision_instruction(job, (body or {}).get('instruction') or '')}


@app.post('/v1/jobs/{jid}/revisions')
async def create_job_revision(jid: str, req: Request):
    """Create a versioned child workspace; the parent directory stays read-only."""
    parent = jpath(jid)
    if not os.path.isdir(parent): return JSONResponse({'ok': False, 'error': '任务不存在'}, 404)
    body = await req.json()
    instruction = str(body.get('instruction') or '').strip()
    if not instruction: return JSONResponse({'ok': False, 'error': '请填写本次修改要求'}, 400)
    parent_meta = read_json(os.path.join(parent, '任务.json'), {})
    tender = os.path.basename(parent_meta.get('tender') or '')
    tender_path = os.path.join(parent, tender)
    if not tender or not os.path.isfile(tender_path):
        return JSONResponse({'ok': False, 'error': '原任务的招标文件不在了'}, 404)
    parent_product = product_meta(parent, parent_meta)
    root_id = parent_product.get('root_job_id') or os.path.basename(jid)
    # Reserve the next version by creating its product metadata while holding a
    # root-series lock. Concurrent double-clicks then see the reservation.
    with _json_lock(os.path.join(jobs_dir(), '.revision-series-' + os.path.basename(root_id))):
        family_versions = [int(item.get('version') or 1) for item in _revision_family(root_id)]
        version = max([int(parent_product.get('version') or 1)] + family_versions) + 1
        child_id = datetime.datetime.now().strftime('%m%d-%H%M%S-') + uuid.uuid4().hex[:4]
        child = jpath(child_id); os.makedirs(child)
        write_json(os.path.join(child, 'product.json'),
                   {'name': '%s · v%d' % (parent_product['name'], version),
                    'project_id': parent_product.get('project_id') or '', 'version': version,
                    'parent_job_id': os.path.basename(jid), 'root_job_id': root_id,
                    'created_at': now(), 'reserved': True})
    try:
        shutil.copy2(tender_path, os.path.join(child, tender))
        for dirname in ('素材', '参考资料'):
            source = os.path.join(parent, dirname)
            if os.path.isdir(source): shutil.copytree(source, os.path.join(child, dirname))
        for filename in list_deliverables(parent):
            source = os.path.join(parent, filename)
            if os.path.isfile(source): shutil.copy2(source, os.path.join(child, filename))
        previous_requirements = ''
        req_path = os.path.join(parent, '你的要求.md')
        if os.path.isfile(req_path):
            try: previous_requirements = open(req_path, encoding='utf-8').read().strip()
            except OSError: previous_requirements = ''
        revision_prompt = ('%s\n\n' % previous_requirements if previous_requirements else '') + \
            '# 第 %d 版修改要求\n\n%s\n\n请基于当前任务目录中的上一版结果修改，不要从头推翻；完成后重新执行目录、偏离表与出件检查。\n' % (version, instruction)
        open(os.path.join(child, '你的要求.md'), 'w', encoding='utf-8').write(revision_prompt)
        baseline = {'docx': {fn: _file_digest(os.path.join(child, fn)) for fn in _body_docxs(child)},
                    'md': {fn: _file_digest(os.path.join(child, fn)) for fn in _body_mds(child)}}
        write_json(os.path.join(child, '任务.json'),
                   {'name': '%s · v%d' % (parent_product['name'], version), 'created_at': now(),
                    'paused': False, 'staged': True, 'tender': tender,
                    'prompt': revision_prompt, 'redo_baseline': baseline,
                    'template_id': parent_meta.get('template_id', ''),
                    'template_snapshot': parent_meta.get('template_snapshot') or
                                         task_template_snapshot(get_task_template(parent_meta.get('template_id')))})
        patch_json(os.path.join(child, 'product.json'), {'reserved': False, 'updated_at': now()})
        emit(child, {'type': 'progress', 'stage': '新版已准备好，等待开始生成',
                     'pct': 0, 'step': 0, 'total': 12})
    except Exception:
        shutil.rmtree(child, ignore_errors=True)
        raise
    if body.get('start', True) is False:
        return {'ok': True, 'job_id': child_id, 'mode': 'staged', 'version': version,
                'parent_job_id': os.path.basename(jid)}
    result = _launch_job(child_id, child, 'auto')
    payload = _launch_http_result(result, include_ok=True)
    if isinstance(payload, dict): payload.update({'version': version, 'parent_job_id': os.path.basename(jid)})
    return payload

RESUME_PROMPT = (
    '接着做这一单。**先看一眼任务目录里已经有什么** —— 已经写好的不要重写、不要推翻,'
    '从没做完的地方往下接,一直做到出 .docx 为止。'
    '如果中途我补过要求或答过你的问题,以最新的为准。')

def resume_worker(job, cmd):
    """「继续做」的执行体。刻意跟正常跑同一套收尾,免得续做完的播报跟跑完的不一样。"""
    result = OC_RUN_INTERRUPTED
    try:
        result = oc_run(job, RESUME_PROMPT, allow_cli_fallback=False)
    except Exception as e:
        append_diagnostic(job, 'resume_worker_exception', str(e), level='error')
    if result == OC_RUN_CANCELLED: return
    if result != OC_RUN_COMPLETED:
        # 这里**不能**回落成从头重跑:用户点的是「继续做」,悄悄从头来会把已有产物覆盖掉
        emit(job, {'type': 'error', 'text': '暂时没能恢复上次进度。已经写出的内容都还在，'
                                            '可以再试一次继续，或新建重跑。',
                   'actions': [{'act': 'resume', 'label': '再试一次继续'},
                               {'act': 'open_log', 'label': '查看诊断详情'},
                               {'act': 'rerun', 'label': '重跑本任务'}]})
        settle(job, stop_reason='已停止（没能续做）')
        return
    finish_job(job)

@app.post('/v1/jobs/{jid}/resume')
async def resume_job(jid: str):
    """接着做:在**原来那个会话**里往下接,它还记得读过什么、写到哪。

    与「重跑本任务」的区别:重跑是另起一单从头来,产物全新;续做是同一单往下写。
    只有 server 模式留下了会话才做得到 —— 做不到就当面说,绝不悄悄改成从头重跑。"""
    job = jpath(jid)
    if not os.path.isdir(job): return JSONResponse({'ok': False, 'error': '任务不存在'}, 404)
    base = os.path.basename(jid)
    owner, reason = _reserve_running_reason(base)
    if not owner:
        return _admission_failure(reason, '这一单正在跑，不用重复续做。')
    try:
        # 先占位再读会话；否则 redo/start 可在读取后轮换 run_id/session，续做会串到旧轮。
        meta = read_json(os.path.join(job, '任务.json'), {})
        pipeline_state = generation_pipeline.load(job)
        if pipeline_state.get('version') == 1:
            pipeline_state = generation_pipeline.migrate(job)
        if pipeline_state.get('version') == generation_pipeline.PIPELINE_VERSION:
            try:
                _pipeline_validate_execution_contract(job, s2_conf(read_json(conf_path(), {})))
            except generation_pipeline.NodeExecutionError:
                _release_running(base, owner)
                return JSONResponse({
                    'ok': False,
                    'error': '任务运行所用的模型网关已变化；请在「模型接入」切回原网关后继续,或重跑本任务。'
                }, 409)
            # delivery_gate 已落盘、outcome 未落盘时是一个可提交的中间态。
            # 这里只重放幂等收尾，不重启模型节点。
            delivery_done = any(
                node.get('id') == 'delivery_gate' and node.get('state') == 'done'
                for node in pipeline_state.get('nodes') or [])
            if pipeline_state.get('state') == 'done' and delivery_done:
                final = settle(job)
                _release_running(base, owner)
                if isinstance(final, dict) and final.get('state') == 'done':
                    return {'ok': True, 'pipeline': True, 'finalized': True}
                return JSONResponse({'ok': False, 'error': '交付收尾未通过，请查看诊断详情。'}, 409)
            recovered = generation_pipeline.recover(job)
            if not recovered.get('recoverable'):
                _release_running(base, owner)
                return JSONResponse({'ok': False,
                                     'error': '当前节点已达到重试上限或被配置问题阻断，请使用“重试当前节点”。'}, 409)
            try: os.unlink(os.path.join(job, 'outcome.json'))
            except FileNotFoundError: pass
            meta['paused'] = False
            write_json(os.path.join(job, '任务.json'), meta)
            emit(job, {'type': 'message', 'role': 'agent',
                       'text': '已从检查点继续：完成节点直接复用，只执行未完成或失败的当前节点。'})
            _start_reserved_worker(base, owner, generation_pipeline_worker, job)
            return {'ok': True, 'pipeline': True}
        sid = meta.get('oc_session')
        if not sid:
            _release_running(base, owner)
            return JSONResponse({'ok': False, 'error': '这一单没有留下可恢复的进度记录，只能“重跑本任务”重新生成。'}, 400)
        snapshot = meta.get('engine_snapshot') if isinstance(meta.get('engine_snapshot'), dict) else {}
        if not session_runtime_compatible(snapshot):
            _release_running(base, owner)
            return JSONResponse({'ok': False,
                                 'error': '生成模式已更改或连接配置已更改，旧会话不能跨环境续写；'
                                          '请使用“重跑本任务”按当前模式重新生成。'}, 409)
        if not oc_serve():
            append_diagnostic(job, 'resume_connection_unavailable', 'OpenCode server unavailable', level='error')
            _release_running(base, owner)
            return JSONResponse({'ok': False, 'error': '暂时无法恢复上次进度，请稍后再试。'}, 400)
        if _cancel_requested(base, owner):
            _release_running(base, owner)
            return JSONResponse({'ok': False, 'stopped': True,
                                 'error': '续做连接期间已收到停止请求，本轮没有启动。'}, 409)
        st, _s = oc_api('/api/session/%s' % sid, timeout=20)
        if st != 200:
            append_diagnostic(job, 'resume_session_missing', 'Session lookup returned HTTP %s' % st,
                              level='error', session_id=sid)
            _release_running(base, owner)
            return JSONResponse({'ok': False, 'error': '上次进度记录已失效，只能“重跑本任务”重新生成。'}, 400)
        if _cancel_requested(base, owner):
            _release_running(base, owner)
            return JSONResponse({'ok': False, 'stopped': True,
                                 'error': '续做连接期间已收到停止请求，本轮没有启动。'}, 409)
        meta['paused'] = False; write_json(os.path.join(job, '任务.json'), meta)
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '好,接着上次的地方往下做。它还记得前面读过什么、写到哪,不会推翻已经写好的。'})
        emit(job, {'type': 'progress', 'stage': '继续做', 'pct': max(1, int(
            (read_json(os.path.join(job, 'progress.json'), {}).get('pct') or 1))), 'step': int(
            read_json(os.path.join(job, 'progress.json'), {}).get('step') or 1), 'total': 12})
        _start_reserved_worker(base, owner, resume_worker, job, None)
        return {'ok': True}
    except Exception:
        _release_running(base, owner)
        raise

def mock_redo(job, instruction):
    """演示引擎的定向重做:三步小流程,产出重做说明,进度回到完成态"""
    try:
        steps = ['定位相关内容', '按指令重写', '重新汇总自检']
        for i, st in enumerate(steps):
            if _cancelled(job): return
            emit(job, {'type': 'progress', 'stage': '定向重做:' + st, 'pct': int((i + 0.5) / 3 * 100), 'step': i + 1, 'total': 3})
            time.sleep(1.2)
        fn = '定向重做说明.md'
        open(os.path.join(job, fn), 'w', encoding='utf-8').write(
            '# 定向重做说明\n\n指令:%s\n\n(演示引擎:真实引擎会按指令改写对应产物并更新自检报告)\n' % instruction)
        emit(job, {'type': 'artifact', 'name': fn})
        settle(job)
    finally:
        pass  # reservation wrapper 负责 owner-scoped 取消清理与释放

@app.get('/v1/jobs/{jid}/log')
def job_log(jid: str, n: int = 20000):
    """agent 文本日志 + 引擎观测到的结构化技能证据状态。

    OpenCode server 模式没有 run.log(CLI 模式才有)。真机反馈:任务失败点「查看运行日志」,
    只看到一句「当前运行方式没有文本日志」——等于没给。事件流 events.jsonl 和
    diagnostics.jsonl 里其实全程都有记录,这里把它们合成一份带时间戳的可读时间线兜底。"""
    job = jpath(jid)
    txt = read_tail(os.path.join(job, 'run.log'), max(1000, min(int(n or 20000), 200000)))
    if not (txt or '').strip():
        lines = []
        try:
            for ln in open(os.path.join(job, 'events.jsonl'), encoding='utf-8', errors='ignore').read().splitlines()[-400:]:
                try: ev2 = json.loads(ln)
                except Exception: continue
                t = ev2.get('type'); ts = str(ev2.get('ts') or '')[:19]
                if t == 'progress':
                    lines.append('%s [进度] 第 %s/%s 步 · %s' % (ts, ev2.get('step') or '?', ev2.get('total') or 12, ev2.get('stage') or ''))
                elif t == 'worklog':
                    for wl in (ev2.get('lines') or [])[:8]: lines.append('%s [动作] %s' % (ts, wl))
                elif t == 'error':
                    lines.append('%s [错误] %s' % (ts, ev2.get('text') or ''))
                elif t == 'message':
                    lines.append('%s [%s] %s' % (ts, '用户' if ev2.get('role') == 'user' else '中标狗',
                                                 str(ev2.get('text') or '').replace('\n', ' ')[:200]))
                elif t == 'health':
                    lines.append('%s [体检] %s' % (ts, ev2.get('summary') or ''))
        except OSError:
            pass
        try:
            for ln in open(os.path.join(job, 'diagnostics.jsonl'), encoding='utf-8', errors='ignore').read().splitlines()[-60:]:
                try: rec2 = json.loads(ln)
                except Exception: continue
                lines.append('%s [诊断·%s] %s:%s' % (str(rec2.get('ts') or '')[:19],
                                                      rec2.get('level') or '', rec2.get('code') or '',
                                                      str(rec2.get('detail') or '')[:200]))
        except OSError:
            pass
        lines.sort()
        txt = '\n'.join(lines)
    ev = skill_evidence(job)
    if ev['state'] == 'unverifiable' and _replay_oc_skill_evidence(job):
        ev = skill_evidence(job)
    return {'ok': True,
            'log': _redact_runtime(txt) or '(这一单还没有任何运行记录)',
            'skill_state': ev['state'], 'skill_used': ev['ok'],
            'hits': ev['hits'], 'why': ev['why']}

@app.get('/v1/jobs/{jid}/pipeline')
def job_pipeline(jid: str):
    job = jpath(jid)
    if not os.path.isdir(job):
        return JSONResponse({'ok': False, 'error': '任务不存在'}, 404)
    state = generation_pipeline.load(job)
    if state.get('version') != generation_pipeline.PIPELINE_VERSION:
        return JSONResponse({'ok': False, 'error': '该任务使用旧版生成流程，没有节点检查点'}, 404)
    return {'ok': True, 'summary': generation_pipeline.summary(job), 'pipeline': state}


@app.post('/v1/jobs/{jid}/nodes/{node_id:path}/retry')
async def retry_pipeline_node(jid: str, node_id: str, request: Request):
    job = jpath(jid)
    if not os.path.isdir(job):
        return JSONResponse({'ok': False, 'error': '任务不存在'}, 404)
    raw_key = str(request.headers.get('idempotency-key') or '')[:128]
    if not re.fullmatch(r'[A-Za-z0-9._:-]{8,128}', raw_key):
        return JSONResponse({'ok': False, 'error': '缺少有效的重试请求标识'}, 400)
    ledger_path = os.path.join(job, '.node_retry_requests.json')
    ledger_key = str(node_id) + '|' + raw_key
    base = os.path.basename(jid)
    retry_reset_required = True
    def lease_is_active(entry):
        try: return float((entry or {}).get('lease_until') or 0) > time.time()
        except (AttributeError, TypeError, ValueError): return False
    with _json_lock(ledger_path):
        ledger = read_json(ledger_path, {})
        if isinstance(ledger.get(ledger_key), dict):
            prior = ledger[ledger_key]
            if prior.get('status') == 'rejected':
                return JSONResponse({'ok': False, 'error': prior.get('error') or '节点不能重试',
                                     'deduplicated': True}, int(prior.get('http_status') or 409))
            if prior.get('status') == 'dispatch_failed':
                return JSONResponse({'ok': False,
                                     'code': 'retry_dispatch_failed',
                                     'error': prior.get('error') or '上次重试派发失败，请使用新的重试操作。',
                                     'deduplicated': True}, int(prior.get('http_status') or 500))
            if prior.get('status') == 'accepted':
                return {'ok': True, 'node_id': node_id, 'deduplicated': True,
                        'accepted_at': prior.get('accepted_at') or ''}
            # `dispatching` 可能是上次进程在写入意图后崩溃留下的。
            # 内存中仍有 owner 才表示真正在派发；否则对账节点并继续事务。
            if prior.get('status') == 'dispatching':
                lease_active = lease_is_active(prior)
                # worker 很快完成时 RUNNING owner 可能已释放，但首个 HTTP 请求尚未
                # 把账本改成 accepted。短租约覆盖这个窗口，避免相同请求再次派发。
                if lease_active or _is_running(base):
                    return JSONResponse({'ok': False, 'node_id': node_id, 'deduplicated': True,
                                         'pending': True, 'error': '重试请求正在派发'}, 202)
            state = generation_pipeline.load(job)
            current = next((item for item in state.get('nodes') or []
                            if item.get('id') == node_id), None)
            if not current:
                return JSONResponse({'ok': False, 'error': '节点不存在',
                                     'deduplicated': True}, 409)
            if current.get('state') == 'pending':
                retry_reset_required = False
            elif current.get('state') not in ('failed', 'blocked'):
                return JSONResponse({'ok': False, 'error': '重试意图与节点状态不一致，请刷新后重试',
                                     'deduplicated': True}, 409)
        active_dispatch = next((entry for key, entry in ledger.items()
                                if key != ledger_key and isinstance(entry, dict)
                                and entry.get('node_id') == node_id
                                and entry.get('status') == 'dispatching'
                                and lease_is_active(entry)), None)
        if active_dispatch:
            return JSONResponse({'ok': False, 'code': 'retry_dispatch_confirmation_pending',
                                 'node_id': node_id, 'deduplicated': True, 'pending': True,
                                 'error': '重试请求已派发，正在确认运行状态'}, 202)
        owner, reason = _reserve_running_reason(base)
        if not owner:
            # 同一节点刚被(另一个 key 的)点击派发过、worker 还在跑:这不是失败,
            # 是重复点击——回 202 pending,前端显示「重试进行中」,不能弹错误吓用户。
            if reason == 'running':
                prior_same_node = next((entry for entry in ledger.values()
                                        if isinstance(entry, dict)
                                        and entry.get('node_id') == node_id
                                        and entry.get('status') in ('accepted', 'dispatching')), None)
                if prior_same_node:
                    return JSONResponse({'ok': True, 'node_id': node_id, 'deduplicated': True,
                                         'pending': True,
                                         'message': '这一节点的重试已经在跑了，稍等片刻看进度即可。'}, 202)
            return _admission_failure(reason, '这一单正在跑，不用重复重试。')
        accepted_at = (str((ledger.get(ledger_key) or {}).get('accepted_at') or '') or now())
        ledger[ledger_key] = {'node_id': node_id, 'accepted_at': accepted_at,
                              'status': 'dispatching', 'dispatch_token': owner,
                              'lease_until': time.time() + 30}
        if len(ledger) > 100:
            ledger = dict(list(ledger.items())[-100:])
        write_json(ledger_path, ledger)
    try:
        if retry_reset_required:
            node = generation_pipeline.retry_node(job, node_id)
        else:
            state = generation_pipeline.load(job)
            node = next(item for item in state.get('nodes') or [] if item.get('id') == node_id)
        try: os.unlink(os.path.join(job, 'outcome.json'))
        except FileNotFoundError: pass
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '已人工重试节点“%s”；只重做这一节点，已完成内容不会重写。' %
                           (node.get('title') or node_id)})
        _start_reserved_worker(base, owner, generation_pipeline_worker, job)
        try:
            with _json_lock(ledger_path):
                ledger = read_json(ledger_path, {})
                ledger[ledger_key] = {'node_id': node_id, 'accepted_at': accepted_at,
                                      'status': 'accepted'}
                write_json(ledger_path, ledger)
        except Exception:
            # worker 已经启动，之后的账本确认失败只能报告“结果待确认”。绝不能
            # 释放 owner/回滚节点/让前端换 key，否则会把同一节点启动两次。
            return JSONResponse({'ok': False, 'code': 'retry_dispatch_confirmation_pending',
                                 'node_id': node_id, 'deduplicated': False, 'pending': True,
                                 'error': '重试已派发，确认写入较慢，请等待状态刷新'}, 202)
        return {'ok': True, 'node_id': node_id, 'deduplicated': False}
    except generation_pipeline.PipelineError as exc:
        _release_running(base, owner)
        with _json_lock(ledger_path):
            ledger = read_json(ledger_path, {})
            ledger[ledger_key] = {'node_id': node_id, 'accepted_at': accepted_at,
                                  'status': 'rejected', 'http_status': 409, 'error': str(exc)}
            write_json(ledger_path, ledger)
        return JSONResponse({'ok': False, 'error': str(exc)}, 409)
    except Exception:
        _release_running(base, owner)
        try:
            generation_pipeline.fail_node(job, node_id, 'node_dispatch_failed', retryable=False)
        except Exception:
            pass
        with _json_lock(ledger_path):
            ledger = read_json(ledger_path, {})
            if isinstance(ledger.get(ledger_key), dict):
                ledger[ledger_key]['status'] = 'dispatch_failed'
                ledger[ledger_key]['http_status'] = 500
                ledger[ledger_key]['error'] = '节点重试派发失败，请使用新的重试操作。'
                write_json(ledger_path, ledger)
        return JSONResponse({'ok': False, 'code': 'retry_dispatch_failed',
                             'error': '节点重试派发失败，请再次点击重试。'}, 500)

@app.post('/v1/jobs/{jid}/chapters/{node_id:path}/rewrite')
async def rewrite_chapter(jid: str, node_id: str, request: Request):
    """单章重写:只重做一个章节节点,其余章节不动。

    机制:把用户补充要求写进节点契约,执行摘要随之变化 → start_node 归档旧稿并
    重跑本章;汇总/质检/Word 导出是摘要驱动的,会自动跟着重做,不需要级联重置。"""
    job = jpath(jid)
    if not os.path.isdir(job):
        return JSONResponse({'ok': False, 'error': '任务不存在'}, 404)
    try:
        body = await request.json()
    except Exception:
        body = {}
    note = str((body or {}).get('note') or '').strip()[:generation_pipeline.REWRITE_NOTE_MAX]
    try:
        state = generation_pipeline.load(job)
    except Exception:
        state = {}
    if state.get('version') != generation_pipeline.PIPELINE_VERSION:
        return JSONResponse({'ok': False, 'error': '该任务使用旧版生成流程，不支持单章重写；请重新生成整单'}, 409)
    node = next((item for item in state.get('nodes') or [] if item.get('id') == node_id), None)
    if not node:
        return JSONResponse({'ok': False, 'error': '章节不存在'}, 404)
    base = os.path.basename(jid)
    owner, reason = _reserve_running_reason(base)
    if not owner:
        return _admission_failure(reason, '这一单正在生成；等它停下或完成后，再重写单章。')
    try:
        # 旧稿另存一份到用户可见的「历史版本/」(流水线内部 .pipeline_attempts 仍是权威归档)
        serial = int(node.get('rewrite_serial') or 0) + 1
        for name in node.get('outputs') or []:
            src = os.path.join(job, os.path.basename(str(name)))
            if os.path.isfile(src):
                hist = os.path.join(job, '历史版本')
                os.makedirs(hist, exist_ok=True)
                stem, ext = os.path.splitext(os.path.basename(src))
                shutil.copy2(src, os.path.join(hist, '%s.v%d%s' % (stem, serial, ext)))
        node = generation_pipeline.rewrite_node(job, node_id, note)
        try: os.unlink(os.path.join(job, 'outcome.json'))
        except FileNotFoundError: pass
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '已开始重写章节「%s」%s。其余章节不动；旧稿已存入「历史版本」文件夹，可随时找回。' % (
                       node.get('title') or node_id, '，并带上你的补充要求' if note else '')})
        _start_reserved_worker(base, owner, generation_pipeline_worker, job)
        return {'ok': True, 'node_id': node_id, 'version': serial + 1}
    except generation_pipeline.PipelineError as exc:
        _release_running(base, owner)
        return JSONResponse({'ok': False, 'error': str(exc)}, 409)
    except Exception:
        _release_running(base, owner)
        raise

COVERAGE_CLEAR_GAPS = ('', '-', '—', '无', '无缺口', '已覆盖', '不适用', '/')

def _coverage_view(job):
    """把《评分点响应矩阵.md》变成结构化覆盖度:规划矩阵 × 章节完成状态。

    一个评分点算「已覆盖」需要同时满足:缺口列干净(没有〔需补充〕类标记)、
    响应位置落到了具体章节、且该章节已经写完(章节节点 done;旧版无流水线的
    任务退化为「整册已生成」)。这是给售前汇报用的数字,口径必须经得起追问。"""
    rows = _md_table_rows(os.path.join(job, '评分点响应矩阵.md'), 6)
    if not rows: return None
    plan = read_json(os.path.join(job, 'response_plan.json'), {})
    if not isinstance(plan, dict): plan = {}
    plan_source = str(plan.get('source') or '')
    chapter_state = {}
    try:
        state = generation_pipeline.load(job)
        if state.get('version') == generation_pipeline.PIPELINE_VERSION:
            for node in state.get('nodes') or []:
                if str(node.get('id') or '').startswith('chapter_write:'):
                    chapter_state[str(node.get('title') or '')] = \
                        {'id': node.get('id'), 'state': node.get('state')}
    except Exception:
        pass
    whole_done = os.path.isfile(os.path.join(job, '投标文件_整册.md'))
    items, covered = [], 0
    for row in rows[:200]:
        requirement, score, location, gap = row[1][:140], row[2][:20], row[3][:80], row[5][:140]
        gap_clear = (gap in COVERAGE_CLEAR_GAPS) or not gap.strip()
        located = bool(location.strip()) and '需补充' not in location
        node_id, node_title, chapter_done = '', '', whole_done
        if chapter_state:
            matched = _plan_match_title(location, [title for title in chapter_state if title])
            hit = (matched, chapter_state[matched]) if matched else None
            if hit:
                node_title, node_id = hit[0], str(hit[1].get('id') or '')
                chapter_done = hit[1].get('state') == 'done'
            else:
                chapter_done = whole_done
        ok = gap_clear and located and chapter_done
        covered += 1 if ok else 0
        # 界面上「为什么没算覆盖」以前是前端靠 node_id 空不空反推的,于是出现过
        # 左边写「落位:技术方案」右边写「未定位到章节」这种自相矛盾。原因由这里给准。
        if ok: reason = ''
        elif not located: reason = 'unlocated'
        elif not gap_clear: reason = 'gap'
        else: reason = 'chapter_pending'
        items.append({'requirement': requirement, 'score': score, 'location': location,
                      'gap': gap, 'covered': ok, 'node_id': node_id, 'chapter': node_title,
                      'reason': reason})
    # plan_source 让界面分得清「真实覆盖率」和「本地候选」:本地索引里没有一条落到章节,
    # 显示 0/N 是把一个没有意义的数字当结论。
    return {'total': len(items), 'covered': covered, 'items': items,
            'plan_source': plan_source, 'plan_model': str(plan.get('model') or ''),
            'plan_notes': [str(n) for n in (plan.get('notes') or []) if str(n).strip()]}

@app.get('/v1/jobs/{jid}/coverage')
def job_coverage(jid: str):
    job = jpath(jid)
    if not os.path.isdir(job):
        return JSONResponse({'ok': False, 'error': '任务不存在'}, 404)
    view = _coverage_view(job)
    if view is None:
        return {'ok': True, 'available': False,
                'note': '响应规划完成后才有评分点矩阵；生成开始后这里会实时更新。'}
    return {'ok': True, 'available': True, **view}

def _redact_config(value, key=''):
    if _sensitive_config_key(key) or str(key).lower() in ('cmd', 'env'):
        return ('s' + 'k' + '-***') if value else ''
    if isinstance(value, dict): return {k: _redact_config(v, k) for k, v in value.items()}
    if isinstance(value, list): return [_redact_config(v, key) for v in value]
    return redact(value)

def diagnostic_bundle(job):
    """构建可直接交给支持人员的全脱敏诊断包；返回 zip bytes。"""
    buf = io.BytesIO()
    meta_files = ('任务.json', 'progress.json', 'pipeline.json', 'source_manifest.json')
    conf = read_json(conf_path(), {})
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as z:
        for fn in meta_files:
            data = read_json(os.path.join(job, fn), {})
            z.writestr(fn, json.dumps(_redact_runtime(data, conf), ensure_ascii=False, indent=2))
        safe_events = []
        try: raw_lines = open(os.path.join(job, 'events.jsonl'), encoding='utf-8', errors='ignore').read().splitlines()
        except Exception: raw_lines = []
        for ln in raw_lines:
            try: safe_events.append(json.dumps(_redact_runtime(json.loads(ln), conf), ensure_ascii=False))
            except Exception: safe_events.append(str(_redact_runtime(ln, conf)))
        z.writestr('events.jsonl', '\n'.join(safe_events) + ('\n' if safe_events else ''))
        z.writestr('run.log', '\n'.join(_tail_lines(os.path.join(job, 'run.log'), 500)))
        z.writestr('config.redacted.json', json.dumps(_redact_config(conf),
                                                     ensure_ascii=False, indent=2))
        system = {'engine_version': ENGINE_VERSION, 'platform': platform.platform(),
                  'python': platform.python_version(), 'created_at': now(),
                  'job_state': job_state(job)}
        z.writestr('system.json', json.dumps(system, ensure_ascii=False, indent=2))
        z.writestr('relay_status.json', json.dumps(_redact_runtime(RELAY_LAST, conf), ensure_ascii=False, indent=2))
        diagnostic_lines = _tail_lines(os.path.join(job, 'diagnostics.jsonl'), 500)
        z.writestr('diagnostics.jsonl', '\n'.join(diagnostic_lines) + ('\n' if diagnostic_lines else ''))
        oc_log = os.path.join(DATA, 'opencode-server.log')
        z.writestr('opencode-server.log', '\n'.join(_tail_lines(oc_log, 500)))
    return buf.getvalue()

def _job_verdict(job, state, outcome, recent, pipe):
    """诊断的第一行必须是人话结论 + 下一步动作。

    真机反馈:任务失败后点「一键诊断」,三项环境检查全绿、技术详情一屏原始 JSON,
    用户完全看不出这单为什么停 —— 环境检查回答的是「引擎还能不能干活」,
    没回答「这一单出了什么事」。这里把停止原因、最后一条错误级诊断、
    卡住的流水线节点拼成一句结论和一条建议,诊断弹窗置顶展示。"""
    reason = str((outcome or {}).get('reason') or '')
    last_err = next((r for r in reversed(recent or [])
                     if isinstance(r, dict) and r.get('level') == 'error'), None)
    problem = (pipe or {}).get('problem') if isinstance(pipe, dict) else None
    parts, suggestion = [], ''
    if state == 'done':
        return {'summary': '这一单已完成。', 'suggestion': '', 'last_error': ''}
    if state == 'running':
        parts.append('这一单正在运行中。')
        suggestion = '如果长时间没有新动态,可以先「停止」再「从已保存内容继续」。'
    elif reason:
        parts.append('这一单停止的直接原因:%s。' % reason.strip('。'))
    if problem:
        parts.append('流水线停在节点「%s」(已试 %s/%s 次%s)。' % (
            problem.get('title') or problem.get('node_id'),
            problem.get('attempt') or 0, problem.get('max_attempts') or 0,
            (',错误码 ' + problem['error_code']) if problem.get('error_code') else ''))
        if not suggestion:
            suggestion = '建议点「重试当前节点」;已完成节点不会重做。反复失败时检查「设置 · 模型接入」的网络与 Key。'
    if last_err and not suggestion:
        code = str(last_err.get('code') or '')
        if 'stall' in code or 'interrupt' in code or '中断' in str(last_err.get('detail') or ''):
            suggestion = '多为模型网络瞬断:点「从已保存内容继续」,已写好的内容不会重写。'
        elif 'export' in code or 'docx' in code:
            suggestion = '正文已生成,点「重试导出 Word」;仍失败时重启应用后再导出。'
        else:
            suggestion = '可点「从已保存内容继续」;反复出现时导出诊断包发给支持人员。'
    if not parts:
        parts.append('没有发现这一单的明确故障记录。')
        suggestion = suggestion or '如任务未完成,可「从已保存内容继续」或「重跑本任务」。'
    return {'summary': ' '.join(parts),
            'suggestion': suggestion,
            'last_error': str((last_err or {}).get('detail') or '')[:300]}

def _diagnostic_snapshot(jid=''):
    """A fast, secret-free support snapshot; expensive network tests stay opt-in."""
    conf = read_json(conf_path(), {})
    setup = setup_status_data(conf)
    eng = conf.get('engine') if isinstance(conf.get('engine'), dict) else {}
    storage_ok = os.path.isdir(_mk(ws_root())) and os.access(ws_root(), os.W_OK)
    kind = str(eng.get('kind') or 's2')
    cli_name = 'opencode' if kind in ('s2', 'opencode') else {'codex': 'codex', 'claude': 'claude', 'sowork': 'sowork'}.get(kind)
    runtime_source = resolve_cli(cli_name, eng) if cli_name else ''
    shell_ok = bool(runtime_source) if cli_name else bool(config_agent_cmd())
    checks = [
        {'id': 'storage', 'label': '数据目录可写', 'status': 'pass' if storage_ok else 'fail',
         'message': '任务与设置可以正常保存' if storage_ok else '数据目录不可写，请检查磁盘权限'},
        {'id': 'connection', 'label': '模型连接',
         'status': 'pass' if setup.get('connected') or setup.get('legacy_skipped') else ('warning' if setup.get('key_set') else 'fail'),
         'message': ('连接已验证' if setup.get('connected') else
                     ('存量配置可用；需要时可重新测试' if setup.get('legacy_skipped') else
                      ('Key 已保存，建议运行连接测试' if setup.get('key_set') else '尚未填写 Key')))},
        {'id': 'runtime', 'label': '生成组件', 'status': 'pass' if shell_ok else 'fail',
         'message': '生成组件已就绪' if shell_ok else '生成组件尚未就绪，可在设置中一键安装'},
    ]
    provision = {key: PROV.get(key) for key in ('state', 'which', 'pct', 'note', 'error')}
    result = {'ok': not any(item['status'] == 'fail' for item in checks),
              'checked_at': now(), 'engine_version': ENGINE_VERSION,
              'checks': checks, 'active_jobs': len(_running_snapshot()),
              'runtime_source': runtime_source or '', 'provision': provision}
    if jid:
        job = jpath(jid)
        if not os.path.isdir(job): return JSONResponse({'ok': False, 'error': '任务不存在'}, 404)
        meta = read_json(os.path.join(job, '任务.json'), {})
        prog = read_json(os.path.join(job, 'progress.json'), {})
        state = job_state(job, meta, prog)
        presentation, current_action = job_presentation(job, state, meta, prog)
        recent = []
        try: lines = open(os.path.join(job, 'diagnostics.jsonl'), encoding='utf-8', errors='ignore').read().splitlines()[-20:]
        except OSError: lines = []
        for line in lines:
            try: recent.append(_redact_runtime(json.loads(line), conf))
            except Exception: recent.append({'detail': _redact_runtime(line, conf)})
        preflight = read_json(os.path.join(job, 'preflight.json'), {})
        last_activity = str(prog.get('ts') or meta.get('updated_at') or meta.get('created_at') or '')
        if not last_activity:
            try: last_activity = datetime.datetime.fromtimestamp(os.path.getmtime(os.path.join(job, 'events.jsonl'))).isoformat(timespec='seconds')
            except OSError: last_activity = now()
        outcome = read_json(os.path.join(job, 'outcome.json'), {})
        pipe = generation_pipeline.summary(job)
        verdict = _job_verdict(job, state, outcome, recent, pipe)
        # 第 4 项检查:本单执行。前三项是环境体检,任务失败时它们照样全绿——
        # 用户点「一键诊断」要的答案是「这一单怎么了」,必须让它参与结论。
        job_ok = state in ('done', 'running') and not (pipe or {}).get('problem')
        checks.append({'id': 'job', 'label': '本单执行',
                       'status': 'pass' if job_ok else 'fail',
                       'message': verdict.get('summary') or ''})
        result['ok'] = result['ok'] and job_ok
        result['job'] = {'job_id': os.path.basename(jid), 'state': state,
                         'presentation': presentation, 'current_action': current_action,
                         'runtime': job_runtime(job, meta), 'delivery': delivery_summary(job),
                         'flow': job_flow(job, state, meta, prog, outcome),
                         'preflight': preflight, 'last_activity': last_activity,
                         'pipeline': pipe,
                         'verdict': verdict,
                         'recent_diagnostics': recent,
                         'bundle_url': '/v1/jobs/%s/bundle' % os.path.basename(jid)}
    # 「技术详情」给人话文本;原始 JSON 只进诊断包,不再糊用户一脸
    lines = ['[%s] %s:%s' % ({'pass': '通过', 'fail': '未通过', 'warning': '警告'}.get(item.get('status'), item.get('status')),
                             item.get('label'), item.get('message')) for item in checks]
    jobinfo = result.get('job') or {}
    if jobinfo:
        v = jobinfo.get('verdict') or {}
        if v.get('summary'): lines.append('结论:%s' % v['summary'])
        if v.get('suggestion'): lines.append('建议:%s' % v['suggestion'])
        if v.get('last_error'): lines.append('最后错误:%s' % v['last_error'])
        for rec2 in (jobinfo.get('recent_diagnostics') or [])[-5:]:
            if isinstance(rec2, dict):
                lines.append('%s %s:%s' % (str(rec2.get('ts') or '')[:19], rec2.get('code') or '',
                                           str(rec2.get('detail') or '')[:160]))
    lines.append('完整数据请点「导出诊断包」。')
    result['technical_detail'] = '\n'.join(lines)
    return _redact_runtime(result, conf)

@app.get('/v1/diagnostics')
def one_click_diagnostics(jid: str = ''):
    return _diagnostic_snapshot(jid)

@app.post('/v1/diagnostics')
async def run_one_click_diagnostics(req: Request):
    body = await req.json()
    return _diagnostic_snapshot(str(body.get('job_id') or body.get('jid') or ''))

@app.get('/v1/jobs/{jid}/bundle')
def job_bundle(jid: str):
    job = jpath(jid)
    if not os.path.isdir(job):
        return JSONResponse({'ok': False, 'error': '任务不存在，无法导出诊断包'}, 404)
    data = diagnostic_bundle(job)
    return StreamingResponse(io.BytesIO(data), media_type='application/zip',
                             headers={'Content-Disposition': 'attachment; filename="bid-dog-diagnostic-%s.zip"' % os.path.basename(jid),
                                      'Content-Length': str(len(data))})

@app.post('/v1/jobs/{jid}/bundle/save')
def save_job_bundle(jid: str):
    """桌面 WebView 不保证 blob download 可用：落到任务目录并在系统文件管理器中定位。"""
    if MULTIUSER:
        return JSONResponse({'ok': False, 'error': '网页模式请使用诊断包下载链接'}, 403)
    job = jpath(jid)
    if not os.path.isdir(job):
        return JSONResponse({'ok': False, 'error': '任务不存在，无法导出诊断包'}, 404)
    name = '中标狗_诊断包_%s.zip' % datetime.datetime.now().strftime('%Y%m%d-%H%M%S')
    path = os.path.join(job, name)
    try:
        with open(path, 'wb') as f: f.write(diagnostic_bundle(job))
        _open_local(path, reveal=True)
        return {'ok': True, 'path': path, 'name': name}
    except Exception:
        return JSONResponse({'ok': False, 'error': '诊断包生成失败，请先打开任务文件夹后重试'}, 500)

@app.post('/v1/open_release')
async def open_release(req: Request):
    """仅允许打开本项目 Release，避免把本机 opener 变成任意 URL 跳转器。"""
    if MULTIUSER:
        return JSONResponse({'ok': False, 'error': '网页模式请直接点击下载链接'}, 403)
    body = await req.json()
    url = str(body.get('url') or '')
    if not re.match(r'^https://github\.com/shandianT/bid-dog/releases(?:/|$)', url, re.I):
        return JSONResponse({'ok': False, 'error': '下载地址未通过安全校验'}, 400)
    try:
        _open_local(url)
        return {'ok': True}
    except Exception:
        return JSONResponse({'ok': False, 'error': '系统浏览器没有打开，请手动前往 GitHub Releases'}, 500)

@app.post('/v1/jobs/{jid}/stop')
def stop_job(jid: str):
    """停止正在跑的任务(保留已生成的产物,不删任务)"""
    base = os.path.basename(jid); job = jpath(jid)
    if not os.path.isdir(job): return JSONResponse({'ok': False, 'error': 'not found'}, 404)
    control, owner = _begin_job_control(base)
    if not control:
        return JSONResponse({'ok': False, 'error': '任务正在执行另一个停止或删除操作，请稍后重试'}, 409)
    try:
        if not owner:
            meta_path = os.path.join(job, '任务.json')
            meta = read_json(meta_path, {})
            if meta.get('paused'):
                # 暂停后 worker 已退出，RUNNING 自然为空。“停止”必须把可续做会话
                # 转成普通停止态，否则它会永久占住全局模型切换锁。
                patch_json(meta_path, {'paused': False})
                emit(job, {'type': 'message', 'role': 'agent',
                           'text': '已结束暂停会话。已生成的内容都保留着，之后可以续做或重跑。'})
                halt(job, '已停止(手动)')
                return {'ok': True, 'note': '已结束暂停会话'}
            return {'ok': True, 'note': '任务本来就没有在运行'}
        ok, requested = _stop_running_owner(job, base, owner)
        if not ok:
            return JSONResponse({'ok': False, 'error': '任务暂时没有完全停止，已有内容仍在安全保留；请稍后重试'}, 502)
        if not requested:
            return {'ok': True, 'note': '任务刚刚已经结束'}
        patch_json(os.path.join(job, '任务.json'), {'paused': False})
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '已按你的要求停止。已生成的产物都保留着,可以「重跑」或「定向重做」。'})
        halt(job, '已停止(手动)')
        return {'ok': True}
    finally:
        _end_job_control(base, control)

def _redo_job_reserved(jid, job, instruction, owner):
    base = os.path.basename(jid)
    if _cancel_requested(base, owner):
        return _cancelled_before_dispatch()
    try: os.unlink(os.path.join(job, 'outcome.json'))
    except FileNotFoundError: pass
    meta0 = read_json(os.path.join(job, '任务.json'), {})
    close_pending_questions(job, meta0)
    # redo 是新一次执行尝试，不能继承旧轮次的技能读取证据或 OpenCode 会话。
    meta0['run_id'] = _new_run_id()
    meta0.pop('skill_manifest', None)
    meta0.pop('oc_session', None)
    meta0.pop('oc_questions', None)
    harvest(job)
    if _cancel_requested(base, owner):
        return _cancelled_before_dispatch()
    conf0 = read_json(conf_path(), {})
    up0 = s2_conf(conf0); eng0 = conf0.get('engine') or {}
    meta0['engine_snapshot'] = {'kind': eng0.get('kind', 's2'), 'model': up0['model'],
                                'base_url': up0['base_url'], 'wire': up0['wire'],
                                'verify_ssl': bool(up0['verify_ssl']),
                                'runtime_fingerprint': oc_config_fingerprint(conf0),
                                'started_at': now()}
    # 所有执行分支（包括 mock）都必须在派发前冻结旧产物；否则旧 Word 会替失败重做开门。
    meta0['redo_baseline'] = {
        'docx': {fn: _file_digest(os.path.join(job, fn)) for fn in _body_docxs(job)},
        'md': {fn: _file_digest(os.path.join(job, fn)) for fn in _body_mds(job)},
    }
    write_json(os.path.join(job, '任务.json'), meta0)
    agent_cmd = config_agent_cmd()
    if _cancel_requested(base, owner):
        return _cancelled_before_dispatch()
    if not agent_cmd:
        _start_reserved_worker(base, owner, mock_redo, job, instruction)
        return {'ok': True, 'mode': 'mock'}
    conf = read_json(conf_path(), {})
    meta = read_json(os.path.join(job, '任务.json'), {})
    tpath = os.path.join(job, meta.get('tender', ''))
    sd = skill_dir_conf(conf)
    if not os.path.isfile(os.path.join(sd, 'SKILL.md')): sd = ensure_skill()
    prompt = (AGENT_PROMPT.replace('{mode}', MODE_AGENTS).replace('{skill}', sd)
              + ' 【本次为定向重做】只执行这条指令:「%s」。保留任务目录内其他既有产物;'
                '改写对应文件后,必须重新执行汇总成册与自检体检,更新自检报告。' % instruction)
    # 定向重做也要用合并后的素材:以前这里写死 assets_dir(),于是重做时反而看不到
    # 这一单专门传进来的材料,改出来的章节和首轮不是一个事实来源
    sub = lambda s: (s.replace('{tender}', tpath).replace('{out}', job)
                      .replace('{materials}', merged_materials(job)).replace('{jobid}', os.path.basename(jid))
                      .replace('{skill}', sd))
    raw = agent_cmd
    if isinstance(raw, list):
        # claude/codex/sowork:命令里的长提示词参数(含 SKILL.md 路径)整体换成定向版
        cmd = [(sub(prompt) if ('SKILL.md' in a or len(a) > 300) else sub(a)) for a in raw]
    else:
        # 自定义命令:模板原样跑,定向指令按执行层合同写进 inbox 由 agent 消费
        open(os.path.join(job, 'inbox.jsonl'), 'a', encoding='utf-8').write(
            json.dumps({'type': 'redo', 'instruction': instruction}, ensure_ascii=False) + '\n')
        cmd = sub(raw)
    cmd = login_shell_wrap(cmd, conf.get('engine') or {})
    _set_skill_manifest(job, sd, cmd, _dispatch_contains_skill(cmd, sd), 'cli')
    emit(job, {'type': 'progress', 'stage': '定向重做启动', 'pct': 5, 'step': 1, 'total': 12})
    if _cancel_requested(base, owner):
        return _cancelled_before_dispatch()
    _start_reserved_worker(base, owner, real_agent, job, cmd)
    return {'ok': True, 'mode': 'agent'}

@app.post('/v1/jobs/{jid}/redo')
async def redo_job(jid: str, req: Request):
    """定向重做:在当前任务里只重做用户指定的部分,其余产物保留,完成后重新汇总自检"""
    gate = _generation_gate_response()
    if gate is not None: return gate
    body = await req.json()
    instruction = (body.get('instruction') or '').strip()
    job = jpath(jid)
    if not os.path.isdir(job): return JSONResponse({'ok': False, 'error': 'not found'}, 404)
    if not instruction: return JSONResponse({'ok': False, 'error': '缺少重做指令'}, 400)
    base = os.path.basename(jid)
    owner, reason = _reserve_running_reason(base)
    if not owner:
        return _admission_failure(reason, '任务正在运行，等它停下后再定向重做。')
    try:
        result = _redo_job_reserved(jid, job, instruction, owner)
        if isinstance(result, dict) and result.get('mode') in ('error', 'cancelled'):
            _release_running(base, owner)
        return _launch_http_result(result)
    except Exception:
        _release_running(base, owner)
        raise

def route_command(job, jid, text, running):
    """对话即遥控器:把自然语言识别成任务指令。可逆的(暂停/继续)直接执行;
    重活(整任务重启/定向重做)回一条带确认按钮的消息,点了才动手——不静默执行。
    疑问句一律不当指令,落回正常问答。"""
    t = (text or '').strip()
    if not t or t.endswith(('?', '?', '吗', '么')): return False
    if running and re.match(r'^(暂停|停一下|先停|暂停一下|先暂停)', t):
        # 聊天路由没有资格只改 paused 标志就宣称真外壳停了；真实暂停必须走 /control，
        # 由它确认 OpenCode interrupt 或明确告知兼容模式无法暂停。
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '收到暂停请求。请点任务标题旁的“暂停按钮”执行；系统会先确认任务真的停下。'
                           '如果当前运行方式不支持暂停，会明确提示你改用“停止”。'})
        return True
    if running and re.match(r'^(继续|接着跑|恢复|继续生成|接着写)', t):
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '请点任务标题旁的“继续做”恢复任务；确认接续成功后才会继续显示生成中。'})
        return True
    targeted = re.search(r'第.{1,6}(章|节|步|部分)|某一步|目录|封面|报价|偏离表|自检', t)
    if re.search(r'重新启动|重新生成|整个重|重跑|再来一遍|全部重来', t) and not targeted:
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '要用同一份招标文件重新生成整个任务吗?会新开一个任务开始跑,当前任务和产物都保留。',
                   'actions': [{'act': 'rerun', 'label': '确认重新生成'}]})
        return True
    if (not running) and (targeted or re.match(r'^(重写|改写|扩写|重做|补写|重新输出|修改)', t)) \
            and re.search(r'重写|改写|扩写|重做|补写|重新输出|修改|重新写', t):
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '收到定向指令:「%s」。确认后我在当前任务里只重做这一部分,其余产物保留,完成后重新汇总并更新自检。' % t[:80],
                   'actions': [{'act': 'redo', 'label': '开始定向重做', 'param': t[:500]}]})
        return True
    return False

@app.post('/v1/jobs/{jid}/messages')
async def message(jid: str, req: Request):
    body = await req.json()
    job = jpath(jid); text = body.get('content', '')
    emit(job, {'type': 'message', 'role': 'user', 'text': text})
    if route_command(job, jid, text, _is_running(os.path.basename(jid))):
        return {'ok': True}
    if not _is_running(os.path.basename(jid)):
        # 任务已结束:不管生成引擎是哪种,都用已配模型基于产出回答(没配模型则给出指引);
        # thinking 状态让前端立刻显示"正在回复…",答案/失败原因随后必到
        emit(job, {'type': 'status', 'state': 'thinking'})
        threading.Thread(target=chat_reply, args=(job, jid, text), daemon=True).start()
    elif not config_agent_cmd():
        emit(job, {'type': 'message', 'role': 'agent', 'text': '收到:「%s」。已纳入当前章节,继续推进。' % text})
    else:
        open(os.path.join(job, 'inbox.jsonl'), 'a', encoding='utf-8').write(json.dumps(body, ensure_ascii=False) + '\n')
        # 运行期插话:server 模式下用 OpenCode 原生的 delivery=queue —— 实测语义是
        # 「排到当前这一步之后,不打断」,正是产品负责人定的「想到什么先记下来、下个环节带上」。
        # 以前这里只写 inbox.jsonl 然后回一句「已送达正在写标书的 agent」,
        # 而那个文件**没有任何消费者**,话 100% 丢失(七条假承诺里的第一条)。
        sid = (read_json(os.path.join(job, '任务.json'), {}) or {}).get('oc_session')
        queued = False
        if sid and OC.get('base'):
            queued, _b = oc_send(sid, text, delivery='queue')
        emit(job, {'type': 'message', 'role': 'sys',
                   'text': ('已排进队列 —— 它会在跑完手上这一步之后带上这条,不会打断当前工作。'
                            if queued else
                            '已记下这条。当前使用稳定运行方式，生成助手中途无法接收——'
                            '它会在下次「重跑」或「定向重做」时带上。')})
    return {'ok': True}

@app.post('/v1/jobs/{jid}/answers')
async def answer(jid: str, req: Request):
    """回答 agent 的提问。

    以前这里只把答案写进 answers.jsonl —— 而那个文件**没有任何消费者**,
    agent 一个字都收不到,界面却回一句「好,按 X 处理,继续」。
    现在走 OpenCode 的 /question/{id}/reply,答案真的进它的上下文;
    回传失败就当面说,不再假装已生效。"""
    body = await req.json(); job = jpath(jid)
    rid = body.get('question_id') or ''
    txt = body.get('choice') or body.get('text', '')
    emit(job, {'type': 'message', 'role': 'user', 'text': txt})
    open(os.path.join(job, 'answers.jsonl'), 'a', encoding='utf-8').write(
        json.dumps(body, ensure_ascii=False) + '\n')     # 留档,便于事后追溯
    if rid.startswith('confirm_parse:'):
        # 解析确认门:确认或修正都在这里落地,修正会进入「你的要求.md」,
        # 响应规划的输入摘要因此变化,重启 worker 后会按修正重新规划再开写。
        accept = txt.strip() in (CONFIRM_PARSE_ACCEPT, '确认无误', '确认', '')
        corrections = '' if accept else txt.strip()[:2000]
        if corrections:
            with open(os.path.join(job, '你的要求.md'), 'a', encoding='utf-8') as handle:
                handle.write('\n\n## 解析确认修正（%s）\n\n%s\n' % (now(), corrections))
        write_json(os.path.join(job, '解析确认.json'),
                   {'confirmed': True, 'ts': now(), 'corrections': corrections})
        emit(job, {'type': 'question_closed', 'id': rid})
        base = os.path.basename(jid)
        owner, reason = _reserve_running_reason(base)
        if not owner:
            if reason == 'running':
                return {'ok': True, 'delivered': True}   # worker 已在跑,确认文件会被下一步读取
            emit(job, {'type': 'message', 'role': 'agent',
                       'text': '确认已记录,但任务当前无法继续(%s)。稍后点「继续」即可从检查点接着跑。' % reason})
            return JSONResponse({'ok': False, 'error': '任务当前无法继续：%s' % reason}, 409)
        try: os.unlink(os.path.join(job, 'outcome.json'))
        except FileNotFoundError: pass
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': ('好，已确认关键信息，开始分章撰写。' if not corrections else
                            '已记下修正并同步进任务要求；会按修正重新核对响应规划后开始撰写。')})
        _start_reserved_worker(base, owner, generation_pipeline_worker, job)
        return {'ok': True, 'delivered': True}
    meta = read_json(os.path.join(job, '任务.json'), {})
    if rid and rid in (meta.get('oc_questions') or {}):
        ok, why = oc_answer(job, rid, txt)
        if ok:
            emit(job, {'type': 'question_closed', 'id': rid})
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': ('好,按「%s」处理,已交给正在写标书的生成助手,它会接着做。' % txt[:40]) if ok
                           else ('⚠ 答案没能交给生成助手(%s)。它可能还在等,或已经按自己的判断继续了——'
                                 '建议记下这条,跑完用「定向重做」把它落实。' % why)})
        return {'ok': ok, 'delivered': ok, 'error': ('' if ok else why)}
    # 不是 agent 主动问的(比如内置演示,或历史遗留的问题)
    emit(job, {'type': 'question_closed', 'id': rid})
    running = job_state(job) == 'running'
    emit(job, {'type': 'message', 'role': 'agent',
               'text': ('好,按「%s」处理,继续。' % txt[:40]) if not running else
                       ('已记下「%s」。这条不是生成助手当前在等的问题,它会在下次重跑或定向重做时带上。' % txt[:40])})
    return {'ok': True, 'delivered': False}

@app.post('/v1/jobs/{jid}/control')
async def control(jid: str, req: Request):
    """暂停 / 继续。

    以前这里只是把 paused 标志位写进 任务.json,然后回一句「已暂停,随时继续」——
    而**只有演示引擎读那个标志位**。真跑的时候 agent 一个字都不知道,照写不误:
    界面说停了,它还在花钱、还在改文件(七条假承诺里的最后一条)。

    真 agent 没有「挂起」这回事,能做到的是「停下来 + 记住 + 接着做」:
    暂停 = 优雅中断当前这轮(半截产物完整落盘),继续 = 回原会话往下接。
    做不到的情况(没有会话/外壳没起来)当面说清,不假装暂停成功。"""
    body = await req.json(); job = jpath(jid)
    base = os.path.basename(jid)
    want_pause = body.get('action') == 'pause'
    if not want_pause:                       # 继续 = 走「接着做」那条路
        return await resume_job(jid)
    control, owner = _begin_job_control(base)
    if not control:
        return JSONResponse({'ok': False, 'error': '任务正在执行另一个停止或删除操作，请稍后重试'}, 409)
    try:
        meta = read_json(os.path.join(job, '任务.json'), {})
        if not owner:
            return JSONResponse({'ok': False, 'error': '任务当前没有在运行，不需要暂停'}, 409)
        runtime = job_runtime(job, meta)
        pause_cap = (runtime.get('capabilities') or {}).get('pause') or {}
        if not pause_cap.get('enabled'):
            reason = pause_cap.get('reason') or '当前运行方式暂不支持暂停；如需中止，请使用停止。'
            emit(job, {'type': 'message', 'role': 'agent', 'text': reason})
            return JSONResponse({'ok': False, 'error': reason, 'capability': 'pause'}, 400)
        if not config_agent_cmd():
            # mock worker 只能轮询 paused 标志且仍占 RUNNING；若在这里假暂停，resume 会因
            # 同一 owner 尚在而永远进不去。演示流程明确不支持暂停，比制造死锁可靠。
            emit(job, {'type': 'message', 'role': 'agent',
                       'text': '内置演示流程不支持暂停；可以让它很快跑完，或使用“停止”结束本单。'})
            return JSONResponse({'ok': False, 'error': '内置演示流程不支持暂停，请改用停止'}, 400)
        sid = meta.get('oc_session')
        if not (sid and OC.get('base') and _interrupt_or_finished(sid)):
            emit(job, {'type': 'message', 'role': 'agent',
                       'text': '⚠ 当前运行方式暂不支持暂停。'
                               '它还在继续写 —— 要真停请用「停止」,那会结束整单。'})
            return JSONResponse({'ok': False, 'error': '当前这轮不支持暂停'}, 400)
        if not _request_cancel(base, owner):
            return {'ok': True, 'note': '任务刚刚已经结束'}
        meta['paused'] = True
        write_json(os.path.join(job, '任务.json'), meta)
        emit(job, {'type': 'message', 'role': 'agent',
                   'text': '已停下来。写到哪儿就停在哪儿,产物都在;点「继续做」它会回到原来的思路往下接。',
                   'actions': [{'act': 'resume', 'label': '继续做'}]})
        halt(job, '已暂停')
        return {'ok': True}
    finally:
        _end_job_control(base, control)

@app.get('/v1/jobs/{jid}/artifacts')
def artifacts(jid: str):
    job = jpath(jid); out = []
    if os.path.isdir(job):
        for fn in list_deliverables(job):
            info = artifact_info(fn)
            out.append({'name': fn, 'url': '/v1/jobs/%s/artifacts/%s' % (jid, fn),
                        'size_kb': round(os.path.getsize(os.path.join(job, fn)) / 1024, 1), **info})
    return sorted(out, key=lambda a: (a['group'], a['rank'], a['name']))

@app.get('/v1/jobs/{jid}/artifacts/{fn}')
def download(jid: str, fn: str):
    path = _artifact_path(jid, fn)
    if not path: return JSONResponse({'ok': False, 'error': '文件不存在'}, 404)
    return FileResponse(path, filename=os.path.basename(path))

def _artifact_path(jid, name):
    """只允许访问任务根目录中已登记的交付物，阻止 ../ 与任意本机路径。"""
    raw = str(name or '')
    safe = os.path.basename(raw)
    job = jpath(jid)
    if not safe or safe != raw or not os.path.isdir(job) or safe not in list_deliverables(job): return None
    path = os.path.realpath(os.path.join(job, safe))
    return path if os.path.dirname(path) == os.path.realpath(job) and os.path.isfile(path) else None

def _open_local(path, reveal=False):
    if sys.platform == 'darwin': subprocess.Popen(['open', '-R', path] if reveal else ['open', path])
    elif os.name == 'nt':
        if reveal: subprocess.Popen(['explorer', '/select,%s' % path])
        else: os.startfile(path)  # noqa
    else: subprocess.Popen(['xdg-open', os.path.dirname(path) if reveal else path])

@app.post('/v1/jobs/{jid}/artifacts/open')
async def open_artifact(jid: str, req: Request):
    if MULTIUSER: return JSONResponse({'ok': False, 'error': '云端模式不支持打开本机文件'}, 403)
    body = await req.json(); path = _artifact_path(jid, body.get('name'))
    if not path: return JSONResponse({'ok': False, 'error': '文件不存在或不属于当前任务'}, 404)
    try:
        _open_local(path, bool(body.get('reveal')))
        return {'ok': True, 'path': path}
    except Exception as e:
        return JSONResponse({'ok': False, 'error': str(e), 'path': path}, 500)

@app.post('/v1/jobs/{jid}/open_folder')
def open_job_folder(jid: str):
    if MULTIUSER: return JSONResponse({'ok': False, 'error': '云端模式不支持打开本机文件夹'}, 403)
    job = jpath(jid)
    if not os.path.isdir(job): return JSONResponse({'ok': False, 'error': '任务目录不存在'}, 404)
    try:
        _open_local(job)
        return {'ok': True, 'path': job}
    except Exception as e:
        return JSONResponse({'ok': False, 'error': str(e), 'path': job}, 500)

@app.get('/v1/providers')
def providers():
    stored = read_json(conf_path(), {'providers': [], 'routing': {}}).get('providers', [])
    allowed = ('id', 'name', 'base_url', 'model', 'vision_model', 'kind', 'verify_ssl')
    public = []
    for item in stored:
        safe = {key: item.get(key) for key in allowed if key in item}
        safe['key_set'] = bool(item.get('api_key'))
        public.append(safe)
    return public

@app.delete('/v1/providers/{pid}')
def del_provider(pid: str):
    conf = read_json(conf_path(), {'providers': [], 'routing': {}})
    conf['providers'] = [p for p in conf['providers'] if p.get('id') != pid]
    if (conf.get('routing') or {}).get('default') == pid:
        conf['routing'] = {'default': conf['providers'][0]['id'], 'model': conf['providers'][0].get('model', '')} if conf['providers'] else {}
    write_json(conf_path(), conf)
    return {'ok': True}

@app.post('/v1/providers')
async def add_provider(req: Request):
    body = await req.json()
    conf = read_json(conf_path(), {'providers': [], 'routing': {}})
    body['id'] = body.get('id') or uuid.uuid4().hex[:8]
    # 同一网关+同一模型视为同一个接入点:覆盖而不是新增(避免重复添加堆积一长串)
    same = lambda p: (p.get('base_url'), p.get('model', '')) == (body.get('base_url'), body.get('model', ''))
    dup = next((p for p in conf['providers'] if same(p)), None)
    if dup: body['id'] = dup['id']
    # 生产环境:api_key 应转存系统钥匙串;此处存本地配置仅为最小可用
    conf['providers'] = [p for p in conf['providers'] if p.get('id') != body['id'] and not same(p)] + [body]
    if not (conf.get('routing') or {}).get('default'):
        conf['routing'] = {'default': body['id'], 'model': body.get('model', '')}
    write_json(conf_path(), conf)
    return {'id': body['id']}

RETRY_WAITS = (0.8, 2.0, 4.5)      # 三次重试的等待,总计约 7 秒——网关瞬时抖动基本都在这个窗口内恢复

def _is_transient(e):
    """能靠重试自愈的网络故障。判定要严:把 401/404/证书错误当瞬时错误重试是白烧时间。

    实测最常见的两个(客户机 → S2 网关):
      <urlopen error [Errno 54] Connection reset by peer>
      <urlopen error [SSL: UNEXPECTED_EOF_WHILE_READING] EOF occurred in violation of protocol>
    两条都是「连接建立/握手阶段被对端掐断」,重开一条连接就好——网关侧短时限流或中间设备干扰。
    """
    if isinstance(e, urllib.error.HTTPError):
        return e.code in (429, 500, 502, 503, 504)
    reason = getattr(e, 'reason', e)
    if isinstance(reason, ssl.SSLCertVerificationError):
        return False                                    # 证书问题重试一万次也一样
    if isinstance(reason, (ConnectionResetError, ConnectionAbortedError, BrokenPipeError,
                           ConnectionRefusedError, socket.gaierror,
                           ssl.SSLEOFError, ssl.SSLZeroReturnError, TimeoutError, socket.timeout)):
        # ConnectionRefused/DNS 失败也算瞬时:路由器重连、VPN 抖动、网关重启的窗口期
        # 都表现为这两种,重试几秒后大多自愈;以前不重试,一次抖动直接把节点打停。
        return True
    if isinstance(reason, (http.client.RemoteDisconnected, http.client.IncompleteRead,
                           http.client.BadStatusLine)):
        return True
    s = str(reason).lower()
    return any(k in s for k in ('connection reset', 'unexpected_eof', 'eof occurred in violation',
                                'remote end closed', 'connection aborted', 'broken pipe',
                                'record layer failure', 'timed out', 'connection refused',
                                'getaddrinfo failed', 'name or service not known',
                                'temporary failure in name resolution'))

def _retry(call, tries=len(RETRY_WAITS) + 1, on_wait=None):
    """瞬时网络故障重试。上游网关偶发掐连接是常态,不重试就等于把整条任务赌在一次握手上。"""
    last = None
    for i in range(tries):
        try:
            return call()
        except Exception as e:
            last = e
            if i == tries - 1 or not _is_transient(e): raise
            if on_wait: on_wait(i + 1, tries - 1, e)
            time.sleep(RETRY_WAITS[i])
    raise last

def net_hint(e, secrets_to_hide=()):
    """把底层网络异常翻译成可操作的中文提示"""
    s = _safe_secret_text(e, secrets_to_hide)
    if 'Connection reset by peer' in s or 'unexpected_eof' in s.lower() \
            or 'EOF occurred in violation' in s or 'RemoteDisconnected' in s:
        return ('连接被对端掐断(已自动重试 %d 次仍未成功)。多半是网关侧短时限流,或公司网络/VPN、'
                '安全软件在拦截长连接。可以:① 过一两分钟点「重跑本任务」;② 换个网络(手机热点)试一次'
                '——能通就是本地网络设备的问题;③ 若你是网关方,看下并发连接数限制。原始信息:%s'
                % (len(RETRY_WAITS), s))
    if 'CERTIFICATE_VERIFY' in s or 'certificate verify failed' in s:
        return ('HTTPS 证书校验失败(已尝试系统证书与内置根证书)。常见于公司网络代理拦截 HTTPS 或内网自签名网关;'
                '若确认该网关可信,勾选「跳过证书校验」后重试。原始信息:%s' % s)
    if 'Name or service not known' in s or 'nodename nor servname' in s or 'getaddrinfo' in s:
        return '域名解析失败:请检查 Base URL 拼写与本机网络/VPN。原始信息:%s' % s
    if 'timed out' in s or 'timeout' in s.lower():
        return '连接超时:网关不可达或网络受限(公司网络可能需代理)。原始信息:%s' % s
    if 'Connection refused' in s:
        return '连接被拒绝:端口/地址不对,或服务未启动。原始信息:%s' % s
    return s

def _ssl_ctx(kind):
    """default=系统/环境 CA;certifi=内置根证书(打包版 macOS 必需);none=不校验(仅内网自签名)"""
    if kind == 'none':
        c = ssl.create_default_context(); c.check_hostname = False; c.verify_mode = ssl.CERT_NONE; return c
    if kind == 'certifi':
        import certifi
        return ssl.create_default_context(cafile=certifi.where())
    return ssl.create_default_context()

def _openai_req(base, key, path, payload=None, timeout=30, verify=True):
    hdr = {'Authorization': 'Bearer ' + (key or ''), 'User-Agent': 'bid-assistant/0.9'}
    data = None
    if payload is not None:
        hdr['Content-Type'] = 'application/json'
        data = json.dumps(payload, ensure_ascii=False).encode('utf-8')
    req = urllib.request.Request((base or '').rstrip('/') + path, data=data, headers=hdr)
    kinds = ['default', 'certifi'] if verify else ['none']
    last = None
    for kind in kinds:
        try:
            ctx = _ssl_ctx(kind)
        except Exception:
            continue  # 该环境没有 certifi
        try:
            # 瞬时故障(连接被重置/握手 EOF/网关 5xx)自动重试,不重试就等于把一次对话赌在一次握手上
            r = _retry(lambda: urllib.request.urlopen(req, timeout=timeout, context=ctx))
            return json.loads(r.read().decode('utf-8', 'ignore'))
        except urllib.error.URLError as e:
            # 仅证书类失败才换 CA 重试;其他错误(401/超时/DNS)直接抛
            if not isinstance(getattr(e, 'reason', None), ssl.SSLCertVerificationError): raise
            last = e
    raise last

@app.post('/v1/providers/probe_models')
async def probe_models(req: Request):
    """不落库:用 base_url+key 拉取该网关可用模型列表(供添加前选择你套餐里的模型)"""
    body = await req.json()
    if not body.get('base_url'): return JSONResponse({'ok': False, 'error': '缺 base_url', 'models': []}, 400)
    base, key, verify = body['base_url'], body.get('api_key', ''), body.get('verify_ssl', True)
    hit = _models_cached(base, key, verify)
    if hit: return hit
    try:
        # 阻塞网络请求丢线程池:这是 async 端点,直接跑会卡住整个事件循环(页面全体转圈的元凶之一)
        data = await asyncio.to_thread(_openai_req, base, key, '/models', timeout=15, verify=verify)
        ids = [m.get('id') for m in (data.get('data') or []) if m.get('id')]
        _models_store(base, key, verify, ids)
        return {'ok': True, 'models': ids}
    except urllib.error.HTTPError as e:
        return {'ok': False, 'error': 'HTTP %s %s' % (e.code, _http_error_detail(e, (key,), 200)), 'models': []}
    except Exception as e:
        return {'ok': False, 'error': net_hint(e, (key,)), 'models': []}

@app.post('/v1/providers/{pid}/test')
def test_provider(pid: str):
    """配了模型→真发一次对话请求(验证你的 token 套餐对该模型可用);没配→仅探 /models 连通"""
    conf = read_json(conf_path(), {'providers': []})
    p = next((x for x in conf['providers'] if x.get('id') == pid), None)
    if not p: return JSONResponse({'ok': False, 'error': 'not found'}, 404)
    t0 = time.time(); vs = p.get('verify_ssl', True)
    try:
        if p.get('model'):
            resp = _openai_req(p.get('base_url'), p.get('api_key'), '/chat/completions',
                               {'model': p['model'], 'messages': [{'role': 'user', 'content': 'ping,只回复:pong'}], 'max_tokens': 8},
                               timeout=45, verify=vs)
            reply = ((resp.get('choices') or [{}])[0].get('message') or {}).get('content', '')
            return {'ok': True, 'latency_ms': int((time.time() - t0) * 1000), 'model': resp.get('model', p['model']), 'reply': (reply or '')[:40]}
        _openai_req(p.get('base_url'), p.get('api_key'), '/models', timeout=15, verify=vs)
        return {'ok': True, 'latency_ms': int((time.time() - t0) * 1000), 'note': '通道连通;未配模型,建议选择模型后重测'}
    except urllib.error.HTTPError as e:
        return {'ok': False, 'error': 'HTTP %s %s' % (e.code, _http_error_detail(e, (p.get('api_key'),), 300))}
    except Exception as e:
        return {'ok': False, 'error': net_hint(e, (p.get('api_key'),))}

# ================= Responses ↔ Chat 中转:让 Codex CLI 直接用我们自己的 S2 网关 =================
# 背景(实测,codex-cli 0.146.0):Codex 支持自定义模型供应商(CODEX_HOME/config.toml 里的
# [model_providers.*] + base_url + env_key),但**只认 Responses API**——`wire_api = "chat"`
# 已被移除,填了会直接报错退出。而我们的 S2 网关(以及绝大多数国产 OpenAI 兼容网关)只有
# /chat/completions。所以这里做一层协议翻译,并且刻意放进「已经随 App 分发的这个引擎」里:
# 不多装一个进程、不多开一个端口、不改客户机器的 ~/.codex —— 客户那边只是多填一个 Key。
S2_DEFAULT_BASE = 'https://api.senseaudio.cn/v1'
S2_DEFAULT_MODEL = 'deepseek-v4-flash'
S2_QUALITY_MODEL = 'senseaudio-s2'
SELF_PORT = int(os.environ.get('PORT', 8080))
RELAY_LAST = {}      # 最近一次中转的结果:出问题时「测试连接」和运行日志能说清卡在哪一层

# 模型列表短缓存:重复打开设置面板不再反复打网关(慢且费额度);TTL 默认 120s,BID_MODELS_TTL=0 关闭
_MODELS_CACHE, _MODELS_LOCK = {}, threading.Lock()
_MODELS_TTL = int(os.environ.get('BID_MODELS_TTL', 120))

def _models_cached(base, key, verify):
    if _MODELS_TTL <= 0: return None
    k = (base or '').rstrip('/') + '|' + hashlib.sha1((key or '').encode()).hexdigest()[:12] + '|' + str(bool(verify))
    with _MODELS_LOCK:
        v = _MODELS_CACHE.get(k)
        if v and time.time() - v[0] < _MODELS_TTL:
            return {'ok': True, 'models': list(v[1]), 'cached': True}
    return None

def _models_store(base, key, verify, ids):
    k = (base or '').rstrip('/') + '|' + hashlib.sha1((key or '').encode()).hexdigest()[:12] + '|' + str(bool(verify))
    with _MODELS_LOCK:
        _MODELS_CACHE[k] = (time.time(), list(ids))

def relay_token():
    """本机中转口令:只有我们自己拉起的 Codex 拿得到,别的进程调不动你的 Key"""
    conf = read_json(conf_path(), {})
    t = (conf.get('relay_token') or '').strip()
    if not t:
        t = secrets.token_hex(16)
        conf['relay_token'] = t
        write_json(conf_path(), conf)
    return t

def s2_conf(conf=None):
    """S2 上游:优先取「生成引擎」里填的;留空就借用「模型接入」里已经配好的那个网关,免得同一个 Key 填两遍"""
    conf = conf if conf is not None else read_json(conf_path(), {})
    eng = conf.get('engine') or {}
    base = (eng.get('s2_base_url') or '').strip()
    key = (eng.get('s2_key') or '').strip()
    model = (eng.get('s2_model') or '').strip()
    if not (base and key):
        ps = conf.get('providers') or []
        rid = (conf.get('routing') or {}).get('default')
        p = next((x for x in ps if x.get('id') == rid), None) or (ps[0] if ps else None)
        if p:
            base, key = base or (p.get('base_url') or '').strip(), key or (p.get('api_key') or '').strip()
            model = model or (p.get('model') or '').strip()
    return {'base_url': (base or S2_DEFAULT_BASE).rstrip('/'), 'api_key': key,
            'model': model or S2_DEFAULT_MODEL, 'verify_ssl': eng.get('s2_verify_ssl', True),
            'wire': (eng.get('s2_wire') or 'auto')}

def _key_fingerprint(key):
    return hashlib.sha256(str(key or '').encode()).hexdigest() if key else ''

def _connection_fingerprint(conf_or_up):
    """Bind a successful model probe to the whole transport identity, not only its Key."""
    up = (conf_or_up if isinstance(conf_or_up, dict) and 'api_key' in conf_or_up
          else s2_conf(conf_or_up if isinstance(conf_or_up, dict) else None))
    identity = {
        # 完整传输身份只以摘要形式持久化；查询参数或 userinfo
        # 变化也会使旧验证失效，同时不会把凭据写入配置摘要。
        'base_url_identity': hashlib.sha256(
            str(up.get('base_url') or '').strip().rstrip('/').encode('utf-8')).hexdigest(),
        'key_fingerprint': _key_fingerprint(up.get('api_key') or ''),
        'verify_ssl': bool(up.get('verify_ssl', True)),
        'wire': str(up.get('wire') or 'auto'),
        'model': str(up.get('model') or ''),
    }
    return hashlib.sha256(json.dumps(identity, sort_keys=True, separators=(',', ':')).encode()).hexdigest()

def _verified_model_ids(conf):
    setup = conf.get('setup') if isinstance(conf.get('setup'), dict) else {}
    tested = str(setup.get('tested_connection_fingerprint') or '')
    current = _connection_fingerprint(conf)
    if not tested or not hmac.compare_digest(current, tested):
        return []
    return [str(value) for value in (setup.get('text_verified_model_ids') or []) if value]


def _record_text_model_verification(expected_conf, model_ids, verified_ids):
    """仅在当前连接仍与测试快照一致时落盘，避免慢测试覆盖新设置。"""
    path = conf_path()
    expected = _connection_fingerprint(expected_conf)
    with _json_lock(path):
        current = read_json(path, {})
        current_fp = _connection_fingerprint(current)
        if not hmac.compare_digest(expected, current_fp):
            return False
        setup = dict(current.get('setup') or {})
        setup.update({
            'tested_key_fingerprint': _key_fingerprint(s2_conf(current).get('api_key') or ''),
            'tested_connection_fingerprint': current_fp,
            'connected_at': now(),
            'model_ids': [str(value) for value in (model_ids or []) if value][:200],
            'text_verified_model_ids': [str(value) for value in (verified_ids or []) if value][:20],
        })
        current['setup'] = setup
        write_json(path, current)
    return True

def _has_existing_jobs():
    try:
        return any(os.path.isdir(jpath(jid)) for jid in os.listdir(jobs_dir()))
    except OSError:
        return False

def _legacy_setup_configured(conf):
    eng = conf.get('engine') if isinstance(conf.get('engine'), dict) else {}
    kind = str(eng.get('kind') or '')
    return bool(s2_conf(conf).get('api_key') or conf.get('providers')
                or kind in ('claude', 'codex', 'sowork', 'custom'))

def setup_status_data(conf=None):
    conf = conf if isinstance(conf, dict) else read_json(conf_path(), {})
    setup = conf.get('setup') if isinstance(conf.get('setup'), dict) else {}
    key_set = bool(s2_conf(conf).get('api_key'))
    current_fp = _connection_fingerprint(conf)
    tested_fp = str(setup.get('tested_connection_fingerprint') or '')
    legacy_skipped = not setup and (_has_existing_jobs() or _legacy_setup_configured(conf))
    connected = bool(key_set and tested_fp and hmac.compare_digest(current_fp, tested_fp))
    completed = bool(legacy_skipped or (setup.get('completed_at') and connected))
    return {'needed': not completed, 'completed': completed, 'connected': connected,
            'key_set': key_set, 'legacy_skipped': legacy_skipped,
            'recommended': {'engine': 'opencode', 'mode': 'fast',
                            'model': S2_DEFAULT_MODEL},
            'steps': ['connect', 'create_first_job']}

def setup_connection_probe(candidate_conf):
    """Validate the candidate credential before any secret is committed."""
    up = s2_conf(candidate_conf)
    if not up.get('api_key'): return False, '请先填写 Key', []
    try:
        data = _openai_req(up['base_url'], up['api_key'], '/models', timeout=20,
                           verify=up.get('verify_ssl', True))
        ids = [str(item.get('id')) for item in (data.get('data') or [])
               if isinstance(item, dict) and item.get('id')]
        generation, _vision = _setup_models(ids, up.get('model') or S2_DEFAULT_MODEL)
        if not generation:
            return False, '网关只返回了向量/语音/图像模型，没有可用的文本生成模型', ids
        targets = []
        if _pipeline_mode(candidate_conf) == 'standard' and S2_DEFAULT_MODEL in ids:
            targets.append(S2_DEFAULT_MODEL)
        if generation not in targets:
            targets.append(generation)
        verified = []
        for model_id in targets:
            probe = _openai_req(
                up['base_url'], up['api_key'], '/chat/completions',
                {'model': model_id,
                 'messages': [{'role': 'user', 'content': 'ping,只回复:pong'}],
                 'max_tokens': 8, 'stream': False},
                timeout=45, verify=up.get('verify_ssl', True))
            choices = probe.get('choices') if isinstance(probe, dict) else None
            reply = (((choices or [{}])[0].get('message') or {}).get('content')
                     if choices else '')
            if not str(reply or '').strip():
                return False, '模型“%s”已连接，但未收到有效文本响应' % model_id, ids
            verified.append(model_id)
        return True, '', ids, verified
    except urllib.error.HTTPError as error:
        return False, {401: 'Key 无效或已停用', 403: 'Key 没有访问权限',
                       429: '当前请求较多，请稍后重试'}.get(error.code, '连接测试失败（HTTP %s）' % error.code), []
    except Exception as error:
        return False, '连接测试失败：%s' % net_hint(error, (up.get('api_key'),)), []

def _setup_models(model_ids, preferred):
    ids = [str(value) for value in (model_ids or []) if value]
    vision = next((value for value in ids if re.search(r'vision|(?:^|[-_])vl(?:[-_]|$)|gpt-4o|gemini', value, re.I)), '')
    excluded = re.compile(r'embed(?:ding)?|rerank|tts|speech|whisper|image|vision|(?:^|[-_])vl(?:[-_]|$)', re.I)
    formal = [value for value in ids if not excluded.search(value)]
    if not ids:
        return '', vision
    generation = (preferred if preferred in formal else
                  (S2_DEFAULT_MODEL if S2_DEFAULT_MODEL in formal else (formal[0] if formal else '')))
    return generation, vision

@app.get('/v1/setup')
def setup_status():
    return setup_status_data()

@app.post('/v1/setup/connect')
async def setup_connect(req: Request):
    if MULTIUSER and not ALLOW_AGENT_CONFIG:
        return JSONResponse({'ok': False, 'error': '当前部署由管理员统一配置'}, 403)
    body = await req.json()
    requested_mode = 'quality' if str(body.get('mode') or '').strip().lower() == 'quality' else 'fast'
    requested_model = str(body.get('model') or (
        S2_QUALITY_MODEL if requested_mode == 'quality' else S2_DEFAULT_MODEL)).strip()
    path = conf_path()
    with _json_lock(path):
        current = read_json(path, {})
        current_up = s2_conf(current)
        key = str(body.get('key') or current_up.get('api_key') or '').strip()
        if not key: return JSONResponse({'ok': False, 'error': '请先填写 Key'}, 400)
        engine = dict(current.get('engine') or {})
        engine.update({'kind': 's2', 'mode': 'agents', 'generation_mode': (
                           'standard' if requested_mode == 'quality' else 'fast'),
                       's2_key': key,
                       's2_base_url': str(body.get('base_url') or S2_DEFAULT_BASE).strip().rstrip('/'),
                       's2_model': requested_model,
                       's2_wire': 'auto', 's2_verify_ssl': bool(body.get('verify_ssl', True)),
                       'login_shell': True})
        candidate = dict(current); candidate['engine'] = engine
        if config_locked_jobs() and oc_config_fingerprint(candidate) != oc_config_fingerprint(current):
            return JSONResponse({'ok': False, 'error': '还有任务正在生成，请结束后再更换连接'}, 409)
        config_before_probe = _sha256_text(json.dumps(current, ensure_ascii=False, sort_keys=True))
    # 网络探测是同步 I/O，必须移出配置锁并放到工作线程；否则 Windows 上一次慢连接
    # 会同时冻住健康检查、进度刷新和设置保存，看起来像“本地引擎断了”。
    probe_result = await asyncio.to_thread(setup_connection_probe, candidate)
    ok, why = probe_result[:2]
    model_ids = probe_result[2] if len(probe_result) > 2 else []
    if not ok:
        safe = _safe_secret_text(why, (key,))
        return JSONResponse({'ok': False, 'connected': False,
                             'error': safe or '连接测试没有通过'}, 400)
    if not model_ids:
        return JSONResponse({'ok': False, 'connected': False,
                             'error': '网关未返回可用模型，设置未保存；请检查 base URL 或网关权限'}, 400)
    generation_model, vision_model = _setup_models(model_ids, engine.get('s2_model') or S2_DEFAULT_MODEL)
    if not generation_model:
        return JSONResponse({'ok': False, 'connected': False,
                             'error': '未验证到可用的文本生成模型，设置未保存'}, 400)
    text_verified_model_ids = (probe_result[3] if len(probe_result) > 3 else [generation_model])
    text_verified_model_ids = [str(value) for value in text_verified_model_ids
                               if str(value) in model_ids]
    if generation_model not in text_verified_model_ids:
        return JSONResponse({'ok': False, 'connected': False,
                             'error': '选定的生成模型没有通过真实文本调用，设置未保存'}, 400)
    with _json_lock(path):
        latest = read_json(path, {})
        latest_digest = _sha256_text(json.dumps(latest, ensure_ascii=False, sort_keys=True))
        if latest_digest != config_before_probe:
            return JSONResponse({'ok': False, 'connected': False,
                                 'error': '测试期间设置已发生变化，请重新点击连接测试'}, 409)
        engine['s2_model'] = generation_model
        candidate['engine'] = engine
        provider = {'id': 'setup-s2', 'name': '中标狗模型服务',
                    'base_url': engine['s2_base_url'], 'api_key': key,
                    'model': generation_model, 'vision_model': vision_model,
                    'kind': 'openai', 'verify_ssl': bool(engine.get('s2_verify_ssl', True))}
        providers = [item for item in (current.get('providers') or [])
                     if isinstance(item, dict) and item.get('id') != 'setup-s2']
        candidate['providers'] = providers + [provider]
        candidate['routing'] = {'default': 'setup-s2', 'model': generation_model}
        fingerprint = _key_fingerprint(key)
        connection_fingerprint = _connection_fingerprint(candidate)
        previous = current.get('setup') if isinstance(current.get('setup'), dict) else {}
        candidate['setup'] = {'tested_key_fingerprint': fingerprint,
                              'tested_connection_fingerprint': connection_fingerprint,
                              'connected_at': now(), 'completed_at': '',
                              'model_ids': model_ids[:200],
                              'text_verified_model_ids': text_verified_model_ids[:20]}
        write_json(path, candidate)
    invalidate_oc_runtime()
    return {'ok': True, 'connected': True, 'key_set': True,
            'key_changed': bool(previous.get('tested_key_fingerprint')
                                and previous.get('tested_key_fingerprint') != fingerprint),
            'engine': 'opencode', 'mode': requested_mode, 'model': generation_model,
            'vision_model': vision_model, 'vision_enabled': bool(vision_model)}

@app.post('/v1/setup/complete')
def setup_complete():
    path = conf_path()
    with _json_lock(path):
        conf = read_json(path, {})
        setup = conf.get('setup') if isinstance(conf.get('setup'), dict) else {}
        current_fp = _connection_fingerprint(conf)
        tested_fp = str(setup.get('tested_connection_fingerprint') or '')
        if not current_fp or not tested_fp or not hmac.compare_digest(current_fp, tested_fp):
            return JSONResponse({'ok': False, 'error': '模型连接已变化，请先重新测试连接'}, 409)
        setup = dict(setup); setup['completed_at'] = now(); conf['setup'] = setup
        write_json(path, conf)
    return {'ok': True, **setup_status_data(conf)}

ROLE_MAP = {'developer': 'system', 'system': 'system', 'user': 'user', 'assistant': 'assistant', 'tool': 'tool'}

def _chat_content(content):
    """Responses 的 content 数组 → chat 的 content。带图时转成多模态数组,否则并成一段纯文本"""
    if isinstance(content, str): return content
    texts, parts, has_img = [], [], False
    for c in (content or []):
        if isinstance(c, str): texts.append(c); parts.append({'type': 'text', 'text': c}); continue
        t = c.get('type') or ''
        if t in ('input_text', 'output_text', 'text', 'summary_text'):
            s = c.get('text') or ''
            texts.append(s); parts.append({'type': 'text', 'text': s})
        elif t in ('input_image', 'image_url'):
            u = c.get('image_url')
            u = u.get('url') if isinstance(u, dict) else u
            if u: has_img = True; parts.append({'type': 'image_url', 'image_url': {'url': u}})
    return parts if has_img else '\n'.join(x for x in texts if x)

def resp_to_chat(body, model):
    """Codex 的 Responses 请求 → chat/completions 请求。返回 (payload, 丢弃的工具类型)"""
    msgs = []
    if body.get('instructions'): msgs.append({'role': 'system', 'content': body['instructions']})
    inp = body.get('input')
    if isinstance(inp, str): inp = [{'type': 'message', 'role': 'user', 'content': inp}]
    for it in (inp or []):
        if isinstance(it, str): msgs.append({'role': 'user', 'content': it}); continue
        t = it.get('type') or 'message'
        if t == 'message':
            c = _chat_content(it.get('content'))
            if c: msgs.append({'role': ROLE_MAP.get(it.get('role') or 'user', 'user'), 'content': c})
        elif t == 'function_call':
            msgs.append({'role': 'assistant', 'content': '', 'tool_calls': [
                {'id': it.get('call_id') or it.get('id') or '', 'type': 'function',
                 'function': {'name': it.get('name') or '', 'arguments': it.get('arguments') or '{}'}}]})
        elif t == 'function_call_output':
            out = it.get('output')
            if isinstance(out, (dict, list)): out = json.dumps(out, ensure_ascii=False)
            msgs.append({'role': 'tool', 'tool_call_id': it.get('call_id') or '', 'content': out or ''})
        # reasoning / web_search_call 等本地专属条目:chat 协议没有对应位置,丢掉不影响续跑
    tools, dropped = [], []
    def add(f):
        tools.append({'type': 'function', 'function': {
            'name': f.get('name') or '', 'description': (f.get('description') or '')[:2048],
            'parameters': f.get('parameters') or {'type': 'object', 'properties': {}}}})
    for t in (body.get('tools') or []):
        ty = t.get('type')
        if ty == 'function': add(t)
        elif ty == 'namespace':                 # 子 agent 之类的命名空间工具:摊平成同名普通函数
            for c in (t.get('tools') or []):
                add(c) if c.get('type') == 'function' else dropped.append(c.get('type') or '?')
        else: dropped.append(ty or '?')         # web_search 等网关侧内建工具:上游没有,只能丢
    seen, uniq = set(), []
    for f in tools:
        n = f['function']['name']
        if n and n not in seen: seen.add(n); uniq.append(f)
    payload = {'model': model, 'messages': msgs, 'stream': True}
    if uniq:
        payload['tools'] = uniq
        payload['tool_choice'] = body.get('tool_choice') if body.get('tool_choice') in ('auto', 'none', 'required') else 'auto'
        if body.get('parallel_tool_calls') is not None: payload['parallel_tool_calls'] = bool(body['parallel_tool_calls'])
    return payload, dropped

def sse(ev, data):
    return ('event: %s\ndata: %s\n\n' % (ev, json.dumps(data, ensure_ascii=False))).encode('utf-8')

def _upstream(base, key, path, payload, timeout, verify):
    """发上游请求并返回未读完的连接(用于流式读取);证书失败时按 _openai_req 的策略换 CA 重试"""
    data = json.dumps(payload, ensure_ascii=False).encode('utf-8')
    req = urllib.request.Request(base.rstrip('/') + path, data=data, headers={
        'Authorization': 'Bearer ' + (key or ''), 'Content-Type': 'application/json',
        'Accept': 'text/event-stream' if payload.get('stream') else 'application/json',
        'User-Agent': 'bid-dog-relay/' + ENGINE_VERSION})
    last = None
    for kind in (['default', 'certifi'] if verify else ['none']):
        try: ctx = _ssl_ctx(kind)
        except Exception: continue
        try:
            # 生成一册标书要打几十上百次上游,握手被掐一次就整条任务失败太脆:瞬时故障自动重开连接
            note = lambda i, n, e: RELAY_LAST.update(
                {'ts': now(), 'retry': '第 %d/%d 次重试上游:%s' %
                 (i, n, _safe_secret_text(getattr(e, 'reason', e), (key,))[:120])})
            return _retry(lambda: urllib.request.urlopen(req, timeout=timeout, context=ctx), on_wait=note)
        except urllib.error.URLError as e:
            if not isinstance(getattr(e, 'reason', None), ssl.SSLCertVerificationError): raise
            last = e
    raise last


def _iter_upstream_chunks(response, streaming=False, size=8192):
    """Yield bytes without buffering a slow SSE response to ``size`` bytes.

    ``HTTPResponse.read(size)`` may wait for the whole requested amount. Models normally emit much
    smaller SSE frames, so that behavior can hide real progress and trigger an idle timeout in the
    caller. ``read1`` returns currently available bytes and keeps the model and UI heartbeat aligned.
    """
    reader = getattr(response, 'read1', None) if streaming else None
    if not callable(reader): reader = response.read
    while True:
        chunk = reader(size)
        if not chunk: return
        yield chunk

def _relay_stream(body, up):
    """一次 Codex 轮次:上游 chat 流 → Responses 事件流。
    任何一层失败都要发 response.failed,不能静默——静默失败是这个产品最贵的 bug 类型。"""
    rid = 'resp_' + uuid.uuid4().hex[:20]
    mid = 'msg_' + uuid.uuid4().hex[:16]
    yield sse('response.created', {'type': 'response.created', 'response': {'id': rid, 'status': 'in_progress'}})
    payload, dropped = resp_to_chat(body, up['model'])
    text, calls, usage, err, opened = [], {}, {}, '', [False]
    def open_msg():
        """先 output_item.added 再发 delta——否则 codex 会刷屏 OutputTextDelta without active item"""
        opened[0] = True
        return sse('response.output_item.added', {'type': 'response.output_item.added', 'output_index': 0,
                   'item': {'type': 'message', 'id': mid, 'role': 'assistant', 'status': 'in_progress', 'content': []}})
    RELAY_LAST.update({'ts': now(), 'model': up['model'], 'msgs': len(payload['messages']),
                       'tools': len(payload.get('tools') or []), 'dropped_tools': dropped, 'error': ''})
    try:
        try:
            r = _upstream(up['base_url'], up['api_key'], '/chat/completions', payload, 900, up['verify_ssl'])
        except urllib.error.HTTPError as e:
            det = _http_error_detail(e, (up['api_key'],), 400)
            # 有的网关不支持工具或不支持流式:退一步再试一次,总比整条任务断掉强
            if e.code in (400, 404, 422) and payload.get('stream'):
                payload['stream'] = False
                r = _upstream(up['base_url'], up['api_key'], '/chat/completions', payload, 900, up['verify_ssl'])
            else:
                raise RuntimeError('HTTP %s %s' % (e.code, det))
        # 有的网关不理会 stream:true、直接回整包 JSON —— 按响应头判断,别按我们的请求猜
        ct = (r.headers.get('Content-Type') or '').lower()
        if payload.get('stream') and 'json' not in ct:
            saw_done = False
            for raw in r:
                ln = raw.decode('utf-8', 'ignore').strip()
                if not ln.startswith('data:'): continue
                ln = ln[5:].strip()
                if ln == '[DONE]': saw_done = True; break
                try: d = json.loads(ln)
                except Exception: continue
                if d.get('usage'): usage = d['usage']
                ch = (d.get('choices') or [{}])[0]
                delta = ch.get('delta') or {}
                if delta.get('content'):
                    if not opened[0]: yield open_msg()
                    text.append(delta['content'])
                    yield sse('response.output_text.delta', {'type': 'response.output_text.delta',
                                                             'item_id': mid, 'output_index': 0, 'delta': delta['content']})
                # reasoning_content(思考过程)不转发:codex 侧本来就 hide_agent_reasoning,转过去只会多一类协议噪声
                for tc in (delta.get('tool_calls') or []):
                    i = tc.get('index', 0)
                    c = calls.setdefault(i, {'id': '', 'name': '', 'args': ''})
                    if tc.get('id'): c['id'] = tc['id']
                    fn = tc.get('function') or {}
                    if fn.get('name'): c['name'] = fn['name']
                    if fn.get('arguments'): c['args'] += fn['arguments']
            if not saw_done:
                raise http.client.IncompleteRead(b'', None)
        else:
            d = json.loads(r.read().decode('utf-8', 'ignore'))
            usage = d.get('usage') or {}
            m = ((d.get('choices') or [{}])[0].get('message') or {})
            if m.get('content'):
                text.append(m['content'])
                yield open_msg()
                yield sse('response.output_text.delta', {'type': 'response.output_text.delta',
                                                         'item_id': mid, 'output_index': 0, 'delta': m['content']})
            for i, tc in enumerate(m.get('tool_calls') or []):
                fn = tc.get('function') or {}
                calls[i] = {'id': tc.get('id') or '', 'name': fn.get('name') or '', 'args': fn.get('arguments') or '{}'}
    except Exception as e:
        # 已经向下游吐过任何字节后绝不能整轮重放（会重复正文，半截 tool arguments 更危险）。
        # 明确 response.failed，让执行外壳走“继续做/重跑”，禁止把半截正文伪装 completed。
        err = '上游流在回复完成前断开：' + net_hint(e, (up['api_key'],))
    if err:
        RELAY_LAST['error'] = err
        yield sse('response.failed', {'type': 'response.failed', 'response': {'id': rid, 'status': 'failed',
                  'error': {'code': 'upstream_error', 'message': '中标狗中转层:调用你的 S2 网关失败 —— ' + err}}})
        return
    idx = 0
    body_txt = ''.join(text).strip()
    if body_txt:
        if not opened[0]: yield open_msg()
        yield sse('response.output_text.done', {'type': 'response.output_text.done', 'item_id': mid,
                                                'output_index': 0, 'text': body_txt})
        yield sse('response.output_item.done', {'type': 'response.output_item.done', 'output_index': idx,
                  'item': {'type': 'message', 'id': mid, 'role': 'assistant',
                           'status': 'completed', 'content': [{'type': 'output_text', 'text': body_txt}]}})
        idx += 1
    for i in sorted(calls):
        c = calls[i]
        if not c['name']: continue
        cid = c['id'] or ('call_' + uuid.uuid4().hex[:16])
        yield sse('response.output_item.done', {'type': 'response.output_item.done', 'output_index': idx,
                  'item': {'type': 'function_call', 'id': 'fc_' + uuid.uuid4().hex[:16], 'call_id': cid,
                           'name': c['name'], 'arguments': c['args'] or '{}', 'status': 'completed'}})
        idx += 1
    it, ot = int(usage.get('prompt_tokens') or 0), int(usage.get('completion_tokens') or 0)
    RELAY_LAST.update({'in': it, 'out': ot, 'calls': [calls[i]['name'] for i in sorted(calls)], 'chars': len(body_txt)})
    yield sse('response.completed', {'type': 'response.completed', 'response': {'id': rid, 'status': 'completed',
              'usage': {'input_tokens': it, 'output_tokens': ot,
                        'input_tokens_details': {'cached_tokens': 0}, 'output_tokens_details': {'reasoning_tokens': 0},
                        'total_tokens': int(usage.get('total_tokens') or (it + ot))}}})

def _relay_passthrough(body, up):
    """原生 Responses 直通，但仍核验终止事件；普通 EOF 不能把半截回复伪装成功。"""
    try:
        r = _upstream(up['base_url'], up['api_key'], '/responses', body, 900, up['verify_ssl'])
    except Exception as e:
        RELAY_LAST.update({'ts': now(), 'error': net_hint(e, (up['api_key'],)), 'mode': 'passthrough'})
        yield sse('response.failed', {'type': 'response.failed', 'response': {'id': 'resp_err', 'status': 'failed',
                  'error': {'code': 'upstream_error', 'message': '中标狗中转层(直通模式):' + net_hint(e, (up['api_key'],))}}})
        return
    tail = b''
    try:
        for chunk in _iter_upstream_chunks(r, streaming=True):
            tail = (tail + chunk)[-65536:]
            yield chunk
    except Exception as e:
        reason = '上游 Responses 流在完整收尾前断开：' + net_hint(e, (up['api_key'],))
        RELAY_LAST.update({'ts': now(), 'error': reason, 'mode': 'passthrough'})
        yield sse('response.failed', {'type': 'response.failed', 'response': {'id': 'resp_err', 'status': 'failed',
                  'error': {'code': 'upstream_error', 'message': '中标狗中转层(直通模式):' + reason}}})
        return
    complete = (b'event: response.completed' in tail or b'"type":"response.completed"' in tail
                or b'"type": "response.completed"' in tail)
    failed = (b'event: response.failed' in tail or b'"type":"response.failed"' in tail
              or b'"type": "response.failed"' in tail)
    if body.get('stream') is True and not complete and not failed:
        reason = '上游 Responses 流提前结束，没有收到完整结束信号'
        RELAY_LAST.update({'ts': now(), 'error': reason, 'mode': 'passthrough'})
        yield sse('response.failed', {'type': 'response.failed', 'response': {'id': 'resp_err', 'status': 'failed',
                  'error': {'code': 'upstream_error', 'message': '中标狗中转层(直通模式):' + reason}}})

@app.post('/v1/relay/responses')
async def relay_responses(req: Request):
    if (req.headers.get('authorization') or '').replace('Bearer ', '').strip() != relay_token():
        RELAY_LAST.update({'ts': now(), 'mode': 'auth',
                           'error': '执行外壳未携带本机中继凭据'})
        return JSONResponse({'error': 'relay token 不匹配'}, 401)
    body = await req.json()
    up = s2_conf()
    if not up['api_key']:
        return JSONResponse({'error': '还没填 S2 API Key(设置 · 生成引擎)'}, 400)
    up['model'] = up['model'] or body.get('model') or S2_DEFAULT_MODEL
    # Responses API 缺省是非流式。原生直通时必须保留 JSON 响应类型，不能硬套 SSE。
    if up['wire'] == 'responses' and body.get('stream') is not True:
        try:
            r = await asyncio.to_thread(_upstream, up['base_url'], up['api_key'], '/responses',
                                        body, 900, up['verify_ssl'])
            raw = await asyncio.to_thread(r.read)
            status = int(getattr(r, 'status', 200) or 200)
            content_type = (r.headers.get('Content-Type') or 'application/json').split(';', 1)[0]
            return Response(content=raw, status_code=status, media_type=content_type,
                            headers={'Cache-Control': 'no-store'})
        except urllib.error.HTTPError as e:
            detail = _http_error_detail(e, (up['api_key'],), 400)
            return JSONResponse({'error': {'message': 'HTTP %s %s' % (e.code, detail),
                                           'type': 'upstream_error'}}, e.code)
        except Exception as e:
            return JSONResponse({'error': {'message': '中转层:' + net_hint(e, (up['api_key'],)),
                                           'type': 'upstream_error'}}, 502)
    gen = _relay_passthrough(body, up) if up['wire'] == 'responses' else _relay_stream(body, up)
    return StreamingResponse(gen, media_type='text/event-stream',
                             headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})

@app.post('/v1/relay/chat/completions')
async def relay_chat(req: Request):
    """纯直通(零翻译):opencode 等原生 OpenAI 兼容外壳 → 这里 → S2 网关。
    存在的唯一理由是 Key 托管:子进程只拿本机随机口令,真 Key 不出引擎进程;顺带统一记账 RELAY_LAST。"""
    if (req.headers.get('authorization') or '').replace('Bearer ', '').strip() != relay_token():
        RELAY_LAST.update({'ts': now(), 'mode': 'auth',
                           'error': '执行外壳未携带本机中继凭据'})
        return JSONResponse({'error': {'message': 'relay token 不匹配', 'type': 'auth_error'}}, 401)
    body = await req.json()
    up = s2_conf()
    if not up['api_key']:
        return JSONResponse({'error': {'message': '还没填 S2 API Key(设置 · 生成引擎)', 'type': 'config_error'}}, 400)
    body['model'] = up['model'] or body.get('model')
    RELAY_LAST.update({'ts': now(), 'mode': 'chat-pass', 'model': body.get('model'),
                       'msgs': len(body.get('messages') or []), 'tools': len(body.get('tools') or []), 'error': ''})
    def gen(r):
        sent = 0
        tail = b''
        reopened = 0
        chunks = iter(_iter_upstream_chunks(r, streaming=bool(body.get('stream'))))
        while True:
            try:
                c = next(chunks, b'')
                if c:
                    sent += len(c); tail = (tail + c)[-256:]
                    yield c
                    continue
                # 标准 OpenAI SSE 必须以 [DONE] 收尾。普通 EOF 没有它就是静默断流。
                if body.get('stream') and b'[DONE]' not in tail:
                    raise http.client.IncompleteRead(tail, None)
                break
            except Exception as e:
                # 首字节前可以安全重开；一旦发过字节就绝不重放，避免重复正文/重复工具调用。
                if not sent and reopened < len(RETRY_WAITS) and _is_transient(e):
                    time.sleep(RETRY_WAITS[reopened]); reopened += 1
                    try:
                        r = _upstream(up['base_url'], up['api_key'], '/chat/completions',
                                      body, 900, up['verify_ssl'])
                        chunks = iter(_iter_upstream_chunks(r, streaming=bool(body.get('stream'))))
                        continue
                    except Exception as open_err:
                        e = open_err
                msg = '上游流在回复完成前断开：' + net_hint(e, (up['api_key'],))
                RELAY_LAST['error'] = msg
                if body.get('stream'):
                    yield ('data: %s\n\n' % json.dumps({'error': {
                        'message': msg, 'type': 'upstream_stream_error', 'code': 'stream_interrupted'
                    }}, ensure_ascii=False)).encode('utf-8')
                return
    try:
        r = await asyncio.to_thread(_upstream, up['base_url'], up['api_key'], '/chat/completions', body, 900, up['verify_ssl'])
    except urllib.error.HTTPError as e:
        det = _http_error_detail(e, (up['api_key'],), 400)
        RELAY_LAST['error'] = 'HTTP %s %s' % (e.code, det)
        try: return JSONResponse(json.loads(det), e.code)
        except Exception:
            return JSONResponse({'error': {'message': 'HTTP %s %s' % (e.code, det), 'type': 'upstream_error'}}, e.code)
    except Exception as e:
        RELAY_LAST['error'] = net_hint(e, (up['api_key'],))
        return JSONResponse({'error': {'message': '中转层:' + net_hint(e, (up['api_key'],)), 'type': 'upstream_error'}}, 502)
    ct = r.headers.get('Content-Type') or 'application/json'
    return StreamingResponse(gen(r), media_type=ct,
                             headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})

@app.get('/v1/relay/models')
def relay_models():
    up = s2_conf()
    hit = _models_cached(up['base_url'], up['api_key'], up['verify_ssl'])
    if hit: return hit
    try:
        d = _openai_req(up['base_url'], up['api_key'], '/models', timeout=15, verify=up['verify_ssl'])
        ids = [m.get('id') for m in (d.get('data') or []) if m.get('id')]
        _models_store(up['base_url'], up['api_key'], up['verify_ssl'], ids)
        return {'ok': True, 'models': ids}
    except Exception as e:
        return JSONResponse({'error': net_hint(e, (up['api_key'],))}, 502)

# ---------- 一键安装执行外壳(Codex CLI):没内置、没装 Node 的机器,点一下按钮就好 ----------
# 版本钉死在我们测过兼容的那个:S2 中转是按它的 Responses 行为实测的,随手升级可能翻车
CODEX_PIN = os.environ.get('BID_CODEX_VERSION', '0.146.0')
OPENCODE_PIN = os.environ.get('BID_OPENCODE_VERSION', '1.18.18')   # 与直连实测配套的版本
# opencode 平台包是普通 npm 包名(非别名):tarball = {reg}/{pkg}/-/{pkg}-{ver}.tgz,包内 package/bin/opencode[.exe]
OPENCODE_PLAT = {
    ('darwin', 'arm64'):  'opencode-darwin-arm64',
    ('darwin', 'x86_64'): 'opencode-darwin-x64',
    ('win32',  'amd64'):  'opencode-windows-x64',
    ('win32',  'arm64'):  'opencode-windows-arm64',
    ('linux',  'x86_64'): 'opencode-linux-x64',
    ('linux',  'aarch64'): 'opencode-linux-arm64',
}

def _plat_key():
    import platform as _pl
    mach = (_pl.machine() or '').lower()
    mach = {'x86_64': 'x86_64', 'amd64': 'amd64' if sys.platform == 'win32' else 'x86_64',
            'arm64': 'arm64', 'aarch64': 'arm64' if sys.platform == 'darwin' else 'aarch64'}.get(mach, mach)
    return (sys.platform if sys.platform != 'cygwin' else 'win32', mach)

def _has_avx2():
    """本机 CPU 支不支持 AVX2。

    opencode 官方安装脚本会探测这一项:不支持就必须下 `-baseline` 变体,
    否则二进制**直接崩**,而且报错完全看不出是 CPU 指令集的问题。
    招投标客户里老工控机/老办公机不少,这条不是理论风险。
    探测失败一律当"不支持"——下 baseline 只是稍慢,下错了是根本跑不起来。"""
    try:
        if sys.platform == 'win32':
            import ctypes
            return bool(ctypes.windll.kernel32.IsProcessorFeaturePresent(40))   # 40 = PF_AVX2_INSTRUCTIONS_AVAILABLE
        if sys.platform == 'darwin':
            import platform as _pl
            if (_pl.machine() or '').lower() in ('arm64', 'aarch64'): return True   # 苹果芯片没有 baseline 变体
            out = subprocess.run(['sysctl', '-n', 'machdep.cpu.leaf7_features'],
                                 capture_output=True, text=True, timeout=4).stdout
            return 'AVX2' in out.upper()
        flags = open('/proc/cpuinfo', encoding='utf-8', errors='ignore').read().lower()
        return bool(re.search(r'(^|\s)avx2(\s|$)', flags, re.M))
    except Exception:
        return False

def opencode_pkg():
    """opencode 的 npm 平台包名。老 CPU(无 AVX2)必须换 -baseline 变体,否则二进制装上了也跑不起来。

    baseline 变体**只有 x64 平台有**——实测 npm:windows-x64/linux-x64/darwin-x64 的 baseline
    都是 200,而 darwin-arm64 / windows-arm64 / linux-arm64 的 baseline 全是 404
    (AVX2 是 x86 的指令集,ARM 上根本没这回事)。
    所以只对 x64 做回退,arm64 一律用标准包,别去拼一个不存在的包名。"""
    pkg = OPENCODE_PLAT.get(_plat_key())
    if not pkg: return None
    if not pkg.endswith(('-x64',)): return pkg      # arm64:没有也不需要 baseline
    return pkg if _has_avx2() else pkg + '-baseline'

def _shell_meta(which):
    # 外壳下载参数:('codex'|'opencode') -> (url构造器, 包内路径, 目标文件名, 手装提示);不支持的平台返回 None
    ext = '.exe' if os.name == 'nt' else ''
    if which == 'opencode':
        pkg = opencode_pkg()
        if not pkg: return None
        return (lambda reg: '%s/%s/-/%s-%s.tgz' % (reg.rstrip('/'), pkg, pkg, OPENCODE_PIN),
                'package/bin/opencode' + ext, 'opencode-cli' + ext, 'npm i -g opencode-ai')
    plat = CODEX_PLAT.get(_plat_key())
    if not plat: return None
    suffix, triple = plat
    return (lambda reg: '%s/@openai/codex/-/codex-%s-%s.tgz' % (reg.rstrip('/'), CODEX_PIN, suffix),
            'package/vendor/%s/bin/codex%s' % (triple, ext), 'codex-cli' + ext, 'npm i -g @openai/codex')
# 平台 → (npm 版本后缀, 包内 vendor 三元组)。npm 的平台包是别名:@openai/codex@<版本>-<平台>
CODEX_PLAT = {
    ('darwin', 'arm64'):  ('darwin-arm64',  'aarch64-apple-darwin'),
    ('darwin', 'x86_64'): ('darwin-x64',    'x86_64-apple-darwin'),
    ('win32',  'amd64'):  ('win32-x64',     'x86_64-pc-windows-msvc'),
    ('win32',  'arm64'):  ('win32-arm64',   'aarch64-pc-windows-msvc'),
    ('linux',  'x86_64'): ('linux-x64',     'x86_64-unknown-linux-musl'),
    ('linux',  'aarch64'): ('linux-arm64',  'aarch64-unknown-linux-musl'),
}
# 国内客户下 npmjs 常年龟速:先走 npmmirror,再退回官源(路径格式两家一致)
CODEX_REGISTRIES = ([os.environ.get('BID_CODEX_REGISTRY')] if os.environ.get('BID_CODEX_REGISTRY') else
                    ['https://registry.npmmirror.com', 'https://registry.npmjs.org'])
PROV = {'state': 'idle', 'pct': 0, 'note': '', 'path': '', 'error': ''}

def codex_platform():
    import platform as _pl
    mach = (_pl.machine() or '').lower()
    mach = {'x86_64': 'x86_64', 'amd64': 'amd64' if sys.platform == 'win32' else 'x86_64',
            'arm64': 'arm64', 'aarch64': 'arm64' if sys.platform == 'darwin' else 'aarch64'}.get(mach, mach)
    return CODEX_PLAT.get((sys.platform if sys.platform != 'cygwin' else 'win32', mach))

def _provision_codex(which='codex'):
    meta = _shell_meta(which)
    if not meta:
        PROV.update({'state': 'error', 'error': '暂不支持这个系统架构，请联系提供方获取兼容安装包'}); return
    mk_url, inner, exe, manual = meta
    dstdir = _mk(os.path.join(DATA, 'bin'))
    dst = os.path.join(dstdir, exe)
    tmp = os.path.join(dstdir, '_shell.tgz')
    last = ''
    for reg in CODEX_REGISTRIES:
        url = mk_url(reg)
        try:
            PROV.update({'state': 'running', 'which': which, 'pct': 0, 'note': '正在下载生成组件(%s)…' % which, 'error': ''})
            req = urllib.request.Request(url, headers={'User-Agent': 'bid-dog/' + ENGINE_VERSION})
            with urllib.request.urlopen(req, timeout=60) as r, open(tmp, 'wb') as f:
                total = int(r.headers.get('Content-Length') or 0)
                got = 0
                while True:
                    chunk = r.read(1 << 18)
                    if not chunk: break
                    f.write(chunk); got += len(chunk)
                    if total: PROV['pct'] = int(got * 80 / total)   # 下载占 80%,解压校验占 20%
            PROV.update({'pct': 82, 'note': '正在解压…'})
            import tarfile
            with tarfile.open(tmp, 'r:gz') as t:
                m = t.getmember(inner)
                m.name = os.path.basename(dst)                      # 只取这一个文件,落到 bin/ 平铺
                t.extract(m, dstdir)
            os.chmod(dst, 0o755)
            PROV.update({'pct': 92, 'note': '正在校验…'})
            r = _tracked_detached_run([dst, '--version'], capture_output=True, text=True, timeout=30,
                                      stdin=subprocess.DEVNULL)
            ver = (r.stdout or r.stderr or '').strip()
            if r.returncode != 0: raise RuntimeError('校验失败:%s' % ver[-200:])
            PROV.update({'state': 'done', 'pct': 100, 'path': dst, 'note': '已安装:%s' % (ver or exe)})
            try: os.remove(tmp)
            except Exception: pass
            return
        except Exception as e:
            last = net_hint(e)
            try: os.remove(tmp)
            except Exception: pass
    PROV.update({'state': 'error', 'error': '下载失败(镜像与官源都试过):%s。也可手动安装:%s' % (last, manual)})

@app.post('/v1/agent/provision')
async def agent_provision(req: Request):
    """一键安装执行外壳到数据目录(默认 OpenCode:压缩约 60MB、解压约 170MB;codex 约 130MB)。不碰系统目录,免管理员权限。
    已有可用外壳(内置或此前装过)则直接返回;文件坏了(跑不动 --version)会自动删掉重下。"""
    # 默认装 OpenCode:「自动(默认)」引擎现在就是它。前端不传 which 时不该再去装 codex ——
    # 那样客户点完「一键安装执行外壳」,装回来的还是个用不上的壳,自检照样报缺外壳。
    try: which = ((await req.json()).get('which') or 'opencode')
    except Exception: which = 'opencode'
    if which not in ('codex', 'opencode'): which = 'opencode'
    if PROV['state'] == 'running': return PROV
    b = bundled_cli('opencode-cli' if which == 'opencode' else 'codex-cli')
    if b:
        try:
            r = _tracked_detached_run([b, '--version'], capture_output=True, text=True, timeout=30,
                                      stdin=subprocess.DEVNULL)
            if r.returncode == 0:
                PROV.update({'state': 'done', 'pct': 100, 'path': b,
                             'note': '已就绪:%s' % (r.stdout or '').strip()[:60]})
                return PROV
        except Exception: pass
        if os.path.dirname(b) == os.path.join(DATA, 'bin'):   # 只清我们自己下的,不动安装包内置的
            try: os.remove(b)
            except Exception: pass
    threading.Thread(target=_provision_codex, args=(which,), daemon=True).start()
    await asyncio.sleep(0.3)   # 让状态先落到 running,前端第一次轮询就有数
    return PROV

@app.get('/v1/agent/provision')
def agent_provision_status():
    return PROV

@app.get('/v1/relay/status')
def relay_status():
    up = s2_conf()
    return {'base_url': up['base_url'], 'model': up['model'], 'has_key': bool(up['api_key']),
            'wire': up['wire'], 'relay_url': relay_base(), 'last': _redact_runtime(RELAY_LAST)}

def relay_base():
    return 'http://127.0.0.1:%d/v1/relay' % SELF_PORT

@app.put('/v1/routing')
async def routing(req: Request):
    conf = read_json(conf_path(), {'providers': [], 'routing': {}})
    conf['routing'] = await req.json(); write_json(conf_path(), conf)
    return {'ok': True}

@app.post('/v1/assets/config')
async def assets_config(req: Request):
    """更换素材库位置(留空 dir = 恢复默认)。原文件不搬迁,只是改指向,提示由前端给出"""
    body = await req.json()
    d = (body.get('dir') or '').strip()
    conf = read_json(conf_path(), {'providers': [], 'routing': {}})
    if d:
        d = os.path.expanduser(d)
        try: os.makedirs(d, exist_ok=True)
        except Exception as e:
            return JSONResponse({'ok': False, 'error': '这个路径建不出来:%s' % e}, 400)
        if not os.path.isdir(d) or not os.access(d, os.W_OK):
            return JSONResponse({'ok': False, 'error': '路径不可写,请换一个位置'}, 400)
        conf['assets_dir'] = d
    else:
        conf.pop('assets_dir', None)
    write_json(conf_path(), conf)
    return {'ok': True, 'folder': assets_dir(), 'default': assets_default(), 'is_default': not d}

@app.get('/v1/assets')
def assets():
    out = {'folder': assets_dir(), 'default': assets_default(),
           'is_default': not (json_quiet(conf_path()).get('assets_dir') or '').strip(), 'items': []}
    for root, _, files in os.walk(assets_dir()):
        for fn in files:
            if fn.startswith('.') or fn == '入库流水.jsonl': continue
            rel = os.path.relpath(os.path.join(root, fn), assets_dir())
            out['items'].append({'path': rel, 'category': rel.split(os.sep)[0] if os.sep in rel else '未分类'})
    try:
        lines = open(os.path.join(assets_dir(), '入库流水.jsonl'), encoding='utf-8').read().splitlines()
        out['recent'] = [json.loads(l) for l in lines[-5:]][::-1]
    except Exception:
        out['recent'] = []
    return out

@app.post('/v1/assets')
async def add_asset(file: UploadFile = File(...), category: str = Form('未分类')):
    # category/filename 都是外部输入:不做净化的话 '..\\..\\x' 能逃出素材库写任意路径
    fn = os.path.basename(file.filename or '素材')
    cat = re.sub(r'[\\/:*?"<>|.]+', '', str(category or ''))[:40] or '未分类'
    d = os.path.join(assets_dir(), cat)
    if not os.path.realpath(d).startswith(os.path.realpath(assets_dir())):
        return JSONResponse({'ok': False, 'error': '分类名不合法'}, 400)
    os.makedirs(d, exist_ok=True)
    open(os.path.join(d, fn), 'wb').write(await file.read())
    return {'ok': True, 'path': os.path.join(cat, fn)}

# ---------- 素材库 · 过往标书自动入库(确定性:拆图/拆章/归类;AI 增强走 AGENT_CMD) ----------
CAT_RULES = [
    ('公司介绍',   ['公司简介', '公司介绍', '企业概况', '企业简介', '关于我们', '公司概况']),
    ('产品资料',   ['产品', '技术方案', '系统架构', '技术架构', '功能', '平台', '解决方案', '总体设计', '系统设计', '技术路线']),
    ('资质与案例', ['资质', '证书', '荣誉', '案例', '业绩', '中标', '成功案例', '项目经验', '合同']),
    ('人员与团队', ['人员', '团队', '项目组', '组织架构', '简历', '项目经理']),
    ('服务承诺',   ['售后', '服务承诺', '培训', '质保', '维保', '服务方案', '运维', '应急']),
    ('商务文件',   ['应标信', '投标函', '授权', '承诺书', '偏离表', '报价']),
]

def classify(title):
    for cat, kws in CAT_RULES:
        if any(k in title for k in kws): return cat
    return '通用章节'

def split_md_sections(text):
    """按 markdown 标题拆章 → [(标题, 含标题的整章内容)];无标题则整文一节"""
    import re
    secs, cur_t, cur = [], None, []
    for ln in text.splitlines():
        m = re.match(r'^(#{1,3})\s+(.+)', ln)
        if m:
            if cur_t is not None or any(l.strip() for l in cur):
                secs.append((cur_t or '未命名章节', '\n'.join(cur).strip()))
            cur_t, cur = m.group(2).strip(), [ln]
        else:
            cur.append(ln)
    if cur_t is not None or any(l.strip() for l in cur):
        secs.append((cur_t or '未命名章节', '\n'.join(cur).strip()))
    return [(t, c) for t, c in secs if c.strip()]

def _style_name(p):
    """安全取段落样式名。真实招标/投标 docx(尤其 .doc 转来的、或 WPS 出的)常有 p.style 为 None,
    直接 p.style.name 会抛 AttributeError 把整个解析打崩——用户给的那份工程招标文件就崩在这里。"""
    try:
        st = p.style
        return (st.name if st is not None and st.name else '') or ''
    except Exception:
        return ''

def docx_to_sections_and_images(path, img_dir):
    """docx → (按标题拆分的章节, 提取出的图片文件名);表格转 markdown 表"""
    import hashlib
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
    doc = Document(path)
    os.makedirs(img_dir, exist_ok=True)
    images = []
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype and not rel.is_external:
            blob = rel.target_part.blob
            if len(blob) < 2048: continue  # 过小的多为分隔线/图标
            ext = os.path.splitext(str(rel.target_part.partname))[1] or '.png'
            fn = '标书图_' + hashlib.md5(blob).hexdigest()[:10] + ext
            fp = os.path.join(img_dir, fn)
            if not os.path.exists(fp): open(fp, 'wb').write(blob)
            images.append(fn)
    out = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            p = Paragraph(child, doc)
            t = p.text.strip()
            if not t: continue
            style = _style_name(p).lower()      # p.style 可能为 None(真实招标文件里常见),不能直接取 .name
            if 'heading 1' in style or style == '标题 1': out.append('# ' + t)
            elif 'heading 2' in style or style == '标题 2': out.append('## ' + t)
            elif 'heading' in style or style.startswith('标题'): out.append('### ' + t)
            else: out.append(t)
        elif child.tag == qn('w:tbl'):
            tb = Table(child, doc)
            rows = ['| ' + ' | '.join(' '.join(c.text.split()) for c in r.cells) + ' |' for r in tb.rows]
            if rows:
                rows.insert(1, '|' + '---|' * len(tb.rows[0].cells))
                out.append('\n'.join(rows))
    return split_md_sections('\n\n'.join(out)), images

_CN_NUM = '一二三四五六七八九十百零0-9'

def promote_headings(text):
    """把中文标书常见编号行提升为 markdown 标题。

    PDF/doc 转出来的纯文本没有样式,split_md_sections 会把整本书当成一节——
    章节模板一条都拆不出来。按「第X章 / 一、 / 1.1」这类编号行(且行足够短,
    避免把正文长句误提升)补上 #/##/###,无样式文本也能按章入库。"""
    if re.search(r'^#{1,3}\s', text, re.M): return text     # 已有 markdown 标题就不动
    out = []
    for ln in text.splitlines():
        s = ln.strip()
        if s and len(s) <= 40:
            if re.match(r'^第[%s]+[章部分篇卷册]' % _CN_NUM, s):
                out.append('# ' + s); continue
            if re.match(r'^[一二三四五六七八九十]+、', s):
                out.append('## ' + s); continue
            if re.match(r'^\d+(\.\d+)+\s*[^.\d]', s):
                out.append('### ' + s); continue
            if re.match(r'^\d+[、.]\s*[^\d.]', s) and len(s) <= 30:
                out.append('## ' + s); continue
        out.append(ln)
    return '\n'.join(out)

def pdf_to_text(path, max_pages=300, max_chars=800_000):
    """pypdf 文本层提取(引擎打包依赖里早就带着 pypdf,招标主件链路一直在用)。"""
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = []
    remaining = max_chars
    for pg in reader.pages[:max_pages]:
        if remaining <= 0: break
        try: t = (pg.extract_text() or '').strip()
        except Exception: t = ''
        if t:
            pages.append(t[:remaining]); remaining -= len(pages[-1])
    return '\n\n'.join(pages)

def html_to_text(blob):
    """HTML 参考资料(资质证书页、网页存档)→ 纯文本。标准库解析,不引第三方。"""
    from html.parser import HTMLParser
    class _T(HTMLParser):
        def __init__(self):
            super().__init__(); self.buf = []; self.skip = 0
        def handle_starttag(self, tag, attrs):
            if tag in ('script', 'style'): self.skip += 1
            if tag in ('p', 'div', 'br', 'tr', 'li', 'h1', 'h2', 'h3', 'h4', 'table'):
                self.buf.append('\n')
        def handle_endtag(self, tag):
            if tag in ('script', 'style'): self.skip = max(0, self.skip - 1)
        def handle_data(self, data):
            if not self.skip and data.strip(): self.buf.append(data)
    t = _T()
    try: t.feed(blob.decode('utf-8', 'ignore'))
    except Exception: pass
    return re.sub(r'\n{3,}', '\n\n', ''.join(t.buf)).strip()

def rtf_to_text(blob):
    """RTF → 纯文本(WPS/老系统把 .doc 存成 RTF 很常见)。标准库实现,不引依赖。
    连续的 \\'xx 十六进制转义按整段解码——GBK 双字节被逐字节解码会全变成乱码。"""
    txt = blob.decode('latin-1', 'ignore')
    txt = re.sub(r'\{\\\*[^{}]*\}', '', txt)
    txt = re.sub(r"(?:\\'[0-9a-fA-F]{2})+",
                 lambda m: bytes.fromhex(m.group(0).replace("\\'", '')).decode('gbk', 'ignore'), txt)
    txt = re.sub(r'\\par[d]?\b', '\n', txt)
    txt = re.sub(r'\\[a-zA-Z]+-?\d*\s?', '', txt)
    txt = txt.replace('{', '').replace('}', '')
    return re.sub(r'\n{3,}', '\n\n', txt).strip()

def convert_doc_to_docx(src):
    """老 .doc → .docx。纯 Python 解不了 OLE2 老格式,借系统自带能力转换:
    macOS 用 textutil(系统内置);Windows 用已装的 Word COM;两端都认 LibreOffice。
    转不了就抛 RuntimeError,错误文案直接告诉用户手动另存为 docx。"""
    dst = os.path.splitext(src)[0] + '.docx'
    if os.path.isfile(dst): return dst
    outdir = os.path.dirname(src)
    tries = []
    if sys.platform == 'darwin':
        tries.append(['textutil', '-convert', 'docx', '-output', dst, src])
    soffice = shutil.which('soffice') or next((p for p in (
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe') if os.path.isfile(p)), None)
    if soffice:
        tries.append([soffice, '--headless', '--convert-to', 'docx', '--outdir', outdir, src])
    for cmd in tries:
        try:
            subprocess.run(cmd, timeout=180, capture_output=True, **({} if os.name != 'nt' else DETACH))
            if os.path.isfile(dst): return dst
        except Exception:
            continue
    if os.name == 'nt':
        # Word COM:多数装了 Office 的 Windows 都走得通;没装 Office 会在 New-Object 处失败
        ps = ("$ErrorActionPreference='Stop';$w=New-Object -ComObject Word.Application;"
              "$w.Visible=$false;$d=$w.Documents.Open('%s',$false,$true);"
              "$d.SaveAs2('%s',16);$d.Close($false);$w.Quit()") % (
                  src.replace("'", "''"), dst.replace("'", "''"))
        try:
            subprocess.run(['powershell', '-NoProfile', '-NonInteractive', '-Command', ps],
                           timeout=300, capture_output=True, **DETACH)
            if os.path.isfile(dst): return dst
        except Exception:
            pass
    raise RuntimeError('本机没有可用的 .doc 转换器(Word/LibreOffice/textutil 都不可用),'
                       '请用 Word/WPS 打开后「另存为 .docx」再上传')

# ---------- 图片 AI 打标(调用已配的视觉模型:OCR 读图 → 分类/图注/落位锚点) ----------
CATS = ['架构图', '功能截图', '资质证书', '案例证明', '人员证书', '其他']
VISION_PROMPT = (
    '这是投标文件素材库里的一张图片。请仔细看图并 OCR 读出图内文字,然后只输出一个 JSON(不要代码块、不要解释):\n'
    '{"category":"从 架构图/功能截图/资质证书/案例证明/人员证书/其他 里选一个",'
    '"id":"用 类别-业务名 命名,如 资质-ISO9001、功能截图-代码扫描、架构-总体架构,只用中文数字英文和连字符",'
    '"caption":"正式标书口径的图注,如 ISO9001质量管理体系认证证书",'
    '"ocr":"图内关键文字摘要,60字以内;看不清写 待确认",'
    '"keywords":"适用场景关键词,逗号分隔",'
    '"anchor":"落位锚点=这张图只允许插到标书哪类位置,如 总体架构章·概述段后 / 资质证明章 / 逐条应答证明材料列"}')
VISION = {'running': False, 'done': 0, 'total': 0, 'updated': 0, 'errors': [], 'ts': ''}
VISION_LOCK = threading.Lock()   # check-then-set 竞态防护:双击/自动触发并发时只允许一个 worker

def try_start_vision(force=False):
    """原子地启动打标 worker;已在跑或没配视觉模型返回 False。"""
    p, err = _vision_provider()
    if err or not p: return False
    with VISION_LOCK:
        if VISION['running']: return False
        VISION.update(running=True, done=0, total=0, updated=0, degraded=0, errors=[])
    threading.Thread(target=vision_worker, args=(force,), daemon=True).start()
    return True

def _vision_provider():
    conf = read_json(conf_path(), {'providers': [], 'routing': {}})
    ps = conf.get('providers') or []
    if not ps: return None, '还没有配置模型接入点'
    pid = (conf.get('routing') or {}).get('default')
    p = next((x for x in ps if x.get('id') == pid), ps[-1])
    vm = p.get('vision_model') or ''
    if not vm:
        m = (p.get('model') or '').lower()
        vm = p.get('model') if ('vl' in m or 'vision' in m or '-v' in m) else ''
    if not vm: return None, '当前接入点没有指定视觉模型:在「模型接入」的视觉模型一栏填一个多模态模型(如 senseaudio-vl-1.0-260319)'
    return {**p, 'vision_model': vm}, ''

def _lenient_json_load(txt):
    """容错解析视觉模型返回的 JSON。

    合同截图这类 OCR 文本长的图,模型返回的 JSON 三种常见坏法(真机 120 张里 7 张全倒在这):
      1) 输出到 max_tokens 截断 → Unterminated string / 缺右括号
      2) 字符串值里混进裸英文引号或换行 → Expecting ',' delimiter
      3) 末尾多写了半个键值对
    先严格解析;失败后逐步保守修复(转义裸控制字符→补关字符串→砍尾部残缺对→补右括号)再试。"""
    try:
        return json.loads(txt)
    except Exception:
        pass
    out, in_str, esc = [], False, False
    for ch in txt:
        if in_str:
            if esc: out.append(ch); esc = False; continue
            if ch == '\\': out.append(ch); esc = True; continue
            if ch == '"': in_str = False; out.append(ch); continue
            if ch == '\n': out.append('\\n'); continue
            if ch == '\r': continue
            if ch == '\t': out.append('\\t'); continue
            out.append(ch); continue
        if ch == '"': in_str = True
        out.append(ch)
    fixed = ''.join(out)
    if in_str: fixed += '"'
    for _ in range(8):
        for cand in (fixed, fixed + '}', fixed.rstrip().rstrip(',') + '}'):
            try:
                d = json.loads(cand)
                if isinstance(d, dict): return d
            except Exception:
                pass
        cut = fixed.rfind(',')
        if cut <= 0: break
        fixed = fixed[:cut]
    raise ValueError('模型返回的打标结果不是有效 JSON')

_IMG_MIME = {'.png': 'image/png', '.webp': 'image/webp', '.gif': 'image/gif',
             '.bmp': 'image/bmp', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg'}

def tag_one_image(path, p):
    import base64
    with open(path, 'rb') as fh:
        raw = fh.read()
    if len(raw) > 8 * 1024 * 1024: raise RuntimeError('图片过大(>8MB),已跳过')
    mime = _IMG_MIME.get(os.path.splitext(path)[1].lower(), 'image/jpeg')
    def ask(extra=''):
        # max_tokens 只给 500 时,合同截图的长 OCR 会把 JSON 顶断在字符串中间——必须留足余量
        payload = {'model': p['vision_model'], 'max_tokens': 1000,
                   'messages': [{'role': 'user', 'content': [
                       {'type': 'text', 'text': VISION_PROMPT + extra},
                       {'type': 'image_url', 'image_url': {'url': 'data:%s;base64,%s' % (mime, base64.b64encode(raw).decode())}}]}]}
        resp = _openai_req(p.get('base_url'), p.get('api_key'), '/chat/completions', payload,
                           timeout=120, verify=p.get('verify_ssl', True))
        txt = ((resp.get('choices') or [{}])[0].get('message') or {}).get('content', '') or ''
        return txt.strip().removeprefix('```json').removeprefix('```').removesuffix('```').strip()
    txt = ask()
    i, j = txt.find('{'), txt.rfind('}')
    try:
        d = _lenient_json_load(txt[i:j + 1] if i >= 0 else txt)
    except Exception:
        # 换个更严的说法重试一次:引号必须转义、ocr 压到 60 字内
        txt = ask('\n注意:严格输出合法 JSON。字符串值里的英文双引号必须写成 \\";'
                  'ocr 只给 60 字以内摘要,宁短勿断。')
        i, j = txt.find('{'), txt.rfind('}')
        d = _lenient_json_load(txt[i:j + 1] if i >= 0 else txt)
    if not isinstance(d, dict): raise ValueError('打标结果不是 JSON 对象')
    if d.get('category') not in CATS: d['category'] = '其他'
    return d

def write_image_index(entries):
    """写 图片索引.md(技能包按这张表自动配图)+ 同名 json 便于增量更新"""
    d = assets_dir()
    write_json(os.path.join(d, '图片索引.json'), entries)
    rows = ['| 图片ID | 文件 | 默认图注 | 类别 | 图内文字摘要 | 适用场景/关键词 | 落位锚点 |',
            '|---|---|---|---|---|---|---|']
    cell = lambda s: str(s or '').replace('|', '/').replace('\n', ' ')[:120]
    for e in entries:
        rows.append('| %s | %s | %s | %s | %s | %s | %s |' % (cell(e.get('id')), cell(e.get('file')),
                    cell(e.get('caption')), cell(e.get('category')), cell(e.get('ocr')),
                    cell(e.get('keywords')), cell(e.get('anchor'))))
    open(os.path.join(d, '图片索引.md'), 'w', encoding='utf-8').write('\n'.join(rows) + '\n')

def vision_worker(force):
    d = assets_dir(); img_dir = os.path.join(d, '图片')
    entries = read_json(os.path.join(d, '图片索引.json'), [])
    # 「标书图」= 入库占位、「待整理」= 上次识别失败的降级行:两者都算未完成,下次继续重试
    done_files = {e.get('file') for e in entries
                  if e.get('id') and not str(e.get('id')).startswith(('标书图', '待整理'))}
    p, err = _vision_provider()
    try:
        if err: VISION['errors'] = [err]; return
        files = [f for f in sorted(os.listdir(img_dir)) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.webp')) and not f.startswith('.')] if os.path.isdir(img_dir) else []
        todo = files if force else [f for f in files if f not in done_files]
        VISION.update(total=len(todo), done=0, updated=0, degraded=0, errors=[])
        used = {e.get('id') for e in entries}
        for fn in todo:
            fp = os.path.join(img_dir, fn)
            info, tag_err = None, ''
            try:
                info = tag_one_image(fp, p)
            except Exception as e:
                tag_err = str(e)[:120]
            if info is None:
                # 识别失败不再让这张图掉出索引:登记降级行(类别=其他、图注=文件名),
                # 撰写时仍可按文件名引用;错误原因留在 errors 里供用户重试排查。
                stem0 = os.path.splitext(fn)[0]
                info = {'category': '其他', 'id': '待整理-' + stem0[-12:],
                        'caption': stem0, 'ocr': '识别失败,待人工确认',
                        'keywords': '', 'anchor': '', '_degraded': True}
                VISION['degraded'] = VISION.get('degraded', 0) + 1
                VISION['errors'] = (VISION['errors'] + ['%s:%s(已降级登记,可重新打标)' % (fn, tag_err)])[-30:]
            try:
                safe = ''.join(c for c in str(info.get('id') or '')[:40] if c not in '\\/:*?"<>|').strip() or (info['category'] + '-未命名')
                base, n = safe, 2
                while safe in used: safe = '%s-%d' % (base, n); n += 1
                used.add(safe)
                ext = os.path.splitext(fn)[1]
                newfn = safe + ext
                if info.get('_degraded'):
                    newfn = fn      # 降级行保留原文件名,重打标时按「未识别」重试
                elif newfn != fn and not os.path.exists(os.path.join(img_dir, newfn)):
                    os.rename(fp, os.path.join(img_dir, newfn))
                else:
                    newfn = fn
                entries = [e for e in entries if e.get('file') not in (fn, newfn)]
                entries.append({'id': safe, 'file': newfn, 'caption': info.get('caption', ''), 'category': info.get('category', '其他'),
                                'ocr': info.get('ocr', ''), 'keywords': info.get('keywords', ''), 'anchor': info.get('anchor', '')})
                if not info.get('_degraded'): VISION['updated'] += 1
                write_image_index(entries)      # 每张即时落盘,中断也不丢
            except Exception as e:
                VISION['errors'] = (VISION['errors'] + ['%s:%s' % (fn, str(e)[:80])])[-30:]
            VISION['done'] += 1
    finally:
        with VISION_LOCK:
            VISION['running'] = False; VISION['ts'] = now()

@app.post('/v1/assets/vision_index')
def start_vision(force: bool = False):
    p, err = _vision_provider()
    if err: return JSONResponse({'ok': False, 'error': err}, 400)
    if not try_start_vision(force):
        return {'ok': True, 'already': True, **VISION}
    return {'ok': True, 'model': p['vision_model']}

@app.get('/v1/assets/vision_index')
def vision_status(): return VISION

@app.post('/v1/assets/open')
def open_assets():
    """在系统文件管理器里打开素材库目录(仅本地模式;云端多租户禁用)"""
    if MULTIUSER: return JSONResponse({'ok': False, 'error': '云端模式不支持'}, 403)
    d = assets_dir()
    try:
        if sys.platform == 'darwin': subprocess.Popen(['open', d])
        elif os.name == 'nt': os.startfile(d)  # noqa
        else: subprocess.Popen(['xdg-open', d])
        return {'ok': True, 'folder': d}
    except Exception as e:
        return {'ok': False, 'error': str(e), 'folder': d}

@app.delete('/v1/assets')
def clear_assets(scope: str = 'ingested'):
    """清空入库产物(章节模板/图片/原始标书/索引/流水);scope=all 时连规范素材一起清"""
    d = assets_dir(); removed = 0
    targets = ['章节模板', '图片', '原始标书']
    files = ['图片索引.md', '入库流水.jsonl']
    if scope == 'all': files += ['公司介绍.md', '产品资料.md', '资质与案例.md', '应答要点.md']
    for t in targets:
        p = os.path.join(d, t)
        if os.path.isdir(p): removed += sum(len(f) for _, _, f in os.walk(p)); shutil.rmtree(p, ignore_errors=True)
    for f in files:
        p = os.path.join(d, f)
        if os.path.isfile(p): os.remove(p); removed += 1
    return {'ok': True, 'removed': removed}

@app.post('/v1/assets/ingest')
async def ingest_asset(file: UploadFile = File(...)):
    """上传过往标书/素材(docx/doc/pdf/md/txt/html/图片)→ 图片入 图片/、章节拆分归类入
    章节模板/<分类>/、缺失的规范素材自动生成;入库出现新图片且配了视觉模型时自动开打标"""
    fn = os.path.basename(file.filename or '过往标书.docx')
    stem, ext = os.path.splitext(fn); ext = ext.lower()
    raw_dir = os.path.join(assets_dir(), '原始标书'); os.makedirs(raw_dir, exist_ok=True)
    blob = await file.read()
    digest = hashlib.md5(blob).hexdigest()
    # 同一份文件重复上传直接跳过(否则章节模板会成倍堆积)
    log_path = os.path.join(assets_dir(), '入库流水.jsonl')
    if os.path.isfile(log_path):
        for ln in open(log_path, encoding='utf-8').read().splitlines():
            try: rec = json.loads(ln)
            except Exception: continue
            if rec.get('md5') == digest:
                return {'ok': True, 'skipped': True, 'source': fn, 'sections': 0, 'images': 0, 'categories': {},
                        'folder': assets_dir(), 'note': '这份文件之前已入库(%s),本次跳过,未重复生成' % rec.get('ts', '')}
    src = os.path.join(raw_dir, fn)
    open(src, 'wb').write(blob)
    img_dir = os.path.join(assets_dir(), '图片')
    sections, images = [], []
    if ext in ('.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp'):
        # 单图直接入 图片/,登记占位行;随后可由「AI 识图打标」补全
        os.makedirs(img_dir, exist_ok=True)
        dst = os.path.join(img_dir, fn)
        if not os.path.exists(dst): shutil.copy2(src, dst)
        images = [fn]
    elif ext in ('.docx', '.doc'):
        # .doc 先嗅魔数:不少「.doc」其实是改了后缀的 docx(zip)或 RTF,能直接解析,
        # 只有真 OLE 老格式才需要借系统转换器(Word/WPS/LibreOffice/textutil)。
        parse_src, rtf_text = src, ''
        if ext == '.doc':
            if blob[:4] == b'PK\x03\x04':
                pass                                    # 伪 doc,按 docx 解析
            elif blob[:5] == b'{\\rtf':
                rtf_text = rtf_to_text(blob)
            else:
                try:
                    parse_src = convert_doc_to_docx(src)
                except Exception as e:
                    return JSONResponse({'ok': False, 'error': str(e)}, 400)
        if rtf_text:
            sections = split_md_sections(promote_headings(rtf_text))
        else:
            try:
                sections, images = docx_to_sections_and_images(parse_src, img_dir)
            except ImportError:
                return JSONResponse({'ok': False, 'error': '解析 docx 需要 python-docx(内置引擎版已自带;源码运行请 pip install python-docx)'}, 501)
            except Exception as e:
                return JSONResponse({'ok': False, 'error': 'docx 解析失败:%s' % e}, 400)
    elif ext == '.pdf':
        try:
            text = pdf_to_text(src)
        except Exception as e:
            return JSONResponse({'ok': False, 'error': 'PDF 解析失败:%s' % e}, 400)
        if len(text.strip()) < 50:
            return JSONResponse({'ok': False, 'error': '这份 PDF 没有可提取的文字层(多半是扫描件)。'
                                 '请用带文字层的 PDF,或用 Word/WPS 转成 .docx 再上传'}, 400)
        sections = split_md_sections(promote_headings(text))
    elif ext in ('.html', '.htm'):
        text = html_to_text(blob)
        if not text.strip():
            return JSONResponse({'ok': False, 'error': 'HTML 里没有可提取的正文'}, 400)
        sections = split_md_sections(promote_headings(text))
    elif ext in ('.md', '.txt'):
        sections = split_md_sections(promote_headings(open(src, encoding='utf-8', errors='ignore').read()))
    else:
        return JSONResponse({'ok': False, 'error': '暂支持 .docx/.doc/.pdf/.md/.txt/.html 及图片'
                             '(png/jpg/webp);其他格式请先转成 Word 再上传'}, 400)
    stamp = datetime.datetime.now().strftime('%m%d')
    safe = lambda s: (''.join(c for c in s if c not in '\\/:*?"<>|').strip() or '未命名')[:40]
    created, cats, canon_new = [], {}, {}
    for title, content in sections:
        cat = classify(title)
        cats[cat] = cats.get(cat, 0) + 1
        d = os.path.join(assets_dir(), '章节模板', cat); os.makedirs(d, exist_ok=True)
        p = os.path.join(d, '%s_%s_%s.md' % (safe(title), safe(stem)[:20], stamp))
        open(p, 'w', encoding='utf-8').write('> 来源:%s · 入库 %s · 整章可复用,引用前请核对时效与项目名\n\n%s\n' % (fn, now(), content))
        created.append(os.path.relpath(p, assets_dir()))
        canon = {'公司介绍': '公司介绍.md', '产品资料': '产品资料.md', '资质与案例': '资质与案例.md'}.get(cat)
        if canon: canon_new.setdefault(canon, []).append(content)
    for canon, parts in canon_new.items():
        cp = os.path.join(assets_dir(), canon)
        if not os.path.exists(cp):  # 已有的规范素材绝不覆盖,只补缺
            open(cp, 'w', encoding='utf-8').write('\n\n'.join(parts) + '\n')
            created.append(canon)
    auto_tagging = False
    if images:
        # 先登记占位行(未识别),随后由「AI 识图打标」补全类别/图注/OCR/落位锚点
        entries = read_json(os.path.join(assets_dir(), '图片索引.json'), [])
        have = {e.get('file') for e in entries}
        for f in images:
            if f not in have:
                entries.append({'id': os.path.splitext(f)[0], 'file': f, 'caption': '〔待AI识图〕', 'category': '未识别',
                                'ocr': '', 'keywords': '来自 %s' % stem, 'anchor': ''})
        write_image_index(entries)
        # 反馈落地:图片自动打标 —— 已配视觉模型就直接开跑,不用用户再手点「开始打标」
        try:
            auto_tagging = try_start_vision(False)
        except Exception:
            pass
    rec = {'ts': now(), 'source': fn, 'md5': digest, 'sections': len(sections), 'images': len(images), 'categories': cats}
    open(log_path, 'a', encoding='utf-8').write(json.dumps(rec, ensure_ascii=False) + '\n')
    return {'ok': True, **rec, 'created': created[:50], 'folder': assets_dir(),
            'engine': 'local-parser', 'auto_tagging': auto_tagging}

@app.post('/v1/transcribe')
def transcribe():
    return JSONResponse({'error': '本机未装转写模型;接 whisper.cpp sidecar 或云端转写'}, 501)

HOST_GONE = False       # 桌面壳已经退出,但还有任务在跑:跑完自己走
_EXIT_TIMER = [None]
RELEASES_API = os.environ.get('BIDDOG_RELEASES_URL', os.environ.get(
    'BID_RELEASES_URL', 'https://api.github.com/repos/shandianT/bid-dog/releases/latest'))
# 发布策略文件(自有域名,国内可达;GitHub API 只作兜底):版本治理的唯一事实源。
# 商业化后同一文件下发 minimum_supported/sunset,老版本按策略进入提示→强提示→停新单。
POLICY_URL = os.environ.get('BIDDOG_POLICY_URL', 'https://bid-dog.vercel.app/release_policy.json')
UPDATE_TIMEOUT = float(os.environ.get('BIDDOG_UPDATE_TIMEOUT', os.environ.get('BID_UPDATE_TIMEOUT', 3)))
UPDATE_RECHECK_S = float(os.environ.get('BIDDOG_UPDATE_RECHECK_SECONDS', 6 * 3600))
# 构建级落日:每个发出去的安装包都带一个"保质期"。到期后若拿不到策略文件的
# 明确许可,新生成需要先升级(已生成内容的查看/导出永不受限)。
# 每次发版把它前移;策略文件可远程延长任何在外版本的寿命,所以不怕误伤。
BUILD_SUNSET = os.environ.get('BIDDOG_BUILD_SUNSET', '2026-11-30')

def _env_true(name): return str(os.environ.get(name, '')).strip().lower() in ('1', 'true', 'yes', 'on')

UPDATE_STATE = {'status': 'disabled' if _env_true('BID_NO_UPDATE_CHECK') else 'checking'}

def _version_tuple(value):
    nums = re.findall(r'\d+', str(value or ''))[:3]
    return tuple((list(map(int, nums)) + [0, 0, 0])[:3])

def _policy_cache_path(): return os.path.join(DATA, 'release_policy_cache.json')

def _fetch_json(url, opener, accept='application/json'):
    req = urllib.request.Request(url, headers={'Accept': accept,
                                               'User-Agent': 'bid-dog/' + ENGINE_VERSION})
    r = (opener or urllib.request.urlopen)(req, timeout=UPDATE_TIMEOUT)
    return json.loads(r.read().decode('utf-8', 'ignore'))

def check_for_update(url=None, opener=None):
    """升级检查:策略文件优先,GitHub Release 兜底。异常只落状态,不打扰任务或页面。
    成功取得的策略持久化到数据目录——断网期间版本治理仍按最后一次已知策略执行。"""
    global UPDATE_STATE
    if _env_true('BID_NO_UPDATE_CHECK'):
        UPDATE_STATE = {'status': 'disabled'}
        return UPDATE_STATE
    opener = opener or urllib.request.urlopen
    try:
        policy = _fetch_json(url or POLICY_URL, opener)
        if not isinstance(policy, dict) or not policy.get('latest'):
            raise ValueError('策略文件缺少 latest')
        latest = str(policy.get('latest'))
        page = str(policy.get('url') or policy.get('notes_url')
                   or 'https://github.com/shandianT/bid-dog/releases')
        record = {'latest': latest, 'url': page,
                  'minimum_supported': str(policy.get('minimum_supported') or ''),
                  'sunset': str(policy.get('sunset') or ''),
                  'message': str(policy.get('message') or '')[:300],
                  'checked_at': now()}
        try: write_json(_policy_cache_path(), record)
        except Exception: pass
        UPDATE_STATE = {'status': 'available' if _version_tuple(latest) > _version_tuple(ENGINE_VERSION)
                        else 'latest', 'latest': latest, 'url': page, 'source': 'policy'}
        return UPDATE_STATE
    except Exception:
        pass
    try:
        data = _fetch_json(url or RELEASES_API, opener, accept='application/vnd.github+json')
        releases = data if isinstance(data, list) else [data]
        rel = next((x for x in releases if isinstance(x, dict) and not x.get('draft') and not x.get('prerelease')), None)
        if not rel: raise ValueError('没有可用正式版本')
        tag = str(rel.get('tag_name') or rel.get('name') or '')
        match = re.search(r'\d+\.\d+\.\d+', tag)
        latest = match.group(0) if match else tag.lstrip('v')
        page = rel.get('html_url') or 'https://github.com/shandianT/bid-dog/releases'
        try:
            cached = read_json(_policy_cache_path(), {})
            cached.update({'latest': latest, 'url': page, 'checked_at': now()})
            write_json(_policy_cache_path(), cached)
        except Exception:
            pass
        UPDATE_STATE = {'status': 'available' if _version_tuple(latest) > _version_tuple(ENGINE_VERSION) else 'latest',
                        'latest': latest, 'url': page}
    except Exception:
        UPDATE_STATE = {'status': 'error'}
    return UPDATE_STATE

def version_gate():
    """当前版本的治理结论,生成入口据此放行/提示/拦截。

    模式:ok(正常)/ notice(有新版,提示)/ required(低于最低支持版,限期内强提示)
        / expired(过期:新生成先升级;查看与导出永不受限)。
    离线不奖励也不惩罚:按最后一次已知策略执行;从未取到策略时只有构建级落日兜底。"""
    if _env_true('BID_NO_UPDATE_CHECK') or _env_true('BIDDOG_NO_SUNSET'):
        return {'mode': 'ok'}
    policy = read_json(_policy_cache_path(), {})
    update_url = str(policy.get('url') or 'https://github.com/shandianT/bid-dog/releases')
    today = now()[:10]
    minimum = str(policy.get('minimum_supported') or '')
    if minimum and _version_tuple(ENGINE_VERSION) < _version_tuple(minimum):
        deadline = str(policy.get('sunset') or '')
        expired = bool(deadline) and today > deadline
        message = policy.get('message') or (
            '当前版本 %s 已低于最低支持版本 %s%s。请更新后继续生成;已生成的内容不受影响。' % (
                ENGINE_VERSION, minimum,
                ('，并已于 %s 停止支持' % deadline) if expired else ('，将于 %s 停止支持' % (deadline or '近期'))))
        return {'mode': 'expired' if expired else 'required',
                'message': message, 'url': update_url, 'deadline': deadline,
                'minimum_supported': minimum}
    # 构建级落日:打包版专用(源码/开发运行不受限)。取到过策略且未被列为过低 → 策略说了算。
    packaged = bool(getattr(sys, 'frozen', False)) or _env_true('BIDDOG_FORCE_SUNSET')
    if packaged and not policy.get('checked_at') and today > str(BUILD_SUNSET):
        return {'mode': 'expired', 'url': update_url, 'deadline': str(BUILD_SUNSET),
                'message': '这个安装包是内测构建,已于 %s 到期且长期未能联网确认版本状态。'
                           '请下载最新版后继续生成;已生成的内容与素材库完全不受影响。' % BUILD_SUNSET}
    if UPDATE_STATE.get('status') == 'available':
        return {'mode': 'notice', 'latest': UPDATE_STATE.get('latest'),
                'url': UPDATE_STATE.get('url') or update_url,
                'message': '有新版本 %s 可用。' % UPDATE_STATE.get('latest')}
    return {'mode': 'ok'}

def _generation_gate_response():
    """过期版本拦「开新单/续跑/重写」这一类会启动生成的动作;其余一律放行。"""
    gate = version_gate()
    if gate.get('mode') != 'expired':
        return None
    return JSONResponse({'ok': False, 'code': 'version_sunset', 'action': 'update',
                         'error': gate.get('message') or '当前版本已停止支持,请更新后继续生成。',
                         'update_url': gate.get('url') or ''}, 426)

def start_update_check():
    if _env_true('BID_NO_UPDATE_CHECK'): return
    def loop():
        while True:
            check_for_update()
            time.sleep(max(300.0, UPDATE_RECHECK_S))
    threading.Thread(target=loop, daemon=True, name='bid-dog-update-check').start()

def _exit_process_cleanly():
    """正常退出必须带走 detached OpenCode server；会话数据库保留供下次续做/验签。"""
    _kill_tracked_detached_children()
    invalidate_oc_runtime()
    os._exit(0)

def _runtime_work_locked():
    """调用方持有 RUNNING_LOCK；控制动作/证据重放也可能正在改文件或 OpenCode 状态。"""
    return bool(RUNNING or JOB_CONTROL or OC_REPLAYING)

def _exit_if_shutdown(generation):
    """退出提交点：generation 与所有状态在同一把锁下复核，重连可在提交前取消。"""
    global EXITING
    with RUNNING_LOCK:
        if (generation != SHUTDOWN_GENERATION or not SHUTTING_DOWN
                or not HOST_GONE or _runtime_work_locked() or EXITING):
            return False
        EXITING = True
        _EXIT_TIMER[0] = None
    # EXITING 是不可逆提交点；health 会返回 503。锁外清理避免和 OC replay 的 finally 锁反转。
    _exit_process_cleanly()
    return True

def _schedule_clean_exit(generation, delay=0.3):
    """调用方持有 RUNNING_LOCK，避免 health 在保存 timer 之前穿过。"""
    timer = threading.Timer(delay, _exit_if_shutdown, args=(generation,))
    timer.daemon = True
    old = _EXIT_TIMER[0]
    if old:
        try: old.cancel()
        except Exception: pass
    _EXIT_TIMER[0] = timer
    timer.start()

@app.post('/v1/shutdown')
def shutdown():
    """桌面壳关窗时调它。有任务在跑就先把任务跑完再退,不当场自尽。

    以前 Tauri 退出是无条件 kill 引擎,而 agent 是我们用 start_new_session=True
    起的独立进程组 —— **杀得掉引擎,杀不掉 agent**。于是 agent 还在后台写文件,
    负责收尾的那段代码(收拢产物、补出 Word、质检、完成播报)却已经死了:
    用户关窗去开个会,回来看到的是一个永远没有结局的任务。
    云端多用户模式下不接受这个请求——那是共享服务,不能被某个客户端关掉。"""
    global HOST_GONE, SHUTTING_DOWN, SHUTDOWN_GENERATION
    if MULTIUSER:
        return JSONResponse({'ok': False, 'error': '云端模式不支持远程关闭'}, 403)
    with RUNNING_LOCK:
        SHUTDOWN_GENERATION += 1
        generation = SHUTDOWN_GENERATION
        SHUTTING_DOWN = True       # 从这一刻起拒绝新派发，封住“看到空闲后又启动一单”的窗口
        HOST_GONE = True
        running = sorted(RUNNING)
        controls = sorted(JOB_CONTROL)
        replaying = bool(OC_REPLAYING)
        busy = _runtime_work_locked()
        if not busy:
            _schedule_clean_exit(generation, 0.3)  # 先把响应发出去，再停 OpenCode 并退出
    if not busy:
        return {'ok': True, 'exiting': True, 'running': 0}
    threading.Thread(target=_exit_when_idle, args=(generation,), daemon=True).start()
    if running:
        note = '还有 %d 个任务在跑，我先把它们跑完、收好尾再退出' % len(running)
    elif controls:
        note = '任务正在完成停止或删除操作，收好尾后自动退出'
    else:
        note = '正在恢复运行证据，完成后自动退出'
    return {'ok': True, 'exiting': False, 'running': len(running), 'jobs': running,
            'controls': controls, 'replaying': replaying, 'note': note}

def _exit_when_idle(generation):
    """桌面壳走了之后守着:任务全部收尾完成就自己退出,不留孤儿进程。"""
    while True:
        time.sleep(5)
        with RUNNING_LOCK:
            if generation != SHUTDOWN_GENERATION or not SHUTTING_DOWN or not HOST_GONE:
                return                     # 用户又把应用打开了(端口复用),继续正常服务
            busy = _runtime_work_locked()
        if busy: continue
        time.sleep(8)                       # 给 finalize/质检/出 Word 收尾留出余量
        if _exit_if_shutdown(generation): return

@app.get('/v1/health')
def health():
    """纯只读存活检查；监控、轮询和诊断不得悄悄改变引擎生命周期。"""
    if EXITING:
        return JSONResponse({'ok': False, 'error': '引擎正在退出，请等待应用自动重连'}, 503)
    return {'ok': True, 'data_dir': DATA, 'agent': bool(config_agent_cmd()),
            'version': ENGINE_VERSION, 'author': AUTHOR, 'features': ENGINE_FEATURES,
            'update': dict(UPDATE_STATE), 'version_gate': version_gate()}

@app.post('/v1/attach')
def attach():
    """桌面壳明确声明重新接管当前引擎，取消上一窗口留下的优雅退出倒计时。"""
    global HOST_GONE, SHUTTING_DOWN, SHUTDOWN_GENERATION
    with RUNNING_LOCK:
        if EXITING:
            return JSONResponse({'ok': False, 'error': '引擎正在退出，请等待应用自动重连'}, 503)
        SHUTDOWN_GENERATION += 1             # 使已排队 timer/idle watcher 全部失效
        HOST_GONE = False
        SHUTTING_DOWN = False
        timer = _EXIT_TIMER[0]
        if timer:
            try: timer.cancel()
            except Exception: pass
            _EXIT_TIMER[0] = None
    return {'ok': True, 'attached': True, 'version': ENGINE_VERSION}

def migrate_conf():
    """存量配置迁移。

    `s2`(界面上的「自动(默认)」)现在就是 OpenCode —— 拿真实招标文件做过对比后转正的。
    历史上存过 kind=opencode 的用户依旧归并到 s2:两者本就共用 Key/网关/模型字段,
    而且现在连执行外壳都是同一个,留两个值只会让「当前生效的是哪个引擎」这件事更难说清。
    未知 kind 一律落安全值。"""
    cp = os.path.join(DATA, 'config.json')
    conf = read_json(cp, None)
    if not conf: return
    eng = conf.get('engine') or {}
    kind = eng.get('kind')
    known = ('mock', 's2', 'sowork', 'claude', 'codex', 'custom')
    if kind == 'opencode':
        eng['kind'] = 's2'
    elif kind and kind not in known:
        eng['kind'] = 's2' if (eng.get('s2_key') or '').strip() else 'mock'
    else:
        return
    conf['engine'] = eng
    write_json(cp, conf)

def reconcile_orphan_jobs():
    """启动对账:引擎/应用上次异常退出(崩溃、睡眠杀进程、强关)留下的「运行中」任务,
    重启后没有 worker、没有终态,会被呈现成「任务意外中断」且重试被 409 拒——
    检查点明明完好,用户却被逼重跑(反馈 6/9 的组合症状之一)。
    这里把 running/retry_wait 节点归位成 pending、补一个可续跑的停止终态,
    界面自动给出「从已保存内容继续」入口。"""
    if MULTIUSER: return
    try:
        root = jobs_dir()
        if not os.path.isdir(root): return
        for base in sorted(os.listdir(root)):
            job = os.path.join(root, base)
            if not os.path.isdir(job) or base.startswith('.'): continue
            outcome = read_json(os.path.join(job, 'outcome.json'), {})
            if outcome.get('state') in ('done', 'stopped'): continue
            try:
                state = generation_pipeline.load(job)
            except Exception:
                continue
            if state.get('version') != generation_pipeline.PIPELINE_VERSION: continue
            hanging = [n for n in state.get('nodes') or []
                       if n.get('state') in ('running', 'retry_wait')]
            if not hanging: continue
            try:
                generation_pipeline.recover(job)
                emit(job, {'type': 'message', 'role': 'agent',
                           'text': '检测到应用上次中途退出,已把任务归位到最近的检查点。'
                                   '已生成的内容全部保留,点「从已保存内容继续」接着跑。'})
                halt(job, '已停止（应用重启，可从检查点继续）')
            except Exception:
                continue
    except Exception:
        pass

def backup_docx_template():
    """把 python-docx 的 default.docx 模板备份到数据目录,并经环境变量交给技能脚本兜底。

    PyInstaller onefile 的 _MEI 临时解压目录放在系统 %TEMP% 下,Windows 磁盘清理/
    存储感知会在引擎长时间运行(隔夜挂机)后把它删掉——真机报障:
    「自动补出 Word 失败(Package not found at ...\\_MEI...\\docx\\templates\\default.docx)」。
    引擎自身启动时模板还在,先复制一份到常驻数据目录;build_tender_docx.py 等脚本
    在 Document() 抛 PackageNotFound 时改用 BIDDOG_DOCX_TEMPLATE 指向的这份备份。"""
    try:
        import docx as _docx
        src = os.path.join(os.path.dirname(_docx.__file__), 'templates', 'default.docx')
        dst = os.path.join(DATA, 'docx_default_template.docx')
        if os.path.isfile(src) and (not os.path.isfile(dst)
                                    or os.path.getsize(dst) != os.path.getsize(src)):
            shutil.copy2(src, dst)
        if os.path.isfile(dst):
            os.environ['BIDDOG_DOCX_TEMPLATE'] = dst
    except Exception:
        pass    # 模板备份是兜底,失败不阻塞启动

ensure_preset()   # 预置配置种子(定制发包用),必须在任何请求处理前完成
migrate_conf()    # 存量引擎配置归并(撤下的选项不能把老用户晾在空白下拉上)
backup_docx_template()  # _MEI 被系统清理后 Word 导出的最后一道保险
reconcile_orphan_jobs() # 上次异常退出留下的「运行中」任务归位到检查点,可续跑
start_update_check()

# 网页/演示模式的前端:优先新界面构建产物,没构建过则回落经典单文件页。
# 桌面版前端由 Tauri 壳自带(frontendDist),不走这里。
_web_next = os.path.join(HERE, '..', 'app-next', 'dist')
web = os.environ.get('BID_WEB_DIR') or (_web_next if os.path.isdir(_web_next) else os.path.join(HERE, '..', 'app', 'src'))
if os.path.isdir(web): app.mount('/', StaticFiles(directory=web, html=True), name='web')

if __name__ == '__main__':
    import uvicorn
    # 云端容器需监听 0.0.0.0;桌面/本机默认只听回环,不对外暴露
    host = os.environ.get('HOST') or ('0.0.0.0' if MULTIUSER else '127.0.0.1')
    _port = int(os.environ.get('PORT', 8080))
    print('[中标狗] 引擎 v%s 启动:http://127.0.0.1:%d  数据目录:%s' % (ENGINE_VERSION, _port, DATA), flush=True)
    # access_log 关掉:SSE 每秒一次轮询把日志刷成瀑布,网页端长跑时 IO 白白吃 CPU
    uvicorn.run(app, host=host, port=_port, access_log=os.environ.get('BID_ACCESS_LOG') == '1',
                log_level='warning', timeout_keep_alive=65)
